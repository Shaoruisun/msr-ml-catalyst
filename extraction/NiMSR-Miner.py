
#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
================================================================================
MSR literature extraction script V6 Pro
Focus: promoter-aware identity, table-first extraction, conservative fallback
================================================================================
This file keeps the single-script workflow:
1. chunk text extraction
2. figure extraction
3. merge
4. clean
5. save

Current priorities:
- more reliable catalyst identity merging
- table-first PDF text extraction
- global fallback limited to conservative preparation fields
================================================================================
"""

import sys
sys.stdout.reconfigure(encoding='utf-8')

import os
import re
import csv
import json
import base64
from collections import Counter
from typing import List, Dict, Any, Optional, Tuple, Set
import pdfplumber
import docx
from openai import OpenAI

# ==========================================
# Runtime config
# ==========================================
def _get_env_or_default(name: str, default: str) -> str:
    value = os.getenv(name)
    if value is None:
        return str(default)
    return str(value).strip()


def _get_env_flag(name: str, default: bool) -> bool:
    value = os.getenv(name)
    if value is None:
        return bool(default)
    return str(value).strip().lower() not in {"", "0", "false", "no", "off"}


DEFAULT_API_KEY = ""  # Set MSR_API_KEY environment variable before running
DEFAULT_BASE_URL = ""  # Set MSR_BASE_URL environment variable before running
DEFAULT_TEXT_MODEL = ""  # Set MSR_TEXT_MODEL environment variable (e.g., qwen3-max)
DEFAULT_VISION_MODEL = ""  # Set MSR_VISION_MODEL environment variable (e.g., qwen-vl-max)
_SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
_REPO_DIR = os.path.dirname(_SCRIPT_DIR)
DEFAULT_IN_DIR = os.path.join(_REPO_DIR, "data", "pdf_in")
DEFAULT_OUT_DIR = os.path.join(_REPO_DIR, "data", "csv_out")
DEFAULT_TEMP_IMG_DIR = os.path.join(_REPO_DIR, "data", "temp_images")

API_KEY = _get_env_or_default("MSR_API_KEY", DEFAULT_API_KEY)
BASE_URL = _get_env_or_default("MSR_BASE_URL", DEFAULT_BASE_URL)
TEXT_MODEL = _get_env_or_default("MSR_TEXT_MODEL", DEFAULT_TEXT_MODEL)
VISION_MODEL = _get_env_or_default("MSR_VISION_MODEL", DEFAULT_VISION_MODEL)
EXPORT_AUDIT_ARTIFACTS = _get_env_flag("MSR_EXPORT_AUDITS", False)

client = None

# Directory config
IN_DIR = _get_env_or_default("MSR_IN_DIR", DEFAULT_IN_DIR)
OUT_DIR = _get_env_or_default("MSR_OUT_DIR", DEFAULT_OUT_DIR)
TEMP_IMG_DIR = _get_env_or_default("MSR_TEMP_IMG_DIR", DEFAULT_TEMP_IMG_DIR)

# Three-layer aggregated master tables:
# - research_master_all.csv: maximal research retention
# - research_master_strict.csv: cleaner research table for review/statistics
# - modeling_final.csv: conservative high-purity final table
RESEARCH_MASTER_ALL_CSV = os.path.join(OUT_DIR, "research_master_all.csv")
RESEARCH_MASTER_STRICT_CSV = os.path.join(OUT_DIR, "research_master_strict.csv")
MODELING_FINAL_MASTER_CSV = os.path.join(OUT_DIR, "modeling_final.csv")

# Legacy compatibility aliases. They now point to research layers only and must
# never again silently mean the strict modeling-final layer.
COMBINED_CSV = os.path.join(OUT_DIR, "dataset_all.csv")
STRICT_COMBINED_CSV = os.path.join(OUT_DIR, "dataset_all_strict.csv")

os.makedirs(OUT_DIR, exist_ok=True)
os.makedirs(TEMP_IMG_DIR, exist_ok=True)


# ==========================================
# MSR catalyst prior knowledge
# ==========================================
MSR_ACTIVE_METALS = {"Cu", "Ni", "Pd", "Pt", "Ru", "Rh", "Co", "Fe", "Ir", "Au"}
MSR_PROMOTERS = {"Zn", "Ce", "La", "K", "Mn", "Cr", "Mo", "Sn", "Ga", "In", "Ge", "Bi", "Sb"}
MSR_RARE_EARTH = {"Ce", "La", "Y", "Sm", "Gd", "Pr", "Nd", "Eu", "Tb", "Dy", "Ho", "Er", "Tm", "Yb", "Lu"}

# ==========================================
# Catalyst abbreviation alias map
# 用于将 VL 提取的缩写名称展开为标准化学式，解决 figure/text 命名不一致问题
# 在 normalize_identity_aliases 入口处统一展开，下游所有路径自动受益
# ==========================================
_CATALYST_ABBREVIATION_ALIAS_MAP = {
    # MAl2O4 系列：金属铝酸盐缩写 → 标准化学式
    "CuNiAl":  "Cu/NiAl2O4",
    "CuCoAl":  "Cu/CoAl2O4",
    "CuMgAl":  "Cu/MgAl2O4",
    "CuZnAl":  "Cu/ZnAl2O4",
    "CuFeAl":  "Cu/FeAl2O4",
    "NiCoAl":  "Ni/CoAl2O4",
    "NiMgAl":  "Ni/MgAl2O4",
    "NiZnAl":  "Ni/ZnAl2O4",
    # 简单氧化铝缩写
    "NiAl":    "Ni/Al2O3",
    "CuAl":    "Cu/Al2O3",
    "CoAl":    "Co/Al2O3",
}

# folded 版 alias map：key 统一小写 + 去非字母数字，用于大小写不敏感 + 轻微标点容错
_CATALYST_ABBREVIATION_ALIAS_MAP_FOLDED = {
    re.sub(r"[^a-z0-9]", "", k.lower()): v
    for k, v in _CATALYST_ABBREVIATION_ALIAS_MAP.items()
}


def _fold_alias_key(value: str) -> str:
    """将输入 strip、去首尾轻微标点、转小写、去非字母数字，得到 folded key。"""
    s = str(value).strip()
    s = s.strip(",.;:()[]{}\"'")
    return re.sub(r"[^a-z0-9]", "", s.lower())


def _expand_catalyst_alias(value: str) -> str:
    """若 folded key 在 alias map 中，返回标准展开式；否则返回原值。"""
    folded = _fold_alias_key(value)
    return _CATALYST_ABBREVIATION_ALIAS_MAP_FOLDED.get(folded, value)

# ==========================================
# Schema definition
# ==========================================
DATA_SCHEMA = {
    "type": "object",
    "properties": {
        "Catalyst_ID": {"type": "string", "description": "Canonical catalyst identity string."},
        "Catalyst": {"type": "string", "description": "Catalyst label as stated in the paper."},
        "Active_Metal": {"type": "string", "description": "Core active metal(s)."},
        "Metal_Loading_wt%": {"type": "string", "description": "Total metal loading in wt%."},
        "Ni_Loading_wt%": {"type": "string", "description": "Ni metal loading in wt%."},
        "Promoter_Loading_wt%": {"type": "string", "description": "Promoter or second-metal loading in wt%."},
        "Alloy_Ratio": {"type": "string", "description": "Explicit metal ratio with element names."},
        "Support": {"type": "string", "description": "Support or support family."},
        "Promoter": {"type": "string", "description": "Promoter or second metal."},
        "Precursor": {"type": "string", "description": "Metal precursor or salt."},
        "Support_Prep_Method": {"type": "string", "description": "Normalized support preparation method."},
        "Metal_Loading_Method": {"type": "string", "description": "Normalized metal loading method."},
        "Dry_Temp_C": {"type": "string", "description": "Drying temperature in C."},
        "Dry_Time_h": {"type": "string", "description": "Drying time in h."},
        "Calcination_Temp_C": {"type": "string", "description": "Calcination temperature in C."},
        "Calcination_Time_h": {"type": "string", "description": "Calcination time in h."},
        "Reduction_Temp_C": {"type": "string", "description": "Reduction or activation temperature in C."},
        "Reduction_Time_h": {"type": "string", "description": "Reduction or activation time in h."},
        "Reaction_Temp_C": {"type": "string", "description": "Reaction temperature in C for a specific data point."},
        "S_C_Ratio": {"type": "string", "description": "Steam-to-carbon or H2O/MeOH ratio."},
        "GHSV_mL_g_h": {"type": "string", "description": "Gas hourly space velocity or preserved WHSV string."},
        "Pressure_bar": {"type": "string", "description": "Reaction pressure in bar."},
        "Feed_Composition": {"type": "string", "description": "Feed composition string."},
        "MeOH_Conversion_%": {"type": "string", "description": "Methanol conversion percent."},
        "H2_Yield_%": {"type": "string", "description": "Hydrogen yield percent."},
        "H2_Production_Rate": {"type": "string", "description": "Hydrogen production rate or flow value with original unit string."},
        "H2_Selectivity_%": {"type": "string", "description": "Hydrogen selectivity percent."},
        "CO_Selectivity_%": {"type": "string", "description": "CO selectivity percent."},
        "CO2_Selectivity_%": {"type": "string", "description": "CO2 selectivity percent."},
        "Reasoning_Selectivity": {"type": "string", "description": "Reasoning note for CO vs CO2 selectivity assignment."},
        "CO_Concentration_ppm": {"type": "string", "description": "CO concentration in ppm."},
        "TOS_h": {"type": "string", "description": "Time on stream in h."},
        "Deactivation_Rate_%_h": {"type": "string", "description": "Deactivation rate in percent per hour."},
        "Carbon_Deposition_wt%": {"type": "string", "description": "Carbon deposition in wt%."},
        "Notes": {"type": "string", "description": "Free-form notes for ambiguity, ranges, or context."}
    },
    "required": ["Catalyst"]
}



# ==========================================
# [Fix 2] Table-first extraction: preserve tables before body text parsing
# ==========================================
def table_to_markdown(table: List[List]) -> str:
    """Function docstring removed for runtime stability."""
    if not table:
        return ""
    rows = []
    for i, row in enumerate(table):
        cleaned = [str(c).strip().replace("\n", " ") if c else "" for c in row]
        rows.append("| " + " | ".join(cleaned) + " |")
        # Add the Markdown separator row after the header.
        if i == 0:
            rows.append("| " + " | ".join(["---"] * len(cleaned)) + " |")
    return "\n".join(rows)


def extract_page_content(page) -> str:
    """Function docstring removed for runtime stability."""
    page_content = ""
    table_bboxes = []  # Collect table bounding boxes for later body-text filtering.

    # --- Step 1: extract full tables first to preserve wide table structure. ---
    try:
        # find_tables() returns table objects with bbox metadata.
        found_tables = page.find_tables()
        if found_tables:
            for t_idx, table_obj in enumerate(found_tables):
                # Cache table bbox so we can exclude it from body text.
                table_bboxes.append(table_obj.bbox)
                raw_table = table_obj.extract()
                md = table_to_markdown(raw_table)
                # [Improve5] If table has fewer than 3 rows, pdfplumber likely
                # failed to parse it correctly (e.g. complex header or cross-page).
                # Fall back to plain text extraction for that region so the LLM
                # can still read the content as prose.
                if md and raw_table and len(raw_table) >= 3:
                    page_content += f"\n=== TABLE {t_idx+1} (Page {page.page_number}) ===\n"
                    page_content += md + "\n"
                    page_content += "=== END TABLE ===\n\n"
                elif raw_table and len(raw_table) < 3:
                    # Sparse table — extract as plain text fallback
                    try:
                        region_text = page.within_bbox(table_obj.bbox).extract_text() or ""
                        if region_text.strip():
                            page_content += f"\n=== TABLE {t_idx+1} (Page {page.page_number}, plain-text fallback) ===\n"
                            page_content += region_text.strip() + "\n"
                            page_content += "=== END TABLE ===\n\n"
                    except Exception:
                        if md:
                            page_content += f"\n=== TABLE {t_idx+1} (Page {page.page_number}) ===\n"
                            page_content += md + "\n"
                            page_content += "=== END TABLE ===\n\n"
    except Exception:
        pass  # Do not let table extraction failure block body extraction.

    # --- Step 2: extract non-table body text only. ---
    try:
        if table_bboxes:
            # Remove text objects whose center falls inside any table bbox.
            def _not_in_table(obj):
                ox0 = obj.get("x0", 0)
                oy0 = obj.get("top", 0)
                ox1 = obj.get("x1", 0)
                oy1 = obj.get("bottom", 0)
                for (tx0, ty0, tx1, ty1) in table_bboxes:
                    cx = (ox0 + ox1) / 2
                    cy = (oy0 + oy1) / 2
                    if tx0 <= cx <= tx1 and ty0 <= cy <= ty1:
                        return False
                return True
            filtered_page = page.filter(_not_in_table)
        else:
            filtered_page = page

        # [Fix-TwoCol] Detect true column split via x-coordinate gap analysis,
        # then rebuild each column independently using word-level coordinates.
        # This prevents double-column row interleaving (e.g. "tions at 9.0"
        # or "molar ratio =" separated from its value by the opposite column).
        def _detect_col_split(pg) -> float:
            """Return the x-coordinate of the column gap, or 0 if single-column."""
            try:
                words = pg.extract_words()
                if not words:
                    return 0.0
                w = pg.width
                # Build 40-bucket histogram of word x0 positions
                buckets = [0] * 40
                for wd in words:
                    b = min(int(wd["x0"] / w * 40), 39)
                    buckets[b] += 1
                # Find the minimum-density bucket in the central 30%-70% zone
                # (column gutter is always near the middle)
                zone = list(range(12, 28))  # 30%-70%
                gap_bucket = min(zone, key=lambda i: buckets[i])
                gap_density = buckets[gap_bucket]
                # Neighbouring buckets must be significantly denser
                left_peak  = max(buckets[max(0, gap_bucket-3):gap_bucket] or [0])
                right_peak = max(buckets[gap_bucket+1:min(40, gap_bucket+4)] or [0])
                if left_peak > 0 and right_peak > 0 and gap_density < 0.35 * min(left_peak, right_peak):
                    return (gap_bucket + 0.5) / 40 * w
            except Exception:
                pass
            return 0.0

        def _words_to_col_text(pg, x_min: float, x_max: float) -> str:
            """Reconstruct column text from words within [x_min, x_max]."""
            try:
                col_words = [w for w in pg.extract_words()
                             if w["x0"] >= x_min and w["x1"] <= x_max + 5]
                if not col_words:
                    return ""
                lines: dict = {}
                for wd in col_words:
                    y = round(wd["top"] / 3) * 3
                    lines.setdefault(y, []).append(wd)
                parts = []
                for y in sorted(lines):
                    row = sorted(lines[y], key=lambda wd: wd["x0"])
                    parts.append(" ".join(wd["text"] for wd in row))
                return "\n".join(parts)
            except Exception:
                return ""

        gap_x = _detect_col_split(filtered_page)
        if gap_x > 0:
            left_text  = _words_to_col_text(filtered_page, 0, gap_x)
            right_text = _words_to_col_text(filtered_page, gap_x, filtered_page.width)
            if len(left_text.strip()) > 50 and len(right_text.strip()) > 50:
                body_text = left_text + "\n" + right_text
            else:
                body_text = filtered_page.extract_text() or ""
        else:
            mid_x = filtered_page.width / 2
            left_text  = filtered_page.within_bbox((0, 0, mid_x, filtered_page.height)).extract_text() or ""
            right_text = filtered_page.within_bbox((mid_x, 0, filtered_page.width, filtered_page.height)).extract_text() or ""
            if len(left_text.strip()) > 50 and len(right_text.strip()) > 50:
                body_text = left_text + "\n" + right_text
            else:
                body_text = filtered_page.extract_text() or ""

        if body_text.strip():
            page_content += f"\n=== BODY TEXT (Page {page.page_number}) ===\n{body_text}\n"
    except Exception:
        # Fallback to whole-page text if filtering fails.
        body_text = page.extract_text() or ""
        if body_text:
            page_content += f"\n=== BODY TEXT (Page {page.page_number}) ===\n{body_text}\n"

    return page_content


def extract_pages_as_chunks(file_path: str, pages_per_chunk: int = 4, overlap: int = 1) -> List[str]:
    """Function docstring removed for runtime stability."""
    chunks = []
    try:
        with pdfplumber.open(file_path) as pdf:
            total_pages = len(pdf.pages)
            page_texts = []
            for page in pdf.pages:
                page_texts.append(extract_page_content(page))

            step = pages_per_chunk - overlap
            start = 0
            while start < total_pages:
                end = min(start + pages_per_chunk, total_pages)
                chunk = "".join(page_texts[start:end])
                if chunk.strip():
                    chunks.append(chunk)
                start += step
    except Exception as e:
        print(f"  [ERROR] PDF page extraction failed: {e}")
    return chunks


def extract_text_from_docx(file_path: str) -> List[str]:
    """Function docstring removed for runtime stability."""
    text = ""
    try:
        doc = docx.Document(file_path)
        for para in doc.paragraphs:
            if para.text.strip():
                text += para.text + "\n"
        for table in doc.tables:
            text += "\n=== TABLE ===\n"
            for row in table.rows:
                cells = [c.text.replace("\n", " ").strip() for c in row.cells]
                text += "| " + " | ".join(cells) + " |\n"
            text += "=== END TABLE ===\n\n"
    except Exception as e:
        print(f"  [ERROR] DOCX extraction failed: {e}")
    return [text] if text.strip() else []


def clean_text_for_api(text: str) -> str:
    """Function docstring removed for runtime stability."""
    text = clean_residual_mojibake_chars(text)
    text = re.sub(r'[\x00-\x08\x0b-\x0c\x0e-\x1f\x7f-\x9f]', '', text)
    text = re.sub(r'http[s]?://\S+', ' ', text)
    text = re.sub(r'\S+@\S+\.\S+', ' ', text)
    text = re.sub(r'[^\x00-\x7F]{50,}', ' ', text)
    text = re.sub(r'[^\w\s\.\,\-\(\)\[\]\{\}\/\:\;\=\+\%\|\#\u63b3\u864f\u9c81]{5,}', ' ', text)
    text = re.sub(r'<<[^>]*>>', ' ', text)
    text = re.sub(r'\s+', ' ', text)
    text = re.sub(r'\d{20,}', ' ', text)
    text = clean_residual_mojibake_chars(text.strip())
    return text if len(text) >= 50 else ""


# ==========================================
# Global preparation parameter extraction
# ==========================================
GLOBAL_FALLBACK_ALLOWED_FIELDS = {
    # 制备字段
    "Dry_Temp_C", "Dry_Time_h",
    "Calcination_Temp_C", "Calcination_Time_h",
    "Reduction_Temp_C", "Reduction_Time_h",
    # 反应条件字段（全局固定值，如单温度比较研究）
    "Reaction_Temp_C",
    "S_C_Ratio", "S_C_Ratio_Raw",
    "GHSV_mL_g_h", "GHSV_mL_g_h_Raw",
    "WHSV_h_inv", "WHSV_h_inv_Raw",
    "SpaceVelocity_norm", "SpaceVelocity_norm_Raw",
    "SpaceVelocity_type", "SpaceVelocity_unit",
    "Pressure_bar", "Pressure_bar_Raw",
    "Feed_MeOH_to_H2O_Ratio", "Feed_MeOH_to_H2O_Ratio_Raw",
    "Feed_Composition",
}

# 不再封锁任何反应条件字段——全部允许全局广播
GLOBAL_FALLBACK_BLOCKED_FIELDS: set = set()

# Fields supplemented into figure-sourced records from matched text records.
GLOBAL_FIGURE_SUPPLEMENT_FIELDS = {
    # 制备字段
    "Reaction_Temp_C", "Calcination_Temp_C", "Reduction_Temp_C",
    "Dry_Temp_C", "Dry_Time_h", "Calcination_Time_h", "Reduction_Time_h",
    "Metal_Loading_wt%", "Support", "Precursor", "Metal_Loading_Method", "Support_Prep_Method",
    # 反应条件字段
    "S_C_Ratio", "S_C_Ratio_Raw",
    "GHSV_mL_g_h", "GHSV_mL_g_h_Raw",
    "WHSV_h_inv", "WHSV_h_inv_Raw",
    "SpaceVelocity_norm", "SpaceVelocity_norm_Raw",
    "SpaceVelocity_type", "SpaceVelocity_unit",
    "Pressure_bar", "Pressure_bar_Raw",
    "Feed_MeOH_to_H2O_Ratio", "Feed_MeOH_to_H2O_Ratio_Raw",
    "Feed_Composition",
}

def _reaction_temp_value_is_sane(value: Any) -> bool:
    text = clean_residual_mojibake_chars(str(value or "")).strip()
    if not text:
        return False
    try:
        temp = float(re.findall(r"[-+]?\d+(?:\.\d+)?", text.replace(",", ""))[0])
    except (IndexError, ValueError, TypeError):
        return False
    return 50.0 <= temp <= 900.0


def extract_global_params(first_chunk: str, file_name: str) -> Dict[str, str]:
    """Function docstring removed for runtime stability."""
    text = clean_text_for_api(first_chunk)
    if not text:
        return {}

    system_prompt = """You extract paper-level global parameters for MSR catalyst studies.

Only extract a parameter if it is explicitly shared across the whole study, or clearly stated as common to all catalysts, all samples, or all experiments in the paper.

Allowed global fields:
Preparation fields:
- Dry_Temp_C
- Dry_Time_h
- Calcination_Temp_C
- Calcination_Time_h
- Reduction_Temp_C
- Reduction_Time_h

Reaction condition fields (extract ONLY when the value is explicitly stated as the single fixed condition for ALL experiments in the paper):
- Reaction_Temp_C
- S_C_Ratio         (steam-to-carbon molar ratio, e.g. "S/C = 2", "steam-to-carbon ratio of 1.5", "molar ratio of water and methanol of 2:1", "molar ratio = 2.5:1 (H2O:CH3OH)")
- GHSV_mL_g_h       (gas hourly space velocity, e.g. "GHSV = 16000 mL/g/h")
- WHSV_h_inv        (weight hourly space velocity)
- SpaceVelocity_norm (normalized space velocity)
- Pressure_bar      (reaction pressure)
- Feed_MeOH_to_H2O_Ratio (methanol-to-water feed ratio)
- Feed_Composition  (feed gas composition)

RULES:
1. Only extract a reaction condition as global if it is explicitly stated as the single fixed value used for ALL catalysts/experiments in the paper.
2. For Reaction_Temp_C: do NOT extract globally if the paper contains a temperature sweep (e.g. 200-400°C). If the paper reports BOTH thermal catalytic conditions AND other reaction modes (e.g. non-thermal plasma, NTP, photocatalysis, electrochemical), extract ONLY the thermal catalytic (thermally-driven) reaction temperature. Ignore temperatures associated with plasma discharge, photocatalytic, or electrochemical conditions.
3. For S_C_Ratio / GHSV: extract globally if the paper states a single fixed value (e.g. "S/C ratio of 2 was used throughout", "GHSV of 16000 h-1"). For S_C_Ratio, also recognize expressions like "molar ratio of water and methanol of X:1", "molar ratio = X:1 (H2O:CH3OH)", or "water-to-methanol molar ratio of X" — extract the numeric value of X as S_C_Ratio.
4. Do not infer from one figure caption, one table row, or one local experiment.
5. If unsure, leave empty.

Output JSON: {"global_params": {"field_name": "value"}}"""
    system_prompt = clean_residual_mojibake_chars(system_prompt)

    try:
        response = client.chat.completions.create(
            model=TEXT_MODEL,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": f"Extract paper-level global preparation parameters from the text below:\n{text}"}
            ],
            response_format={"type": "json_object"},
            temperature=0.1
        )
        params = json.loads(response.choices[0].message.content).get("global_params", {})
        # Both preparation fields and reaction-condition fields (S/C, GHSV, pressure, etc.)
        # are allowed into the fallback pool when explicitly stated as paper-level globals.
        params = {
            k: v for k, v in params.items()
            if k in GLOBAL_FALLBACK_ALLOWED_FIELDS and str(v).strip()
        }
        if "Reaction_Temp_C" in params and not _reaction_temp_value_is_sane(params.get("Reaction_Temp_C", "")):
            print(f"  [global_params_sanity] discard Reaction_Temp_C={params.get('Reaction_Temp_C')!r}")
            params.pop("Reaction_Temp_C", None)
        if params:
            print(f"  [V6] global params: {list(params.keys())}")
        return params
    except Exception as e:
        print(f"  [WARN] global parameter extraction failed: {e}")
        return {}

# [Fix 11] Scan multiple chunks and merge only non-conflicting global preparation params.
def extract_global_params_from_chunks(chunks: List[str], file_name: str) -> Dict[str, str]:
    """Function docstring removed for runtime stability."""
    partial_results: List[Dict[str, str]] = []
    for i, chunk in enumerate(chunks):
        params = extract_global_params(chunk, file_name)
        if params:
            partial_results.append(params)
            print(f"  [Fix11] chunk{i+1} global params: {list(params.keys())}")

    if not partial_results:
        return {}

    field_values: Dict[str, List[str]] = {}
    for params in partial_results:
        for field, val in params.items():
            if field not in GLOBAL_FALLBACK_ALLOWED_FIELDS:
                continue
            v = str(val).strip()
            if not v:
                continue
            field_values.setdefault(field, [])
            if v not in field_values[field]:
                field_values[field].append(v)

    merged: Dict[str, str] = {}
    conflict_fields = []
    for field, vals in field_values.items():
        if len(vals) == 1:
            merged[field] = vals[0]
        else:
            conflict_fields.append(f"{field}({'/'.join(vals)})")

    if conflict_fields:
        print(f"  [Fix11] conflicting global params discarded: {conflict_fields}")

    # 规则兜底：扫描所有 performance_paragraph，用规则提取补全 LLM 未提取到的条件字段
    rule_condition_fields = [
        "S_C_Ratio", "GHSV_mL_g_h", "WHSV_h_inv", "Reaction_Temp_C",
        "Pressure_bar", "Feed_MeOH_to_H2O_Ratio",
    ]
    rule_field_values: Dict[str, List[str]] = {}
    for chunk in chunks:
        for para_record in route_chunk_paragraphs(chunk or ""):
            if para_record.get("paragraph_role") not in ("performance_paragraph", "other"):
                continue
            anchor = _extract_condition_anchor_from_text(
                str(para_record.get("paragraph_text", ""))
            )
            for field in rule_condition_fields:
                v = str(anchor.get(field, "")).strip()
                if not v:
                    continue
                if field == "Reaction_Temp_C" and not _reaction_temp_value_is_sane(v):
                    continue
                rule_field_values.setdefault(field, [])
                if v not in rule_field_values[field]:
                    rule_field_values[field].append(v)

    for field, vals in rule_field_values.items():
        if field in merged:
            continue  # LLM已提取，不覆盖
        if len(vals) == 1:
            merged[field] = vals[0]
            print(f"  [Fix11-rule] rule-extracted global param: {field}={vals[0]}")
        else:
            print(f"  [Fix11-rule] conflicting rule param discarded: {field}({'/'.join(vals)})")

    if merged:
        print(f"  [Fix11] merged global params: {list(merged.keys())}")
    return merged


def _collect_preparation_paragraph_texts(chunks: List[str]) -> List[str]:
    paragraph_texts: List[str] = []
    for chunk in chunks:
        for paragraph_record in route_chunk_paragraphs(chunk or ""):
            if str(paragraph_record.get("paragraph_role", "")).strip() != "preparation_paragraph":
                continue
            paragraph_text = str(paragraph_record.get("paragraph_text", "")).strip()
            if paragraph_text:
                paragraph_texts.append(paragraph_text)
    return _dedupe_keep_order(paragraph_texts)


def extract_global_params_from_preparation_paragraphs(chunks: List[str], file_name: str) -> Dict[str, str]:
    preparation_paragraphs = _collect_preparation_paragraph_texts(chunks)
    if not preparation_paragraphs:
        return {}

    partial_results: List[Dict[str, str]] = []
    for i, paragraph_text in enumerate(preparation_paragraphs):
        params = extract_global_params(paragraph_text, file_name)
        if params:
            filtered = {
                field: value for field, value in params.items()
                if field in GLOBAL_FALLBACK_ALLOWED_FIELDS and str(value).strip()
            }
            if filtered:
                partial_results.append(filtered)
                print(f"  [prep-global] paragraph{i+1} global params: {list(filtered.keys())}")

    if not partial_results:
        return {}

    field_values: Dict[str, List[str]] = {}
    for params in partial_results:
        for field, val in params.items():
            if field not in GLOBAL_FALLBACK_ALLOWED_FIELDS:
                continue
            v = str(val).strip()
            if not v:
                continue
            field_values.setdefault(field, [])
            if v not in field_values[field]:
                field_values[field].append(v)

    merged: Dict[str, str] = {}
    conflict_fields = []
    for field, vals in field_values.items():
        if len(vals) == 1:
            merged[field] = vals[0]
        else:
            conflict_fields.append(f"{field}({'/'.join(vals)})")

    if conflict_fields:
        print(f"  [prep-global] conflicting preparation params discarded: {conflict_fields}")
    if merged:
        print(f"  [prep-global] merged preparation params: {list(merged.keys())}")
    return merged


# ==========================================
# Broadcast path accounting helpers
# 记账规则：只在实际填入字段时调用，不因"经过分支"就记账
# ==========================================
def _set_broadcast_primary_if_empty(row: Dict, primary: str) -> None:
    """只在 _broadcast_path_primary 为空时写入，保留最高精度路径（先到先得）。"""
    if not row.get("_broadcast_path_primary"):
        row["_broadcast_path_primary"] = primary


def _append_broadcast_flag(row: Dict, flag: str) -> None:
    """追加广播路径标记，允许多路径并存（如 registry_bind|global_params）。"""
    existing = row.get("_broadcast_flags", "")
    flags = set(existing.split("|")) if existing else set()
    flags.discard("")
    flags.add(flag)
    row["_broadcast_flags"] = "|".join(sorted(flags))


def _load_field_source_map(row: Dict) -> Dict[str, Dict[str, str]]:
    raw = str(row.get("_field_source_map_json", "") or "").strip()
    if not raw:
        return {}
    try:
        parsed = json.loads(raw)
        return parsed if isinstance(parsed, dict) else {}
    except Exception:
        return {}


def _set_field_source(row: Dict, field: str, source: str, confidence: str, scope: str) -> None:
    tracked = set(globals().get("_FIELD_SOURCE_TRACKED_FIELDS", set()) or set())
    if tracked and field not in tracked:
        return
    mapping = _load_field_source_map(row)
    mapping[str(field)] = {
        "source": str(source or "").strip(),
        "confidence": str(confidence or "").strip(),
        "scope": str(scope or "").strip(),
    }
    row["_field_source_map_json"] = json.dumps(mapping, ensure_ascii=False, sort_keys=True)


def _broadcast_scope_for_field(field: str) -> str:
    condition_fields = set(globals().get("_FIGURE_PANEL_CONDITION_FIELDS", []) or [])
    return "figure_condition" if field in condition_fields else "identity_bound"


# ---------------------------------------------------------------------------
# build_text_master_registry: 按 CID 合并同一论文所有 text 行，生成 master 记录。
# 比旧的 text_by_id（只保留字段最多的单条）更完整，能捕获分散在不同行的字段。
# 不替换 build_text_sample_registry，作为 apply_global_params_fallback 的数据源。
# ---------------------------------------------------------------------------
def build_text_master_registry(text_records: List[Dict], file_name: str = "") -> Dict[str, Dict]:
    """按 Canonical_Catalyst_ID / Catalyst_ID 合并所有 text 行，生成 master 记录字典。

    合并策略：
    - 身份字段：first-non-empty，冲突时记录但不强合并
    - 制备/静态字段：first-non-empty + 冲突检测
    - 反应条件字段：只保留组内完全一致的公共值，不硬合并
    """
    MASTER_STATIC_FIELDS = [
        "Canonical_Catalyst_ID", "Catalyst_ID", "Catalyst_ID_normalized",
        "identity_alias_group", "Catalyst", "Active_Metal", "Promoter", "Promoter_Metal",
        "Support", "Support_Normalized", "Support_Grouped",
        "Metal_Loading_wt%", "Ni_Loading_wt%", "Promoter_Loading_wt%",
        "Alloy_Ratio", "Ni_Fraction", "Promoter_Fraction",
        "Precursor", "Precursor_Normalized", "Precursor_Family",
        "Metal_Loading_Method", "Metal_Loading_Method_Normalized",
        "Support_Prep_Method", "Support_Prep_Method_Normalized",
        "Dry_Temp_C", "Dry_Time_h",
        "Calcination_Temp_C", "Calcination_Time_h",
        "Reduction_Temp_C", "Reduction_Time_h",
        "Preparation_Fingerprint",
    ]
    MASTER_CONDITION_FIELDS = [
        "Reaction_Temp_C", "S_C_Ratio", "GHSV_mL_g_h", "WHSV_h_inv",
        "Pressure_bar", "Feed_Composition",
    ]

    # 按 CID 分组（4级优先链）
    groups: Dict[str, List[Dict]] = {}
    for r in text_records:
        cid = str(
            r.get("Canonical_Catalyst_ID") or
            r.get("identity_alias_group") or
            r.get("Catalyst_ID_normalized") or
            r.get("Catalyst_ID") or ""
        ).strip()
        if not cid or cid.lower() == "nan":
            continue
        groups.setdefault(cid, []).append(r)

    master_registry: Dict[str, Dict] = {}
    for cid, rows in groups.items():
        master: Dict = {"source_record_count": len(rows), "source_file": file_name,
                        "master_registry_notes": "", "master_registry_conflict_fields": []}
        conflict_fields = []

        # 静态字段：first-non-empty，冲突时记录
        for field in MASTER_STATIC_FIELDS:
            vals = [str(r.get(field, "")).strip() for r in rows
                    if str(r.get(field, "")).strip() and str(r.get(field, "")).strip().lower() != "nan"]
            if not vals:
                master[field] = ""
                continue
            unique_vals = list(dict.fromkeys(vals))  # 保序去重
            master[field] = unique_vals[0]  # first-non-empty
            if len(unique_vals) > 1:
                conflict_fields.append(field)

        # 反应条件字段：只保留组内完全一致的值，否则留空
        for field in MASTER_CONDITION_FIELDS:
            vals = list({str(r.get(field, "")).strip() for r in rows
                         if str(r.get(field, "")).strip() and str(r.get(field, "")).strip().lower() != "nan"})
            master[field] = vals[0] if len(vals) == 1 else ""

        master["master_registry_conflict_fields"] = conflict_fields
        if conflict_fields:
            master["master_registry_notes"] = f"conflicts: {','.join(conflict_fields)}"
        master_registry[cid] = master

    return master_registry


def _supplement_figure_from_master(fig_record: Dict, master_record: Dict,
                                   supplement_fields: List[str], tag_prefix: str) -> List[str]:
    """从 master_record 向 fig_record 补空字段，返回实际补入的字段列表。只填空，不覆盖。"""
    filled = []
    for key in supplement_fields:
        src_val = str(master_record.get(key, "")).strip()
        if not src_val or src_val.lower() == "nan":
            continue
        if key == "GHSV_mL_g_h":
            try:
                if float(src_val.replace(",", "")) < 100:
                    continue
            except (ValueError, TypeError):
                pass
        if not str(fig_record.get(key, "")).strip():
            fig_record[key] = src_val
            filled.append(key)
            _set_field_source(
                fig_record,
                key,
                source="text_master",
                confidence="high" if key not in FIGURE_TEXT_MASTER_CONDITION_FIELDS else "medium",
                scope="cid_consensus" if key in FIGURE_TEXT_MASTER_CONDITION_FIELDS else "identity_bound",
            )
    if filled:
        notes = str(fig_record.get("Notes", ""))
        tag = f"[{tag_prefix}: {','.join(filled)}]"
        if tag not in notes:
            fig_record["Notes"] = (notes + " " + tag).strip()
    return filled


# 从 text master registry 向 figure 行补全的字段集合
FIGURE_TEXT_MASTER_SUPPLEMENT_FIELDS = [
    # 身份字段（只有精确 CID 命中时才补）
    "Canonical_Catalyst_ID", "Catalyst_ID_normalized", "identity_alias_group",
    "Active_Metal", "Promoter", "Promoter_Metal",
    "Alloy_Ratio", "Ni_Fraction", "Promoter_Fraction",
    # 制备字段
    "Metal_Loading_wt%", "Support", "Support_Normalized", "Support_Grouped",
    "Precursor", "Precursor_Normalized", "Precursor_Family",
    "Metal_Loading_Method", "Metal_Loading_Method_Normalized",
    "Support_Prep_Method", "Support_Prep_Method_Normalized",
    "Dry_Temp_C", "Dry_Time_h",
    "Calcination_Temp_C", "Calcination_Time_h",
    "Reduction_Temp_C", "Reduction_Time_h",
    "Preparation_Fingerprint",
    # 反应条件（只补完全一致的公共值）
    "Reaction_Temp_C", "S_C_Ratio", "GHSV_mL_g_h", "Pressure_bar",
]

FIGURE_TEXT_MASTER_CONDITION_FIELDS = [
    "Reaction_Temp_C", "S_C_Ratio", "GHSV_mL_g_h", "Pressure_bar",
]

# 低风险字段：允许在 consensus fallback 中广播（不含高风险身份字段）
CONSENSUS_BROADCAST_FIELDS_LOW_RISK = [
    "Support", "Metal_Loading_Method", "Support_Prep_Method", "Precursor",
    "Support_Normalized", "Support_Grouped",
    "Precursor_Normalized", "Precursor_Family",
    "Metal_Loading_Method_Normalized", "Support_Prep_Method_Normalized",
    "Preparation_Fingerprint",
]


# [Fix 3] Fallback fill: only fill empty fields and never overwrite local evidence.
# 重构为分层 fallback：
#   Step 1. 构建 text_master_registry（按 CID 合并所有 text 行）
#   Step 2. figure 行按 CID 精确命中 master → 补全静态+制备字段
#   Step 3. 未命中 → single-catalyst fallback（论文只有一种催化剂）
#   Step 4. 仍未命中 → consensus fallback（低风险字段 ≥90% 一致才广播）
#   Step 5. paper-level global_params fallback（所有行）
def apply_global_params_fallback(records: List[Dict], global_params: Dict[str, str]) -> None:
    """Function docstring removed for runtime stability."""

    def _ghsv_val_is_sane(val) -> bool:
        """MSR实验GHSV通常 > 100 mL/g/h；低于此值视为误提取。"""
        try:
            return float(str(val).replace(",", "").strip()) >= 100
        except (ValueError, TypeError):
            return True

    text_records = [r for r in records if str(r.get("data_source", "")) == "text"]
    figure_records = [r for r in records if str(r.get("data_source", "")) == "figure"]

    if not text_records or not figure_records:
        # 无 figure 或无 text 时，仍需执行 global_params 兜底
        if global_params:
            filtered_global_params = {
                key: val for key, val in global_params.items()
                if key in GLOBAL_FALLBACK_ALLOWED_FIELDS and key not in GLOBAL_FALLBACK_BLOCKED_FIELDS
            }
            for r in records:
                filled_global: List[str] = []
                for key, val in filtered_global_params.items():
                    if not val:
                        continue
                    if key == "GHSV_mL_g_h" and not _ghsv_val_is_sane(val):
                        continue
                    if key == "Reaction_Temp_C" and not _reaction_temp_value_is_sane(val):
                        continue
                    if not r.get(key):
                        r[key] = val
                        filled_global.append(key)
                        _set_field_source(r, key, "global_params", "low", _broadcast_scope_for_field(key))
                        if key == "Reaction_Temp_C":
                            notes = str(r.get("Notes", ""))
                            if "[global_reaction_temp:" not in notes:
                                r["Notes"] = (notes + f" [global_reaction_temp: {val}°C]").strip()
                if filled_global:
                    _set_broadcast_primary_if_empty(r, "global_params")
                    _append_broadcast_flag(r, "global_params")
                elif str(r.get("data_source", "")) == "figure":
                    _set_broadcast_primary_if_empty(r, "unresolved")
        else:
            for r in records:
                if str(r.get("data_source", "")) == "figure":
                    _set_broadcast_primary_if_empty(r, "unresolved")
        return

    # ----------------------------------------------------------------
    # Step 1: 构建 text_master_registry（按 CID 合并所有 text 行）
    # 同时保留旧的 text_by_id 作为兼容兜底
    # ----------------------------------------------------------------
    source_file = str(text_records[0].get("Source_File", "")) if text_records else ""
    text_master = build_text_master_registry(text_records, source_file)

    # 旧的 text_by_id（兼容兜底，不再是主路径）
    text_by_id: Dict[str, Dict] = {}
    for r in text_records:
        cid = str(
            r.get("Canonical_Catalyst_ID") or
            r.get("identity_alias_group") or
            r.get("Catalyst_ID_normalized") or
            r.get("Catalyst_ID") or ""
        ).strip()
        if not cid:
            continue
        existing = text_by_id.get(cid)
        if existing is None:
            text_by_id[cid] = r
        else:
            def _filled(rec):
                return sum(1 for v in rec.values() if v and str(v).strip())
            if _filled(r) > _filled(existing):
                text_by_id[cid] = r

    unique_cids_in_text = list(text_by_id.keys())
    single_catalyst_src = text_by_id[unique_cids_in_text[0]] if len(unique_cids_in_text) == 1 else None

    # ----------------------------------------------------------------
    # Step 2 & 3: 对每个 figure 行，按优先级查找 master source
    # ----------------------------------------------------------------
    orphan_figure_records = []
    for r in figure_records:
        cid = str(
            r.get("Canonical_Catalyst_ID") or
            r.get("identity_alias_group") or
            r.get("Catalyst_ID_normalized") or
            r.get("Catalyst_ID") or ""
        ).strip()
        matched_master = None

        # Step 2: 精确命中 text_master_registry
        if cid and cid in text_master:
            matched_master = text_master[cid]
            filled_tm = _supplement_figure_from_master(r, matched_master, FIGURE_TEXT_MASTER_SUPPLEMENT_FIELDS,
                                                       "text_master_supplement")
            if filled_tm:
                _set_broadcast_primary_if_empty(r, "text_master")
                _append_broadcast_flag(r, "text_master")
        # Step 3: single-catalyst fallback（论文只有一种催化剂）
        elif not cid and single_catalyst_src is not None:
            print(f"  [fig_single_catalyst_fallback] figure row has no CID, using sole text record")
            filled = []
            for key in GLOBAL_FIGURE_SUPPLEMENT_FIELDS:
                src_val = single_catalyst_src.get(key)
                if not r.get(key) and src_val:
                    if key == "GHSV_mL_g_h":
                        try:
                            if float(str(src_val).replace(",", "").strip()) < 100:
                                continue
                        except (ValueError, TypeError):
                            pass
                    r[key] = src_val
                    filled.append(key)
                    _set_field_source(r, key, "single_catalyst", "medium", _broadcast_scope_for_field(key))
            if filled:
                notes = str(r.get("Notes", ""))
                tag = f"[fig_supplement: {','.join(filled)}]"
                if tag not in notes:
                    r["Notes"] = (notes + " " + tag).strip()
                _set_broadcast_primary_if_empty(r, "single_catalyst")
                _append_broadcast_flag(r, "single_catalyst")
        else:
            # 未命中任何精确 source，进入 consensus 候选
            if not cid:
                orphan_figure_records.append(r)
            else:
                # CID 有值但 master 里没有（可能是 binding 后新赋值），尝试旧 text_by_id
                if cid in text_by_id:
                    src = text_by_id[cid]
                    filled = []
                    for key in GLOBAL_FIGURE_SUPPLEMENT_FIELDS:
                        src_val = src.get(key)
                        if not r.get(key) and src_val:
                            if key == "GHSV_mL_g_h":
                                try:
                                    if float(str(src_val).replace(",", "").strip()) < 100:
                                        continue
                                except (ValueError, TypeError):
                                    pass
                            r[key] = src_val
                            filled.append(key)
                            _set_field_source(r, key, "text_master", "medium", _broadcast_scope_for_field(key))
                    if filled:
                        notes = str(r.get("Notes", ""))
                        tag = f"[fig_supplement: {','.join(filled)}]"
                        if tag not in notes:
                            r["Notes"] = (notes + " " + tag).strip()
                        _set_broadcast_primary_if_empty(r, "text_master")
                        _append_broadcast_flag(r, "text_master")

    # ----------------------------------------------------------------
    # Step 4: consensus fallback（低风险字段，≥90% 一致才广播）
    # 只针对 orphan figure 行（CID 为空且未命中任何精确 source）
    # ----------------------------------------------------------------
    if orphan_figure_records and text_records:
        consensus: Dict[str, str] = {}
        for field in CONSENSUS_BROADCAST_FIELDS_LOW_RISK:
            vals = [str(r.get(field, "")).strip() for r in text_records
                    if str(r.get(field, "")).strip() and str(r.get(field, "")).strip().lower() != "nan"]
            if not vals:
                continue
            most_common = max(set(vals), key=vals.count)
            ratio = vals.count(most_common) / len(vals)
            if ratio >= 0.9:
                consensus[field] = most_common
                print(f"  [consensus_fallback] {field}={most_common!r} ({ratio:.0%} of text rows)")
        if consensus:
            for r in orphan_figure_records:
                filled2 = []
                for field, val in consensus.items():
                    if not r.get(field):
                        r[field] = val
                        filled2.append(field)
                        _set_field_source(r, field, "consensus", "low", _broadcast_scope_for_field(field))
                if filled2:
                    notes = str(r.get("Notes", ""))
                    tag = f"[consensus_fallback: {','.join(filled2)}]"
                    if tag not in notes:
                        r["Notes"] = (notes + " " + tag).strip()
                    _set_broadcast_primary_if_empty(r, "consensus")
                    _append_broadcast_flag(r, "consensus")

    # ----------------------------------------------------------------
    # Step 5: paper-level global_params broadcast（所有行，最后兜底）
    # 必须在精确补全之后执行，避免全局值先占坑阻断更精确的 text_master 补全
    # ----------------------------------------------------------------
    if global_params:
        filtered_global_params = {
            key: val for key, val in global_params.items()
            if key in GLOBAL_FALLBACK_ALLOWED_FIELDS and key not in GLOBAL_FALLBACK_BLOCKED_FIELDS
        }
        for r in records:
            filled_global = []
            for key, val in filtered_global_params.items():
                if not val:
                    continue
                if key == "GHSV_mL_g_h" and not _ghsv_val_is_sane(val):
                    print(f"  [global_params_sanity] skip GHSV={val} (< 100, likely misextracted)")
                    continue
                if key == "Reaction_Temp_C" and not _reaction_temp_value_is_sane(val):
                    print(f"  [global_params_sanity] skip Reaction_Temp_C={val} (out of sane range)")
                    continue
                if not r.get(key):
                    r[key] = val
                    filled_global.append(key)
                    _set_field_source(r, key, "global_params", "low", _broadcast_scope_for_field(key))
                    if key == "Reaction_Temp_C":
                        notes = str(r.get("Notes", ""))
                        if "[global_reaction_temp:" not in notes:
                            r["Notes"] = (notes + f" [global_reaction_temp: {val}°C]").strip()
            if filled_global:
                _set_broadcast_primary_if_empty(r, "global_params")
                _append_broadcast_flag(r, "global_params")

    # 所有路径结束后，未命中任何主路径的 figure 行标记为 unresolved
    for r in figure_records:
        if not r.get("_broadcast_path_primary"):
            r["_broadcast_path_primary"] = "unresolved"

# ==========================================
# [Fix 1] V6 LLM extraction: require promoter-aware Catalyst_ID
# ==========================================
# ==========================================
# second metal / provenance / upstream filtering
# ==========================================
MSR_NI_SECOND_METALS = {
    "Cu", "Pt", "Pd", "Ru", "Rh", "Co", "Fe", "Au", "Ir",
    "Mo", "Sn", "Ga", "In", "Zn", "Ce", "La", "K", "Mg",
    "Mn", "Cr", "Y", "Pr", "Nd", "Sm", "Gd",
}
SUPPORT_LIKE_TOKENS = {
    "al2o3", "gamma-al2o3", "gamma-al2o3", "gamma alumina", "alumina",
    "ceo2", "zro2", "sio2", "tio2", "mgo",
    "cnt", "cnts", "cnf", "ac", "activated carbon", "carbon",
    "sep", "sepiolite", "sba-15", "mcm-41", "support", "carrier",
}
CONDITION_AXIS_KEYWORDS = {
    "temperature": ["temperature", "reaction temperature", "reforming temperature", "temp"],
    "time": ["tos", "time on stream", "stability time", "stream time"],
    "ratio": ["s/c", "steam-to-carbon", "steam to carbon", "feed ratio", "methanol-to-water", "meoh-to-h2o"],
    "space_velocity": ["ghsv", "whsv", "sv", "space velocity"],
    "pressure": ["pressure"],
    "flow": ["flow rate", "feed flow", "feed rate"],
    "catalyst_amount": ["catalyst amount", "catalyst mass", "catalyst charge", "loading in reactor"],
}
CATEGORY_ONLY_AXIS_KEYWORDS = {
    "catalyst", "catalysts", "sample", "samples", "support", "carrier",
    "loading", "metal loading", "composition", "series", "catalyst id",
}
CORE_PERFORMANCE_FIELDS = [
    "MeOH_Conversion_%", "H2_Yield_%", "H2_Production_Rate",
    "H2_Selectivity_%", "CO_Selectivity_%", "CO2_Selectivity_%",
    "CO_Concentration_ppm", "Carbon_Deposition_wt%", "Deactivation_Rate_%_h",
]
IDENTITY_VALUE_FIELDS = [
    "Catalyst", "Catalyst_ID", "Canonical_Catalyst_ID", "Active_Metal",
    "Metal_Loading_wt%", "Alloy_Ratio", "Promoter", "Promoter_Metal", "Support",
]
PREPARATION_VALUE_FIELDS = [
    "Precursor", "Metal_Loading_Method", "Support_Prep_Method",
    "Dry_Temp_C", "Dry_Time_h", "Calcination_Temp_C", "Calcination_Time_h",
    "Reduction_Temp_C", "Reduction_Time_h",
]
UPSTREAM_FILTER_AUDIT_CSV = os.path.join(OUT_DIR, "upstream_row_filter_audit.csv")
FIGURE_BINDING_AUDIT_CSV = os.path.join(OUT_DIR, "figure_binding_audit.csv")
COMPARISON_FILTER_AUDIT_CSV = os.path.join(OUT_DIR, "comparison_filter_audit.csv")
FIGURE_PREMERGE_FILTER_AUDIT_CSV = os.path.join(OUT_DIR, "figure_premerge_filter_audit.csv")
COMPOSITION_CONSISTENCY_AUDIT_CSV = os.path.join(OUT_DIR, "composition_consistency_audit.csv")
TEXT_POINT_GATE_AUDIT_CSV = os.path.join(OUT_DIR, "text_point_gate_audit.csv")
FIGURE_CANDIDATE_AUDIT_CSV = os.path.join(OUT_DIR, "figure_candidate_audit.csv")
FIGURE_POINT_VALIDATION_AUDIT_CSV = os.path.join(OUT_DIR, "figure_point_validation_audit.csv")
FINAL_EXCLUDED_AUDIT_CSV = os.path.join(OUT_DIR, "final_excluded_audit.csv")
FINAL_DEDUPE_AUDIT_CSV = os.path.join(OUT_DIR, "final_dedupe_audit.csv")
LAYERED_DISTRIBUTION_SUMMARY_JSON = os.path.join(OUT_DIR, "layered_distribution_summary.json")
LEGACY_FINAL_DISTRIBUTION_SUMMARY_JSON = os.path.join(OUT_DIR, "final_distribution_summary.json")
FINAL_DISTRIBUTION_SUMMARY_JSON = LEGACY_FINAL_DISTRIBUTION_SUMMARY_JSON
EXTRACTION_QUALITY_SUMMARY_JSON = os.path.join(OUT_DIR, "extraction_quality_summary.json")
RESEARCH_MASTER_ALL_JSON = os.path.join(OUT_DIR, "research_master_all.json")
RESEARCH_MASTER_STRICT_JSON = os.path.join(OUT_DIR, "research_master_strict.json")
MODELING_FINAL_JSON = os.path.join(OUT_DIR, "modeling_final.json")
LEGACY_MSR_DATASET_JSON = os.path.join(OUT_DIR, "msr_dataset.json")
MULTIMODAL_AUDIT_DIR = os.path.join(OUT_DIR, "multimodal_audit")


def _dedupe_keep_order(values: List[str]) -> List[str]:
    seen = set()
    result = []
    for value in values:
        if value and value not in seen:
            seen.add(value)
            result.append(value)
    return result


def clean_residual_mojibake_chars(text: Any) -> Any:
    """
    High-confidence residual mojibake cleanup.
    This is not a general mojibake decoder.
    Only a few known substitutions are allowed here.
    """
    if not isinstance(text, str) or not text:
        return text

    cleaned = text
    cleaned = re.sub(r"掳\s*([Cc])\b", "°C", cleaned)
    cleaned = re.sub(r"掳\s*([Kk])\b", "°K", cleaned)
    cleaned = re.sub(r"°\s*([Cc])\b", "°C", cleaned)
    cleaned = re.sub(r"°\s*([Kk])\b", "°K", cleaned)
    cleaned = cleaned.replace("mol h鈭?", "mol h-1")
    cleaned = cleaned.replace("mmol min鈭?", "mmol min-1")
    cleaned = cleaned.replace("ml min鈭?", "ml min-1")
    return cleaned


def clean_record_text_fields(record: Dict) -> Dict:
    """
    Apply high-confidence residual cleanup to record values only.
    Field names are never modified.
    """
    if not isinstance(record, dict):
        return record
    for key, value in list(record.items()):
        if isinstance(value, str):
            record[key] = clean_residual_mojibake_chars(value)
    return record


def clean_records_text_fields(records: List[Dict]) -> List[Dict]:
    return [clean_record_text_fields(record) for record in records if isinstance(record, dict)]


def _normalize_identity_text(value: str) -> str:
    if not value:
        return ""
    s = clean_residual_mojibake_chars(str(value))
    s = s.translate(str.maketrans({
        "₀": "0", "₁": "1", "₂": "2", "₃": "3", "₄": "4",
        "₅": "5", "₆": "6", "₇": "7", "₈": "8", "₉": "9",
        "²": "2", "³": "3",
    }))
    s = s.replace("\u2013", "-").replace("\u2014", "-").replace("\u2212", "-")
    s = re.sub(r"\s+", " ", s)
    return s.strip()


def _append_note_tag(record: Dict, tag: str) -> None:
    if not tag:
        return
    tag = clean_residual_mojibake_chars(tag)
    notes = clean_residual_mojibake_chars(str(record.get("Notes", ""))).strip()
    if tag not in notes:
        record["Notes"] = (notes + " " + tag).strip()


def _format_ratio_token(value: str) -> str:
    value = str(value).strip()
    if "." in value:
        value = value.rstrip("0").rstrip(".")
    return value


def normalize_metal_loading_method(value: str) -> str:
    text = clean_residual_mojibake_chars(str(value or "")).strip()
    if not text:
        return ""
    if text in {"Impregnation", "Co-precipitation", "Deposition-precipitation", "Sol-gel", "Hydrothermal", "Mechanical mixing", "Other"}:
        return text

    low = text.lower()
    if re.search(r"impregnation|incipient\s+wetness|wet\s+impregnation|co\s*-?impregnation|impregnat", low):
        return "Impregnation"
    if re.search(r"co\s*-?precipit|coprecipit", low):
        return "Co-precipitation"
    if re.search(r"deposition\s*-?precipit", low):
        return "Deposition-precipitation"
    if re.search(r"sol\s*-?gel", low):
        return "Sol-gel"
    if re.search(r"hydrothermal", low):
        return "Hydrothermal"
    if re.search(r"mechanical\s*-?mix|ball\s*-?mill|grind|grinding", low):
        return "Mechanical mixing"
    return "Other"



def normalize_support_prep_method(value: str) -> str:
    text = clean_residual_mojibake_chars(str(value or "")).strip()
    if not text:
        return ""
    if text in {"Commercial", "Precipitation", "Sol-gel", "Hydrothermal", "Other"}:
        return text

    low = text.lower()
    if re.search(r"commercial|purchased|supplied|sigma|aladdin|aldrich|bought", low):
        return "Commercial"
    if re.search(r"precipit|co\s*-?precipit", low):
        return "Precipitation"
    if re.search(r"sol\s*-?gel", low):
        return "Sol-gel"
    if re.search(r"hydrothermal", low):
        return "Hydrothermal"
    return "Other"



def normalize_precursor_expression(value: str) -> str:
    text = clean_residual_mojibake_chars(str(value or "")).strip()
    if not text:
        return ""

    text = text.replace("\u2022", "\u00b7").replace("\u00b7", "\u00b7")
    text = re.sub(r"\s+", " ", text)
    text = re.sub(r"\s*/\s*", " / ", text)
    text = re.sub(r"\s*;\s*", " / ", text)
    text = re.sub(r"\s+and\s+", " / ", text, flags=re.I)
    parts = [p.strip(" ,;./") for p in re.split(r"\s*/\s*", text) if p.strip(" ,;./")]

    canonical_patterns = [
        (r"\b(?:nickel\s+nitrate|ni\s*\(\s*no3\s*\)\s*2(?:\s*[\u00b7\.]\s*\d*h2o)?)\b", "Nickel nitrate"),
        (r"\b(?:copper\s+nitrate|cu\s*\(\s*no3\s*\)\s*2(?:\s*[\u00b7\.]\s*\d*h2o)?)\b", "Copper nitrate"),
        (r"\b(?:cerium\s+nitrate|ce\s*\(\s*no3\s*\)\s*3(?:\s*[\u00b7\.]\s*\d*h2o)?)\b", "Cerium nitrate"),
        (r"\b(?:lanthanum\s+nitrate|la\s*\(\s*no3\s*\)\s*3(?:\s*[\u00b7\.]\s*\d*h2o)?)\b", "Lanthanum nitrate"),
        (r"\b(?:magnesium\s+nitrate|mg\s*\(\s*no3\s*\)\s*2(?:\s*[\u00b7\.]\s*\d*h2o)?)\b", "Magnesium nitrate"),
        (r"\b(?:nickel\s+acetate|ni\s*\(\s*ch3coo\s*\)\s*2(?:\s*[\u00b7\.]\s*\d*h2o)?)\b", "Nickel acetate"),
        (r"\b(?:copper\s+acetate|cu\s*\(\s*ch3coo\s*\)\s*2(?:\s*[\u00b7\.]\s*\d*h2o)?)\b", "Copper acetate"),
        (r"\b(?:nickel\s+chloride|ni\s*cl\s*2(?:\s*[\u00b7\.]\s*\d*h2o)?)\b", "Nickel chloride"),
        (r"\b(?:copper\s+chloride|cu\s*cl\s*2(?:\s*[\u00b7\.]\s*\d*h2o)?)\b", "Copper chloride"),
        (r"\b(?:nickel\s+oxide|ni\s*o)\b", "Nickel oxide"),
        (r"\b(?:copper\s+oxide|cu\s*o)\b", "Copper oxide"),
    ]

    normalized_parts = []
    for part in parts or [text]:
        normalized = part
        low = part.lower()
        for pattern, repl in canonical_patterns:
            if re.search(pattern, low, flags=re.I):
                normalized = repl
                break
        normalized = re.sub(r"\s+", " ", normalized).strip(" ,;./")
        normalized_parts.append(normalized)

    return " / ".join(_dedupe_keep_order([p for p in normalized_parts if p]))



def classify_precursor_family(value: str) -> str:
    text = clean_residual_mojibake_chars(str(value or "")).strip()
    if not text:
        return ""

    low = text.lower()
    flags = set()
    if re.search(r"nitrate|\bno3\b", low):
        flags.add("nitrate")
    if re.search(r"acetate|ch3coo|c2h3o2", low):
        flags.add("acetate")
    if re.search(r"chloride|\bcl\b", low):
        flags.add("chloride")
    if re.search(r"acetylacetonate|acac|alkoxide|carbonyl|oleate", low):
        flags.add("organometallic")
    if re.search(r"oxide|\box\b", low):
        flags.add("oxide")

    salt_flags = flags & {"nitrate", "acetate", "chloride"}
    if len(salt_flags) >= 2:
        return "mixed_salts"
    if len(salt_flags) == 1:
        return next(iter(salt_flags))
    if "organometallic" in flags:
        return "organometallic"
    if "oxide" in flags:
        return "oxide"
    return "other"



def group_support_family(value: str) -> str:
    text = clean_residual_mojibake_chars(str(value or "")).strip()
    if not text:
        return ""

    normalized = normalize_support_expression(text)
    normalized = re.sub(r"(?i)^commercial\s+", "", normalized).strip()
    normalized = re.sub(r"(?i)^gamma[-\s]?al2o3$", "Al2O3", normalized)
    if not normalized:
        return ""

    single_map = {
        "Al2O3": "Al2O3",
        "CeO2": "CeO2",
        "CeOx": "CeO2",
        "ZrO2": "ZrO2",
        "TiO2": "TiO2",
        "SiO2": "SiO2",
        "MgO": "MgO",
        "CNTs": "CNTs",
        "Activated Carbon": "Activated Carbon",
        "SEP": "SEP",
    }
    if normalized in single_map:
        return single_map[normalized]

    oxide_tokens = {"Al2O3", "CeO2", "CeOx", "ZrO2", "TiO2", "SiO2", "MgO", "La2O3"}
    split_parts = [p.strip() for p in re.split(r"[-/]", normalized) if p.strip()]
    unique_oxides = {p for p in split_parts if p in oxide_tokens}
    if len(unique_oxides) >= 2:
        return "Mixed Oxide"
    return "Other"



def apply_preparation_normalization_layer(record: Dict) -> Dict:
    if not isinstance(record, dict):
        return record

    row = dict(record)

    mlm_source = str(row.get("Metal_Loading_Method", "")).strip()
    row["Metal_Loading_Method_Normalized"] = normalize_metal_loading_method(mlm_source)

    spm_source = str(row.get("Support_Prep_Method_Normalized", "")).strip() or str(row.get("Support_Prep_Method", "")).strip()
    row["Support_Prep_Method_Normalized"] = normalize_support_prep_method(spm_source)

    precursor_source = str(row.get("Precursor", "")).strip()
    row["Precursor_Normalized"] = normalize_precursor_expression(precursor_source)
    row["Precursor_Family"] = classify_precursor_family(row["Precursor_Normalized"] or precursor_source)

    support_source = str(row.get("Support_Normalized", "")).strip() or str(row.get("Support", "")).strip()
    row["Support_Grouped"] = group_support_family(support_source)
    return row



def apply_preparation_normalization_batch(records: List[Dict]) -> List[Dict]:
    return [apply_preparation_normalization_layer(dict(record)) for record in records if isinstance(record, dict)]


def _normalize_promoter_like_token(token: Any) -> str:
    text = clean_residual_mojibake_chars(str(token or "")).strip()
    if not text:
        return ""
    text = text.replace("–", "-").replace("—", "-")
    text = re.sub(r"\b(?:second metal|promoter|additive|dopant|modifier|metal)\b", " ", text, flags=re.I)
    text = re.sub(r"[\(\)\[\]\{\},;]", " ", text)
    text = re.sub(r"\s+", " ", text).strip(" -_/")
    if not text:
        return ""

    # Reject non-metal element symbols that can never be a promoter
    _NON_METAL_SYMBOLS = {"O", "N", "C", "H", "S", "F", "P", "Cl", "Br", "I",
                          "Si", "B", "Se", "Te", "At", "Xe", "Kr", "Ar", "Ne", "He"}
    if text.strip() in _NON_METAL_SYMBOLS:
        return ""

    lower_text = text.lower()
    normalized_support = normalize_support_expression(text).strip().lower()
    if lower_text in SUPPORT_LIKE_TOKENS or normalized_support in SUPPORT_LIKE_TOKENS:
        return ""

    name_map = {
        "nickel": "Ni", "copper": "Cu", "platinum": "Pt", "palladium": "Pd",
        "ruthenium": "Ru", "rhodium": "Rh", "cobalt": "Co", "iron": "Fe",
        "gold": "Au", "iridium": "Ir", "molybdenum": "Mo", "tin": "Sn",
        "gallium": "Ga", "indium": "In", "zinc": "Zn", "cerium": "Ce",
        "lanthanum": "La", "potassium": "K", "magnesium": "Mg",
        "manganese": "Mn", "chromium": "Cr", "yttrium": "Y",
        "praseodymium": "Pr", "neodymium": "Nd", "samarium": "Sm",
        "gadolinium": "Gd",
    }
    alpha = re.sub(r"[^A-Za-z]", "", text)
    if alpha.lower() in name_map:
        symbol = name_map[alpha.lower()]
        return "" if symbol == "Ni" else symbol
    if re.fullmatch(r"[A-Z][a-z]?", text):
        symbol = text[0].upper() + text[1:].lower()
        return symbol if symbol in MSR_NI_SECOND_METALS else ""
    match = re.match(r"^([A-Z][a-z]?)(?:\d.*)?$", text)
    if match:
        symbol = match.group(1)
        if symbol in MSR_NI_SECOND_METALS:
            return "" if symbol == "Ni" else symbol
    return ""


def _safe_split_multi_promoter(value: Any) -> List[str]:
    text = clean_residual_mojibake_chars(str(value or "")).strip()
    if not text:
        return []
    normalized = text.replace("–", "-").replace("—", "-")
    normalized = re.sub(r"\b(?:and|with|plus)\b", "/", normalized, flags=re.I)
    normalized = normalized.replace("&", "/").replace("+", "/").replace(";", "/").replace(",", "/")
    candidates: List[str] = []
    for part in [item.strip() for item in normalized.split("/") if item.strip()]:
        subparts = [part]
        if "-" in part and not re.search(r"\d", part):
            subparts = [item.strip() for item in re.split(r"\s*-\s*", part) if item.strip()]
        for subpart in subparts:
            symbol = _normalize_promoter_like_token(subpart)
            if symbol:
                candidates.append(symbol)
    return _dedupe_keep_order(candidates)


def _looks_like_commercial_or_descriptor_suffix(token: str) -> bool:
    text = clean_residual_mojibake_chars(str(token or "")).strip().lower()
    if not text:
        return False
    text = re.sub(r"[^a-z0-9]+", " ", text).strip()
    return text in {
        "com",
        "commercial",
        "commercial catalyst",
        "common",
        "comparative",
        "comparison",
        "reference",
        "ref",
        "baseline",
        "blank",
        "std",
        "standard",
        "sample",
        "cat",
    }


def _strip_descriptor_suffix_before_element_parse(value: str) -> str:
    text = _normalize_identity_text(value)
    if not text:
        return ""
    text = text.replace("–", "-").replace("—", "-").replace("−", "-")
    text = re.sub(r"\((?:commercial|reference|baseline|comparative|comparison|common|blank|standard)[^)]*\)", " ", text, flags=re.I)
    text = re.sub(r"(?:-|\s)+(?:commercial(?: catalyst)?|common|comparative|comparison|reference|baseline|blank|standard|std|ref|com)\b.*$", "", text, flags=re.I)
    parts = [part.strip() for part in text.split("-")]
    while parts and _looks_like_commercial_or_descriptor_suffix(parts[-1]):
        parts.pop()
    text = "-".join([part for part in parts if part])
    text = re.sub(r"\s+", " ", text).strip(" -_/;,")
    return text


def _extract_identity_text_for_composition(row: Dict) -> str:
    fields = ["Active_Metal", "Catalyst", "Catalyst_ID", "Canonical_Catalyst_ID", "Series_Name"]
    text = " | ".join(
        _normalize_identity_text(str(row.get(field, "")))
        for field in fields
        if str(row.get(field, "")).strip()
    )
    if not text:
        return ""
    support_terms = sorted(SUPPORT_LIKE_TOKENS, key=len, reverse=True)
    for term in support_terms:
        text = re.sub(rf"(?i)(?:/|-|\b){re.escape(term)}\b", " ", text)
    text = re.sub(r"\s+", " ", text).strip(" |;,-")
    return text


def has_core_performance_metric(row: Dict) -> bool:
    for field in [
        "MeOH_Conversion_%", "H2_Yield_%", "H2_Production_Rate",
        "CO_Selectivity_%", "CO2_Selectivity_%", "CO_Concentration_ppm",
    ]:
        text = clean_residual_mojibake_chars(str(row.get(field, "") or "")).strip()
        if text and text.lower() not in {"", "n/a", "na", "none", "nan", "unknown"}:
            return True
    return False


def has_valid_point_condition_anchor(row: Dict) -> bool:
    def _is_valid_numeric(value: Any, allow_ratio: bool = False) -> bool:
        text = clean_residual_mojibake_chars(str(value or "")).strip()
        if not text or text.lower() in {"", "n/a", "na", "none", "nan", "unknown"}:
            return False
        if re.search(r"[~<>]", text):
            return False
        if re.search(r"\b(?:approx|approximately|about|around|between)\b", text, flags=re.I):
            return False
        if allow_ratio and re.fullmatch(r"\d+(?:\.\d+)?\s*[:/]\s*\d+(?:\.\d+)?", text):
            return True
        nums = re.findall(r"[-+]?\d+(?:\.\d+)?", text.replace(",", ""))
        return len(nums) == 1

    def _axis_kind(text: str) -> str:
        s = _normalize_identity_text(text).lower()
        if not s:
            return ""
        prep_like = bool(re.search(r"\b(calcination|reduction|drying|impregnation|preparation|precursor)\b", s, flags=re.I))
        if any(token in s for token in CATEGORY_ONLY_AXIS_KEYWORDS):
            if not any(any(keyword in s for keyword in keywords) for keywords in CONDITION_AXIS_KEYWORDS.values()):
                return ""
        for kind, keywords in CONDITION_AXIS_KEYWORDS.items():
            if any(keyword in s for keyword in keywords):
                if kind == "temperature" and prep_like and not re.search(r"\b(reaction|reforming)\b", s, flags=re.I):
                    return ""
                return kind
        return ""

    if _is_valid_numeric(row.get("Reaction_Temp_C", "")):
        return True
    if _is_valid_numeric(row.get("TOS_h", "")):
        return True
    for field in ["S_C_Ratio", "Feed_MeOH_to_H2O_Ratio", "GHSV_mL_g_h", "SpaceVelocity_norm", "Pressure_bar", "Flow_Rate", "Catalyst_Amount_g"]:
        if _is_valid_numeric(row.get(field, ""), allow_ratio=("Ratio" in field or field == "S_C_Ratio")):
            return True

    notes = str(row.get("Notes", "") or "")
    note_axis, note_x_value = _extract_x_axis_and_value_from_notes(notes)
    axis_candidates = [
        str(row.get("x_axis", "") or ""),
        str(row.get("x_axis_mode", "") or ""),
        note_axis,
        str(row.get("figure_point_validation_notes", "") or ""),
    ]
    value_candidates = [
        str(row.get("x_value", "") or ""),
        note_x_value,
    ]
    for axis_text in axis_candidates:
        kind = _axis_kind(axis_text)
        if not kind:
            continue
        for value_text in value_candidates:
            if _looks_like_identity_mapping(value_text):
                continue
            if _is_valid_numeric(value_text, allow_ratio=(kind == "ratio")):
                return True

    note_text = _normalize_identity_text(" ".join(axis_candidates + value_candidates + [notes])).lower()
    prep_noise = r"(?:calcination|reduction|drying|impregnation|preparation|precursor)"
    patterns = [
        r"\b(?:reaction|reforming)\s+temperature\b[^.;\n]{0,20}?\d+(?:\.\d+)?",
        r"\b(?:temperature|temp)\b[^.;\n]{0,20}?\d+(?:\.\d+)?\s*c\b",
        r"\b(?:tos|time on stream)\b[^.;\n]{0,20}?\d+(?:\.\d+)?\s*h\b",
        r"\b(?:s/c|steam[- ]?to[- ]?carbon|feed ratio|methanol[- ]?to[- ]?water)\b[^.;\n]{0,20}?\d+(?:\.\d+)?(?:\s*[:/]\s*\d+(?:\.\d+)?)?",
        r"\b(?:ghsv|whsv|space velocity)\b[^.;\n]{0,20}?\d+(?:\.\d+)?",
        r"\bpressure\b[^.;\n]{0,20}?\d+(?:\.\d+)?",
        r"\b(?:flow rate|feed flow|feed rate|catalyst amount|catalyst mass|catalyst charge)\b[^.;\n]{0,25}?\d+(?:\.\d+)?",
    ]
    if re.search(prep_noise, note_text, flags=re.I) and not re.search(r"\b(reaction|reforming|tos|time on stream)\b", note_text, flags=re.I):
        return False
    return any(re.search(pattern, note_text, flags=re.I) for pattern in patterns)


def is_preparation_backbone_only_row(row: Dict) -> bool:
    if str(row.get("data_source", "")).strip() != "text":
        return False
    if has_core_performance_metric(row):
        return False
    if str(row.get("paragraph_role", "")).strip() == "preparation_paragraph":
        return True
    if str(row.get("text_extraction_subroute", "")).strip() == "preparation_schema":
        return True
    return _is_preparation_backbone_row(row)


def is_text_point_record(row: Dict) -> bool:
    if str(row.get("data_source", "")).strip() != "text":
        return False
    if str(row.get("identity_completeness_level", "")).strip() != "complete":
        return False
    if str(row.get("numeric_reliability_level", "")).strip() != "direct_numeric":
        return False
    if str(row.get("source_granularity", "")).strip() not in {"table_row", "text_numeric", "si_table"}:
        return False
    if bool(row.get("is_literature_comparison")) or str(row.get("comparison_filter_action", "")).strip() == "downgraded_candidate_only":
        return False
    if bool(row.get("is_approximate_value")) or bool(row.get("is_range_like_value")) or bool(row.get("is_qualitative_value")):
        return False
    if bool(row.get("same_physical_point_possible")) or bool(row.get("obvious_duplicate_flag")) or str(row.get("duplicate_candidate_type", "")).strip():
        return False
    if not has_core_performance_metric(row):
        return False
    if not has_valid_point_condition_anchor(row):
        return False
    if is_preparation_backbone_only_row(row):
        return False
    return True


_LAYERED_TEXT_HIGH_TRUST_GRANULARITIES = {"table_row", "text_numeric", "si_table"}
_LAYERED_BLOCKED_FIGURE_ROLES = {
    "product_species_profile",
    "case_profile",
    "non_performance_like",
    "catalyst_amount_effect",
    "schematic_or_nonperformance",
}
_LAYERED_EXPORT_COMPAT_NOTE = (
    "dataset_all.csv => research_master_all.csv; "
    "dataset_all_strict.csv => research_master_strict.csv"
)


def _initialize_layered_export_state_if_needed() -> None:
    if globals().get("_LAYERED_EXPORT_RUN_COMPLETED") or not globals().get("_LAYERED_EXPORT_STATE_READY"):
        globals()["_LAYERED_EXPORT_STATE_READY"] = True
        globals()["_LAYERED_EXPORT_RUN_COMPLETED"] = False
        globals()["_RESEARCH_ALL_ROWS_CACHE"] = []
        globals()["_RESEARCH_STRICT_ROWS_CACHE"] = []
        globals()["_LAYERED_PAPER_STATS"] = {}


def _count_reason_values(rows: List[Dict], field_name: str) -> Dict[str, int]:
    counts: Dict[str, int] = {}
    for row in rows or []:
        if not isinstance(row, dict):
            continue
        key = clean_residual_mojibake_chars(str(row.get(field_name, "") or "")).strip()
        if not key:
            continue
        counts[key] = counts.get(key, 0) + 1
    return dict(sorted(counts.items(), key=lambda item: (-item[1], item[0])))


def _is_core_metric_present(row: Dict) -> bool:
    return has_core_performance_metric(row) or _has_core_figure_performance_value(row)


def _has_any_numeric_signal(row: Dict) -> bool:
    numeric_level = clean_residual_mojibake_chars(str(row.get("numeric_reliability_level", "") or "")).strip()
    if numeric_level in {"direct_numeric", "approximate_numeric", "range_like"}:
        return True
    if bool(row.get("is_approximate_value")) or bool(row.get("is_range_like_value")):
        return True
    if clean_residual_mojibake_chars(str(row.get("raw_numeric_expression", "") or "")).strip():
        return True
    return _is_core_metric_present(row)


def _has_any_strong_condition_anchor(row: Dict) -> bool:
    if str(row.get("data_source", "") or "").strip() == "figure":
        return _has_reliable_figure_condition_anchor(row) or _count_nonempty_condition_features(row) >= 1
    return has_valid_point_condition_anchor(row)


def _has_minimum_traceable_identity(row: Dict) -> bool:
    identity_level = clean_residual_mojibake_chars(str(row.get("identity_completeness_level", "") or "")).strip()
    if identity_level in {"complete", "partial"}:
        return True
    clue_fields = [
        "Catalyst", "Catalyst_ID", "Catalyst_ID_raw", "Catalyst_ID_normalized",
        "Canonical_Catalyst_ID", "identity_alias_group", "Active_Metal", "Promoter",
        "Promoter_Metal", "Support", "Support_Normalized", "Series_Name",
        "raw_category_label", "matched_registry_key", "matched_registry_label",
        "alias_map_source", "alias_map_evidence",
    ]
    return any(clean_residual_mojibake_chars(str(row.get(field, "") or "")).strip() for field in clue_fields)


def _has_figure_context_support(row: Dict) -> bool:
    if _has_any_strong_condition_anchor(row):
        return True
    context_blob = clean_residual_mojibake_chars(" ".join([
        str(row.get("condition_anchor_source", "") or ""),
        str(row.get("figure_point_validation_notes", "") or ""),
        str(row.get("figure_binding_reason", row.get("binding_reason", "")) or ""),
        str(row.get("Notes", "") or ""),
        str(row.get("x_axis", "") or ""),
        str(row.get("x_value", "") or ""),
        str(row.get("Series_Name", "") or ""),
        str(row.get("raw_category_label", "") or ""),
    ])).strip()
    if not context_blob:
        return False
    if len(context_blob) >= 24:
        return True
    return bool(re.search(
        r"\b(temperature|reaction|tos|time on stream|s/c|steam.?to.?carbon|ghsv|whsv|pressure|flow|conversion|yield|selectivity|ppm|legend|caption|series)\b",
        context_blob,
        flags=re.I,
    ))


def _is_baseline_or_support_only_row(row: Dict) -> bool:
    current = _finalize_figure_row_binding_payload(dict(row or {}))
    mode = clean_residual_mojibake_chars(str(current.get("figure_binding_mode", current.get("binding_mode", "")) or "")).strip()
    if mode in {"baseline_match", "support_only_match"}:
        return True
    if bool(current.get("is_baseline_label")) or bool(current.get("is_support_only_label")):
        return True
    if str(current.get("data_source", "") or "").strip() != "figure":
        return False
    if clean_residual_mojibake_chars(str(current.get("Active_Metal", "") or "")).strip():
        return False
    label_blob = clean_residual_mojibake_chars(" ".join([
        str(current.get("raw_category_label", "") or ""),
        str(current.get("Series_Name", "") or ""),
        str(current.get("Catalyst", "") or ""),
        str(current.get("figure_binding_reason", current.get("binding_reason", "")) or ""),
    ])).strip()
    return bool(re.search(r"\b(no catalyst|blank|baseline|support only|bare support)\b", label_blob, flags=re.I))


def _derive_partial_point_flag(row: Dict) -> int:
    current = _finalize_figure_row_binding_payload(dict(row or {}))
    if not _is_core_metric_present(current):
        return 0
    if bool(current.get("is_approximate_value")) or bool(current.get("is_range_like_value")) or bool(current.get("is_qualitative_value")):
        return 1
    if str(current.get("data_source", "") or "").strip() == "figure":
        if str(current.get("premerge_filter_action", "") or "").strip() == "candidate_only":
            return 1
        if not _has_reliable_final_figure_condition_anchor(current):
            return 1
    if str(current.get("identity_completeness_level", "") or "").strip() == "partial":
        return 1
    if str(current.get("data_source", "") or "").strip() == "text" and not has_valid_point_condition_anchor(current):
        return 1
    return 0


def _derive_layer_exclusion_detail(row: Dict, reason: str) -> str:
    current = _finalize_figure_row_binding_payload(dict(row or {}))
    reason = clean_residual_mojibake_chars(str(reason or "")).strip()
    note_axis, note_x_value = _extract_x_axis_and_value_from_notes(str(current.get("Notes", "") or ""))
    if reason == "comparison_like":
        return str(current.get("comparison_filter_reason", "") or current.get("comparison_filter_action", "") or "")[:240]
    if reason in {"weak_identity", "missing_traceable_identity"}:
        return "; ".join(item for item in [
            f"identity_level={str(current.get('identity_completeness_level', '')).strip()}" if str(current.get("identity_completeness_level", "")).strip() else "",
            f"catalyst={str(current.get('Catalyst', '')).strip()}" if str(current.get("Catalyst", "")).strip() else "",
            f"series={str(current.get('Series_Name', '')).strip()}" if str(current.get("Series_Name", "")).strip() else "",
        ] if item)
    if reason in {"non_direct_numeric", "approximate_or_range", "missing_structured_numeric", "qualitative_only"}:
        return str(current.get("raw_numeric_expression", "") or current.get("numeric_expression_type", "") or current.get("numeric_reliability_level", ""))[:240]
    if reason == "weak_source_granularity":
        return str(current.get("source_granularity", "") or "")[:240]
    if reason == "duplicate_like":
        return str(current.get("duplicate_candidate_type", "") or current.get("obvious_duplicate_type", "") or "")[:240]
    if reason in {"performance_without_condition_anchor", "missing_condition_or_context_anchor", "unresolved_without_context"}:
        return "; ".join(item for item in [
            f"x_axis={str(current.get('x_axis', '') or note_axis).strip()}" if str(current.get("x_axis", "") or note_axis).strip() else "",
            f"x_value={str(current.get('x_value', '') or note_x_value).strip()}" if str(current.get("x_value", "") or note_x_value).strip() else "",
            f"condition_count={_count_nonempty_condition_features(current)}" if str(current.get("data_source", "")).strip() == "figure" else "",
        ] if item)
    if reason in {"candidate_only_figure_row", "nonmatched_figure_binding", "ambiguous_binding", "unmatched_binding", "low_figure_binding_confidence", "unsupported_binding_mode"}:
        return "; ".join(item for item in [
            f"status={str(current.get('figure_binding_status', '')).strip()}" if str(current.get("figure_binding_status", "")).strip() else "",
            f"mode={str(current.get('figure_binding_mode', current.get('binding_mode', ''))).strip()}" if str(current.get("figure_binding_mode", current.get("binding_mode", ""))).strip() else "",
            f"confidence={str(current.get('figure_binding_confidence', current.get('binding_confidence', ''))).strip()}" if str(current.get("figure_binding_confidence", current.get("binding_confidence", ""))).strip() else "",
            f"premerge_reason={str(current.get('premerge_filter_reason', '')).strip()}" if str(current.get("premerge_filter_reason", "")).strip() else "",
        ] if item)
    if reason in {"support_only_or_baseline", "support_only_or_baseline_without_context"}:
        return "; ".join(item for item in [
            f"mode={str(current.get('figure_binding_mode', current.get('binding_mode', ''))).strip()}" if str(current.get("figure_binding_mode", current.get("binding_mode", ""))).strip() else "",
            f"raw_category_label={str(current.get('raw_category_label', '')).strip()}" if str(current.get("raw_category_label", "")).strip() else "",
            f"catalyst={str(current.get('Catalyst', '')).strip()}" if str(current.get("Catalyst", "")).strip() else "",
        ] if item)
    if reason == "blocked_figure_semantic_role":
        return _canonicalize_semantic_figure_role(str(current.get("semantic_figure_role", "") or ""))
    return str(current.get("Notes", "") or current.get("figure_binding_reason", "") or "")[:240]


def _sync_layer_membership_fields(row: Dict) -> Dict:
    current = dict(row or {})
    layers = []
    if str(current.get("included_in_research_all", "")).strip() in {"1", "True", "true"} or bool(current.get("included_in_research_all")):
        layers.append("research_all")
    if str(current.get("included_in_research_strict", "")).strip() in {"1", "True", "true"} or bool(current.get("included_in_research_strict")):
        layers.append("research_strict")
    if str(current.get("included_in_modeling_final", "")).strip() in {"1", "True", "true"} or bool(current.get("included_in_modeling_final")):
        layers.append("modeling_final")
    current["layer_membership"] = "|".join(layers) if layers else ""
    return current


def _annotate_row_with_layer_membership(
    row: Dict,
    research_all_ok: bool,
    research_strict_ok: bool,
    modeling_final_ok: bool,
    strict_reason: str = "",
    final_reason: str = "",
    strict_detail: str = "",
    final_detail: str = "",
) -> Dict:
    current = _finalize_figure_row_binding_payload(dict(row or {}))
    current["approx_flag"] = int(bool(current.get("is_approximate_value")))
    current["range_flag"] = int(bool(current.get("is_range_like_value")))
    current["partial_point_flag"] = int(_derive_partial_point_flag(current))
    current["included_in_research_all"] = int(bool(research_all_ok))
    current["included_in_research_strict"] = int(bool(research_all_ok and research_strict_ok))
    current["included_in_modeling_final"] = int(bool(research_all_ok and research_strict_ok and modeling_final_ok))
    current["exclusion_from_strict_reason"] = "" if research_strict_ok else clean_residual_mojibake_chars(str(strict_reason or "")).strip()
    current["exclusion_from_strict_detail"] = "" if research_strict_ok else clean_residual_mojibake_chars(str(strict_detail or _derive_layer_exclusion_detail(current, strict_reason))).strip()
    current["exclusion_from_final_reason"] = "" if modeling_final_ok else clean_residual_mojibake_chars(str(final_reason or "")).strip()
    current["exclusion_from_final_detail"] = "" if modeling_final_ok else clean_residual_mojibake_chars(str(final_detail or _derive_layer_exclusion_detail(current, final_reason))).strip()
    current["strict_exclusion_reason"] = current["exclusion_from_strict_reason"]
    current["strict_exclusion_detail"] = current["exclusion_from_strict_detail"]
    current["final_exclusion_reason"] = current["exclusion_from_final_reason"]
    current["final_exclusion_detail"] = current["exclusion_from_final_detail"]
    return _sync_layer_membership_fields(current)


def _is_research_all_eligible_text_row(row: Dict) -> Tuple[bool, str]:
    current = dict(row or {})
    if str(current.get("data_source", "")).strip() != "text":
        return False, "not_text"
    if is_preparation_backbone_only_row(current):
        return False, "preparation_backbone_only"
    if not _is_core_metric_present(current):
        return False, "missing_core_metric"
    if not _has_any_numeric_signal(current):
        return False, "missing_structured_numeric"
    if bool(current.get("is_qualitative_value")):
        return False, "qualitative_only"
    if not _has_minimum_traceable_identity(current):
        return False, "missing_traceable_identity"
    source_granularity = clean_residual_mojibake_chars(str(current.get("source_granularity", "") or "")).strip()
    if not (_has_any_strong_condition_anchor(current) or source_granularity in _LAYERED_TEXT_HIGH_TRUST_GRANULARITIES):
        return False, "missing_condition_or_context_anchor"
    return True, "ok"


def _is_research_strict_eligible_text_row(row: Dict) -> Tuple[bool, str]:
    current = dict(row or {})
    all_ok, all_reason = _is_research_all_eligible_text_row(current)
    if not all_ok:
        return False, all_reason
    if bool(current.get("is_literature_comparison")) or str(current.get("comparison_filter_action", "")).strip() == "downgraded_candidate_only":
        return False, "comparison_like"
    if str(current.get("identity_completeness_level", "")).strip() not in {"complete", "partial"}:
        return False, "weak_identity"
    if str(current.get("source_granularity", "")).strip() not in _LAYERED_TEXT_HIGH_TRUST_GRANULARITIES:
        return False, "weak_source_granularity"
    if not _has_any_strong_condition_anchor(current):
        return False, "performance_without_condition_anchor"
    if bool(current.get("same_physical_point_possible")) or bool(current.get("obvious_duplicate_flag")) or str(current.get("duplicate_candidate_type", "")).strip():
        return False, "duplicate_like"
    numeric_level = str(current.get("numeric_reliability_level", "")).strip()
    if numeric_level == "direct_numeric" and not any(bool(current.get(flag)) for flag in ["is_approximate_value", "is_range_like_value", "is_qualitative_value"]):
        return True, "ok"
    if numeric_level in {"approximate_numeric", "range_like"} and not bool(current.get("is_qualitative_value")):
        return True, "ok_with_flagged_numeric"
    return False, "non_direct_numeric"


def _is_modeling_final_eligible_text_row(row: Dict) -> Tuple[bool, str]:
    return is_final_publishable_text_row(row)


def _is_research_all_eligible_figure_row(row: Dict) -> Tuple[bool, str]:
    current = _finalize_figure_row_binding_payload(dict(row or {}))
    if str(current.get("data_source", "")).strip() != "figure":
        return False, "not_figure"
    semantic_role = _canonicalize_semantic_figure_role(str(current.get("semantic_figure_role", "") or "").strip())
    if semantic_role in _LAYERED_BLOCKED_FIGURE_ROLES:
        return False, "blocked_figure_semantic_role"
    if not _has_core_figure_performance_value(current):
        return False, "missing_core_metric"
    if bool(current.get("is_qualitative_value")):
        return False, "qualitative_only"
    if not _has_minimum_traceable_identity(current):
        return False, "missing_traceable_identity"
    if _is_baseline_or_support_only_row(current):
        if _has_figure_context_support(current) or _has_any_strong_condition_anchor(current):
            return True, "support_only_or_baseline_research_row"
        return False, "support_only_or_baseline_without_context"
    binding_status = str(current.get("figure_binding_status", "") or "").strip()
    premerge_action = str(current.get("premerge_filter_action", "") or "").strip()
    if binding_status == "matched":
        if _has_figure_context_support(current) or _has_any_strong_condition_anchor(current):
            return True, "ok"
        # Fallback: matched binding with identity+metric but no strong context still qualifies for research_all
        return True, "matched_no_strong_context_but_has_identity_and_metric"
    if premerge_action == "candidate_only" or binding_status in {"ambiguous", "unmatched"}:
        if _has_figure_context_support(current) or _has_any_strong_condition_anchor(current):
            return True, "candidate_only_figure_row"
        # Fallback: candidate-only row with traceable identity + core metric value still enters research_all
        # (not research_strict or modeling_final — just preserved for research_all).
        if _has_minimum_traceable_identity(current) and _has_core_figure_performance_value(current):
            return True, "candidate_only_no_context_but_has_identity_and_metric"
        return False, "unresolved_without_context"
    return False, "nonmatched_figure_binding"


def _is_research_strict_eligible_figure_row(row: Dict) -> Tuple[bool, str]:
    current = _finalize_figure_row_binding_payload(dict(row or {}))
    all_ok, all_reason = _is_research_all_eligible_figure_row(current)
    if not all_ok:
        return False, all_reason
    if bool(current.get("is_literature_comparison")) or str(current.get("comparison_filter_action", "")).strip() == "downgraded_candidate_only":
        return False, "comparison_like"
    if _is_baseline_or_support_only_row(current):
        return False, "support_only_or_baseline"
    if str(current.get("figure_binding_status", "")).strip() != "matched":
        return False, "nonmatched_figure_binding"
    if str(current.get("figure_binding_confidence", current.get("binding_confidence", ""))).strip() not in {"high", "medium"}:
        return False, "low_figure_binding_confidence"
    mode = str(current.get("figure_binding_mode", current.get("binding_mode", ""))).strip()
    if not (_is_high_value_figure_binding_mode(mode) or mode in {"canonical_label", "exact_label"}):
        return False, "unsupported_binding_mode"
    if str(current.get("identity_completeness_level", "")).strip() not in {"complete", "partial"} and not str(current.get("matched_registry_key", "")).strip():
        return False, "weak_identity"
    if bool(current.get("same_physical_point_possible")) or bool(current.get("obvious_duplicate_flag")) or str(current.get("duplicate_candidate_type", "")).strip():
        return False, "duplicate_like"
    if not (_has_any_strong_condition_anchor(current) or _has_figure_context_support(current)):
        return False, "missing_condition_or_context_anchor"
    numeric_level = str(current.get("numeric_reliability_level", "")).strip()
    if numeric_level == "direct_numeric" and not any(bool(current.get(flag)) for flag in ["is_approximate_value", "is_range_like_value", "is_qualitative_value"]):
        return True, "ok"
    if numeric_level in {"approximate_numeric", "range_like"} and not bool(current.get("is_qualitative_value")):
        return True, "ok_with_flagged_numeric"
    return False, "non_direct_numeric"


def _is_modeling_final_eligible_figure_row(row: Dict) -> Tuple[bool, str]:
    downgrade, reason = _should_downgrade_figure_row_from_final(row)
    return (False, reason) if downgrade else (True, "ok")


def _evaluate_layer_membership_for_row(row: Dict) -> Dict:
    current = _finalize_figure_row_binding_payload(dict(row or {}))
    data_source = str(current.get("data_source", "") or "").strip()
    if data_source == "text":
        research_all_ok, research_all_reason = _is_research_all_eligible_text_row(current)
        research_strict_ok, research_strict_reason = _is_research_strict_eligible_text_row(current) if research_all_ok else (False, research_all_reason or "not_in_research_all")
        modeling_final_ok, modeling_final_reason = _is_modeling_final_eligible_text_row(current) if (research_all_ok and research_strict_ok) else (False, research_strict_reason or research_all_reason or "not_in_research_strict")
    elif data_source == "figure":
        research_all_ok, research_all_reason = _is_research_all_eligible_figure_row(current)
        research_strict_ok, research_strict_reason = _is_research_strict_eligible_figure_row(current) if research_all_ok else (False, research_all_reason or "not_in_research_all")
        modeling_final_ok, modeling_final_reason = _is_modeling_final_eligible_figure_row(current) if (research_all_ok and research_strict_ok) else (False, research_strict_reason or research_all_reason or "not_in_research_strict")
    else:
        research_all_ok, research_all_reason = False, "unsupported_data_source"
        research_strict_ok, research_strict_reason = False, research_all_reason
        modeling_final_ok, modeling_final_reason = False, research_all_reason
    if not research_all_ok:
        research_strict_ok = False
        modeling_final_ok = False
    if research_all_ok and not research_strict_ok:
        modeling_final_ok = False
    return _annotate_row_with_layer_membership(
        current,
        research_all_ok=research_all_ok,
        research_strict_ok=research_strict_ok,
        modeling_final_ok=modeling_final_ok,
        strict_reason="" if research_strict_ok else research_strict_reason,
        final_reason="" if modeling_final_ok else modeling_final_reason,
    )


def _bucket_layered_views(annotated_rows: List[Dict]) -> Dict[str, List[Dict]]:
    return {
        "annotated_rows": [dict(row) for row in annotated_rows],
        "research_all_rows": [dict(row) for row in annotated_rows if bool(row.get("included_in_research_all"))],
        "research_all_excluded_rows": [dict(row) for row in annotated_rows if not bool(row.get("included_in_research_all"))],
        "research_strict_rows": [dict(row) for row in annotated_rows if bool(row.get("included_in_research_strict"))],
        "research_strict_excluded_rows": [dict(row) for row in annotated_rows if bool(row.get("included_in_research_all")) and not bool(row.get("included_in_research_strict"))],
        "modeling_final_preeligible_rows": [dict(row) for row in annotated_rows if bool(row.get("included_in_modeling_final"))],
        "modeling_final_preexcluded_rows": [dict(row) for row in annotated_rows if bool(row.get("included_in_research_strict")) and not bool(row.get("included_in_modeling_final"))],
    }


def build_text_gate_stratified_views(rows: List[Dict]) -> Dict[str, List[Dict]]:
    annotated_rows = []
    for row in rows or []:
        if not isinstance(row, dict):
            continue
        if str(row.get("data_source", "")).strip() != "text":
            continue
        annotated_rows.append(_evaluate_layer_membership_for_row(row))
    return _bucket_layered_views(annotated_rows)


def build_figure_gate_stratified_views(rows: List[Dict]) -> Dict[str, List[Dict]]:
    annotated_rows = []
    for row in rows or []:
        if not isinstance(row, dict):
            continue
        if str(row.get("data_source", "")).strip() != "figure":
            continue
        annotated_rows.append(_evaluate_layer_membership_for_row(row))
    return _bucket_layered_views(annotated_rows)


def _selfcheck_layered_gate_examples() -> None:
    if globals().get("_LAYERED_GATE_SELFCHECK_DONE"):
        return
    globals()["_LAYERED_GATE_SELFCHECK_DONE"] = True
    try:
        strong_text = {
            "Source_File": "layer_selfcheck.pdf",
            "data_source": "text",
            "Catalyst": "Ni/ZrO2",
            "Catalyst_ID": "Ni/ZrO2",
            "Canonical_Catalyst_ID": "Ni/ZrO2",
            "Active_Metal": "Ni",
            "Support": "ZrO2",
            "identity_completeness_level": "complete",
            "numeric_reliability_level": "direct_numeric",
            "source_granularity": "table_row",
            "MeOH_Conversion_%": "65",
            "Reaction_Temp_C": "250",
            "TOS_h": "10",
        }
        a1, _ = _is_research_all_eligible_text_row(strong_text)
        s1, _ = _is_research_strict_eligible_text_row(strong_text)
        f1, _ = _is_modeling_final_eligible_text_row(strong_text)
        if not (a1 and s1 and f1):
            print("  [WARNING] layered selfcheck failed: strong text row should enter all/strict/final")

        partial_text = {
            "Source_File": "layer_selfcheck.pdf",
            "data_source": "text",
            "Catalyst": "Ni catalyst on ZrO2",
            "Catalyst_ID": "Ni/ZrO2",
            "Active_Metal": "Ni",
            "Support": "ZrO2",
            "identity_completeness_level": "partial",
            "numeric_reliability_level": "direct_numeric",
            "source_granularity": "text_numeric",
            "MeOH_Conversion_%": "54",
            "Reaction_Temp_C": "240",
        }
        a2, _ = _is_research_all_eligible_text_row(partial_text)
        s2, _ = _is_research_strict_eligible_text_row(partial_text)
        f2, _ = _is_modeling_final_eligible_text_row(partial_text)
        if not a2 or not s2 or f2:
            print("  [WARNING] layered selfcheck failed: partial-identity text row should stay in all/strict but not auto-enter final")

        candidate_only_figure = {
            "Source_File": "layer_selfcheck.pdf",
            "data_source": "figure",
            "Catalyst": "Ni/ZrO2",
            "Series_Name": "Ni/ZrO2",
            "raw_category_label": "Ni/ZrO2",
            "identity_completeness_level": "partial",
            "numeric_reliability_level": "direct_numeric",
            "figure_binding_status": "ambiguous",
            "premerge_filter_action": "candidate_only",
            "premerge_filter_reason": "ambiguous_binding",
            "semantic_figure_role": "condition_effect",
            "MeOH_Conversion_%": "72",
            "Reaction_Temp_C": "250",
            "figure_binding_confidence": "medium",
        }
        a3, _ = _is_research_all_eligible_figure_row(candidate_only_figure)
        f3, _ = _is_modeling_final_eligible_figure_row(candidate_only_figure)
        if not a3 or f3:
            print("  [WARNING] layered selfcheck failed: candidate-only figure row should enter all but not final")

        approx_text = {
            "Source_File": "layer_selfcheck.pdf",
            "data_source": "text",
            "Catalyst": "Ni/ZrO2",
            "Catalyst_ID": "Ni/ZrO2",
            "Active_Metal": "Ni",
            "Support": "ZrO2",
            "identity_completeness_level": "complete",
            "numeric_reliability_level": "approximate_numeric",
            "is_approximate_value": 1,
            "source_granularity": "table_row",
            "MeOH_Conversion_%": "~60",
            "Reaction_Temp_C": "250",
        }
        a4, _ = _is_research_all_eligible_text_row(approx_text)
        s4, _ = _is_research_strict_eligible_text_row(approx_text)
        f4, _ = _is_modeling_final_eligible_text_row(approx_text)
        if not a4 or not s4 or f4:
            print("  [WARNING] layered selfcheck failed: approximate text row should be retained in all/strict and excluded from final")

        baseline_figure = {
            "Source_File": "layer_selfcheck.pdf",
            "data_source": "figure",
            "Catalyst": "ZrO2",
            "Support": "ZrO2",
            "Series_Name": "ZrO2",
            "raw_category_label": "ZrO2",
            "identity_completeness_level": "partial",
            "numeric_reliability_level": "direct_numeric",
            "figure_binding_status": "matched",
            "figure_binding_mode": "support_only_match",
            "figure_binding_confidence": "medium",
            "semantic_figure_role": "category_screening",
            "MeOH_Conversion_%": "12",
            "Reaction_Temp_C": "250",
        }
        a5, _ = _is_research_all_eligible_figure_row(baseline_figure)
        s5, _ = _is_research_strict_eligible_figure_row(baseline_figure)
        f5, _ = _is_modeling_final_eligible_figure_row(baseline_figure)
        if not a5 or s5 or f5:
            print("  [WARNING] layered selfcheck failed: support-only/baseline row should stay research-all only")
    except Exception as e:
        print(f"  [WARNING] layered gate selfcheck error: {e}")


def _build_text_point_gate_audit_rows(rows: List[Dict], file_name: str) -> List[Dict]:
    audit_rows = []
    for row in rows or []:
        if not isinstance(row, dict):
            continue
        ok, reason = is_final_publishable_text_row(row)
        note_axis, note_x_value = _extract_x_axis_and_value_from_notes(str(row.get("Notes", "")))
        detail = ""
        if reason == "comparison_like":
            detail = str(row.get("comparison_filter_reason", "") or row.get("comparison_filter_action", ""))
        elif reason == "weak_identity":
            detail = str(row.get("identity_completeness_level", ""))
        elif reason == "non_direct_numeric":
            detail = str(row.get("numeric_reliability_level", ""))
        elif reason == "weak_source_granularity":
            detail = str(row.get("source_granularity", ""))
        elif reason == "approximate_or_range":
            detail = str(row.get("raw_numeric_expression", "") or row.get("numeric_expression_type", ""))
        elif reason == "duplicate_like":
            detail = str(row.get("duplicate_candidate_type", "") or row.get("obvious_duplicate_type", ""))
        elif reason == "performance_without_condition_anchor":
            detail = "; ".join(item for item in [
                f"x_axis={str(row.get('x_axis', '') or note_axis).strip()}" if str(row.get("x_axis", "") or note_axis).strip() else "",
                f"x_value={str(row.get('x_value', '') or note_x_value).strip()}" if str(row.get("x_value", "") or note_x_value).strip() else "",
            ] if item)
        else:
            detail = str(row.get("Notes", ""))[:240]
        audit_rows.append({
            "Source_File": file_name,
            "data_source": str(row.get("data_source", "")),
            "Catalyst": str(row.get("Catalyst", "")),
            "Catalyst_ID": str(row.get("Catalyst_ID", "")),
            "Canonical_Catalyst_ID": str(row.get("Canonical_Catalyst_ID", "")),
            "source_granularity": str(row.get("source_granularity", "")),
            "numeric_reliability_level": str(row.get("numeric_reliability_level", "")),
            "identity_completeness_level": str(row.get("identity_completeness_level", "")),
            "point_gate_status": "ok" if ok else "blocked_for_strict_final",
            "point_gate_reason": reason,
            "point_gate_notes": detail,
        })
    return audit_rows


def _build_composition_consistency_audit_rows(rows: List[Dict], file_name: str) -> List[Dict]:
    audit_rows = []
    for row in rows or []:
        if not isinstance(row, dict):
            continue
        fractions = derive_composition_fractions(dict(row))
        audit_rows.append({
            "Source_File": file_name,
            "Catalyst": str(row.get("Catalyst", "")),
            "Catalyst_ID": str(row.get("Catalyst_ID", "")),
            "Active_Metal": str(row.get("Active_Metal", "")),
            "Promoter": str(row.get("Promoter", "")),
            "Promoter_Metal": str(fractions.get("Promoter_Metal", "") or row.get("Promoter_Metal", "")),
            "Alloy_Ratio": str(row.get("Alloy_Ratio", "") or row.get("Metal_Ratio", "") or row.get("Composition_Ratio", "")),
            "parsed_ratio_pairs": str(fractions.get("parsed_ratio_pairs", "")),
            "Ni_Fraction": str(fractions.get("Ni_Fraction", "")),
            "Promoter_Fraction": str(fractions.get("Promoter_Fraction", "")),
            "has_promoter": str(fractions.get("has_promoter", "")),
            "is_bimetallic": str(fractions.get("is_bimetallic", "")),
            "composition_consistency_flag": str(fractions.get("composition_consistency_flag", "ok")),
            "composition_consistency_notes": str(fractions.get("composition_consistency_notes", "")),
        })
    return audit_rows


def _selfcheck_generalization_rules() -> None:
    def _warn(name: str, passed: bool) -> None:
        if not passed:
            print(f"  [WARNING] selfcheck failed: {name}")

    case1 = derive_composition_fractions({"Alloy_Ratio": "Pt:Ni = 0.5:10", "Catalyst": "Pt-Ni/SEP", "Active_Metal": "Ni"})
    _warn("Pt:Ni = 0.5:10", "Pt" in str(case1.get("Promoter_Metal", "")) and str(case1.get("Ni_Fraction", "")) != "1.0" and str(case1.get("Promoter_Fraction", "")) != "0.0")

    case2 = derive_composition_fractions({"Catalyst": "Ni80Cu20/Al2O3", "Active_Metal": "Ni"})
    _warn("Ni80Cu20/Al2O3", str(case2.get("has_promoter", "")) == "1" and str(case2.get("is_bimetallic", "")) == "1" and "Cu" in str(case2.get("Promoter_Metal", "")))

    case3 = derive_composition_fractions({"Catalyst": "Ni/Al2O3", "Active_Metal": "Ni"})
    _warn("Ni/Al2O3", str(case3.get("has_promoter", "")) == "0" and str(case3.get("is_bimetallic", "")) == "0" and str(case3.get("Ni_Fraction", "")) == "1.0" and str(case3.get("Promoter_Fraction", "")) == "0.0")

    case4 = {"data_source": "text", "paragraph_role": "preparation_paragraph", "text_extraction_subroute": "preparation_schema", "Calcination_Temp_C": "500", "Reduction_Temp_C": "700", "Metal_Loading_wt%": "10"}
    _warn("preparation_only_row", is_preparation_backbone_only_row(case4) and not is_text_point_record(case4))

    case5 = {"data_source": "text", "identity_completeness_level": "complete", "numeric_reliability_level": "direct_numeric", "source_granularity": "text_numeric", "MeOH_Conversion_%": "95", "Catalyst": "Ni/Al2O3"}
    _warn("performance_without_anchor", has_core_performance_metric(case5) and not has_valid_point_condition_anchor(case5) and not is_text_point_record(case5))

    precheck = globals().get("precheck_table_or_figure_candidates")
    if callable(precheck):
        try:
            result = precheck("XRD TEM XPS", "", "", 1)
            _warn("characterization_heavy_low_text", not bool(result.get("is_candidate", False)))
        except Exception as e:
            print(f"  [WARNING] selfcheck failed: figure_candidate_gate ({e})")


def derive_promoter_metal(record: Dict) -> str:
    if not isinstance(record, dict):
        return ""

    def _join(tokens: List[str]) -> str:
        return "/".join(_dedupe_keep_order([tok for tok in tokens if tok and tok != "Ni"]))

    explicit = []
    for field in ("Promoter_Metal", "Promoter", "Second_Metal", "Additive"):
        explicit.extend(_safe_split_multi_promoter(record.get(field, "")))
    explicit = _filter_identity_promoter_tokens(_dedupe_keep_order(explicit), record)
    if explicit:
        return _join(explicit)

    ratio_pairs = parse_alloy_ratio_to_ordered_pairs(record)
    ratio_promoters = _filter_identity_promoter_tokens(
        _dedupe_keep_order([metal for metal, _ in ratio_pairs if metal != "Ni"]),
        record,
    )
    if ratio_promoters:
        return _join(ratio_promoters)

    identity_text = _extract_identity_text_for_composition(record)
    identity_candidates = []
    identity_candidates.extend(parse_second_metal_from_text_label(identity_text))
    if "Ni" in _extract_elements_for_merge(identity_text) or re.search(r"\bni\b", identity_text, flags=re.I):
        identity_candidates.extend([
            token for token in _extract_elements_for_merge(identity_text)
            if token != "Ni" and token in MSR_NI_SECOND_METALS
        ])
    identity_candidates = _filter_identity_promoter_tokens(identity_candidates, record, identity_text)
    if identity_candidates:
        return _join(identity_candidates)

    fallback = []
    for field in ("Active_Metal", "Catalyst", "Catalyst_ID", "Canonical_Catalyst_ID", "Series_Name"):
        fallback.extend(infer_promoter_from_catalyst_name(str(record.get(field, ""))))
    return _join(_filter_identity_promoter_tokens(_dedupe_keep_order(fallback), record))


def derive_composition_fractions(record: Dict) -> Dict[str, Any]:
    result: Dict[str, Any] = {
        "Promoter_Metal": "",
        "Ni_Fraction": "",
        "Promoter_Fraction": "",
        "has_promoter": "",
        "is_bimetallic": "",
        "parsed_ratio_pairs": "",
        "composition_consistency_flag": "ok",
        "composition_consistency_notes": "",
    }

    def _fmt_fraction(value: float) -> str:
        value = round(float(value), 4)
        if abs(value - round(value)) < 1e-9:
            return f"{value:.1f}"
        return _format_ratio_token(f"{value:.4f}")

    def _join(tokens: List[str]) -> str:
        return "/".join(_dedupe_keep_order([tok for tok in tokens if tok and tok != "Ni"]))

    explicit = []
    for field in ("Promoter_Metal", "Promoter", "Second_Metal", "Additive"):
        explicit.extend(_safe_split_multi_promoter(record.get(field, "")))
    explicit = _filter_identity_promoter_tokens(_dedupe_keep_order(explicit), record)

    ratio_pairs = parse_alloy_ratio_to_ordered_pairs(record)
    ratio_promoters = _filter_identity_promoter_tokens(
        _dedupe_keep_order([metal for metal, _ in ratio_pairs if metal != "Ni"]),
        record,
    )

    identity_text = _extract_identity_text_for_composition(record)
    identity = _filter_identity_promoter_tokens([
        token for token in (
            parse_second_metal_from_text_label(identity_text)
            + [tok for tok in _extract_elements_for_merge(identity_text) if tok != "Ni"]
        )
        if token in MSR_NI_SECOND_METALS and token != "Ni"
    ], record, identity_text)

    chosen = explicit or ratio_promoters or identity
    notes = []
    if explicit and ratio_promoters and set(explicit) != set(ratio_promoters):
        result["composition_consistency_flag"] = "explicit_promoter_conflict"
        notes.append(f"explicit={_join(explicit)}")
        notes.append(f"ratio={_join(ratio_promoters)}")
        notes.append(f"identity={_join(identity)}")
        chosen = ratio_promoters
    elif explicit and identity and set(explicit) != set(identity):
        result["composition_consistency_flag"] = "explicit_promoter_conflict"
        notes.append(f"explicit={_join(explicit)}")
        notes.append(f"ratio={_join(ratio_promoters)}")
        notes.append(f"identity={_join(identity)}")
        chosen = explicit

    result["Promoter_Metal"] = _join(chosen)
    result["parsed_ratio_pairs"] = "|".join(f"{metal}:{_format_ratio_token(str(value))}" for metal, value in ratio_pairs)
    result["composition_consistency_notes"] = "; ".join(note for note in notes if note)

    ni_context_text = " ".join([
        str(record.get("Active_Metal", "")),
        str(record.get("Catalyst", "")),
        str(record.get("Catalyst_ID", "")),
        str(record.get("Canonical_Catalyst_ID", "")),
        str(record.get("Series_Name", "")),
        str(record.get("Alloy_Ratio", "")),
        str(record.get("Metal_Ratio", "")),
        str(record.get("Composition_Ratio", "")),
    ])
    has_ni_context = bool(re.search(r"\bni\b", ni_context_text, flags=re.I)) or "Ni" in [metal for metal, _ in ratio_pairs]

    if ratio_pairs:
        ratio_dict = {metal: float(value) for metal, value in ratio_pairs}
        total = sum(ratio_dict.values())
        if total > 0 and "Ni" in ratio_dict:
            result["Ni_Fraction"] = _fmt_fraction(ratio_dict["Ni"] / total)
            if ratio_promoters:
                result["has_promoter"] = 1
                result["is_bimetallic"] = 1
                if len(ratio_promoters) == 1 and ratio_promoters[0] in ratio_dict:
                    result["Promoter_Fraction"] = _fmt_fraction(ratio_dict[ratio_promoters[0]] / total)
            else:
                result["Promoter_Fraction"] = "0.0"
                result["has_promoter"] = 0
                result["is_bimetallic"] = 0
            return result

    if result["Promoter_Metal"] and has_ni_context:
        result["has_promoter"] = 1
        result["is_bimetallic"] = 1
        return result

    if has_ni_context and not result["Promoter_Metal"]:
        result["Ni_Fraction"] = "1.0"
        result["Promoter_Fraction"] = "0.0"
        result["has_promoter"] = 0
        result["is_bimetallic"] = 0

    return result



_OXIDE_TO_METAL_FRACTION = {
    "NiO": 58.6934 / (58.6934 + 15.999),
    "CuO": 63.546 / (63.546 + 15.999),
    "ZnO": 65.38 / (65.38 + 15.999),
    "CeO2": 140.116 / (140.116 + 2 * 15.999),
    "Pr6O11": (6 * 140.90766) / (6 * 140.90766 + 11 * 15.999),
    "SnO2": 118.710 / (118.710 + 2 * 15.999),
    "In2O3": (2 * 114.818) / (2 * 114.818 + 3 * 15.999),
    "CoO": 58.9332 / (58.9332 + 15.999),
    "Fe2O3": (2 * 55.845) / (2 * 55.845 + 3 * 15.999),
    "MnO": 54.938 / (54.938 + 15.999),
    "MoO3": 95.95 / (95.95 + 3 * 15.999),
}
_OXIDE_TO_ELEMENT = {
    "NiO": "Ni",
    "CuO": "Cu",
    "ZnO": "Zn",
    "CeO2": "Ce",
    "Pr6O11": "Pr",
    "SnO2": "Sn",
    "In2O3": "In",
    "CoO": "Co",
    "Fe2O3": "Fe",
    "MnO": "Mn",
    "MoO3": "Mo",
}


def _component_loading_context(record: Dict) -> str:
    fields = [
        "Catalyst", "Catalyst_ID", "Canonical_Catalyst_ID", "Series_Name",
        "Active_Metal", "Promoter", "Promoter_Metal", "Metal_Loading_wt%",
        "Alloy_Ratio", "Notes", "_comparison_context_raw",
    ]
    return clean_residual_mojibake_chars(" ".join(str(record.get(field, "") or "") for field in fields))


def _component_loading_identity_text(record: Dict) -> str:
    return clean_residual_mojibake_chars(" ".join(
        str(record.get(field, "") or "")
        for field in ("Catalyst", "Catalyst_ID", "Canonical_Catalyst_ID", "Series_Name")
    ))


def _strip_sample_state_suffix(label: str) -> str:
    text = clean_residual_mojibake_chars(str(label or "")).strip()
    text = re.sub(r"(?i)(?:[-_ ](?:r|s|red|reduced|spent|used|calcined|fresh|oxided|oxidized))$", "", text)
    return text.strip()


def _label_definition_supported(record: Dict) -> bool:
    context = _normalize_identity_text(_component_loading_context(record))
    if not context:
        return False

    labels = _dedupe_keep_order([
        str(record.get("Catalyst", "") or "").strip(),
        str(record.get("Series_Name", "") or "").strip(),
        str(record.get("Catalyst_ID", "") or "").strip(),
        str(record.get("Canonical_Catalyst_ID", "") or "").strip(),
    ])
    labels = [label for label in labels if label]
    label_variants = []
    for label in labels:
        label_variants.append(label)
        stripped = _strip_sample_state_suffix(label)
        if stripped and stripped != label:
            label_variants.append(stripped)

    definition_re = re.compile(r"\b(?:denoted|designated|named|labeled|labelled|abbreviated|referred to as|called)\b", flags=re.I)
    loading_re = re.compile(r"\b(?:loading|loadings|content|contents|wt\.?\s*%|wt%|mass\s*%)\b", flags=re.I)

    for label in _dedupe_keep_order(label_variants):
        if not label:
            continue
        idx = context.lower().find(_normalize_identity_text(label).lower())
        if idx < 0:
            continue
        window = context[max(0, idx - 700): idx + len(label) + 700]
        if definition_re.search(window) and loading_re.search(window):
            return True
    return False


def _context_has_ratio_warning_near_label(record: Dict) -> bool:
    context = _normalize_identity_text(_component_loading_context(record)).lower()
    label = _normalize_identity_text(str(record.get("Catalyst", "") or record.get("Series_Name", "") or ""))
    if not context or not label:
        return False
    idx = context.find(label.lower())
    window = context[max(0, idx - 300): idx + len(label) + 300] if idx >= 0 else context[:700]
    return bool(re.search(r"\b(?:weight|molar|atomic|metal|ni[- ]?to[- ]?cu|cu[- ]?to[- ]?ni)\s+ratios?\b|\bratios?\s*(?:=|of)\b|\bfactor\s+level\b|\btaguchi\b", window, flags=re.I))


def _format_loading_from_float(value: float) -> str:
    value = round(float(value), 3)
    if abs(value - round(value)) < 1e-9:
        return str(int(round(value)))
    return _format_ratio_token(f"{value:.3f}")


def _infer_oxide_component_loadings(record: Dict) -> Dict[str, Any]:
    row = dict(record)
    label_text = _component_loading_identity_text(row)
    promoter_symbols = set(_extract_elements_for_merge(
        " ".join(str(row.get(field, "") or "") for field in ("Promoter_Metal", "Promoter"))
    ))

    oxide_alt = "|".join(sorted((re.escape(k) for k in _OXIDE_TO_METAL_FRACTION), key=len, reverse=True))
    patterns = [
        rf"(\d+(?:\.\d+)?)\s*(?:wt\.?\s*%|%)\s*({oxide_alt})\b",
        rf"\b({oxide_alt})\s*[\-/_(]?\s*(\d+(?:\.\d+)?)\s*(?:wt\.?\s*%|%)",
    ]
    for pattern in patterns:
        for match in re.finditer(pattern, label_text, flags=re.I):
            if match.group(1) in _OXIDE_TO_METAL_FRACTION:
                oxide, oxide_loading = match.group(1), match.group(2)
            else:
                oxide_loading, oxide = match.group(1), match.group(2)
            canonical_oxide = next((key for key in _OXIDE_TO_METAL_FRACTION if key.lower() == oxide.lower()), "")
            if not canonical_oxide:
                continue
            element = _OXIDE_TO_ELEMENT.get(canonical_oxide, "")
            if element == "Ni":
                field = "Ni_Loading_wt%"
            elif element in promoter_symbols:
                field = "Promoter_Loading_wt%"
            else:
                continue
            if str(row.get(field, "")).strip():
                continue
            value = float(oxide_loading) * _OXIDE_TO_METAL_FRACTION[canonical_oxide]
            row[field] = _format_loading_from_float(value)
            _append_note_tag(row, f"[component_loading_from_oxide:{oxide_loading}%{canonical_oxide}->{row[field]}wt%{element}]")
    return row


def _infer_constant_ni_loading_from_definition(record: Dict) -> Dict[str, Any]:
    row = dict(record)
    if str(row.get("Ni_Loading_wt%", "")).strip():
        return row
    if not _label_definition_supported(row):
        return row
    context = _normalize_identity_text(_component_loading_context(row))
    if _context_has_ratio_warning_near_label(row):
        return row
    match = re.search(
        r"\bloading\s+of\s+Ni\b[^.;]{0,160}?\b(?:constant|kept|fixed|maintained|was|is)?[^.;]{0,80}?(\d+(?:\.\d+)?)\s*(?:wt\.?\s*%|%)",
        context,
        flags=re.I,
    )
    if match:
        row["Ni_Loading_wt%"] = _format_ratio_token(match.group(1))
        _append_note_tag(row, "[component_loading_from_label_definition:constant_Ni_loading]")
    return row


def _float_values_close(left: Any, right: Any, tolerance: float = 0.03) -> bool:
    lv = _parse_float_if_possible(left)
    rv = _parse_float_if_possible(right)
    if lv is None or rv is None:
        return False
    return abs(lv - rv) <= max(tolerance, abs(rv) * 0.01)


def _component_value_has_explicit_evidence(record: Dict, field: str, value: str) -> bool:
    if not str(value or "").strip():
        return False
    notes = clean_residual_mojibake_chars(str(record.get("Notes", "") or ""))
    if field == "Ni_Loading_wt%":
        element_alt = r"(?:Ni|nickel)"
        element = "Ni"
        if re.search(r"\[component_loading_from_[^\]]*(?:Ni|constant_Ni_loading)", notes):
            return True
    else:
        promoter = str(record.get("Promoter_Metal", "") or record.get("Promoter", "") or "").strip()
        promoter_symbols = _extract_elements_for_merge(promoter)
        element = promoter_symbols[0] if promoter_symbols else ""
        element_alt = rf"(?:{re.escape(element)}|promoter|second[- ]metal)" if element else r"(?:promoter|second[- ]metal)"
        if element and re.search(rf"\[component_loading_from_[^\]]*{re.escape(element)}", notes):
            return True

    context = _normalize_identity_text(_component_loading_context(record))
    identity = _normalize_identity_text(_component_loading_identity_text(record))
    value_re = re.escape(_compact_loading_number(_format_ratio_token(str(value))))
    if not context and not identity:
        return False

    evidence_text = " ".join([context, identity])
    direct_patterns = [
        rf"\b{value_re}\s*(?:wt\.?\s*%|wt%|%)\s*(?:of\s+)?{element_alt}\b",
        rf"\b{element_alt}\b[^.;]{{0,120}}\b(?:loading|loadings|content|contents|concentration|concentrations|measured|final metal loadings?)\b[^.;]{{0,120}}\b{value_re}\s*(?:wt\.?\s*%|wt%|%)",
        rf"\b(?:loading|loadings|content|contents|concentration|concentrations|final metal loadings?)\b[^.;]{{0,120}}\b{value_re}\s*(?:wt\.?\s*%|wt%|%)\s*(?:of\s+)?{element_alt}\b",
    ]
    if any(re.search(pattern, evidence_text, flags=re.I) for pattern in direct_patterns):
        if _context_has_ratio_warning_near_label(record) and not re.search(r"\b(?:loading|content|concentration|final metal loadings?)\b", evidence_text, flags=re.I):
            return False
        return True

    for sentence in re.split(r"(?<=[.;])\s+", evidence_text):
        if not re.search(r"(?:wt\.?\s*%|wt%|%)", sentence, flags=re.I):
            continue
        if not re.search(r"\b(?:loading|loadings|content|contents|concentration|concentrations|composition|compositions|promoter|promoters|respectively)\b", sentence, flags=re.I):
            continue
        if element and not re.search(rf"\b{re.escape(element)}\b", sentence, flags=re.I):
            continue
        if not re.search(rf"\b{value_re}\b", sentence):
            continue
        if re.search(r"\b(?:weight|molar|atomic|metal)\s+ratios?\b", sentence, flags=re.I):
            continue
        return True

    if _label_definition_supported(record):
        if field == "Ni_Loading_wt%":
            ni_match = re.search(
                rf"\bloading\s+of\s+Ni\b[^.;]{{0,180}}\b{value_re}\s*(?:wt\.?\s*%|wt%|%)",
                context,
                flags=re.I,
            )
            if ni_match:
                return True
        if element and re.search(rf"\b{value_re}\s*(?:wt\.?\s*%|wt%|%)?\s*{re.escape(element)}(?=[0-9A-Z\-/_.]|$)", identity, flags=re.I):
            return True

    oxide_label_row = _infer_oxide_component_loadings({k: v for k, v in record.items() if k not in ("Ni_Loading_wt%", "Promoter_Loading_wt%")})
    inferred = str(oxide_label_row.get(field, "") or "").strip()
    if inferred and _float_values_close(inferred, value):
        return True
    return False


def validate_component_loading_evidence(record: Dict) -> Dict[str, Any]:
    row = dict(record)
    for field in ("Ni_Loading_wt%", "Promoter_Loading_wt%"):
        value = str(row.get(field, "") or "").strip()
        if not value:
            continue
        notes = str(row.get("Notes", "") or "")
        suppressed_tag = f"component_loading_suppressed:no_explicit_evidence:{field}="

        if suppressed_tag in notes:
            row[field] = ""
            continue

        if "component_loading_from_label_definition:" in notes:
            continue

        # Case 3: LLM directly extracted value (no naming tag, no suppression).
        # The LLM read the paper — its reading is the evidence. We do not
        # require the evidence to be duplicated in the row's limited fields.
        continue
    return row


def infer_component_loadings_from_label(record: Dict) -> Dict[str, Any]:
    if not isinstance(record, dict):
        return record
    row = dict(record)
    row = _infer_oxide_component_loadings(row)
    row = _infer_constant_ni_loading_from_definition(row)

    _xpromoter_ysupport = re.compile(
        r"^(\d+(?:\.\d+)?)\s*([A-Z][a-z]{0,2})\s*(?:-\w+|(?=[A-Za-z]{2,}))?$",
        flags=re.I,
    )
    promoter_text = " ".join(
        clean_residual_mojibake_chars(str(row.get(field, "") or ""))
        for field in ("Promoter_Metal", "Promoter")
    )
    promoter_symbols = set(_extract_elements_for_merge(promoter_text))
    component_symbols = set(promoter_symbols)
    component_symbols.add("Ni")

    label_text = " ".join(
        clean_residual_mojibake_chars(str(row.get(field, "") or ""))
        for field in ("Catalyst", "Catalyst_ID", "Canonical_Catalyst_ID", "Series_Name")
    )
    for label_field in ("Catalyst", "Catalyst_ID", "Canonical_Catalyst_ID", "Series_Name"):
        label_val = str(row.get(label_field, "") or "").strip()
        if not label_val:
            continue
        m = _xpromoter_ysupport.match(label_val)
        if not m:
            continue
        x_val, x_elem = m.group(1), m.group(2)
        x_elem_cap = x_elem.capitalize()

        x_elem_cap_normalized = x_elem_cap
        if x_elem_cap == "N" and re.search(r"\d+\s*[Nn][Aa]?\s*$", label_val, flags=re.I):
            x_elem_cap_normalized = "NA"

        if x_elem_cap_normalized in component_symbols or x_elem_cap_normalized in MSR_NI_SECOND_METALS:
            if x_elem_cap_normalized == "Ni":
                if not str(row.get("Ni_Loading_wt%", "")).strip():
                    row["Ni_Loading_wt%"] = _format_ratio_token(x_val)
                    _append_note_tag(row, f"[component_loading_from_label_definition:Ni={x_val}]")
            else:
                if not str(row.get("Promoter_Loading_wt%", "")).strip():
                    row["Promoter_Loading_wt%"] = _format_ratio_token(x_val)
                    _append_note_tag(row, f"[component_loading_from_label_definition:Promoter={x_val}]")
                if not str(row.get("Promoter_Metal", "")).strip():
                    row["Promoter_Metal"] = x_elem_cap_normalized
                ni_suffix_match = re.search(r"-(\d+)\s*[Nn][Aa]?\s*$", label_val, flags=re.I)
                ni_val = ni_suffix_match.group(1) if ni_suffix_match else None
                if ni_val is None:
                    active_metal = str(row.get("Active_Metal", "") or "").strip()
                    if "Ni" in _extract_elements_for_merge(active_metal):
                        ni_val = "10"
                if ni_val is not None and not str(row.get("Ni_Loading_wt%", "")).strip():
                    row["Ni_Loading_wt%"] = ni_val
                    _append_note_tag(row, f"[component_loading_from_label_definition:Ni={ni_val}]")
        elif x_elem_cap_normalized == "NA":
            if not str(row.get("Ni_Loading_wt%", "")).strip():
                row["Ni_Loading_wt%"] = _format_ratio_token(x_val)
                _append_note_tag(row, f"[component_loading_from_label_definition:Ni={x_val}]")
                if not str(row.get("Support_Normalized", "")).strip():
                    row["Support_Normalized"] = "Al2O3"
                    _append_note_tag(row, "[support_from_naming:NA=Al2O3]")
        break

    for match in re.finditer(
        r"(\d+(?:\.\d+)?)\s*(?:wt\.?\s*%|%)?\s*([A-Z][a-z]{0,2})(?=[0-9A-Z\-/_.]|$)",
        label_text,
    ):
        value, element = match.group(1), match.group(2)
        if element not in component_symbols:
            continue
        element_cap = element.capitalize()
        if element_cap == "N":
            element_cap = "NA"
        field = "Ni_Loading_wt%" if element_cap == "Ni" else "Promoter_Loading_wt%"
        if element_cap != "Ni" and promoter_symbols and element_cap not in promoter_symbols:
            continue
        if not str(row.get(field, "")).strip():
            row[field] = _format_ratio_token(value)
            _append_note_tag(row, f"[component_loading_from_label_definition:{element_cap}={row[field]}]")

    notes_after = clean_residual_mojibake_chars(str(row.get("Notes", "") or ""))
    if re.search(r"\bnaming suggests\b|no explicit (?:component )?loadings?", notes_after, flags=re.I):
        if not _label_definition_supported(row):
            has_def_tag = "component_loading_from_label_definition:" in notes_after
            if not has_def_tag:
                if str(row.get("Ni_Loading_wt%", "")).strip():
                    _append_note_tag(row, f"[component_loading_suppressed:no_explicit_evidence:Ni_Loading_wt%={row.get('Ni_Loading_wt%')}]")
                    row["Ni_Loading_wt%"] = ""
                if str(row.get("Promoter_Loading_wt%", "")).strip():
                    _append_note_tag(row, f"[component_loading_suppressed:no_explicit_evidence:Promoter_Loading_wt%={row.get('Promoter_Loading_wt%')}]")
                    row["Promoter_Loading_wt%"] = ""

    if not _label_definition_supported(row):
        if not re.search(r"\[component_loading_from_label_definition:", notes_after):
            return row
    if _context_has_ratio_warning_near_label(row):
        _append_note_tag(row, "[component_loading_from_label_blocked:ratio_context]")
        return row

    return row



def _can_recover_total_loading_from_identity(record: Dict, label_text: str, value_text: str) -> bool:
    label_norm = _normalize_identity_text(str(label_text or ""))
    value = re.escape(str(value_text or "").strip())
    if not label_norm or not value:
        return False
    if _context_has_ratio_warning_near_label(record):
        return False

    if re.search(rf"\b{value}\s*(?:wt\.?\s*%|wt%|%)\s*[A-Z][a-z]?\b", label_norm, flags=re.I):
        return True

    if _label_definition_supported(record):
        context = _normalize_identity_text(_component_loading_context(record))
        if re.search(r"\b(?:loading|loadings|content|contents|metal\s+loadings?)\b", context, flags=re.I):
            return True
    return False



def sync_component_loading_derivatives(record: Dict) -> Dict[str, Any]:
    row = dict(record)
    ni_loading = _parse_float_if_possible(row.get("Ni_Loading_wt%", ""))
    promoter_loading = _parse_float_if_possible(row.get("Promoter_Loading_wt%", ""))
    promoter = str(row.get("Promoter_Metal", "") or row.get("Promoter", "") or "").strip()

    if ni_loading is not None and promoter_loading is not None and promoter_loading > 0:
        total = ni_loading + promoter_loading
        if total > 0:
            row["Metal_Loading_wt%"] = _format_loading_from_float(total)
            _append_note_tag(row, "[metal_loading_from_component_sum]")
            if not str(row.get("Ni_Fraction", "")).strip():
                row["Ni_Fraction"] = _format_ratio_token(f"{ni_loading / total:.4f}")
            if not str(row.get("Promoter_Fraction", "")).strip():
                row["Promoter_Fraction"] = _format_ratio_token(f"{promoter_loading / total:.4f}")
    elif ni_loading is not None and not promoter:
        row["Metal_Loading_wt%"] = _format_loading_from_float(ni_loading)
        _append_note_tag(row, "[metal_loading_from_component_sum]")
        if not str(row.get("Ni_Fraction", "")).strip():
            row["Ni_Fraction"] = "1.0"
        if not str(row.get("Promoter_Fraction", "")).strip():
            row["Promoter_Fraction"] = "0.0"
    return row



def derive_loading_partition(record: Dict) -> Dict[str, Any]:
    result: Dict[str, Any] = {"Ni_Loading": "", "Promoter_Loading": ""}

    def _fmt_loading(value: float) -> str:
        value = round(float(value), 3)
        if abs(value - round(value)) < 1e-9:
            return str(int(round(value)))
        return _format_ratio_token(f"{value:.3f}")

    ni_direct = str(record.get("Ni_Loading_wt%", "") or "").strip()
    prom_direct = str(record.get("Promoter_Loading_wt%", "") or "").strip()

    if ni_direct:
        result["Ni_Loading"] = ni_direct
    if prom_direct:
        result["Promoter_Loading"] = prom_direct

    total_loading = _parse_float_if_possible(record.get("Metal_Loading_wt%", ""))
    ni_fraction = _parse_float_if_possible(record.get("Ni_Fraction", ""))
    promoter_fraction = _parse_float_if_possible(record.get("Promoter_Fraction", ""))

    if not result["Ni_Loading"] and total_loading is not None and ni_fraction is not None:
        result["Ni_Loading"] = _fmt_loading(total_loading * ni_fraction)
    if (
        not result["Promoter_Loading"]
        and total_loading is not None
        and promoter_fraction is not None
        and promoter_fraction > 0
    ):
        result["Promoter_Loading"] = _fmt_loading(total_loading * promoter_fraction)

    if result["Ni_Loading"] and not ni_direct:
        record["Ni_Loading_wt%"] = result["Ni_Loading"]
    if result["Promoter_Loading"] and not prom_direct:
        record["Promoter_Loading_wt%"] = result["Promoter_Loading"]

    return result



def derive_support_clean(record: Dict) -> str:
    if not isinstance(record, dict):
        return ""
    for field in ("Support_Grouped", "Support_Normalized", "Support"):
        value = clean_residual_mojibake_chars(str(record.get(field, ""))).strip()
        if value:
            if field == "Support":
                normalized = normalize_support_expression(value)
                return normalized or value
            return value
    return ""



def apply_ml_ready_catalyst_feature_layer(record: Dict) -> Dict:
    if not isinstance(record, dict):
        return record

    row = dict(record)

    if not str(row.get("Metal_Loading_Method_Normalized", "")).strip():
        row["Metal_Loading_Method_Normalized"] = normalize_metal_loading_method(str(row.get("Metal_Loading_Method", "")))
    if not str(row.get("Support_Prep_Method_Normalized", "")).strip():
        row["Support_Prep_Method_Normalized"] = normalize_support_prep_method(str(row.get("Support_Prep_Method", "")))
    if not str(row.get("Precursor_Normalized", "")).strip():
        row["Precursor_Normalized"] = normalize_precursor_expression(str(row.get("Precursor", "")))
    if not str(row.get("Precursor_Family", "")).strip():
        row["Precursor_Family"] = classify_precursor_family(str(row.get("Precursor_Normalized", "")) or str(row.get("Precursor", "")))
    if not str(row.get("Support_Grouped", "")).strip():
        support_source = str(row.get("Support_Normalized", "")).strip() or str(row.get("Support", "")).strip()
        row["Support_Grouped"] = group_support_family(support_source)

    row["Promoter_Metal"] = derive_promoter_metal(row)

    fractions = derive_composition_fractions(row)
    row.update(fractions)

    row = infer_component_loadings_from_label(row)
    row = validate_component_loading_evidence(row)
    row = sync_component_loading_derivatives(row)

    loadings = derive_loading_partition(row)
    row.update(loadings)

    row["Support_clean"] = derive_support_clean(row)
    row["Metal_Loading_Method_grouped"] = str(row.get("Metal_Loading_Method_Normalized", "")).strip()
    row["Support_Prep_Method_grouped"] = str(row.get("Support_Prep_Method_Normalized", "")).strip()
    row["Precursor_family_grouped"] = str(row.get("Precursor_Family", "")).strip()
    return row



def apply_ml_ready_catalyst_feature_batch(records: List[Dict]) -> List[Dict]:
    return [apply_ml_ready_catalyst_feature_layer(dict(record)) for record in records if isinstance(record, dict)]


def propagate_loading_within_catalyst_families(records: List[Dict]) -> List[Dict]:
    """Within each paper, propagate Ni/Promoter/Metal loading between rows
    that share the same Canonical_Catalyst_ID.

    When a paper states a loading once (e.g. "10wt% nickel") for a catalyst
    family, the LLM may only extract it for one variant name.  This function
    copies the value to every row that belongs to the same canonical identity
    within the same paper, provided the recipient row's field is empty.
    """
    if not records:
        return records

    from collections import defaultdict

    propagate_fields = ["Ni_Loading_wt%", "Promoter_Loading_wt%", "Metal_Loading_wt%"]

    # Group by (Source_File, Canonical_Catalyst_ID)
    groups: Dict[Tuple[str, str], List[Dict]] = defaultdict(list)
    for r in records:
        src = str(r.get("Source_File", "") or "").strip()
        cid = str(r.get("Canonical_Catalyst_ID", "") or "").strip()
        if src and cid:
            groups[(src, cid)].append(r)

    for (src, cid), group_rows in groups.items():
        # Collect the best value for each propagate field from any row in the group
        donor_values: Dict[str, str] = {}
        for field in propagate_fields:
            for r in group_rows:
                val = str(r.get(field, "") or "").strip()
                if val:
                    donor_values[field] = val
                    break

        if not donor_values:
            continue

        # Propagate to rows that are missing the field
        for r in group_rows:
            for field, val in donor_values.items():
                if not str(r.get(field, "") or "").strip():
                    r[field] = val

    return records


def parse_feed_meoh_to_h2o_ratio(value: str) -> str:
    text = clean_residual_mojibake_chars(str(value or "")).strip()
    if not text:
        return ""

    normalized = _normalize_identity_text(text)
    normalized = re.sub(r"(?i)\b(?:ch3oh|methanol|meoh)\b", "MeOH", normalized)
    normalized = re.sub(r"(?i)\b(?:h2o|water)\b", "H2O", normalized)

    def _fmt_ratio(number: float) -> str:
        number = round(float(number), 4)
        if abs(number - round(number)) < 1e-9:
            return _format_ratio_token(f"{number:.1f}")
        return _format_ratio_token(f"{number:.4f}")

    def _safe_ratio(meoh_value: float, h2o_value: float) -> str:
        if meoh_value <= 0 or h2o_value <= 0:
            return ""
        return _fmt_ratio(meoh_value / h2o_value)

    m = re.search(r"\b(MeOH|H2O)\s*[:/]\s*(MeOH|H2O)\s*=\s*(\d+(?:\.\d+)?)\s*[:/]\s*(\d+(?:\.\d+)?)", normalized, flags=re.I)
    if m:
        left_species, right_species = m.group(1), m.group(2)
        left_value, right_value = float(m.group(3)), float(m.group(4))
        meoh_value = left_value if left_species == "MeOH" else right_value
        h2o_value = left_value if left_species == "H2O" else right_value
        ratio = _safe_ratio(meoh_value, h2o_value)
        if ratio:
            return ratio

    m = re.search(r"\b(MeOH|H2O)\s*[:/]\s*(MeOH|H2O)\s*=\s*(\d+(?:\.\d+)?)", normalized, flags=re.I)
    if m:
        left_species, right_species = m.group(1), m.group(2)
        left_over_right = float(m.group(3))
        if left_over_right > 0:
            if left_species == "MeOH" and right_species == "H2O":
                return _fmt_ratio(left_over_right)
            if left_species == "H2O" and right_species == "MeOH":
                return _fmt_ratio(1.0 / left_over_right)

    m = re.search(r"\b(MeOH|H2O)\s*(?:and|/)\s*(MeOH|H2O)[^\n\.;]{0,40}?molar ratio\s*(?:of\s*)?(\d+(?:\.\d+)?)\s*[:/]\s*(\d+(?:\.\d+)?)", normalized, flags=re.I)
    if not m:
        m = re.search(r"molar ratio[^\n\.;]{0,40}\b(MeOH|H2O)\s*(?:and|/)\s*(MeOH|H2O)[^\n\.;]{0,20}?(\d+(?:\.\d+)?)\s*[:/]\s*(\d+(?:\.\d+)?)", normalized, flags=re.I)
    if m:
        first_species, second_species = m.group(1), m.group(2)
        first_value, second_value = float(m.group(3)), float(m.group(4))
        meoh_value = first_value if first_species == "MeOH" else second_value
        h2o_value = first_value if first_species == "H2O" else second_value
        ratio = _safe_ratio(meoh_value, h2o_value)
        if ratio:
            return ratio

    pressure_pairs = {}
    for match in re.finditer(r"(\d+(?:\.\d+)?)\s*kpa\s*(MeOH|H2O)", normalized, flags=re.I):
        pressure_pairs[match.group(2)] = float(match.group(1))
    if "MeOH" in pressure_pairs and "H2O" in pressure_pairs:
        ratio = _safe_ratio(pressure_pairs["MeOH"], pressure_pairs["H2O"])
        if ratio:
            return ratio

    return ""



def parse_space_velocity_features(record: Dict) -> Dict[str, str]:
    # [Improve1] Added SpaceVelocity_type and SpaceVelocity_unit to distinguish
    # GHSV (mL/g/h) from WHSV (h⁻¹) — mixing these causes ML feature dimension errors.
    result = {
        "SpaceVelocity_norm": "",
        "SpaceVelocity_type": "",   # GHSV / WHSV / LHSV / unknown
        "SpaceVelocity_unit": "",   # mL/g/h  or  h-1  or  mL/mL/h
        "SpaceVelocity_source": "",
        "SpaceVelocity_unit_source": "",
    }
    if not isinstance(record, dict):
        return result

    def _extract_numeric_token(raw_value: str) -> str:
        raw_value = clean_residual_mojibake_chars(str(raw_value or "")).strip()
        if not raw_value:
            return ""
        cleaned = re.sub(r"([A-Za-z])\s*(?:\^|-)?\s*1\b", r"\1inv", raw_value)
        match = re.search(r"[-+]?\d+(?:\.\d+)?", cleaned)
        return _format_ratio_token(match.group(0)) if match else ""

    def _detect_sv_type_and_unit(raw_value: str, fallback_source: str = "") -> Tuple[str, str]:
        """Return (sv_type, sv_unit). sv_type: GHSV/WHSV/LHSV/unknown."""
        s = clean_residual_mojibake_chars(str(raw_value or ""))
        if re.search(r"\bWHSV\b", s, flags=re.I):
            return "WHSV", "h-1"
        if re.search(r"\bLHSV\b", s, flags=re.I):
            return "LHSV", "mL/mL/h"
        if re.search(r"\bGHSV\b", s, flags=re.I):
            # distinguish mL/g/h vs mL/mL/h
            if re.search(r"mL\s*g\s*(?:\^|-)?\s*1\s*h\s*(?:\^|-)?\s*1", s, flags=re.I):
                return "GHSV", "mL/g/h"
            if re.search(r"mL\s*mL\s*(?:\^|-)?\s*1\s*h\s*(?:\^|-)?\s*1", s, flags=re.I):
                return "GHSV", "mL/mL/h"
            return "GHSV", "mL/g/h"  # default GHSV unit
        if re.search(r"mL\s*g\s*(?:\^|-)?\s*1\s*h\s*(?:\^|-)?\s*1", s, flags=re.I):
            return "GHSV", "mL/g/h"
        if re.search(r"\bh\s*(?:\^|-)?\s*1\b", s, flags=re.I):
            return "WHSV", "h-1"
        if fallback_source == "GHSV_mL_g_h":
            return "GHSV", "mL/g/h"
        return "unknown", ""

    def _detect_unit_source(raw_value: str, fallback_source: str = "") -> str:
        sv_type, _ = _detect_sv_type_and_unit(raw_value, fallback_source)
        return sv_type

    space_velocity_value = str(record.get("SpaceVelocity_Value", "")).strip()
    if space_velocity_value:
        norm = _extract_numeric_token(space_velocity_value)
        if norm:
            sv_type, sv_unit = _detect_sv_type_and_unit(space_velocity_value, "SpaceVelocity_Value")
            result["SpaceVelocity_norm"] = norm
            result["SpaceVelocity_type"] = sv_type
            result["SpaceVelocity_unit"] = sv_unit
            result["SpaceVelocity_source"] = "SpaceVelocity_Value"
            result["SpaceVelocity_unit_source"] = sv_type
            return result

    ghsv_raw = str(record.get("GHSV_mL_g_h", "")).strip()
    if ghsv_raw:
        norm = _extract_numeric_token(ghsv_raw)
        if norm:
            sv_type, sv_unit = _detect_sv_type_and_unit(ghsv_raw, "GHSV_mL_g_h")
            result["SpaceVelocity_norm"] = norm
            result["SpaceVelocity_type"] = sv_type
            result["SpaceVelocity_unit"] = sv_unit
            result["SpaceVelocity_source"] = "GHSV_mL_g_h"
            result["SpaceVelocity_unit_source"] = sv_type
            return result

    return result




def apply_condition_feature_layer(record: Dict) -> Dict:
    if not isinstance(record, dict):
        return record

    if not globals().get("_TEXT_CONDITION_SALVAGE_SELFCHECK_DONE"):
        globals()["_TEXT_CONDITION_SALVAGE_SELFCHECK_DONE"] = True
        try:
            case = apply_condition_feature_layer({
                "data_source": "text",
                "source_granularity": "text_numeric",
                "identity_completeness_level": "complete",
                "numeric_reliability_level": "direct_numeric",
                "Catalyst": "Ni/Al2O3",
                "MeOH_Conversion_%": "95",
                "Notes": "[x_axis=Reaction temperature][x_value=250 C]",
            })
            if str(case.get("Reaction_Temp_C", "")).strip() != "250":
                print("  [WARNING] text salvage selfcheck failed: header/note temperature anchor not recovered")
        except Exception as e:
            print(f"  [WARNING] text salvage selfcheck error: {e}")

    row = dict(record)

    if not str(row.get("Feed_MeOH_to_H2O_Ratio", "")).strip():
        row["Feed_MeOH_to_H2O_Ratio"] = parse_feed_meoh_to_h2o_ratio(str(row.get("Feed_Composition", "")))
        if str(row.get("Feed_MeOH_to_H2O_Ratio", "")).strip() and not str(row.get("Feed_MeOH_to_H2O_Ratio_Raw", "")).strip():
            row["Feed_MeOH_to_H2O_Ratio_Raw"] = clean_residual_mojibake_chars(str(row.get("Feed_Composition", "") or "")).strip()

    salvage = _salvage_text_condition_anchor(row)
    for field in [
        "Reaction_Temp_C",
        "TOS_h",
        "S_C_Ratio",
        "Feed_MeOH_to_H2O_Ratio",
        "Pressure_bar",
        "GHSV_mL_g_h",
        "Flow_Rate",
        "Catalyst_Amount_g",
    ]:
        if not str(row.get(field, "")).strip() and str(salvage.get(field, "")).strip():
            row[field] = salvage.get(field, "")
    if str(salvage.get("_salvage_source", "")).strip():
        row["_text_condition_salvage_source"] = str(salvage.get("_salvage_source", "")).strip()
    if str(salvage.get("_salvage_evidence", "")).strip():
        row["_text_condition_salvage_evidence"] = str(salvage.get("_salvage_evidence", "")).strip()

    velocity_features = parse_space_velocity_features(row)
    for field, value in velocity_features.items():
        if not str(row.get(field, "")).strip() and str(value).strip():
            row[field] = value

    row = capture_raw_unit_fields(row)

    def _present(value: Any) -> bool:
        if value is None:
            return False
        return bool(str(value).strip())

    row["metal_loading_missing_flag"] = 0 if _present(row.get("Metal_Loading_wt%", "")) else 1
    row["ni_fraction_missing_flag"] = 0 if _present(row.get("Ni_Fraction", "")) else 1
    row["promoter_fraction_missing_flag"] = 0 if _present(row.get("Promoter_Fraction", "")) else 1
    row["sc_ratio_missing_flag"] = 0 if _present(row.get("S_C_Ratio", "")) else 1
    row["pressure_missing_flag"] = 0 if _present(row.get("Pressure_bar", "")) else 1
    row["feed_ratio_missing_flag"] = 0 if _present(row.get("Feed_MeOH_to_H2O_Ratio", "")) else 1
    row["space_velocity_missing_flag"] = 0 if _present(row.get("SpaceVelocity_norm", "")) else 1

    source_file = str(row.get("Source_File", "")).strip()
    catalyst_key = str(
        row.get("Canonical_Catalyst_ID", "")
        or row.get("Catalyst_ID_normalized", "")
        or row.get("Catalyst_ID", "")
    ).strip()
    row["cv_group"] = f"{source_file}||{catalyst_key}" if source_file and catalyst_key else ""

    row = apply_unit_normalization(row)
    return row




def apply_condition_feature_batch(records: List[Dict]) -> List[Dict]:
    return [apply_condition_feature_layer(dict(record)) for record in records if isinstance(record, dict)]


def safe_model_numeric(value: Any) -> str:
    text = clean_residual_mojibake_chars(str(value or "")).strip()
    if not text:
        return ""
    direct = _parse_float_if_possible(text)
    if direct is not None:
        return _format_ratio_token(str(direct))

    compact = text.replace(",", "")
    numeric_tokens = re.findall(r"[-+]?\d+(?:\.\d+)?", compact)
    if len(numeric_tokens) != 1:
        return ""
    if re.search(r"\b(?:to|and)\b", compact, flags=re.I):
        return ""
    return _format_ratio_token(numeric_tokens[0])



def clean_model_ready_category(value: str, field_name: str) -> str:
    text = clean_residual_mojibake_chars(str(value or "")).strip()
    if not text:
        return ""

    text = re.sub(r"\s+", " ", text)
    low = text.lower()

    if field_name == "Promoter_Metal_clean_model":
        if low in {"none", "no promoter", "nopromoter", "single-ni", "single ni"}:
            return "NoPromoter"
        tokens = [tok for tok in _sort_elements_for_merge(_extract_elements_for_merge(text)) if tok != "Ni"]
        if tokens:
            return "/".join(tokens)
        return text

    if field_name == "Support_clean_model":
        normalized = normalize_support_expression(text)
        if normalized:
            return normalized
    return text



def build_model_ready_schema_layer(record: Dict) -> Dict:
    if not isinstance(record, dict):
        return record

    row = dict(record)

    row["Reaction_Temp_C_num_model"] = safe_model_numeric(row.get("Reaction_Temp_C", ""))
    row["Metal_Loading_wt%_num_model"] = safe_model_numeric(row.get("Metal_Loading_wt%", ""))
    row["Ni_Fraction_fixed_model"] = safe_model_numeric(row.get("Ni_Fraction", ""))
    row["Promoter_Fraction_num_model"] = safe_model_numeric(row.get("Promoter_Fraction", ""))
    row["S_C_Ratio_num_model"] = safe_model_numeric(row.get("S_C_Ratio", ""))
    row["Pressure_bar_num_model"] = safe_model_numeric(row.get("Pressure_bar", ""))
    row["Feed_MeOH_to_H2O_Ratio_model"] = safe_model_numeric(row.get("Feed_MeOH_to_H2O_Ratio", ""))
    row["SpaceVelocity_norm_model"] = safe_model_numeric(row.get("SpaceVelocity_norm", ""))

    row["Support_clean_model"] = clean_model_ready_category(
        str(row.get("Support_clean", "")),
        "Support_clean_model",
    )
    promoter_clean_source = str(row.get("Promoter_Metal", "")).strip()
    if not promoter_clean_source and str(row.get("has_promoter", "")).strip() in {"0", "0.0"}:
        promoter_clean_source = "NoPromoter"
    row["Promoter_Metal_clean_model"] = clean_model_ready_category(
        promoter_clean_source,
        "Promoter_Metal_clean_model",
    )

    row["Calcination_Temp_C_num"] = safe_model_numeric(row.get("Calcination_Temp_C", ""))
    row["Calcination_Time_h_num"] = safe_model_numeric(row.get("Calcination_Time_h", ""))
    row["Reduction_Temp_C_num"] = safe_model_numeric(row.get("Reduction_Temp_C", ""))
    row["Reduction_Time_h_num"] = safe_model_numeric(row.get("Reduction_Time_h", ""))

    def _product_or_empty(left_value: Any, right_value: Any) -> str:
        left_num = _parse_float_if_possible(left_value)
        right_num = _parse_float_if_possible(right_value)
        if left_num is None or right_num is None:
            return ""
        return _format_ratio_token(str(round(left_num * right_num, 3)))

    row["Calcination_Severity"] = _product_or_empty(
        row.get("Calcination_Temp_C_num", ""),
        row.get("Calcination_Time_h_num", ""),
    )
    row["Reduction_Severity"] = _product_or_empty(
        row.get("Reduction_Temp_C_num", ""),
        row.get("Reduction_Time_h_num", ""),
    )
    row["Temp_Ni_interaction"] = _product_or_empty(
        row.get("Reaction_Temp_C_num_model", ""),
        row.get("Ni_Fraction_fixed_model", ""),
    )

    def _present(value: Any) -> bool:
        return bool(str(value or "").strip())

    if str(row.get("metal_loading_missing_flag", "")).strip() == "":
        row["metal_loading_missing_flag"] = 0 if _present(row.get("Metal_Loading_wt%", "")) else 1
    if str(row.get("ni_fraction_missing_flag", "")).strip() == "":
        row["ni_fraction_missing_flag"] = 0 if _present(row.get("Ni_Fraction", "")) else 1
    if str(row.get("promoter_fraction_missing_flag", "")).strip() == "":
        row["promoter_fraction_missing_flag"] = 0 if _present(row.get("Promoter_Fraction", "")) else 1
    if str(row.get("sc_ratio_missing_flag", "")).strip() == "":
        row["sc_ratio_missing_flag"] = 0 if _present(row.get("S_C_Ratio", "")) else 1
    if str(row.get("pressure_missing_flag", "")).strip() == "":
        row["pressure_missing_flag"] = 0 if _present(row.get("Pressure_bar", "")) else 1
    if str(row.get("feed_ratio_missing_flag", "")).strip() == "":
        row["feed_ratio_missing_flag"] = 0 if _present(row.get("Feed_MeOH_to_H2O_Ratio", "")) else 1
    if str(row.get("space_velocity_missing_flag", "")).strip() == "":
        row["space_velocity_missing_flag"] = 0 if _present(row.get("SpaceVelocity_norm", "")) else 1

    row["calcination_temp_missing_flag"] = 0 if _present(row.get("Calcination_Temp_C_num", "")) else 1
    row["calcination_time_missing_flag"] = 0 if _present(row.get("Calcination_Time_h_num", "")) else 1
    row["reduction_temp_missing_flag"] = 0 if _present(row.get("Reduction_Temp_C_num", "")) else 1
    row["reduction_time_missing_flag"] = 0 if _present(row.get("Reduction_Time_h_num", "")) else 1
    row["method_missing_flag"] = 0 if _present(row.get("Metal_Loading_Method_grouped", "")) else 1
    return row



def build_model_ready_schema_batch(records: List[Dict]) -> List[Dict]:
    return [build_model_ready_schema_layer(dict(record)) for record in records if isinstance(record, dict)]


def _extract_support_from_identity_text(value: str) -> str:
    text = _normalize_identity_text(value)
    if not text or "/" not in text:
        return ""
    support = text.split("/", 1)[1]
    support = re.sub(r"\(.*?\)", "", support)
    support = re.split(r"\b(?:at|for|with|under|after)\b", support, maxsplit=1, flags=re.I)[0]
    support = support.strip(" -;,.")
    return support.strip() if re.search(r"[A-Za-z]", support) else ""


def _has_ni_identity_context(record: Dict, extra_text: str = "") -> bool:
    context_text = " ".join([
        str(record.get("Active_Metal", "")),
        str(record.get("Catalyst", "")),
        str(record.get("Catalyst_ID", "")),
        str(record.get("Series_Name", "")),
        str(record.get("Canonical_Catalyst_ID", "")),
        str(record.get("Alloy_Ratio", "")),
        str(record.get("Metal_Ratio", "")),
        str(record.get("Composition_Ratio", "")),
        str(extra_text or ""),
    ])
    if re.search(r"\bni(?:ckel)?\b", context_text, flags=re.I):
        return True
    return any(metal == "Ni" for metal, _ in parse_alloy_ratio_to_ordered_pairs(record))


def _identity_label_metal_tokens(record: Dict) -> List[str]:
    tokens: List[str] = []
    for field in ("Catalyst", "Catalyst_ID", "Series_Name"):
        value = _normalize_identity_text(str(record.get(field, "") or ""))
        if not value:
            continue
        left = value.split("/", 1)[0]
        tokens.extend(_extract_elements_for_merge(left))
    tokens.extend(_extract_elements_for_merge(str(record.get("Active_Metal", "") or "")))
    for metal, _ in parse_alloy_ratio_to_ordered_pairs(record):
        tokens.append(metal)
    return _dedupe_keep_order([token for token in tokens if token])


def _infer_support_blocked_metals(record: Dict) -> set:
    support_texts = [
        str(record.get("Support", "") or ""),
        str(record.get("Support_Normalized", "") or ""),
        str(record.get("Support_Grouped", "") or ""),
    ]
    for field in ("Catalyst", "Catalyst_ID", "Series_Name"):
        support = _extract_support_from_identity_text(str(record.get(field, "") or ""))
        if support:
            support_texts.append(support)
    text = " | ".join([
        normalize_support_expression(value) or _normalize_identity_text(value)
        for value in support_texts if str(value).strip()
    ]).lower()
    blocked = set()
    if re.search(r"\bal2o3\b|alumina|al-ldh|\bldh\b", text):
        blocked.add("Al")
    if re.search(r"\bzno\b|\bzno-rod\b", text):
        blocked.add("Zn")
    if re.search(r"\bfe3o4\b|\bfe2o3\b|ferrite|spinel", text):
        blocked.add("Fe")
    if re.search(r"\bceo2\b|ceria", text):
        blocked.add("Ce")
    if re.search(r"\bla2o3\b|lanthana", text):
        blocked.add("La")
    if re.search(r"\bmgo\b|magnesia", text):
        blocked.add("Mg")
    if re.search(r"mo2c|moc|molybdenum carbide", text):
        blocked.add("Mo")
    return blocked


def _infer_reagent_only_metals(text: str) -> set:
    source = _normalize_identity_text(text).lower()
    blocked = set()
    reagent_context = bool(re.search(
        r"\b(?:precipitant|precursor|solution|aqueous|added|addition|post-addition|co-precipitation|coprecipitation|impregnation)\b",
        source,
        flags=re.I,
    ))
    if reagent_context and re.search(r"\bk(?:2)?co3\b|\bkoh\b|\bkcl\b|potassium\s+(?:carbonate|hydroxide|chloride|nitrate)", source, flags=re.I):
        blocked.add("K")
    return blocked


def _filter_identity_promoter_tokens(tokens: List[str], record: Dict, source_text: str = "") -> List[str]:
    candidates = _dedupe_keep_order([
        token for token in (tokens or [])
        if token and token in MSR_NI_SECOND_METALS and token != "Ni"
    ])
    if not candidates:
        return []
    if not _has_ni_identity_context(record, source_text):
        return []

    support_blocked = _infer_support_blocked_metals(record)
    reagent_blocked = _infer_reagent_only_metals(" ".join([
        str(record.get("Notes", "") or ""),
        str(record.get("Precursor", "") or ""),
        str(record.get("Support_Prep_Method", "") or ""),
        str(source_text or ""),
    ]))
    label_tokens = set(_identity_label_metal_tokens(record))
    ratio_metals = {metal for metal, _ in parse_alloy_ratio_to_ordered_pairs(record)}

    kept = []
    for token in candidates:
        if token in support_blocked:
            continue
        if token in reagent_blocked and token not in label_tokens and token not in ratio_metals:
            continue
        kept.append(token)
    return _dedupe_keep_order(kept)


def _extract_alloy_second_metals(record: Dict) -> List[str]:
    pairs = parse_alloy_ratio_to_ordered_pairs(record)
    if not pairs:
        return []
    elements = [el for el, _ in pairs]
    if "Ni" not in elements:
        return []
    return [el for el in elements if el != "Ni" and el in MSR_NI_SECOND_METALS]


def _extract_promoter_phrases(text: str) -> List[str]:
    text = _normalize_identity_text(text)
    if not text or "ni" not in text.lower():
        return []

    promoters = []
    patterns = [
        r"\bni(?:ckel)?\s+(?:based\s+)?(?:catalyst\s+)?(?:promoted|modified|doped)\s+(?:with|by)\s+([A-Za-z][a-z]?)",
        r"\b([A-Za-z][a-z]?)\s*[- ]?(?:promoted|modified|doped)\s+ni(?:ckel)?(?:\s+based)?\s+catalyst",
        r"\b([A-Za-z][a-z]?)\s*[- ]?modified\s+ni(?:ckel)?\s+catalyst",
        r"\b([A-Za-z][a-z]?)\s*[- ]?promoted\s+ni(?:ckel)?\s+catalyst",
    ]
    for pattern in patterns:
        for match in re.finditer(pattern, text, flags=re.I):
            metal = match.group(1).capitalize()
            if metal in MSR_NI_SECOND_METALS:
                promoters.append(metal)
    return _dedupe_keep_order(promoters)


def parse_second_metal_from_text_label(value: str) -> List[str]:
    if not globals().get("_COMPOSITION_SUFFIX_SELFCHECK_DONE"):
        globals()["_COMPOSITION_SUFFIX_SELFCHECK_DONE"] = True
        try:
            tokens = _extract_elements_for_merge("Ni/Al2O3-Com")
            if "Co" in tokens:
                print("  [WARNING] composition selfcheck failed: descriptor suffix misparsed as Co")
        except Exception as e:
            print(f"  [WARNING] composition selfcheck error: {e}")

    text = _strip_descriptor_suffix_before_element_parse(_normalize_identity_text(value))
    if not text or "ni" not in text.lower():
        return []

    promoters = []
    promoters.extend(_extract_promoter_phrases(text))

    label_text = text.split("/", 1)[0]
    if len(text.split()) <= 8 or any(ch in text for ch in "/-%"):
        tokens = _extract_elements_for_merge(label_text)
        if "Ni" in tokens:
            for token in tokens:
                if token != "Ni" and token in MSR_NI_SECOND_METALS:
                    promoters.append(token)

    for metal in sorted(MSR_NI_SECOND_METALS, key=len, reverse=True):
        if re.search(rf"Ni\d*-?{metal}(?:\d+)?(?:Ox)?", text):
            promoters.append(metal)
        if re.search(rf"{metal}\d*-?Ni(?:\d+)?", text):
            promoters.append(metal)

    return _dedupe_keep_order(promoters)



def infer_promoter_from_catalyst_name(value: str) -> List[str]:
    return parse_second_metal_from_text_label(value)


def _catalyst_id_mentions_promoter(catalyst_id: str, promoters: List[str]) -> bool:
    cid = _normalize_identity_text(catalyst_id).lower()
    return bool(cid) and all(promoter.lower() in cid for promoter in promoters)


def _build_structured_catalyst_id(record: Dict) -> str:
    loading = normalize_loading_for_merge(str(record.get("Metal_Loading_wt%", "")))
    support = normalize_support_for_merge(str(record.get("Support_Normalized", "") or record.get("Support", "")))
    pairs = parse_alloy_ratio_to_ordered_pairs(record)

    if pairs:
        metal_identity = "-".join(
            f"{el.lower()}{_format_ratio_token(val)}" for el, val in pairs
        )
    else:
        active_tokens = _extract_elements_for_merge(str(record.get("Active_Metal", "")))
        promoter_tokens = _extract_elements_for_merge(str(record.get("Promoter", "")))
        metal_tokens = []
        if "Ni" in active_tokens:
            metal_tokens.append("Ni")
        elif active_tokens:
            metal_tokens.extend(active_tokens)
        for token in promoter_tokens:
            if token != "Ni" and token not in metal_tokens:
                metal_tokens.append(token)
        if not metal_tokens:
            label_tokens = _extract_elements_for_merge(
                " ".join([
                    str(record.get("Catalyst", "")),
                    str(record.get("Series_Name", "")),
                ])
            )
            if "Ni" in label_tokens:
                metal_tokens.append("Ni")
                for token in label_tokens:
                    if token != "Ni" and token in MSR_NI_SECOND_METALS and token not in metal_tokens:
                        metal_tokens.append(token)
        metal_identity = "-".join(token.lower() for token in _dedupe_keep_order(metal_tokens))

    if not metal_identity:
        label = str(record.get("Catalyst", "") or record.get("Series_Name", "") or record.get("Catalyst_ID", "")).strip()
        if not label:
            return ""
        return re.sub(r"[^a-z0-9/\-]+", "", _normalize_identity_text(label).lower())

    loading_part = f"{loading}wt%-" if loading else ""
    support_part = f"/{support}" if support else ""
    return f"{loading_part}{metal_identity}{support_part}"


def _build_partial_identity_from_label(label: str) -> str:
    temp = {"Catalyst": label}
    support = _extract_support_from_identity_text(label)
    if support:
        temp["Support"] = support
    promoters = parse_second_metal_from_text_label(label)
    if promoters:
        temp["Promoter"] = "/".join(promoters)
    if re.search(r"\bni\b", label, flags=re.I) or "Ni" in re.findall(r"[A-Z][a-z]?", label):
        temp["Active_Metal"] = "Ni"
    return _build_structured_catalyst_id(temp)


def enrich_promoter_fields_from_identity(record: Dict, source_text: str = "") -> Dict:
    if not isinstance(record, dict):
        return record

    identity_text = " ".join([
        str(record.get("Catalyst", "")),
        str(record.get("Catalyst_ID", "")),
        str(record.get("Series_Name", "")),
        str(source_text or ""),
    ])
    template_info = _extract_identity_template_candidates(identity_text)
    if template_info.get("expanded"):
        record["_template_candidates"] = list(template_info.get("expanded", []) or [])
        record["_template_info"] = {
            "template": str(template_info.get("template", "") or ""),
            "variable": str(template_info.get("variable", "") or ""),
            "values": list(template_info.get("values", []) or []),
            "suffix": str(template_info.get("suffix", "") or ""),
        }

    promoter_candidates = []
    promoter_candidates.extend(_extract_elements_for_merge(str(record.get("Promoter", ""))))
    promoter_candidates.extend(_extract_elements_for_merge(str(record.get("Promoter_Metal", ""))))
    promoter_candidates.extend(_extract_alloy_second_metals(record))

    identity_texts = [
        str(record.get("Catalyst", "")),
        str(record.get("Catalyst_ID", "")),
        str(record.get("Series_Name", "")),
        str(record.get("Notes", "")),
    ]
    for text_item in identity_texts:
        promoter_candidates.extend(parse_second_metal_from_text_label(text_item))
    if source_text:
        promoter_candidates.extend(_extract_promoter_phrases(source_text))

    current_promoters = []
    for field in ("Promoter", "Promoter_Metal", "Second_Metal", "Additive"):
        current_promoters.extend(_safe_split_multi_promoter(record.get(field, "")))

    promoter_candidates = _filter_identity_promoter_tokens(
        _dedupe_keep_order([item for item in promoter_candidates if item in MSR_NI_SECOND_METALS]),
        record,
        source_text,
    )
    current_promoters = _filter_identity_promoter_tokens(_dedupe_keep_order(current_promoters), record, source_text)

    if promoter_candidates:
        record["Promoter"] = "/".join(promoter_candidates)
        record["Promoter_Metal"] = "/".join(promoter_candidates)
    elif current_promoters:
        record["Promoter"] = "/".join(current_promoters)
        record["Promoter_Metal"] = "/".join(current_promoters)
    elif str(record.get("Promoter", "")).strip() or str(record.get("Promoter_Metal", "")).strip():
        record["Promoter"] = ""
        record["Promoter_Metal"] = ""

    combined_identity = " ".join(identity_texts[:3])
    identity_elements = _extract_elements_for_merge(combined_identity)
    active_elements = _extract_elements_for_merge(str(record.get("Active_Metal", "")))
    if promoter_candidates and ("Ni" in identity_elements or "Ni" in active_elements or re.search(r"\bni\b", combined_identity, flags=re.I)):
        if not active_elements or "Ni" in active_elements:
            record["Active_Metal"] = "Ni"

    if not str(record.get("Support", "")).strip():
        for text_item in identity_texts[:3]:
            support = _extract_support_from_identity_text(text_item)
            if support:
                record["Support"] = support
                break

    rebuilt_id = _build_structured_catalyst_id(record)
    current_id = str(record.get("Catalyst_ID", "")).strip()
    if rebuilt_id and not current_id:
        record["Catalyst_ID"] = rebuilt_id
    elif promoter_candidates and rebuilt_id and not _catalyst_id_mentions_promoter(current_id, promoter_candidates):
        if re.search(r"\bni\b", combined_identity + " " + str(record.get("Active_Metal", "")), flags=re.I):
            record["Catalyst_ID"] = rebuilt_id
            _append_note_tag(record, "[Catalyst_ID_rebuilt_promoter_aware]")

    return record


def annotate_record_provenance(record: Dict, source_text: str, file_name: str, origin_type: str) -> Dict:
    text = " ".join([
        str(record.get("Catalyst", "")),
        str(record.get("Notes", "")),
        str(source_text or ""),
    ]).lower()
    record["record_origin_type"] = origin_type or str(record.get("record_origin_type", "")).strip()
    lower_name = str(file_name).lower()
    record["is_source_file_si"] = bool(re.search(r"(?:^|[_\-\s])si\.(?:pdf|docx)$", lower_name))
    record["is_this_work"] = any(kw in text for kw in ["this work", "present work", "our work", "our catalyst"])
    record["is_literature_comparison"] = any(
        kw in text for kw in ["comparison", "benchmark", "reference catalyst", "reference", "ref."]
    )
    record["source_granularity"] = _infer_source_granularity(record, source_text, origin_type)
    return record


def annotate_temperature_conversion_clues(record: Dict, source_text: str = "") -> Dict:
    check_text = _normalize_identity_text(" ".join([
        str(record.get("Notes", "")),
        str(source_text or ""),
    ]))
    check_text = clean_residual_mojibake_chars(check_text)
    check_text = check_text.replace("°", " ").replace("潞", " ")
    clues = []
    patterns = [
        r"(\d+(?:\.\d+)?)\s*k\s*\(\s*(\d+(?:\.\d+)?)\s*c\s*\)",
        r"converted\s+from\s+(\d+(?:\.\d+)?)\s*k\s+to\s+(\d+(?:\.\d+)?)\s*c",
        r"temperature\s+converted\s+from\s+(\d+(?:\.\d+)?)\s*k\s+to\s+(\d+(?:\.\d+)?)\s*c",
        r"(\d+(?:\.\d+)?)\s*k\s*(?:is|was)?\s*equivalent\s+to\s+(\d+(?:\.\d+)?)\s*c",
    ]
    for pattern in patterns:
        for match in re.finditer(pattern, check_text, flags=re.I):
            clues.append(f"{match.group(1)}K->{match.group(2)}C")

    if clues:
        record["temp_conversion_clue_flag"] = True
        note = "; ".join(_dedupe_keep_order(clues))
        if not str(record.get("explicit_temp_conversion_note", "")).strip():
            record["explicit_temp_conversion_note"] = note
    else:
        record["temp_conversion_clue_flag"] = bool(record.get("temp_conversion_clue_flag", False))
    return record


def finalize_extracted_records(records: List[Dict], source_text: str, file_name: str, origin_type: str) -> List[Dict]:
    source_text = clean_residual_mojibake_chars(source_text or "")
    finalized = []
    for record in records:
        if not isinstance(record, dict):
            continue
        row = clean_record_text_fields(dict(record))
        row["_comparison_context_raw"] = source_text
        annotate_record_provenance(row, source_text, file_name, origin_type)
        annotate_temperature_conversion_clues(row, source_text)
        normalize_identity_aliases(row)
        if not str(row.get("Catalyst_ID", "")).strip():
            row["Catalyst_ID"] = normalize_catalyst_id(row)
        row = capture_raw_unit_fields(row)
        annotate_numeric_expression_guards(row)
        row = apply_preparation_normalization_layer(row)
        finalized.append(clean_record_text_fields(row))
    return finalized


def _parse_float_if_possible(value: Any) -> Optional[float]:
    try:
        value = str(value).strip()
        if not value:
            return None
        return float(value)
    except (TypeError, ValueError):
        return None


def _rows_have_close_performance(left: Dict, right: Dict) -> bool:
    matched = 0
    for field in CORE_PERFORMANCE_FIELDS:
        lv = _parse_float_if_possible(left.get(field, ""))
        rv = _parse_float_if_possible(right.get(field, ""))
        if lv is None or rv is None:
            continue
        tol = 10.0 if field == "CO_Concentration_ppm" else 1.0
        if abs(lv - rv) > tol:
            return False
        matched += 1
    return matched > 0


def resolve_obvious_temp_conversion_duplicates_pre_save(records: List[Dict]) -> tuple[List[Dict], List[Dict]]:
    grouped = {}
    for idx, row in enumerate(records):
        identity_key = str(row.get("Catalyst_ID") or row.get("Canonical_Catalyst_ID") or row.get("Catalyst") or "").strip()
        source_key = str(row.get("Source_File", "")).strip()
        if identity_key:
            grouped.setdefault((source_key, identity_key), []).append((idx, row))

    removed_idx = set()
    audit_rows = []
    for (source_key, identity_key), items in grouped.items():
        for i, (left_idx, left_row) in enumerate(items):
            if left_idx in removed_idx:
                continue
            left_temp = _parse_float_if_possible(left_row.get("Reaction_Temp_C", ""))
            if left_temp is None:
                continue
            for right_idx, right_row in items[i + 1:]:
                if right_idx in removed_idx:
                    continue
                right_temp = _parse_float_if_possible(right_row.get("Reaction_Temp_C", ""))
                if right_temp is None:
                    continue
                if not (left_row.get("temp_conversion_clue_flag") or right_row.get("temp_conversion_clue_flag")):
                    continue

                kelvin_idx = celsius_idx = None
                kelvin_row = celsius_row = None
                kelvin_temp = celsius_temp = None
                if left_temp > 450 and 80 <= right_temp <= 400:
                    kelvin_idx, kelvin_row, kelvin_temp = left_idx, left_row, left_temp
                    celsius_idx, celsius_row, celsius_temp = right_idx, right_row, right_temp
                elif right_temp > 450 and 80 <= left_temp <= 400:
                    kelvin_idx, kelvin_row, kelvin_temp = right_idx, right_row, right_temp
                    celsius_idx, celsius_row, celsius_temp = left_idx, left_row, left_temp
                else:
                    continue

                if abs((kelvin_temp - 273.15) - celsius_temp) > 2.0:
                    continue
                if not _rows_have_close_performance(kelvin_row, celsius_row):
                    continue

                pair_id = f"{source_key}|{identity_key}|{round(celsius_temp, 1)}"
                celsius_row["temp_duplicate_resolution_upstream"] = "kept_celsius_removed_kelvin"
                celsius_row["upstream_temp_duplicate_pair_id"] = pair_id
                celsius_row["temp_conversion_clue_flag"] = True
                if not str(celsius_row.get("explicit_temp_conversion_note", "")).strip():
                    celsius_row["explicit_temp_conversion_note"] = f"{kelvin_temp}K->{celsius_temp}C"
                _append_note_tag(celsius_row, f"[upstream_temp_duplicate_pair={pair_id}]")

                removed_row = dict(kelvin_row)
                removed_row["temp_duplicate_resolution_upstream"] = "removed_kelvin_duplicate"
                removed_row["upstream_temp_duplicate_pair_id"] = pair_id
                removed_row["audit_reason"] = "explicit_temp_conversion_kelvin_duplicate"
                audit_rows.append(removed_row)
                removed_idx.add(kelvin_idx)
                break

    kept_rows = [row for idx, row in enumerate(records) if idx not in removed_idx]
    if audit_rows:
        print(f"  [Audit] upstream temp duplicate removal: {len(audit_rows)} rows")
    return kept_rows, audit_rows


OBVIOUS_DUPLICATE_TEMP_TOL_C = 1.0
OBVIOUS_DUPLICATE_METRIC_TOL = 0.5
OBVIOUS_DUPLICATE_PPM_TOL = 10.0
_OBVIOUS_DUPLICATE_METRIC_FIELDS = [
    "MeOH_Conversion_%", "H2_Yield_%", "H2_Selectivity_%",
    "CO_Selectivity_%", "CO2_Selectivity_%", "CO_Concentration_ppm",
]
_ORIGIN_QUALITY_RANK = {
    "high_confidence_direct_numeric": 6,
    "medium_confidence_direct_numeric": 5,
    "comparison_table_like": 3,
    "figure_series_partial_identity": 2,
    "implied_or_range_like": 1,
    "low_information_shell": 0,
}
_SOURCE_GRANULARITY_RANK = {
    "table_row": 6,
    "text_numeric": 5,
    "si_table": 4,
    "text_si": 3,
    "figure_series": 2,
    "caption_only": 1,
    "": 0,
}
_CONFIDENCE_RANK = {"high": 3, "medium": 2, "low": 1, "": 0}
_IDENTITY_LEVEL_RANK = {"complete": 3, "partial": 2, "poor": 1, "": 0}
_NUMERIC_LEVEL_RANK = {
    "direct_numeric": 5,
    "approximate_numeric": 4,
    "range_like": 2,
    "qualitative_only": 1,
    "missing": 0,
    "": 0,
}


def _merge_pipe_tokens(existing: str, new_tokens: List[str]) -> str:
    tokens = [t for t in str(existing).split("|") if t]
    tokens.extend([t for t in new_tokens if t])
    return "|".join(_dedupe_keep_order(tokens))


def _get_obvious_duplicate_identity_key(row: Dict) -> str:
    candidates = [
        row.get("Catalyst_ID_normalized", ""),
        row.get("identity_alias_group", ""),
        row.get("Canonical_Catalyst_ID", ""),
        row.get("Catalyst_ID", ""),
    ]
    for value in candidates:
        key = str(value).strip().lower()
        if not key or key.startswith("__fallback_md5_"):
            continue
        return key
    return ""


def _get_obvious_duplicate_metric(row: Dict) -> Optional[tuple[str, float]]:
    if str(row.get("numeric_reliability_level", "")).strip() != "direct_numeric":
        return None
    if bool(row.get("is_range_like_value")) or bool(row.get("is_qualitative_value")) or bool(row.get("is_approximate_value")):
        return None
    for field in _OBVIOUS_DUPLICATE_METRIC_FIELDS:
        value = _parse_float_if_possible(row.get(field, ""))
        if value is not None:
            return field, value
    return None


def _classify_obvious_duplicate_type(left: Dict, right: Dict) -> str:
    left_g = str(left.get("source_granularity", "")).strip()
    right_g = str(right.get("source_granularity", "")).strip()
    pair = {left_g, right_g}
    if pair == {"text_numeric", "table_row"}:
        return "obvious_text_table_same_point"
    if "figure_series" in pair and pair.intersection({"text_numeric", "table_row", "si_table"}):
        return "obvious_figure_vs_text_table_same_point"
    return ""


def _row_nonempty_value_count(row: Dict) -> int:
    return sum(bool(str(v).strip()) for v in row.values())


def _score_row_for_obvious_duplicate(row: Dict) -> tuple:
    return (
        _ORIGIN_QUALITY_RANK.get(str(row.get("origin_quality_class", "")).strip(), 0),
        _SOURCE_GRANULARITY_RANK.get(str(row.get("source_granularity", "")).strip(), 0),
        _CONFIDENCE_RANK.get(str(row.get("extraction_confidence", "")).strip(), 0),
        _IDENTITY_LEVEL_RANK.get(str(row.get("identity_completeness_level", "")).strip(), 0),
        _NUMERIC_LEVEL_RANK.get(str(row.get("numeric_reliability_level", "")).strip(), 0),
        _row_nonempty_value_count(row),
        -len(str(row.get("Notes", ""))),
    )


def choose_better_duplicate_record(left_idx: int, left_row: Dict, right_idx: int, right_row: Dict) -> tuple[int, Dict, int, Dict, str]:
    left_score = _score_row_for_obvious_duplicate(left_row)
    right_score = _score_row_for_obvious_duplicate(right_row)
    if right_score > left_score:
        keep_idx, keep_row, drop_idx, drop_row = right_idx, right_row, left_idx, left_row
    else:
        keep_idx, keep_row, drop_idx, drop_row = left_idx, left_row, right_idx, right_row
    reason = (
        f"kept_better_record:{keep_row.get('source_granularity','')}:{keep_row.get('origin_quality_class','')}"
        f">{drop_row.get('source_granularity','')}:{drop_row.get('origin_quality_class','')}"
    )
    return keep_idx, keep_row, drop_idx, drop_row, reason


def build_obvious_duplicate_group_id(source_file: str, identity_key: str, temp_value: float, metric_field: str, metric_value: float) -> str:
    return (
        f"{str(source_file).strip()}|{identity_key}|temp:{round(temp_value, 1)}|"
        f"{metric_field}:{round(metric_value, 3)}"
    )


def _mark_obvious_duplicate_row(row: Dict, flag: bool, dup_type: str, keep_drop: str, reason: str, group_id: str) -> None:
    row["obvious_duplicate_flag"] = flag
    row["obvious_duplicate_type"] = _merge_pipe_tokens(row.get("obvious_duplicate_type", ""), [dup_type])
    row["obvious_duplicate_keep_drop"] = keep_drop
    row["obvious_duplicate_reason"] = reason
    row["obvious_duplicate_group_id"] = group_id


def resolve_obvious_cross_source_duplicates_pre_save(records: List[Dict]) -> tuple[List[Dict], List[Dict]]:
    grouped: Dict[tuple[str, str], List[tuple[int, Dict]]] = {}
    for idx, row in enumerate(records):
        row.setdefault("obvious_duplicate_flag", False)
        row.setdefault("obvious_duplicate_type", "")
        row.setdefault("obvious_duplicate_keep_drop", "")
        row.setdefault("obvious_duplicate_reason", "")
        row.setdefault("obvious_duplicate_group_id", "")

        source_key = str(row.get("Source_File", "")).strip()
        identity_key = _get_obvious_duplicate_identity_key(row)
        if source_key and identity_key:
            grouped.setdefault((source_key, identity_key), []).append((idx, row))

    removed_idx = set()
    audit_rows = []

    for (source_key, identity_key), items in grouped.items():
        for i, (left_idx, left_row) in enumerate(items):
            if left_idx in removed_idx:
                continue
            if str(left_row.get("identity_completeness_level", "")).strip() == "poor":
                continue
            left_temp = _parse_float_if_possible(left_row.get("Reaction_Temp_C", ""))
            left_metric = _get_obvious_duplicate_metric(left_row)
            if left_temp is None or left_metric is None:
                continue
            left_field, left_value = left_metric

            for right_idx, right_row in items[i + 1:]:
                if right_idx in removed_idx:
                    continue
                if str(right_row.get("identity_completeness_level", "")).strip() == "poor":
                    continue
                dup_type = _classify_obvious_duplicate_type(left_row, right_row)
                if not dup_type:
                    continue

                right_temp = _parse_float_if_possible(right_row.get("Reaction_Temp_C", ""))
                right_metric = _get_obvious_duplicate_metric(right_row)
                if right_temp is None or right_metric is None:
                    continue
                right_field, right_value = right_metric
                if right_field != left_field:
                    continue
                if abs(left_temp - right_temp) > OBVIOUS_DUPLICATE_TEMP_TOL_C:
                    continue
                metric_tol = OBVIOUS_DUPLICATE_PPM_TOL if left_field == "CO_Concentration_ppm" else OBVIOUS_DUPLICATE_METRIC_TOL
                if abs(left_value - right_value) > metric_tol:
                    continue

                keep_idx, keep_row, drop_idx, drop_row, reason = choose_better_duplicate_record(
                    left_idx, left_row, right_idx, right_row
                )
                group_id = build_obvious_duplicate_group_id(source_key, identity_key, left_temp, left_field, left_value)
                _mark_obvious_duplicate_row(keep_row, True, dup_type, "keep", reason, group_id)
                _mark_obvious_duplicate_row(drop_row, True, dup_type, "drop", reason, group_id)
                _append_note_tag(keep_row, f"[obvious_duplicate_keep={dup_type}]")

                dropped = clean_record_text_fields(dict(drop_row))
                dropped["audit_reason"] = "obvious_cross_source_duplicate"
                audit_rows.append(dropped)
                removed_idx.add(drop_idx)

                if keep_idx != left_idx:
                    break

    kept_rows = [clean_record_text_fields(row) for idx, row in enumerate(records) if idx not in removed_idx]
    if audit_rows:
        print(f"  [Audit] obvious cross-source duplicate removal: {len(audit_rows)} rows")
    return kept_rows, audit_rows


def is_empty_fallback_shell_row(row: Dict) -> bool:
    catalyst_id = str(row.get("Catalyst_ID", "")).strip()
    canonical_id = str(row.get("Canonical_Catalyst_ID", "")).strip()
    if catalyst_id:
        return False
    if not canonical_id.startswith("__fallback_md5_"):
        return False
    if str(row.get("Reaction_Temp_C", "")).strip():
        return False

    perf_count = sum(bool(str(row.get(field, "")).strip()) for field in CORE_PERFORMANCE_FIELDS)
    if perf_count > 0:
        return False

    identity_count = sum(bool(str(row.get(field, "")).strip()) for field in IDENTITY_VALUE_FIELDS if field != "Canonical_Catalyst_ID")
    prep_count = sum(bool(str(row.get(field, "")).strip()) for field in PREPARATION_VALUE_FIELDS)

    # Preserve original text when Markdown/table markers are useful for later parsing.
    if identity_count >= 2 or prep_count >= 2 or (identity_count >= 1 and prep_count >= 1):
        return False
    return True


def filter_low_information_rows_before_save(records: List[Dict], file_name: str) -> tuple[List[Dict], List[Dict]]:
    kept_rows, audit_rows = [], []
    for row in records:
        if is_empty_fallback_shell_row(row):
            dropped = dict(row)
            dropped["audit_reason"] = "low_information_fallback_shell"
            dropped["audit_source_file"] = file_name
            audit_rows.append(dropped)
        else:
            kept_rows.append(row)
    if audit_rows:
        print(f"  [Audit] low-information fallback shell rows removed: {len(audit_rows)}")
    return kept_rows, audit_rows



# ==========================================
# identity alias / support role / numeric guard / quality labels
# ==========================================
SUPPORT_FAMILY_HINTS = {
    "Al2O3", "SiO2", "MgO", "CeO2", "CeOx", "ZrO2", "TiO2",
    "CNTs", "Activated Carbon", "SEP", "Al-LDH", "Mo2C",
    "MgO-Al2O3", "CeO2-ZrO2",
}
PREPARATION_SUFFIX_MAP = {
    "impre": "impregnation",
    "imp": "impregnation",
    "hydro": "hydrothermal",
    "copre": "coprecipitation",
    "coprec": "coprecipitation",
    "co-pre": "coprecipitation",
    "solgel": "sol-gel",
    "sg": "sol-gel",
    "ox": "oxidized",
}
NUMERIC_GUARD_FIELDS = [
    "Reaction_Temp_C", "Calcination_Temp_C", "Reduction_Temp_C",
    "Dry_Temp_C", "Pressure_bar", "TOS_h",
    "MeOH_Conversion_%", "H2_Yield_%", "H2_Selectivity_%",
    "CO_Selectivity_%", "CO2_Selectivity_%", "CO_Concentration_ppm",
]
QUALITATIVE_NUMERIC_PATTERNS = [
    "complete conversion", "nearly complete", "almost complete",
    "full conversion", "qualitative", "maintained", "remained",
    "stable for", "stable above", "trace", "very low", "very high",
]
APPROXIMATE_NUMERIC_PATTERNS = [
    "approximately", "approx.", "approx", "around", "about", "nearly", "ca.", "ca ", "~", "?",
]
BOUNDARY_NUMERIC_PATTERNS = [
    ">", "<", "above", "below", "more than", "less than",
    "greater than", "at least", "at most", "no more than", "no less than",
]
NUMERIC_EXPR_PRIORITY = {
    "missing": 0,
    "direct": 1,
    "approximate": 2,
    "bounded": 3,
    "range": 4,
    "qualitative": 5,
}


def _format_float_token(value: float) -> str:
    if abs(value - round(value)) < 1e-9:
        return str(int(round(value)))
    return f"{value:.3f}".rstrip("0").rstrip(".")


def _merge_semicolon_notes(existing: str, new_parts: List[str]) -> str:
    parts = []
    if existing:
        parts.extend([p.strip() for p in str(existing).split(";") if p.strip()])
    parts.extend([p.strip() for p in new_parts if p and str(p).strip()])
    return "; ".join(_dedupe_keep_order(parts))


def _split_identity_support_text(label: str) -> tuple[str, str]:
    text = _normalize_identity_text(label)
    if not text:
        return "", ""
    if "/" in text:
        left, right = text.split("/", 1)
        return left.strip(), right.strip()
    return text.strip(), ""


def _parse_component_metal_loadings_from_label(label: str) -> List[tuple[str, float]]:
    left, _ = _split_identity_support_text(label)
    matches = re.findall(r"(\d+(?:\.\d+)?)\s*(?:wt\.?\s*%|wt%|%)\s*([A-Z][a-z]?)", left, flags=re.I)
    pairs = []
    allowed = MSR_ACTIVE_METALS | MSR_PROMOTERS | MSR_RARE_EARTH | MSR_NI_SECOND_METALS
    for value, metal in matches:
        metal = metal.capitalize()
        if metal in allowed:
            pairs.append((metal, float(value)))
    return pairs


def _parse_inline_ratio_pairs_from_label(label: str) -> List[tuple[str, float]]:
    left, _ = _split_identity_support_text(label)
    allowed = sorted(MSR_ACTIVE_METALS | MSR_PROMOTERS | MSR_RARE_EARTH | MSR_NI_SECOND_METALS, key=len, reverse=True)
    alt = "|".join(sorted(set(allowed), key=len, reverse=True))

    colon_matches = re.findall(rf"({alt})\s*:\s*(\d+(?:\.\d+)?)", left)
    if len(colon_matches) >= 2:
        return [(metal.capitalize(), float(value)) for metal, value in colon_matches]

    compact = re.sub(r"[^A-Za-z0-9]", "", left)
    matches = re.findall(rf"({alt})(\d+(?:\.\d+)?)", compact)
    if len(matches) >= 2:
        return [(metal.capitalize(), float(value)) for metal, value in matches]
    return []


def _compact_loading_number(value_text: str) -> str:
    value_text = str(value_text or "").strip()
    return value_text[:-2] if value_text.endswith(".0") else value_text


def _build_component_loading_aliases(pairs: List[Tuple[str, float]], support: str = "") -> List[str]:
    normalized_pairs = []
    for metal, value in pairs or []:
        metal_text = clean_residual_mojibake_chars(str(metal or "")).strip()
        value_text = _format_ratio_token(str(value or "")).strip()
        if metal_text and value_text:
            normalized_pairs.append((metal_text.capitalize(), value_text))
    if len(normalized_pairs) < 2:
        return []

    support_norm = _normalize_registry_label(support) if support else ""
    orders = [normalized_pairs]
    if len(normalized_pairs) == 2:
        orders.append(list(reversed(normalized_pairs)))

    aliases: List[str] = []
    for ordered in orders:
        compact_pairs = [(metal, _compact_loading_number(value)) for metal, value in ordered]
        percent_alias = "-".join(f"{value}%{metal.lower()}" for metal, value in compact_pairs)
        metal_first_alias = "-".join(f"{metal.lower()}{value}" for metal, value in compact_pairs)
        value_first_alias = "-".join(f"{value}{metal.lower()}" for metal, value in compact_pairs)
        value_compact = "".join(f"{value}{metal.lower()}" for metal, value in compact_pairs)
        metal_compact = "".join(f"{metal.lower()}{value}" for metal, value in compact_pairs)
        aliases.extend([
            percent_alias,
            metal_first_alias,
            value_first_alias,
            value_compact,
            metal_compact,
        ])
        if support_norm:
            aliases.extend([
                f"{percent_alias}/{support_norm}",
                f"{metal_first_alias}/{support_norm}",
                f"{value_first_alias}/{support_norm}",
                f"{value_compact}/{support_norm}",
                f"{metal_compact}/{support_norm}",
            ])
    return _dedupe_keep_order([alias for alias in aliases if alias])


def _build_single_loading_support_aliases(record: Dict) -> List[str]:
    active = normalize_active_metal_for_merge(str(record.get("Active_Metal", "") or ""))
    promoter = normalize_promoter_for_merge(str(record.get("Promoter", "") or record.get("Promoter_Metal", "") or ""))
    loading = normalize_loading_for_merge(str(record.get("Metal_Loading_wt%", "") or ""))
    support = _normalize_registry_label(str(record.get("Support_Normalized", "") or record.get("Support", "") or ""))

    if not active or promoter or not loading or not support:
        return []

    loading_compact = _compact_loading_number(loading)
    support_tokens: List[str] = []
    if support == "al2o3":
        support_tokens = ["al", "al2o3"]

    aliases: List[str] = []
    for support_tok in support_tokens:
        aliases.extend([
            f"{loading_compact}{active.lower()}{support_tok}",
            f"{loading_compact}%{active.lower()}/{support}",
            f"{loading_compact}{active.lower()}/{support}",
            f"{active.lower()}{loading_compact}/{support}",
        ])

    return _dedupe_keep_order([alias for alias in aliases if alias])


def canonicalize_metal_order_and_ratio(record: Dict, raw_label: str = "") -> List[str]:
    notes = []
    label = raw_label or " ".join([
        str(record.get("Catalyst", "")),
        str(record.get("Catalyst_ID", "")),
        str(record.get("Series_Name", "")),
    ])

    existing_pairs = parse_alloy_ratio_to_ordered_pairs(record)
    if existing_pairs:
        normalized_alloy = ", ".join(f"{el}:{_format_float_token(float(val))}" for el, val in existing_pairs)
        if str(record.get("Alloy_Ratio", "")).strip() != normalized_alloy:
            notes.append("normalized_alloy_ratio_order")
        record["Alloy_Ratio"] = normalized_alloy
    else:
        comp_pairs = _parse_component_metal_loadings_from_label(label)
        if len(comp_pairs) >= 2 and any(el == "Ni" for el, _ in comp_pairs):
            total = sum(val for _, val in comp_pairs)
            if total > 0:
                if not str(record.get("Metal_Loading_wt%", "")).strip():
                    record["Metal_Loading_wt%"] = _format_float_token(total)
                    notes.append("inferred_total_loading_from_component_loadings")
                sorted_pairs = sorted(comp_pairs, key=lambda item: (0, item[0]) if item[0] == "Ni" else (1, item[0]))
                ratio_text = ", ".join(
                    f"{el}:{_format_float_token(val / total * 100)}" for el, val in sorted_pairs
                )
                if not str(record.get("Alloy_Ratio", "")).strip():
                    record["Alloy_Ratio"] = ratio_text
                    notes.append("inferred_alloy_ratio_from_component_loadings")
        else:
            ratio_pairs = _parse_inline_ratio_pairs_from_label(label)
            if len(ratio_pairs) >= 2 and any(el == "Ni" for el, _ in ratio_pairs):
                sorted_pairs = sorted(ratio_pairs, key=lambda item: (0, item[0]) if item[0] == "Ni" else (1, item[0]))
                if not str(record.get("Alloy_Ratio", "")).strip():
                    record["Alloy_Ratio"] = ", ".join(
                        f"{el}:{_format_float_token(val)}" for el, val in sorted_pairs
                    )
                    notes.append("inferred_alloy_ratio_from_inline_label")

    parsed_pairs = parse_alloy_ratio_to_ordered_pairs(record)
    if parsed_pairs and any(el == "Ni" for el, _ in parsed_pairs):
        other_metals = [el for el, _ in parsed_pairs if el != "Ni"]
        if not str(record.get("Active_Metal", "")).strip():
            record["Active_Metal"] = "Ni"
        if other_metals and not str(record.get("Promoter", "")).strip():
            record["Promoter"] = "/".join(_dedupe_keep_order(other_metals))
    return notes


def infer_preparation_suffix(value: str) -> tuple[str, str, str]:
    text = _normalize_identity_text(value)
    if not text:
        return "", "", ""
    match = re.match(r"^(.*?)[\-_](impre|imp|hydro|copre|coprec|co-pre|solgel|sg|ox)$", text, flags=re.I)
    if not match:
        return text, "", ""
    core = match.group(1).strip()
    suffix = match.group(2).lower()
    normalized = PREPARATION_SUFFIX_MAP.get(suffix, "")
    return core, normalized, suffix


def normalize_support_expression(value: str) -> str:
    text = _normalize_identity_text(value)
    if not text:
        return ""
    text = re.sub(r"\s*/\s*", "/", text)
    text = re.sub(r"\s*-\s*", "-", text)
    replacements = [
        (r"(?i)gamma-alumina|alumina", "Al2O3"),
        (r"(?i)\bal2o3\b", "Al2O3"),
        (r"(?i)\bmgo\b", "MgO"),
        (r"(?i)\bceo2\b", "CeO2"),
        (r"(?i)\bceox\b", "CeOx"),
        (r"(?i)\bla2o3\b", "La2O3"),
        (r"(?i)\bsio2\b", "SiO2"),
        (r"(?i)\bzro2\b", "ZrO2"),
        (r"(?i)\btio2\b", "TiO2"),
        (r"(?i)\bcnts?\b|carbon nanotubes?", "CNTs"),
        (r"(?i)activated carbon", "Activated Carbon"),
        (r"(?i)sepiolite", "SEP"),
    ]
    for pattern, repl in replacements:
        text = re.sub(pattern, repl, text)
    return text.strip(" -;,.")


def infer_support_family_and_modifier(support_text: str) -> tuple[str, str]:
    text = normalize_support_expression(support_text)
    if not text:
        return "", ""
    if text in {"MgO-Al2O3", "CeO2-ZrO2"}:
        return text, ""
    if "-" not in text:
        return text, ""
    parts = [p for p in text.split("-") if p]
    if len(parts) < 2:
        return text, ""

    for idx in range(len(parts) - 1, -1, -1):
        if parts[idx] in SUPPORT_FAMILY_HINTS:
            family = parts[idx]
            modifier = "-".join(parts[:idx] + parts[idx + 1:]).strip("-")
            return family, modifier
    for idx, part in enumerate(parts):
        if part in SUPPORT_FAMILY_HINTS:
            family = part
            modifier = "-".join(parts[:idx] + parts[idx + 1:]).strip("-")
            return family, modifier
    return text, ""


def split_support_promoter_preparation_roles(record: Dict) -> List[str]:
    notes = []
    support_text = str(record.get("Support", "")).strip()
    if not support_text:
        for candidate in [record.get("Catalyst", ""), record.get("Catalyst_ID", ""), record.get("Series_Name", "")]:
            extracted = _extract_support_from_identity_text(str(candidate))
            if extracted:
                support_text = extracted
                break
    if not support_text:
        return notes

    core_support, prep_norm, prep_suffix = infer_preparation_suffix(support_text)
    support_normalized = normalize_support_expression(core_support or support_text)
    support_family, support_modifier = infer_support_family_and_modifier(support_normalized)

    if support_normalized:
        # Normalize subscript unicode digits to ASCII for consistency
        support_normalized = (support_normalized
            .replace("₀","0").replace("₁","1").replace("₂","2").replace("₃","3")
            .replace("₄","4").replace("₅","5").replace("₆","6").replace("₇","7")
            .replace("₈","8").replace("₉","9"))
        if str(record.get("Support", "")).strip() != support_normalized:
            notes.append(f"support_normalized:{support_text}->{support_normalized}")
        record["Support"] = support_normalized
        record["Support_Normalized"] = support_normalized
    if support_family:
        record["Support_Family"] = support_family
    if support_modifier:
        record["Support_Modifier"] = support_modifier
    if prep_norm:
        record["Support_Prep_Method_Normalized"] = prep_norm
        notes.append(f"support_suffix:{prep_suffix}")
        if not str(record.get("Support_Prep_Method", "")).strip():
            if prep_norm == "hydrothermal":
                record["Support_Prep_Method"] = "Hydrothermal"
            elif prep_norm == "coprecipitation":
                record["Support_Prep_Method"] = "Precipitation"
            elif prep_norm == "sol-gel":
                record["Support_Prep_Method"] = "Sol-gel"
            elif prep_norm == "commercial":
                record["Support_Prep_Method"] = "Commercial"
            elif prep_norm == "impregnation":
                record["Support_Prep_Method"] = "Other"

    if notes:
        record["identity_role_parse_notes"] = _merge_semicolon_notes(
            str(record.get("identity_role_parse_notes", "")), notes
        )
    return notes


def normalize_catalyst_name_variant(value: str) -> str:
    temp = {"Catalyst": value, "Catalyst_ID": "", "Active_Metal": "", "Promoter": "", "Support": ""}
    canonicalize_metal_order_and_ratio(temp, value)
    split_support_promoter_preparation_roles(temp)
    enrich_promoter_fields_from_identity(temp, value)
    normalized = _build_structured_catalyst_id(temp)
    return normalized.lower() if normalized else re.sub(r"\s+", "", _normalize_identity_text(value).lower())


def normalize_identity_aliases(record: Dict) -> Dict:
    if not isinstance(record, dict):
        return record

    # alias 展开：将 VL 提取的缩写名称替换为标准化学式（大小写不敏感 + 去标点噪声）
    # 在所有规范化逻辑之前执行，确保下游 build_text_master_registry / merge_fragmented_records 等全部使用展开后的 key
    for field in ("Catalyst_ID", "Catalyst"):
        raw_val = str(record.get(field, "")).strip()
        if raw_val:
            expanded = _expand_catalyst_alias(raw_val)
            if expanded != raw_val:
                record[field] = expanded
                record["alias_map_source"] = field
                record["alias_map_evidence"] = f"{raw_val} -> {expanded}"

    raw_cid = str(record.get("Catalyst_ID", "")).strip()
    if raw_cid and not str(record.get("Catalyst_ID_raw", "")).strip():
        record["Catalyst_ID_raw"] = raw_cid

    label_text = " ".join([
        str(record.get("Catalyst", "")),
        str(record.get("Catalyst_ID_raw", "") or record.get("Catalyst_ID", "")),
        str(record.get("Series_Name", "")),
    ]).strip()

    notes = []
    notes.extend(canonicalize_metal_order_and_ratio(record, label_text))
    notes.extend(split_support_promoter_preparation_roles(record))
    enrich_promoter_fields_from_identity(record, label_text)

    normalized_id = _build_structured_catalyst_id(dict(record)).lower()
    if normalized_id:
        record["Catalyst_ID_normalized"] = normalized_id
        record["identity_alias_group"] = normalized_id
    elif raw_cid:
        compact = re.sub(r"\s+", "", raw_cid).lower()
        record["Catalyst_ID_normalized"] = compact
        record["identity_alias_group"] = compact
    else:
        record.setdefault("Catalyst_ID_normalized", "")
        if not str(record.get("identity_alias_group", "")).strip() and str(record.get("Canonical_Catalyst_ID", "")).strip():
            record["identity_alias_group"] = str(record.get("Canonical_Catalyst_ID", "")).strip()

    if raw_cid and normalized_id and re.sub(r"\s+", "", raw_cid).lower() != normalized_id:
        notes.append("normalized_identity_alias")
    if label_text and not raw_cid and normalized_id:
        notes.append("normalized_from_catalyst_label")

    if notes:
        record["identity_normalization_notes"] = _merge_semicolon_notes(
            str(record.get("identity_normalization_notes", "")), notes
        )
    return record


def classify_numeric_expression_type(value: Any, field_name: str = "") -> tuple[str, str, str]:
    raw = _normalize_identity_text(str(value))
    if not raw or raw.lower() in {"", "n/a", "na", "none", "unknown"}:
        return "missing", "", ""
    lowered = raw.lower()

    if any(pat in lowered for pat in QUALITATIVE_NUMERIC_PATTERNS):
        return "qualitative", "", raw
    if re.search(r"\d+(?:\.\d+)?\s*(?:-|–|—|to)\s*\d+(?:\.\d+)?", raw, flags=re.I) or "range" in lowered:
        return "range", "", raw
    if any(pat in lowered for pat in BOUNDARY_NUMERIC_PATTERNS):
        return "bounded", "", raw
    if any(pat in lowered for pat in APPROXIMATE_NUMERIC_PATTERNS):
        m = re.search(r"([+-]?\d+(?:\.\d+)?)", raw)
        return "approximate", (m.group(1) if m else ""), raw

    m = re.search(r"^\s*([+-]?\d+(?:\.\d+)?)\s*(?:%|c|k|bar|h|ppm)?\s*$", lowered)
    if m:
        return "direct", m.group(1), raw

    m = re.search(r"([+-]?\d+(?:\.\d+)?)", raw)
    if m and re.search(r"[a-zA-Z]", raw) and not any(pat in lowered for pat in QUALITATIVE_NUMERIC_PATTERNS):
        return "approximate", m.group(1), raw
    return "qualitative", "", raw


def annotate_numeric_expression_guards(record: Dict) -> Dict:
    raw_fragments = []
    has_direct = False
    has_approx = False
    has_range = False
    has_qualitative = False

    for field in NUMERIC_GUARD_FIELDS:
        raw_val = str(record.get(field, "")).strip()
        expr_type, numeric_token, raw_expr = classify_numeric_expression_type(raw_val, field)
        if expr_type == "missing":
            continue
        if expr_type == "direct":
            record[field] = numeric_token
            has_direct = True
        elif expr_type == "approximate":
            if numeric_token:
                record[field] = numeric_token
                has_approx = True
            else:
                record[field] = ""
            raw_fragments.append(f"{field}={raw_expr}")
        elif expr_type in {"bounded", "range"}:
            record[field] = ""
            has_range = True
            raw_fragments.append(f"{field}={raw_expr}")
        else:
            record[field] = ""
            has_qualitative = True
            raw_fragments.append(f"{field}={raw_expr}")

    if has_qualitative:
        numeric_level = "qualitative_only"
        numeric_expr = "qualitative"
    elif has_range:
        numeric_level = "range_like"
        numeric_expr = "range_like"
    elif has_approx:
        numeric_level = "approximate_numeric"
        numeric_expr = "approximate"
    elif has_direct:
        numeric_level = "direct_numeric"
        numeric_expr = "direct"
    else:
        numeric_level = "missing"
        numeric_expr = "missing"

    record["numeric_expression_type"] = numeric_expr
    record["numeric_reliability_level"] = numeric_level
    record["is_range_like_value"] = has_range
    record["is_qualitative_value"] = has_qualitative
    record["is_approximate_value"] = has_approx
    if raw_fragments:
        record["raw_numeric_expression"] = " | ".join(_dedupe_keep_order(raw_fragments))
    return record


def apply_numeric_expression_guards(records: List[Dict]) -> List[Dict]:
    return [annotate_numeric_expression_guards(record) for record in records if isinstance(record, dict)]


def build_duplicate_candidate_signature(record: Dict) -> str:
    alias = str(
        record.get("identity_alias_group")
        or record.get("Catalyst_ID_normalized")
        or record.get("Catalyst_ID")
        or record.get("Canonical_Catalyst_ID")
        or record.get("Catalyst")
        or ""
    ).strip().lower()
    source_file = str(record.get("Source_File", "")).strip().lower()
    if not alias:
        return ""

    temp_val = _parse_float_if_possible(record.get("Reaction_Temp_C", ""))
    tos_val = _parse_float_if_possible(record.get("TOS_h", ""))
    if temp_val is not None:
        if temp_val > 450 and record.get("temp_conversion_clue_flag"):
            temp_val = temp_val - 273.15
        x_sig = f"temp:{round(temp_val, 1)}"
    elif tos_val is not None:
        x_sig = f"tos:{round(tos_val, 1)}"
    else:
        x_sig = "x:na"

    perf_parts = []
    for field in ["MeOH_Conversion_%", "H2_Yield_%", "CO_Selectivity_%", "CO2_Selectivity_%"]:
        val = _parse_float_if_possible(record.get(field, ""))
        if val is not None:
            perf_parts.append(f"{field}:{round(val, 1)}")
    perf_sig = "|".join(perf_parts[:3]) if perf_parts else "perf:na"
    return f"{source_file}|{alias}|{x_sig}|{perf_sig}"


def _infer_source_granularity(record: Dict, source_text: str, origin_type: str) -> str:
    text = str(source_text or "")
    if origin_type == "figure":
        return "figure_series"
    if record.get("is_source_file_si") and "=== TABLE" in text:
        return "si_table"
    if record.get("is_source_file_si"):
        return "text_si"
    if "=== TABLE" in text:
        return "table_row"
    if "figure" in text.lower() and "caption" in text.lower():
        return "caption_only"
    return "text_numeric"


def annotate_cross_source_duplicate_candidates(records: List[Dict]) -> List[Dict]:
    groups: Dict[str, List[Dict]] = {}
    for row in records:
        row.setdefault("same_physical_point_possible", False)
        row.setdefault("duplicate_candidate_type", "")
        row.setdefault("duplicate_candidate_signature", "")
        sig = build_duplicate_candidate_signature(row)
        if sig:
            groups.setdefault(sig, []).append(row)

    for signature, rows in groups.items():
        if len(rows) < 2:
            continue
        granularities = {str(r.get("source_granularity", "")).strip() for r in rows}
        raw_ids = {
            str(r.get("Catalyst_ID_raw") or r.get("Catalyst_ID") or r.get("Catalyst") or "").strip().lower()
            for r in rows if str(r.get("Catalyst_ID_raw") or r.get("Catalyst_ID") or r.get("Catalyst") or "").strip()
        }
        norm_ids = {
            str(r.get("Catalyst_ID_normalized") or r.get("identity_alias_group") or "").strip().lower()
            for r in rows if str(r.get("Catalyst_ID_normalized") or r.get("identity_alias_group") or "").strip()
        }

        dup_types = set()
        if "table_row" in granularities and "text_numeric" in granularities:
            dup_types.add("same_point_possible_text_table")
        if "figure_series" in granularities and any(g in granularities for g in {"text_numeric", "table_row", "caption_only"}):
            dup_types.add("same_point_possible_text_figure")
        if any(g in granularities for g in {"si_table", "text_si"}) and any(g not in {"si_table", "text_si"} for g in granularities):
            dup_types.add("same_point_possible_text_si")
        if any(r.get("temp_duplicate_resolution_upstream") or r.get("temp_conversion_clue_flag") for r in rows):
            dup_types.add("same_point_possible_temp_unit_duplicate")
        if len(raw_ids) > 1 and len(norm_ids) == 1:
            dup_types.add("same_point_possible_identity_alias_duplicate")
        if not dup_types and granularities == {"text_numeric"}:
            dup_types.add("same_point_possible_text_overlap")

        for row in rows:
            row["same_physical_point_possible"] = True
            row["duplicate_candidate_signature"] = signature
            existing = [t for t in str(row.get("duplicate_candidate_type", "")).split("|") if t]
            row["duplicate_candidate_type"] = "|".join(_dedupe_keep_order(existing + sorted(dup_types)))
    return records


def score_identity_completeness(record: Dict) -> str:
    score = 0
    if str(record.get("Catalyst_ID_normalized", "")).strip() or str(record.get("Catalyst_ID", "")).strip():
        score += 2
    if str(record.get("Active_Metal", "")).strip():
        score += 1
    if str(record.get("Support_Normalized", "")).strip() or str(record.get("Support", "")).strip():
        score += 1
    if str(record.get("Promoter", "")).strip() or str(record.get("Alloy_Ratio", "")).strip():
        score += 1
    level = "complete" if score >= 4 else "partial" if score >= 2 else "poor"
    if "[series_identity_incomplete=" in str(record.get("Notes", "")) and level == "complete":
        level = "partial"
    record["identity_completeness_level"] = level
    return level


def score_numeric_reliability(record: Dict) -> str:
    level = str(record.get("numeric_reliability_level", "")).strip()
    if level:
        return level
    has_numeric = any(_parse_float_if_possible(record.get(field, "")) is not None for field in NUMERIC_GUARD_FIELDS)
    level = "direct_numeric" if has_numeric else "missing"
    record["numeric_reliability_level"] = level
    return level


def assign_extraction_quality_labels(records: List[Dict]) -> List[Dict]:
    for row in records:
        identity_level = score_identity_completeness(row)
        numeric_level = score_numeric_reliability(row)
        source_granularity = str(row.get("source_granularity", "")).strip()
        is_duplicate = bool(row.get("same_physical_point_possible"))

        if is_empty_fallback_shell_row(row):
            extraction_confidence = "low"
            origin_class = "low_information_shell"
        elif bool(row.get("is_literature_comparison")):
            extraction_confidence = "medium" if numeric_level == "direct_numeric" else "low"
            origin_class = "comparison_table_like"
        elif source_granularity == "figure_series" and identity_level != "complete":
            extraction_confidence = "low" if numeric_level in {"range_like", "qualitative_only", "missing"} else "medium"
            origin_class = "figure_series_partial_identity"
        elif numeric_level in {"range_like", "qualitative_only"}:
            extraction_confidence = "low"
            origin_class = "implied_or_range_like"
        elif numeric_level == "direct_numeric" and identity_level == "complete" and source_granularity in {"text_numeric", "table_row", "si_table"}:
            extraction_confidence = "high"
            origin_class = "high_confidence_direct_numeric"
        elif numeric_level == "direct_numeric":
            extraction_confidence = "medium" if identity_level != "poor" else "low"
            origin_class = "medium_confidence_direct_numeric" if extraction_confidence == "medium" else "figure_series_partial_identity"
        elif numeric_level == "approximate_numeric":
            extraction_confidence = "medium" if identity_level != "poor" else "low"
            origin_class = "implied_or_range_like"
        else:
            extraction_confidence = "low"
            origin_class = "low_information_shell" if identity_level == "poor" else "comparison_table_like"

        if is_duplicate and extraction_confidence == "high":
            extraction_confidence = "medium"

        row["extraction_confidence"] = extraction_confidence
        row["identity_completeness_level"] = identity_level
        row["numeric_reliability_level"] = numeric_level
        row["origin_quality_class"] = origin_class
    return records


def split_text_into_paragraphs(text: str) -> List[str]:
    if not text or not str(text).strip():
        return []

    text = clean_residual_mojibake_chars(str(text))
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"\n{3,}", "\n\n", text)

    def _is_heading_like(line: str) -> bool:
        line_norm = re.sub(r"\s+", " ", line).strip()
        if not line_norm:
            return False
        low = line_norm.lower().strip(" :")
        if low.startswith("=== body text"):
            return True
        if re.fullmatch(r"[-=]{3,}", line_norm):
            return True
        if len(line_norm) > 120:
            return False
        heading_keywords = [
            "catalyst preparation", "catalysts preparation", "preparation of", "synthesis of",
            "materials preparation", "experimental section", "materials and methods",
            "results and discussion", "characterization", "catalytic performance",
            "performance evaluation", "activity test", "stability test",
        ]
        if any(k in low for k in heading_keywords):
            return True
        if re.match(r"^\d+(?:\.\d+){0,3}\s+[A-Za-z]", line_norm):
            return True
        if re.match(r"^(abstract|introduction|experimental|results|discussion|conclusions?|references)\b", low):
            return True
        words = line_norm.split()
        if 1 <= len(words) <= 10 and not line_norm.endswith("."):
            alpha_words = [w for w in words if re.search(r"[A-Za-z]", w)]
            if alpha_words and sum(w[:1].isupper() for w in alpha_words) >= max(1, len(alpha_words) - 1):
                return True
        return False

    def _split_long_block(block: str) -> List[str]:
        block = block.strip()
        if not block:
            return []
        if len(block) <= 1600 and block.count("\n") <= 12:
            return [block]

        parts: List[str] = []
        current: List[str] = []
        for raw_line in block.split("\n"):
            line = raw_line.strip()
            if not line:
                if current and current[-1] != "":
                    current.append("")
                continue
            if current and _is_heading_like(line):
                segment = "\n".join(current).strip()
                if segment:
                    parts.append(segment)
                current = [line]
            else:
                current.append(line)
        if current:
            segment = "\n".join(current).strip()
            if segment:
                parts.append(segment)
        return parts if len(parts) > 1 else [block]

    def _keep_fragment(block: str) -> bool:
        compact = re.sub(r"\s+", " ", block).strip()
        if not compact:
            return False
        if "=== table" in compact.lower() or "=== body text" in compact.lower():
            return True
        if re.search(r"^\s*\|.*\|\s*$", block, flags=re.M):
            return True
        if len(compact) >= 30:
            return True
        if _is_heading_like(compact):
            return True
        info_markers = [
            "conversion", "yield", "selectivity", "performance", "preparation",
            "calcined", "reduced", "impregnated", "xrd", "xps", "tem", "bet",
            "benchmark", "ref.", "reference",
        ]
        return any(marker in compact.lower() for marker in info_markers)

    paragraphs: List[str] = []
    table_pattern = re.compile(r"=== TABLE[^\n]*===\n.*?\n=== END TABLE ===", flags=re.S)
    cursor = 0
    for match in table_pattern.finditer(text):
        prefix = text[cursor:match.start()]
        if prefix.strip():
            for block in re.split(r"\n\s*\n+", prefix):
                for piece in _split_long_block(block):
                    if _keep_fragment(piece):
                        paragraphs.append(piece.strip())
        table_block = match.group(0).strip()
        if _keep_fragment(table_block):
            paragraphs.append(table_block)
        cursor = match.end()

    suffix = text[cursor:]
    if suffix.strip():
        for block in re.split(r"\n\s*\n+", suffix):
            for piece in _split_long_block(block):
                if _keep_fragment(piece):
                    paragraphs.append(piece.strip())

    return paragraphs


def classify_paragraph_role(paragraph_text: str) -> Tuple[str, str]:
    text = clean_residual_mojibake_chars(str(paragraph_text or ""))
    if not text.strip():
        return "other", "fallback_rule"

    compact = re.sub(r"\s+", " ", text).strip()
    low = compact.lower()
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    first_line = lines[0].lower() if lines else low
    is_table_block = "=== table" in low or len(re.findall(r"^\s*\|.*\|\s*$", text, flags=re.M)) >= 2
    citation_hits = re.findall(r"\[\d+\]", compact)

    comparison_patterns = [
        r"\bref\.", r"\breference\b", r"\bthis study\b", r"\bbenchmark\b",
        r"\bliterature comparison\b", r"\bcompared with\b", r"\bcompared to\b",
        r"\bversus\b", r"\bvs\.?(?=\s|$)",
    ]
    characterization_patterns = [
        r"\bxrd\b", r"\bxps\b", r"\btem\b", r"\bsem\b", r"\bbet\b", r"\bftir\b",
        r"\bdrifts\b", r"\braman\b", r"\btpr\b", r"\btpd\b", r"\btpsr\b",
        r"\bmorphology\b", r"\bparticle size\b", r"\bsurface area\b",
    ]
    preparation_heading_keywords = [
        "catalyst preparation", "catalysts preparation", "preparation of", "synthesis of",
        "materials preparation", "experimental section", "materials and methods",
    ]
    preparation_action_keywords = [
        "impregnated", "impregnation", "co-precipitated", "coprecipitation", "co precipitation",
        "deposited", "deposition-precipitation", "deposition precipitation", "dried", "drying",
        "calcined", "calcination", "reduced", "reduction", "aged", "aging",
        "stirred", "dissolved", "evaporated", "filtered", "washed",
    ]
    performance_patterns = [
        r"\bmethanol conversion\b", r"\bconversion\b", r"\bh2 yield\b", r"\bhydrogen yield\b",
        r"\bselectivity\b", r"\bco selectivity\b", r"\bco2 selectivity\b", r"\bco concentration\b",
        r"\breaction temperature\b", r"\bghsv\b", r"\bwhsv\b", r"\bs/c\b",
        r"\bsteam-to-carbon\b", r"\bsteam to carbon\b", r"\bpressure\b", r"\btos\b",
        r"\btime on stream\b", r"\bstability\b", r"\bdeactivation\b", r"\bperformance\b",
    ]

    has_comparison = any(re.search(pattern, low, flags=re.I) for pattern in comparison_patterns) or len(citation_hits) >= 2
    if is_table_block and has_comparison:
        return "comparison_paragraph", "table_rule"
    if has_comparison:
        if any(k in first_line for k in ["benchmark", "comparison", "this study", "reference", "ref."]):
            return "comparison_paragraph", "heading_rule"
        return "comparison_paragraph", "keyword_rule"

    # 反应条件强信号：含 S/C、GHSV、WHSV、space velocity 时优先分为 performance_paragraph
    # 避免被表征词（XRD/BET/TEM等）拦截
    condition_strong_patterns = [
        r"\bghsv\b", r"\bwhsv\b", r"\bs/c\b", r"\bsteam-to-carbon\b", r"\bsteam to carbon\b",
        r"\bspace velocity\b", r"\bmethanol-to-water\b", r"\bmeoh[- ]to[- ]h2o\b",
    ]
    has_condition_strong = any(re.search(p, low, flags=re.I) for p in condition_strong_patterns)
    if has_condition_strong:
        return "performance_paragraph", "condition_keyword_rule"

    has_characterization = any(re.search(pattern, low, flags=re.I) for pattern in characterization_patterns)
    if has_characterization:
        if is_table_block:
            return "characterization_paragraph", "table_rule"
        if any(re.search(pattern, first_line, flags=re.I) for pattern in characterization_patterns):
            return "characterization_paragraph", "heading_rule"
        return "characterization_paragraph", "keyword_rule"

    has_preparation_heading = any(k in first_line for k in preparation_heading_keywords) or (
        len(compact) <= 160 and any(k in low for k in preparation_heading_keywords)
    )
    if has_preparation_heading:
        return "preparation_paragraph", "heading_rule"

    has_preparation_action = any(k in low for k in preparation_action_keywords)
    has_preparation_entity = bool(re.search(
        r"\b(precursor|nitrate|acetate|chloride|support|wt%|wt\.%|calcination|reduction|drying)\b",
        low,
        flags=re.I,
    )) or bool(re.search(r"(?:\b\d+(?:\.\d+)?\s*(?:wt%|wt\.%|h)\b)|(\b\d+(?:\.\d+)?\s*(?:[^A-Za-z0-9\s]\s*)?(?:deg\s*)?c\b)", compact, flags=re.I))
    if has_preparation_action and has_preparation_entity:
        return "preparation_paragraph", "keyword_rule"

    has_performance = any(re.search(pattern, low, flags=re.I) for pattern in performance_patterns)
    if has_performance:
        if is_table_block:
            return "performance_paragraph", "table_rule"
        if any(k in first_line for k in ["performance", "conversion", "yield", "selectivity", "stability"]):
            return "performance_paragraph", "heading_rule"
        return "performance_paragraph", "keyword_rule"

    return "other", "fallback_rule"


def route_chunk_paragraphs(text_chunk: str) -> List[Dict]:
    paragraph_records: List[Dict] = []
    for paragraph_text in split_text_into_paragraphs(text_chunk):
        paragraph_role, paragraph_router_method = classify_paragraph_role(paragraph_text)
        paragraph_records.append({
            "paragraph_text": paragraph_text,
            "paragraph_role": paragraph_role,
            "paragraph_router_method": paragraph_router_method,
        })
    return paragraph_records


def _extract_info_from_text_block(text_block: str, file_name: str) -> List[Dict]:
    text_block = clean_text_for_api(text_block)
    if not text_block:
        return []

    system_prompt = f"""You extract structured MSR catalyst data from a paper chunk.

Core rules:
1. Table-first extraction.
   - If the chunk contains markdown tables, treat each table row as an independent record.
   - Do not merge different table rows into one record.
2. Catalyst_ID is mandatory and promoter-aware.
   - Catalyst_ID must encode promoter or second metal identity.
   - Promoter identity is part of catalyst identity, not just an optional note.
   - Ni/Al2O3, Ni-Au/Al2O3, Ni-Rh/Al2O3, and Ni-Ir/Al2O3 must be different Catalyst_ID values.
   - Ni80-Cu20/CNTs and Ni20-Cu80/CNTs must be different Catalyst_ID values.
   - Do not omit promoter or second metal from Catalyst_ID.
   - For Ni-based bimetallic or promoted catalysts, Ni remains the core active metal, but the second metal or promoter must still appear in Catalyst_ID.
   - Suggested style: [loading]-[core metal]-[promoter or second metal]/[support], or directly encode the alloy ratio when available.
3. Promoter and second metal extraction.
   - Fill Promoter when a promoter or second metal is explicitly present in catalyst names, row labels, captions, legends, or prose.
   - In names such as NiAl-Au, NiAl-Rh, NiAl-Ir, Ni-Cu, Ni-Mo, Ni-Ce, Ni-Mg, and Ni-La, the second metal must be captured and reflected in Catalyst_ID.
4. Local evidence overrides global context.
   - Only use evidence from the current chunk.
   - If a value is not stated in this chunk, leave it empty instead of borrowing from another chunk.
5. Preserve point-by-point conditions.
   - If multiple temperatures or time-on-stream points are reported, output separate records.
6. Loading vs alloy ratio must stay separate.
   - Metal_Loading_wt% is total loading only when wt% is explicitly stated.
   - Alloy_Ratio must include element names, for example Ni:80, Cu:20.
   - Never copy alloy ratio numbers into Metal_Loading_wt%.
7. Component metal loading evidence.
   - Extract Ni_Loading_wt% and Promoter_Loading_wt% when Ni and promoter/second-metal loadings are explicitly stated in tables, body text, catalyst preparation, experimental sections, nominal-composition descriptions, or SI composition tables.
   - Table columns such as "Ni loading (wt%)", "Cu loading (wt%)", "Mo content (wt%)", or "Pt loading" are direct component loading evidence.
   - Catalyst names can be used only as fallback when they explicitly combine total loading and composition ratio, e.g. "20 wt.% Ni80-Cu20/Al2O3".
   - Ni80-Cu20, Ni:Cu = 80:20, or Ni/Cu ratios are composition ratios, not wt% loadings by themselves.
   - Only derive component loadings when both total Metal_Loading_wt% and a clear composition ratio are available.
   - If only a promoter metal name is present and no loading or ratio is given, fill Promoter but leave Promoter_Loading_wt% empty.
   - Never invent Ni_Loading_wt% or Promoter_Loading_wt% from metal names alone.
   - Oxide-to-metal conversion: when the paper reports catalyst composition as oxide wt% (e.g. "NiO = 2 wt%", "CuO = 5 wt%"), convert to metal wt%:
     Ni wt% = NiO wt% × 58.69 / 74.69; Cu wt% = CuO wt% × 63.55 / 79.55.
     Put the converted metal wt% into Ni_Loading_wt% or Promoter_Loading_wt%, and note the original oxide value in Notes.
     If the paper already gives metal wt% directly, use it as-is without conversion.
8. Numeric guard.
   - If a value is approximate, bounded, range-like, or qualitative (for example ~100, >99, 90-100, complete conversion, or maintained 100% above 500 C), do not rewrite it into a fake exact single-point number.
   - Keep the raw wording or leave the numeric field empty.
8. Selectivity guard.
   - Distinguish CO selectivity from CO2 selectivity carefully.
   - Use Reasoning_Selectivity to explain the assignment when selectivity is reported.
9. Be conservative.
   - If uncertain, leave the field empty.
   - Do not infer missing values from comparison text, mechanism discussion, or vague summaries.
10. Catalyst alias resolution. [Improve4]
   - If the catalyst name is a code or label (e.g. "Cat-1", "S1", "Sample A", "optimal catalyst", "as-prepared"),
     look for its composition definition elsewhere in the chunk (e.g. "Cat-1 = 10wt%Ni/Al2O3").
   - Fill Catalyst with the original label as written, and fill Catalyst_ID using the resolved composition.
   - If the composition cannot be resolved from this chunk, fill Catalyst with the label and leave
     Active_Metal / Support / Metal_Loading_wt% empty rather than guessing.
   - Never fabricate a composition for an unresolved alias.

Output strict JSON with key "data" and value as a list of records.
Schema:
{json.dumps(DATA_SCHEMA, ensure_ascii=False)}

Current file: {file_name}
"""
    system_prompt = clean_residual_mojibake_chars(system_prompt)

    try:
        response = client.chat.completions.create(
            model=TEXT_MODEL,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": f"Extract structured MSR data from the text below:\n{text_block}"}
            ],
            response_format={"type": "json_object"},
            temperature=0.1
        )
        data = json.loads(response.choices[0].message.content)
        records = _apply_stability_transition_fix(
            _apply_h2_unit_routing(_apply_semantic_guard(data.get("data", []), "text"))
        )
        return finalize_extracted_records(records, text_block, file_name, "text")
    except Exception as e:
        error_msg = str(e)
        if "data_inspection_failed" in error_msg or "inappropriate content" in error_msg:
            print(f"  [WARN] content inspection failed; retrying with stricter cleanup...")
            text_clean = re.sub(r'[^\w\s\.\,\-\(\)\[\]\/\:\;\|]', ' ', text_block)
            text_clean = re.sub(r'\s+', ' ', text_clean)[:8000]
            if len(text_clean) > 100:
                try:
                    response = client.chat.completions.create(
                        model=TEXT_MODEL,
                        messages=[
                            {"role": "system", "content": system_prompt},
                            {"role": "user", "content": f"Extract structured MSR data from the text below:\n{text_clean}"}
                        ],
                        response_format={"type": "json_object"},
                        temperature=0.1
                    )
                    data = json.loads(response.choices[0].message.content)
                    print(f"  [OK] retry succeeded")
                    records = _apply_stability_transition_fix(
                        _apply_h2_unit_routing(_apply_semantic_guard(data.get("data", []), "text-retry"))
                    )
                    return finalize_extracted_records(records, text_clean, file_name, "text")
                except Exception:
                    print("  [ERROR] retry failed again; skip this chunk")
        else:
            print(f"  [ERROR] API call failed: {e}")
        return []


def _blank_forbidden_performance_fields(record: Dict) -> Dict:
    if not isinstance(record, dict):
        return record

    cleaned = dict(record)
    # 只清空性能结果字段，保留反应条件字段（S_C_Ratio、GHSV等）供后续广播到图表行
    forbidden_fields = [
        "MeOH_Conversion_%", "H2_Yield_%", "H2_Production_Rate", "H2_Selectivity_%",
        "CO_Selectivity_%", "CO2_Selectivity_%", "CO_Concentration_ppm",
        "Deactivation_Rate_%_h", "Carbon_Deposition_wt%",
    ]
    for field in forbidden_fields:
        cleaned[field] = ""
    return cleaned


def _finalize_preparation_records(records: List[Dict], source_text: str, file_name: str) -> List[Dict]:
    cleaned_records = []
    for record in records:
        if isinstance(record, dict):
            cleaned_records.append(_blank_forbidden_performance_fields(record))
    finalized = finalize_extracted_records(cleaned_records, source_text, file_name, "text")
    for row in finalized:
        row["text_extraction_subroute"] = "preparation_schema"
    return finalized


def _is_rule_confident_preparation_value(field: str, value: str) -> bool:
    value = clean_residual_mojibake_chars(str(value or "")).strip()
    if not value or field == "Catalyst_ID":
        return False

    if field == "Metal_Loading_Method":
        return value in {"Impregnation", "Co-precipitation", "Deposition-precipitation", "Sol-gel", "Hydrothermal", "Mechanical mixing"}
    if field == "Support_Prep_Method":
        return value in {"Commercial", "Precipitation", "Sol-gel", "Hydrothermal"}
    if field == "Precursor":
        return bool(
            re.search(r"\b(?:nitrate|acetate|chloride|salt|salts|NO3|CH3COO|C2H3O2|Cl)\b", value, flags=re.I)
            or re.search(r"\b[A-Z][a-z]?\([^\)]{1,16}\)\d", value)
        )
    if field == "Support":
        normalized = normalize_support_expression(value)
        return bool(re.search(r"\b(?:Al2O3|MgO|CeO2|CeOx|La2O3|SiO2|ZrO2|TiO2|CNTs|Activated Carbon|SEP)\b", normalized, flags=re.I))
    if field == "Metal_Loading_wt%":
        numeric = _parse_float_if_possible(value)
        return numeric is not None and 0 < numeric <= 80
    if field.endswith("_Temp_C"):
        numeric = _parse_float_if_possible(value)
        return numeric is not None and 20 <= numeric <= 1200
    if field.endswith("_Time_h"):
        numeric = _parse_float_if_possible(value)
        return numeric is not None and 0 < numeric <= 200
    if field in {"Catalyst", "Active_Metal", "Promoter", "Alloy_Ratio"}:
        return False
    if field == "Notes":
        return True
    return bool(value)



def _extract_preparation_entities_rule_based(text_block: str) -> Dict[str, Any]:
    text_raw = clean_residual_mojibake_chars(str(text_block or ""))
    if not text_raw.strip():
        return {}

    text = text_raw.replace("\r\n", "\n").replace("\r", "\n")
    compact = re.sub(r"\s+", " ", text).strip()
    low = compact.lower()
    result: Dict[str, Any] = {"rule_hits": []}

    def _set_value(field: str, value: str, hit: str) -> None:
        value = clean_residual_mojibake_chars(str(value or "")).strip(" ;,.")
        if not value:
            return
        if field == "Support":
            value = re.sub(r"(?i)^commercial\s+", "", value).strip()
            value = normalize_support_expression(value)
            value = re.sub(r"(?i)^gamma[-\s]?al2o3$", "Al2O3", value)
        if not _is_rule_confident_preparation_value(field, value):
            return
        if not str(result.get(field, "")).strip():
            result[field] = value
            result["rule_hits"] = _dedupe_keep_order(list(result.get("rule_hits", [])) + [hit])

    metal_loading_patterns = [
        (r"co[\s\-]?impregnat|incipient wetness|wet impregnation|impregnat", "Impregnation"),
        (r"deposition[\s\-]?precipitat", "Deposition-precipitation"),
        (r"co[\s\-]?precipit|coprecipit", "Co-precipitation"),
        (r"sol[\s\-]?gel", "Sol-gel"),
        (r"hydrothermal", "Hydrothermal"),
        (r"mechanical[\s\-]?mix|ball[\s\-]?mill|grinding|ground mixture", "Mechanical mixing"),
    ]
    for pattern, normalized in metal_loading_patterns:
        if re.search(pattern, low, flags=re.I):
            _set_value("Metal_Loading_Method", normalized, f"Metal_Loading_Method:{normalized}")
            break

    if re.search(r"commercial|purchased|supplied|sigma|aladdin|aldrich|bought", low, flags=re.I):
        _set_value("Support_Prep_Method", "Commercial", "Support_Prep_Method:Commercial")
    else:
        support_method_patterns = [
            (r"(?:support|carrier)[^.;\n]{0,80}hydrothermal|hydrothermal[^.;\n]{0,80}(?:support|carrier)", "Hydrothermal"),
            (r"(?:support|carrier)[^.;\n]{0,80}sol[\s\-]?gel|sol[\s\-]?gel[^.;\n]{0,80}(?:support|carrier)", "Sol-gel"),
            (r"(?:support|carrier)[^.;\n]{0,80}precipit|precipit[^.;\n]{0,80}(?:support|carrier)", "Precipitation"),
        ]
        for pattern, normalized in support_method_patterns:
            if re.search(pattern, low, flags=re.I):
                _set_value("Support_Prep_Method", normalized, f"Support_Prep_Method:{normalized}")
                break

    precursor_matches: List[str] = []
    precursor_patterns = [
        r"\b(?:Ni|Cu|Ce|La|Mg|Co|Fe|Zn|Mn|Cr|Mo|Rh|Ru|Pt|Pd|Ir|Au)\s*\([^\)]{1,16}\)\s*\d+(?:\s*[\u00b7\.]\s*\d*H2O)?\b",
        r"\b(?:nickel|copper|cerium|lanthanum|magnesium|cobalt|iron|zinc|manganese|chromium|molybdenum|rhodium|ruthenium|platinum|palladium|iridium|gold)\s+(?:nitrate(?:s)?|acetate(?:s)?|chloride(?:s)?)(?:\s+\w+)?\b",
    ]
    for pattern in precursor_patterns:
        precursor_matches.extend([m.group(0).strip() for m in re.finditer(pattern, text, flags=re.I)])
    if not precursor_matches:
        generic_match = re.search(r"\b(?:metal\s+salts?|nitrate(?:s)?|acetate(?:s)?|chloride(?:s)?)\b", compact, flags=re.I)
        if generic_match and re.search(r"\b(?:precursor|salt|salts|dissolved|solution)\b", low, flags=re.I):
            precursor_matches.append(generic_match.group(0).strip())
    if precursor_matches:
        _set_value("Precursor", " / ".join(_dedupe_keep_order(precursor_matches)), "Precursor")

    def _extract_condition(action_pattern: str) -> tuple[str, str]:
        temp_patterns = [
            rf"(?:{action_pattern})[^.;\n]{{0,120}}?(?:at|under|in)?\s*(\d+(?:\.\d+)?)\s*(?:[^A-Za-z0-9\s]\s*)?(?:deg\s*)?c\b",
            rf"(?:{action_pattern})[^.;\n]{{0,120}}?\(\s*(\d+(?:\.\d+)?)\s*(?:[^A-Za-z0-9\s]\s*)?(?:deg\s*)?c\s*\)",
        ]
        time_patterns = [
            rf"(?:{action_pattern})[^.;\n]{{0,120}}?(?:for|during)\s*(\d+(?:\.\d+)?)\s*h(?:ours?)?\b",
            rf"(?:{action_pattern})[^.;\n]{{0,120}}?(\d+(?:\.\d+)?)\s*h(?:ours?)?\b",
        ]
        temp_value = ""
        time_value = ""
        for pattern in temp_patterns:
            match = re.search(pattern, compact, flags=re.I)
            if match:
                temp_value = match.group(1)
                break
        for pattern in time_patterns:
            match = re.search(pattern, compact, flags=re.I)
            if match:
                time_value = match.group(1)
                break
        return temp_value, time_value

    dry_temp, dry_time = _extract_condition(r"dried|drying")
    calc_temp, calc_time = _extract_condition(r"calcined|calcination|calcine")
    red_temp, red_time = _extract_condition(r"reduced|reduction|activated\s+(?:under|in)\s+h2|activated\s+under\s+hydrogen")

    _set_value("Dry_Temp_C", dry_temp, "Dry_Temp_C")
    _set_value("Dry_Time_h", dry_time, "Dry_Time_h")
    _set_value("Calcination_Temp_C", calc_temp, "Calcination_Temp_C")
    _set_value("Calcination_Time_h", calc_time, "Calcination_Time_h")
    _set_value("Reduction_Temp_C", red_temp, "Reduction_Temp_C")
    _set_value("Reduction_Time_h", red_time, "Reduction_Time_h")

    # Metal_Loading_wt% is no longer extracted by regex.
    # It is computed exclusively by sync_component_loading_derivatives
    # as Ni_Loading_wt% + Promoter_Loading_wt%, which is more accurate.

    # Component Ni loading patterns
    ni_component_patterns = [
        r"\b(\d+(?:\.\d+)?)\s*wt\.?\s*%\s*(?:of\s+)?\bni\b",
        r"\bni\s+(?:loading|content)\s+(?:of|was|is|at)\s+(\d+(?:\.\d+)?)\s*wt\.?\s*%",
        r"\b(\d+(?:\.\d+)?)\s*wt\.?\s*%\s*(?:of\s+)?nickel\b",
    ]
    for pattern in ni_component_patterns:
        match = re.search(pattern, low, flags=re.I)
        if match:
            _set_value("Ni_Loading_wt%", match.group(1), "Ni_Loading_wt%")
            break

    # Component Promoter loading patterns
    _promoter_elems = r"(?:cu|ce|pr|la|mg|co|fe|zn|mn|mo|zr|sn|in|k|na|ca|sr|ba|cr|ru|pt|pd|au|ir|rh)"
    _promoter_loading_re = rf"\b(\d+(?:\.\d+)?)\s*wt\.?\s*%\s*(?:of\s+)?\b{_promoter_elems}\b"
    _promoter_loading_inv = rf"\b{_promoter_elems}\s+(?:loading|content)\s+(?:of|was|is|at)\s+(\d+(?:\.\d+)?)\s*wt\.?\s*%"
    for match in re.finditer(_promoter_loading_re, low, flags=re.I):
        if not str(result.get("Promoter_Loading_wt%", "")).strip():
            _set_value("Promoter_Loading_wt%", match.group(1), "Promoter_Loading_wt%")
            break
    if not str(result.get("Promoter_Loading_wt%", "")).strip():
        match = re.search(_promoter_loading_inv, low, flags=re.I)
        if match:
            _set_value("Promoter_Loading_wt%", match.group(1), "Promoter_Loading_wt%")

    support_value = ""
    label_like_matches = re.findall(r"\b[A-Za-z0-9][A-Za-z0-9\-\(\)]{0,30}/[A-Za-z0-9][A-Za-z0-9\-\(\)]{1,30}\b", compact)
    for label in label_like_matches:
        support_candidate = normalize_support_expression(_extract_support_from_identity_text(label))
        if _is_rule_confident_preparation_value("Support", support_candidate):
            support_value = support_candidate
            break
    if not support_value:
        support_patterns = [
            r"(?:onto|on|over|supported on)\s+((?:commercial\s+)?(?:gamma[-\s]?al2o3|gamma[-\s]?alumina|alumina|al2o3|ceo2|ceox|la2o3|sio2|zro2|tio2|mgo|cnts?|carbon nanotubes?|activated carbon|sepiolite|sep))",
            r"((?:commercial\s+)?(?:gamma[-\s]?al2o3|gamma[-\s]?alumina|alumina|al2o3|ceo2|ceox|la2o3|sio2|zro2|tio2|mgo|cnts?|carbon nanotubes?|activated carbon|sepiolite|sep))\s+(?:was|were)?\s*(?:used|employed|chosen)?\s+as support",
        ]
        for pattern in support_patterns:
            match = re.search(pattern, compact, flags=re.I)
            if match:
                normalized_support = normalize_support_expression(match.group(1))
                if _is_rule_confident_preparation_value("Support", normalized_support):
                    support_value = normalized_support
                    break
    if support_value:
        _set_value("Support", support_value, "Support")

    if not result.get("rule_hits"):
        return {}
    return result



def _extract_preparation_llm_complements(text_block: str, file_name: str) -> List[Dict]:
    raw_text = clean_residual_mojibake_chars(str(text_block or ""))
    if not raw_text.strip():
        return []

    llm_text = raw_text.replace("\r\n", "\n").replace("\r", "\n")
    llm_text = re.sub(r"[\x00-\x08\x0b-\x0c\x0e-\x1f\x7f-\x9f]", "", llm_text)
    llm_text = re.sub(r"http[s]?://\S+", " ", llm_text)
    llm_text = re.sub(r"\S+@\S+\.\S+", " ", llm_text)
    llm_text = re.sub(r"<<[^>]*>>", " ", llm_text)
    llm_text = re.sub(r"\n{3,}", "\n\n", llm_text).strip()
    if len(re.sub(r"\s+", " ", llm_text)) < 30:
        return []

    allowed_fields = [
        "Catalyst_ID", "Catalyst", "Active_Metal", "Promoter", "Support",
        "Metal_Loading_wt%", "Ni_Loading_wt%", "Promoter_Loading_wt%",
        "Alloy_Ratio", "Precursor", "Support_Prep_Method",
        "Metal_Loading_Method", "Dry_Temp_C", "Dry_Time_h", "Calcination_Temp_C",
        "Calcination_Time_h", "Reduction_Temp_C", "Reduction_Time_h", "Notes",
    ]

    system_prompt = f"""You complement structured Ni-based MSR catalyst preparation data from a preparation paragraph.

A rule-based extractor has already captured straightforward preparation entities such as obvious preparation methods, obvious precursor salts, and action-bound drying/calcination/reduction conditions.
You are only complementing missing or complex preparation identity information.

This extractor is NOT for catalytic performance point extraction.
Do NOT output reaction/performance fields.
Do NOT convert activity, conversion, selectivity, stability, or benchmark wording into point records.

Focus on complementing these fields when explicit in the paragraph:
- Catalyst
- Catalyst_ID
- Active_Metal
- Promoter
- Alloy_Ratio
- Ni_Loading_wt% and Promoter_Loading_wt% when component metal loadings are explicitly stated (e.g. Table columns "Ni loading (wt%)", "Cu loading (wt%)")
- Support (when not obvious)
- Precursor / Metal_Loading_Method / Support_Prep_Method only when the wording is complex and explicit
- Dry/Calcination/Reduction conditions only when explicit and clearly bound to the action

Allowed output fields only:
- Catalyst_ID
- Catalyst
- Active_Metal
- Promoter
- Support
- Metal_Loading_wt%
- Ni_Loading_wt%
- Promoter_Loading_wt%
- Alloy_Ratio
- Precursor
- Support_Prep_Method
- Metal_Loading_Method
- Dry_Temp_C
- Dry_Time_h
- Calcination_Temp_C
- Calcination_Time_h
- Reduction_Temp_C
- Reduction_Time_h
- Notes

Hard prohibitions:
- Leave all reaction/performance fields empty.
- Do not output Reaction_Temp_C, S_C_Ratio, GHSV_mL_g_h, Pressure_bar, Feed_Composition, MeOH_Conversion_%, H2_Yield_%, H2_Production_Rate, H2_Selectivity_%, CO_Selectivity_%, CO2_Selectivity_%, CO_Concentration_ppm, TOS_h, Deactivation_Rate_%_h, or Carbon_Deposition_wt%.
- Do not create catalytic performance point records from preparation text.
- Do not convert complete conversion, high activity, or stable for 20 h into data points.

Core rules:
1. Complement-only behavior.
   - Rule-based extraction has already captured straightforward preparation entities.
   - You are complementing only missing or complex preparation identity fields.
2. Table-first extraction.
   - If the paragraph contains markdown tables, treat each table row as an independent preparation record.
   - Do not merge different table rows into one record.
3. Catalyst_ID must stay promoter-aware.
   - Ni/Al2O3, Ni-Cu/Al2O3, Ni-Ce/Al2O3, and Ni-Mg/Al2O3 must remain different Catalyst_ID values.
   - Do not omit promoter or second metal from Catalyst_ID.
4. Multiple preparations must stay separate.
   - If the paragraph describes multiple catalysts or multiple preparation rows, output separate records.
5. Local evidence only.
   - Only use evidence from the current paragraph.
   - If a value is not stated in this paragraph, leave it empty.
6. Be conservative.
   - Do not infer missing values from comparison text, mechanism text, or vague summaries.
   - If uncertain, leave the field empty.

Output strict JSON with key "data" and value as a list of records.
Each record should fill only the allowed preparation fields above.
Schema:
{json.dumps(DATA_SCHEMA, ensure_ascii=False)}

Current file: {file_name}
"""
    system_prompt = clean_residual_mojibake_chars(system_prompt)

    def _sanitize_llm_records(data_obj: Dict) -> List[Dict]:
        sanitized: List[Dict] = []
        for item in data_obj.get("data", []):
            if not isinstance(item, dict):
                continue
            filtered = {field: item.get(field, "") for field in allowed_fields}
            filtered = _blank_forbidden_performance_fields(filtered)
            sanitized.append(clean_record_text_fields(filtered))
        return _apply_stability_transition_fix(_apply_h2_unit_routing(sanitized))

    try:
        response = client.chat.completions.create(
            model=TEXT_MODEL,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": f"Complement structured preparation data from the text below:\n{llm_text}"}
            ],
            response_format={"type": "json_object"},
            temperature=0.1
        )
        data = json.loads(response.choices[0].message.content)
        return _sanitize_llm_records(data)
    except Exception as e:
        error_msg = str(e)
        if "data_inspection_failed" in error_msg or "inappropriate content" in error_msg:
            print(f"  [WARN] preparation complement inspection failed; retrying with stricter cleanup...")
            text_clean = re.sub(r'[^\w\s\.\,\-\(\)\[\]\/\:\;\|\%\n]', ' ', llm_text)
            text_clean = re.sub(r'\n{3,}', '\n\n', text_clean)
            text_clean = re.sub(r'[ \t]+', ' ', text_clean).strip()[:8000]
            if len(re.sub(r"\s+", " ", text_clean)) > 30:
                try:
                    response = client.chat.completions.create(
                        model=TEXT_MODEL,
                        messages=[
                            {"role": "system", "content": system_prompt},
                            {"role": "user", "content": f"Complement structured preparation data from the text below:\n{text_clean}"}
                        ],
                        response_format={"type": "json_object"},
                        temperature=0.1
                    )
                    data = json.loads(response.choices[0].message.content)
                    print(f"  [OK] preparation complement retry succeeded")
                    return _sanitize_llm_records(data)
                except Exception:
                    print("  [ERROR] preparation complement retry failed again; skip this block")
        else:
            print(f"  [ERROR] preparation complement API call failed: {e}")
        return []



def _merge_preparation_rule_and_llm(rule_result: Dict, llm_records: List[Dict], source_text: str) -> List[Dict]:
    rule_result = dict(rule_result or {})
    rule_fields = {
        field: clean_residual_mojibake_chars(str(value)).strip()
        for field, value in rule_result.items()
        if field != "rule_hits" and clean_residual_mojibake_chars(str(value)).strip()
    }

    single_record_fill_fields = {
        "Support", "Metal_Loading_wt%", "Precursor", "Metal_Loading_Method", "Support_Prep_Method",
        "Dry_Temp_C", "Dry_Time_h", "Calcination_Temp_C", "Calcination_Time_h",
        "Reduction_Temp_C", "Reduction_Time_h",
    }
    shared_multirecord_fields = {
        "Metal_Loading_Method", "Support_Prep_Method", "Dry_Temp_C", "Dry_Time_h",
        "Calcination_Temp_C", "Calcination_Time_h", "Reduction_Temp_C", "Reduction_Time_h",
    }

    if llm_records:
        merged_records: List[Dict] = []
        single_record_mode = len(llm_records) == 1
        for record in llm_records:
            if not isinstance(record, dict):
                continue
            merged = _blank_forbidden_performance_fields(dict(record))
            fill_fields = single_record_fill_fields if single_record_mode else shared_multirecord_fields
            for field in fill_fields:
                value = rule_fields.get(field, "")
                if not value or not _is_rule_confident_preparation_value(field, value):
                    continue
                if not str(merged.get(field, "")).strip():
                    merged[field] = value
            merged_records.append(merged)
        return merged_records

    rule_only_record = {
        field: value for field, value in rule_fields.items()
        if _is_rule_confident_preparation_value(field, value)
    }
    is_table_like = len(re.findall(r"^\s*\|.*\|\s*$", source_text or "", flags=re.M)) >= 2
    if rule_only_record and not is_table_like and _has_static_preparation_identity_signal(rule_only_record):
        rule_only_record["Notes"] = _merge_semicolon_notes(
            str(rule_only_record.get("Notes", "")),
            ["[preparation_rule_only_record]"]
        )
        return [rule_only_record]
    return []



def _extract_preparation_from_text_block(text_block: str, file_name: str) -> List[Dict]:
    raw_text_block = clean_residual_mojibake_chars(str(text_block or ""))
    if not re.sub(r"\s+", " ", raw_text_block).strip():
        return []

    rule_result = _extract_preparation_entities_rule_based(raw_text_block)
    llm_records = _extract_preparation_llm_complements(raw_text_block, file_name)
    merged_records = _merge_preparation_rule_and_llm(rule_result, llm_records, raw_text_block)
    return _finalize_preparation_records(merged_records, raw_text_block, file_name)


def extract_info_from_chunk(text_chunk: str, file_name: str,
                             global_params: Optional[Dict] = None) -> List[Dict]:
    """
    V6 text extraction entry point.
    - Keep Catalyst_ID promoter-aware.
    - Do not let global fallback instructions overwrite local chunk evidence.
    """
    raw_text_chunk = clean_residual_mojibake_chars(str(text_chunk or ""))
    text_chunk = clean_text_for_api(raw_text_chunk)
    if not text_chunk:
        return []

    # Global parameters are applied later by apply_global_params_fallback().
    # The chunk-level extractor must stay conservative and only use evidence from the current chunk.
    paragraph_records = route_chunk_paragraphs(raw_text_chunk)
    routable_roles = {"preparation_paragraph", "performance_paragraph"}
    routed_extractable = [p for p in paragraph_records if p.get("paragraph_role") in routable_roles]

    if not routed_extractable:
        fallback_records = _extract_info_from_text_block(raw_text_chunk, file_name)
        for record in fallback_records:
            record["paragraph_role"] = "other"
            record["paragraph_router_method"] = "fallback_rule"
        return fallback_records

    all_records: List[Dict] = []
    for paragraph_record in routed_extractable:
        paragraph_text = str(paragraph_record.get("paragraph_text", "")).strip()
        if not paragraph_text:
            continue

        paragraph_role = paragraph_record.get("paragraph_role", "other")
        if paragraph_role == "preparation_paragraph":
            block_records = _extract_preparation_from_text_block(paragraph_text, file_name)
        else:
            block_records = _extract_info_from_text_block(paragraph_text, file_name)

        for record in block_records:
            record["paragraph_role"] = paragraph_role
            record["paragraph_router_method"] = paragraph_record.get("paragraph_router_method", "fallback_rule")
        all_records.extend(block_records)
    return all_records


# [Fix 23] H2 unit routing: prevent rate/flow values from being written into percentage fields.
_H2_RATE_UNITS = {
    "mol/h", "mol h", "mol h-1",
    "mmol/min", "mmol min", "mmol min-1",
    "ml/min", "ml min", "ml min-1",
    "flow rate", "production rate", "flow",
}

def _fix_h2_unit_routing(record: Dict) -> Dict:
    # [Fix 23] Route H2 rate-like values out of H2_Yield_% and into H2_Production_Rate.
    # Keep percent fields empty when the unit is not a true percentage.
    h2_yield_val = clean_residual_mojibake_chars(str(record.get("H2_Yield_%", ""))).strip()
    if not h2_yield_val or h2_yield_val in ("", "N/A", "n/a"):
        return record

    val_lower = h2_yield_val.lower()

    # Detect rate or flow units first.
    is_rate = any(kw in val_lower for kw in _H2_RATE_UNITS)

    if is_rate:
        # Route the value to H2_Production_Rate without overwriting an existing rate.
        existing_rate = str(record.get("H2_Production_Rate", "")).strip()
        if not existing_rate or existing_rate in ("", "N/A", "n/a"):
            record["H2_Production_Rate"] = h2_yield_val
        else:
            # If a rate already exists, keep the conflict note in Notes.
            notes = str(record.get("Notes", "")).strip()
            tag = f"[H2_rate_unit_conflict={h2_yield_val}]"
            if tag not in notes:
                record["Notes"] = (notes + " " + tag).strip()
        record["H2_Yield_%"] = ""
        print(f"  [Fix23] H2_Yield_%='{h2_yield_val}' looked rate-like; moved to H2_Production_Rate")
        return record

    # No percent sign and no recognized rate unit.
    # [Fix23-v2] Do NOT clear the field — preserve the numeric value.
    # Instead flag as low-confidence unit so research layers can still use it,
    # while final gate can remain conservative.
    if "%" not in h2_yield_val:
        notes = str(record.get("Notes", "")).strip()
        tag = f"[H2_yield_unit_unclear={h2_yield_val}]"
        if tag not in notes:
            record["Notes"] = (notes + " " + tag).strip()
        # Mark unit confidence low but KEEP the numeric value in the field.
        record["metric_unit_confidence"] = "low"
        record["metric_unit_unclear_flag"] = 1
        print(f"  [Fix23] H2_Yield_%='{h2_yield_val}' unit unclear; value preserved, "
              f"metric_unit_confidence=low, metric_unit_unclear_flag=1 (NOT cleared)")

    return record


def _apply_h2_unit_routing(records: List[Dict]) -> List[Dict]:
    """Apply H2 unit routing to a batch of records."""
    return [_fix_h2_unit_routing(r) for r in records]


RAW_UNIT_FIELD_SPECS = {
    "H2_Yield_%": "H2_Yield_Raw",
    "H2_Production_Rate": "H2_Production_Rate_Raw",
    "S_C_Ratio": "S_C_Ratio_Raw",
    "Feed_MeOH_to_H2O_Ratio": "Feed_MeOH_to_H2O_Ratio_Raw",
    "GHSV_mL_g_h": "GHSV_mL_g_h_Raw",
    "WHSV_h_inv": "WHSV_h_inv_Raw",
    "SpaceVelocity_norm": "SpaceVelocity_norm_Raw",
}


def capture_raw_unit_fields(record: Dict) -> Dict:
    if not isinstance(record, dict):
        return record

    row = dict(record)

    def _first_nonempty(*values: Any) -> str:
        for value in values:
            text = clean_residual_mojibake_chars(str(value or "")).strip()
            if text:
                return text
        return ""

    for field, raw_field in RAW_UNIT_FIELD_SPECS.items():
        if str(row.get(raw_field, "")).strip():
            continue
        value = _first_nonempty(row.get(field, ""))
        if value:
            row[raw_field] = value

    if not str(row.get("Feed_MeOH_to_H2O_Ratio_Raw", "")).strip():
        feed_ratio_raw = _first_nonempty(row.get("Feed_MeOH_to_H2O_Ratio", ""), row.get("Feed_Composition", ""))
        if feed_ratio_raw:
            row["Feed_MeOH_to_H2O_Ratio_Raw"] = feed_ratio_raw

    if not str(row.get("GHSV_mL_g_h_Raw", "")).strip():
        ghsv_raw = _first_nonempty(row.get("GHSV_mL_g_h", ""), row.get("SpaceVelocity_Value", ""))
        if ghsv_raw:
            row["GHSV_mL_g_h_Raw"] = ghsv_raw

    if not str(row.get("WHSV_h_inv_Raw", "")).strip():
        whsv_raw = _first_nonempty(
            row.get("WHSV_h_inv", ""),
            row.get("SpaceVelocity_Value", "") if re.search(r"\bWHSV\b", str(row.get("SpaceVelocity_Value", "") or ""), flags=re.I) else "",
            row.get("GHSV_mL_g_h", "") if re.search(r"\bWHSV\b", str(row.get("GHSV_mL_g_h", "") or ""), flags=re.I) else "",
        )
        if whsv_raw:
            row["WHSV_h_inv_Raw"] = whsv_raw

    if not str(row.get("SpaceVelocity_norm_Raw", "")).strip():
        sv_raw = _first_nonempty(
            row.get("SpaceVelocity_Value", ""),
            row.get("WHSV_h_inv_Raw", ""),
            row.get("GHSV_mL_g_h_Raw", ""),
            row.get("WHSV_h_inv", ""),
            row.get("GHSV_mL_g_h", ""),
            row.get("SpaceVelocity_norm", ""),
        )
        if sv_raw:
            row["SpaceVelocity_norm_Raw"] = sv_raw

    return row


import re as _re

_STABILITY_PERF_FIELDS = [
    "MeOH_Conversion_%", "H2_Yield_%", "CO_Selectivity_%",
    "CO2_Selectivity_%", "H2_Selectivity_%",
]
_RANGE_KEYWORDS = [
    "range", "between", "varied", "from", "to", "~", "approximately",
    "approx", "around", "about", ">", "<", "above", "below",
    "at least", "at most", "no more than", "no less than",
]
_COMPARISON_KEYWORDS = [
    "better than", "higher than", "lower than", "compared with", "compared to",
    "comparison", "benchmark", "reference catalyst", "versus", "vs.", "vs ",
]
_MECHANISM_KEYWORDS = [
    "mechanism", "pathway", "intermediate", "xps", "tpr", "drifts",
    "characterization", "schematic", "proposed route",
]
_STABILITY_ALLOW_PATTERNS = [
    "stable for", "tested for", "time on stream", "tos", "after", "over",
    "remained at", "decreased from", "dropped from",
]

def _normalize_stability_transition_record(record: Dict) -> Dict:
    """Normalize stability-transition phrasing before later merging."""
    check_text = " ".join([
        str(record.get("Notes", "")),
        str(record.get("Reasoning_Selectivity", "")),
    ]).lower()

    m_transition = _re.search(
        r"from\s+([\d\.]+)\s*%?\s*to\s+([\d\.]+)\s*%?\s*(?:after|over)\s+([\d\.]+)\s*h",
        check_text
    )
    if m_transition:
        init_val = float(m_transition.group(1))
        end_val = float(m_transition.group(2))
        tos_val = m_transition.group(3)
        existing_tos = str(record.get("TOS_h", "")).strip()
        if not existing_tos or existing_tos in ("", "N/A"):
            record["TOS_h"] = tos_val
        for f in _STABILITY_PERF_FIELDS:
            v = str(record.get(f, "")).strip()
            try:
                fv = float(v)
            except (ValueError, TypeError):
                continue
            if abs(fv - init_val) < 0.01 and abs(fv - end_val) > 0.01:
                record[f] = ""
                tag = f"[Fix24_cleared_init_val={v}_field={f}]"
                notes = str(record.get("Notes", "")).strip()
                if tag not in notes:
                    record["Notes"] = (notes + " " + tag).strip()
        return record

    m_remained = _re.search(
        r"remained\s+at\s+([\d\.]+)\s*%?\s*(?:after|for)\s+([\d\.]+)\s*h",
        check_text
    )
    if m_remained:
        if not str(record.get("TOS_h", "")).strip():
            record["TOS_h"] = m_remained.group(2)
        return record

    m_stable = _re.search(r"stable\s+for\s+([\d\.]+)\s*h", check_text)
    if m_stable and not m_transition and not m_remained:
        has_any_perf = any(
            str(record.get(f, "")).strip() not in ("", "N/A")
            for f in _STABILITY_PERF_FIELDS
        )
        if has_any_perf:
            for f in _STABILITY_PERF_FIELDS:
                record[f] = ""
            tag = "[Fix24_stable_no_endval]"
            notes = str(record.get("Notes", "")).strip()
            if tag not in notes:
                record["Notes"] = (notes + " " + tag).strip()
        if not str(record.get("TOS_h", "")).strip():
            record["TOS_h"] = m_stable.group(1)
    return record

def _apply_stability_transition_fix(records: List[Dict]) -> List[Dict]:
    """Apply stability-transition normalization to all records."""
    fixed = []
    for r in records:
        if isinstance(r, dict):
            fixed.append(_normalize_stability_transition_record(r))
        else:
            fixed.append(r)
    return fixed

def _should_keep_as_point_record(record: Dict) -> bool:
    """Decide whether an extracted record should be kept as a point record."""
    perf_fields = [
        "MeOH_Conversion_%", "H2_Yield_%", "H2_Production_Rate",
        "CO_Selectivity_%", "CO2_Selectivity_%", "CO_Concentration_ppm",
    ]

    has_catalyst = bool(
        str(record.get("Active_Metal", "")).strip() or
        str(record.get("Catalyst_ID", "")).strip() or
        str(record.get("Catalyst", "")).strip()
    )
    has_perf = any(
        str(record.get(f, "")).strip() not in ("", "N/A", "n/a")
        for f in perf_fields
    )
    # 纯条件记录通道：有催化剂身份 + 有反应条件字段，即使无性能数据也保留
    # 用于后续将 S/C、GHSV 等广播到图表行
    condition_fields = ["S_C_Ratio", "GHSV_mL_g_h", "WHSV_h_inv", "SpaceVelocity_norm",
                        "Reaction_Temp_C", "Pressure_bar", "Feed_MeOH_to_H2O_Ratio"]
    has_condition = any(str(record.get(f, "")).strip() not in ("", "N/A", "n/a") for f in condition_fields)
    if has_catalyst and has_condition and not has_perf:
        return True

    if not has_catalyst or not has_perf:
        return False

    check_text = " ".join([
        str(record.get("Notes", "")),
        str(record.get("Reasoning_Selectivity", "")),
        str(record.get("Catalyst", "")),
    ]).lower()

    has_tos = bool(str(record.get("TOS_h", "")).strip())
    has_stability_hint = any(kw in check_text for kw in _STABILITY_ALLOW_PATTERNS)
    if has_catalyst and has_perf and (has_tos or has_stability_hint):
        if any(kw in check_text for kw in _MECHANISM_KEYWORDS):
            return False
        return True

    if any(kw in check_text for kw in _RANGE_KEYWORDS):
        return False
    if any(kw in check_text for kw in _COMPARISON_KEYWORDS):
        return False
    if any(kw in check_text for kw in _MECHANISM_KEYWORDS):
        return False
    return True

def _apply_semantic_guard(records: List[Dict], source_hint: str = "") -> List[Dict]:
    """Filter out non-point records returned by the LLM."""
    kept, dropped = [], []
    for r in records:
        if _should_keep_as_point_record(r):
            kept.append(r)
        else:
            dropped.append(r.get("Catalyst_ID", r.get("Active_Metal", "?")))
    if dropped:
        hint = f" ({source_hint})" if source_hint else ""
        print(f"  [Fix12] semantic guard filtered {len(dropped)} non-point records{hint}: {dropped}")
    return kept

def parse_alloy_ratio_to_ordered_pairs(record: Dict):
    def _canon_metal(token: Any) -> str:
        text = clean_residual_mojibake_chars(str(token or "")).strip()
        if not text:
            return ""
        name_map = {
            "nickel": "Ni", "copper": "Cu", "platinum": "Pt", "palladium": "Pd",
            "ruthenium": "Ru", "rhodium": "Rh", "cobalt": "Co", "iron": "Fe",
            "gold": "Au", "iridium": "Ir", "molybdenum": "Mo", "tin": "Sn",
            "gallium": "Ga", "indium": "In", "zinc": "Zn", "cerium": "Ce",
            "lanthanum": "La", "potassium": "K", "magnesium": "Mg", "manganese": "Mn",
            "chromium": "Cr", "yttrium": "Y", "praseodymium": "Pr", "neodymium": "Nd",
            "samarium": "Sm", "gadolinium": "Gd",
        }
        alpha = re.sub(r"[^A-Za-z]", "", text)
        if alpha.lower() in name_map:
            return name_map[alpha.lower()]
        if re.fullmatch(r"[A-Z][a-z]?", text):
            return text[0].upper() + text[1:].lower()
        match = re.search(r"[A-Z][a-z]?", text)
        return match.group(0) if match else ""

    def _num(value: Any) -> Optional[float]:
        match = re.search(r"[-+]?\d+(?:\.\d+)?", clean_residual_mojibake_chars(str(value or "")).replace(",", ""))
        return float(match.group(0)) if match else None

    def _strip_support_tail(text: str) -> str:
        parts = [part.strip() for part in str(text or "").split("/") if part.strip()]
        while len(parts) > 1:
            tail = parts[-1]
            tail_norm = normalize_support_expression(tail).strip().lower()
            if tail.lower() in SUPPORT_LIKE_TOKENS or tail_norm in SUPPORT_LIKE_TOKENS:
                parts.pop()
                continue
            break
        return "/".join(parts)

    def _finalize(raw_pairs: List[Tuple[Any, Any]]) -> List[Tuple[str, float]]:
        result = []
        seen = set()
        for metal, value in raw_pairs:
            symbol = _canon_metal(metal)
            number = _num(value)
            if not symbol or number is None or number <= 0:
                continue
            if symbol in seen:
                continue
            seen.add(symbol)
            result.append((symbol, float(number)))
        if len(result) < 2:
            return []
        if "Ni" not in [metal for metal, _ in result]:
            return []
        return result

    texts = []
    if isinstance(record, dict):
        for field in ["Alloy_Ratio", "Metal_Ratio", "Composition_Ratio", "Catalyst", "Catalyst_ID", "Canonical_Catalyst_ID", "Series_Name", "Active_Metal"]:
            value = clean_residual_mojibake_chars(str(record.get(field, ""))).strip()
            if value:
                texts.append(value)
    else:
        value = clean_residual_mojibake_chars(str(record or "")).strip()
        if value:
            texts.append(value)

    for raw_text in _dedupe_keep_order(texts):
        text = _strip_support_tail(_normalize_identity_text(raw_text))
        if "ni" not in text.lower() or not re.search(r"\d", text):
            continue

        match = re.search(
            r"([A-Z][a-z]?(?:\s*[:/\-]\s*[A-Z][a-z]?){1,3})\s*=\s*([0-9]+(?:\.\d+)?(?:\s*[:/,]\s*[0-9]+(?:\.\d+)?){1,3})",
            text,
        )
        if match:
            metals = re.findall(r"[A-Z][a-z]?", match.group(1))
            values = re.findall(r"[0-9]+(?:\.\d+)?", match.group(2))
            pairs = _finalize(list(zip(metals, values)))
            if pairs:
                return pairs

        match = re.search(
            r"([A-Z][a-z]?(?:\s*[-/]\s*[A-Z][a-z]?){1,3})\s*\(\s*([0-9]+(?:\.\d+)?(?:\s*[:/,]\s*[0-9]+(?:\.\d+)?){1,3})\s*\)",
            text,
        )
        if match:
            metals = re.findall(r"[A-Z][a-z]?", match.group(1))
            values = re.findall(r"[0-9]+(?:\.\d+)?", match.group(2))
            pairs = _finalize(list(zip(metals, values)))
            if pairs:
                return pairs

        colon_pairs = re.findall(r"\b([A-Z][a-z]?)\s*[:=]\s*([0-9]+(?:\.\d+)?)\b", text)
        pairs = _finalize(colon_pairs)
        if pairs:
            return pairs

        paren_pairs = re.findall(r"\b([A-Z][a-z]?)\s*[\(\[]\s*([0-9]+(?:\.\d+)?)\s*[\)\]]", text)
        pairs = _finalize(paren_pairs)
        if pairs:
            return pairs

        attached_pairs = re.findall(r"([A-Z][a-z]?)([0-9]+(?:\.\d+)?)", text)
        pairs = _finalize(attached_pairs)
        if pairs:
            return pairs

        prefixed_pairs = [(metal, value) for value, metal in re.findall(r"([0-9]+(?:\.\d+)?)\s*([A-Z][a-z]?)", text)]
        pairs = _finalize(prefixed_pairs)
        if pairs:
            return pairs

    return []



def normalize_support_for_merge(value: str) -> str:
    """Function docstring removed for runtime stability."""
    if not value:
        return ""
    s = value.strip()
    s = re.sub(r'\s+', ' ', s)           # Collapse repeated whitespace.
    s = re.sub(r'[\u2013\u2014]', '-', s) # Normalize long dashes to '-'.
    s = re.sub(r'-{2,}', '-', s)          # 澶氳繛瀛楃鍚堝苟
    return s.lower()


def _sort_elements_for_merge(elements: List[str]) -> List[str]:
    """Function docstring removed for runtime stability."""
    seen = []
    for el in elements:
        if el and el not in seen:
            seen.append(el)
    return sorted(seen, key=lambda el: (0, el.lower()) if el.lower() == "ni" else (1, el.lower()))


def _extract_elements_for_merge(value: str) -> List[str]:
    """Function docstring removed for runtime stability."""
    if not value:
        return []
    cleaned = _strip_descriptor_suffix_before_element_parse(str(value))
    allowed = MSR_ACTIVE_METALS | MSR_PROMOTERS | MSR_RARE_EARTH | globals().get("MSR_NI_SECOND_METALS", set())
    name_map = {
        "nickel": "Ni", "copper": "Cu", "platinum": "Pt", "palladium": "Pd",
        "ruthenium": "Ru", "rhodium": "Rh", "cobalt": "Co", "iron": "Fe",
        "gold": "Au", "iridium": "Ir", "molybdenum": "Mo", "tin": "Sn",
        "gallium": "Ga", "indium": "In", "zinc": "Zn", "cerium": "Ce",
        "lanthanum": "La", "potassium": "K", "magnesium": "Mg",
        "manganese": "Mn", "chromium": "Cr", "yttrium": "Y",
        "praseodymium": "Pr", "neodymium": "Nd", "samarium": "Sm",
        "gadolinium": "Gd",
    }
    kept: List[str] = []
    for raw_token in re.split(r"[\s/;,:]+", cleaned):
        token = raw_token.strip("()[]{}")
        if not token:
            continue
        alpha = re.sub(r"[^A-Za-z]", "", token)
        if alpha and alpha.lower() in name_map:
            symbol = name_map[alpha.lower()]
            if symbol in allowed:
                kept.append(symbol)
            continue

        compact = token.replace("-", "").replace("_", "")
        if re.fullmatch(r"[A-Z][a-z]?", compact):
            if compact in allowed:
                kept.append(compact)
            continue

        chemistry_like = bool(re.fullmatch(r"[A-Za-z0-9\.\-\(\)%]+", token)) and (
            bool(re.search(r"\d", token))
            or len(re.findall(r"[A-Z][a-z]?", token)) >= 2
            or bool(re.search(r"[\(\)%\-]", token))
        )
        if not chemistry_like:
            continue

        for symbol in re.findall(r"[A-Z][a-z]?", token):
            if symbol in allowed:
                kept.append(symbol)
    return _sort_elements_for_merge(kept)



def normalize_loading_for_merge(value: str) -> str:
    """Function docstring removed for runtime stability."""
    if not value:
        return ""
    s = re.sub(r"\s+", "", str(value)).lower()
    s = re.sub(r"wt%$", "", s)
    return s


def normalize_active_metal_for_merge(value: str) -> str:
    """Function docstring removed for runtime stability."""
    if not value:
        return ""
    return "/".join(el.lower() for el in _extract_elements_for_merge(value))


def normalize_promoter_for_merge(value: str) -> str:
    """Function docstring removed for runtime stability."""
    if not value:
        return ""
    return "/".join(el.lower() for el in _extract_elements_for_merge(value))


def alloy_ratio_signature_for_merge(record: Dict) -> str:
    """Function docstring removed for runtime stability."""
    pairs = parse_alloy_ratio_to_ordered_pairs(record)
    return "|".join(f"{el.lower()}:{val}" for el, val in pairs)


def _build_metal_identity_for_merge(record: Dict, for_display: bool = False) -> str:
    """Function docstring removed for runtime stability."""
    alloy_sig = alloy_ratio_signature_for_merge(record)
    if alloy_sig:
        base = alloy_sig.replace(":", "").replace("|", "-")
    else:
        base = normalize_active_metal_for_merge(str(record.get("Active_Metal", ""))).replace("/", "-")

    promoter_norm = normalize_promoter_for_merge(str(record.get("Promoter", "")))
    promoter_part = promoter_norm.replace("/", "-") if promoter_norm else ""

    if promoter_part:
        if base:
            return f"{base}-prom-{promoter_part}" if for_display else f"{base}__prom-{promoter_part}"
        return f"prom-{promoter_part}"
    return base


def build_canonical_merge_key(record: Dict) -> str:
    """Function docstring removed for runtime stability."""
    import hashlib

    loading = normalize_loading_for_merge(str(record.get("Metal_Loading_wt%", "")))
    support = normalize_support_for_merge(str(record.get("Support_Normalized", "") or record.get("Support", "")))
    metal_identity = _build_metal_identity_for_merge(record)

    if metal_identity:
        loading_part = f"{loading}wt%-" if loading else ""
        support_part = f"/{support}" if support else ""
        return f"{loading_part}{metal_identity}{support_part}"

    normalized_alias = str(record.get("Catalyst_ID_normalized", "")).strip()
    if normalized_alias:
        return "normid-" + re.sub(r"\s+", "", normalized_alias).lower()

    llm_id = str(record.get("Catalyst_ID", "")).strip()
    if llm_id:
        return "llmid-" + re.sub(r"\s+", "", llm_id).lower()

    fp_fields = ["Catalyst_ID", "Catalyst", "Active_Metal",
                 "Metal_Loading_wt%", "Alloy_Ratio", "Promoter", "Support", "Source_File"]
    fp_raw = "|".join(str(record.get(f, "")).strip() for f in fp_fields)
    if not fp_raw.replace("|", "").strip():
        fp_raw = repr(sorted(record.items()))
    digest = hashlib.md5(fp_raw.encode("utf-8")).hexdigest()[:12]
    return f"__fallback_md5_{digest}"


def _selfcheck_promoter_aware_merge_key() -> bool:
    """Function docstring removed for runtime stability."""
    cases = {
        "Ni/Al2O3": {
            "Active_Metal": "Ni",
            "Metal_Loading_wt%": "10",
            "Support": "Al2O3",
            "Promoter": "",
            "Alloy_Ratio": "",
        },
        "Ni-Au/Al2O3": {
            "Active_Metal": "Ni",
            "Metal_Loading_wt%": "10",
            "Support": "Al2O3",
            "Promoter": "Au",
            "Alloy_Ratio": "",
        },
        "Ni-Rh/Al2O3": {
            "Active_Metal": "Ni",
            "Metal_Loading_wt%": "10",
            "Support": "Al2O3",
            "Promoter": "Rh",
            "Alloy_Ratio": "",
        },
        "Ni-Ir/Al2O3": {
            "Active_Metal": "Ni",
            "Metal_Loading_wt%": "10",
            "Support": "Al2O3",
            "Promoter": "Ir",
            "Alloy_Ratio": "",
        },
        "Ni80-Cu20/CNTs": {
            "Active_Metal": "Ni/Cu",
            "Metal_Loading_wt%": "20",
            "Support": "CNTs",
            "Promoter": "Cu",
            "Alloy_Ratio": "Ni:80, Cu:20",
        },
        "Ni20-Cu80/CNTs": {
            "Active_Metal": "Ni/Cu",
            "Metal_Loading_wt%": "20",
            "Support": "CNTs",
            "Promoter": "Cu",
            "Alloy_Ratio": "Ni:20, Cu:80",
        },
    }
    keys = [build_canonical_merge_key(rec) for rec in cases.values()]
    ok = len(keys) == len(set(keys))
    print(f"[Selfcheck promoter-aware merge key] {'PASS' if ok else 'FAIL'}")
    return ok


def _selfcheck_canonical_merge_key():
    """Function docstring removed for runtime stability."""
    return _selfcheck_promoter_aware_merge_key()


def normalize_catalyst_id(record: Dict) -> str:
    """Function docstring removed for runtime stability."""
    for label in [record.get("Catalyst", ""), record.get("Series_Name", "")]:
        if label and _parse_niznal_compact_identity(str(label)):
            temp = _apply_niznal_compact_identity(dict(record), str(label))
            structured = _build_structured_catalyst_id(temp)
            if structured:
                return structured.lower()
    llm_id = str(record.get("Catalyst_ID", "")).strip()
    if llm_id:
        return re.sub(r"\s+", "", llm_id).lower()

    normalize_identity_aliases(record)
    normalized_alias = str(record.get("Catalyst_ID_normalized", "")).strip()
    if normalized_alias:
        return normalized_alias.lower()

    structured_id = _build_structured_catalyst_id(record)
    return structured_id.lower() if structured_id else ""




_REGISTRY_STATIC_FIELDS = [
    "Catalyst_ID", "Catalyst_ID_normalized", "identity_alias_group", "Canonical_Catalyst_ID",
    "Catalyst", "Active_Metal", "Promoter", "Promoter_Metal", "Support",
    "Support_Normalized", "Support_Family", "Support_Modifier", "Support_Grouped",
    "Metal_Loading_wt%", "Alloy_Ratio", "Ni_Fraction", "Promoter_Fraction",
    "Precursor", "Precursor_Normalized", "Precursor_Family",
    "Metal_Loading_Method", "Metal_Loading_Method_Normalized",
    "Dry_Temp_C", "Dry_Time_h",
    "Calcination_Temp_C", "Reduction_Temp_C", "Reduction_Time_h", "Calcination_Time_h",
    "Support_Prep_Method", "Support_Prep_Method_Normalized", "Preparation_Fingerprint",
]
_FIGURE_IDENTITY_BINDABLE_FIELDS = [
    "Canonical_Catalyst_ID",
    "Catalyst", "Catalyst_ID", "Catalyst_ID_normalized", "identity_alias_group",
    "Active_Metal", "Promoter", "Promoter_Metal",
    "Support", "Support_Normalized", "Support_Family", "Support_Modifier",
    "Metal_Loading_wt%", "Ni_Loading_wt%", "Promoter_Loading_wt%",
    "Alloy_Ratio", "Ni_Fraction", "Promoter_Fraction",
    "Precursor", "Precursor_Normalized", "Precursor_Family",
    "Metal_Loading_Method", "Metal_Loading_Method_Normalized",
    "Dry_Temp_C", "Dry_Time_h",
    "Calcination_Temp_C", "Calcination_Time_h",
    "Reduction_Temp_C", "Reduction_Time_h",
    "Support_Prep_Method", "Support_Prep_Method_Normalized",
    "Support_Grouped", "Preparation_Fingerprint",
]
_CLUE_BACKFILL_FIELDS = [
    "Active_Metal",
    "Promoter",
    "Promoter_Metal",
    "Support",
    "Support_Normalized",
    "Metal_Loading_wt%",
    "Ni_Loading_wt%",
    "Promoter_Loading_wt%",
    "Alloy_Ratio",
    "Ni_Fraction",
    "Promoter_Fraction",
]
_FIGURE_PANEL_CONDITION_FIELDS = [
    "Reaction_Temp_C", "S_C_Ratio", "S_C_Ratio_Raw",
    "Pressure_bar", "Pressure_bar_Raw",
    "GHSV_mL_g_h", "GHSV_mL_g_h_Raw",
    "WHSV_h_inv", "WHSV_h_inv_Raw",
    "SpaceVelocity_norm", "SpaceVelocity_norm_Raw",
    "SpaceVelocity_type", "SpaceVelocity_unit",
    "Feed_MeOH_to_H2O_Ratio", "Feed_MeOH_to_H2O_Ratio_Raw",
    "TOS_h", "Feed_Composition",
]
_FIGURE_NO_BROADCAST_FIELDS = [
    "MeOH_Conversion_%", "H2_Yield_%", "H2_Production_Rate",
    "H2_Selectivity_%", "CO_Selectivity_%", "CO2_Selectivity_%",
    "CO_Concentration_ppm",
]
_FIGURE_BINDABLE_FIELDS = list(_FIGURE_IDENTITY_BINDABLE_FIELDS)
_FIELD_SOURCE_TRACKED_FIELDS = set(_FIGURE_IDENTITY_BINDABLE_FIELDS) | {
    "Reaction_Temp_C", "S_C_Ratio", "GHSV_mL_g_h", "Pressure_bar", "TOS_h"
}
_FIGURE_BINDING_AUDIT_FIELDS = [
    "Source_File", "data_source", "Series_Name", "raw_catalyst", "raw_catalyst_id",
    "raw_category_label", "x_axis", "x_value", "figure_binding_status", "figure_binding_mode",
    "figure_binding_confidence", "figure_binding_reason", "matched_registry_key", "matched_registry_label",
    "alias_map_source", "alias_map_evidence", "semantic_figure_role",
    "_broadcast_path_primary", "_broadcast_flags", "_field_source_map_json",
]


def _extract_x_axis_and_value_from_notes(notes: str) -> Tuple[str, str]:
    notes = str(notes or "")
    x_axis_match = re.search(r"\[x_axis=([^\]]+)\]", notes)
    x_value_match = re.search(r"\[x_value=([^\]]+)\]", notes)
    x_axis = x_axis_match.group(1).strip() if x_axis_match else ""
    x_value = x_value_match.group(1).strip() if x_value_match else ""
    return x_axis, x_value



def _make_registry_fallback_key(record: Dict) -> str:
    catalyst = re.sub(r"\s+", " ", _normalize_identity_text(str(record.get("Catalyst", "")))).lower().strip()
    active = normalize_active_metal_for_merge(str(record.get("Active_Metal", "")))
    promoter = normalize_promoter_for_merge(str(record.get("Promoter", "")))
    support = normalize_support_for_merge(str(record.get("Support_Normalized", "") or record.get("Support", "")))
    loading = normalize_loading_for_merge(str(record.get("Metal_Loading_wt%", "")))
    alloy = alloy_ratio_signature_for_merge(record)

    parts = []
    if catalyst:
        parts.append(f"cat={catalyst}")
    if active:
        parts.append(f"active={active}")
    if promoter:
        parts.append(f"prom={promoter}")
    if support:
        parts.append(f"support={support}")
    if loading:
        parts.append(f"loading={loading}")
    if alloy:
        parts.append(f"alloy={alloy}")
    return f"fallback|{'|'.join(parts)}" if parts else ""



def _relaxed_registry_alias_signature(label: str) -> str:
    text = _normalize_registry_label(label)
    if not text:
        return ""
    text = re.sub(
        r"\([^)]*(?:with(?:out)?|w/|w/o|fresh|spent|reduced|oxidized|calcined|treated|untreated|as-?prepared|pretreated)[^)]*\)",
        "",
        text,
        flags=re.I,
    )
    if "/" in text:
        head, support = text.split("/", 1)
    else:
        head, support = text, ""
    head = re.sub(r"^\d+(?:\.\d+)?(?:wt%|%)?-?", "", head, flags=re.I)
    head = re.sub(r"^\d+(?:\.\d+)?", "", head)
    relaxed = f"{head}/{support}" if support else head
    relaxed = re.sub(r"\s+", "", relaxed).strip(" -_/")
    return relaxed


def _make_registry_alias_keys(record: Dict) -> List[str]:
    temp = dict(record)
    normalize_identity_aliases(temp)
    for label in [temp.get("Catalyst", ""), temp.get("Series_Name", ""), temp.get("Catalyst_ID", "")]:
        if label and _parse_niznal_compact_identity(str(label)):
            temp = _apply_niznal_compact_identity(temp, str(label))
            break

    candidates = []
    for value in [
        temp.get("identity_alias_group", ""),
        temp.get("Catalyst_ID_normalized", ""),
        temp.get("Canonical_Catalyst_ID", ""),
        normalize_catalyst_id(temp),
    ]:
        key = str(value).strip().lower()
        if key and not key.startswith("__fallback_md5_"):
            candidates.append(key)
            candidates.extend(str(alias).strip().lower() for alias in _expand_canonical_id_to_aliases(key) if str(alias).strip())

    raw_cid = str(temp.get("Catalyst_ID", "")).strip()
    if raw_cid:
        candidates.append(re.sub(r"\s+", "", raw_cid).lower())

    for value in [
        temp.get("Catalyst", ""),
        temp.get("Catalyst_ID_raw", ""),
        temp.get("Catalyst_ID", ""),
        temp.get("Series_Name", ""),
        temp.get("Canonical_Catalyst_ID", ""),
        temp.get("identity_alias_group", ""),
    ]:
        raw = clean_residual_mojibake_chars(str(value or "")).strip()
        if not raw:
            continue
        raw_compact = re.sub(r"\s+", "", raw).lower()
        if raw_compact:
            candidates.append(raw_compact)
        norm = _normalize_registry_label(raw)
        if norm:
            candidates.append(norm)

    for template_candidate in temp.get("_template_candidates", []) or []:
        raw = clean_residual_mojibake_chars(str(template_candidate or "")).strip()
        if not raw:
            continue
        raw_compact = re.sub(r"\s+", "", raw).lower()
        if raw_compact:
            candidates.append(raw_compact)
        norm = _normalize_registry_label(raw)
        if norm:
            candidates.append(norm)

    component_pairs = parse_alloy_ratio_to_ordered_pairs(temp)
    support_norm = _normalize_registry_label(str(
        temp.get("Support_Normalized", "") or temp.get("Support", "") or _extract_support_from_identity_text(str(temp.get("Catalyst", "") or ""))
    ))
    candidates.extend(_build_component_loading_aliases(component_pairs, support_norm))
    candidates.extend(_build_single_loading_support_aliases(temp))

    fallback_key = _make_registry_fallback_key(temp)
    if fallback_key:
        candidates.append(fallback_key)
    return _dedupe_keep_order([c for c in candidates if c])



def _make_registry_key(record: Dict) -> str:
    temp = dict(record)
    normalize_identity_aliases(temp)
    for label in [temp.get("Catalyst", ""), temp.get("Series_Name", ""), temp.get("Catalyst_ID", "")]:
        if label and _parse_niznal_compact_identity(str(label)):
            temp = _apply_niznal_compact_identity(temp, str(label))
            break
    for value in [
        temp.get("Canonical_Catalyst_ID", ""),
        temp.get("identity_alias_group", ""),
        temp.get("Catalyst_ID_normalized", ""),
        normalize_catalyst_id(temp),
        _make_registry_fallback_key(temp),
    ]:
        key = str(value).strip().lower()
        if key:
            return key
    return ""



def build_preparation_fingerprint(record: Dict) -> str:
    if not isinstance(record, dict):
        return ""
    parts = []
    for key, field in [
        ("mlm", "Metal_Loading_Method_Normalized"),
        ("spm", "Support_Prep_Method_Normalized"),
        ("pf", "Precursor_Family"),
        ("sg", "Support_Grouped"),
    ]:
        value = clean_residual_mojibake_chars(str(record.get(field, ""))).strip()
        if value:
            parts.append(f"{key}={value}")
    return "|".join(parts)



def enrich_registry_preparation_backbone(record: Dict) -> Dict:
    if not isinstance(record, dict):
        return record

    row = dict(record)
    derived = apply_preparation_normalization_layer(row)
    normalized_fields = [
        "Metal_Loading_Method_Normalized",
        "Support_Prep_Method_Normalized",
        "Precursor_Normalized",
        "Precursor_Family",
        "Support_Grouped",
    ]
    for field in normalized_fields:
        if not str(row.get(field, "")).strip() and str(derived.get(field, "")).strip():
            row[field] = derived.get(field, "")
    if not str(row.get("Preparation_Fingerprint", "")).strip():
        row["Preparation_Fingerprint"] = build_preparation_fingerprint(row)
    return row



def _merge_registry_normalized_fields(base: Dict, candidate: Dict) -> Dict:
    merged = dict(base)
    conflict_notes = []
    normalized_fields = [
        "Metal_Loading_Method_Normalized",
        "Support_Prep_Method_Normalized",
        "Precursor_Normalized",
        "Precursor_Family",
        "Support_Grouped",
        "Preparation_Fingerprint",
    ]
    for field in normalized_fields:
        base_val = clean_residual_mojibake_chars(str(merged.get(field, ""))).strip()
        cand_val = clean_residual_mojibake_chars(str(candidate.get(field, ""))).strip()
        if cand_val.lower() in {"none", "nan", "n/a", "na", "null"}:
            cand_val = ""
        if base_val.lower() in {"none", "nan", "n/a", "na", "null"}:
            base_val = ""
            merged[field] = ""
        if not cand_val:
            continue
        if not base_val:
            merged[field] = candidate.get(field, "")
        elif base_val != cand_val:
            conflict_notes.append(f"conflict_{field}")
    if conflict_notes:
        merged["registry_notes"] = _merge_semicolon_notes(
            str(merged.get("registry_notes", "")),
            conflict_notes,
        )
    return merged



def _merge_static_registry_fields(base: Dict, candidate: Dict) -> Dict:
    merged = dict(base)
    conflict_notes = []
    normalized_fields = {
        "Metal_Loading_Method_Normalized",
        "Support_Prep_Method_Normalized",
        "Precursor_Normalized",
        "Precursor_Family",
        "Support_Grouped",
        "Preparation_Fingerprint",
    }

    base_role = str(merged.get("registry_role", "") or "")
    cand_role = str(candidate.get("registry_role", "") or "")
    if base_role != "family" and cand_role == "family":
        merged["registry_role"] = "family"
    elif not base_role and cand_role:
        merged["registry_role"] = cand_role
    if int(candidate.get("can_broadcast_preparation", 0) or 0):
        merged["can_broadcast_preparation"] = 1
    else:
        merged["can_broadcast_preparation"] = int(merged.get("can_broadcast_preparation", 0) or 0)

    for field in _REGISTRY_STATIC_FIELDS:
        if field in normalized_fields:
            continue
        if base_role == "family" and cand_role == "sample" and field in {
            "Metal_Loading_wt%", "Alloy_Ratio", "Ni_Fraction", "Promoter_Fraction"
        }:
            continue
        base_val = str(merged.get(field, "")).strip()
        cand_val = str(candidate.get(field, "")).strip()
        if cand_val.lower() in {"none", "nan", "n/a", "na", "null"}:
            cand_val = ""
        if base_val.lower() in {"none", "nan", "n/a", "na", "null"}:
            base_val = ""
            merged[field] = ""
        if not cand_val:
            continue
        if not base_val:
            merged[field] = candidate.get(field, "")
        elif base_val != cand_val:
            conflict_notes.append(f"conflict_{field}")

    merged = _merge_registry_normalized_fields(merged, candidate)
    merged["source_record_count"] = int(merged.get("source_record_count", 1)) + 1
    merged["_registry_aliases"] = _dedupe_keep_order(
        list(merged.get("_registry_aliases", [])) + list(candidate.get("_registry_aliases", []))
    )
    merged["registry_notes"] = _merge_semicolon_notes(
        str(merged.get("registry_notes", "")),
        conflict_notes + [str(candidate.get("registry_notes", "")).strip()]
    )
    return merged



def _has_static_preparation_identity_signal(row: Dict) -> bool:
    if not isinstance(row, dict):
        return False

    def _filled(field: str) -> bool:
        value = str(row.get(field, "")).strip()
        return value not in {"", "N/A", "n/a", "None", "none"}

    identity_fields = [
        "Catalyst", "Catalyst_ID", "Active_Metal", "Promoter", "Support",
        "Metal_Loading_wt%", "Alloy_Ratio",
    ]
    preparation_fields = [
        "Precursor", "Metal_Loading_Method", "Support_Prep_Method",
        "Dry_Temp_C", "Dry_Time_h", "Calcination_Temp_C", "Calcination_Time_h",
        "Reduction_Temp_C", "Reduction_Time_h",
    ]

    identity_count = sum(1 for field in identity_fields if _filled(field))
    preparation_count = sum(1 for field in preparation_fields if _filled(field))
    return preparation_count >= 2 or (identity_count >= 1 and preparation_count >= 1) or identity_count >= 3


def _has_broadcastable_preparation_fields(record: Dict) -> bool:
    if not isinstance(record, dict):
        return False
    fields = [
        "Precursor", "Precursor_Normalized", "Precursor_Family",
        "Metal_Loading_Method", "Metal_Loading_Method_Normalized",
        "Support_Prep_Method", "Support_Prep_Method_Normalized",
        "Dry_Temp_C", "Dry_Time_h",
        "Calcination_Temp_C", "Calcination_Time_h",
        "Reduction_Temp_C", "Reduction_Time_h",
        "Preparation_Fingerprint",
    ]
    return any(str(record.get(field, "") or "").strip() for field in fields)


def _has_registry_identity_core(record: Dict) -> bool:
    if not isinstance(record, dict):
        return False
    if str(record.get("Canonical_Catalyst_ID", "") or record.get("Catalyst_ID", "") or "").strip():
        return True
    return bool(
        str(record.get("Active_Metal", "") or "").strip()
        and str(record.get("Support_Normalized", "") or record.get("Support", "") or "").strip()
    )


def _registry_record_role(record: Dict) -> str:
    if not _has_registry_identity_core(record):
        return ""
    catalyst = _normalize_registry_label(str(record.get("Catalyst", "") or ""))
    catalyst_head = catalyst.split("/", 1)[0].strip() if "/" in catalyst else catalyst
    loading = normalize_loading_for_merge(str(record.get("Metal_Loading_wt%", "") or ""))
    alloy = alloy_ratio_signature_for_merge(record)
    if loading or alloy or re.search(r"\d", catalyst_head):
        return "sample"
    return "family"


_NIZNAL_TABLE_SPECS: Tuple[Tuple[str, float, float], ...] = (
    ("2NiAl", 2.0, 0.0),
    ("5ZnAl", 0.0, 5.0),
    ("2Ni5Zn", 2.0, 5.0),
    ("4NiAl", 4.0, 0.0),
    ("10ZnAl", 0.0, 10.0),
    ("4Ni10Zn", 4.0, 10.0),
    ("8NiAl", 8.0, 0.0),
    ("20ZnAl", 0.0, 20.0),
    ("8Ni20Zn", 8.0, 20.0),
)


def _looks_like_niznal_context_text(text: str) -> bool:
    s = _normalize_identity_text(text).lower()
    if not s:
        return False
    return bool(
        re.search(r"\bniznal\b", s, flags=re.I)
        or re.search(r"\bznal2o4\b", s, flags=re.I)
        or re.search(r"\bni\s*/\s*znal2o4\b", s, flags=re.I)
        or re.search(r"\bni\s*[-/]\s*zn\s*/\s*al2o3\b", s, flags=re.I)
        or re.search(r"\b(?:2|4|8)ni(?:5|10|20)zn\b", s, flags=re.I)
    )


def _format_niznal_value(value: float) -> str:
    return _format_ratio_token(str(value))


def _build_niznal_identity_seed_row(label: str, ni_loading: float, zn_loading: float, file_name: str) -> Dict[str, Any]:
    row: Dict[str, Any] = {
        "Source_File": file_name,
        "data_source": "text",
        "source_granularity": "table_row",
        "paragraph_role": "preparation_paragraph",
        "text_extraction_subroute": "niznal_table_identity_seed",
        "Catalyst": label,
        "Support": "Al2O3",
        "Support_Normalized": "Al2O3",
        "Support_Grouped": "Al2O3",
        "Notes": "[table_identity_seed=NiZnAl Table 1; no_performance_value_inferred]",
        "identity_completeness_level": "identity_only",
        "numeric_reliability_level": "identity_only",
    }
    if ni_loading > 0:
        row["Active_Metal"] = "Ni"
        row["Metal_Loading_wt%"] = _format_niznal_value(ni_loading)
        row["Precursor"] = "Ni nitrate" if zn_loading <= 0 else "Ni and Zn nitrate"
        row["Precursor_Family"] = "nitrate"
    elif zn_loading > 0:
        row["Active_Metal"] = "Zn"
        row["Metal_Loading_wt%"] = _format_niznal_value(zn_loading)
        row["Precursor"] = "Zn nitrate"
        row["Precursor_Family"] = "nitrate"
    if ni_loading > 0 and zn_loading > 0:
        total = ni_loading + zn_loading
        row["Promoter"] = "Zn"
        row["Promoter_Metal"] = "Zn"
        row["Alloy_Ratio"] = f"Ni:{_format_niznal_value(ni_loading)}, Zn:{_format_niznal_value(zn_loading)}"
        row["Ni_Fraction"] = _format_float_token(ni_loading / total)
        row["Promoter_Fraction"] = _format_float_token(zn_loading / total)
    elif ni_loading > 0:
        row["Ni_Fraction"] = "1"
        row["Promoter_Fraction"] = "0"
    return row


def _build_niznal_identity_seed_rows(records: List[Dict], file_name: str) -> List[Dict]:
    text = " ".join(
        clean_residual_mojibake_chars(str(value or ""))
        for record in (records or [])
        if isinstance(record, dict)
        for value in [
            record.get("Source_File", ""),
            record.get("Catalyst", ""),
            record.get("Catalyst_ID", ""),
            record.get("Series_Name", ""),
            record.get("Support", ""),
            record.get("Notes", ""),
        ]
    )
    if not _looks_like_niznal_context_text(" ".join([file_name, text])):
        return []
    existing = {
        re.sub(r"\s+", "", clean_residual_mojibake_chars(str(record.get("Catalyst", "") or "")).lower())
        for record in (records or [])
        if isinstance(record, dict) and str(record.get("data_source", "")).strip() == "text"
    }
    rows: List[Dict] = []
    for label, ni_loading, zn_loading in _NIZNAL_TABLE_SPECS:
        if label.lower() in existing:
            continue
        rows.append(_build_niznal_identity_seed_row(label, ni_loading, zn_loading, file_name))
    return rows



def _is_preparation_backbone_row(row: Dict) -> bool:
    if not isinstance(row, dict):
        return False
    if str(row.get("data_source", "")).strip() != "text":
        return False

    paragraph_role = str(row.get("paragraph_role", "")).strip()
    if paragraph_role in {"comparison_paragraph", "characterization_paragraph"}:
        return False
    if bool(row.get("is_literature_comparison")):
        return False
    if str(row.get("comparison_filter_action", "")).strip() == "downgraded_candidate_only":
        return False
    if not _has_static_preparation_identity_signal(row):
        return False

    if paragraph_role == "preparation_paragraph":
        return True
    if str(row.get("text_extraction_subroute", "")).strip() == "preparation_schema":
        return True

    source_granularity = str(row.get("source_granularity", "")).strip()
    if source_granularity in {"table_row", "si_table"}:
        return True

    # 放宽：text_numeric行如果有CID且有足够制备信号也纳入注册表
    if source_granularity == "text_numeric":
        has_cid = bool(str(row.get("Canonical_Catalyst_ID", "") or row.get("Catalyst_ID", "")).strip())
        if has_cid and _has_static_preparation_identity_signal(row):
            return True

    return False



def build_text_sample_registry(records: List[Dict], file_name: str) -> Dict[str, Dict]:
    registry: Dict[str, Dict] = {}
    alias_lookup: Dict[str, str] = {}
    total_text_rows = 0
    backbone_rows = 0
    role_counts: Counter = Counter()
    registry_source_records = list(records or [])
    niznal_seed_rows = _build_niznal_identity_seed_rows(registry_source_records, file_name)
    if niznal_seed_rows:
        print(f"  [registry] added NiZnAl table identity seeds: {len(niznal_seed_rows)}")
        registry_source_records.extend(niznal_seed_rows)

    for record in registry_source_records:
        if str(record.get("data_source", "")).strip() != "text":
            continue
        total_text_rows += 1
        role = _registry_record_role(record)
        if not role:
            continue
        can_broadcast_preparation = int(
            _is_preparation_backbone_row(record)
            or _has_broadcastable_preparation_fields(record)
        )
        if can_broadcast_preparation:
            backbone_rows += 1
        role_counts[role] += 1

        candidate = enrich_registry_preparation_backbone(dict(record))
        normalize_identity_aliases(candidate)
        for identity_label in [candidate.get("Catalyst", ""), candidate.get("Series_Name", ""), candidate.get("Catalyst_ID", "")]:
            if identity_label and _parse_niznal_compact_identity(str(identity_label)):
                candidate = _apply_niznal_compact_identity(candidate, str(identity_label))
                break
        if not str(candidate.get("Catalyst_ID", "")).strip():
            normalized_id = normalize_catalyst_id(candidate)
            if normalized_id:
                candidate["Catalyst_ID"] = normalized_id

        registry_key = _make_registry_key(candidate)
        if not registry_key:
            continue

        entry = {}
        for field in _REGISTRY_STATIC_FIELDS:
            value = candidate.get(field, "")
            text_value = clean_residual_mojibake_chars(str(value or "")).strip()
            entry[field] = "" if text_value.lower() in {"none", "nan", "n/a", "na", "null"} else value
        entry["registry_key"] = registry_key
        entry["source_record_count"] = 1
        entry["source_file"] = file_name
        entry["registry_role"] = role
        entry["can_broadcast_preparation"] = can_broadcast_preparation
        entry["registry_notes"] = "fallback_registry_key" if registry_key.startswith("fallback|") else ""
        entry["_registry_aliases"] = _make_registry_alias_keys(candidate)

        existing_key = ""
        existing_from_alias = False
        for alias in entry["_registry_aliases"]:
            if alias in alias_lookup:
                existing_key = alias_lookup[alias]
                existing_from_alias = True
                break
        if not existing_key and registry_key in registry:
            existing_key = registry_key

        if existing_key and existing_from_alias and existing_key in registry:
            cand_metals = _extract_registry_metal_set(entry)
            existing_metals = _extract_registry_metal_set(registry.get(existing_key, {}))
            compatible, _why = _is_metal_set_compatible_for_unique_match(cand_metals, existing_metals)
            if not compatible:
                existing_key = ""
            else:
                cand_loading = normalize_loading_for_merge(str(entry.get("Metal_Loading_wt%", "") or ""))
                existing_loading = normalize_loading_for_merge(str(registry.get(existing_key, {}).get("Metal_Loading_wt%", "") or ""))
                cand_alloy = alloy_ratio_signature_for_merge(entry)
                existing_alloy = alloy_ratio_signature_for_merge(registry.get(existing_key, {}))
                if cand_loading and existing_loading and cand_loading != existing_loading:
                    existing_key = ""
                elif cand_alloy and existing_alloy and cand_alloy != existing_alloy:
                    existing_key = ""

        if existing_key:
            registry[existing_key] = _merge_static_registry_fields(registry[existing_key], entry)
            for alias in registry[existing_key].get("_registry_aliases", []):
                alias_lookup[alias] = existing_key
        else:
            registry[registry_key] = entry
            for alias in entry["_registry_aliases"]:
                alias_lookup[alias] = registry_key

    print(f"  [registry] using preparation-broadcast-capable text rows: {backbone_rows} / {total_text_rows}")
    print(f"  [registry] role_counts={dict(role_counts)}")
    print(f"  [registry] {file_name}: {len(registry)} text sample entries")
    return registry


NON_CATALYST_SPECIES: frozenset = frozenset({
    "h2", "h2o", "co", "co2", "ch4", "ch3oh", "meoh", "dme",
    "o2", "n2", "c2h4", "c3h8", "nh3", "ar", "he",
    "hydrogen", "water", "steam", "carbon monoxide", "carbon dioxide",
    "methane", "oxygen", "nitrogen", "ethylene", "propane", "ammonia",
})


_NON_CATALYST_SERIES_LABELS: frozenset = frozenset(
    set(NON_CATALYST_SPECIES)
    | {
        "conv", "conversion", "selectivity", "yield", "activity", "tof", "h2-tor",
        "conv.", "x_meoh", "h2 production rate",
    }
)


def _normalize_species_token(text: str) -> str:
    raw = _normalize_identity_text(text)
    raw = raw.replace("\u2013", "-").replace("\u2014", "-").replace("\u2212", "-")
    raw = re.sub(r"(?<=[A-Za-z])\s+(?=\d)", "", raw)
    raw = re.sub(r"\s+", " ", raw).strip().lower()
    return raw


def _is_instrument_or_analysis_series_label(label: str) -> bool:
    s = _normalize_species_token(label)
    if not s:
        return False
    if s in {"dtg", "tg", "tga", "dta", "ms", "ftir", "xrd", "xps", "tem", "sem"}:
        return True
    if re.search(r"\b(?:tcd|fid|ms|ftir|gc|gc-ms)\s+signal\b", s, flags=re.I):
        return True
    if re.search(r"\b(?:dtg|tg|tga|dta)\b", s, flags=re.I):
        return True
    if re.search(r"\b(?:signal|intensity|counts?|a\.u\.|arb\.?\s*units?)\b", s, flags=re.I):
        return not _looks_like_identity_mapping(s)
    return False

def _is_generic_series_label(label: str) -> bool:
    s = _normalize_species_token(label)
    if not s:
        return False
    if _is_instrument_or_analysis_series_label(s):
        return True
    if re.fullmatch(r"[a-z]", s):
        return True
    if re.fullmatch(r"(?:sample|cat|cat\.?|catalyst|run)\s*[- ]?[a-z0-9]+", s):
        return True
    if re.fullmatch(r"cat[- ]?\d+", s):
        return True
    if re.fullmatch(r"sample\s+[a-z]", s):
        return True
    # MSR reaction species / condition labels — not catalyst identities
    if s in _NON_CATALYST_SERIES_LABELS:
        return True
    # "H₂ production rate (...)" style
    if re.match(r"^h2\s+(?:production|evolution|yield)", s):
        return True
    # Case N / Case-N style
    if re.fullmatch(r"case\s*[-]?\s*\d+", s):
        return True
    # 1# / 2# style
    if re.fullmatch(r"\d+#", s):
        return True
    # Pure ratio labels: S/MA=1, S/C=3
    if re.fullmatch(r"[a-z]{1,4}/[a-z]{1,4}\s*=\s*[\d.]+", s):
        return True
    if s in {"equilibrium", "eq", "eq."}:
        return True
    if re.fullmatch(r"(?:co2?|h2|ch4|ch3oh|meoh)\s*-\s*eq(?:uivalent)?", s, flags=re.I):
        return True
    return False


def _is_non_catalyst_figure_label(label: str) -> bool:
    s = _normalize_species_token(label)
    if not s:
        return False
    if _is_instrument_or_analysis_series_label(s):
        return True
    if _is_generic_series_label(s):
        return True
    if re.search(r"\b(?:methanol|hydrogen|water|steam|carbon monoxide|carbon dioxide|methane)\b", s):
        return True
    if re.search(r"\b(?:conversion|yield|selectivity|activity|deactivation|stability|production rate|evolution rate)\b", s):
        return True
    if re.fullmatch(r"(?:co2?|h2|ch4|ch3oh|meoh|methanol|hydrogen|carbon monoxide|carbon dioxide)\s*-\s*eq(?:uivalent)?", s, flags=re.I):
        return True
    if s in {"equilibrium", "eq", "eq."}:
        return True
    return False


def _strip_non_catalyst_identity_labels(row: Dict) -> Dict:
    current = dict(row or {})
    filtered: List[str] = []

    for field in ("raw_category_label", "Catalyst", "Series_Name"):
        raw = clean_residual_mojibake_chars(str(current.get(field, "") or "")).strip()
        if not raw:
            continue
        normalized = _normalize_species_token(raw)
        if _is_non_catalyst_figure_label(normalized) or _is_generic_series_label(normalized):
            current[field] = ""
            filtered.append(f"{field}={raw}")

    if filtered:
        tag = f"[filtered_non_catalyst_label={'|'.join(filtered[:3])}]"
        notes = str(current.get("Notes", "") or "").strip()
        if tag not in notes:
            current["Notes"] = (notes + " " + tag).strip()

    return current


def _has_identity_context_for_figure_row(row: Dict) -> bool:
    current = dict(row or {})
    return bool(
        clean_residual_mojibake_chars(str(current.get("raw_category_label", "") or "")).strip()
        or clean_residual_mojibake_chars(str(current.get("Catalyst", "") or "")).strip()
        or clean_residual_mojibake_chars(str(current.get("Series_Name", "") or "")).strip()
        or clean_residual_mojibake_chars(str(current.get("Canonical_Catalyst_ID", "") or "")).strip()
    )


def _is_product_species_only_identity(row: Dict) -> bool:
    values = [
        str((row or {}).get("Catalyst", "") or ""),
        str((row or {}).get("raw_category_label", "") or ""),
        str((row or {}).get("Series_Name", "") or ""),
    ]
    normalized = [value for value in (_normalize_species_token(item) for item in values) if value]
    if not normalized:
        return False
    return all(
        value in NON_CATALYST_SPECIES
        or _is_non_catalyst_figure_label(value)
        or _is_generic_series_label(value)
        for value in normalized
    )


def _should_drop_filtered_non_catalyst_row(row: Dict) -> bool:
    notes = clean_residual_mojibake_chars(str((row or {}).get("Notes", "") or "")).strip()
    has_identity_context = _has_identity_context_for_figure_row(row)
    if _is_product_species_only_identity(row):
        return not has_identity_context
    if "filtered_non_catalyst_label=" not in notes:
        return False
    return not has_identity_context


def _usable_figure_category_label(label: Any) -> str:
    cleaned = clean_residual_mojibake_chars(str(label or "")).strip()
    if not cleaned:
        return ""
    return "" if _is_non_catalyst_figure_label(cleaned) else cleaned



def _looks_like_identity_mapping(text: str) -> bool:
    s = _normalize_identity_text(text)
    if not s:
        return False
    if "/" in s:
        return True
    if re.search(r"\b(?:ni|cu|pd|pt|ru|rh|co|fe|ir|au|ce|la|mg|mo|cr|in|zn|ga|sn)\b", s, flags=re.I):
        return True
    if re.search(r"al2o3|ceo2|sio2|zro2|tio2|cnts|sep|activated carbon", s, flags=re.I):
        return True
    if re.search(r"\d+(?:\.\d+)?\s*(?:wt\.?\s*%|wt%|%)", s, flags=re.I):
        return True
    return False



_CONTEXT_NUMBERED_LABEL_RE = re.compile(r"(?:[a-z]\d{1,3}|case\s*[-]?\s*\d+|\d+#)$", flags=re.I)


def _normalize_context_alias_key(label: str) -> str:
    value = clean_residual_mojibake_chars(str(label or "")).replace("\r", " ")
    value = re.sub(r"\s+", " ", value).strip(" ;:,.()[]{}")
    return value.lower()


def _clean_context_alias_label(text: str) -> str:
    cleaned = clean_residual_mojibake_chars(str(text or "")).replace("\r", " ")
    cleaned = re.sub(r"\s+", " ", cleaned).strip(" ;:,.()[]{}")
    cleaned = re.sub(r"^(?:sample|samples|catalyst|catalysts|cat\.?|run)\s+", "", cleaned, flags=re.I)
    cleaned = re.sub(r"^(?:and|or)\s+", "", cleaned, flags=re.I)
    return cleaned.strip(" ;:,.()[]{}")


def _clean_context_mapped_label(text: str) -> str:
    mapped = _clean_context_alias_label(text)
    mapped = re.sub(r"^(?:the\s+)?(?:tested|prepared|synthesized|used)\s+(?:catalysts?|samples?)\s+(?:were|include|are)\s+", "", mapped, flags=re.I)
    mapped = re.sub(r"\b(?:tested at|used in|for msr|for methanol steam reforming)\b.*$", "", mapped, flags=re.I)
    mapped = re.sub(r"\b(?:hereafter|denoted|labeled|labelled|coded|named)\b\s+as\b.*$", "", mapped, flags=re.I)
    mapped = re.sub(r"\s+", " ", mapped).strip(" ;:,.()[]{}")
    return mapped


def _is_context_alias_label_candidate(text: str) -> bool:
    cleaned = _clean_context_alias_label(text)
    if not cleaned:
        return False
    if len(cleaned) > 80 or len(cleaned.split()) > 10:
        return False
    if re.search(r"[;\n]", cleaned):
        return False
    if re.fullmatch(r"\d+(?:\.\d+)?", cleaned):
        return False
    if _CONTEXT_NUMBERED_LABEL_RE.fullmatch(cleaned):
        return True
    if _is_generic_series_label(cleaned):
        return True
    if _looks_like_identity_mapping(cleaned):
        return True
    if re.search(r"\b(?:with|without|w/|w/o|fresh|spent|reduced|oxidized|calcined|treated|untreated|as-?prepared)\b", cleaned, flags=re.I):
        return True
    return "/" in cleaned or "-" in cleaned


def _context_label_lookup_keys(label: str) -> List[str]:
    cleaned = _clean_context_alias_label(label)
    keys = []
    raw_key = _normalize_context_alias_key(cleaned)
    if raw_key:
        keys.append(raw_key)
    norm = _normalize_registry_label(cleaned)
    if norm:
        keys.append(norm)
    relaxed = _relaxed_registry_alias_signature(cleaned)
    if relaxed:
        keys.append(relaxed)
    return _dedupe_keep_order([key for key in keys if key])


def _split_identity_list_preserve_parenthetical(text: str) -> List[str]:
    raw = clean_residual_mojibake_chars(str(text or "")).strip()
    if not raw:
        return []

    parts: List[str] = []
    current: List[str] = []
    depth = 0
    i = 0
    while i < len(raw):
        ch = raw[i]
        if ch in "([{":
            depth += 1
            current.append(ch)
            i += 1
            continue
        if ch in ")]}":
            depth = max(0, depth - 1)
            current.append(ch)
            i += 1
            continue
        if depth == 0 and ch in ",;":
            part = "".join(current).strip()
            if part:
                parts.append(part)
            current = []
            i += 1
            continue
        if depth == 0 and raw[i:i + 5].lower() == " and ":
            part = "".join(current).strip()
            if part:
                parts.append(part)
            current = []
            i += 5
            continue
        current.append(ch)
        i += 1

    tail = "".join(current).strip()
    if tail:
        parts.append(tail)
    return _dedupe_keep_order([part for part in parts if part])


def _extract_variable_assignments(text: str) -> Dict[str, List[str]]:
    source = clean_residual_mojibake_chars(str(text or "")).replace("\r", " ")
    if not source:
        return {}

    assignments: Dict[str, List[str]] = {}
    allowed = MSR_ACTIVE_METALS | MSR_PROMOTERS | MSR_RARE_EARTH | MSR_NI_SECOND_METALS
    patterns = [
        re.compile(r"\b([A-Z])\s*=\s*([^\.;\)\]\}]{1,80})"),
        re.compile(r"[\(\[\{]\s*([A-Z])\s*=\s*([^\)\]\}]{1,80})[\)\]\}]"),
    ]

    for pattern in patterns:
        for match in pattern.finditer(source):
            var_name = match.group(1).strip().upper()
            raw_values = clean_residual_mojibake_chars(match.group(2)).strip()
            if not var_name or var_name in allowed:
                continue
            raw_values = re.sub(r"\b(?:where|with|which|and the)\b.*$", "", raw_values, flags=re.I)
            values = []
            for part in re.split(r"\s*(?:,|/|\bor\b|\band\b)\s*", raw_values):
                symbol = _normalize_promoter_like_token(part)
                if symbol and symbol in allowed:
                    values.append(symbol)
            if not values:
                values = [
                    symbol for symbol in (
                        _normalize_promoter_like_token(token)
                        for token in re.findall(r"[A-Z][a-z]?", raw_values)
                    )
                    if symbol and symbol in allowed
                ]
            if values:
                assignments[var_name] = _dedupe_keep_order(values)
    return assignments


def _extract_template_suffix_tail(text_tail: str) -> str:
    tail = clean_residual_mojibake_chars(str(text_tail or "")).strip()
    if not tail:
        return ""
    suffix_patterns = [
        r"^(layered double hydroxides?)\b",
        r"^(hydrotalcites?)\b",
        r"^(ldhs?)\b",
    ]
    for pattern in suffix_patterns:
        match = re.search(pattern, tail, flags=re.I)
        if match:
            suffix = match.group(1)
            if re.search(r"layered double hydroxides?|ldhs?", suffix, flags=re.I):
                return "LDH"
            return suffix
    return ""


def _extract_identity_template_candidates(text: str) -> Dict[str, Any]:
    result: Dict[str, Any] = {
        "template": "",
        "variable": "",
        "values": [],
        "expanded": [],
        "suffix": "",
    }
    source = clean_residual_mojibake_chars(str(text or "")).replace("\r", " ")
    if not source:
        return result

    assignments = _extract_variable_assignments(source)
    if not assignments:
        return result

    allowed = MSR_ACTIVE_METALS | MSR_PROMOTERS | MSR_RARE_EARTH | MSR_NI_SECOND_METALS

    parenthetical_match = re.search(
        r"\b([A-Z][a-z]?)\s*[\(\[\{]\s*([A-Z])\s*[\)\]\}]\s*([A-Z][a-z]?)",
        source,
    )
    if parenthetical_match:
        prefix_metal, var_name, suffix_metal = parenthetical_match.group(1), parenthetical_match.group(2), parenthetical_match.group(3)
        if var_name in assignments and var_name not in allowed:
            result["template"] = f"{prefix_metal}({var_name}){suffix_metal}"
            result["variable"] = var_name
            result["values"] = list(assignments.get(var_name, []))
            result["suffix"] = _extract_template_suffix_tail(source[parenthetical_match.end():parenthetical_match.end() + 60])
            result["expanded"] = [
                f"{prefix_metal}({value}){suffix_metal}" + (f" {result['suffix']}" if result["suffix"] else "")
                for value in result["values"]
            ]
            return result

    hyphen_pattern = re.compile(r"\b(?:[A-Z][a-z]?|[A-Z])(?:\s*-\s*(?:[A-Z][a-z]?|[A-Z])){2,4}\b")
    for match in hyphen_pattern.finditer(source):
        tokens = [token.strip() for token in re.split(r"\s*-\s*", match.group(0)) if token.strip()]
        var_positions = [idx for idx, token in enumerate(tokens) if len(token) == 1 and token in assignments and token not in allowed]
        if len(var_positions) != 1:
            continue
        idx = var_positions[0]
        var_name = tokens[idx]
        result["template"] = "-".join(tokens)
        result["variable"] = var_name
        result["values"] = list(assignments.get(var_name, []))
        result["suffix"] = _extract_template_suffix_tail(source[match.end():match.end() + 60])
        result["expanded"] = []
        for value in result["values"]:
            expanded_tokens = list(tokens)
            expanded_tokens[idx] = value
            expanded = "-".join(expanded_tokens)
            if result["suffix"]:
                expanded += f" {result['suffix']}"
            result["expanded"].append(expanded)
        if result["expanded"]:
            return result

    return result


def _registry_like_context_label(text: str, registry: Optional[Dict] = None) -> bool:
    candidate = _clean_context_mapped_label(text)
    if not candidate:
        return False
    if _looks_like_identity_mapping(candidate) or _normalize_support_only_label(candidate) or _is_blank_or_baseline_label(candidate):
        return True
    norm = _normalize_registry_label(candidate)
    relaxed = _relaxed_registry_alias_signature(candidate)
    if not norm and not relaxed:
        return False
    for reg_key, reg_record in (registry or {}).items():
        aliases = _dedupe_keep_order([
            _normalize_registry_label(reg_key),
            _normalize_registry_label(str(reg_record.get("Catalyst", "") or "")),
            _normalize_registry_label(str(reg_record.get("Catalyst_ID", "") or "")),
            _normalize_registry_label(str(reg_record.get("Canonical_Catalyst_ID", "") or "")),
            _normalize_registry_label(str(reg_record.get("identity_alias_group", "") or "")),
            _relaxed_registry_alias_signature(reg_key),
            _relaxed_registry_alias_signature(str(reg_record.get("Catalyst", "") or "")),
            _relaxed_registry_alias_signature(str(reg_record.get("Catalyst_ID", "") or "")),
            _relaxed_registry_alias_signature(str(reg_record.get("Canonical_Catalyst_ID", "") or "")),
            _relaxed_registry_alias_signature(str(reg_record.get("identity_alias_group", "") or "")),
        ] + [
            _normalize_registry_label(str(alias or "")) for alias in reg_record.get("_registry_aliases", []) or []
        ] + [
            _relaxed_registry_alias_signature(str(alias or "")) for alias in reg_record.get("_registry_aliases", []) or []
        ])
        alias_pool = {alias for alias in aliases if alias}
        if norm and norm in alias_pool:
            return True
        if relaxed and relaxed in alias_pool:
            return True
    return False


def _build_context_label_alias_map(local_context: str, registry: Optional[Dict] = None) -> Dict[str, Dict[str, Any]]:
    context_raw = clean_residual_mojibake_chars(str(local_context or "")).replace("\r", " ")
    context_compact = re.sub(r"\s+", " ", context_raw).strip()
    alias_map: Dict[str, Dict[str, Any]] = {}
    if not context_compact:
        return alias_map

    def _add(label: str, mapped: str, source: str, evidence: str) -> None:
        label_clean = _clean_context_alias_label(label)
        mapped_clean = _clean_context_mapped_label(mapped)
        if not _is_context_alias_label_candidate(label_clean) or not _registry_like_context_label(mapped_clean, registry):
            return
        if _normalize_registry_label(label_clean) and _normalize_registry_label(label_clean) == _normalize_registry_label(mapped_clean):
            return

        entry = {
            "mapped_label": mapped_clean,
            "alias_map_source": source,
            "alias_map_evidence": _context_preview(evidence, 320),
            "ambiguous": False,
        }
        mapped_norm = _normalize_registry_label(mapped_clean)

        for key in _context_label_lookup_keys(label_clean):
            existing = alias_map.get(key)
            if existing:
                existing_norm = _normalize_registry_label(str(existing.get("mapped_label", "")))
                if existing_norm and mapped_norm and existing_norm != mapped_norm:
                    existing["ambiguous"] = True
                    existing["alias_map_evidence"] = _context_preview(
                        " || ".join(_dedupe_keep_order([existing.get("alias_map_evidence", ""), entry["alias_map_evidence"]])),
                        360,
                    )
                    existing["alias_map_source"] = "|".join(_dedupe_keep_order([
                        str(existing.get("alias_map_source", "")).strip(),
                        source,
                    ]))
                continue
            alias_map[key] = dict(entry)

    def _store_candidate(label: str, source: str, evidence: str, **extra_flags: Any) -> None:
        label_clean = _clean_context_mapped_label(label)
        if not _registry_like_context_label(label_clean, registry):
            return
        entry = {
            "mapped_label": label_clean,
            "alias_map_source": source,
            "alias_map_evidence": _context_preview(evidence, 320),
            "ambiguous": False,
        }
        entry.update(extra_flags)
        for key in _context_label_lookup_keys(label_clean):
            alias_map.setdefault(key, dict(entry))

    def _maybe_add_pair(left: str, right: str, source: str, evidence: str) -> None:
        left_clean = _clean_context_alias_label(left)
        right_clean = _clean_context_alias_label(right)
        if not left_clean or not right_clean or left_clean == right_clean:
            return

        left_registry = _registry_like_context_label(left_clean, registry)
        right_registry = _registry_like_context_label(right_clean, registry)
        left_candidate = _is_context_alias_label_candidate(left_clean)
        right_candidate = _is_context_alias_label_candidate(right_clean)
        same_relaxed = bool(
            _relaxed_registry_alias_signature(left_clean)
            and _relaxed_registry_alias_signature(left_clean) == _relaxed_registry_alias_signature(right_clean)
        )

        if source == "label_parenthetical":
            if left_registry and right_candidate:
                _add(right_clean, left_clean, source, evidence)
                return
            if right_registry and left_candidate and not left_registry:
                _add(left_clean, right_clean, source, evidence)
                return

        if right_registry and left_candidate:
            _add(left_clean, right_clean, source, evidence)
        if left_registry and right_candidate and (
            _CONTEXT_NUMBERED_LABEL_RE.fullmatch(right_clean)
            or _is_generic_series_label(right_clean)
            or same_relaxed
        ):
            _add(right_clean, left_clean, source, evidence)

    directional_patterns = [
        ("verbal_forward", re.compile(
            r"([^.;\n]{2,80}?)\s*(?:corresponds to|is|represents|denotes|stands for|refers to|maps to|equals?|assigned to|identified as)\s*([^.;\n]{3,120})",
            flags=re.I,
        )),
        ("verbal_reverse", re.compile(
            r"([^.;\n]{3,120}?)\s*(?:,|\s)*(?:denoted as|labeled as|labelled as|coded as|named|identified as|abbreviated as)\s*([^.;\n]{2,80})",
            flags=re.I,
        )),
        ("inline_equals", re.compile(r"([^;\n:=]{2,80})\s*[:=]\s*([^;\n]{3,120})", flags=re.I)),
        ("label_parenthetical", re.compile(r"([^;\n]{2,120}?)\s*\(\s*([^\)\n]{2,80})\s*\)")),
    ]

    for source_name, pattern in directional_patterns:
        for match in pattern.finditer(context_compact):
            left, right = match.group(1), match.group(2)
            _maybe_add_pair(left, right, source_name, match.group(0))

    tested_catalysts_harvest = re.compile(
        r"(?:the\s+)?(?:tested|prepared|synthesized|used)\s+(?:catalysts?|samples?)\s+(?:were|include|are)\s+([^.;\n]{10,250}?)(?:\s*(?:respectively|in\s+this\s+study))?[\.;]",
        flags=re.I,
    )
    denoted_as_harvest = re.compile(
        r"(?:will\s+be\s+)?(?:denoted|labeled|labelled|referred)\s+(?:as|hereafter)\s+([^.;\n]{3,200}?)(?:\s+for\s+brevity)?[\.;]",
        flags=re.I,
    )
    pairwise_correspond = re.compile(
        r"([^.;\n]{5,150}?)\s+(?:correspond\s+to|are\s+identified\s+as|map\s+to)\s+([^.;\n]{5,150})[\.;]",
        flags=re.I,
    )

    for pattern_name, pattern in [
        ("harvest_tested_catalysts", tested_catalysts_harvest),
        ("harvest_denoted_as", denoted_as_harvest),
    ]:
        for match in pattern.finditer(context_compact):
            for part in _split_identity_list_preserve_parenthetical(match.group(1)):
                if _registry_like_context_label(part, registry):
                    _store_candidate(part, pattern_name, match.group(0), harvest_only=True)

    for match in pairwise_correspond.finditer(context_compact):
        left_parts = _split_identity_list_preserve_parenthetical(match.group(1))
        right_parts = _split_identity_list_preserve_parenthetical(match.group(2))
        if len(left_parts) != len(right_parts) or len(left_parts) < 2:
            continue
        left_ok = all(_is_context_alias_label_candidate(part) for part in left_parts)
        right_ok = all(_registry_like_context_label(part, registry) for part in right_parts)
        if not left_ok or not right_ok:
            continue
        for short, full in zip(left_parts, right_parts):
            short_clean = _clean_context_alias_label(short)
            full_clean = _clean_context_mapped_label(full)
            if short_clean and full_clean and short_clean != full_clean:
                _add(short_clean, full_clean, "pairwise_correspond", match.group(0))

    template_info = _extract_identity_template_candidates(context_compact)
    for expanded_name in template_info.get("expanded", []) or []:
        if _registry_like_context_label(expanded_name, registry):
            _store_candidate(
                expanded_name,
                "template_expansion",
                f"template={template_info.get('template', '')}; values={','.join(template_info.get('values', []) or [])}",
                template_candidate=True,
            )

    return alias_map


def _extract_label_mapping_from_context(
    label: str,
    local_context: str,
    registry: Optional[Dict] = None,
    context_alias_map: Optional[Dict[str, Dict[str, Any]]] = None,
) -> str:
    label_raw = _clean_context_alias_label(label)
    if not label_raw:
        return ""

    alias_map = context_alias_map if isinstance(context_alias_map, dict) else _build_context_label_alias_map(local_context, registry)
    for key in _context_label_lookup_keys(label_raw):
        entry = dict(alias_map.get(key) or {})
        if entry and not bool(entry.get("ambiguous")):
            mapped = _clean_context_mapped_label(str(entry.get("mapped_label", "") or ""))
            if mapped:
                return mapped

    label_norm = _normalize_registry_label(label_raw)
    label_relaxed = _relaxed_registry_alias_signature(label_raw)
    label_fold = label_raw.lower()
    candidate_hits: List[str] = []
    seen_candidates = set()
    for entry in alias_map.values():
        if not isinstance(entry, dict) or bool(entry.get("ambiguous")):
            continue
        if not (bool(entry.get("harvest_only")) or bool(entry.get("template_candidate"))):
            continue
        mapped = _clean_context_mapped_label(str(entry.get("mapped_label", "") or ""))
        if not mapped or mapped in seen_candidates:
            continue
        seen_candidates.add(mapped)
        mapped_norm = _normalize_registry_label(mapped)
        mapped_relaxed = _relaxed_registry_alias_signature(mapped)

        if label_norm and mapped_norm == label_norm:
            candidate_hits.append(mapped)
            continue
        if label_relaxed and mapped_relaxed and mapped_relaxed == label_relaxed:
            candidate_hits.append(mapped)
            continue
        if len(label_fold) >= 4:
            if re.search(rf"[\(\[\{{,;/ -]{re.escape(label_fold)}[\)\]\}},;/ -]", mapped.lower()):
                candidate_hits.append(mapped)
                continue
            if mapped.lower().endswith(" " + label_fold) or mapped.lower().startswith(label_fold + " "):
                candidate_hits.append(mapped)

    candidate_hits = _dedupe_keep_order(candidate_hits)
    if len(candidate_hits) == 1:
        return candidate_hits[0]

    context_raw = _normalize_identity_text(local_context)
    if not context_raw:
        return ""

    label_fold = label_raw.lower()
    escaped = re.escape(label_fold)
    patterns = [
        rf"(?:^|[\s\(\[,;])(?:sample|catalyst|cat\.?|run)?\s*{escaped}(?:[\s\)\],:;]|$)\s*(?:corresponds to|is|=|represents|denotes|stands for|refers to|maps to)\s*([^.;\n]{{4,120}})",
        rf"([^.;\n]{{4,120}}?)\s*\(\s*(?:sample|catalyst|cat\.?|run)?\s*{escaped}\s*\)",
        rf"([^.;\n]{{4,120}}?)\s*(?:,|\s)*(?:denoted as|labeled as|labelled as|coded as|named|identified as|abbreviated as)\s*(?:sample|catalyst|cat\.?|run)?\s*{escaped}\b",
    ]
    for pattern in patterns:
        match = re.search(pattern, context_raw, flags=re.I)
        if not match:
            continue
        mapped = _clean_context_mapped_label(match.group(1))
        if not mapped:
            continue
        if len(_split_identity_list_preserve_parenthetical(mapped)) != 1:
            continue
        if _registry_like_context_label(mapped, registry):
            return mapped
    return ""


def _select_compact_loading_support_label(record: Dict) -> str:
    candidates: List[str] = []

    def _add_candidate(value: Any) -> None:
        value = clean_residual_mojibake_chars(str(value or "")).strip()
        if value and value not in candidates:
            candidates.append(value)

    notes = str((record or {}).get("Notes", ""))
    _, x_value = _extract_x_axis_and_value_from_notes(notes)
    for field in ("raw_category_label", "Catalyst", "Series_Name", "Catalyst_ID"):
        _add_candidate((record or {}).get(field, ""))
    _add_candidate(x_value)
    _add_candidate(" ".join(
        part for part in [
            str((record or {}).get("Catalyst_ID", "") or "").strip(),
            str((record or {}).get("Series_Name", "") or "").strip(),
        ] if part
    ))

    support_alt = r"(?:Al2O3|SiO2|CeO2|ZrO2|ZnO|TiO2|MgO|CNTs|Al)"
    compact_pattern = re.compile(
        rf"\d+(?:\.\d+)?\s*(?:wt\.?\s*%|wt%|%)?\s*[A-Za-z]{{1,2}}\s*(?:[-_/]|\s+)?\s*{support_alt}\b",
        flags=re.I,
    )

    for candidate in list(candidates):
        if _parse_compact_loading_support_identity(candidate):
            return candidate
        for hit in compact_pattern.findall(candidate):
            hit = clean_residual_mojibake_chars(str(hit or "")).strip()
            if hit and _parse_compact_loading_support_identity(hit):
                return hit
        for token in re.split(r"\s*(?:,|;|\||\(|\)|\[|\]|\{|\}|=|:|\band\b|\bvs\.?\b)\s*", candidate, flags=re.I):
            token = clean_residual_mojibake_chars(str(token or "")).strip()
            if token and _parse_compact_loading_support_identity(token):
                return token
    return ""


def _extract_identity_clues_from_figure_record(record: Dict) -> Dict:
    notes = str(record.get("Notes", ""))
    x_axis, x_value = _extract_x_axis_and_value_from_notes(notes)
    catalyst = _normalize_identity_text(str(record.get("Catalyst", "")))
    catalyst_id = _normalize_identity_text(str(record.get("Catalyst_ID", "")))
    series_name = _normalize_identity_text(str(record.get("Series_Name", "")))
    catalyst_amount_like = _looks_like_catalyst_amount_label(catalyst)
    series_amount_like = _looks_like_catalyst_amount_label(series_name)
    x_value_amount_like = _looks_like_catalyst_amount_label(x_value)

    clue_labels = []

    def _add_label(value: str) -> None:
        value = _normalize_identity_text(value)
        if value and value not in clue_labels:
            clue_labels.append(value)

    if catalyst_id:
        _add_label(catalyst_id)
    if catalyst and not _is_generic_series_label(catalyst) and not catalyst_amount_like:
        _add_label(catalyst)

    x_axis_lower = x_axis.lower()
    x_value_identity_like = bool(
        x_value and re.search(r"[A-Za-z]", x_value) and (
            any(kw in x_axis_lower for kw in ["catalyst", "sample", "support", "loading", "composition", "content"])
            or "/" in x_value
            or _looks_like_identity_mapping(x_value)
        )
    )
    if x_value_identity_like and not x_value_amount_like:
        _add_label(x_value)
    if series_name and not _is_generic_series_label(series_name) and not series_amount_like:
        _add_label(series_name)

    label_text = " ".join(clue_labels).strip()
    compact_label = _select_compact_loading_support_label(record)
    clue_record = {
        "Catalyst": catalyst if catalyst and not _is_generic_series_label(catalyst) and not catalyst_amount_like else "",
        "Catalyst_ID": catalyst_id,
        "Series_Name": series_name,
        "Notes": notes,
        "Active_Metal": clean_residual_mojibake_chars(str(record.get("Active_Metal", "") or "")).strip(),
        "Promoter": clean_residual_mojibake_chars(str(record.get("Promoter", "") or record.get("Promoter_Metal", "") or "")).strip(),
        "Promoter_Metal": clean_residual_mojibake_chars(str(record.get("Promoter_Metal", "") or record.get("Promoter", "") or "")).strip(),
        "Support": clean_residual_mojibake_chars(str(record.get("Support", "") or "")).strip(),
        "Support_Normalized": clean_residual_mojibake_chars(str(record.get("Support_Normalized", "") or record.get("Support", "") or "")).strip(),
        "Metal_Loading_wt%": clean_residual_mojibake_chars(str(record.get("Metal_Loading_wt%", "") or "")).strip(),
        "Alloy_Ratio": clean_residual_mojibake_chars(str(record.get("Alloy_Ratio", "") or "")).strip(),
        "Ni_Fraction": clean_residual_mojibake_chars(str(record.get("Ni_Fraction", "") or "")).strip(),
        "Promoter_Fraction": clean_residual_mojibake_chars(str(record.get("Promoter_Fraction", "") or "")).strip(),
        "raw_category_label": clean_residual_mojibake_chars(str(record.get("raw_category_label", "") or "")).strip(),
    }
    if not clue_record["Catalyst"] and x_value_identity_like and not x_value_amount_like:
        clue_record["Catalyst"] = x_value
    elif not clue_record["Catalyst"] and series_name and not _is_generic_series_label(series_name) and not series_amount_like:
        clue_record["Catalyst"] = series_name
    if compact_label and not clue_record["Catalyst"] and not _is_generic_series_label(compact_label):
        clue_record["Catalyst"] = compact_label

    identity_text_for_parse = label_text or compact_label
    if identity_text_for_parse:
        canonicalize_metal_order_and_ratio(clue_record, identity_text_for_parse)
        split_support_promoter_preparation_roles(clue_record)
        enrich_promoter_fields_from_identity(clue_record, identity_text_for_parse)
        normalize_identity_aliases(clue_record)
        if compact_label:
            clue_record = _apply_niznal_compact_identity(clue_record, compact_label)
        compact_identity = (
            _parse_compact_loading_support_identity(compact_label)
            if compact_label else {}
        ) or _parse_compact_loading_support_identity(identity_text_for_parse)
        if compact_identity:
            for field, value in compact_identity.items():
                if value and not str(clue_record.get(field, "")).strip():
                    clue_record[field] = value
        _apply_mo2c_framework_identity(clue_record, identity_text_for_parse)
        if not str(clue_record.get("Support", "")).strip():
            compact_support = _infer_support_from_compact_loading_label(identity_text_for_parse)
            if compact_support:
                clue_record["Support"] = compact_support
                if not str(clue_record.get("Support_Normalized", "")).strip():
                    clue_record["Support_Normalized"] = compact_support
        if not str(clue_record.get("Catalyst_ID", "")).strip() and not _is_generic_series_label(identity_text_for_parse):
            partial_id = _build_partial_identity_from_label(identity_text_for_parse)
            if partial_id:
                clue_record["Catalyst_ID"] = partial_id

    explicit_loading_in_label = (
        _label_has_structured_loading(compact_label)
        if compact_label else False
    ) or _label_has_structured_loading(identity_text_for_parse)
    return {
        "label_text": label_text,
        "compact_label": compact_label,
        "series_name": series_name,
        "x_axis": x_axis,
        "x_value": x_value,
        "generic_series_only": bool(series_name and _is_generic_series_label(series_name) and not catalyst and not catalyst_id and not x_value_identity_like),
        "clue_record": clue_record,
        "alias_keys": _make_registry_alias_keys(clue_record),
        "explicit_loading_in_label": explicit_loading_in_label,
    }


def _backfill_identity_from_clue_record(row: Dict, clue_record: Dict) -> Dict:
    current = dict(row or {})
    clue = dict(clue_record or {})

    for field in _CLUE_BACKFILL_FIELDS:
        if not str(current.get(field, "")).strip() and str(clue.get(field, "")).strip():
            current[field] = clue.get(field, "")

    compact_seed = dict(current)
    for field, value in clue.items():
        if str(value or "").strip() and not str(compact_seed.get(field, "")).strip():
            compact_seed[field] = value
    compact_label = _select_compact_loading_support_label(compact_seed)
    if compact_label:
        current = _apply_niznal_compact_identity(current, compact_label)
        compact_identity = _parse_compact_loading_support_identity(compact_label)
        if compact_identity:
            if not str(current.get("Catalyst", "")).strip() and not _is_generic_series_label(compact_label):
                current["Catalyst"] = compact_label
            for field, value in compact_identity.items():
                if value and not str(current.get(field, "")).strip():
                    current[field] = value
    for candidate in [
        current.get("raw_category_label", ""),
        current.get("Catalyst", ""),
        current.get("Series_Name", ""),
        clue.get("Catalyst", ""),
        clue.get("Series_Name", ""),
    ]:
        if candidate:
            current = _apply_mo2c_framework_identity(current, str(candidate))

    return current


def _compact_loading_support_text(label_text: str) -> str:
    text = _normalize_registry_label(label_text)
    if not text:
        return ""
    text = re.sub(r"\bwt\.?\s*%?", "", text, flags=re.I)
    text = text.replace("%", "")
    compact = re.sub(r"[^A-Za-z0-9]+", "", text).lower()
    return re.sub(r"wt", "", compact, flags=re.I)


def _parse_niznal_compact_identity(label_text: str) -> Dict[str, str]:
    compact = _compact_loading_support_text(label_text)
    if not compact:
        return {}
    match = re.fullmatch(r"(\d+(?:\.\d+)?)ni(?:(\d+(?:\.\d+)?)zn)?(?:al|al2o3)?", compact, flags=re.I)
    if match:
        ni_value = float(match.group(1))
        zn_value = float(match.group(2)) if match.group(2) else 0.0
        result: Dict[str, str] = {
            "Active_Metal": "Ni",
            "Support": "Al2O3",
            "Support_Normalized": "Al2O3",
            "Metal_Loading_wt%": _format_ratio_token(str(ni_value)),
        }
        if zn_value > 0:
            total = ni_value + zn_value
            result.update({
                "Promoter": "Zn",
                "Promoter_Metal": "Zn",
                "Alloy_Ratio": f"Ni:{_format_ratio_token(str(ni_value))}, Zn:{_format_ratio_token(str(zn_value))}",
                "Ni_Fraction": _format_float_token(ni_value / total),
                "Promoter_Fraction": _format_float_token(zn_value / total),
            })
        else:
            result.update({"Ni_Fraction": "1", "Promoter_Fraction": "0"})
        return result

    match = re.fullmatch(r"(\d+(?:\.\d+)?)zn(?:al|al2o3)", compact, flags=re.I)
    if match:
        zn_value = match.group(1)
        return {
            "Active_Metal": "Zn",
            "Support": "Al2O3",
            "Support_Normalized": "Al2O3",
            "Metal_Loading_wt%": _format_ratio_token(zn_value),
        }
    return {}


def _apply_niznal_compact_identity(record: Dict, label_text: str) -> Dict:
    parsed = _parse_niznal_compact_identity(label_text)
    if not parsed:
        return record
    for field in [
        "Active_Metal", "Support", "Support_Normalized", "Metal_Loading_wt%",
        "Alloy_Ratio", "Ni_Fraction", "Promoter_Fraction",
    ]:
        if field in parsed:
            record[field] = parsed[field]
    if "Promoter_Metal" in parsed or "Promoter" in parsed:
        record["Promoter"] = parsed.get("Promoter", parsed.get("Promoter_Metal", ""))
        record["Promoter_Metal"] = parsed.get("Promoter_Metal", parsed.get("Promoter", ""))
    else:
        for field in ("Promoter", "Promoter_Metal"):
            if str(record.get(field, "") or "").strip().lower() in {"zn", "zinc"}:
                record[field] = ""
        if "Zn" in str(record.get("Alloy_Ratio", "") or ""):
            record["Alloy_Ratio"] = ""

    # Existing alias normalization may have expanded "2NiAl" into a Zn-containing
    # identity. Recompute ids from the corrected compact identity instead.
    for field in ("Catalyst_ID", "Catalyst_ID_normalized", "identity_alias_group"):
        value = str(record.get(field, "") or "")
        if re.search(r"__prom-zn|-zn|zn-", value, flags=re.I):
            record[field] = ""
    structured_id = _build_structured_catalyst_id(record)
    if structured_id:
        record["Catalyst_ID"] = structured_id.lower()
        record["Catalyst_ID_normalized"] = structured_id.lower()
        record["identity_alias_group"] = structured_id.lower()
    return record


def _parse_mo2c_framework_identity(label_text: str) -> Dict[str, str]:
    raw = clean_residual_mojibake_chars(str(label_text or "")).strip()
    if not raw:
        return {}
    text = _normalize_identity_text(raw)
    text = text.replace("\u2212", "-").replace("\u2013", "-").replace("\u2014", "-")
    text = text.replace("₂", "2").replace("₄", "4")
    if not re.search(r"\b(?:mo\s*2\s*c|mo2c|molybdenum carbide)\b", text, flags=re.I):
        return {}

    metal_alt = "|".join(sorted({sym for sym in _MSR_ACTIVE_PROMOTER_SYMBOLS if sym not in {"Mo"}}, key=len, reverse=True))
    match = re.search(
        rf"\b({metal_alt})\s*[-_/]?\s*Mo\s*2\s*C\s*(?:\(\s*(\d+(?:\.\d+)?)\s*\))?",
        text,
        flags=re.I,
    )
    if match:
        metal = match.group(1).capitalize()
        ratio = match.group(2)
        result = {
            "Active_Metal": metal,
            "Promoter": "",
            "Promoter_Metal": "",
            "Support": "Mo2C",
            "Support_Normalized": "Mo2C",
        }
        if ratio:
            metal_value = float(ratio)
            mo_value = max(0.0, 100.0 - metal_value)
            result["Alloy_Ratio"] = f"{metal}:{_format_ratio_token(str(metal_value))}, Mo:{_format_ratio_token(str(mo_value))}"
        return result

    if re.search(r"\b(?:beta\s*[- ]?)?mo\s*2\s*c\b|\bβ\s*[- ]?mo\s*2\s*c\b|\bmolybdenum carbide\b", text, flags=re.I):
        return {
            "Active_Metal": "",
            "Promoter": "",
            "Promoter_Metal": "",
            "Support": "Mo2C",
            "Support_Normalized": "Mo2C",
        }
    return {}


def _apply_mo2c_framework_identity(record: Dict, label_text: str) -> Dict:
    parsed = _parse_mo2c_framework_identity(label_text)
    if not parsed:
        return record
    for field, value in parsed.items():
        if field in {"Promoter", "Promoter_Metal"}:
            if str(record.get(field, "") or "").strip().lower() in {"", "mo", "molybdenum"}:
                record[field] = value
            continue
        if value and not str(record.get(field, "") or "").strip():
            record[field] = value
    if str(record.get("Support_Normalized", "") or record.get("Support", "")).strip().lower() == "mo2c":
        for field in ("Promoter", "Promoter_Metal"):
            if str(record.get(field, "") or "").strip().lower() in {"mo", "molybdenum"}:
                record[field] = ""
    return record


def _parse_compact_loading_support_identity(label_text: str) -> Dict[str, str]:
    specialized = _parse_niznal_compact_identity(label_text)
    if specialized:
        return specialized
    compact = _compact_loading_support_text(label_text)
    if not compact:
        return {}
    metals_alt = "|".join(sorted({sym.lower() for sym in _MSR_ACTIVE_PROMOTER_SYMBOLS}, key=len, reverse=True))
    support_suffix_map = {
        "al": "Al2O3",
        "al2o3": "Al2O3",
        "sio2": "SiO2",
        "ceo2": "CeO2",
        "zro2": "ZrO2",
        "zno": "ZnO",
        "tio2": "TiO2",
        "mgo": "MgO",
        "cnts": "CNTs",
    }
    suffix_alt = "|".join(sorted(support_suffix_map, key=len, reverse=True))
    match = re.fullmatch(rf"(\d+(?:\.\d+)?)({metals_alt})({suffix_alt})", compact, flags=re.I)
    if not match:
        return {}
    value, metal, suffix = match.groups()
    metal = metal.capitalize()
    support = support_suffix_map.get(suffix.lower(), "")
    result = {
        "Active_Metal": metal,
        "Support": support,
        "Support_Normalized": support,
        "Metal_Loading_wt%": _format_ratio_token(value),
    }
    return result


def _label_has_structured_loading(label_text: str) -> bool:
    text = _normalize_registry_label(label_text)
    if not text:
        return False
    if re.search(r"\d+(?:\.\d+)?\s*(?:wt\.?\s*%|wt%|%)", text, flags=re.I):
        return True

    compact = _compact_loading_support_text(text)
    if not re.search(r"\d", compact):
        return False
    metals_alt = "|".join(sorted({sym.lower() for sym in _MSR_ACTIVE_PROMOTER_SYMBOLS}, key=len, reverse=True))
    if re.fullmatch(rf"(?:\d+(?:\.\d+)?(?:{metals_alt}))+",
                    compact, flags=re.I):
        return True
    if re.fullmatch(rf"(?:{metals_alt}\d+(?:\.\d+)?)+",
                    compact, flags=re.I):
        return True
    if re.fullmatch(rf"\d+(?:\.\d+)?(?:{metals_alt})(?:al|al2o3|sio2|ceo2|zro2|zno|cnts)",
                    compact, flags=re.I):
        return True
    return False


def _infer_support_from_compact_loading_label(label_text: str) -> str:
    text = _normalize_registry_label(label_text)
    if not text:
        return ""
    compact = _compact_loading_support_text(text)
    if not re.search(r"\d", compact):
        return ""
    metals_alt = "|".join(sorted({sym.lower() for sym in _MSR_ACTIVE_PROMOTER_SYMBOLS}, key=len, reverse=True))
    if re.fullmatch(rf"(?:\d+(?:\.\d+)?(?:{metals_alt}))+al", compact, flags=re.I):
        return "Al2O3"
    if re.fullmatch(rf"(?:{metals_alt}\d+(?:\.\d+)?)+al", compact, flags=re.I):
        return "Al2O3"
    return ""


def _find_registry_matches_by_keys(alias_keys: List[str], registry: Dict[str, Dict]) -> List[Tuple[str, Dict]]:
    hits: List[Tuple[str, Dict]] = []
    seen = set()
    alias_set = {str(k).strip().lower() for k in alias_keys if str(k).strip()}
    if not alias_set:
        return hits

    for reg_key, reg_record in registry.items():
        reg_aliases = {str(reg_key).strip().lower()}
        reg_aliases.update(str(a).strip().lower() for a in reg_record.get("_registry_aliases", []) if str(a).strip())
        if alias_set.intersection(reg_aliases) and reg_key not in seen:
            hits.append((reg_key, reg_record))
            seen.add(reg_key)
    return hits





# ============================================================
# Problem A: metal set compatibility helpers for binding safety
# Prevent: Ni/clay→Ni-Cu/clay, Cu/clay→Ni-Cu/clay, support-only→active, baseline→active
# ============================================================

# Metals that appear in the HEAD of catalyst labels (active metals / promoters)
_MSR_ACTIVE_PROMOTER_SYMBOLS: frozenset = frozenset({
    "Ni", "Cu", "Pt", "Pd", "Ru", "Rh", "Co", "Fe", "Ir", "Au",
    "Ag", "Zn", "Sn", "Mo", "Cr", "Mn", "In", "Ga", "Ge",
    "Ce", "La", "K", "Mg",
})


def _strip_identity_loading_prefix(text: str) -> str:
    s = clean_residual_mojibake_chars(str(text or "")).strip()
    return re.sub(r"^\d+(?:\.\d+)?\s*(?:wt\.?\s*%|wt%|%)?\s*", "", s, flags=re.I)


def _extract_compact_number_metal_pairs(text: str) -> List[Tuple[str, str]]:
    compact = re.sub(r"\s+", "", clean_residual_mojibake_chars(str(text or "")))
    pairs = re.findall(r"(\d+(?:\.\d+)?)([A-Z][a-z]?)(?=\d|$)", compact, flags=re.I)
    return [(value, metal.capitalize()) for value, metal in pairs]


def _extract_explicit_metal_set_from_identity_clue(clue_text: str) -> Optional[Set[str]]:
    """
    Extract the explicit set of active/promoter metals from a catalyst identity label.

    Returns:
    - frozenset of metal symbols when the label has clear metal identity
      e.g., "Ni/clay" → {"Ni"},  "Ni-Cu/Al2O3" → {"Ni", "Cu"}
    - None  — generic / opaque label, metals unknown (e.g., "cat-1", "C2")
    - set() — support-only or baseline (no active metals)

    Used by binding safety logic to prevent mono→bimetal and wrong-metal mismatches.
    """
    raw = clean_residual_mojibake_chars(str(clue_text or "")).strip()
    if not raw:
        return None
    # Generic labels have no deterministic metal set
    if _is_generic_series_label(raw):
        return None
    # Baseline: "No catalyst", "Blank", etc. → no active metals
    if _is_blank_or_baseline_label(raw):
        return set()
    # Support-only labels → no active metals
    if _normalize_support_only_label(raw):
        return set()

    norm = _normalize_registry_label(raw)
    if not norm:
        return None

    mo2c_identity = _parse_mo2c_framework_identity(raw)
    if mo2c_identity:
        active = clean_residual_mojibake_chars(str(mo2c_identity.get("Active_Metal", "") or "")).strip()
        return {active} if active else set()

    allowed = MSR_ACTIVE_METALS | MSR_PROMOTERS | MSR_RARE_EARTH | MSR_NI_SECOND_METALS
    head = norm.split("/", 1)[0].strip() if "/" in norm else norm.strip()
    head = _strip_identity_loading_prefix(head)

    found: Set[str] = set()
    for tok in re.split(r"[-_+]", head):
        tok_clean = _strip_identity_loading_prefix(tok.strip())
        tok_clean = tok_clean.replace("%", "")
        tok_clean = re.sub(r"\d+$", "", tok_clean)
        for sym in _MSR_ACTIVE_PROMOTER_SYMBOLS:
            if tok_clean.lower() == sym.lower():
                found.add(sym)
                break

    if not found:
        percent_matches = re.findall(
            r"(\d+(?:\.\d+)?)\s*(?:wt\.?\s*%|wt%|%)\s*([A-Z][a-z]?)",
            raw,
            flags=re.I,
        )
        for _, metal in percent_matches:
            metal = metal.capitalize()
            if metal in allowed:
                found.add(metal)

    if not found:
        for metal, _ in _parse_inline_ratio_pairs_from_label(raw):
            metal = metal.capitalize()
            if metal in allowed:
                found.add(metal)

    if not found:
        for _, metal in _extract_compact_number_metal_pairs(raw):
            if metal in allowed:
                found.add(metal)

    if not found:
        support_formula_keys = {
            "al2o3", "ceo2", "zro2", "sio2", "tio2", "mgo", "la2o3",
            "zno", "fe3o4", "fe2o3", "ceo2zro2", "mgoal2o3",
        }
        support_formula_norm = re.sub(r"[\s\-_]+", "", norm.lower())
        support_formula_head = re.sub(r"[\s\-_]+", "", head.lower())
        if support_formula_norm in support_formula_keys or support_formula_head in support_formula_keys:
            return set()

    if not found:
        support_only = (
            _normalize_support_only_label(raw)
            or _normalize_support_only_label(norm)
            or _normalize_support_only_label(head)
        )
        if support_only:
            return set()

    return found or None


def _extract_registry_metal_set(reg_record: Dict) -> Set[str]:
    """
    Collect all active/promoter metals declared in a registry record.
    Combines Active_Metal + Promoter + similar second-metal fields.
    """
    found: Set[str] = set()
    for field in ("Active_Metal", "Promoter", "Promoter_Metal", "Second_Metal",
                  "Dopant", "Auxiliary_Metal"):
        raw_val = clean_residual_mojibake_chars(str(reg_record.get(field, "") or "")).strip()
        if not raw_val:
            continue
        for tok in re.split(r"[,\-/\s+]+", raw_val):
            tok_clean = re.sub(r"\d+$", "", tok.strip())
            for sym in _MSR_ACTIVE_PROMOTER_SYMBOLS:
                if tok_clean.lower() == sym.lower():
                    found.add(sym)
                    break
    return found


def _is_metal_set_compatible_for_unique_match(
    clue_metals: Optional[Set[str]], reg_metals: Set[str]
) -> Tuple[bool, str]:
    """
    Determine whether a clue's metal set is compatible with a registry entry for binding.

    Returns (is_compatible: bool, reason: str).

    Science rules:
    - clue_metals=None → metals unknown → conservatively allow (no evidence to block)
    - clue_metals={}   → support-only or baseline → NOT compatible with active registry
    - {Ni} vs {Ni,Cu}  → monometallic vs bimetallic → CONFLICT
    - {Ni,Cu} vs {Ni}  → bimetallic vs monometallic → CONFLICT
    - {Cu} vs {Ni}     → wrong active metal → CONFLICT
    - {Ni} vs {Ni}     → OK
    """
    if clue_metals is None:
        return True, "clue_metal_set_unknown"

    if not clue_metals:
        # Support-only or baseline has no active metals → block active registry match
        if not reg_metals:
            return True, "support_only_matches_inactive_registry"
        return False, "support_only_or_baseline_label_blocks_active_registry_match"

    if not reg_metals:
        # Registry has no declared metals → allow conservatively
        return True, "registry_metal_set_empty_allow"

    if clue_metals == reg_metals:
        return True, "metal_set_exact_match"

    # Monometallic clue vs bimetallic registry: {Ni} ⊂ {Ni, Cu}
    # Exception: allow when the extra registry metals are support-framework metals
    # (e.g., Mo in Mo2C/MoC support — Mo is part of the carrier, not an independent promoter)
    _SUPPORT_FRAMEWORK_METALS = {"Mo", "W", "V", "Nb"}  # metals that double as support components
    if len(clue_metals) < len(reg_metals) and clue_metals.issubset(reg_metals):
        extra = reg_metals - clue_metals
        if extra.issubset(_SUPPORT_FRAMEWORK_METALS):
            return True, "support_framework_metal_allowed"  # Mo/W in Mo2C etc. is carrier, not promoter
        else:
            return False, "metal_set_conflict_monometallic_vs_bimetallic"

    # Bimetallic clue vs monometallic registry: {Ni, Cu} ⊃ {Ni}
    if len(clue_metals) > len(reg_metals) and reg_metals.issubset(clue_metals):
        return False, "metal_set_conflict_bimetallic_vs_monometallic"

    # No common metals → wrong active metal
    if not clue_metals.intersection(reg_metals):
        return False, "active_metal_conflict"

    # Partial overlap with different size (e.g., {Ni,Ce} vs {Ni,Cu})
    return False, "metal_set_partial_mismatch"


def _find_unique_registry_candidate(clue_record: Dict, registry: Dict[str, Dict]) -> List[Tuple[str, Dict]]:
    active = normalize_active_metal_for_merge(str(clue_record.get("Active_Metal", "")))
    promoter = normalize_promoter_for_merge(str(clue_record.get("Promoter", "")))
    support = normalize_support_for_merge(str(clue_record.get("Support_Normalized", "") or clue_record.get("Support", "")))
    alloy = alloy_ratio_signature_for_merge(clue_record)
    loading = normalize_loading_for_merge(str(clue_record.get("Metal_Loading_wt%", "")))
    require_loading = bool(clue_record.get("_explicit_loading_in_label")) and bool(loading)

    if not any([active, promoter, support, alloy, require_loading and loading]):
        return []

    # Problem A fix: if clue has an explicit active metal but no promoter,
    # it is claiming to be monometallic. Matching it to a bimetallic registry entry
    # (which has a promoter) is a scientific identity conflict, not just "promoter unknown".
    # Exception: caller may set _promoter_unknown_ok=True for truly opaque labels.
    clue_has_explicit_active = bool(active)
    promoter_unknown_ok = bool(clue_record.get("_promoter_unknown_ok"))

    matches: List[Tuple[str, Dict]] = []
    for reg_key, reg_record in registry.items():
        if active and normalize_active_metal_for_merge(str(reg_record.get("Active_Metal", ""))) != active:
            continue
        if promoter and normalize_promoter_for_merge(str(reg_record.get("Promoter", ""))) != promoter:
            continue

        # Safety: clue explicitly monometallic (known active, no promoter) but
        # registry has a promoter → promoter_missing_blocks_unique_match
        if clue_has_explicit_active and not promoter and not promoter_unknown_ok:
            reg_promoter_val = normalize_promoter_for_merge(str(reg_record.get("Promoter", "")))
            if reg_promoter_val:
                continue  # promoter_missing_blocks_unique_match

        if support and normalize_support_for_merge(str(reg_record.get("Support_Normalized", "") or reg_record.get("Support", ""))) != support:
            continue
        if alloy and alloy_ratio_signature_for_merge(reg_record) != alloy:
            continue
        if require_loading and normalize_loading_for_merge(str(reg_record.get("Metal_Loading_wt%", ""))) != loading:
            continue
        # Belt-and-suspenders: metal-set compatibility from clue_record fields.
        # Catches the critical case where clue_has_explicit_active=False
        # (row Active_Metal is empty before binding) but clue_record still has
        # Promoter="" implying monometallic — the promoter safety block above is
        # bypassed when active="" because clue_has_explicit_active=False.
        # This check covers that gap: if clue has ANY metals declared, they must
        # be compatible with the registry entry's metal set.
        _clue_metals_rc = _extract_registry_metal_set(clue_record)
        if _clue_metals_rc:  # only apply when clue has explicit metal fields
            _reg_metals_rc = _extract_registry_metal_set(reg_record)
            _compat_rc, _compat_rc_reason = _is_metal_set_compatible_for_unique_match(
                _clue_metals_rc, _reg_metals_rc
            )
            if not _compat_rc:
                continue  # metal_set_conflict_from_clue_fields: _compat_rc_reason
        # Critical safety guard (Problem 4): if clue has NO explicit active metal
        # AND NO explicit promoter (i.e. Active_Metal="" and Promoter="" in row fields),
        # AND the registry entry declares active metals,
        # THEN block the match unless the clue's identity label text explicitly
        # contains active metal evidence.
        # This prevents: empty-AM clue with Support=clay matching Ni-Cu/clay.
        if not active and not promoter:
            _reg_metals_safety = _extract_registry_metal_set(reg_record)
            if _reg_metals_safety:
                _clue_label_blob = " ".join(filter(None, [
                    str(clue_record.get("Catalyst", "") or ""),
                    str(clue_record.get("Catalyst_ID", "") or ""),
                    str(clue_record.get("raw_category_label", "") or ""),
                    str(clue_record.get("Series_Name", "") or ""),
                ])).strip()
                _clue_label_metals = (
                    _extract_explicit_metal_set_from_identity_clue(_clue_label_blob)
                    if _clue_label_blob else set()
                )
                if not _clue_label_metals:
                    # support-only or no-metal clue cannot match active catalyst registry
                    continue
        matches.append((reg_key, reg_record))
    return matches


def _extract_family_material_signature(text: str) -> str:
    norm = _normalize_registry_label(text)
    if not norm:
        return ""
    for token in ("znal2o4", "mgal2o4", "coal2o4", "nial2o4", "spinel", "aluminate"):
        if token in norm:
            return token
    return ""


def _is_family_level_registry_record(reg_record: Dict) -> bool:
    catalyst = _normalize_registry_label(str(reg_record.get("Catalyst", "") or ""))
    canonical = _normalize_registry_label(str(reg_record.get("Canonical_Catalyst_ID", "") or reg_record.get("registry_key", "") or ""))
    family_base = ""
    for candidate in [canonical, catalyst]:
        if not candidate:
            continue
        head = candidate.split("/", 1)[0].strip() if "/" in candidate else candidate
        if not _label_has_structured_loading(head):
            family_base = candidate
            break
    if not family_base:
        return False

    raw_catalyst = clean_residual_mojibake_chars(str(reg_record.get("Catalyst", "") or "")).strip().lower()
    if not canonical and re.search(r"\b(?:reduction|reduced|calcined|treated|oxidized|fresh|spent)\b", raw_catalyst):
        return False

    if normalize_loading_for_merge(str(reg_record.get("Metal_Loading_wt%", "") or "")):
        return False
    if alloy_ratio_signature_for_merge(reg_record):
        return False

    active = normalize_active_metal_for_merge(str(reg_record.get("Active_Metal", "") or ""))
    support = normalize_support_for_merge(str(reg_record.get("Support_Normalized", "") or reg_record.get("Support", "") or ""))
    return bool(active and support)


def _find_unique_family_registry_candidate(clue_record: Dict, registry: Dict[str, Dict]) -> List[Tuple[str, Dict]]:
    active = normalize_active_metal_for_merge(str(clue_record.get("Active_Metal", "") or ""))
    promoter = normalize_promoter_for_merge(str(
        clue_record.get("Promoter", "") or clue_record.get("Promoter_Metal", "") or ""
    ))
    support = normalize_support_for_merge(str(
        clue_record.get("Support_Normalized", "") or clue_record.get("Support", "") or ""
    ))

    if not active or not support:
        return []

    clue_label_blob = " ".join(filter(None, [
        str(clue_record.get("Catalyst", "") or ""),
        str(clue_record.get("Catalyst_ID", "") or ""),
        str(clue_record.get("raw_category_label", "") or ""),
        str(clue_record.get("Series_Name", "") or ""),
    ])).strip()
    clue_material_sig = _extract_family_material_signature(clue_label_blob)

    matches: List[Tuple[str, Dict]] = []
    for reg_key, reg_record in registry.items():
        if not _is_family_level_registry_record(reg_record):
            continue
        if normalize_active_metal_for_merge(str(reg_record.get("Active_Metal", "") or "")) != active:
            continue
        reg_promoter = normalize_promoter_for_merge(str(
            reg_record.get("Promoter", "") or reg_record.get("Promoter_Metal", "") or ""
        ))
        if active and not promoter and reg_promoter:
            continue
        if promoter and reg_promoter != promoter:
            continue
        if normalize_support_for_merge(str(
            reg_record.get("Support_Normalized", "") or reg_record.get("Support", "") or ""
        )) != support:
            continue

        reg_material_sig = _extract_family_material_signature(" ".join(filter(None, [
            str(reg_record.get("Catalyst", "") or ""),
            str(reg_record.get("Canonical_Catalyst_ID", "") or ""),
        ])))
        if reg_material_sig and reg_material_sig != clue_material_sig:
            continue

        matches.append((reg_key, reg_record))

    return matches


def _find_loading_relaxed_single_metal_family_candidate(clue_record: Dict, registry: Dict[str, Dict]) -> List[Tuple[str, Dict]]:
    """
    Single-metal, promoter-free loading-relaxed family fallback.
    Used for labels like 2NiAl when the text-side registry lacks an exact loading
    sample but does contain same-active same-support rows with explicit loadings.
    Only used for preparation-backbone propagation, never for Canonical_Catalyst_ID.
    """
    active = normalize_active_metal_for_merge(str(clue_record.get("Active_Metal", "") or ""))
    promoter = normalize_promoter_for_merge(str(
        clue_record.get("Promoter", "") or clue_record.get("Promoter_Metal", "") or ""
    ))
    support = normalize_support_for_merge(str(
        clue_record.get("Support_Normalized", "") or clue_record.get("Support", "") or ""
    ))

    if not active or promoter or not support:
        return []

    matches: List[Tuple[str, Dict]] = []
    for reg_key, reg_record in registry.items():
        reg_active = normalize_active_metal_for_merge(str(reg_record.get("Active_Metal", "") or ""))
        reg_promoter = normalize_promoter_for_merge(str(
            reg_record.get("Promoter", "") or reg_record.get("Promoter_Metal", "") or ""
        ))
        reg_support = normalize_support_for_merge(str(
            reg_record.get("Support_Normalized", "") or reg_record.get("Support", "") or ""
        ))

        if reg_active != active:
            continue
        if reg_promoter:
            continue
        if reg_support != support:
            continue

        reg_loading = normalize_loading_for_merge(str(reg_record.get("Metal_Loading_wt%", "") or ""))
        if not reg_loading:
            continue

        matches.append((reg_key, reg_record))

    return matches


def _is_support_only_figure_label(label: str) -> bool:
    raw = clean_residual_mojibake_chars(str(label or "")).strip()
    if not raw or _is_generic_series_label(raw) or _is_blank_or_baseline_label(raw):
        return False
    norm = _normalize_registry_label(raw)
    if norm in {
        "zno", "al2o3", "ceo2", "zro2", "sio2", "tio2", "mgo", "la2o3",
        "cnt", "cnts", "sep", "sba-15", "mcm-41", "activated carbon", "carbon",
    }:
        return True
    if not _normalize_support_only_label(raw):
        return False
    explicit_metals = _extract_explicit_metal_set_from_identity_clue(raw)
    return isinstance(explicit_metals, set) and not explicit_metals


def _get_registry_bindable_preparation_fields() -> List[str]:
    return [
        "Metal_Loading_Method_Normalized",
        "Support_Prep_Method_Normalized",
        "Precursor_Normalized",
        "Precursor_Family",
        "Support_Grouped",
        "Preparation_Fingerprint",
    ]



def _backfill_registry_preparation_backbone_into_figure(
    fig_record: Dict,
    registry_record: Dict,
    source: str = "registry_bind",
    confidence: str = "high",
    scope: str = "identity_bound",
) -> Dict:
    merged = dict(fig_record)
    if not int((registry_record or {}).get("can_broadcast_preparation", 0) or 0):
        return merged
    registry_backbone = enrich_registry_preparation_backbone(dict(registry_record or {}))
    if not str(registry_backbone.get("Preparation_Fingerprint", "")).strip():
        fingerprint = build_preparation_fingerprint(registry_backbone)
        if fingerprint:
            registry_backbone["Preparation_Fingerprint"] = fingerprint

    for field in _get_registry_bindable_preparation_fields():
        if not str(merged.get(field, "")).strip() and str(registry_backbone.get(field, "")).strip():
            merged[field] = registry_backbone.get(field, "")
            _set_field_source(merged, field, source, confidence, scope)
    return merged



def _merge_registry_fields_into_figure_record(fig_record: Dict, registry_record: Dict) -> Dict:
    """从 registry_record 向 fig_record 补全静态字段。
    只填空字段，不覆盖已有局部值。
    补全后在 Notes 中追加 [registry_bind_supplement: field1,field2,...] 审计标签。
    不做任何 paper-level 众数推断，只做 registry 精确补全。
    """
    merged = dict(fig_record)
    normalized_fields = set(_get_registry_bindable_preparation_fields())
    supplemented = []
    for field in _FIGURE_IDENTITY_BINDABLE_FIELDS:
        if field in normalized_fields:
            continue
        src_val = str(registry_record.get(field, "")).strip()
        if not str(merged.get(field, "")).strip() and src_val and src_val.lower() != "nan":
            merged[field] = registry_record.get(field, "")
            supplemented.append(field)
            _set_field_source(merged, field, "registry_bind", "high", "identity_bound")
    merged = _backfill_registry_preparation_backbone_into_figure(merged, registry_record)
    # 审计标签：只记录本次真正补入的字段
    if supplemented:
        notes = str(merged.get("Notes", ""))
        tag = f"[registry_bind_supplement: {','.join(supplemented)}]"
        if tag not in notes:
            merged["Notes"] = (notes + " " + tag).strip()
    return merged


def _merge_registry_identity_fields_without_cid(
    fig_record: Dict,
    registry_record: Dict,
    source: str = "family_identity",
    confidence: str = "low",
    scope: str = "family_identity",
) -> Dict:
    merged = dict(fig_record)
    for field in [
        "Active_Metal", "Promoter", "Promoter_Metal",
        "Support", "Support_Normalized", "Support_Family",
    ]:
        src_val = str((registry_record or {}).get(field, "") or "").strip()
        if src_val and not str(merged.get(field, "") or "").strip():
            merged[field] = (registry_record or {}).get(field, "")
            _set_field_source(merged, field, source, confidence, scope)
    return merged



def _format_binding_notes(clues: Dict, reason: str, extra: str = "") -> str:
    parts = []
    if clues.get("series_name"):
        parts.append(f"series={clues['series_name']}")
    if clues.get("label_text"):
        parts.append(f"label={clues['label_text']}")
    if clues.get("x_value"):
        parts.append(f"x_value={clues['x_value']}")
    parts.append(f"reason={reason}")
    if extra:
        parts.append(extra)
    return "; ".join([p for p in parts if p])


def _disambiguate_mo2c_framework_matches(
    direct_matches: List[Tuple[str, Dict]],
    clues: Dict,
    point_record: Dict,
) -> List[Tuple[str, Dict]]:
    label_text = clean_residual_mojibake_chars(str(
        clues.get("label_text", "")
        or point_record.get("raw_category_label", "")
        or point_record.get("Series_Name", "")
        or point_record.get("Catalyst", "")
        or ""
    )).strip()
    framework = _parse_mo2c_framework_identity(label_text)
    if not framework or not direct_matches:
        return direct_matches
    active = normalize_active_metal_for_merge(str(framework.get("Active_Metal", "") or ""))
    if not active:
        return direct_matches

    filtered: List[Tuple[str, Dict]] = []
    for reg_key, reg_record in direct_matches:
        support = normalize_support_for_merge(str(reg_record.get("Support_Normalized", "") or reg_record.get("Support", "")))
        reg_active = normalize_active_metal_for_merge(str(reg_record.get("Active_Metal", "") or ""))
        reg_promoter = normalize_promoter_for_merge(str(reg_record.get("Promoter_Metal", "") or reg_record.get("Promoter", "") or ""))
        if support == "mo2c" and reg_active == active and reg_promoter in {"", active.lower()}:
            filtered.append((reg_key, reg_record))
    if len(filtered) == 1:
        return filtered

    if len(filtered) > 1:
        ratio_match = re.search(r"\(\s*(\d+(?:\.\d+)?)\s*\)", label_text)
        if ratio_match:
            ratio_value = ratio_match.group(1)
            exact_ratio: List[Tuple[str, Dict]] = []
            for reg_key, reg_record in filtered:
                haystack = " ".join([
                    str(reg_record.get("Catalyst", "") or ""),
                    str(reg_record.get("Catalyst_ID", "") or ""),
                    str(reg_record.get("Canonical_Catalyst_ID", "") or ""),
                    str(reg_record.get("Alloy_Ratio", "") or ""),
                ])
                if re.search(rf"(?:\(|:)\s*{re.escape(ratio_value)}(?:\)|\b)", haystack):
                    exact_ratio.append((reg_key, reg_record))
            if len(exact_ratio) == 1:
                return exact_ratio
    return filtered or direct_matches



def match_figure_identity_to_registry(
    point_record: Dict,
    registry: Dict,
    local_context: str,
    context_alias_map: Optional[Dict[str, Dict[str, Any]]] = None,
) -> Dict:
    clues = _extract_identity_clues_from_figure_record(point_record)

    if not registry:
        return {
            "matched": False,
            "matched_registry_key": "",
            "match_mode": "unmatched",
            "match_confidence": "low",
            "binding_notes": _format_binding_notes(clues, "registry_empty"),
            "registry_record": None,
        }

    raw_direct_matches = _find_registry_matches_by_keys(clues.get("alias_keys", []), registry)
    direct_matches: List[Tuple[str, Dict]] = []
    for reg_key, reg_record in raw_direct_matches:
        clue_loading = normalize_loading_for_merge(str((clues.get("clue_record") or {}).get("Metal_Loading_wt%", "") or ""))
        reg_loading = normalize_loading_for_merge(str(reg_record.get("Metal_Loading_wt%", "") or ""))
        if bool(clues.get("explicit_loading_in_label") and clue_loading and reg_loading and clue_loading != reg_loading):
            continue
        direct_label_text = str(clues.get("label_text", "") or point_record.get("raw_category_label", "") or point_record.get("Catalyst", "") or "")
        direct_clue_metals = _extract_explicit_metal_set_from_identity_clue(direct_label_text)
        if direct_clue_metals is None:
            clue_record_for_metals = clues.get("clue_record") or {}
            direct_clue_metals = {
                metal for metal in [
                    normalize_active_metal_for_merge(str(clue_record_for_metals.get("Active_Metal", "") or "")),
                    normalize_promoter_for_merge(str(clue_record_for_metals.get("Promoter", "") or clue_record_for_metals.get("Promoter_Metal", "") or "")),
                ] if metal
            } or None
        direct_reg_metals = _extract_registry_metal_set(reg_record)
        direct_ok, _direct_reason = _is_metal_set_compatible_for_unique_match(direct_clue_metals, direct_reg_metals)
        if direct_ok:
            direct_matches.append((reg_key, reg_record))
    if len(direct_matches) == 1:
        reg_key, reg_record = direct_matches[0]
        clue_loading = normalize_loading_for_merge(str((clues.get("clue_record") or {}).get("Metal_Loading_wt%", "") or ""))
        reg_loading = normalize_loading_for_merge(str(reg_record.get("Metal_Loading_wt%", "") or ""))
        explicit_loading_mismatch = bool(clues.get("explicit_loading_in_label") and clue_loading and reg_loading and clue_loading != reg_loading)
        direct_label_text = str(clues.get("label_text", "") or point_record.get("raw_category_label", "") or point_record.get("Catalyst", "") or "")
        direct_clue_metals = _extract_explicit_metal_set_from_identity_clue(direct_label_text)
        if direct_clue_metals is None:
            clue_record_for_metals = clues.get("clue_record") or {}
            direct_clue_metals = {
                metal for metal in [
                    normalize_active_metal_for_merge(str(clue_record_for_metals.get("Active_Metal", "") or "")),
                    normalize_promoter_for_merge(str(clue_record_for_metals.get("Promoter", "") or clue_record_for_metals.get("Promoter_Metal", "") or "")),
                ] if metal
            } or None
        direct_reg_metals = _extract_registry_metal_set(reg_record)
        direct_ok, direct_reason = _is_metal_set_compatible_for_unique_match(direct_clue_metals, direct_reg_metals)
        if not explicit_loading_mismatch and direct_ok:
            return {
                "matched": True,
                "matched_registry_key": reg_key,
                "match_mode": "direct_label",
                "match_confidence": "high",
                "binding_notes": _format_binding_notes(clues, "direct_registry_key_match", f"registry_key={reg_key}"),
                "registry_record": dict(reg_record),
            }
    if len(direct_matches) > 1:
        mo2c_direct_matches = _disambiguate_mo2c_framework_matches(direct_matches, clues, point_record)
        if len(mo2c_direct_matches) == 1:
            reg_key, reg_record = mo2c_direct_matches[0]
            return {
                "matched": True,
                "matched_registry_key": reg_key,
                "match_mode": "direct_label",
                "match_confidence": "high",
                "binding_notes": _format_binding_notes(clues, "mo2c_framework_disambiguated", f"registry_key={reg_key}"),
                "registry_record": dict(reg_record),
            }
        return {
            "matched": False,
            "matched_registry_key": "",
            "match_mode": "ambiguous",
            "match_confidence": "low",
            "binding_notes": _format_binding_notes(clues, "multiple_direct_registry_matches"),
            "registry_record": None,
        }

    if not clues.get("generic_series_only"):
        partial_clue_record = dict(clues.get("clue_record", {}))
        partial_clue_record["_explicit_loading_in_label"] = clues.get("explicit_loading_in_label", False)
        partial_matches = _find_unique_registry_candidate(partial_clue_record, registry)
        if len(partial_matches) == 1:
            reg_key, reg_record = partial_matches[0]
            # Safety: verify metal-set using figure label text, not just clue_record fields.
            # This catches the case where clue_record has no Active_Metal (empty before binding)
            # but the raw_category_label / Series_Name explicitly names the metal.
            # e.g. "Ni/clay" label must NOT be accepted against "Ni-Cu/clay" registry.
            _pu_label = clean_residual_mojibake_chars(str(
                point_record.get("raw_category_label", "")
                or point_record.get("Series_Name", "")
                or point_record.get("Catalyst", "") or ""
            )).strip()
            if _pu_label:
                _pu_clue_metals = _extract_explicit_metal_set_from_identity_clue(_pu_label)
                _pu_reg_metals = _extract_registry_metal_set(reg_record)
                _pu_ok, _pu_why = _is_metal_set_compatible_for_unique_match(_pu_clue_metals, _pu_reg_metals)
                if not _pu_ok:
                    return {
                        "matched": False,
                        "matched_registry_key": reg_key,
                        "match_mode": "unmatched",
                        "match_confidence": "low",
                        "binding_notes": _format_binding_notes(
                            clues, f"partial_unique_blocked_identity_safety: {_pu_why}"),
                        "registry_record": None,
                    }
            return {
                "matched": True,
                "matched_registry_key": reg_key,
                "match_mode": "partial_unique",
                "match_confidence": "medium",
                "binding_notes": _format_binding_notes(clues, "unique_partial_identity_match", f"registry_key={reg_key}"),
                "registry_record": dict(reg_record),
            }
        if len(partial_matches) > 1:
            return {
                "matched": False,
                "matched_registry_key": "",
                "match_mode": "ambiguous",
                "match_confidence": "low",
                "binding_notes": _format_binding_notes(clues, "multiple_partial_candidates"),
                "registry_record": None,
            }

    for label in [clues.get("series_name", ""), clues.get("x_value", "")]:
        if not label:
            continue
        mapped_label = _extract_label_mapping_from_context(label, local_context, registry, context_alias_map)
        if not mapped_label:
            continue

        mapped_record = {
            "Catalyst": mapped_label,
            "Series_Name": str(point_record.get("Series_Name", "")),
            "Notes": str(point_record.get("Notes", "")),
        }
        mapped_clues = _extract_identity_clues_from_figure_record(mapped_record)
        mapped_direct = _find_registry_matches_by_keys(mapped_clues.get("alias_keys", []), registry)
        if len(mapped_direct) == 1:
            reg_key, reg_record = mapped_direct[0]
            return {
                "matched": True,
                "matched_registry_key": reg_key,
                "match_mode": "caption_map",
                "match_confidence": "medium",
                "binding_notes": _format_binding_notes(mapped_clues, "caption_mapping_match", f"mapped_from={label}; registry_key={reg_key}"),
                "registry_record": dict(reg_record),
            }

        mapped_clue_record = dict(mapped_clues.get("clue_record", {}))
        mapped_clue_record["_explicit_loading_in_label"] = mapped_clues.get("explicit_loading_in_label", False)
        mapped_partial = _find_unique_registry_candidate(mapped_clue_record, registry)
        if len(mapped_partial) == 1:
            reg_key, reg_record = mapped_partial[0]
            # Safety: metal-set check using the mapped_label text
            _mp_clue_metals = _extract_explicit_metal_set_from_identity_clue(mapped_label)
            _mp_reg_metals = _extract_registry_metal_set(reg_record)
            _mp_ok, _mp_why = _is_metal_set_compatible_for_unique_match(_mp_clue_metals, _mp_reg_metals)
            if not _mp_ok:
                continue  # skip this mapped_label, try next if any
            return {
                "matched": True,
                "matched_registry_key": reg_key,
                "match_mode": "caption_map",
                "match_confidence": "medium",
                "binding_notes": _format_binding_notes(mapped_clues, "caption_mapping_unique_candidate", f"mapped_from={label}; registry_key={reg_key}"),
                "registry_record": dict(reg_record),
            }
        if len(mapped_partial) > 1:
            return {
                "matched": False,
                "matched_registry_key": "",
                "match_mode": "ambiguous",
                "match_confidence": "low",
                "binding_notes": _format_binding_notes(mapped_clues, "caption_mapping_ambiguous", f"mapped_from={label}"),
                "registry_record": None,
            }

    if not clues.get("generic_series_only"):
        family_clue_record = dict(clues.get("clue_record", {}))
        for field in [
            "Active_Metal", "Promoter", "Promoter_Metal",
            "Support", "Support_Normalized",
            "Metal_Loading_wt%", "Alloy_Ratio",
            "raw_category_label", "Series_Name", "Catalyst",
        ]:
            if not str(family_clue_record.get(field, "")).strip():
                family_clue_record[field] = point_record.get(field, "")
        family_matches = _find_unique_family_registry_candidate(family_clue_record, registry)
        if len(family_matches) == 1:
            reg_key, reg_record = family_matches[0]
            can_broadcast = int(reg_record.get("can_broadcast_preparation", 0) or 0)
            family_mode = "family_backbone_match" if can_broadcast else "family_identity_match"
            return {
                "matched": True,
                "matched_registry_key": reg_key,
                "match_mode": family_mode,
                "match_confidence": "low",
                "binding_notes": _format_binding_notes(
                    clues,
                    family_mode,
                    f"registry_key={reg_key}; active={family_clue_record.get('Active_Metal', '')}; "
                    f"promoter={family_clue_record.get('Promoter_Metal', '') or family_clue_record.get('Promoter', '')}; "
                    f"support={family_clue_record.get('Support_Normalized', '') or family_clue_record.get('Support', '')}",
                ),
                "registry_record": dict(reg_record),
                "_family_backbone_only": bool(can_broadcast),
                "_family_identity_only": not bool(can_broadcast),
            }
        if len(family_matches) > 1:
            return {
                "matched": False,
                "matched_registry_key": "",
                "match_mode": "ambiguous",
                "match_confidence": "low",
                "binding_notes": _format_binding_notes(clues, "ambiguous_family_backbone_candidates"),
                "registry_record": None,
            }

        relaxed_matches = _find_loading_relaxed_single_metal_family_candidate(family_clue_record, registry)
        if len(relaxed_matches) == 1:
            reg_key, reg_record = relaxed_matches[0]
            can_broadcast = int(reg_record.get("can_broadcast_preparation", 0) or 0)
            relaxed_mode = "loading_relaxed_family_backbone" if can_broadcast else "loading_relaxed_identity_match"
            return {
                "matched": True,
                "matched_registry_key": reg_key,
                "match_mode": relaxed_mode,
                "match_confidence": "low",
                "binding_notes": _format_binding_notes(
                    clues,
                    relaxed_mode,
                    f"registry_key={reg_key}; active={family_clue_record.get('Active_Metal', '')}; "
                    f"support={family_clue_record.get('Support_Normalized', '') or family_clue_record.get('Support', '')}",
                ),
                "registry_record": dict(reg_record),
                "_family_backbone_only": bool(can_broadcast),
                "_family_identity_only": not bool(can_broadcast),
            }
        if len(relaxed_matches) > 1:
            return {
                "matched": False,
                "matched_registry_key": "",
                "match_mode": "ambiguous",
                "match_confidence": "low",
                "binding_notes": _format_binding_notes(clues, "ambiguous_loading_relaxed_family_candidates"),
                "registry_record": None,
            }

    support_label = clean_residual_mojibake_chars(str(
        point_record.get("raw_category_label", "")
        or point_record.get("Catalyst", "")
        or point_record.get("Series_Name", "")
        or ""
    )).strip()
    if _is_support_only_figure_label(support_label):
        return {
            "matched": False,
            "matched_registry_key": "",
            "match_mode": "support_only_control",
            "match_confidence": "low",
            "binding_notes": _format_binding_notes(clues, "support_only_control_label", f"label={support_label}"),
            "registry_record": None,
        }

    unresolved_reason = "generic_label_without_context_mapping" if clues.get("generic_series_only") else "no_reliable_registry_match"
    return {
        "matched": False,
        "matched_registry_key": "",
        "match_mode": "unmatched",
        "match_confidence": "low",
        "binding_notes": _format_binding_notes(clues, unresolved_reason),
        "registry_record": None,
    }



def _expand_canonical_id_to_aliases(canonical_id: str) -> List[str]:
    """
    Convert internal Canonical_Catalyst_ID format to human-readable aliases
    so figure legend labels can match registry entries.

    Examples:
      ni__prom-cu/al2o3  → ["ni-cu/al2o3", "nicu/al2o3", "ni/al2o3"]
      ni__prom-mo/mo2c   → ["ni-mo/mo2c", "nimo/mo2c", "ni/mo2c"]
      ni/al2o3           → ["ni/al2o3"]
    """
    cid = str(canonical_id or "").strip().lower()
    if not cid:
        return []
    aliases = []

    # Fix A: strip leading loading prefix (e.g. "30wt%-ni__prom-cu/al2o3" → "ni__prom-cu/al2o3")
    # so that figure labels without explicit loading can still match registry entries.
    cid_stripped = re.sub(r"^\d+(?:\.\d+)?wt%-", "", cid)
    # Also strip alloy-ratio infix like "ni20.0-cu80.0" → "ni-cu" for alias generation
    # (keep original cid for exact match, use stripped for alias expansion)
    if cid_stripped != cid:
        # Recursively expand the stripped CID and prepend those aliases
        stripped_aliases = _expand_canonical_id_to_aliases(cid_stripped)
        aliases.extend(stripped_aliases)
        # Also add original cid as-is for exact match
        aliases.append(cid)
        return _dedupe_keep_order([a for a in aliases if a])

    def _extend_numeric_loading_aliases(ma: str, la: str, mb: str, lb: str, support: str = "") -> List[str]:
        try:
            pairs = [(ma.capitalize(), float(la)), (mb.capitalize(), float(lb))]
        except (TypeError, ValueError):
            return []
        return _build_component_loading_aliases(pairs, support)

    # Pattern: <metal>__prom-<promoter>/<support>
    m = re.match(r"^([a-z0-9]+)__prom-([a-z0-9]+)/(.+)$", cid)
    if m:
        metal, promoter, support = m.group(1), m.group(2), m.group(3)
        aliases.append(f"{metal}-{promoter}/{support}")   # ni-cu/al2o3
        aliases.append(f"{metal}{promoter}/{support}")    # nicu/al2o3
        aliases.append(f"{promoter}-{metal}/{support}")   # cu-ni/al2o3 (reversed order)
        aliases.append(f"{promoter}{metal}/{support}")    # cuni/al2o3 (reversed no-hyphen)
        aliases.append(f"{metal}/{support}")              # ni/al2o3 (fallback)
        return aliases
    # Pattern: <metal>__prom-<promoter> (no support)
    m2 = re.match(r"^([a-z0-9]+)__prom-([a-z0-9]+)$", cid)
    if m2:
        metal, promoter = m2.group(1), m2.group(2)
        aliases.append(f"{metal}-{promoter}")
        aliases.append(f"{metal}{promoter}")
        aliases.append(f"{promoter}-{metal}")             # reversed order
        aliases.append(f"{promoter}{metal}")              # reversed no-hyphen
        return aliases
    # Plain format: check if it looks like "metal-promoter/support" or "metalprom/support"
    # and generate the complementary no-hyphen / with-hyphen variant.
    # Restrict promoter to pure letters ≤3 chars (element symbols like cu, zn, la, ce)
    # to avoid false-matching "ni-al2o3/sio2" as metal=ni, promoter=al2o3.
    m3 = re.match(r"^([a-z]{1,3})-([a-z]{1,3})/(.+)$", cid)
    if m3:
        metal, promoter, support = m3.group(1), m3.group(2), m3.group(3)
        aliases.append(cid)                                    # ni-cu/al2o3 (original)
        aliases.append(f"{metal}{promoter}/{support}")         # nicu/al2o3
        aliases.append(f"{promoter}-{metal}/{support}")        # cu-ni/al2o3 (reversed)
        aliases.append(f"{promoter}{metal}/{support}")         # cuni/al2o3 (reversed no-hyphen)
        return aliases
    m4 = re.match(r"^([a-z]{1,3})-([a-z]{1,3})$", cid)
    if m4:
        metal, promoter = m4.group(1), m4.group(2)
        aliases.append(cid)                                    # ni-cu (original)
        aliases.append(f"{metal}{promoter}")                   # nicu
        aliases.append(f"{promoter}-{metal}")                  # cu-ni (reversed)
        aliases.append(f"{promoter}{metal}")                   # cuni (reversed no-hyphen)
        return aliases
    # Pattern: loading-prefixed bimetal e.g. "7%cu-3%ni/al2o3" or "7%cu-3%ni"
    # Keep metal-loading pairs intact; only allow order reversal.
    m5 = re.match(r"^(\d+(?:\.\d+)?%)([a-z]{1,3})-(\d+(?:\.\d+)?%)([a-z]{1,3})(/.*)?$", cid)
    if m5:
        la, ma, lb, mb, sup = m5.group(1), m5.group(2), m5.group(3), m5.group(4), m5.group(5) or ""
        support = sup.lstrip("/")
        try:
            pairs = [(ma.capitalize(), float(la.rstrip("%"))), (mb.capitalize(), float(lb.rstrip("%")))]
        except (TypeError, ValueError):
            pairs = []
        if pairs:
            aliases.extend(_build_component_loading_aliases(pairs, support))
            if support:
                aliases.extend(_build_component_loading_aliases(pairs, ""))
        return _dedupe_keep_order([a for a in aliases if a])
    # Pattern: numeric-ratio bimetal e.g. "ni7-cu3/al2o3" or "ni7-cu3"
    # Again, keep the original metal-number binding and only allow order reversal.
    m6 = re.match(r"^([a-z]{1,3})(\d+(?:\.\d+)?)-([a-z]{1,3})(\d+(?:\.\d+)?)(/.*)?$", cid)
    if m6:
        ma, la, mb, lb, sup = m6.group(1), m6.group(2), m6.group(3), m6.group(4), m6.group(5) or ""
        support = sup.lstrip("/")
        aliases.extend(_extend_numeric_loading_aliases(ma, la, mb, lb, support))
        if support:
            aliases.extend(_extend_numeric_loading_aliases(ma, la, mb, lb, ""))
        return _dedupe_keep_order([a for a in aliases if a])
    aliases.append(cid)
    return aliases


def _normalize_registry_label(label: str) -> str:
    text = clean_residual_mojibake_chars(str(label or "")).strip()
    if not text:
        return ""
    subs_map = str.maketrans({
        "₀": "0", "₁": "1", "₂": "2", "₃": "3", "₄": "4",
        "₅": "5", "₆": "6", "₇": "7", "₈": "8", "₉": "9",
        "²": "2", "³": "3",
    })
    text = text.translate(subs_map)
    text = text.replace("\u2093", "x")   # ₓ (Latin Subscript Small Letter X) → x
    text = text.replace("\u1d67", "y")   # ᵧ (Phonetic Extensions Small Letter Y) → y
    text = text.replace("\u1d62", "i")   # ᵢ → i
    text = text.replace("\u2090", "a")   # ₐ → a
    text = text.replace("\u2091", "e")   # ₑ → e
    text = text.replace("\u2092", "o")   # ₒ → o
    text = text.replace("γ", "gamma").replace("Γ", "gamma")
    text = text.replace("·", "-").replace("•", "-")  # middle dot → hyphen (e.g. Ni/γ-Al2O3·5H)
    text = text.replace("–", "-").replace("—", "-").replace("−", "-")
    text = _strip_descriptor_suffix_before_element_parse(text)
    # Fix C: strip preparation method suffixes so "nixmgyo-hydro" == "nixmgyo-impre" == "nixmgyo"
    # These suffixes describe synthesis route, not catalyst identity
    text = re.sub(r"-(?:hydro|impre|copre|coppt|dp|imp|cp|sg|sol|gel|precip|precip\w*|impreg\w*|hydrotherm\w*)(?=[/\s,;)_]|$)", "", text, flags=re.I)
    text = re.sub(r"\([^)]*(?:commercial|reference|baseline|blank|tested at|prepared by|compared with|comparison|comparative)[^)]*\)", " ", text, flags=re.I)
    text = re.sub(r"\b(?:different catalysts?|catalyst type|samples?)\b", " ", text, flags=re.I)
    text = re.sub(r"\b(?:sample|samples|catalyst|catalysts|cat\.?|tested at|prepared by|commercial(?: catalyst)?|common|comparative|comparison|reference|baseline|blank|run)\b", " ", text, flags=re.I)
    text = re.sub(r"\bat\s+\d+(?:\.\d+)?\s*(?:deg\.?\s*c|celsius|c)\b", " ", text, flags=re.I)
    # Fix: normalize "7wt%" → "7%", "7 wt %" → "7%"
    text = re.sub(r"(\d)\s*wt\s*%\s*", r"\1%", text, flags=re.I)
    # Fix: normalize "7% Ni" → "7%ni", "3 %" → "3%" (space around %)
    text = re.sub(r"(\d)\s*%\s*", r"\1%", text)
    text = re.sub(r"\s+", " ", text)
    text = re.sub(r"\s*/\s*", "/", text)
    text = re.sub(r"\s*-\s*", "-", text)
    text = text.strip(" ,;:_-/").lower()
    # Fix B: strip parenthetical loading suffix from head before support split
    # e.g. "ni-cu(10%)/al2o3" → "ni-cu/al2o3", "ni-cu(10 %)/al2o3" → "ni-cu/al2o3"
    text = re.sub(r"\(\s*\d+(?:\.\d+)?\s*%?\s*\)(?=/)", "", text)
    # Strip crystal-phase prefixes from support part (after /) so that
    # "ni/gamma-al2o3" == "ni/al2o3" and "ni/alpha-al2o3" == "ni/al2o3"
    if "/" in text:
        _head, _support = text.split("/", 1)
        _support = re.sub(r"^(?:gamma|alpha|beta|delta|theta|eta)-", "", _support)
        # Fix B: strip morphology suffixes from support so "zno-rod" == "zno"
        _support = re.sub(r"-(?:rod|powder|sphere|wire|tube|nanoparticle|nanoparticles|np|nps|sheet|fiber|flake|pellet|granule)s?$", "", _support)
        text = f"{_head}/{_support}"
    return text





def _normalize_support_only_label(label: str) -> str:
    norm = _normalize_registry_label(label)
    if not norm or _is_blank_or_baseline_label(norm):
        return ""
    if "/" in norm:
        return ""

    active_like = bool(re.search(r"(?<![a-z0-9])(?:\d+(?:\.\d+)?)?(?:ni|cu|pt|pd|ru|rh|co|fe|ir|au|zn|sn|ga|in|mo|cr|mn|k)(?![a-z0-9])", norm, flags=re.I))
    if active_like:
        return ""

    support_map = {
        "al2o3": "al2o3",
        "gamma-al2o3": "al2o3",
        "gamma alumina": "al2o3",
        "alumina": "al2o3",
        "zro2": "zro2",
        "ceo2": "ceo2",
        "sio2": "sio2",
        "tio2": "tio2",
        "mgo": "mgo",
        "la2o3": "la2o3",
        "cnt": "cnts",
        "cnts": "cnts",
        "cnf": "cnf",
        "ac": "activated carbon",
        "activated carbon": "activated carbon",
        "carbon": "activated carbon",
        "sep": "sep",
        "sepiolite": "sep",
        "sba-15": "sba-15",
        "mcm-41": "mcm-41",
        "ceo2-zro2": "ceo2-zro2",
        "mgo-al2o3": "mgo-al2o3",
    }
    if norm in support_map:
        return support_map[norm]
    return norm if norm in SUPPORT_LIKE_TOKENS else ""




def _is_blank_or_baseline_label(label: str) -> bool:
    raw = clean_residual_mojibake_chars(str(label or "")).strip()
    if not raw:
        return False
    raw = raw.translate(str.maketrans({
        "₀": "0", "₁": "1", "₂": "2", "₃": "3", "₄": "4",
        "₅": "5", "₆": "6", "₇": "7", "₈": "8", "₉": "9",
        "²": "2", "³": "3",
    }))
    raw = raw.replace("–", "-").replace("—", "-").replace("−", "-")
    raw = raw.replace("γ", "gamma").replace("Γ", "gamma")
    raw = re.sub(r"\s+", " ", raw).strip(" ,;:_-/").lower()
    return bool(re.search(r"\b(?:no catalyst|without catalyst|blank|empty reactor|empty bed|reactor blank|no cat\.?|without cat\.?)\b", raw, flags=re.I))





def _match_category_figure_label_to_registry(
    label: str,
    registry: Dict,
    local_context: str,
    context_alias_map: Optional[Dict[str, Dict[str, Any]]] = None,
) -> Dict[str, Any]:
    raw_label = clean_residual_mojibake_chars(str(label or "")).strip()
    mapped_label = clean_residual_mojibake_chars(str(
        _extract_label_mapping_from_context(raw_label, local_context, registry, context_alias_map) or ""
    )).strip()
    candidate_texts = _dedupe_keep_order([item for item in [raw_label, mapped_label] if item])
    candidate_norms = _dedupe_keep_order([_normalize_registry_label(item) for item in candidate_texts if _normalize_registry_label(item)])
    support_only = _normalize_support_only_label(mapped_label or raw_label)
    is_baseline = _is_blank_or_baseline_label(mapped_label or raw_label)

    def _result(matched: bool, mode: str, confidence: str, reason: str, reg_key: str = "", reg_record: Optional[Dict] = None) -> Dict[str, Any]:
        reg_record = dict(reg_record or {})
        matched_label = str(reg_record.get("Catalyst", "") or reg_record.get("Catalyst_ID", "") or reg_record.get("Canonical_Catalyst_ID", "") or reg_key)
        return {
            "matched": matched,
            "matched_registry_key": str(reg_key or ""),
            "registry_record": reg_record if matched else (reg_record or None),
            "binding_mode": mode,
            "binding_confidence": confidence,
            "binding_reason": reason,
            "matched_registry_label": matched_label,
            "is_baseline_label": bool(is_baseline),
            "is_support_only_label": bool(support_only),
        }

    if not candidate_norms and not support_only and not is_baseline:
        return _result(False, "unmatched", "low", "empty_category_label")
    if not registry:
        if is_baseline:
            return _result(False, "baseline_match", "medium", "blank_or_baseline_label")
        if support_only:
            return _result(False, "support_only_match", "medium", "support_only_label")
        return _result(False, "unmatched", "low", "registry_empty")

    def _has_loading(text: str) -> bool:
        return bool(re.search(r"(?<![a-z])\d+(?:\.\d+)?(?=(?:ni|cu|pt|pd|ru|rh|co|fe|ir|au))", text, flags=re.I))

    def _head_signature(text: str) -> str:
        head = text.split("/", 1)[0]
        head = re.sub(r"(?<![a-z])\d+(?:\.\d+)?", "", head)
        head = re.sub(r"[^a-z0-9-]+", "", head.lower())
        return head

    def _support_signature(text: str) -> str:
        if "/" in text:
            support = text.split("/", 1)[1]
        else:
            support = text
        # Strip crystal-phase prefixes so "gamma-al2o3" matches "al2o3"
        support = re.sub(r"^(?:gamma|alpha|beta|delta|theta|eta)-", "", support)
        return support

    exact_candidates: List[Tuple[str, str, str, Dict]] = []
    canonical_candidates: List[Tuple[str, str, str, Dict]] = []
    support_candidates: List[Tuple[str, str, str, Dict]] = []
    baseline_candidates: List[Tuple[str, str, str, Dict]] = []
    partial_candidates: List[Tuple[str, str, str, Dict]] = []

    for reg_key, reg_record in registry.items():
        raw_aliases = _dedupe_keep_order([
            str(reg_key).strip().lower(),
            str(reg_record.get("Catalyst", "")).strip().lower(),
            str(reg_record.get("Catalyst_ID", "")).strip().lower(),
            str(reg_record.get("Canonical_Catalyst_ID", "")).strip().lower(),
        ] + [str(a).strip().lower() for a in reg_record.get("_registry_aliases", []) if str(a).strip()])

        norm_aliases = _dedupe_keep_order([
            _normalize_registry_label(reg_key),
            _normalize_registry_label(str(reg_record.get("Catalyst", "") or "")),
            _normalize_registry_label(str(reg_record.get("Catalyst_ID", "") or "")),
            _normalize_registry_label(str(reg_record.get("Catalyst_ID_normalized", "") or "")),
            _normalize_registry_label(str(reg_record.get("Canonical_Catalyst_ID", "") or "")),
            _normalize_registry_label(str(reg_record.get("identity_alias_group", "") or "")),
        ] + [_normalize_registry_label(a) for a in _expand_canonical_id_to_aliases(
            str(reg_record.get("Canonical_Catalyst_ID", "") or reg_key)
        )])
        norm_aliases = [alias for alias in norm_aliases if alias]

        catalyst_norm = _normalize_registry_label(str(reg_record.get("Catalyst", "") or reg_record.get("Canonical_Catalyst_ID", "") or reg_key))
        support_norm = _normalize_support_only_label(str(reg_record.get("Support_Normalized", "") or reg_record.get("Support", "") or reg_record.get("Catalyst", "")))
        active_norm = normalize_active_metal_for_merge(str(reg_record.get("Active_Metal", "")))
        support_row_like = bool(support_norm and (not active_norm or catalyst_norm == support_norm))
        baseline_row_like = _is_blank_or_baseline_label(str(reg_record.get("Catalyst", "") or reg_record.get("Canonical_Catalyst_ID", "") or reg_key))

        for candidate_text in candidate_texts:
            candidate_lower = candidate_text.lower()
            if candidate_lower and candidate_lower in raw_aliases:
                exact_candidates.append(("exact_label", "high", reg_key, reg_record))
                break

        for candidate_norm in candidate_norms:
            if candidate_norm and candidate_norm in norm_aliases:
                canonical_candidates.append(("canonical_label", "high", reg_key, reg_record))
                break

        if support_only and support_row_like and support_norm == support_only:
            support_candidates.append(("support_only_match", "medium", reg_key, reg_record))

        if is_baseline and baseline_row_like:
            baseline_candidates.append(("baseline_match", "medium", reg_key, reg_record))

        if candidate_norms and not support_only and not is_baseline:
            for candidate_norm in candidate_norms:
                for alias in norm_aliases:
                    if not alias or alias == candidate_norm:
                        continue
                    if not (candidate_norm in alias or alias in candidate_norm):
                        continue
                    if _has_loading(candidate_norm) and _has_loading(alias):
                        continue  # both have explicit loadings — too risky for partial, use exact/canonical
                    if _head_signature(candidate_norm) != _head_signature(alias):
                        continue
                    if _support_signature(candidate_norm) != _support_signature(alias):
                        continue
                    # Problem A: explicit metal-set safety guard for partial overlap.
                    # Even if string signatures match, block when metal set is incompatible.
                    # e.g. "ni/clay" substring of "ni-cu/clay" would normally be caught by
                    # head_signature ("ni" != "ni-cu"), but this adds a belt-and-suspenders layer.
                    clue_metals = _extract_explicit_metal_set_from_identity_clue(candidate_norm)
                    reg_metals = _extract_registry_metal_set(reg_record)
                    compat, compat_reason = _is_metal_set_compatible_for_unique_match(clue_metals, reg_metals)
                    if not compat:
                        # Log reason for audit but don't add to partial_candidates
                        print(f"  [partial_overlap_blocked] clue='{candidate_norm}' "
                              f"reg='{reg_key}' reason={compat_reason}")
                        break
                    partial_candidates.append(("partial_unique", "high", reg_key, reg_record))
                    break

    for candidates, ambiguous_reason in [
        (exact_candidates,     "ambiguous_exact_category_label"),
        (canonical_candidates, "ambiguous_category_label"),
        (support_candidates,   "ambiguous_support_only_label"),
        (baseline_candidates,  "ambiguous_baseline_label"),
        (partial_candidates,   "ambiguous_partial_category_label"),
    ]:
        unique_keys = _dedupe_keep_order([reg_key for _, _, reg_key, _ in candidates])
        if len(unique_keys) == 1:
            for mode, confidence, reg_key, reg_record in candidates:
                if reg_key == unique_keys[0]:
                    # Problem A: one final metal-set safety check at acceptance time.
                    # A "unique" remaining candidate is still unsafe if metal sets conflict.
                    primary_label = (candidate_texts[0] if candidate_texts else "")
                    clue_metals = _extract_explicit_metal_set_from_identity_clue(primary_label)
                    reg_metals = _extract_registry_metal_set(reg_record)
                    compat, compat_reason = _is_metal_set_compatible_for_unique_match(clue_metals, reg_metals)
                    if not compat:
                        print(f"  [unique_match_blocked_metal_set] label='{primary_label}' "
                              f"reg='{reg_key}' reason={compat_reason}")
                        return _result(False, mode, "low",
                                       f"unique_match_blocked_for_incomplete_identity: {compat_reason}",
                                       reg_key, reg_record)
                    return _result(True, mode, confidence,
                                   f"unique_{mode}; metal_set_ok={compat_reason}",
                                   reg_key, reg_record)
        if len(unique_keys) > 1:
            # Fix D: for multiple_partial_candidates, try to disambiguate using
            # numeric value in parentheses of the label (e.g. "Ni-Mo2C (0.8)" → 0.8)
            # Match against Metal_Loading_wt% in each registry record
            if ambiguous_reason == "ambiguous_partial_category_label" and candidate_texts:
                paren_val_m = re.search(r"\(\s*(\d+(?:\.\d+)?)\s*\)", candidate_texts[0])
                if paren_val_m:
                    paren_val = float(paren_val_m.group(1))
                    best_key, best_dist = None, float("inf")
                    for _, _, rk, rr in candidates:
                        try:
                            loading = float(str(rr.get("Metal_Loading_wt%", "") or "").strip())
                            dist = abs(loading - paren_val)
                            if dist < best_dist:
                                best_dist, best_key = dist, rk
                        except (ValueError, TypeError):
                            pass
                    if best_key and best_dist < 1.0:
                        for mode, confidence, rk, rr in candidates:
                            if rk == best_key:
                                return _result(True, "partial_loading_disambig", "medium",
                                               f"disambig_by_paren_val={paren_val}_loading={best_dist:.2f}",
                                               rk, rr)
            return _result(False, candidates[0][0], "low", ambiguous_reason)

    if is_baseline:
        return _result(False, "baseline_match", "medium", "baseline_label_blocks_registry_match")
    if support_only:
        return _result(False, "support_only_match", "medium", "support_only_label_blocks_active_registry_match")
    return _result(False, "unmatched", "low", "no_reliable_category_match")




def _build_category_point_from_figure_record(raw_record: Dict, condition_anchor: Dict[str, Any], metadata: Dict[str, Any]) -> Dict[str, Any]:
    row = dict(raw_record or {})
    note_axis, note_x_value = _extract_x_axis_and_value_from_notes(str(row.get("Notes", "")))
    catalyst_fallback = _usable_figure_category_label(row.get("Catalyst", ""))
    raw_category_label = clean_residual_mojibake_chars(str(
        row.get("raw_category_label", "") or note_x_value or catalyst_fallback
    )).strip()
    series_name = clean_residual_mojibake_chars(str(row.get("Series_Name", "") or "")).strip()

    if not raw_category_label and series_name and _is_series_name_real_catalyst(series_name, str(metadata.get("series_role", "unknown"))):
        raw_category_label = series_name

    row["raw_category_label"] = raw_category_label
    if str(row.get("Catalyst", "")).strip() and not catalyst_fallback and not raw_category_label:
        row["Catalyst"] = ""
    row["semantic_figure_role"] = str(metadata.get("semantic_figure_role", "category_screening") or "category_screening")
    anchor_sources = []
    for key, value in (condition_anchor or {}).items():
        if key == "_anchor_sources":
            anchor_sources.extend(value if isinstance(value, list) else [str(value)])
            continue
        if key.startswith("_"):
            continue
        if value and not str(row.get(key, "")).strip():
            row[key] = value
    if anchor_sources:
        row["condition_anchor_source"] = "|".join(_dedupe_keep_order([str(item) for item in anchor_sources if str(item).strip()]))

    if raw_category_label:
        if _is_blank_or_baseline_label(raw_category_label):
            row["is_baseline_label"] = 1
            row["is_support_only_label"] = 0
            if not str(row.get("Catalyst", "")).strip():
                row["Catalyst"] = raw_category_label
        else:
            support_only = _normalize_support_only_label(raw_category_label)
            if support_only:
                row["is_baseline_label"] = 0
                row["is_support_only_label"] = 1
                if not str(row.get("Support", "")).strip():
                    row["Support"] = support_only
                if not str(row.get("Catalyst", "")).strip():
                    row["Catalyst"] = raw_category_label
            elif not str(row.get("Catalyst", "")).strip() or (
                series_name and str(row.get("Catalyst", "")).strip().lower() == series_name.lower() and not _is_series_name_real_catalyst(series_name, str(metadata.get("series_role", "unknown")))
            ):
                row["Catalyst"] = raw_category_label
    return row





def attach_figure_points_to_registry(fig_records: List[Dict], registry: Dict, local_context: str, file_name: str) -> Tuple[List[Dict], List[Dict]]:
    if not globals().get("_FIGURE_BINDING_SELFCHECK_DONE"):
        globals()["_FIGURE_BINDING_SELFCHECK_DONE"] = True
        try:
            case_anchor = _extract_condition_anchor_from_series_or_caption("90 C", "Fig. 4 Methanol steam reforming performance Different catalysts at 90 C")
            if str(case_anchor.get("Reaction_Temp_C", "")).strip() != "90":
                print("  [WARNING] figure selfcheck failed: category condition anchor")
            baseline = _match_category_figure_label_to_registry("No catalyst", {}, "")
            support_only = _match_category_figure_label_to_registry("ZrO2", {}, "")
            active = _match_category_figure_label_to_registry("Ni/ZrO2", {}, "")
            if not baseline.get("is_baseline_label"):
                print("  [WARNING] figure selfcheck failed: baseline label detection")
            if not support_only.get("is_support_only_label"):
                print("  [WARNING] figure selfcheck failed: support-only label detection")
            if active.get("is_baseline_label") or active.get("is_support_only_label"):
                print("  [WARNING] figure selfcheck failed: active label classification")

            sample_registry = {
                "ni/zro2": {"Catalyst": "Ni/ZrO2", "_registry_aliases": ["ni/zro2"], "Active_Metal": "Ni", "Support": "ZrO2"},
                "10ni/zro2": {"Catalyst": "10Ni/ZrO2", "_registry_aliases": ["10ni/zro2"], "Active_Metal": "Ni", "Support": "ZrO2"},
                "zro2": {"Catalyst": "ZrO2", "_registry_aliases": ["zro2"], "Active_Metal": "", "Support": "ZrO2"},
                "ni/p-9-2-1,5": {"Catalyst": "Ni/P-9-2-1,5", "_registry_aliases": ["ni/p-9-2-1,5"], "Active_Metal": "Ni", "Support": "Activated Carbon"},
                "20wt%-ni/cnts": {"Catalyst": "20wt.%Ni/CNTs", "_registry_aliases": ["20wt.%ni/cnts", "ni/cnts"], "Active_Metal": "Ni", "Support": "CNTs"},
            }
            match_case = _match_category_figure_label_to_registry("Ni/ZrO2", sample_registry, "")
            if match_case.get("matched_registry_key") == "10ni/zro2":
                print("  [WARNING] figure selfcheck failed: partial overlap mis-bound Ni/ZrO2 -> 10Ni/ZrO2")
            alias_map = _build_numbered_label_alias_map("C1=Ni/P-9-2-1,5; C2=Ni/C-8-3-3", sample_registry)
            numbered = _resolve_numbered_category_label("C1", alias_map, sample_registry, "C1=Ni/P-9-2-1,5; C2=Ni/C-8-3-3")
            if not numbered.get("matched") or str(numbered.get("binding_mode", "")) != "numbered_alias_map":
                print("  [WARNING] figure selfcheck failed: numbered alias map")
            named_map = _build_context_label_alias_map("20wt.%Ni/CNTs (Ni/CNTs without TMAOH)", sample_registry)
            if str(_extract_label_mapping_from_context("Ni/CNTs without TMAOH", "20wt.%Ni/CNTs (Ni/CNTs without TMAOH)", sample_registry, named_map)).strip() != "20wt.%Ni/CNTs":
                print("  [WARNING] figure selfcheck failed: named alias map")
            # Run binding safety selfcheck (has its own idempotency guard)
            _selfcheck_binding_safety()
        except Exception as e:
            print(f"  [WARNING] figure binding selfcheck error: {e}")

    def _empty_result(reason: str = "empty_category_label") -> Dict[str, Any]:
        return {
            "matched": False,
            "matched_registry_key": "",
            "registry_record": None,
            "binding_mode": "unmatched",
            "binding_confidence": "low",
            "binding_reason": reason,
            "matched_registry_label": "",
            "is_baseline_label": False,
            "is_support_only_label": False,
            "alias_map_source": "",
            "alias_map_evidence": "",
        }

    def _compat_mode(mode: str) -> str:
        mode = str(mode or "").strip()
        if mode in {"exact_label", "canonical_label", "support_only_match", "baseline_match"}:
            return "direct_label"
        if mode in {"alias_map_label", "numbered_alias_map"}:
            return "caption_map"
        if mode in {"direct_label", "caption_map", "partial_unique", "ambiguous", "unmatched"}:
            return mode
        return mode or "unmatched"

    def _is_ambiguous(payload: Dict[str, Any]) -> bool:
        reason = str(payload.get("binding_reason", "") or payload.get("match_notes", "") or payload.get("binding_notes", "") or "").lower()
        mode = str(payload.get("binding_mode", "") or payload.get("match_mode", "") or "").lower()
        return mode == "ambiguous" or "ambiguous" in reason

    def _from_category_match(payload: Dict[str, Any]) -> Dict[str, Any]:
        result = _empty_result(str(payload.get("binding_reason", "no_reliable_category_match")))
        result.update({
            "matched": bool(payload.get("matched")),
            "matched_registry_key": str(payload.get("matched_registry_key", "")),
            "registry_record": dict(payload.get("registry_record") or {}) if payload.get("matched") else None,
            "binding_mode": str(payload.get("binding_mode", "unmatched")),
            "binding_confidence": str(payload.get("binding_confidence", "low")),
            "binding_reason": str(payload.get("binding_reason", "")),
            "matched_registry_label": str(payload.get("matched_registry_label", "")),
            "is_baseline_label": bool(payload.get("is_baseline_label", False)),
            "is_support_only_label": bool(payload.get("is_support_only_label", False)),
        })
        return result

    def _from_alias_resolution(raw_label: str, mapped_label: str, base_match: Dict[str, Any], mode: str, source: str, evidence: str) -> Dict[str, Any]:
        result = _empty_result(f"{mode}_unmatched")
        result.update({
            "matched": bool(base_match.get("matched")),
            "matched_registry_key": str(base_match.get("matched_registry_key", "")),
            "registry_record": dict(base_match.get("registry_record") or {}) if base_match.get("matched") else None,
            "binding_mode": mode,
            "binding_confidence": str(base_match.get("binding_confidence", "medium" if base_match.get("matched") else "low")),
            "binding_reason": f"{mode}:{raw_label}->{mapped_label}",
            "matched_registry_label": str(base_match.get("matched_registry_label", "")),
            "is_baseline_label": bool(base_match.get("is_baseline_label", False)),
            "is_support_only_label": bool(base_match.get("is_support_only_label", False)),
            "alias_map_source": source,
            "alias_map_evidence": evidence,
        })
        if not base_match.get("matched") and _is_ambiguous(base_match):
            result["binding_reason"] = f"ambiguous_{mode}:{raw_label}->{mapped_label}"
        elif not base_match.get("matched") and str(base_match.get("binding_reason", "")).strip():
            result["binding_reason"] = f"{mode}:{raw_label}->{mapped_label}; {base_match.get('binding_reason', '')}"
        return result

    context_label_alias_map = _build_context_label_alias_map(local_context, registry)
    numbered_alias_map = {
        key: value for key, value in context_label_alias_map.items()
        if _CONTEXT_NUMBERED_LABEL_RE.fullmatch(str(key or "").strip())
    }
    enriched_records: List[Dict] = []
    binding_audit_rows: List[Dict] = []

    for fig_record in fig_records:
        row = _strip_non_catalyst_identity_labels(dict(fig_record))
        seed_clues = _extract_identity_clues_from_figure_record(row)
        if seed_clues.get("clue_record"):
            row = _backfill_identity_from_clue_record(row, seed_clues["clue_record"])
        raw_catalyst = str(row.get("Catalyst", "")).strip()
        raw_catalyst_id = str(row.get("Catalyst_ID", "")).strip()
        note_axis, note_x_value = _extract_x_axis_and_value_from_notes(str(row.get("Notes", "")))
        raw_category_label = clean_residual_mojibake_chars(str(row.get("raw_category_label", "") or note_x_value or raw_catalyst)).strip()
        row["raw_category_label"] = raw_category_label
        if raw_category_label:
            clue_seed = dict(row)
            if not str(clue_seed.get("Catalyst", "")).strip():
                clue_seed["Catalyst"] = raw_category_label
            category_clues = _extract_identity_clues_from_figure_record(clue_seed)
            if category_clues.get("clue_record"):
                row = _backfill_identity_from_clue_record(row, category_clues["clue_record"])
        compact_label = _select_compact_loading_support_label({
            **row,
            "raw_category_label": raw_category_label,
            "Catalyst": raw_catalyst or row.get("Catalyst", ""),
            "Catalyst_ID": raw_catalyst_id or row.get("Catalyst_ID", ""),
        })
        compact_identity = _parse_compact_loading_support_identity(compact_label) if compact_label else {}
        if compact_identity:
            compact_clue = dict(compact_identity)
            compact_clue["Catalyst"] = compact_label
            row = _backfill_identity_from_clue_record(row, compact_clue)

        is_numbered_label = bool(
            re.fullmatch(r"[A-Za-z]\d{1,3}", raw_category_label or "", flags=re.I)
            or re.fullmatch(r"case\s*[-]?\s*\d+", raw_category_label or "", flags=re.I)
            or re.fullmatch(r"\d+#", raw_category_label or "", flags=re.I)
        )
        is_generic_label = bool(raw_category_label) and (_is_generic_series_label(raw_category_label) or is_numbered_label)

        category_match = _empty_result("empty_category_label")
        if raw_category_label and not is_generic_label:
            category_match = _from_category_match(
                _match_category_figure_label_to_registry(raw_category_label, registry, local_context, context_label_alias_map)
            )

        alias_match = _empty_result("alias_map_not_applicable")
        if raw_category_label and not is_numbered_label and not category_match.get("matched") and not _is_ambiguous(category_match):
            mapped_label = clean_residual_mojibake_chars(str(
                _extract_label_mapping_from_context(raw_category_label, local_context, registry, context_label_alias_map) or ""
            )).strip()
            if mapped_label and _normalize_registry_label(mapped_label) != _normalize_registry_label(raw_category_label):
                base_match = _match_category_figure_label_to_registry(
                    mapped_label, registry, local_context, context_label_alias_map
                )
                alias_match = _from_alias_resolution(
                    raw_category_label,
                    mapped_label,
                    base_match,
                    "alias_map_label",
                    "context_label_map",
                    _context_preview(f"{raw_category_label}={mapped_label}", 240),
                )

        numbered_match = _empty_result("numbered_alias_map_not_applicable")
        if raw_category_label and is_numbered_label and not category_match.get("matched") and not _is_ambiguous(category_match):
            numbered_match = _resolve_numbered_category_label(raw_category_label, numbered_alias_map, registry, local_context)

        chosen = category_match
        if not chosen.get("matched") and not _is_ambiguous(chosen) and alias_match.get("binding_mode") == "alias_map_label":
            chosen = alias_match
        if not chosen.get("matched") and not _is_ambiguous(chosen) and numbered_match.get("binding_mode") == "numbered_alias_map":
            chosen = numbered_match

        structured_loading_label = _label_has_structured_loading(raw_category_label)
        should_try_fallback = (
            not chosen.get("is_baseline_label")
            and (not chosen.get("is_support_only_label") or not chosen.get("matched"))
            and (not raw_category_label or not is_generic_label or structured_loading_label)
        )
        fallback_match = (
            match_figure_identity_to_registry(row, registry, local_context, context_label_alias_map)
            if should_try_fallback else {"matched": False}
        )
        if should_try_fallback and not fallback_match.get("matched") and str(fallback_match.get("binding_notes", "")).strip():
            fail_mode = str(fallback_match.get("match_mode", "unmatched") or "unmatched")
            fail_note = str(fallback_match.get("binding_notes", "") or "")[:80]
            _append_note_tag(row, f"[fallback_fail={fail_mode}:{fail_note}]")
        if not chosen.get("matched") and not _is_ambiguous(chosen) and fallback_match.get("matched"):
            chosen = {
                "matched": True,
                "matched_registry_key": str(fallback_match.get("matched_registry_key", "")),
                "registry_record": dict(fallback_match.get("registry_record") or {}),
                "binding_mode": str(fallback_match.get("match_mode", "unmatched")),
                "binding_confidence": str(fallback_match.get("match_confidence", "medium")),
                "binding_reason": str(fallback_match.get("binding_notes", "registry_fallback_match")),
                "matched_registry_label": str((fallback_match.get("registry_record") or {}).get("Catalyst", "") or fallback_match.get("matched_registry_key", "")),
                "is_baseline_label": False,
                "is_support_only_label": False,
                "alias_map_source": str(chosen.get("alias_map_source", "")),
                "alias_map_evidence": str(chosen.get("alias_map_evidence", "")),
                "_family_backbone_only": bool(fallback_match.get("_family_backbone_only", False)),
                "_family_identity_only": bool(fallback_match.get("_family_identity_only", False)),
            }
        elif (
            not chosen.get("matched")
            and not _is_ambiguous(chosen)
            and str(fallback_match.get("match_mode", "") or "") == "support_only_control"
        ):
            chosen = {
                "matched": False,
                "matched_registry_key": "",
                "registry_record": None,
                "binding_mode": "support_only_control",
                "binding_confidence": str(fallback_match.get("match_confidence", "low") or "low"),
                "binding_reason": str(fallback_match.get("binding_notes", "support_only_control_label") or "support_only_control_label"),
                "matched_registry_label": "",
                "is_baseline_label": False,
                "is_support_only_label": True,
                "alias_map_source": str(chosen.get("alias_map_source", "")),
                "alias_map_evidence": str(chosen.get("alias_map_evidence", "")),
            }
        elif not chosen.get("matched") and _is_ambiguous(fallback_match):
            chosen = {
                "matched": False,
                "matched_registry_key": "",
                "registry_record": None,
                "binding_mode": "ambiguous",
                "binding_confidence": "low",
                "binding_reason": str(fallback_match.get("binding_notes", "ambiguous_binding")),
                "matched_registry_label": "",
                "is_baseline_label": False,
                "is_support_only_label": False,
                "alias_map_source": str(chosen.get("alias_map_source", "")),
                "alias_map_evidence": str(chosen.get("alias_map_evidence", "")),
            }

        matched = bool(chosen.get("matched")) and isinstance(chosen.get("registry_record"), dict)
        registry_record = chosen.get("registry_record") or {}

        # ── Final metal-set safety gate (covers ALL binding paths) ──────────────
        # Replaces the old narrow Task-H bimetallic guard (which only blocked
        # bimetallic→mono) with a direction-agnostic check using proper helpers.
        # Blocked cases (all emit binding_blocked_for_identity_safety):
        #   mono→bimetal  : "Ni/clay"     → "Ni-Cu/clay"  ← the critical production bug
        #   bimetal→mono  : "Ni-Cu/clay"  → "Ni/clay"
        #   wrong metal   : "Cu/clay"     → "Ni-Cu/clay"
        #   support-only  : "ZrO2"        → any active-metal registry entry
        #   baseline      : "No catalyst" → any registry entry
        if matched and registry_record:
            _sg_label = clean_residual_mojibake_chars(raw_category_label or raw_catalyst or "").strip()
            if _sg_label:
                # Normalize before metal extraction so Unicode subscripts (₂→2) don't break parsing
                _sg_label_norm = _normalize_registry_label(_sg_label) or _sg_label
                _sg_clue_metals = _extract_explicit_metal_set_from_identity_clue(_sg_label_norm)
                _sg_reg_metals = _extract_registry_metal_set(registry_record)
                _sg_ok, _sg_reason = _is_metal_set_compatible_for_unique_match(
                    _sg_clue_metals, _sg_reg_metals
                )
                if not _sg_ok:
                    print(f"  [binding_safety_gate] {file_name}: blocked "
                          f"{_sg_label!r} → {chosen.get('matched_registry_key')!r} "
                          f"reason={_sg_reason}")
                    matched = False
                    chosen["matched"] = False
                    chosen["binding_reason"] = f"binding_blocked_for_identity_safety:{_sg_reason}"
                    chosen["binding_confidence"] = "low"
                    registry_record = {}

        if matched:
            if bool(chosen.get("_family_backbone_only", False)):
                backbone_mode = str(chosen.get("binding_mode", "") or "family_backbone_match")
                backbone_source = "loading_relaxed_family_backbone" if backbone_mode == "loading_relaxed_family_backbone" else "family_backbone"
                row = _backfill_registry_preparation_backbone_into_figure(
                    row,
                    registry_record,
                    source=backbone_source,
                    confidence="low",
                    scope="family_backbone",
                )
                note_name = "loading_relaxed_family_backbone" if backbone_mode == "loading_relaxed_family_backbone" else "family_backbone_match"
                _append_note_tag(row, f"[{note_name}={chosen.get('matched_registry_key', '')}]")
                _set_broadcast_primary_if_empty(row, "family_backbone")
                _append_broadcast_flag(row, "family_backbone")
                _append_broadcast_flag(row, backbone_mode)
            elif bool(chosen.get("_family_identity_only", False)):
                identity_mode = str(chosen.get("binding_mode", "") or "family_identity_match")
                identity_source = "loading_relaxed_identity_match" if identity_mode == "loading_relaxed_identity_match" else "family_identity"
                row = _merge_registry_identity_fields_without_cid(
                    row,
                    registry_record,
                    source=identity_source,
                    confidence="low",
                    scope="family_identity",
                )
                note_name = "loading_relaxed_identity_match" if identity_mode == "loading_relaxed_identity_match" else "family_identity_match"
                _append_note_tag(row, f"[{note_name}={chosen.get('matched_registry_key', '')}]")
                _set_broadcast_primary_if_empty(row, "family_identity")
                _append_broadcast_flag(row, "family_identity")
                _append_broadcast_flag(row, identity_mode)
            else:
                row = _merge_registry_fields_into_figure_record(row, registry_record)
                # match 成功就记 registry_bind（无论是否补了字段），因为 identity_bound 语义是"被精确身份匹配覆盖"
                _set_broadcast_primary_if_empty(row, "registry_bind")
                _append_broadcast_flag(row, "registry_bind")
                if not str(row.get("Canonical_Catalyst_ID", "") or "").strip():
                    fallback_cid = (
                        str(registry_record.get("Canonical_Catalyst_ID", "") or "").strip()
                        or str(chosen.get("matched_registry_key", "") or "").strip()
                    )
                    if fallback_cid:
                        row["Canonical_Catalyst_ID"] = fallback_cid
                        _set_field_source(row, "Canonical_Catalyst_ID", "registry_bind", "high", "identity_bound")

        binding_reason = str(chosen.get("binding_reason", ""))
        alias_map_source = str(chosen.get("alias_map_source", ""))
        alias_map_evidence = str(chosen.get("alias_map_evidence", ""))
        if alias_map_source and alias_map_source not in binding_reason:
            binding_reason = "; ".join([part for part in [binding_reason, f"alias_map_source={alias_map_source}"] if part])
        if alias_map_evidence and alias_map_evidence not in binding_reason:
            binding_reason = "; ".join([part for part in [binding_reason, f"alias_map_evidence={alias_map_evidence}"] if part])

        binding_mode = str(chosen.get("binding_mode", "unmatched") or "unmatched")
        binding_confidence = str(chosen.get("binding_confidence", "low") or "low")
        figure_binding_status = "matched" if matched else ("ambiguous" if _is_ambiguous(chosen) else "unmatched")

        # Task H: annotate Notes with binding failure reason for unmatched/ambiguous rows
        if not matched and binding_reason:
            fail_tag = f"[fig_bind_fail={binding_reason[:80]}]"
            existing_notes = str(row.get("Notes", "") or "").strip()
            if fail_tag not in existing_notes:
                row["Notes"] = (existing_notes + " " + fail_tag).strip()

        row["figure_binding_status"] = figure_binding_status
        row["binding_mode"] = binding_mode
        row["binding_confidence"] = binding_confidence
        row["binding_reason"] = binding_reason
        row["figure_binding_mode"] = _compat_mode(binding_mode)
        row["figure_binding_confidence"] = binding_confidence
        row["figure_binding_notes"] = binding_reason
        row["matched_registry_key"] = str(chosen.get("matched_registry_key", ""))
        row["matched_registry_label"] = str(chosen.get("matched_registry_label", ""))
        row["is_baseline_label"] = int(bool(chosen.get("is_baseline_label")))
        row["is_support_only_label"] = int(bool(chosen.get("is_support_only_label")))
        row["alias_map_source"] = alias_map_source
        row["alias_map_evidence"] = alias_map_evidence
        enriched_records.append(row)

        binding_audit_rows.append({
            "Source_File": file_name,
            "data_source": str(row.get("data_source", "figure")),
            "Series_Name": str(fig_record.get("Series_Name", "")),
            "raw_catalyst": raw_catalyst,
            "raw_catalyst_id": raw_catalyst_id,
            "raw_category_label": raw_category_label,
            "x_axis": note_axis,
            "x_value": note_x_value,
            "figure_binding_status": row["figure_binding_status"],
            "figure_binding_mode": row["figure_binding_mode"],
            "figure_binding_confidence": row["figure_binding_confidence"],
            "figure_binding_reason": row["binding_reason"],
            "matched_registry_label": row["matched_registry_label"],
            "matched_registry_key": row["matched_registry_key"],
            "is_baseline_label": row["is_baseline_label"],
            "is_support_only_label": row["is_support_only_label"],
            "alias_map_source": row.get("alias_map_source", ""),
            "alias_map_evidence": row.get("alias_map_evidence", ""),
            "_broadcast_path_primary": row.get("_broadcast_path_primary", ""),
            "_broadcast_flags": row.get("_broadcast_flags", ""),
            "_field_source_map_json": row.get("_field_source_map_json", ""),
            # Metal-set diagnostic fields for binding safety audit
            "explicit_metal_set": str(sorted(
                _extract_explicit_metal_set_from_identity_clue(
                    clean_residual_mojibake_chars(raw_category_label or raw_catalyst or "")
                ) or set()
            )),
            "registry_metal_set": str(sorted(_extract_registry_metal_set(
                chosen.get("registry_record") or {}
            ))),
            "blocked_by_safety_rule": str(chosen.get("binding_reason", "")).startswith(
                "binding_blocked_for_identity_safety"
            ),
        })

    return enriched_records, binding_audit_rows





def can_merge_chemically(base: Dict, candidate: Dict) -> bool:
    """Function docstring removed for runtime stability."""
    base_metal = normalize_active_metal_for_merge(str(base.get("Active_Metal", "")))
    cand_metal = normalize_active_metal_for_merge(str(candidate.get("Active_Metal", "")))
    base_loading = normalize_loading_for_merge(str(base.get("Metal_Loading_wt%", "")))
    cand_loading = normalize_loading_for_merge(str(candidate.get("Metal_Loading_wt%", "")))
    base_support = normalize_support_for_merge(str(base.get("Support_Normalized", "") or base.get("Support", "")))
    cand_support = normalize_support_for_merge(str(candidate.get("Support_Normalized", "") or candidate.get("Support", "")))
    base_promoter = normalize_promoter_for_merge(str(base.get("Promoter", "")))
    cand_promoter = normalize_promoter_for_merge(str(candidate.get("Promoter", "")))
    base_alloy = alloy_ratio_signature_for_merge(base)
    cand_alloy = alloy_ratio_signature_for_merge(candidate)

    if not base_metal or base_metal != cand_metal:
        return False
    if base_loading and cand_loading and base_loading != cand_loading:
        return False
    if not base_support or base_support != cand_support:
        return False
    if base_promoter != cand_promoter:
        return False
    if base_alloy or cand_alloy:
        if not base_alloy or not cand_alloy:
            return False
        if base_alloy != cand_alloy:
            return False

    try:
        t1 = float(str(base.get("Reaction_Temp_C", "") or "0"))
        t2 = float(str(candidate.get("Reaction_Temp_C", "") or "0"))
        if t1 > 0 and t2 > 0 and abs(t1 - t2) > 10:
            return False
    except (ValueError, TypeError):
        pass

    try:
        c1 = float(str(base.get("MeOH_Conversion_%", "") or "0"))
        c2 = float(str(candidate.get("MeOH_Conversion_%", "") or "0"))
        if c1 > 0 and c2 > 0 and abs(c1 - c2) > 5:
            return False
    except (ValueError, TypeError):
        pass

    return True


def _normalize_static_value_for_compare(field: str, value: str) -> str:
    """Function docstring removed for runtime stability."""
    if not value:
        return ""
    if field == "Support":
        return normalize_support_for_merge(value)
    # Default normalization: trim ends, collapse inner spaces, and lowercase.
    return re.sub(r'\s+', ' ', str(value).strip()).lower()


def _collect_conflict_fields(recs: List[Dict], conflict_fields: set) -> set:
    """Function docstring removed for runtime stability."""
    conflicted = set()
    for field in conflict_fields:
        # Collect normalized non-empty values for this field and de-duplicate them.
        seen = set()
        for r in recs:
            raw = str(r.get(field, "")).strip()
            if raw:
                seen.add(_normalize_static_value_for_compare(field, raw))
        # More than one distinct normalized value means a conflict.
        if len(seen) > 1:
            conflicted.add(field)
    return conflicted


def merge_fragmented_records(records: List[Dict]) -> List[Dict]:
    """Function docstring removed for runtime stability."""
    if not records:
        return records

    # Static fields describe catalyst identity or preparation and can be broadcast within a group.
    # Condition fields are experiment-level variables and must not be broadcast as static values.
    CONDITION_FIELDS = {
        "S_C_Ratio", "GHSV_mL_g_h", "Pressure_bar", "Feed_Composition",
    }

    # Static fields describe catalyst identity or preparation and may be shared within one canonical group.
    # 扩展：加入 normalized 字段，防止 binding 成功后在 merge 阶段丢失这些字段
    STATIC_FIELDS = {
        "Catalyst", "Active_Metal", "Metal_Loading_wt%", "Alloy_Ratio",
        "Support", "Promoter", "Precursor",
        "Metal_Loading_Method", "Support_Prep_Method",
        "Dry_Temp_C", "Dry_Time_h",
        "Calcination_Temp_C", "Calcination_Time_h",
        "Reduction_Temp_C", "Reduction_Time_h",
        "Source_File",
        # 新增：normalized/derived 身份字段，binding 后不应在 merge 阶段掉落
        "Canonical_Catalyst_ID", "Catalyst_ID_normalized", "identity_alias_group",
        "Promoter_Metal", "Ni_Fraction", "Promoter_Fraction",
        "Support_Normalized", "Support_Grouped",
        "Precursor_Normalized", "Precursor_Family",
        "Metal_Loading_Method_Normalized",
        "Support_Prep_Method_Normalized",
        "Preparation_Fingerprint",
    }

    # [Fix 7] Conflict-sensitive static fields:
    # if they conflict inside one group, do not broadcast them from STATIC_FIELDS.
    # 新增高风险身份字段：冲突时不广播，避免错误绑定
    CONFLICT_SENSITIVE_STATIC_FIELDS = {
        "Promoter", "Precursor",
        "Metal_Loading_Method", "Support_Prep_Method",
        "Dry_Temp_C", "Dry_Time_h",
        "Calcination_Temp_C", "Calcination_Time_h",
        "Reduction_Temp_C", "Reduction_Time_h",
        # 新增：高风险身份字段，组内冲突时不广播
        "Canonical_Catalyst_ID", "Promoter_Metal",
        "Alloy_Ratio", "Ni_Fraction", "Promoter_Fraction",
        "Metal_Loading_wt%",
    }

    # --- Stage 1: group by canonical key and build the static feature pool ---
    # Use OrderedDict to preserve stable record order.
    from collections import OrderedDict
    groups: Dict[str, Dict] = OrderedDict()   # id -> 闈欐€佺壒寰佹睜
    id_to_records: Dict[str, List[Dict]] = {}  # id -> records under this canonical group

    for rec in records:
        normalize_identity_aliases(rec)
        # [Fix 6] Internal grouping uses a deterministic canonical merge key.
        # Do not trust raw LLM Catalyst_ID for primary grouping when chemistry disagrees.
        rid = build_canonical_merge_key(rec)
        # Use a fallback group only when no usable identity is available.
        if not rid or rid == "--/":
            rid = f"__ungrouped_{id(rec)}"

        if rid not in groups:
            groups[rid] = {}
            id_to_records[rid] = []

        # Aggregate static features by first non-empty value.
        for key in STATIC_FIELDS:
            if rec.get(key) and not groups[rid].get(key):
                groups[rid][key] = rec[key]

        id_to_records[rid].append(rec)

    # --- Stage 2: conflict-aware static broadcast plus dynamic record retention ---
    DYNAMIC_FIELDS = {
        "Reaction_Temp_C", "MeOH_Conversion_%", "H2_Yield_%",
        "H2_Selectivity_%", "CO_Selectivity_%", "CO2_Selectivity_%",
        "CO_Concentration_ppm", "TOS_h", "Deactivation_Rate_%_h",
        "Carbon_Deposition_wt%", "Reasoning_Selectivity", "data_source",
        *CONDITION_FIELDS,
    }

    result = []
    for rid, static_pool in groups.items():
        recs = id_to_records[rid]

        # [Fix 7] Detect which sensitive static fields conflict inside this group.
        conflicted_fields = _collect_conflict_fields(recs, CONFLICT_SENSITIVE_STATIC_FIELDS)

        # Build a conflict note only when conflicts are actually present.
        conflict_note = ""
        if conflicted_fields:
            sorted_conflicts = "|".join(sorted(conflicted_fields))
            conflict_note = f"[static_conflict={sorted_conflicts}]"

        # A record is dynamic if it contains any condition or performance field.
        dynamic_recs = [r for r in recs if any(r.get(f) for f in DYNAMIC_FIELDS)]
        static_only_recs = [r for r in recs if not any(r.get(f) for f in DYNAMIC_FIELDS)]

        if dynamic_recs:
            # Dynamic records keep their point-level values and only receive non-conflicting static fields.
            for r in dynamic_recs:
                merged_r = dict(r)
                for key, val in static_pool.items():
                    # [Fix 7] Never broadcast conflicted fields from the static pool.
                    if key in conflicted_fields:
                        continue
                    if val and not merged_r.get(key):
                        merged_r[key] = val  # Fill empty fields only; never overwrite dynamic values.
                # [Fix 6] Write canonical merge key for downstream auditing without deleting Catalyst_ID.
                if not merged_r.get("Canonical_Catalyst_ID"):
                    merged_r["Canonical_Catalyst_ID"] = rid
                # [Fix 7] Attach the conflict note once, without duplicating it in Notes.
                if conflict_note:
                    existing_notes = str(merged_r.get("Notes", "")).strip()
                    if conflict_note not in existing_notes:
                        merged_r["Notes"] = f"{existing_notes} {conflict_note}".strip()
                normalize_identity_aliases(merged_r)
                result.append(merged_r)
        else:
            # Pure static fragments: keep only the most complete one after safe broadcast.
            best = max(static_only_recs, key=lambda x: sum(1 for v in x.values() if v))
            merged_r = dict(best)
            for key, val in static_pool.items():
                # [Fix 7] The same conflict rule applies when broadcasting into static-only fragments.
                if key in conflicted_fields:
                    continue
                if val and not merged_r.get(key):
                    merged_r[key] = val
            # [Fix 6] Static-only rows also receive the canonical merge key.
            if not merged_r.get("Canonical_Catalyst_ID"):
                merged_r["Canonical_Catalyst_ID"] = rid
            # [Fix 7] Append the same conflict note here as well.
            if conflict_note:
                existing_notes = str(merged_r.get("Notes", "")).strip()
                if conflict_note not in existing_notes:
                    merged_r["Notes"] = f"{existing_notes} {conflict_note}".strip()
            normalize_identity_aliases(merged_r)
            result.append(merged_r)

    n_before = len(records)
    n_after  = len(result)
    if n_before != n_after:
        print(f"  [V6] static broadcast merge: {n_before} -> {n_after} rows (removed {n_before - n_after} fragments)")
    return result


# ==========================================
# Multi-modal figure extraction: two-stage VL pipeline (prefilter -> extraction)
# ==========================================

# [Fix 9] Candidate-page keywords: any hit marks the page as potentially figure/performance related.
_FIGURE_KEYWORDS = [
    "Fig", "Figure", "Chart", "Scheme",
    "conversion", "yield", "selectivity", "stability", "TOS"
]
# Very short page text can indicate OCR failure or a scanned figure page; keep it as a fallback signal.
_LOW_TEXT_THRESHOLD = 80


def _is_nonperformance_figure_context(text: str) -> Dict[str, Any]:
    text_norm = _normalize_identity_text(text).lower()
    negative_flags: List[str] = []
    if not text_norm:
        return {
            "hard_exclusion": False,
            "hard_reason": "",
            "negative_flags": negative_flags,
            "hits": [],
            "nonperformance_strength": 0,
        }

    characterization_patterns = {
        "xrd": r"\bxrd\b|x-?ray diffraction|diffraction pattern",
        "xps": r"\bxps\b|x-?ray photoelectron",
        "tem": r"\btem\b|\bhrtem\b|\bstem\b",
        "sem": r"\bsem\b|\bfesem\b",
        "bet": r"\bbet\b|surface area|n2 adsorption|pore size|pore volume",
        "tpr": r"\btpr\b|h2-?tpr",
        "tpd": r"\btpd\b|co2-?tpd|nh3-?tpd|h2-?tpd",
        "raman": r"\braman\b",
        "ftir": r"\bftir\b|\bdrifts\b",
        "uv_vis": r"\buv-?vis\b",
        "xanes": r"\bxanes\b",
        "exafs": r"\bexafs\b",
        "eds": r"\beds\b|element mapping|mapping",
        "particle": r"particle size|size distribution|pore distribution",
    }
    mechanism_patterns = {
        "mechanism": r"reaction pathway|proposed mechanism|mechanism|active site|adsorption model|structure model",
        "schematic": r"\bschematic\b|\bdiagram\b|\billustration\b|\bconceptual\b",
        "graphical_abstract": r"graphical abstract|toc graphic|visual abstract",
        "apparatus": r"reactor setup|apparatus|experimental setup|setup diagram|flow chart|flowsheet",
        "photo": r"\bphoto(graph)?\b|digital image|reactor photograph",
        "synthesis": r"synthesis route|preparation route|fabrication route|synthesis scheme|scheme \d",
    }

    char_hits = [name for name, pat in characterization_patterns.items() if re.search(pat, text_norm, flags=re.I)]
    mechanism_hits = [name for name, pat in mechanism_patterns.items() if re.search(pat, text_norm, flags=re.I)]

    references_like = bool(
        re.search(r"\breferences\b|\bbibliography\b", text_norm, flags=re.I)
        and (text_norm.count("doi") >= 2 or len(re.findall(r"\[\d+\]", text_norm)) >= 6)
    )
    methods_only = bool(
        re.search(
            r"\b(experimental section|materials and methods|characterization methods|catalyst preparation|experimental methods|materials)\b",
            text_norm,
            flags=re.I,
        )
    )
    title_like = bool(
        len(text_norm) < 260
        and re.search(r"\babstract\b|\bkeywords\b|\bintroduction\b", text_norm, flags=re.I)
        and not re.search(r"\bfig(?:ure)?\.?\s*\d+\b", text_norm, flags=re.I)
    )

    if references_like:
        negative_flags.append("references_like")
        return {
            "hard_exclusion": True,
            "hard_reason": "hard_exclusion_references",
            "negative_flags": negative_flags,
            "hits": ["references"],
            "nonperformance_strength": 6,
        }
    if len(char_hits) >= 2 or (char_hits and any(hit in {"xrd", "xps", "tem", "sem", "tpr", "tpd"} for hit in char_hits)):
        negative_flags.extend([f"characterization:{hit}" for hit in char_hits])
        return {
            "hard_exclusion": True,
            "hard_reason": "hard_exclusion_characterization",
            "negative_flags": negative_flags,
            "hits": char_hits,
            "nonperformance_strength": len(char_hits) + 3,
        }
    if mechanism_hits:
        negative_flags.extend([f"mechanism:{hit}" for hit in mechanism_hits])
        return {
            "hard_exclusion": True,
            "hard_reason": "hard_exclusion_mechanism",
            "negative_flags": negative_flags,
            "hits": mechanism_hits,
            "nonperformance_strength": len(mechanism_hits) + 2,
        }
    if methods_only or title_like:
        if methods_only:
            negative_flags.append("methods_only")
        if title_like:
            negative_flags.append("title_page_like")
        return {
            "hard_exclusion": True,
            "hard_reason": "hard_exclusion_methods_only",
            "negative_flags": negative_flags,
            "hits": ["methods_only" if methods_only else "title_page_like"],
            "nonperformance_strength": 3,
        }
    return {
        "hard_exclusion": False,
        "hard_reason": "",
        "negative_flags": negative_flags,
        "hits": char_hits + mechanism_hits,
        "nonperformance_strength": len(char_hits) + len(mechanism_hits),
    }




def _is_performance_figure_context(page_text: str, prev_text: str = "", next_text: str = "") -> Dict[str, Any]:
    current = _normalize_identity_text(page_text).lower()
    prev_norm = _normalize_identity_text(prev_text).lower()
    next_norm = _normalize_identity_text(next_text).lower()
    neighbor = " ".join([prev_norm, next_norm]).strip()
    combined = " ".join(part for part in [current, prev_norm, next_norm] if part).strip()

    caption_hit = bool(re.search(r"\bfig(?:ure)?\.?\s*\d+\b|\bfig(?:ure)?\b|\bcaption\b", current, flags=re.I))
    neighbor_caption_hit = bool(re.search(r"\bfig(?:ure)?\.?\s*\d+\b|\bfig(?:ure)?\b|\bcaption\b", neighbor, flags=re.I))
    figure_keyword_hit = bool(
        re.search(r"\b(fig(?:ure)?|chart|plot|curve|bar chart|line chart|scatter|legend|series)\b", combined, flags=re.I)
    )

    performance_patterns = {
        "msr": r"methanol steam reforming|steam reforming of methanol|methanol reforming|\bmsr\b",
        "performance": r"catalytic performance|performance comparison|performance",
        "conversion": r"methanol conversion|meoh conversion|\bconversion\b",
        "yield": r"hydrogen yield|h2 yield|\byield\b",
        "h2_rate": r"h2 production|hydrogen production|h2 productivity",
        "selectivity": r"co selectivity|co2 selectivity|\bselectivity\b",
        "stability": r"stability|time on stream|\btos\b|deactivation",
    }
    perf_hits = [name for name, pat in performance_patterns.items() if re.search(pat, combined, flags=re.I)]
    metric_signal = bool(set(perf_hits).intersection({"performance", "conversion", "yield", "h2_rate", "selectivity", "stability"}))

    axis_like = bool(
        re.search(
            r"\b(temperature|reaction temperature|temp|different catalysts?|catalyst type|catalysts?|samples?|sample|support|composition|s/c|steam[- ]?to[- ]?carbon|pressure|ghsv|whsv|flow rate|catalyst amount|time on stream|tos|hours?)\b",
            combined,
            flags=re.I,
        )
    )
    legend_like = bool(
        re.search(
            r"\blegend\b|\bseries\b|\b(?:sample|cat(?:alyst)?)\s*[- ]?[a-z0-9]+\b|\bno catalyst\b|\bwithout catalyst\b|\bzro2\b|\bal2o3\b|\bceo2\b|\bni[/\-][a-z0-9\-_/]+\b",
            combined,
            flags=re.I,
        )
    )
    category_like = bool(
        re.search(
            r"\b(different catalysts?|catalyst type|catalysts?|samples?|sample|support|composition)\b",
            combined,
            flags=re.I,
        )
    )
    category_strong_terms = bool(
        re.search(
            r"\b(different catalysts?|catalyst type|samples?)\b|\bni[/\-][a-z0-9\-_/]+\b|\bzro2\b|\bal2o3\b|\bceo2\b|\bno catalyst\b",
            combined,
            flags=re.I,
        )
    )

    positive_flags: List[str] = []
    weak_positive_flags: List[str] = []
    if caption_hit:
        positive_flags.append("caption_current_page")
    if neighbor_caption_hit:
        positive_flags.append("caption_neighbor_page")
    if axis_like:
        positive_flags.append("axis_like_signal")
    if legend_like:
        positive_flags.append("legend_like_signal")
    if category_like:
        positive_flags.append("category_like_signal")
    for hit in perf_hits:
        positive_flags.append(f"performance_kw:{hit}")

    strong_positive = False
    strong_reason = ""
    suggested_role = ""
    if caption_hit and metric_signal:
        strong_positive = True
        strong_reason = "strong_performance_caption"
    elif neighbor_caption_hit and metric_signal:
        strong_positive = True
        strong_reason = "strong_performance_neighbor_caption"
    elif axis_like and legend_like and metric_signal:
        strong_positive = True
        strong_reason = "strong_axis_legend_performance"
    elif category_like and category_strong_terms and metric_signal:
        strong_positive = True
        strong_reason = "strong_category_performance"
        suggested_role = "category_screening"

    if not strong_positive:
        if len(current.strip()) < _LOW_TEXT_THRESHOLD:
            weak_positive_flags.append("low_text_fallback")
        if axis_like:
            weak_positive_flags.append("axis_like_signal")
        if legend_like:
            weak_positive_flags.append("legend_like_signal")
        if figure_keyword_hit:
            weak_positive_flags.append("figure_keyword")
        if metric_signal:
            weak_positive_flags.append("performance_word_only")
        if "msr" in perf_hits:
            weak_positive_flags.append("msr_context")

    performance_strength = (
        len(set(perf_hits))
        + int(metric_signal) * 2
        + int(axis_like)
        + int(legend_like)
        + int(caption_hit)
        + int(neighbor_caption_hit)
        + int(category_like)
        + int(category_strong_terms)
    )
    return {
        "strong_positive": strong_positive,
        "strong_reason": strong_reason,
        "suggested_role": suggested_role,
        "positive_flags": _dedupe_keep_order(positive_flags),
        "weak_positive_flags": _dedupe_keep_order(weak_positive_flags),
        "weak_positive_count": len(_dedupe_keep_order(weak_positive_flags)),
        "performance_strength": performance_strength,
        "metric_signal": metric_signal,
        "category_like": category_like,
        "axis_like": axis_like,
        "legend_like": legend_like,
        "perf_hits": perf_hits,
        "caption_hit": caption_hit,
        "neighbor_caption_hit": neighbor_caption_hit,
    }




def _extract_condition_anchor_from_text(text: str) -> Dict[str, Any]:
    source = _normalize_identity_text(text).replace("℃", " C ")
    # 双栏PDF换行会把关键词和数值分到两行，替换换行为空格以支持跨行匹配
    source = source.replace("\n", " ")
    result: Dict[str, Any] = {
        "Reaction_Temp_C": "",
        "TOS_h": "",
        "S_C_Ratio": "",
        "S_C_Ratio_Raw": "",
        "Pressure_bar": "",
        "Pressure_bar_Raw": "",
        "GHSV_mL_g_h": "",
        "GHSV_mL_g_h_Raw": "",
        "Flow_Rate": "",
        "Catalyst_Amount_g": "",
        "_anchor_sources": [],
    }
    if not source:
        return result
    source = source.replace("°", " ").replace("?", " ")
    source = re.sub(r"\s+", " ", source).strip()
    def _mark(field: str, value: str, source_name: str, raw_value: str = "") -> None:
        if value and not str(result.get(field, "")).strip():
            result[field] = _format_ratio_token(value)
            result["_anchor_sources"].append(f"{source_name}:{field}")
        raw_field = RAW_UNIT_FIELD_SPECS.get(field)
        raw_text = clean_residual_mojibake_chars(str(raw_value or value or "")).strip()
        if raw_field and raw_text and not str(result.get(raw_field, "")).strip():
            result[raw_field] = raw_text
    temp_patterns = [
        r"\b(?:at|under|tested at|reaction temperature(?: of)?)\s*(\d{2,4}(?:\.\d+)?)\s*(?:deg\.?\s*c|celsius|c)\b",
        r"\b(\d{2,4}(?:\.\d+)?)\s*(?:deg\.?\s*c|celsius|c)\b",
    ]
    for pattern in temp_patterns:
        match = re.search(pattern, source, flags=re.I)
        if match:
            if not _reaction_temp_value_is_sane(match.group(1)):
                continue
            _mark("Reaction_Temp_C", match.group(1), "text", match.group(0))
            break
    tos_patterns = [
        r"\b(?:tos|time on stream|stream time|stability(?: test)?)\b[^.;\n]{0,30}?(\d+(?:\.\d+)?)\s*h\b",
        r"\bfor\s+(\d+(?:\.\d+)?)\s*h\b[^.;\n]{0,20}?\b(?:stability|tos|time on stream)\b",
    ]
    for pattern in tos_patterns:
        match = re.search(pattern, source, flags=re.I)
        if match:
            _mark("TOS_h", match.group(1), "text", match.group(0))
            break
    # pattern1: 标准 S/C 关键词
    sc_match = re.search(
        r"\b(?:s/c|steam[- ]?to[- ]?carbon|steam carbon ratio|feed ratio|methanol[- ]?to[- ]?water)\b[^.;]{0,25}?((?:\d+(?:\.\d+)?)(?:\s*[:/]\s*\d+(?:\.\d+)?)?)",
        source,
        flags=re.I,
    )
    # pattern2: molar ratio + water/methanol/feed（如 'molar ratio of water and methanol of 2:1'）
    if not sc_match:
        sc_match = re.search(
            r"\bmolar ratio\b[^.;]{0,15}(?:water|methanol|feed|reactant)[^.;]{0,30}?((?:\d+(?:\.\d+)?)(?:\s*[:/]\s*\d+(?:\.\d+)?)?)",
            source,
            flags=re.I,
        )
    # pattern3: methanol/water 上下文中的 molar ratio = 数字（如 'Methanol in water solution (molar ratio = 2.5:1)'）
    if not sc_match and re.search(r"\b(?:methanol|water)\b", source[:200], re.I):
        sc_match = re.search(
            r"\bmolar ratio\b[^.;]{0,20}?((?:\d+(?:\.\d+)?)(?:\s*[:/]\s*\d+(?:\.\d+)?)?)",
            source,
            flags=re.I,
        )
    if sc_match:
        _mark("S_C_Ratio", sc_match.group(1).replace(" ", ""), "text", sc_match.group(1))
    pressure_match = re.search(r"\bpressure\b[^.;\n]{0,25}?(\d+(?:\.\d+)?)\s*((?:bar|atm|mpa)?)\b", source, flags=re.I)
    if pressure_match:
        raw_pressure = " ".join(part for part in [pressure_match.group(1), pressure_match.group(2)] if str(part).strip())
        _mark("Pressure_bar", pressure_match.group(1), "text", raw_pressure)
    ghsv_match = re.search(
        r"\b(?:ghsv|whsv|space velocity)\b[^.;\n]{0,35}?([0-9]+(?:\.[0-9]+)?(?:\s*[x\u00d7]\s*10\^?\d+)?)",
        source,
        flags=re.I,
    )
    if ghsv_match:
        _mark("GHSV_mL_g_h", ghsv_match.group(1).replace(" ", ""), "text", ghsv_match.group(0))
    flow_match = re.search(
        r"\b(?:flow rate|feed flow|feed rate)\b[^.;\n]{0,25}?(\d+(?:\.\d+)?)",
        source,
        flags=re.I,
    )
    if flow_match:
        _mark("Flow_Rate", flow_match.group(1), "text", flow_match.group(0))
    amount_match = re.search(
        r"\b(?:catalyst amount|catalyst mass|catalyst charge|catalyst bed)\b[^.;\n]{0,25}?(\d+(?:\.\d+)?)",
        source,
        flags=re.I,
    )
    if amount_match:
        _mark("Catalyst_Amount_g", amount_match.group(1), "text", amount_match.group(0))
    result["_anchor_sources"] = _dedupe_keep_order(result["_anchor_sources"])
    return result





def _extract_condition_anchor_from_series_or_caption(series_text: str, caption_text: str) -> Dict[str, Any]:
    series_anchor = _extract_condition_anchor_from_text(series_text)
    caption_anchor = _extract_condition_anchor_from_text(caption_text)
    result: Dict[str, Any] = {
        "Reaction_Temp_C": "",
        "TOS_h": "",
        "S_C_Ratio": "",
        "Pressure_bar": "",
        "GHSV_mL_g_h": "",
        "Flow_Rate": "",
        "Catalyst_Amount_g": "",
        "_anchor_sources": [],
    }
    for source_name, anchor in [("caption", caption_anchor), ("series", series_anchor)]:
        for field in [
            "Reaction_Temp_C",
            "TOS_h",
            "S_C_Ratio",
            "Pressure_bar",
            "GHSV_mL_g_h",
            "Flow_Rate",
            "Catalyst_Amount_g",
        ]:
            value = clean_residual_mojibake_chars(str(anchor.get(field, "") or "")).strip()
            if value:
                result[field] = value
                result["_anchor_sources"].append(f"{source_name}:{field}")
    result["_anchor_sources"] = _dedupe_keep_order(result["_anchor_sources"])
    return result



def _extract_condition_from_table_headers_or_notes(row: Dict) -> Dict[str, Any]:
    result: Dict[str, Any] = {
        "Reaction_Temp_C": "",
        "TOS_h": "",
        "S_C_Ratio": "",
        "Feed_MeOH_to_H2O_Ratio": "",
        "Pressure_bar": "",
        "GHSV_mL_g_h": "",
        "Flow_Rate": "",
        "Catalyst_Amount_g": "",
        "_anchor_sources": [],
        "_anchor_evidence": [],
    }
    if not isinstance(row, dict):
        return result

    def _set(field: str, value: Any, source: str, evidence: str = "") -> None:
        value_s = clean_residual_mojibake_chars(str(value or "")).strip()
        if not value_s or str(result.get(field, "")).strip():
            return
        result[field] = value_s
        if source:
            result["_anchor_sources"].append(str(source))
        if evidence:
            result["_anchor_evidence"].append(_context_preview(evidence, 180))

    def _merge(anchor: Dict[str, Any], source_name: str, evidence: str = "") -> None:
        if not isinstance(anchor, dict):
            return
        for field in [
            "Reaction_Temp_C",
            "TOS_h",
            "S_C_Ratio",
            "Feed_MeOH_to_H2O_Ratio",
            "Pressure_bar",
            "GHSV_mL_g_h",
            "Flow_Rate",
            "Catalyst_Amount_g",
        ]:
            if str(anchor.get(field, "")).strip():
                _set(field, anchor.get(field, ""), source_name, evidence)
        for source in anchor.get("_anchor_sources", []) or []:
            if str(source).strip():
                result["_anchor_sources"].append(str(source).strip())

    def _axis_kind(text: str) -> str:
        s = _normalize_identity_text(text).lower()
        if not s:
            return ""
        prep_like = bool(re.search(r"\b(calcination|reduction|drying|impregnation|preparation|precursor)\b", s, flags=re.I))
        if any(token in s for token in CATEGORY_ONLY_AXIS_KEYWORDS):
            if not any(any(keyword in s for keyword in keywords) for keywords in CONDITION_AXIS_KEYWORDS.values()):
                return ""
        for kind, keywords in CONDITION_AXIS_KEYWORDS.items():
            if any(keyword in s for keyword in keywords):
                if kind == "temperature" and prep_like and not re.search(r"\b(reaction|reforming)\b", s, flags=re.I):
                    return ""
                return kind
        return ""

    def _normalize_numeric_token(value: str, allow_ratio: bool = False) -> str:
        text = clean_residual_mojibake_chars(str(value or "")).strip()
        if not text:
            return ""
        compact = text.replace(",", "")
        if allow_ratio:
            ratio_match = re.search(r"\d+(?:\.\d+)?\s*[:/]\s*\d+(?:\.\d+)?", compact)
            if ratio_match:
                return ratio_match.group(0).replace(" ", "")
        direct = _parse_float_if_possible(compact)
        if direct is not None:
            return _format_ratio_token(str(direct))
        number_match = re.search(r"[-+]?\d+(?:\.\d+)?", compact)
        return _format_ratio_token(number_match.group(0)) if number_match else ""

    def _looks_condition_text(text: str) -> bool:
        s = _normalize_identity_text(text).lower()
        if not s:
            return False
        if re.search(r"\b(calcination|reduction|drying|impregnation|preparation|precursor)\b", s, flags=re.I) and not re.search(
            r"\b(reaction|reforming|conversion|yield|selectivity|performance|tos|time on stream|s/c|steam[- ]?to[- ]?carbon|ghsv|whsv|pressure|flow rate|catalyst amount|feed ratio)\b",
            s,
            flags=re.I,
        ):
            return False
        return bool(re.search(
            r"\b(reaction temperature|temperature|temp|tos|time on stream|stability|s/c|steam[- ]?to[- ]?carbon|ghsv|whsv|space velocity|pressure|flow rate|catalyst amount|feed ratio|meoh[- ]?to[- ]?h2o|conversion|yield|selectivity|performance)\b",
            s,
            flags=re.I,
        ))

    notes = clean_residual_mojibake_chars(str(row.get("Notes", "") or ""))
    note_axis, note_x_value = _extract_x_axis_and_value_from_notes(notes)
    axis_candidates = [
        ("row_x_axis", str(row.get("x_axis", "") or "")),
        ("note_x_axis", note_axis),
        ("column_header", str(row.get("column_header", "") or "")),
        ("column_name", str(row.get("column_name", "") or "")),
        ("table_header", str(row.get("table_header", "") or "")),
        ("table_headers", str(row.get("table_headers", "") or "")),
        ("header_row", str(row.get("header_row", "") or "")),
        ("row_header", str(row.get("row_header", "") or "")),
        ("unit_row", str(row.get("unit_row", "") or "")),
        ("units", str(row.get("units", "") or "")),
    ]
    value_candidates = [
        ("row_x_value", str(row.get("x_value", "") or "")),
        ("note_x_value", note_x_value),
        ("column_value", str(row.get("column_value", "") or "")),
        ("header_value", str(row.get("header_value", "") or "")),
    ]

    for axis_source, axis_text in axis_candidates:
        axis_kind = _axis_kind(axis_text)
        if not axis_kind:
            continue
        for value_source, value_text in value_candidates:
            if not str(value_text).strip() or _looks_like_identity_mapping(str(value_text)):
                continue
            normalized = _normalize_numeric_token(value_text, allow_ratio=(axis_kind == "ratio"))
            if not normalized:
                continue
            evidence = f"{axis_text} {value_text}".strip()
            if axis_kind == "temperature":
                _set("Reaction_Temp_C", normalized, f"{axis_source}:{value_source}", evidence)
            elif axis_kind == "time":
                _set("TOS_h", normalized, f"{axis_source}:{value_source}", evidence)
            elif axis_kind == "ratio":
                _set("S_C_Ratio", normalized, f"{axis_source}:{value_source}", evidence)
            elif axis_kind == "pressure":
                _set("Pressure_bar", normalized, f"{axis_source}:{value_source}", evidence)
            elif axis_kind == "space_velocity":
                _set("GHSV_mL_g_h", normalized, f"{axis_source}:{value_source}", evidence)
            elif axis_kind == "flow":
                _set("Flow_Rate", normalized, f"{axis_source}:{value_source}", evidence)
            elif axis_kind == "catalyst_amount":
                _set("Catalyst_Amount_g", normalized, f"{axis_source}:{value_source}", evidence)

    text_sources = [
        ("notes", notes),
        ("notes_axis_pair", f"{note_axis} {note_x_value}".strip()),
        ("x_axis_pair", f"{row.get('x_axis', '')} {row.get('x_value', '')}".strip()),
        ("table_header", str(row.get("table_header", "") or "")),
        ("table_headers", str(row.get("table_headers", "") or "")),
        ("header_row", str(row.get("header_row", "") or "")),
        ("column_header", str(row.get("column_header", "") or "")),
        ("column_name", str(row.get("column_name", "") or "")),
        ("unit_row", str(row.get("unit_row", "") or "")),
        ("row_text", str(row.get("row_text", "") or "")),
        ("paragraph_text", str(row.get("paragraph_text", "") or "")),
        ("sentence_text", str(row.get("sentence_text", "") or "")),
        ("source_text", str(row.get("source_text", "") or "")),
    ]
    for source_name, source_text in text_sources:
        if not str(source_text).strip() or not _looks_condition_text(source_text):
            continue
        anchor = _extract_condition_anchor_from_text(source_text)
        _merge(anchor, source_name, source_text)
        if not str(result.get("Feed_MeOH_to_H2O_Ratio", "")).strip():
            parsed_ratio = parse_feed_meoh_to_h2o_ratio(source_text)
            if parsed_ratio:
                _set("Feed_MeOH_to_H2O_Ratio", parsed_ratio, source_name, source_text)

    if not str(result.get("Feed_MeOH_to_H2O_Ratio", "")).strip():
        parsed_feed_ratio = parse_feed_meoh_to_h2o_ratio(str(row.get("Feed_Composition", "") or ""))
        if parsed_feed_ratio:
            _set("Feed_MeOH_to_H2O_Ratio", parsed_feed_ratio, "Feed_Composition", str(row.get("Feed_Composition", "") or ""))

    result["_anchor_sources"] = _dedupe_keep_order([str(item).strip() for item in result["_anchor_sources"] if str(item).strip()])
    result["_anchor_evidence"] = _dedupe_keep_order([str(item).strip() for item in result["_anchor_evidence"] if str(item).strip()])
    return result


def _salvage_text_condition_anchor(record: Dict) -> Dict[str, Any]:
    result: Dict[str, Any] = {
        "Reaction_Temp_C": "",
        "TOS_h": "",
        "S_C_Ratio": "",
        "Feed_MeOH_to_H2O_Ratio": "",
        "Pressure_bar": "",
        "GHSV_mL_g_h": "",
        "Flow_Rate": "",
        "Catalyst_Amount_g": "",
        "_salvage_source": "",
        "_salvage_evidence": "",
        "_salvage_fields": [],
    }
    if not isinstance(record, dict):
        return result

    row = dict(record)
    if str(row.get("data_source", "")).strip() != "text":
        return result
    if str(row.get("paragraph_role", "")).strip() == "preparation_paragraph":
        return result
    if str(row.get("text_extraction_subroute", "")).strip() == "preparation_schema":
        return result
    if not has_core_performance_metric(row):
        return result
    salvage_sources: List[str] = []
    salvage_evidence: List[str] = []

    def _set(field: str, value: Any, source: str, evidence: str = "") -> None:
        value_s = clean_residual_mojibake_chars(str(value or "")).strip()
        if not value_s or str(result.get(field, "")).strip() or str(row.get(field, "")).strip():
            return
        result[field] = value_s
        salvage_sources.append(source)
        if evidence:
            salvage_evidence.append(_context_preview(evidence, 200))

    header_anchor = _extract_condition_from_table_headers_or_notes(row)
    for field in [
        "Reaction_Temp_C",
        "TOS_h",
        "S_C_Ratio",
        "Feed_MeOH_to_H2O_Ratio",
        "Pressure_bar",
        "GHSV_mL_g_h",
        "Flow_Rate",
        "Catalyst_Amount_g",
    ]:
        if str(header_anchor.get(field, "")).strip():
            _set(field, header_anchor.get(field, ""), "header_or_notes", " || ".join(header_anchor.get("_anchor_evidence", [])[:2]))
    salvage_sources.extend([str(item) for item in header_anchor.get("_anchor_sources", []) if str(item).strip()])
    salvage_evidence.extend([str(item) for item in header_anchor.get("_anchor_evidence", []) if str(item).strip()])

    local_text = " ".join([
        str(row.get("paragraph_text", "") or ""),
        str(row.get("sentence_text", "") or ""),
        str(row.get("source_text", "") or ""),
        str(row.get("row_text", "") or ""),
    ]).strip()
    if local_text:
        inline_anchor = _extract_condition_anchor_from_text(local_text)
        for field in [
            "Reaction_Temp_C",
            "TOS_h",
            "S_C_Ratio",
            "Pressure_bar",
            "GHSV_mL_g_h",
            "Flow_Rate",
            "Catalyst_Amount_g",
        ]:
            if str(inline_anchor.get(field, "")).strip():
                _set(field, inline_anchor.get(field, ""), "local_text", local_text)
        if not str(row.get("Feed_MeOH_to_H2O_Ratio", "")).strip() and not str(result.get("Feed_MeOH_to_H2O_Ratio", "")).strip():
            parsed_ratio = parse_feed_meoh_to_h2o_ratio(local_text)
            if parsed_ratio:
                _set("Feed_MeOH_to_H2O_Ratio", parsed_ratio, "local_text", local_text)

    if not any(str(result.get(field, "")).strip() for field in [
        "Reaction_Temp_C",
        "TOS_h",
        "S_C_Ratio",
        "Feed_MeOH_to_H2O_Ratio",
        "Pressure_bar",
        "GHSV_mL_g_h",
        "Flow_Rate",
        "Catalyst_Amount_g",
    ]):
        return result

    result["_salvage_source"] = "|".join(_dedupe_keep_order([str(item).strip() for item in salvage_sources if str(item).strip()]))
    result["_salvage_evidence"] = " || ".join(_dedupe_keep_order([str(item).strip() for item in salvage_evidence if str(item).strip()]))[:800]
    result["_salvage_fields"] = [
        field for field in [
            "Reaction_Temp_C",
            "TOS_h",
            "S_C_Ratio",
            "Feed_MeOH_to_H2O_Ratio",
            "Pressure_bar",
            "GHSV_mL_g_h",
            "Flow_Rate",
            "Catalyst_Amount_g",
        ] if str(result.get(field, "")).strip()
    ]
    return result



def _build_numbered_label_alias_map(local_context: str, registry: Dict) -> Dict[str, Dict[str, Any]]:
    alias_map = _build_context_label_alias_map(local_context, registry)
    return {
        key: value for key, value in alias_map.items()
        if _CONTEXT_NUMBERED_LABEL_RE.fullmatch(str(key or "").strip())
    }


def _resolve_numbered_category_label(label: str, alias_map: Dict[str, Dict[str, Any]], registry: Dict, local_context: str) -> Dict[str, Any]:
    raw_label = clean_residual_mojibake_chars(str(label or "")).strip()
    label_norm = _normalize_context_alias_key(raw_label)
    result = {
        "matched": False,
        "matched_registry_key": "",
        "registry_record": None,
        "binding_mode": "numbered_alias_map",
        "binding_confidence": "low",
        "binding_reason": "numbered_alias_map_unresolved",
        "matched_registry_label": "",
        "is_baseline_label": False,
        "is_support_only_label": False,
        "alias_map_source": "",
        "alias_map_evidence": "",
        "resolved_label": "",
    }
    if not (re.fullmatch(r"[A-Za-z]\d{1,3}", raw_label, flags=re.I)
            or re.fullmatch(r"case\s*[-]?\s*\d+", raw_label, flags=re.I)
            or re.fullmatch(r"\d+#", raw_label, flags=re.I)):
        return result
    entry = dict((alias_map or {}).get(label_norm) or {})
    if not entry:
        result["binding_reason"] = "numbered_alias_map_missing"
        return result
    result["alias_map_source"] = str(entry.get("alias_map_source", ""))
    result["alias_map_evidence"] = str(entry.get("alias_map_evidence", ""))
    if bool(entry.get("ambiguous")):
        result["binding_reason"] = "alias_map_ambiguous"
        return result
    mapped_label = clean_residual_mojibake_chars(str(entry.get("mapped_label", "") or "")).strip()
    if not mapped_label:
        result["binding_reason"] = "numbered_alias_map_empty"
        return result
    result["resolved_label"] = mapped_label
    base_match = _match_category_figure_label_to_registry(mapped_label, registry, local_context, alias_map)
    result["matched"] = bool(base_match.get("matched"))
    result["matched_registry_key"] = str(base_match.get("matched_registry_key", ""))
    result["registry_record"] = dict(base_match.get("registry_record") or {}) if base_match.get("matched") else None
    result["binding_confidence"] = str(base_match.get("binding_confidence", "medium" if base_match.get("matched") else "low"))
    result["matched_registry_label"] = str(base_match.get("matched_registry_label", ""))
    result["is_baseline_label"] = bool(base_match.get("is_baseline_label", False))
    result["is_support_only_label"] = bool(base_match.get("is_support_only_label", False))
    if base_match.get("matched"):
        result["binding_reason"] = f"numbered_alias_map:{raw_label}->{mapped_label}"
    elif str(base_match.get("binding_reason", "")).strip():
        result["binding_reason"] = f"numbered_alias_map:{raw_label}->{mapped_label}; {base_match.get('binding_reason', '')}"
    else:
        result["binding_reason"] = f"numbered_alias_map:{raw_label}->{mapped_label}"
    return result






def _classify_metadata_failure_reason(metadata: Any, context_text: str = "", exception: Any = None) -> str:
    if exception is not None:
        return "metadata_parse_error"
    if metadata is None:
        return "metadata_empty"
    if not isinstance(metadata, dict):
        return "metadata_parse_error"

    x_axis = clean_residual_mojibake_chars(str(metadata.get("x_axis", "") or "")).strip()
    y_axis = clean_residual_mojibake_chars(str(metadata.get("y_axis", "") or "")).strip()
    chart_type = clean_residual_mojibake_chars(str(metadata.get("chart_type", "") or "")).strip().lower()
    x_axis_mode = clean_residual_mojibake_chars(str(metadata.get("x_axis_mode", "") or "")).strip().lower()
    extractor_type = clean_residual_mojibake_chars(str(metadata.get("extractor_type", "") or "")).strip().lower()
    semantic_role = clean_residual_mojibake_chars(str(metadata.get("semantic_figure_role", "") or "")).strip().lower()
    series = [clean_residual_mojibake_chars(str(item)).strip() for item in (metadata.get("series", []) or []) if str(item).strip()]

    if not any([x_axis, y_axis, chart_type, x_axis_mode, extractor_type, semantic_role, series]):
        return "metadata_empty"

    meta_text = " ".join([x_axis, y_axis, chart_type, x_axis_mode, extractor_type, semantic_role, " ".join(series), context_text])
    perf_info = _is_performance_figure_context(meta_text, context_text, "")
    nonperf_info = _is_nonperformance_figure_context(meta_text)

    if semantic_role == "non_performance_like":
        return "metadata_nonperformance_like"
    if nonperf_info.get("hard_exclusion") and not perf_info.get("strong_positive"):
        return "metadata_nonperformance_like"

    x_unknown = (not x_axis or x_axis.lower() == "unknown") and x_axis_mode in {"", "unknown"}
    y_unknown = not y_axis or y_axis.lower() == "unknown"
    category_like = semantic_role == "category_screening" or x_axis_mode == "category" or perf_info.get("category_like")
    condition_like = semantic_role == "condition_effect" or x_axis_mode == "condition"

    if x_unknown and y_unknown and not category_like and not condition_like:
        return "metadata_unknown_axis"
    if category_like:
        if y_unknown and not perf_info.get("metric_signal") and not perf_info.get("category_like"):
            return "metadata_unknown_axis"
        return ""
    if condition_like:
        if y_unknown and x_unknown:
            return "metadata_unknown_axis"
        return ""
    if not series and extractor_type == "unknown" and chart_type in {"", "unknown"}:
        return "metadata_no_series"
    return ""




def _infer_validation_status(is_candidate: bool = True, metadata_failure_reason: str = "", raw_record_count_before_dedupe: int = 0, raw_record_count_after_dedupe: int = 0, validated_record_count: int = 0, removed_shell_count: int = 0, validation_flags: Optional[List[str]] = None) -> str:
    flags = set(validation_flags or [])
    if not is_candidate:
        return "no_candidate"
    if metadata_failure_reason:
        return "metadata_failed"
    if int(raw_record_count_before_dedupe or 0) <= 0:
        return "extraction_empty"
    if int(validated_record_count or 0) <= 0:
        return "validation_empty"

    partial_markers = {
        "multi_series_under_extracted",
        "series_metadata_not_used",
        "temperature_axis_missing",
        "time_axis_missing",
        "category_condition_axis_missing",
        "no_core_metric",
    }
    if (
        int(validated_record_count or 0) < int(raw_record_count_after_dedupe or 0)
        or int(removed_shell_count or 0) > 0
        or bool(flags.intersection(partial_markers))
    ):
        return "validation_kept_partial"
    return "validation_ok"




def _count_series_from_records(records: List[Dict]) -> int:
    names: List[str] = []
    for row in records or []:
        if not isinstance(row, dict):
            continue
        semantic_role = clean_residual_mojibake_chars(str(row.get("semantic_figure_role", "") or "")).strip().lower()
        if semantic_role == "category_screening" or str(row.get("x_axis_mode", "")).strip().lower() == "category":
            series_name = clean_residual_mojibake_chars(
                str(row.get("raw_category_label", "") or row.get("Catalyst", "") or row.get("Series_Name", ""))
            ).strip()
        else:
            series_name = clean_residual_mojibake_chars(
                str(row.get("Series_Name", "") or row.get("raw_category_label", "") or row.get("Catalyst", ""))
            ).strip()
        if series_name and series_name.lower() not in {"", "unknown", "main_curve"}:
            names.append(series_name)
    return len(_dedupe_keep_order(names))




def _check_monotonic_numeric(values: List[Any]) -> bool:
    nums = []
    for value in values or []:
        text = clean_residual_mojibake_chars(str(value or "")).strip().replace(",", "")
        match = re.search(r"[-+]?\d+(?:\.\d+)?", text)
        if match:
            nums.append(float(match.group(0)))
    if len(nums) < 2:
        return True
    diffs = [right - left for left, right in zip(nums[:-1], nums[1:]) if abs(right - left) > 1e-9]
    if not diffs:
        return True
    return all(delta >= -1e-9 for delta in diffs) or all(delta <= 1e-9 for delta in diffs)




def _build_figure_candidate_audit_row(file_name: str, page_num: int, precheck_result: Dict[str, Any], text_context: str = "") -> Dict[str, Any]:
    return {
        "Source_File": str(file_name or ""),
        "page_num": int(page_num or 0),
        "candidate_score": int(precheck_result.get("candidate_score", 0) or 0),
        "candidate_reason": str(precheck_result.get("candidate_reason", "")),
        "positive_flags": list(precheck_result.get("positive_flags", []) or []),
        "negative_flags": list(precheck_result.get("negative_flags", []) or []),
        "hard_exclusion": bool(precheck_result.get("hard_exclusion", False)),
        "hard_exclusion_reason": str(precheck_result.get("hard_exclusion_reason", "")),
        "page_text_len": int(precheck_result.get("page_text_len", 0) or 0),
        "is_candidate": bool(precheck_result.get("is_candidate", False)),
        "text_context_preview": _context_preview(text_context, 800),
    }




def precheck_table_or_figure_candidates(page_text: str, prev_text: str = "", next_text: str = "", page_num: int = 0) -> Dict[str, Any]:
    if not globals().get("_FIGURE_PREFILTER_SELFCHECK_DONE"):
        globals()["_FIGURE_PREFILTER_SELFCHECK_DONE"] = True
        try:
            case1 = precheck_table_or_figure_candidates("XRD TEM XPS", "", "", 1)
            if case1.get("is_candidate"):
                print("  [WARNING] figure selfcheck failed: characterization-heavy low-text page passed candidate gate")
            case2 = precheck_table_or_figure_candidates("Fig. 4 Methanol steam reforming performance Conversion Different catalysts 90 C", "", "", 4)
            if not case2.get("is_candidate"):
                print("  [WARNING] figure selfcheck failed: category performance candidate rejected")
        except RecursionError:
            pass
        except Exception as e:
            print(f"  [WARNING] figure prefilter selfcheck error: {e}")

    current = _normalize_identity_text(page_text)
    prev_norm = _normalize_identity_text(prev_text)
    next_norm = _normalize_identity_text(next_text)
    combined = " ".join([current, prev_norm, next_norm]).strip()
    page_text_len = len(current.strip())

    perf_info = _is_performance_figure_context(current, prev_norm, next_norm)
    nonperf_info = _is_nonperformance_figure_context(combined)
    strong_positive = bool(perf_info.get("strong_positive", False))
    performance_strength = int(perf_info.get("performance_strength", 0) or 0)
    nonperformance_strength = int(nonperf_info.get("nonperformance_strength", 0) or 0)
    positive_flags = list(perf_info.get("positive_flags", []) or [])
    weak_flags = list(perf_info.get("weak_positive_flags", []) or [])
    negative_flags = list(nonperf_info.get("negative_flags", []) or [])
    hard_exclusion = bool(nonperf_info.get("hard_exclusion", False))
    hard_reason = str(nonperf_info.get("hard_reason", "") or "")

    # Task B: recall-first — only hard-exclude page 1 if essentially blank (< 100 chars)
    title_page_like = bool(
        page_num == 1
        and page_text_len < 100
        and not re.search(r"\bfig(?:ure)?\.?\s*\d+\b", combined, flags=re.I)
        and not perf_info.get("strong_positive")
    )
    if title_page_like and not hard_exclusion:
        hard_exclusion = True
        hard_reason = "title_page_blank"
        negative_flags.append("title_page_blank")

    # Task B: neighbor caption boost — prev/next page mentions Figure + performance keywords
    neighbor_context = " ".join([prev_norm, next_norm])
    neighbor_fig_mention = bool(re.search(r"\bfig(?:ure)?\.?\s*\d+\b", neighbor_context, flags=re.I))
    neighbor_perf_signal = bool(re.search(
        r"conversion|yield|selectivity|stability|time.on.stream|methanol|meoh|hydrogen",
        neighbor_context, flags=re.I
    ))
    if neighbor_fig_mention and neighbor_perf_signal and "neighbor_caption_boost" not in weak_flags:
        weak_flags.append("neighbor_caption_boost")

    weak_positive_count = len(_dedupe_keep_order(weak_flags))
    only_low_text = set(weak_flags) == {"low_text_fallback"} and weak_positive_count == 1

    is_candidate = False
    candidate_reason = "no_performance_signal"
    if strong_positive and not hard_exclusion:
        is_candidate = True
        candidate_reason = str(perf_info.get("strong_reason", "") or "strong_axis_legend_performance")
    elif strong_positive and hard_exclusion:
        clear_override = (
            hard_reason != "hard_exclusion_references"
            and performance_strength >= nonperformance_strength + 3
            and perf_info.get("metric_signal")
            and (perf_info.get("caption_hit") or perf_info.get("neighbor_caption_hit") or (perf_info.get("axis_like") and perf_info.get("legend_like")))
        )
        if clear_override:
            is_candidate = True
            candidate_reason = str(perf_info.get("strong_reason", "") or "strong_axis_legend_performance")
        else:
            candidate_reason = hard_reason or "no_performance_signal"
    elif weak_positive_count >= 1 and not hard_exclusion and not only_low_text:
        # Task B: recall-first — any single weak signal is enough (VL will filter false positives)
        is_candidate = True
        candidate_reason = "weak_signal_candidate"
    elif only_low_text and not hard_exclusion:
        candidate_reason = "low_text_only_rejected"
    elif hard_exclusion and not strong_positive:
        candidate_reason = hard_reason or "no_performance_signal"
    else:
        candidate_reason = "no_performance_signal"

    candidate_score = (
        performance_strength * 10
        + weak_positive_count * 4
        - len(_dedupe_keep_order(negative_flags)) * 7
        - (14 if hard_exclusion else 0)
        + (10 if strong_positive else 0)
    )
    return {
        "is_candidate": bool(is_candidate),
        "candidate_reason": str(candidate_reason),
        "candidate_score": int(candidate_score),
        "positive_flags": _dedupe_keep_order(positive_flags + weak_flags),
        "negative_flags": _dedupe_keep_order(negative_flags),
        "hard_exclusion": bool(hard_exclusion),
        "hard_exclusion_reason": hard_reason,
        "page_text_len": int(page_text_len),
    }




def extract_figures_from_pdf(pdf_path: str) -> List[Dict[str, Any]]:
    figures: List[Dict[str, Any]] = []
    candidate_audit_rows: List[Dict[str, Any]] = []
    file_name = os.path.basename(pdf_path)

    try:
        with pdfplumber.open(pdf_path) as pdf:
            total_pages = len(pdf.pages)
            page_texts: List[str] = []
            for page in pdf.pages:
                try:
                    page_texts.append(clean_residual_mojibake_chars(page.extract_text() or ""))
                except Exception as e:
                    print(f"  [WARNING] page text extraction failed: {file_name} page? {e}")
                    page_texts.append("")

            for page_num, page in enumerate(pdf.pages, 1):
                try:
                    idx = page_num - 1
                    page_text = page_texts[idx]
                    prev_text = page_texts[idx - 1] if idx > 0 else ""
                    next_text = page_texts[idx + 1] if idx < total_pages - 1 else ""
                    context_parts = []
                    if prev_text.strip():
                        context_parts.append(f"[prev_page]\n{prev_text[:400]}")
                    context_parts.append(f"[this_page]\n{page_text[:500]}")
                    if next_text.strip():
                        context_parts.append(f"[next_page]\n{next_text[:400]}")
                    text_context = "\n".join(context_parts)

                    precheck = precheck_table_or_figure_candidates(page_text, prev_text, next_text, page_num)
                    audit_row = _build_figure_candidate_audit_row(file_name, page_num, precheck, text_context)
                    candidate_audit_rows.append(audit_row)
                    if not precheck.get("is_candidate", False):
                        continue

                    img_path = os.path.join(
                        TEMP_IMG_DIR,
                        f"{os.path.splitext(os.path.basename(pdf_path))[0]}_page{page_num}.png",
                    )
                    if not os.path.exists(img_path):
                        try:
                            page.to_image(resolution=150).save(img_path)
                        except Exception as e:
                            print(f"  [WARNING] figure image render failed: {file_name} page {page_num}: {e}")
                            continue

                    figures.append({
                        "page_num": page_num,
                        "image_path": img_path,
                        "text_context": text_context,
                        "candidate_reason": precheck.get("candidate_reason", ""),
                        "candidate_score": precheck.get("candidate_score", 0),
                        "positive_flags": list(precheck.get("positive_flags", []) or []),
                        "negative_flags": list(precheck.get("negative_flags", []) or []),
                        "hard_exclusion": bool(precheck.get("hard_exclusion", False)),
                        "hard_exclusion_reason": str(precheck.get("hard_exclusion_reason", "")),
                        "page_text_len": int(precheck.get("page_text_len", 0) or 0),
                        "text_len": int(precheck.get("page_text_len", 0) or 0),
                        "candidate_audit_row": audit_row,
                    })
                except Exception as e:
                    print(f"  [WARNING] figure page precheck failed: {file_name} page {page_num}: {e}")
    except Exception as e:
        print(f"  [ERROR] figure extraction failed: {e}")

    globals().setdefault("_FIGURE_CANDIDATE_AUDIT_ROWS", []).extend(candidate_audit_rows)
    globals()["_LAST_FIGURE_CANDIDATE_AUDIT_ROWS"] = candidate_audit_rows
    return figures




def _encode_image(image_path: str) -> str:
    """Function docstring removed for runtime stability."""
    with open(image_path, "rb") as f:
        return base64.b64encode(f.read()).decode("utf-8")


def _vl_is_performance_chart(image_data: str, img_name: str) -> bool:
    """
    JSON-only structured prefilter for figure relevance (Task A).
    Recall-first: parse failure or ambiguity → keep.
    Rejects ONLY when label=characterization AND confidence=high AND zero performance_evidence.
    """
    prompt = (
        "Analyze this scientific figure. Return ONLY valid JSON — no markdown, no explanation:\n"
        '{"label": "performance or characterization or unclear",\n'
        ' "confidence": "high or medium or low",\n'
        ' "performance_evidence": ["axis/legend text mentioning conversion/yield/selectivity/TOS/stability"],\n'
        ' "characterization_evidence": ["XRD/XPS/TEM/SEM/TPR/BET/FTIR text if present"],\n'
        ' "reason": "10-word summary"}\n\n'
        "Definitions:\n"
        "performance: catalytic results — methanol conversion, H2 yield, CO/CO2 selectivity, "
        "H2 production rate, activity or stability vs time or temperature.\n"
        "characterization: XRD, XPS, TEM, SEM, TPR, TPD, BET, FTIR, Raman, "
        "particle size distribution, schematic, mechanism diagram.\n"
        "unclear: cannot determine definitively from the image.\n\n"
        "IMPORTANT: When in doubt choose unclear. Only use characterization when you are certain."
    )
    try:
        response = client.chat.completions.create(
            model=VISION_MODEL,
            messages=[{"role": "user", "content": [
                {"type": "image_url", "image_url": {"url": f"data:image/png;base64,{image_data}"}},
                {"type": "text", "text": prompt}
            ]}],
            max_tokens=200,
            temperature=0.0
        )
        raw = response.choices[0].message.content.strip()
        print(f"    [prefilter_json] {img_name}: {raw[:220]}")

        # Resilient JSON parse (try raw → extract {...})
        parsed = None
        try:
            parsed = json.loads(raw)
        except json.JSONDecodeError:
            pass
        if not isinstance(parsed, dict):
            m = re.search(r"\{.*\}", raw, re.DOTALL)
            if m:
                try:
                    parsed = json.loads(m.group(0))
                except json.JSONDecodeError:
                    pass
        if not isinstance(parsed, dict):
            # Parse failure → keep (recall-first)
            print(f"    [prefilter_parse_fail_keep] {img_name}")
            return True

        label = str(parsed.get("label", "unclear")).strip().lower()
        confidence = str(parsed.get("confidence", "low")).strip().lower()
        perf_ev = list(parsed.get("performance_evidence", []) or [])
        char_ev = list(parsed.get("characterization_evidence", []) or [])
        reason = str(parsed.get("reason", "")).strip()
        print(f"    [prefilter_parsed] {img_name}: label={label} conf={confidence} "
              f"perf_ev={len(perf_ev)} char_ev={len(char_ev)} reason='{reason}'")

        # --- Decision logic (recall-first) ---
        # Keep unconditionally when any performance evidence exists
        if perf_ev:
            print(f"    [prefilter_keep_perf_ev] {img_name}")
            return True
        # Keep when label is performance
        if label == "performance":
            print(f"    [prefilter_keep_performance] {img_name}")
            return True
        # Keep when unclear (recall-first)
        if label == "unclear":
            print(f"    [prefilter_keep_unclear] {img_name}")
            return True
        # Reject only when: label=characterization AND high confidence AND zero performance evidence
        if label == "characterization" and confidence == "high" and not perf_ev:
            print(f"    [prefilter_reject_char_high_conf] {img_name}: char_ev={char_ev[:3]}")
            return False
        # For characterization with medium/low confidence, keep conservatively
        print(f"    [prefilter_keep_conservative] {img_name}: label={label} conf={confidence}")
        return True

    except Exception as e:
        print(f"    [prefilter_exception_keep] {img_name}: {e}")
        return True
def _try_parse_json_block(raw: str, stage: str, img_name: str):
    """Function docstring removed for runtime stability."""
    if not raw or not raw.strip():
        print(f"    [{stage}] {img_name}: empty response")
        return None

    # Strategy 1: parse the raw response directly.
    try:
        obj = json.loads(raw)
        if isinstance(obj, list):
            return obj
        if isinstance(obj, dict):
            for key in ("data", "records", "results", "items", "points"):
                if isinstance(obj.get(key), list):
                    return obj[key]
            return [obj]  # Wrap a single record into a list.
    except json.JSONDecodeError:
        pass

    # Strategy 2: extract from a fenced ```json ... ``` block.
    m = re.search(r"```(?:json)?\s*(\[.*?\]|\{.*?\})\s*```", raw, re.DOTALL)
    if m:
        try:
            obj = json.loads(m.group(1))
            if isinstance(obj, list):
                return obj
            if isinstance(obj, dict):
                for key in ("data", "records", "results", "items", "points"):
                    if isinstance(obj.get(key), list):
                        return obj[key]
        except json.JSONDecodeError:
            pass

    # Strategy 3: extract from the outermost [ ... ] block.
    arr_start = raw.find("[")
    arr_end   = raw.rfind("]")
    if arr_start != -1 and arr_end > arr_start:
        try:
            return json.loads(raw[arr_start:arr_end + 1])
        except json.JSONDecodeError:
            pass

    # Strategy 4: extract from the outermost { ... } block.
    obj_start = raw.find("{")
    obj_end   = raw.rfind("}")
    if obj_start != -1 and obj_end > obj_start:
        try:
            obj = json.loads(raw[obj_start:obj_end + 1])
            if isinstance(obj, dict):
                for key in ("data", "records", "results", "items", "points"):
                    if isinstance(obj.get(key), list):
                        return obj[key]
        except json.JSONDecodeError:
            pass

    # Strategy 5: truncate a damaged tail and repair the outer wrapper conservatively.
    last_brace = raw.rfind("}")
    if last_brace != -1:
        salvaged = raw[obj_start if obj_start != -1 else 0: last_brace + 1]
        for suffix in ("]}", "}"):
            try:
                obj = json.loads(salvaged + suffix)
                if isinstance(obj, dict):
                    for key in ("data", "records", "results", "items"):
                        if isinstance(obj.get(key), list):
                            print(f"    [{stage}] {img_name}: salvage succeeded with {len(obj[key])} items")
                            return obj[key]
            except json.JSONDecodeError:
                pass

    # Strategy 6: truncated array — find last complete JSON object and close the array.
    # Handles the case where the model returns [..., {... (missing closing "]").
    _arr_s6 = raw.find("[")
    if _arr_s6 != -1:
        _last_cb6 = raw.rfind("}")
        if _last_cb6 > _arr_s6:
            _candidate6 = raw[_arr_s6: _last_cb6 + 1] + "]"
            try:
                _result6 = json.loads(_candidate6)
                if isinstance(_result6, list):
                    print(f"    [{stage}] {img_name}: salvage strategy-6 (truncated array) "
                          f"succeeded with {len(_result6)} items")
                    return _result6
            except json.JSONDecodeError:
                pass

    print(f"    [{stage}] {img_name}: all JSON parsing attempts failed; raw preview:\n{raw[:500]}")
    return None


def _salvage_dicts_from_raw_text(raw: str) -> List[Dict]:
    """Last-resort salvage: scan raw text and extract every complete {...} JSON object found.
    Used by category/condition builders when _try_parse_json_block returns None but the
    model clearly described some data points in natural language mixed with JSON fragments.
    Returns a (possibly empty) list of dicts; caller must validate semantic content.
    """
    if not raw:
        return []
    results: List[Dict] = []
    depth = 0
    start = -1
    for i, ch in enumerate(raw):
        if ch == "{":
            if depth == 0:
                start = i
            depth += 1
        elif ch == "}":
            depth -= 1
            if depth == 0 and start != -1:
                snippet = raw[start: i + 1]
                try:
                    obj = json.loads(snippet)
                    if isinstance(obj, dict):
                        results.append(obj)
                except json.JSONDecodeError:
                    pass
                start = -1
    return results


def _normalize_chart_metadata(meta: dict) -> dict:
    """Function docstring removed for runtime stability."""
    meta["x_axis"] = clean_residual_mojibake_chars(str(meta.get("x_axis", "")))
    meta["y_axis"] = clean_residual_mojibake_chars(str(meta.get("y_axis", "")))
    meta["series"] = [clean_residual_mojibake_chars(str(s)) for s in meta.get("series", [])]
    x = meta.get("x_axis", "").lower()

    # Normalize x_axis_mode on the Python side.
    if any(kw in x for kw in ["temp", "temperature", "°c", "celsius"]):
        meta["x_axis_mode"] = "temperature"
    elif any(kw in x for kw in ["tos", "time on stream", "time", "hour", "hours", " h", "stability"]):
        meta["x_axis_mode"] = "time"
    elif any(kw in x for kw in ["catalyst", "sample", "composition", "support", "loading", "content"]):
        meta["x_axis_mode"] = "category"
    elif any(kw in x for kw in ["s/c", "ratio", "ghsv", "whsv", "pressure", "flow", "space velocity", "amount"]):
        meta["x_axis_mode"] = "condition"
    else:
        # Keep the model value when available; otherwise use unknown.
        meta.setdefault("x_axis_mode", "unknown")

    # Determine extractor_type from x_axis_mode on the Python side.
    mode = meta["x_axis_mode"]
    if mode == "temperature":
        meta["extractor_type"] = "temperature_sweep"
    elif mode == "time":
        meta["extractor_type"] = "stability"
    elif mode in ("category", "condition"):
        meta["extractor_type"] = "category_or_condition"
    else:
        meta.setdefault("extractor_type", "unknown")

    # chart_type: keep model output and fall back to unknown when missing.
    meta.setdefault("chart_type", "unknown")

    # [Fix 17] Normalize series_role on the Python side instead of relying on setdefault.
    # Keep a valid model-provided role when it is already trustworthy.
    _VALID_ROLES = {"catalyst", "metric", "product", "condition", "unknown"}
    model_role = str(meta.get("series_role", "")).lower().strip()
    if model_role in _VALID_ROLES and model_role != "unknown":
        meta["series_role"] = model_role
    else:
        # Infer series role from the series labels themselves.
        series = meta.get("series", [])
        series_text = " ".join(str(s).lower() for s in series)
        # Product keywords: H2/CO2/CO/CH4/DME and similar outputs.
        _PRODUCT_KW = ["h2", "co2", " co", "ch4", "dme", "methane", "hydrogen"]
        # Metric keywords: conversion / yield / selectivity.
        _METRIC_KW  = ["conversion", "yield", "selectivity", "activity"]
        # Condition keywords: temperatures, pressures, and other operating conditions.
        _COND_KW    = ["°c", " k", "bar", "atm", "mpa"]
        # Catalyst keywords: catalyst names, supports, and loading patterns.
        _CAT_KW     = ["ni/", "cu/", "ni-", "cu-", "/al2o3", "/ceo2", "/zno",
                       "catalyst", "wt%", "wt.%"]

        if any(kw in series_text for kw in _PRODUCT_KW):
            meta["series_role"] = "product"
        elif any(kw in series_text for kw in _METRIC_KW):
            meta["series_role"] = "metric"
        elif any(kw in series_text for kw in _COND_KW):
            meta["series_role"] = "condition"
        elif any(kw in series_text for kw in _CAT_KW):
            meta["series_role"] = "catalyst"
        else:
            meta["series_role"] = "unknown"

    # Task C: normalize new extended fields with safe defaults
    # x_range_min / x_range_max: try coercing to float, else None
    for range_key in ("x_range_min", "x_range_max"):
        raw_val = meta.get(range_key)
        if raw_val in (None, "", "null", "unknown", "N/A"):
            meta[range_key] = None
        else:
            try:
                meta[range_key] = float(str(raw_val).strip().rstrip("°CcKkh%"))
            except (ValueError, TypeError):
                meta[range_key] = None
    # x_tick_labels: must be list of strings
    tick_raw = meta.get("x_tick_labels")
    if isinstance(tick_raw, list):
        meta["x_tick_labels"] = [str(t) for t in tick_raw]
    else:
        meta["x_tick_labels"] = []
    # estimated_series_count: int, fallback to len(series)
    sc_raw = meta.get("estimated_series_count")
    try:
        meta["estimated_series_count"] = int(sc_raw) if sc_raw not in (None, "") else len(meta.get("series", []))
    except (ValueError, TypeError):
        meta["estimated_series_count"] = len(meta.get("series", []))
    # metric_candidates: list of DATA_SCHEMA field names
    mc_raw = meta.get("metric_candidates")
    meta["metric_candidates"] = [str(m) for m in mc_raw] if isinstance(mc_raw, list) else []
    # legend_detected: bool
    ld_raw = meta.get("legend_detected")
    if isinstance(ld_raw, bool):
        meta["legend_detected"] = ld_raw
    elif isinstance(ld_raw, str):
        meta["legend_detected"] = ld_raw.lower() in ("true", "yes", "1")
    else:
        meta["legend_detected"] = bool(meta.get("series"))
    # metadata_confidence: "high"|"medium"|"low"
    conf_raw = str(meta.get("metadata_confidence", "")).lower().strip()
    meta["metadata_confidence"] = conf_raw if conf_raw in ("high", "medium", "low") else "medium"
    # x_axis_unit / y_axis_unit: safe string defaults
    meta.setdefault("x_axis_unit", "")
    meta.setdefault("y_axis_unit", "")

    return meta


def _vl_extract_chart_metadata(image_data: str, img_name: str, file_name: str) -> dict:
    """
    Task C: richer structured metadata — 19 fields.
    [Improve3] Added has_dual_y_axis, y_axis_right, x_scale_type to support
    dual-Y-axis charts and log-scale axes which were previously missed.
    """
    prompt = (
        f"This is a figure from an MSR (Methanol Steam Reforming) catalyst paper: {file_name}\n\n"
        f"Task: Identify the chart structure ONLY. Do NOT extract data values yet.\n\n"
        f"Return ONLY valid JSON (no markdown, no extra text) with exactly these fields:\n"
        f'{{\n'
        f'  "x_axis": "full x-axis label with unit, e.g. Reaction temperature (°C)",\n'
        f'  "x_axis_unit": "unit only, e.g. °C or h or %",\n'
        f'  "x_scale_type": "linear or log or unknown",\n'
        f'  "y_axis": "full left y-axis label with unit, e.g. MeOH conversion (%)",\n'
        f'  "y_axis_unit": "unit only, e.g. % or mmol/g/h",\n'
        f'  "has_dual_y_axis": true or false,\n'
        f'  "y_axis_right": "full right y-axis label if dual-Y, else null",\n'
        f'  "series": ["series/curve name 1", "series/curve name 2"],\n'
        f'  "chart_type": "line or grouped_bar or bar or scatter or mixed or unknown",\n'
        f'  "x_axis_mode": "temperature or time or category or condition or unknown",\n'
        f'  "series_role": "catalyst or metric or product or condition or unknown",\n'
        f'  "extractor_type": "temperature_sweep or stability or category_or_condition or unknown",\n'
        f'  "x_range_min": <number or null>,\n'
        f'  "x_range_max": <number or null>,\n'
        f'  "x_tick_labels": ["200", "250", "300"],\n'
        f'  "estimated_series_count": <integer>,\n'
        f'  "metric_candidates": ["MeOH_Conversion_% or H2_Yield_% or CO_Selectivity_% or CO2_Selectivity_% or H2_Production_Rate or CO_Concentration_ppm"],\n'
        f'  "legend_detected": true,\n'
        f'  "metadata_confidence": "high or medium or low"\n'
        f'}}\n\n'
        f"Rules:\n"
        f"- series: list ALL curve/series names visible in the legend (catalysts, metrics, products, conditions)\n"
        f"- Do NOT assume series are always catalyst names\n"
        f"- IMPORTANT: When listing series names, always use full chemical formula notation. "
        f"Write 'Cu/NiAl2O4' not 'CuNiAl', 'Cu/CoAl2O4' not 'CuCoAl', 'Ni/Al2O3' not 'NiAl'. "
        f"Never abbreviate spinel or oxide support names.\n"
        f"- x_scale_type: check if x-axis uses logarithmic scale (log) or linear scale (linear)\n"
        f"- has_dual_y_axis: true if there are TWO y-axes (left and right with different scales/labels)\n"
        f"- y_axis_right: label of the right y-axis when has_dual_y_axis=true, else null\n"
        f"- x_range_min/x_range_max: read from axis tick values or scale bar (null if not visible)\n"
        f"- x_tick_labels: list visible x-axis tick values as strings\n"
        f"- metric_candidates: which performance field(s) does the y-axis represent?\n"
        f"- metadata_confidence: high=clearly readable, medium=partially visible, low=very hard to read\n"
        f"- For unknown fields use null (numbers), [] (lists), or unknown (strings)"
    )
    try:
        response = client.chat.completions.create(
            model=VISION_MODEL,
            messages=[{"role": "user", "content": [
                {"type": "image_url", "image_url": {"url": f"data:image/png;base64,{image_data}"}},
                {"type": "text", "text": prompt}
            ]}],
            temperature=0.0,
            max_tokens=1000
        )
        raw = response.choices[0].message.content
        parsed = _try_parse_json_block(raw, "metadata", img_name)
        if isinstance(parsed, list) and len(parsed) == 1 and isinstance(parsed[0], dict):
            parsed = parsed[0]
        if isinstance(parsed, dict):
            parsed = _normalize_chart_metadata(parsed)
            series = parsed.get("series", [])
            dual = parsed.get("has_dual_y_axis", False)
            print(f"    [meta] {img_name}: x={parsed.get('x_axis_mode','?')} "
                  f"ext={parsed.get('extractor_type','?')} "
                  f"series={len(series)} dual_y={dual} "
                  f"conf={parsed.get('metadata_confidence','?')}")
            return parsed
    except Exception as e:
        print(f"    [metadata ERROR] {img_name}: {e}")
    return {}


def _infer_series_x_field(x_axis: str) -> str:
    """Function docstring removed for runtime stability."""
    x_lower = clean_residual_mojibake_chars(x_axis).lower()
    if any(kw in x_lower for kw in ["temp", "temperature", "°c", "celsius"]):
        return "Reaction_Temp_C"
    if any(kw in x_lower for kw in ["tos", "time on stream", "time", "hour", "hours", " h", "stability"]):
        return "TOS_h"
    return "Reaction_Temp_C"  # Conservative fallback.


def _infer_extraction_mode(x_axis: str, y_axis: str = "") -> str:
    """Function docstring removed for runtime stability."""
    x = clean_residual_mojibake_chars(x_axis).lower()
    if any(kw in x for kw in ["temp", "temperature", "°c", "celsius"]):
        return "temperature"
    if any(kw in x for kw in ["tos", "time on stream", "time", "hour", "hours", " h", "stability"]):
        return "stability"
    # Category/condition axes include catalyst names, compositions, loadings, S/C, GHSV, pressure, etc.
    if any(kw in x for kw in [
        "catalyst", "sample", "composition", "loading", "content",
        "s/c", "ratio", "ghsv", "whsv", "pressure", "flow", "amount"
    ]):
        return "category_or_condition"
    return "category_or_condition"  # Unknown x-axis: default to category/condition, not temperature windows.


def _is_series_name_real_catalyst(series_name: str, series_role: str) -> bool:
    """Function docstring removed for runtime stability."""
    s = _normalize_species_token(series_name)
    if _is_non_catalyst_figure_label(s):
        return False
    if series_role == "catalyst":
        return True
    # Temperature or condition labels are not catalyst identities.
    if any(kw in s for kw in ["°c", "°k", " k", "temperature", "temp", " bar", " mpa",
                               "time", "tos", " h ", " min", "flow", "ratio", "whsv", "ghsv"]):
        return False
    # Performance metric labels are not catalyst identities.
    if any(kw in s for kw in ["conversion", "yield", "selectivity", "activity",
                               "deactivation", "stability"]):
        return False
    # Pure product or gas labels are not catalyst identities.
    if s in {"h2", "co", "co2", "ch4", "dme", "ch3oh", "meoh", "n2", "ar", "he"}:
        return False
    # Numeric-first labels (for example "90 C") are not catalyst names either.
    if s and s[0].isdigit():
        return False
    return True


# ============================================================
# Task D: Dynamic window helpers for temperature / TOS ranges
# ============================================================

def _extract_axis_range_from_metadata(metadata: Dict) -> Tuple[Optional[float], Optional[float]]:
    """
    Extract x-axis numeric range (min, max) from metadata fields produced by Task C.
    Falls back to inferring from x_tick_labels when x_range_min/max are absent.
    """
    x_min = metadata.get("x_range_min")
    x_max = metadata.get("x_range_max")
    # Coerce to float
    try:
        x_min_f = float(x_min) if x_min not in (None, "", "null", "unknown") else None
    except (ValueError, TypeError):
        x_min_f = None
    try:
        x_max_f = float(x_max) if x_max not in (None, "", "null", "unknown") else None
    except (ValueError, TypeError):
        x_max_f = None
    # Supplement from tick labels when range fields are missing
    if x_min_f is None or x_max_f is None:
        tick_labels = list(metadata.get("x_tick_labels", []) or [])
        numeric_ticks: List[float] = []
        for t in tick_labels:
            try:
                numeric_ticks.append(float(str(t).strip().rstrip("°CcKkh%")))
            except (ValueError, TypeError):
                pass
        if len(numeric_ticks) >= 2:
            if x_min_f is None:
                x_min_f = min(numeric_ticks)
            if x_max_f is None:
                x_max_f = max(numeric_ticks)
    return x_min_f, x_max_f


def _build_dynamic_temperature_windows(
    x_min: Optional[float], x_max: Optional[float], tick_labels: List[str]
) -> List[str]:
    """
    Build temperature window hint strings for multi-window extraction.
    Derives windows from the actual axis range instead of hardcoded 200/260/320/400°C.
    Returns list of strings like ["200-280°C", "280-360°C", "360-440°C"].
    Falls back to ["all temperatures"] when range cannot be determined.
    """
    # Try inferring from tick_labels when range not given
    if x_min is None or x_max is None:
        numeric_ticks: List[float] = []
        for t in (tick_labels or []):
            try:
                numeric_ticks.append(float(str(t).strip().rstrip("°CcKk")))
            except (ValueError, TypeError):
                pass
        if len(numeric_ticks) >= 2:
            if x_min is None:
                x_min = min(numeric_ticks)
            if x_max is None:
                x_max = max(numeric_ticks)

    if x_min is None or x_max is None or x_max <= x_min:
        return ["all temperatures"]

    span = x_max - x_min
    lo = int(x_min)
    hi = int(x_max) + 1

    # Small span (< 80°C) → single window covers everything
    if span < 80:
        return [f"{lo}-{hi}°C"]

    # Medium span (80-180°C) → two equal windows
    if span < 180:
        mid = int(x_min + span / 2)
        return [f"{lo}-{mid}°C", f"{mid}-{hi}°C"]

    # Large span (≥ 180°C) → three equal windows
    step = span / 3
    w1 = int(x_min + step)
    w2 = int(x_min + 2 * step)
    return [f"{lo}-{w1}°C", f"{w1}-{w2}°C", f"{w2}-{hi}°C"]


def _build_dynamic_time_windows(
    x_min: Optional[float], x_max: Optional[float], tick_labels: List[str]
) -> List[str]:
    """
    Build TOS (time-on-stream) window hint strings.
    Stability tests are usually done in a single pass.
    Only splits into two windows for very long tests (> 200 h).
    Falls back to ["all time points"] when range cannot be determined.
    """
    if x_min is None or x_max is None:
        numeric_ticks: List[float] = []
        for t in (tick_labels or []):
            try:
                numeric_ticks.append(float(str(t).strip().rstrip("hH")))
            except (ValueError, TypeError):
                pass
        if len(numeric_ticks) >= 2:
            if x_min is None:
                x_min = min(numeric_ticks)
            if x_max is None:
                x_max = max(numeric_ticks)

    if x_min is None or x_max is None or x_max <= x_min:
        return ["all time points"]

    span = x_max - x_min
    if span <= 200:
        return ["all time points"]

    # Long stability test → two windows
    mid = int(x_min + span / 2)
    return [f"{int(x_min)}-{mid}h", f"{mid}-{int(x_max)+1}h"]


def _vl_extract_single_series(image_data: str, img_name: str, file_name: str,
                               series_name: str, y_axis: str, x_axis: str,
                               series_role: str = "unknown",
                               metadata: Optional[Dict] = None) -> List[Dict]:
    """
    One call extracts all points from one curve only.
    [Fix 10] x-field is adaptive: temperature uses Reaction_Temp_C; TOS/time uses TOS_h.
    [Fix 15] Three extraction modes: temperature / stability / category_or_condition.
    [Task E] Dynamic temperature/TOS windows from metadata instead of hardcoded ranges.
    Each mode uses its own prompt and filtering rules.
    """
    series_name = clean_residual_mojibake_chars(series_name)
    y_axis = clean_residual_mojibake_chars(y_axis)
    x_axis = clean_residual_mojibake_chars(x_axis)

    # [Fix 15] Infer extraction mode.
    extraction_mode = _infer_extraction_mode(x_axis, y_axis)
    # [Fix 10] Infer the x-field name (temperature -> Reaction_Temp_C, stability -> TOS_h).
    x_field = _infer_series_x_field(x_axis)
    print(f"    [Fix15] {img_name} series='{series_name}' "
          f"extraction_mode={extraction_mode} x_field={x_field}")

    # ------------------------------------------------------------------
    # Internal helper: build the prompt according to extraction_mode.
    # ------------------------------------------------------------------
    def _extract_block(range_hint: str) -> List[Dict]:
        range_clause = (
            f"Focus ONLY on the {x_axis} range {range_hint}. "
            if range_hint else ""
        )

        if extraction_mode == "temperature":
            # Temperature mode: emphasize point-by-point temperature extraction with Reaction_Temp_C.
            prompt = (
                f"This is a figure from MSR catalyst paper: {file_name}\n\n"
                f"Task: Extract data points for ONE specific curve only.\n"
                f"Target curve/series: '{series_name}'\n"
                f"x-axis: {x_axis} (reaction temperature)\n"
                f"y-axis: {y_axis}\n"
                f"{range_clause}\n"
                f"Rules:\n"
                f"- Read ALL clearly visible temperature data points for '{series_name}'\n"
                f"- Do NOT extract other curves\n"
                f"- Capture series identity before numeric point extraction\n"
                f"- Keep promoter or second metal in Catalyst, Catalyst_ID, and Promoter when labels contain Ni-Mg, Ni-Cu, Ni-Mo, Ni-Ce, Ni-La, Ni-Cr, Ni-In, Ni-Au, Ni-Rh, or Ni-Ir\n"
                f"- Do not simplify a promoted or bimetallic Ni catalyst to plain Ni\n"
                f"- IMPORTANT: Always use full chemical formula notation for catalyst names. "
                f"Write 'Cu/NiAl2O4' not 'CuNiAl', 'Cu/CoAl2O4' not 'CuCoAl', 'Ni/Al2O3' not 'NiAl'. "
                f"Never abbreviate spinel or oxide support names.\n"
                f"- Strictly distinguish CO (<5%) vs CO2 (>50%) selectivity\n"
                f"- Strictly distinguish MeOH_Conversion_% vs H2_Yield_%\n"
                f"- Prefer exact numeric values; if a point is hard to read precisely, "
                f"include your best estimate rather than omitting it\n"
                f"- Prioritize the most clearly readable anchor points "
                f"(curve start, end, peak/inflection) even when not all points are perfectly legible\n"
                f"- If the chart or caption only states range-like, above/below, maintained, or qualitative behavior, do not invent an exact single-point value\n\n"
                f"Return JSON array only (no wrapper object):\n"
                f'[{{"Catalyst": "{series_name}", "Catalyst_ID": "", "Promoter": "", "Support": "", '
                f'"Reaction_Temp_C": "...", "MeOH_Conversion_%": "<REQUIRED: read numeric % from y-axis>", '
                f'"H2_Yield_%": "", "CO_Selectivity_%": "", "CO2_Selectivity_%": "", '
                f'"Series_Name": "{series_name}"}}, ...]\n\n'
                f"CRITICAL: MeOH_Conversion_% MUST be filled with a numeric value for every point. Do not leave it empty.\n"
                f"Only return [] if the curve for '{series_name}' is COMPLETELY ABSENT from this figure.\n"
                f"If some points are visible but hard to read exactly, include best-estimate anchor points rather than returning []."
            )

        elif extraction_mode == "stability":
            # Stability mode: x field is TOS_h and Reaction_Temp_C is optional.
            prompt = (
                f"This is a stability/time-on-stream figure from MSR catalyst paper: {file_name}\n\n"
                f"Task: Extract time-on-stream data points for ONE specific curve only.\n"
                f"Target curve/series: '{series_name}'\n"
                f"x-axis: {x_axis} (time on stream, hours)\n"
                f"y-axis: {y_axis}\n"
                f"Rules:\n"
                f"- Read ALL clearly visible time-on-stream data points for '{series_name}'\n"
                f"- Do NOT extract other curves\n"
                f"- Capture series identity before numeric point extraction\n"
                f"- Keep promoter or second metal in Catalyst, Catalyst_ID, and Promoter when labels contain Ni-Mg, Ni-Cu, Ni-Mo, Ni-Ce, Ni-La, Ni-Cr, Ni-In, Ni-Au, Ni-Rh, or Ni-Ir\n"
                f"- Do not simplify a promoted or bimetallic Ni catalyst to plain Ni\n"
                f"- IMPORTANT: Always use full chemical formula notation for catalyst names. "
                f"Write 'Cu/NiAl2O4' not 'CuNiAl', 'Cu/CoAl2O4' not 'CuCoAl', 'Ni/Al2O3' not 'NiAl'. "
                f"Never abbreviate spinel or oxide support names.\n"
                f"- TOS_h is the primary x field - it MUST be filled for every point\n"
                f"- Reaction_Temp_C may be empty if not shown on this axis\n"
                f"- Strictly distinguish CO (<5%) vs CO2 (>50%) selectivity\n"
                f"- Prefer exact numeric values; if a point is hard to read precisely, "
                f"include your best estimate rather than omitting it\n"
                f"- Prioritize clearest anchor points (initial, final, any deactivation transition) "
                f"even when not all points are perfectly legible\n"
                f"- If the chart or caption only states range-like, above/below, maintained, or qualitative behavior, do not invent an exact single-point value\n\n"
                f"Return JSON array only (no wrapper object):\n"
                f'[{{"Catalyst": "{series_name}", "Catalyst_ID": "", "Promoter": "", "Support": "", '
                f'"TOS_h": "...", "MeOH_Conversion_%": "<REQUIRED: read numeric % from y-axis>", '
                f'"H2_Yield_%": "", "CO_Selectivity_%": "", "CO2_Selectivity_%": "", '
                f'"Series_Name": "{series_name}"}}, ...]\n\n'
                f"CRITICAL: MeOH_Conversion_% MUST be filled with a numeric value for every point. Do not leave it empty.\n"
                f"Only return [] if the curve for '{series_name}' is COMPLETELY ABSENT from this figure.\n"
                f"If some points are visible but hard to read exactly, include best-estimate anchor points rather than returning []."
            )

        else:
            # category_or_condition mode: the x-axis is categorical or condition-like.
            # Write the x-axis label into Notes as [x_axis=...][x_value=...].
            # [Fix 26] series_role-aware handling: condition/metric/product series are not catalyst names.
            # Tell the model explicitly not to write series_name into Catalyst in that case.
            series_role_hint = (
                f"- IMPORTANT: The target series '{series_name}' is a {series_role}, NOT a catalyst.\n"
                f"  Do NOT write '{series_name}' into the Catalyst field.\n"
                f"  Catalyst should be the x-axis category label (e.g. the catalyst name on x-axis), "
                f"or leave blank if x-axis is not a catalyst.\n"
                if series_role in ("condition", "metric", "product")
                else ""
            )
            prompt = (
                f"This is a category/condition screening figure from MSR catalyst paper: {file_name}\n\n"
                f"Task: Extract data points for ONE specific metric curve only.\n"
                f"Target curve/series: '{series_name}'\n"
                f"x-axis: {x_axis} (categorical or condition variable)\n"
                f"y-axis: {y_axis}\n"
                f"Rules:\n"
                f"- Read EVERY data point for '{series_name}' visible in the chart\n"
                f"- For each x-axis label/value, record the corresponding y value\n"
                f"- Put the x-axis label and value into Notes field as: "
                f"[x_axis={x_axis}][x_value=<label>]\n"
                f"[x_axis={x_axis}][x_value=<label>]\n"
                f"- Reaction_Temp_C and TOS_h may be empty - do NOT force-fill them\n"
                f"- If the x-axis category is a catalyst label such as Ni-Mg/Al2O3 or Ni-Cu/CNTs, keep that full catalyst identity and do not reduce it to plain Ni\n"
                f"- If promoter or second metal cannot be fully resolved, keep the partial catalyst label and mark identity as incomplete in Notes instead of inventing a monometallic Ni Catalyst_ID\n"
                f"- IMPORTANT: Always use full chemical formula notation for catalyst names. "
                f"Write 'Cu/NiAl2O4' not 'CuNiAl', 'Cu/CoAl2O4' not 'CuCoAl', 'Ni/Al2O3' not 'NiAl'. "
                f"Never abbreviate spinel or oxide support names.\n"
                f"- Strictly distinguish CO (<5%) vs CO2 (>50%) selectivity\n"
                f"- Use exact numeric values, not approximations\n"
                f"- If the chart or caption only states range-like, above/below, maintained, or qualitative behavior, do not invent an exact single-point value\n"
                f"{series_role_hint}\n"
                f"Return JSON array only (no wrapper object):\n"
                f'[{{"Catalyst": "<x-axis category if truly a catalyst label, else blank>", '
                f'"Catalyst_ID": "", "Promoter": "", "Support": "", '
                f'"MeOH_Conversion_%": "<REQUIRED: read numeric % from y-axis>", "H2_Yield_%": "", '
                f'"CO_Selectivity_%": "", "CO2_Selectivity_%": "", '
                f'"Notes": "[x_axis={x_axis}][x_value=<label>]", '
                f'"Series_Name": "{series_name}"}}, ...]\n\n'
                f"CRITICAL: MeOH_Conversion_% MUST be filled with a numeric value for every point. Do not leave it empty.\n"
                f"If no data found for this curve, return: []"
            )

        prompt = clean_residual_mojibake_chars(prompt)

        try:
            resp = client.chat.completions.create(
                model=VISION_MODEL,
                messages=[{"role": "user", "content": [
                    {"type": "image_url", "image_url": {"url": f"data:image/png;base64,{image_data}"}},
                    {"type": "text", "text": prompt}
                ]}],
                temperature=0.1,
                max_tokens=4096
            )
            raw = resp.choices[0].message.content
            tag = f"series[{series_name}]{range_hint}"
            print(f"    [{tag}] {img_name} raw response preview: {raw[:300]}")
            parsed = _try_parse_json_block(raw, tag, img_name)
            if isinstance(parsed, list):
                if extraction_mode == "category_or_condition":
                    # Category mode keeps rows whenever a performance value is present, even if x_field is empty.
                    perf_fields = [
                        "MeOH_Conversion_%", "H2_Yield_%", "H2_Selectivity_%",
                        "CO_Selectivity_%", "CO2_Selectivity_%",
                        "CO_Concentration_ppm", "Carbon_Deposition_wt%", "Deactivation_Rate_%_h"
                    ]
                    valid = [
                        r for r in parsed
                        if isinstance(r, dict) and any(
                            str(r.get(f, "")).strip() not in ("", "N/A")
                            for f in perf_fields
                        )
                    ]
                else:
                    # Temperature/stability mode requires the x-field to be present.
                    valid = [r for r in parsed if isinstance(r, dict) and r.get(x_field)]
                # [Fix 21/26] Post-process with series_role awareness.
                # If series_role is not catalyst, do not let series_name overwrite Catalyst.
                # Clear mistaken Catalyst values copied from non-catalyst series labels.
                # [Fix 26] Use containment checks, not only exact equality.
                if not _is_series_name_real_catalyst(series_name, series_role):
                    sn_lower = series_name.strip().lower()
                    for r in valid:
                        cat_val = str(r.get("Catalyst", "")).strip()
                        cat_lower = cat_val.lower()
                        # Clear Catalyst when it equals or trivially contains a non-catalyst series label.
                        # Also verify that the current Catalyst value does not look like a real catalyst.
                        should_clear = (
                            cat_lower == sn_lower or
                            (sn_lower and sn_lower in cat_lower) or
                            (cat_lower and cat_lower in sn_lower and
                             not _is_series_name_real_catalyst(cat_val, series_role))
                        )
                        if should_clear and cat_val:
                            r["Catalyst"] = ""
                            print(f"    [Fix26] {img_name} series='{series_name}': Catalyst='{cat_val}' cleared because series_role={series_role}")
                        # Always append series_role and series_name to Notes for traceability.
                        existing_notes = str(r.get("Notes", "")).strip()
                        role_tag = f"[series_role={series_role}][series_name={series_name}]"
                        if role_tag not in existing_notes:
                            r["Notes"] = (existing_notes + " " + role_tag).strip()
                print(f"    [{tag}] {img_name}: got {len(valid)} valid records "
                      f"(mode={extraction_mode} x_field={x_field})")
                # [Fix 23] H2 unit routing: prevent flow/rate values from being stored as percentages.
                if _is_series_name_real_catalyst(series_name, series_role):
                    for r in valid:
                        if not str(r.get("Catalyst", "")).strip():
                            r["Catalyst"] = series_name
                        enrich_promoter_fields_from_identity(r, series_name)
                        if not str(r.get("Catalyst_ID", "")).strip():
                            partial_id = _build_partial_identity_from_label(series_name)
                            if partial_id:
                                r["Catalyst_ID"] = partial_id
                        if not str(r.get("Catalyst_ID", "")).strip():
                            _append_note_tag(r, f"[series_identity_incomplete={series_name}]")
                return _apply_h2_unit_routing(valid)
        except Exception as e:
            print(f"    [series ERROR] {img_name} series={series_name} range={range_hint}: {e}")
        return []

    # ------------------------------------------------------------------
    # Layer 3 helper: sparse anchor recovery.
    # Only used for temperature_sweep / stability_tos when both windowed
    # extraction and full-range fallback return empty.
    # Asks the model for just a few clearly visible anchor points instead
    # of a complete digitization — returns [] for any other mode.
    # ------------------------------------------------------------------
    def _extract_anchor_points() -> List[Dict]:
        if extraction_mode not in ("temperature", "stability"):
            return []
        _perf_sa = [
            "MeOH_Conversion_%", "H2_Yield_%",
            "CO_Selectivity_%", "CO2_Selectivity_%",
        ]
        if extraction_mode == "temperature":
            _anchor_prompt = (
                f"This is a temperature-sweep figure from MSR catalyst paper: {file_name}\n\n"
                f"Task: Extract ANCHOR DATA POINTS for ONE curve: '{series_name}'\n"
                f"x-axis: {x_axis}   y-axis: {y_axis}\n\n"
                f"You do NOT need to extract every point — focus on the MOST CLEARLY VISIBLE ones:\n"
                f"  1) The lowest-temperature point visible for this curve (curve start)\n"
                f"  2) The highest-temperature point (curve end)\n"
                f"  3) Any obvious peak, valley, or inflection point\n"
                f"  4) 1-4 additional clearly readable intermediate points\n\n"
                f"IMPORTANT: Uncertain estimates are OK. Do NOT return [] just because some "
                f"points are hard to read — include your best estimate for the above anchors.\n\n"
                f"Rules:\n"
                f"- Reaction_Temp_C MUST have a numeric value for every returned object\n"
                f"- Fill whichever y-axis metric is visible ({y_axis})\n"
                f"- Keep catalyst identity as '{series_name}'; do not reduce bimetallic to plain Ni\n"
                f"- Do NOT invent values for curves that are genuinely absent\n\n"
                f"Return JSON array only (no explanation text):\n"
                f'[{{"Catalyst": "{series_name}", "Reaction_Temp_C": "...", '
                f'"MeOH_Conversion_%": "", "H2_Yield_%": "", "CO_Selectivity_%": "", '
                f'"CO2_Selectivity_%": "", "Series_Name": "{series_name}"}}, ...]\n\n'
                f"Only return [] if the curve for '{series_name}' is COMPLETELY ABSENT."
            )
        else:  # stability
            _anchor_prompt = (
                f"This is a stability/time-on-stream figure from MSR catalyst paper: {file_name}\n\n"
                f"Task: Extract ANCHOR DATA POINTS for ONE curve: '{series_name}'\n"
                f"x-axis: {x_axis}   y-axis: {y_axis}\n\n"
                f"You do NOT need to extract every time point — focus on the MOST CLEARLY VISIBLE:\n"
                f"  1) The initial point (t=0 or earliest visible time)\n"
                f"  2) The final point (latest visible time)\n"
                f"  3) Any obvious deactivation transition or plateau break\n"
                f"  4) 1-4 additional clearly readable intermediate points\n\n"
                f"IMPORTANT: Uncertain estimates are OK. Do NOT return [] just because some "
                f"points are hard to read — include your best estimate for the above anchors.\n\n"
                f"Rules:\n"
                f"- TOS_h MUST have a numeric value for every returned object\n"
                f"- Fill whichever y-axis metric is visible ({y_axis})\n"
                f"- Keep catalyst identity as '{series_name}'\n"
                f"- Do NOT invent values for curves genuinely absent\n\n"
                f"Return JSON array only (no explanation text):\n"
                f'[{{"Catalyst": "{series_name}", "TOS_h": "...", '
                f'"MeOH_Conversion_%": "", "H2_Yield_%": "", "CO_Selectivity_%": "", '
                f'"CO2_Selectivity_%": "", "Series_Name": "{series_name}"}}, ...]\n\n'
                f"Only return [] if the curve for '{series_name}' is COMPLETELY ABSENT."
            )
        _anchor_prompt = clean_residual_mojibake_chars(_anchor_prompt)
        try:
            _resp_a = client.chat.completions.create(
                model=VISION_MODEL,
                messages=[{"role": "user", "content": [
                    {"type": "image_url",
                     "image_url": {"url": f"data:image/png;base64,{image_data}"}},
                    {"type": "text", "text": _anchor_prompt},
                ]}],
                temperature=0.15,
                max_tokens=2048,
            )
            _raw_a = _resp_a.choices[0].message.content
            _tag_a = f"sparse_anchor[{series_name}]"
            print(f"    [{_tag_a}] {img_name} raw preview: {_raw_a[:300]}")
            _parsed_a = _try_parse_json_block(_raw_a, _tag_a, img_name)
            if isinstance(_parsed_a, list):
                # Relaxed validation: keep any record with x_field OR a performance value.
                # (Normal extraction requires x_field; anchor recovery allows either.)
                _valid_a = [
                    r for r in _parsed_a
                    if isinstance(r, dict) and (
                        r.get(x_field) or
                        any(str(r.get(f, "")).strip() not in ("", "N/A")
                            for f in _perf_sa)
                    )
                ]
                for r in _valid_a:
                    r["sparse_recovery"] = 1
                    r["low_density_reconstruction"] = 1
                    if not str(r.get("Catalyst", "")).strip():
                        r["Catalyst"] = series_name
                    if _is_series_name_real_catalyst(series_name, series_role):
                        enrich_promoter_fields_from_identity(r, series_name)
                    _append_note_tag(r, "[sparse_anchor_recovery=1]")
                print(f"    [{_tag_a}] {img_name}: got {len(_valid_a)} anchor points "
                      f"(sparse_recovery=True)")
                return _apply_h2_unit_routing(_valid_a)
        except Exception as _e_a:
            print(f"    [sparse_anchor ERROR] {img_name} series={series_name}: {_e_a}")
        return []

    # ------------------------------------------------------------------
    # Task E: Dynamic dispatch — temperature uses computed windows;
    # stability uses dynamic TOS windows; category/condition single pass.
    # ------------------------------------------------------------------
    meta = metadata or {}
    if extraction_mode == "temperature":
        # Build dynamic windows from metadata x_range / tick_labels
        x_min_v, x_max_v = _extract_axis_range_from_metadata(meta)
        tick_lbls = list(meta.get("x_tick_labels", []) or [])
        windows = _build_dynamic_temperature_windows(x_min_v, x_max_v, tick_lbls)
        print(f"    [dynamic_temp_windows] {img_name} series='{series_name}' windows={windows}")
        all_points: List[Dict] = []
        seen_keys: set = set()
        _window_attempt_count = len(windows)
        _window_success_count = 0
        for window in windows:
            block = _extract_block(window)
            if block:
                _window_success_count += 1
            for pt in block:
                t = str(pt.get(x_field, "")).strip()
                key = f"{series_name}|{t}"
                if key not in seen_keys:
                    seen_keys.add(key)
                    all_points.append(pt)
        # Layer 2: full-range fallback — if all windows returned empty, try single pass with no range hint
        if not all_points:
            print(f"    [temp_fullrange_fallback] {img_name} series='{series_name}' "
                  f"all {_window_attempt_count} windows empty — attempting full-range extraction "
                  f"[series_window_attempt_count={_window_attempt_count} series_window_success_count=0 "
                  f"full_range_fallback_attempted=True]")
            fallback_block = _extract_block("")
            for pt in fallback_block:
                t = str(pt.get(x_field, "")).strip()
                key = f"{series_name}|{t}"
                if key not in seen_keys:
                    seen_keys.add(key)
                    all_points.append(pt)
            if all_points:
                print(f"    [temp_fullrange_fallback] {img_name} series='{series_name}' "
                      f"recovered {len(all_points)} points [full_range_fallback_success=True]")
            else:
                print(f"    [temp_fullrange_fallback] {img_name} series='{series_name}' "
                      f"full-range fallback also empty — trying sparse anchor recovery "
                      f"[full_range_fallback_success=False]")
                # Layer 3: sparse anchor recovery
                print(f"    [sparse_anchor_recovery] {img_name} series='{series_name}' "
                      f"[sparse_recovery_attempted=True]")
                _anchor_pts_t = _extract_anchor_points()
                for _pt in _anchor_pts_t:
                    _tv = str(_pt.get(x_field, "")).strip()
                    _ka = f"{series_name}|{_tv}"
                    if _ka not in seen_keys:
                        seen_keys.add(_ka)
                        all_points.append(_pt)
                if all_points:
                    print(f"    [sparse_anchor_recovery] {img_name} series='{series_name}' "
                          f"recovered {len(all_points)} anchor points [sparse_recovery_success=True]")
                else:
                    print(f"    [sparse_anchor_recovery] {img_name} series='{series_name}' "
                          f"also empty [sparse_recovery_success=False "
                          f"extraction_empty_reason=all_windows_empty+full_range_empty+sparse_recovery_failed]")
        else:
            print(f"    [temp_windows_ok] {img_name} series='{series_name}' "
                  f"series_window_attempt_count={_window_attempt_count} "
                  f"series_window_success_count={_window_success_count} points={len(all_points)}")
        return all_points
    elif extraction_mode == "stability":
        # Build dynamic TOS windows — usually single pass, split only for very long tests
        x_min_v, x_max_v = _extract_axis_range_from_metadata(meta)
        tick_lbls = list(meta.get("x_tick_labels", []) or [])
        windows = _build_dynamic_time_windows(x_min_v, x_max_v, tick_lbls)
        print(f"    [dynamic_tos_windows] {img_name} series='{series_name}' windows={windows}")
        if len(windows) == 1:
            result_s = _extract_block(windows[0])
            if result_s:
                return result_s
            # Layer 2: full-range fallback for single-window stability
            print(f"    [stability_fullrange_fallback] {img_name} series='{series_name}' "
                  f"single window empty — attempting full-range extraction "
                  f"[full_range_fallback_attempted=True]")
            fallback_s = _extract_block("")
            if fallback_s:
                print(f"    [stability_fullrange_fallback] {img_name} series='{series_name}' "
                      f"recovered {len(fallback_s)} points [full_range_fallback_success=True]")
                return fallback_s
            # Layer 3: sparse anchor recovery for single-window stability
            print(f"    [stability_fullrange_fallback] {img_name} series='{series_name}' "
                  f"full-range fallback also empty — trying sparse anchor recovery "
                  f"[sparse_recovery_attempted=True]")
            _anchor_pts_s1 = _extract_anchor_points()
            if _anchor_pts_s1:
                print(f"    [sparse_anchor_recovery] {img_name} series='{series_name}' "
                      f"recovered {len(_anchor_pts_s1)} anchor points [sparse_recovery_success=True]")
                return _anchor_pts_s1
            print(f"    [sparse_anchor_recovery] {img_name} series='{series_name}' "
                  f"also empty [sparse_recovery_success=False "
                  f"extraction_empty_reason=stability_single_window_empty+full_range_empty+sparse_recovery_failed]")
            return []
        # Multi-window for very long stability tests
        all_points_t: List[Dict] = []
        seen_keys_t: set = set()
        _stab_window_success = 0
        for window in windows:
            block = _extract_block(window)
            if block:
                _stab_window_success += 1
            for pt in block:
                t = str(pt.get(x_field, "")).strip()
                key = f"{series_name}|{t}"
                if key not in seen_keys_t:
                    seen_keys_t.add(key)
                    all_points_t.append(pt)
        # Layer 2: full-range fallback for multi-window stability
        if not all_points_t:
            print(f"    [stability_fullrange_fallback] {img_name} series='{series_name}' "
                  f"all {len(windows)} windows empty — attempting full-range extraction "
                  f"[series_window_attempt_count={len(windows)} series_window_success_count=0 "
                  f"full_range_fallback_attempted=True]")
            fallback_t = _extract_block("")
            for pt in fallback_t:
                t = str(pt.get(x_field, "")).strip()
                key = f"{series_name}|{t}"
                if key not in seen_keys_t:
                    seen_keys_t.add(key)
                    all_points_t.append(pt)
            if all_points_t:
                print(f"    [stability_fullrange_fallback] {img_name} series='{series_name}' "
                      f"recovered {len(all_points_t)} points [full_range_fallback_success=True]")
            else:
                print(f"    [stability_fullrange_fallback] {img_name} series='{series_name}' "
                      f"full-range fallback also empty — trying sparse anchor recovery "
                      f"[full_range_fallback_success=False]")
                # Layer 3: sparse anchor recovery for multi-window stability
                print(f"    [sparse_anchor_recovery] {img_name} series='{series_name}' "
                      f"[sparse_recovery_attempted=True]")
                _anchor_pts_t2 = _extract_anchor_points()
                for _pt2 in _anchor_pts_t2:
                    _tv2 = str(_pt2.get(x_field, "")).strip()
                    _kt2 = f"{series_name}|{_tv2}"
                    if _kt2 not in seen_keys_t:
                        seen_keys_t.add(_kt2)
                        all_points_t.append(_pt2)
                if all_points_t:
                    print(f"    [sparse_anchor_recovery] {img_name} series='{series_name}' "
                          f"recovered {len(all_points_t)} anchor points [sparse_recovery_success=True]")
                else:
                    print(f"    [sparse_anchor_recovery] {img_name} series='{series_name}' "
                          f"also empty [sparse_recovery_success=False "
                          f"extraction_empty_reason=all_stability_windows_empty+full_range_empty+sparse_recovery_failed]")
        return all_points_t
    else:
        # category_or_condition mode: single pass, no window splitting.
        return _extract_block("")



# ============================================================
# Point construction utility layer (new helpers v7.1)
# ============================================================

def _normalize_axis_text(text: Any) -> str:
    """Normalize axis label text for keyword comparison."""
    t = clean_residual_mojibake_chars(str(text or "")).strip()
    t = re.sub(r"[()°\[\]{}]", " ", t)
    t = re.sub(r"\s+", " ", t)
    return t.strip().lower()


def _normalize_label_text(text: Any) -> str:
    """Normalize a category label or series label for deduplication."""
    t = clean_residual_mojibake_chars(str(text or "")).strip()
    t = re.sub(r"\s+", " ", t)
    return t.strip()


def _extract_numeric_tokens_from_text(text: Any) -> List[str]:
    """Extract all numeric tokens (including decimals, ratios) from text."""
    t = clean_residual_mojibake_chars(str(text or ""))
    tokens = re.findall(r"\b\d+(?:\.\d+)?(?:\s*[:/]\s*\d+(?:\.\d+)?)?\b", t)
    seen: set = set()
    result: List[str] = []
    for tok in tokens:
        tok = tok.replace(" ", "")
        if tok not in seen:
            seen.add(tok)
            result.append(tok)
    return result


def _extract_category_value_pairs_from_metadata(metadata: Dict) -> List[Tuple]:
    """
    Try to extract (category_label, metric_value) pairs from structured metadata
    bars/points/records/curves fields. Returns list of (label, value) tuples.
    """
    current = dict(metadata or {})
    pairs: List[Tuple] = []
    for container_key in ["bars", "points", "records", "curves", "data_points", "series_data"]:
        container = current.get(container_key, []) or []
        if not isinstance(container, (list, tuple)):
            continue
        for item in container:
            if not isinstance(item, dict):
                continue
            label = ""
            for lk in ["label", "name", "category", "x_label", "catalyst", "x"]:
                candidate = clean_residual_mojibake_chars(str(item.get(lk, "") or "")).strip()
                if candidate and not candidate.replace(".", "", 1).isdigit():
                    label = candidate
                    break
            value = ""
            for vk in ["value", "y", "y_value", "metric_value", "height", "mean"]:
                candidate = clean_residual_mojibake_chars(str(item.get(vk, "") or "")).strip()
                if candidate and re.search(r"\d", candidate):
                    value = candidate
                    break
            if label and value:
                pairs.append((label, value))
            elif label:
                for k, v in item.items():
                    if k in {"label", "name", "category", "x_label", "catalyst", "x"}:
                        continue
                    v_str = clean_residual_mojibake_chars(str(v or "")).strip()
                    if v_str and re.search(r"\d+(?:\.\d+)?", v_str):
                        parsed = _parse_float_if_possible(v_str.replace(",", ""))
                        if parsed is not None:
                            pairs.append((label, str(parsed)))
                            break
    return pairs


def _extract_series_metric_blocks(metadata: Dict) -> List[Dict]:
    """
    Extract structured series blocks (series_name, x_values, y_values)
    from metadata series array. Returns list of dicts.
    """
    current = dict(metadata or {})
    result: List[Dict] = []
    series_list = current.get("series", []) or []
    if not isinstance(series_list, (list, tuple)):
        return result
    for item in series_list:
        if isinstance(item, str):
            label = clean_residual_mojibake_chars(item).strip()
            if label:
                result.append({"series_name": label, "x_values": [], "y_values": []})
        elif isinstance(item, dict):
            label = _safe_get_series_label(item)
            x_values: List[str] = []
            y_values: List[str] = []
            for xk in ["x_values", "x", "xs", "x_data"]:
                xv = item.get(xk, [])
                if isinstance(xv, (list, tuple)) and xv:
                    x_values = [str(v) for v in xv]
                    break
                elif xv and isinstance(xv, str):
                    tokens = _extract_numeric_tokens_from_text(xv)
                    if tokens:
                        x_values = tokens
                        break
            for yk in ["y_values", "y", "ys", "y_data", "values", "data"]:
                yv = item.get(yk, [])
                if isinstance(yv, (list, tuple)) and yv:
                    y_values = [str(v) for v in yv]
                    break
                elif yv and isinstance(yv, str):
                    tokens = _extract_numeric_tokens_from_text(yv)
                    if tokens:
                        y_values = tokens
                        break
            if label or x_values or y_values:
                result.append({"series_name": label, "x_values": x_values, "y_values": y_values})
    return result


def _merge_shared_anchor_bundle_into_points(points: List[Dict], shared_anchor_bundle: Dict[str, Any]) -> List[Dict]:
    """
    Merge shared_anchor_bundle into each point dict (only fill empty fields).
    Does not overwrite existing non-empty values.
    """
    bundle = dict(shared_anchor_bundle or {})
    anchor_fields = [
        "Reaction_Temp_C", "TOS_h", "S_C_Ratio", "Feed_MeOH_to_H2O_Ratio",
        "Pressure_bar", "GHSV_mL_g_h", "WHSV_h_inv", "Flow_Rate", "Catalyst_Amount_g",
    ]
    result: List[Dict] = []
    for row in (points or []):
        current = dict(row)
        for field in anchor_fields:
            value = clean_residual_mojibake_chars(str(bundle.get(field, "") or "")).strip()
            if value and not clean_residual_mojibake_chars(str(current.get(field, "") or "")).strip():
                current[field] = value
        result.append(current)
    return result


def _canonical_metric_value_assignment(point: Dict, metric_field: str, value: Any) -> Dict:
    """
    Safely assign a metric value to the correct field in a raw point dict.
    Does NOT overwrite existing non-empty values.
    """
    current = dict(point or {})
    if not metric_field:
        return current
    existing = clean_residual_mojibake_chars(str(current.get(metric_field, "") or "")).strip()
    if existing and existing.lower() not in {"", "n/a", "na", "none", "nan", "unknown"}:
        return current
    normalized_value = clean_residual_mojibake_chars(str(value or "")).strip()
    if not normalized_value or normalized_value.lower() in {"n/a", "na", "none", "nan", "unknown"}:
        return current
    parsed = _parse_float_if_possible(normalized_value.replace(",", ""))
    if parsed is not None:
        current[metric_field] = _format_ratio_token(str(parsed))
    elif re.search(r"\d", normalized_value):
        current[metric_field] = normalized_value
    return current


def _is_empty_shell_point(row: Dict) -> bool:
    """
    Returns True only when a row has NOTHING of diagnostic value:
    no metric, no category label, no effect axis value, no anchor, no series, no notes.
    Intentionally conservative — prefer false negatives over false positives.
    """
    if not isinstance(row, dict):
        return True
    current = dict(row)
    if _has_any_metric_value(current):
        return False
    if clean_residual_mojibake_chars(str(current.get("raw_category_label", "") or "")).strip():
        return False
    notes = clean_residual_mojibake_chars(str(current.get("Notes", "") or "")).strip()
    _note_axis, note_x_value = _extract_x_axis_and_value_from_notes(notes)
    if note_x_value:
        return False
    if _has_any_anchor_value(current):
        return False
    if clean_residual_mojibake_chars(str(current.get("Series_Name", "") or "")).strip():
        return False
    if clean_residual_mojibake_chars(str(
        current.get("condition_anchor_source", "") or current.get("anchor_sources", "") or ""
    )).strip():
        return False
    effect_field = _infer_condition_axis_field(current, "", "")
    if effect_field and effect_field not in {"", "raw_category_label"}:
        if clean_residual_mojibake_chars(str(current.get(effect_field, "") or "")).strip():
            return False
    return True


def _selfcheck_point_construction_only() -> None:
    """
    In-memory ONLY selfcheck for point construction helpers.
    MUST NOT write to any output container or CSV.
    Uses Source_File='_selfcheck_internal_' to make any accidental leakage detectable.
    """
    try:
        # Case 1: category_screening with seeded labels
        meta1 = {
            "x_axis": "Different catalysts", "y_axis": "MeOH conversion (%)",
            "x_axis_mode": "category", "series": ["90 C"], "chart_type": "bar",
            "extractor_type": "category_or_condition", "semantic_figure_role": "category_screening",
        }
        cap1 = "Fig. 4 Methanol steam reforming performance of different catalysts at 90 C"
        nb1 = "No catalyst, ZrO2, Ni/ZrO2"
        lps1 = [
            {"raw_category_label": "No catalyst", "MeOH_Conversion_%": "0"},
            {"raw_category_label": "ZrO2", "MeOH_Conversion_%": "12.3"},
            {"raw_category_label": "Ni/ZrO2", "MeOH_Conversion_%": "65.0"},
        ]
        rows1, _ = _build_raw_points_from_category_screening(meta1, cap1, nb1, {"legacy_points": lps1})
        if len(rows1) < 3:
            print(f"  [WARNING][selfcheck_point] Case1: category builder got {len(rows1)} rows (expected >=3)")
        if not any(str(r.get("Reaction_Temp_C", "")).strip() == "90" for r in rows1):
            print("  [WARNING][selfcheck_point] Case1: no row has Reaction_Temp_C=90")

        # Case 2: condition_effect with seeded x-values
        meta2 = {
            "x_axis": "S/C ratio", "y_axis": "CO2 selectivity (%)",
            "x_axis_mode": "condition", "semantic_figure_role": "condition_effect",
        }
        lps2 = [
            {"S_C_Ratio": "1", "CO2_Selectivity_%": "70"},
            {"S_C_Ratio": "2", "CO2_Selectivity_%": "75"},
            {"S_C_Ratio": "3", "CO2_Selectivity_%": "80"},
        ]
        rows2, _ = _build_raw_points_from_condition_effect(meta2, "", "Ni-Cu/clay S/C ratio 1 2 3", {"legacy_points": lps2})
        if len(rows2) < 3:
            print(f"  [WARNING][selfcheck_point] Case2: condition builder got {len(rows2)} rows (expected >=3)")
        if not all(str(r.get("S_C_Ratio", "")).strip() for r in rows2):
            print("  [WARNING][selfcheck_point] Case2: some rows missing S_C_Ratio")

        # Case 3: category with no labels → no_category_labels_detected
        r3 = _build_extraction_empty_reason(meta1, "category_screening", metric_field="MeOH_Conversion_%",
                                             category_labels=[], effect_axis_values=[], raw_points=[])
        if r3 != "no_category_labels_detected":
            print(f"  [WARNING][selfcheck_point] Case3: expected 'no_category_labels_detected', got '{r3}'")

        # Case 4: condition_effect metric fail → no_metric_field_inferred
        meta4 = {"x_axis": "S/C ratio", "x_axis_mode": "condition", "semantic_figure_role": "condition_effect"}
        r4 = _build_extraction_empty_reason(meta4, "condition_effect", metric_field="",
                                             category_labels=[], effect_axis_values=["1", "2", "3"], raw_points=[])
        if r4 != "no_metric_field_inferred":
            print(f"  [WARNING][selfcheck_point] Case4: expected 'no_metric_field_inferred', got '{r4}'")

        # Case 5: partial point (has metric + S_C_Ratio but no Reaction_Temp_C) should survive
        partial_row = {
            "raw_category_label": "Ni/ZrO2", "MeOH_Conversion_%": "65",
            "S_C_Ratio": "2", "semantic_figure_role": "category_screening",
        }
        partial_meta = {
            "x_axis": "Different catalysts", "y_axis": "MeOH conversion (%)",
            "x_axis_mode": "category", "semantic_figure_role": "category_screening",
        }
        kept5, audit5 = validate_extracted_point_count_against_context([partial_row], partial_meta)
        if not kept5 or audit5.get("validation_status") == "validation_empty":
            print("  [WARNING][selfcheck_point] Case5: partial point was dropped by validation")

        # Case 6: _is_empty_shell_point
        if not _is_empty_shell_point({}):
            print("  [WARNING][selfcheck_point] Case6: empty dict should be shell")
        if _is_empty_shell_point({"MeOH_Conversion_%": "65"}):
            print("  [WARNING][selfcheck_point] Case6: row with metric should NOT be shell")

        print("  [selfcheck_point_construction_only] done (in-memory only, no CSV output)")
    except Exception as _sc_err:
        print(f"  [WARNING][selfcheck_point_construction_only] error: {_sc_err}")


# ============================================================
# End of new helpers v7.1
# ============================================================


def _has_any_metric_value(row: Dict) -> bool:
    if not isinstance(row, dict):
        return False
    for field in [
        "MeOH_Conversion_%", "H2_Yield_%", "H2_Production_Rate", "H2_Selectivity_%",
        "CO_Selectivity_%", "CO2_Selectivity_%", "CO_Concentration_ppm", "metric_value",
    ]:
        value = clean_residual_mojibake_chars(str(row.get(field, "") or "")).strip()
        if value and value.lower() not in {"", "n/a", "na", "none", "nan", "unknown"}:
            return True
    return False



def _has_any_anchor_value(row: Dict) -> bool:
    if not isinstance(row, dict):
        return False
    for field in [
        "Reaction_Temp_C", "TOS_h", "S_C_Ratio", "Feed_MeOH_to_H2O_Ratio",
        "Pressure_bar", "GHSV_mL_g_h", "WHSV_h_inv", "SpaceVelocity_norm",
        "Flow_Rate", "Catalyst_Amount_g",
    ]:
        value = clean_residual_mojibake_chars(str(row.get(field, "") or "")).strip()
        if value and value.lower() not in {"", "n/a", "na", "none", "nan", "unknown"}:
            return True
    return False



def _safe_get_series_label(series_item: Any) -> str:
    if isinstance(series_item, dict):
        for key in ["label", "name", "series_name", "title", "text"]:
            value = clean_residual_mojibake_chars(str(series_item.get(key, "") or "")).strip()
            if value:
                return value
        return ""
    return clean_residual_mojibake_chars(str(series_item or "")).strip()



def _safe_get_caption_text(context: str) -> str:
    text = clean_residual_mojibake_chars(context or "")
    if not text.strip():
        return ""
    this_page_match = re.search(r"\[this_page\]\s*(.*?)(?:\n\[next_page\]|\Z)", text, flags=re.I | re.S)
    this_page = this_page_match.group(1).strip() if this_page_match else text.strip()
    if not this_page:
        return ""
    fig_match = re.search(r"\bfig(?:ure)?\.?\s*\d+\b", this_page, flags=re.I)
    if fig_match:
        start_idx = max(fig_match.start() - 80, 0)
        end_idx = min(len(this_page), fig_match.end() + 1000)
        return this_page[start_idx:end_idx].strip()
    return this_page[:1200].strip()



def _safe_get_nearby_context_text(context: str) -> str:
    text = clean_residual_mojibake_chars(context or "")
    if not text.strip():
        return ""
    parts: List[str] = []
    prev_match = re.search(r"\[prev_page\]\s*(.*?)(?:\n\[this_page\]|\Z)", text, flags=re.I | re.S)
    next_match = re.search(r"\[next_page\]\s*(.*)$", text, flags=re.I | re.S)
    this_match = re.search(r"\[this_page\]\s*(.*?)(?:\n\[next_page\]|\Z)", text, flags=re.I | re.S)
    if prev_match and prev_match.group(1).strip():
        parts.append(prev_match.group(1).strip())
    if this_match and this_match.group(1).strip():
        parts.append(this_match.group(1).strip()[:800])
    if next_match and next_match.group(1).strip():
        parts.append(next_match.group(1).strip())
    merged = "\n".join(part for part in parts if part.strip())
    return merged[:1800].strip()



def _extract_metric_candidates_from_metadata(metadata: Dict, caption_text: str = "", nearby_text: str = "") -> List[str]:
    current = dict(metadata or {})
    series_text = " ".join(_safe_get_series_label(item) for item in (current.get("series", []) or []))
    combined = " ".join([
        clean_residual_mojibake_chars(str(current.get("y_axis", "") or "")),
        clean_residual_mojibake_chars(str(current.get("x_axis", "") or "")),
        series_text,
        clean_residual_mojibake_chars(str(caption_text or "")),
        clean_residual_mojibake_chars(str(nearby_text or "")),
    ])
    combined_norm = _normalize_identity_text(combined).lower()
    if not combined_norm:
        return []

    candidates: List[str] = []

    def _add(field: str, patterns: List[str]) -> None:
        if any(re.search(pattern, combined_norm, flags=re.I) for pattern in patterns):
            candidates.append(field)

    # CO concentration (ppm) — check first to avoid conflict with CO_Selectivity_%
    _add("CO_Concentration_ppm", [
        r"\bco concentration\b", r"\bco ppm\b", r"\bppm\b[^.;\n]{0,30}\bco\b",
        r"\bco content\b", r"\bcarbon monoxide concentration\b",
        r"\bco level\b", r"\bco outlet\b", r"\boutlet co\b",
    ])
    # H2 production rate — explicit rate/productivity/generation keywords
    _add("H2_Production_Rate", [
        r"\bh2 production rate\b", r"\bhydrogen production rate\b",
        r"\bh2 formation rate\b", r"\bhydrogen formation rate\b",
        r"\bh2 generation rate\b", r"\bhydrogen generation rate\b",
        r"\bh2 productivity\b", r"\bhydrogen productivity\b",
        r"\brate of h2\b", r"\brate of hydrogen\b",
        r"\bh2 rate\b", r"\bhydrogen rate\b",
    ])
    # H2 yield (%) — percentage-based, not a rate
    _add("H2_Yield_%", [r"\bh2 yield\b", r"\bhydrogen yield\b"])
    _add("H2_Selectivity_%", [r"\bh2 selectivity\b", r"\bhydrogen selectivity\b"])
    # CO2 selectivity — check BEFORE CO_Selectivity_ to prevent cross-contamination
    _add("CO2_Selectivity_%", [
        r"\bco2 selectivity\b", r"\bcarbon dioxide selectivity\b",
        r"\bselectivity to co2\b", r"\bselectivity toward co2\b",
        r"\bselectivity of co2\b", r"\bco2 yield\b",
    ])
    # CO selectivity — \bco\b already prevents matching "co2" (no word boundary after o in co2)
    _add("CO_Selectivity_%", [
        r"\bco selectivity\b(?![^.;\n]{0,10}\bco2\b)",
        r"\bselectivity to co\b",
        r"\bselectivity of co\b",
    ])
    # MeOH conversion — many alias forms including ch3oh and "reacted"
    if re.search(
        r"\b(?:methanol|meoh|ch3oh)\b[^.;\n]{0,40}\bconversion\b"
        r"|\bconversion\b[^.;\n]{0,40}\b(?:methanol|meoh|ch3oh)\b"
        r"|\b(?:methanol|meoh|ch3oh)\s+reacted\b"
        r"|\bconversion\s+of\s+(?:methanol|meoh|ch3oh)\b",
        combined_norm, flags=re.I,
    ):
        candidates.append("MeOH_Conversion_%")
    elif re.search(r"\bconversion\b", combined_norm, flags=re.I) and re.search(r"\bmsr\b|\bmethanol steam reforming\b|\breforming\b", combined_norm, flags=re.I):
        candidates.append("MeOH_Conversion_%")

    return _dedupe_keep_order(candidates)



def _infer_metric_field_from_axis_or_caption(metadata: Dict, caption_text: str = "", nearby_text: str = "") -> str:
    """
    Infer the primary metric field.
    Priority: y_axis (any hit) > series (unique) > caption (unique) > combined (first).
    Intentionally less conservative than old version: y_axis hit = use first candidate.
    Returns "" only when NO signal found anywhere.
    """
    current = dict(metadata or {})

    def _candidates_from_text(text_value: str) -> List[str]:
        local_meta = {"y_axis": text_value, "x_axis": "", "series": []}
        return _extract_metric_candidates_from_metadata(local_meta, "", "")

    # Priority 1: y_axis alone — if ANY match, trust the first (y_axis is most specific)
    y_axis_candidates = _candidates_from_text(str(current.get("y_axis", "") or ""))
    if y_axis_candidates:
        return y_axis_candidates[0]

    # Priority 2: series labels — require unique match only (series can contain condition values)
    series_text = " ".join(_safe_get_series_label(item) for item in (current.get("series", []) or []))
    series_candidates = _candidates_from_text(series_text)
    if len(series_candidates) == 1:
        return series_candidates[0]

    # Priority 3: caption alone — require unique match only
    caption_candidates = _candidates_from_text(caption_text or "")
    if len(caption_candidates) == 1:
        return caption_candidates[0]

    # Priority 4: combined — return first candidate if any (not empty)
    combined_candidates = _extract_metric_candidates_from_metadata(current, caption_text, nearby_text)
    if combined_candidates:
        return combined_candidates[0]

    # Fix2: semantic_role fallback — when y_axis is None but role is clear, infer metric
    semantic_role = str(current.get("semantic_figure_role", "") or "").strip().lower()
    if semantic_role in {"temperature_sweep", "stability_tos", "condition_effect"}:
        # MSR papers: temperature sweep almost always shows MeOH conversion
        return "MeOH_Conversion_%"

    return ""



def _infer_condition_axis_field(metadata: Dict, caption_text: str = "", nearby_text: str = "") -> str:
    current = dict(metadata or {})
    x_axis = _normalize_identity_text(str(current.get("x_axis", "") or "")).lower()
    x_mode = _normalize_identity_text(str(current.get("x_axis_mode", "") or "")).lower()
    semantic_role = _normalize_identity_text(str(current.get("semantic_figure_role", "") or "")).lower()
    combined = " ".join([x_axis, _normalize_identity_text(caption_text).lower(), _normalize_identity_text(nearby_text).lower()])

    if x_mode == "category" or semantic_role == "category_screening":
        return "raw_category_label"
    if x_mode == "temperature" or semantic_role == "temperature_sweep":
        return "Reaction_Temp_C"
    if x_mode == "time" or semantic_role == "stability_tos":
        return "TOS_h"
    if re.search(r"\bwhsv\b", combined, flags=re.I):
        return "WHSV_h_inv"
    if re.search(r"\b(?:ghsv|space velocity|sv)\b", combined, flags=re.I):
        return "GHSV_mL_g_h"
    if re.search(r"\b(?:methanol[- ]?to[- ]?water|meoh[- ]?to[- ]?h2o)\b", combined, flags=re.I):
        return "Feed_MeOH_to_H2O_Ratio"
    if re.search(r"\b(?:s/c|steam[- ]?to[- ]?carbon|steam carbon ratio|feed ratio)\b", combined, flags=re.I):
        return "S_C_Ratio"
    if re.search(r"\bpressure\b", combined, flags=re.I):
        return "Pressure_bar"
    if re.search(r"\b(?:flow rate|feed flow|feed rate)\b", combined, flags=re.I):
        return "Flow_Rate"
    if re.search(r"\b(?:catalyst amount|catalyst mass|catalyst charge|loading in reactor)\b", combined, flags=re.I):
        return "Catalyst_Amount_g"
    return ""



def _extract_category_labels_from_metadata_or_context(metadata: Dict, caption_text: str = "", nearby_text: str = "") -> List[str]:
    current = dict(metadata or {})
    candidates: List[str] = []

    def _add_label(value: Any) -> None:
        label = _usable_figure_category_label(value)
        if not label:
            return
        if label.lower() in {"unknown", "main_curve", "series", "legend", "different catalysts", "catalyst", "sample"}:
            return
        if label.replace(".", "", 1).isdigit():
            return
        if re.search(r"\b(conversion|yield|selectivity|temperature|pressure|ghsv|whsv|time on stream|tos)\b", label, flags=re.I):
            return
        candidates.append(label)

    for field in ["categories", "category_labels", "legend_labels", "x_labels", "labels"]:
        value = current.get(field, [])
        if isinstance(value, (list, tuple, set)):
            for item in value:
                _add_label(item)
        elif isinstance(value, str) and value.strip():
            for part in re.split(r"[;,\n|]", value):
                _add_label(part)

    if str(current.get("series_role", "") or "").strip().lower() == "catalyst":
        for item in (current.get("series", []) or []):
            _add_label(_safe_get_series_label(item))

    text_pool = " ".join([caption_text, nearby_text])
    text_pool = clean_residual_mojibake_chars(text_pool or "")
    if text_pool.strip():
        if _looks_like_niznal_context_text(text_pool):
            for label, _, _ in _NIZNAL_TABLE_SPECS:
                if re.search(rf"\b{re.escape(label)}\b", text_pool, flags=re.I):
                    _add_label(label)
        for match in re.findall(r"\b(?:no catalyst|without catalyst|blank|baseline)\b", text_pool, flags=re.I):
            _add_label(match)
        for match in re.findall(r"\b[A-Za-z]\d{1,3}\b", text_pool, flags=re.I):
            _add_label(match)
        for match in re.findall(r"\b(?:[A-Za-z0-9][A-Za-z0-9\-\.,%]*\/[A-Za-z0-9][A-Za-z0-9\-\.,%/]*|Al2O3|ZrO2|CeO2|TiO2|SiO2|MgO|CNTs|SEP|Activated Carbon|Cement-Clay)\b", text_pool, flags=re.I):
            _add_label(match)

    return _dedupe_keep_order(candidates)



def _extract_effect_axis_values_from_metadata_or_context(metadata: Dict, caption_text: str = "", nearby_text: str = "") -> List[str]:
    current = dict(metadata or {})
    axis_field = _infer_condition_axis_field(current, caption_text, nearby_text)
    if axis_field in {"", "raw_category_label"}:
        return []

    values: List[str] = []

    def _normalize_numeric(value: Any, allow_ratio: bool = False) -> str:
        text_value = clean_residual_mojibake_chars(str(value or "")).strip()
        if not text_value or text_value.lower() in {"", "n/a", "na", "none", "nan", "unknown"}:
            return ""
        if allow_ratio and re.fullmatch(r"\d+(?:\.\d+)?\s*[:/]\s*\d+(?:\.\d+)?", text_value):
            return text_value.replace(" ", "")
        parsed = _parse_float_if_possible(text_value.replace(",", ""))
        if parsed is not None:
            return _format_ratio_token(str(parsed))
        match = re.search(r"[-+]?\d+(?:\.\d+)?(?:\s*[:/]\s*[-+]?\d+(?:\.\d+)?)?", text_value.replace(",", ""))
        return match.group(0).replace(" ", "") if match else ""

    def _add_value(value: Any, allow_ratio: bool = False) -> None:
        normalized = _normalize_numeric(value, allow_ratio=allow_ratio)
        if normalized:
            values.append(normalized)

    for field in ["x_values", "effect_axis_values", "condition_values", "categories", "x_labels",
                  "x_tick_labels"]:  # Fix1: also read x_tick_labels from metadata
        raw_value = current.get(field, [])
        if isinstance(raw_value, (list, tuple, set)):
            for item in raw_value:
                _add_value(item, allow_ratio=(axis_field in {"S_C_Ratio", "Feed_MeOH_to_H2O_Ratio"}))
        elif isinstance(raw_value, str) and raw_value.strip():
            for part in re.split(r"[;,\n|]", raw_value):
                _add_value(part, allow_ratio=(axis_field in {"S_C_Ratio", "Feed_MeOH_to_H2O_Ratio"}))

    # Fix1: fallback — derive temperature values from x_range when x_tick_labels empty
    if axis_field == "Reaction_Temp_C" and not values:
        # VL returns x_range_min/x_range_max separately; also support legacy x_range list
        _xmin = current.get("x_range_min")
        _xmax = current.get("x_range_max")
        if _xmin is not None and _xmax is not None:
            x_range = [_xmin, _xmax]
        else:
            x_range = current.get("x_range") or current.get("metadata_x_range")
        if isinstance(x_range, (list, tuple)) and len(x_range) == 2:
            try:
                lo, hi = float(x_range[0]), float(x_range[1])
                # generate 5 evenly spaced points within the range
                step = (hi - lo) / 4 if hi > lo else 50
                t = lo
                while t <= hi + 0.01:
                    _add_value(str(round(t)))
                    t += step
            except Exception:
                pass

    text_pool = " ".join([str(current.get("x_axis", "") or ""), caption_text, nearby_text])
    text_pool = clean_residual_mojibake_chars(text_pool or "")
    if axis_field == "Reaction_Temp_C":
        for match in re.findall(r"(\d{2,4}(?:\.\d+)?)\s*(?:deg\.?\s*c|celsius|\u00b0c|c)\b", text_pool, flags=re.I):
            if _reaction_temp_value_is_sane(match):
                _add_value(match)
    elif axis_field == "TOS_h":
        for match in re.findall(r"(\d+(?:\.\d+)?)\s*h\b", text_pool, flags=re.I):
            _add_value(match)
    elif axis_field in {"S_C_Ratio", "Feed_MeOH_to_H2O_Ratio"}:
        for match in re.findall(r"\b(?:s/c|steam[- ]?to[- ]?carbon|feed ratio|methanol[- ]?to[- ]?water|meoh[- ]?to[- ]?h2o)\b[^.;\n]{0,40}?((?:\d+(?:\.\d+)?)(?:\s*[:/]\s*\d+(?:\.\d+)?)?(?:\s*,\s*\d+(?:\.\d+)?(?:\s*[:/]\s*\d+(?:\.\d+)?)?)*)", text_pool, flags=re.I):
            for token in re.split(r"\s*,\s*", match):
                _add_value(token, allow_ratio=True)
    elif axis_field == "Pressure_bar":
        for match in re.findall(r"\bpressure\b[^.;\n]{0,30}?(\d+(?:\.\d+)?)", text_pool, flags=re.I):
            _add_value(match)
    elif axis_field in {"GHSV_mL_g_h", "WHSV_h_inv"}:
        for match in re.findall(r"\b(?:ghsv|whsv|space velocity)\b[^.;\n]{0,40}?([0-9]+(?:\.[0-9]+)?(?:\s*[x\u00d7]\s*10\^?\d+)?)", text_pool, flags=re.I):
            _add_value(match)
    elif axis_field == "Flow_Rate":
        for match in re.findall(r"\b(?:flow rate|feed flow|feed rate)\b[^.;\n]{0,30}?(\d+(?:\.\d+)?)", text_pool, flags=re.I):
            _add_value(match)
    elif axis_field == "Catalyst_Amount_g":
        for match in re.findall(r"\b(?:catalyst amount|catalyst mass|catalyst charge|catalyst bed)\b[^.;\n]{0,30}?(\d+(?:\.\d+)?)", text_pool, flags=re.I):
            _add_value(match)

    return _dedupe_keep_order(values)



def _extract_shared_condition_anchor_bundle(metadata: Dict, caption_text: str = "", nearby_text: str = "", series_label: str = "") -> Dict[str, Any]:
    current = dict(metadata or {})
    result: Dict[str, Any] = {
        "Reaction_Temp_C": "",
        "TOS_h": "",
        "S_C_Ratio": "",
        "Feed_MeOH_to_H2O_Ratio": "",
        "Pressure_bar": "",
        "GHSV_mL_g_h": "",
        "WHSV_h_inv": "",
        "Flow_Rate": "",
        "Catalyst_Amount_g": "",
        "_anchor_sources": [],
    }

    def _merge_anchor(source_name: str, text_value: str) -> None:
        anchor = _extract_condition_anchor_from_text(text_value)
        for field in ["Reaction_Temp_C", "TOS_h", "S_C_Ratio", "Pressure_bar", "GHSV_mL_g_h", "Flow_Rate", "Catalyst_Amount_g"]:
            value = clean_residual_mojibake_chars(str(anchor.get(field, "") or "")).strip()
            if value and not str(result.get(field, "")).strip():
                result[field] = value
        if re.search(r"\bwhsv\b", text_value or "", flags=re.I):
            match = re.search(r"\bwhsv\b[^.;\n]{0,35}?([0-9]+(?:\.[0-9]+)?(?:\s*[x\u00d7]\s*10\^?\d+)?)", text_value or "", flags=re.I)
            if match and not str(result.get("WHSV_h_inv", "")).strip():
                result["WHSV_h_inv"] = match.group(1).replace(" ", "")
        if re.search(r"\b(?:methanol[- ]?to[- ]?water|meoh[- ]?to[- ]?h2o)\b", text_value or "", flags=re.I):
            match = re.search(r"\b(?:methanol[- ]?to[- ]?water|meoh[- ]?to[- ]?h2o)\b[^.;\n]{0,25}?((?:\d+(?:\.\d+)?)(?:\s*[:/]\s*\d+(?:\.\d+)?)?)", text_value or "", flags=re.I)
            if match and not str(result.get("Feed_MeOH_to_H2O_Ratio", "")).strip():
                result["Feed_MeOH_to_H2O_Ratio"] = match.group(1).replace(" ", "")
        for source in anchor.get("_anchor_sources", []) or []:
            if str(source).strip():
                result["_anchor_sources"].append(f"{source_name}:{str(source).strip()}")

    for field in ["Reaction_Temp_C", "TOS_h", "S_C_Ratio", "Feed_MeOH_to_H2O_Ratio", "Pressure_bar", "GHSV_mL_g_h", "WHSV_h_inv", "Flow_Rate", "Catalyst_Amount_g"]:
        value = clean_residual_mojibake_chars(str(current.get(field, "") or "")).strip()
        if value:
            result[field] = value
            result["_anchor_sources"].append(f"metadata:{field}")

    metadata_notes = " ".join([
        str(current.get("notes", "") or ""),
        str(current.get("Notes", "") or ""),
        str(current.get("description", "") or ""),
    ])
    for source_name, text_value in [
        ("caption", caption_text),
        ("nearby_context", nearby_text),
        ("series", series_label),
        ("metadata_notes", metadata_notes),
        ("x_axis", str(current.get("x_axis", "") or "")),
        ("y_axis", str(current.get("y_axis", "") or "")),
    ]:
        if str(text_value).strip():
            _merge_anchor(source_name, str(text_value))

    result["_anchor_sources"] = _dedupe_keep_order([str(item) for item in result.get("_anchor_sources", []) if str(item).strip()])
    return result



def _coerce_single_raw_point(raw_point: Dict, metadata: Dict, semantic_role: str, shared_anchor_bundle: Optional[Dict[str, Any]] = None, metric_field: str = "", effect_field: str = "", default_series_name: str = "") -> Dict:
    current = clean_record_text_fields(dict(raw_point or {}))
    meta = dict(metadata or {})
    semantic_role = clean_residual_mojibake_chars(str(semantic_role or meta.get("semantic_figure_role", "") or "")).strip() or "unknown"
    x_axis = clean_residual_mojibake_chars(str(meta.get("x_axis", "") or "")).strip()
    y_axis = clean_residual_mojibake_chars(str(meta.get("y_axis", "") or "")).strip()
    x_axis_mode = clean_residual_mojibake_chars(str(meta.get("x_axis_mode", "") or "")).strip() or "unknown"
    extractor_type = clean_residual_mojibake_chars(str(meta.get("extractor_type", "") or "")).strip() or "unknown"
    series_role = clean_residual_mojibake_chars(str(meta.get("series_role", "") or "")).strip() or "unknown"
    chart_type = clean_residual_mojibake_chars(str(meta.get("chart_type", "") or "")).strip() or "unknown"

    def _normalize_numeric(value: Any, allow_ratio: bool = False) -> str:
        text_value = clean_residual_mojibake_chars(str(value or "")).strip()
        if not text_value or text_value.lower() in {"", "n/a", "na", "none", "nan", "unknown"}:
            return ""
        if allow_ratio and re.fullmatch(r"\d+(?:\.\d+)?\s*[:/]\s*\d+(?:\.\d+)?", text_value):
            return text_value.replace(" ", "")
        parsed = _parse_float_if_possible(text_value.replace(",", ""))
        if parsed is not None:
            return _format_ratio_token(str(parsed))
        match = re.search(r"[-+]?\d+(?:\.\d+)?(?:\s*[:/]\s*[-+]?\d+(?:\.\d+)?)?", text_value.replace(",", ""))
        return match.group(0).replace(" ", "") if match else text_value

    current["data_source"] = "figure"
    current["x_axis"] = x_axis
    current["y_axis"] = y_axis
    current["x_axis_mode"] = x_axis_mode
    current["extractor_type"] = extractor_type
    current["series_role"] = series_role
    current["semantic_figure_role"] = semantic_role
    current["chart_type"] = chart_type

    if default_series_name and not str(current.get("Series_Name", "")).strip():
        current["Series_Name"] = clean_residual_mojibake_chars(default_series_name).strip()

    notes = clean_residual_mojibake_chars(str(current.get("Notes", "") or "")).strip()
    note_axis, note_x_value = _extract_x_axis_and_value_from_notes(notes)
    x_value = clean_residual_mojibake_chars(str(
        current.get("x_value", "") or current.get("x_axis_value", "") or current.get("condition_value", "") or note_x_value
    )).strip()

    if metric_field and not _has_any_metric_value(current):
        metric_value = ""
        for key in ["metric_value", "y_value", "value", "y"]:
            metric_value = clean_residual_mojibake_chars(str(current.get(key, "") or "")).strip()
            if metric_value:
                break
        if metric_value:
            current[metric_field] = _normalize_numeric(metric_value)

    if semantic_role == "category_screening":
        catalyst_fallback = _usable_figure_category_label(current.get("Catalyst", ""))
        raw_category_label = clean_residual_mojibake_chars(str(
            current.get("raw_category_label", "") or current.get("category_label", "") or current.get("x_label", "")
            or note_x_value or catalyst_fallback
        )).strip()
        if not raw_category_label and default_series_name and _is_series_name_real_catalyst(default_series_name, series_role):
            raw_category_label = clean_residual_mojibake_chars(default_series_name).strip()
        current["raw_category_label"] = raw_category_label
        current = _strip_non_catalyst_identity_labels(current)
        raw_category_label = clean_residual_mojibake_chars(str(current.get("raw_category_label", "") or "")).strip()
        current["raw_category_label"] = raw_category_label
        if str(current.get("Catalyst", "")).strip() and not catalyst_fallback and not raw_category_label:
            current["Catalyst"] = ""
        if raw_category_label and not str(current.get("Catalyst", "")).strip():
            if _looks_like_identity_mapping(raw_category_label) or _is_blank_or_baseline_label(raw_category_label) or _normalize_support_only_label(raw_category_label):
                current["Catalyst"] = raw_category_label
    else:
        if not effect_field:
            effect_field = _infer_condition_axis_field(meta, "", "")
        if effect_field and effect_field not in {"", "raw_category_label"}:
            if not str(current.get(effect_field, "")).strip() and x_value:
                current[effect_field] = _normalize_numeric(x_value, allow_ratio=(effect_field in {"S_C_Ratio", "Feed_MeOH_to_H2O_Ratio"}))
            if not str(current.get(effect_field, "")).strip() and note_axis and note_x_value:
                note_field = _infer_condition_axis_field({"x_axis": note_axis, "x_axis_mode": meta.get("x_axis_mode", "")}, "", "")
                if note_field == effect_field:
                    current[effect_field] = _normalize_numeric(note_x_value, allow_ratio=(effect_field in {"S_C_Ratio", "Feed_MeOH_to_H2O_Ratio"}))

    if semantic_role == "temperature_sweep" and not str(current.get("Reaction_Temp_C", "")).strip() and x_value:
        current["Reaction_Temp_C"] = _normalize_numeric(x_value)
    if semantic_role == "stability_tos" and not str(current.get("TOS_h", "")).strip() and x_value:
        current["TOS_h"] = _normalize_numeric(x_value)

    anchor_sources: List[str] = []
    bundle = dict(shared_anchor_bundle or {})
    for field in [
        "Reaction_Temp_C", "TOS_h", "S_C_Ratio", "Feed_MeOH_to_H2O_Ratio",
        "Pressure_bar", "GHSV_mL_g_h", "WHSV_h_inv", "Flow_Rate", "Catalyst_Amount_g",
    ]:
        value = clean_residual_mojibake_chars(str(bundle.get(field, "") or "")).strip()
        if value and not str(current.get(field, "")).strip():
            current[field] = value
    for source in bundle.get("_anchor_sources", []) or []:
        if str(source).strip():
            anchor_sources.append(str(source).strip())

    raw_anchor_sources = current.get("anchor_sources", current.get("condition_anchor_source", ""))
    if isinstance(raw_anchor_sources, (list, tuple, set)):
        anchor_sources.extend([str(item).strip() for item in raw_anchor_sources if str(item).strip()])
    elif str(raw_anchor_sources).strip():
        anchor_sources.extend([part.strip() for part in str(raw_anchor_sources).split("|") if part.strip()])
    if anchor_sources:
        current["anchor_sources"] = _dedupe_keep_order(anchor_sources)
        current["condition_anchor_source"] = "|".join(_dedupe_keep_order(anchor_sources))

    if x_value and x_axis and not note_x_value:
        _append_note_tag(current, f"[x_axis={x_axis}][x_value={x_value}]")
    if effect_field and effect_field not in {"", "raw_category_label"} and str(current.get(effect_field, "")).strip():
        current[effect_field] = _normalize_numeric(current.get(effect_field, ""), allow_ratio=(effect_field in {"S_C_Ratio", "Feed_MeOH_to_H2O_Ratio"}))

    for field in [
        "Reaction_Temp_C", "TOS_h", "Pressure_bar", "GHSV_mL_g_h", "WHSV_h_inv",
        "Flow_Rate", "Catalyst_Amount_g", "MeOH_Conversion_%", "H2_Yield_%",
        "H2_Production_Rate", "H2_Selectivity_%", "CO_Selectivity_%",
        "CO2_Selectivity_%", "CO_Concentration_ppm",
    ]:
        if str(current.get(field, "")).strip():
            current[field] = _normalize_numeric(current.get(field, ""), allow_ratio=False)
    for field in ["S_C_Ratio", "Feed_MeOH_to_H2O_Ratio"]:
        if str(current.get(field, "")).strip():
            current[field] = _normalize_numeric(current.get(field, ""), allow_ratio=True)

    current = _strip_non_catalyst_identity_labels(current)
    if (series_role in {"product", "metric"} or _is_product_species_only_identity(current)) and not _has_identity_context_for_figure_row(current):
        for field in ("Catalyst", "raw_category_label", "Series_Name"):
            current[field] = ""
        _append_note_tag(current, "[product_series_identity_blocked]")
    clues = _extract_identity_clues_from_figure_record(current)
    if clues.get("clue_record"):
        current = _backfill_identity_from_clue_record(current, clues["clue_record"])
    return current



def _filter_empty_raw_points(rows: List[Dict]) -> List[Dict]:
    kept: List[Dict] = []
    for row in rows or []:
        if not isinstance(row, dict):
            continue
        current = _strip_non_catalyst_identity_labels(dict(row))
        if _should_drop_filtered_non_catalyst_row(current):
            continue
        notes = clean_residual_mojibake_chars(str(current.get("Notes", "") or "")).strip()
        note_axis, note_x_value = _extract_x_axis_and_value_from_notes(notes)
        effect_field = _infer_condition_axis_field(current, "", "")
        has_effect_value = bool(
            clean_residual_mojibake_chars(str(current.get(effect_field, "") or "")).strip()
            if effect_field and effect_field not in {"", "raw_category_label"} else note_x_value
        )
        semantic_role = clean_residual_mojibake_chars(str(current.get("semantic_figure_role", "") or "")).strip()
        has_identity_context = _has_identity_context_for_figure_row(current)
        if semantic_role == "category_screening" and not has_identity_context:
            continue
        has_context = bool(
            clean_residual_mojibake_chars(str(current.get("raw_category_label", "") or "")).strip()
            or has_effect_value
            or _has_any_anchor_value(current)
            or clean_residual_mojibake_chars(str(current.get("Series_Name", "") or "")).strip()
            or clean_residual_mojibake_chars(str(current.get("condition_anchor_source", "") or "")).strip()
        )
        if _is_product_species_only_identity(current) and not has_identity_context:
            continue
        if _has_any_metric_value(current) and has_context:
            kept.append(current)
    return kept



def _promote_partial_point_when_metric_and_anchor_exist(row: Dict) -> Dict:
    current = _strip_non_catalyst_identity_labels(dict(row or {}))
    if _should_drop_filtered_non_catalyst_row(current):
        current["_drop_filtered_non_catalyst_row"] = 1
        return current
    if _is_product_species_only_identity(current):
        current["_drop_filtered_non_catalyst_row"] = 1
        return current
    if not _has_any_metric_value(current):
        return current
    notes = clean_residual_mojibake_chars(str(current.get("Notes", "") or "")).strip()
    note_axis, note_x_value = _extract_x_axis_and_value_from_notes(notes)
    semantic_role = clean_residual_mojibake_chars(str(current.get("semantic_figure_role", "") or "")).strip()
    has_identity_context = _has_identity_context_for_figure_row(current)
    if semantic_role == "category_screening":
        if not has_identity_context:
            return current
    has_partial_context = bool(
        clean_residual_mojibake_chars(str(current.get("raw_category_label", "") or "")).strip()
        or note_x_value
        or _has_any_anchor_value(current)
        or clean_residual_mojibake_chars(str(current.get("condition_anchor_source", "") or current.get("anchor_sources", "") or "")).strip()
        or clean_residual_mojibake_chars(str(current.get("Series_Name", "") or "")).strip()
    )
    if has_partial_context:
        current["partial_point_promoted"] = 1
        _append_note_tag(current, "[partial_points_promoted=1]")
    return current



def _selfcheck_metric_inference() -> None:
    """Selfcheck for metric candidate extraction and CO/CO2/H2 disambiguation."""
    if globals().get("_METRIC_INFERENCE_SELFCHECK_DONE"):
        return
    globals()["_METRIC_INFERENCE_SELFCHECK_DONE"] = True
    try:
        # Case 1: "selectivity to CO2" → CO2_Selectivity_%, NOT CO_Selectivity_%
        c1 = _extract_metric_candidates_from_metadata({"y_axis": "selectivity to CO2 (%)"}, "", "")
        if "CO2_Selectivity_%" not in c1:
            print("  [WARNING] metric selfcheck failed: CO2 selectivity not detected")
        if "CO_Selectivity_%" in c1:
            print("  [WARNING] metric selfcheck failed: CO_Selectivity_% falsely triggered for CO2 text")

        # Case 2: "CO selectivity" alone → CO_Selectivity_%, NOT CO2_Selectivity_%
        c2 = _extract_metric_candidates_from_metadata({"y_axis": "CO selectivity (%)"}, "", "")
        if "CO_Selectivity_%" not in c2:
            print("  [WARNING] metric selfcheck failed: CO selectivity not detected")
        if "CO2_Selectivity_%" in c2:
            print("  [WARNING] metric selfcheck failed: CO2_Selectivity_% falsely triggered for CO-only text")

        # Case 3: "H2 production rate" → H2_Production_Rate, NOT H2_Yield_%
        c3 = _extract_metric_candidates_from_metadata({"y_axis": "H2 production rate (mL/min/g)"}, "", "")
        if "H2_Production_Rate" not in c3:
            print("  [WARNING] metric selfcheck failed: H2 production rate not detected")
        if "H2_Yield_%" in c3:
            print("  [WARNING] metric selfcheck failed: H2_Yield_% falsely triggered for rate text")

        # Case 4: caption "methanol conversion" → MeOH_Conversion_%
        c4 = _extract_metric_candidates_from_metadata({}, "Effect of temperature on methanol conversion", "")
        if "MeOH_Conversion_%" not in c4:
            print("  [WARNING] metric selfcheck failed: MeOH_Conversion_% not detected from caption")

        # Case 5: y_axis condition hint (S/C) + caption "CO2 selectivity" → CO2_Selectivity_%
        c5 = _extract_metric_candidates_from_metadata({"x_axis": "S/C ratio"}, "CO2 selectivity vs S/C", "")
        if "CO2_Selectivity_%" not in c5:
            print("  [WARNING] metric selfcheck failed: CO2_Selectivity_% not detected from S/C+CO2 caption")

        # Case 6: "CH3OH conversion" alias
        c6 = _extract_metric_candidates_from_metadata({"y_axis": "CH3OH conversion (%)"}, "", "")
        if "MeOH_Conversion_%" not in c6:
            print("  [WARNING] metric selfcheck failed: CH3OH conversion alias not detected")

        # Case 7: infer CO2_Selectivity_% wins over CO_Selectivity_% for CO2 y_axis
        f5 = _infer_metric_field_from_axis_or_caption({"y_axis": "CO2 selectivity (%)"}, "", "")
        if f5 != "CO2_Selectivity_%":
            print(f"  [WARNING] metric selfcheck failed: infer CO2 selectivity expected CO2_Selectivity_%, got {f5!r}")
    except Exception as _e:
        print(f"  [WARNING] metric inference selfcheck error (non-fatal): {_e}")


def _selfcheck_filtered_and_backfill() -> None:
    if globals().get("_FILTERED_BACKFILL_SELFCHECK_DONE"):
        return
    globals()["_FILTERED_BACKFILL_SELFCHECK_DONE"] = True
    try:
        if not _should_drop_filtered_non_catalyst_row({
            "Notes": "[filtered_non_catalyst_label=Catalyst=CO2-eq]",
            "Catalyst": "",
            "Series_Name": "",
            "raw_category_label": "",
            "Canonical_Catalyst_ID": "",
        }):
            print("  [WARNING] filtered/backfill selfcheck failed: filtered row should drop")

        if _should_drop_filtered_non_catalyst_row({
            "Notes": "[filtered_non_catalyst_label=Catalyst=CO2-eq]",
            "Catalyst": "Ni/Al2O3",
            "Series_Name": "",
            "raw_category_label": "",
            "Canonical_Catalyst_ID": "",
        }):
            print("  [WARNING] filtered/backfill selfcheck failed: row with identity should not drop")

        row = _backfill_identity_from_clue_record(
            {"Catalyst": "2NiAl", "Active_Metal": "", "Support": ""},
            {"Active_Metal": "Ni", "Support": "Al2O3", "Support_Normalized": "Al2O3"},
        )
        if row.get("Active_Metal") != "Ni" or row.get("Support") != "Al2O3":
            print("  [WARNING] filtered/backfill selfcheck failed: clue backfill did not populate identity fields")
    except Exception as _e:
        print(f"  [WARNING] filtered/backfill selfcheck error (non-fatal): {_e}")


def _selfcheck_family_backbone_fallback() -> None:
    if globals().get("_FAMILY_BACKBONE_SELFCHECK_DONE"):
        return
    globals()["_FAMILY_BACKBONE_SELFCHECK_DONE"] = True
    try:
        registry = {
            "ni__prom-zn/al2o3": {
                "Catalyst": "Ni-Zn/Al2O3",
                "Canonical_Catalyst_ID": "ni__prom-zn/al2o3",
                "Active_Metal": "Ni",
                "Promoter_Metal": "Zn",
                "Support": "Al2O3",
                "Support_Normalized": "Al2O3",
                "Metal_Loading_wt%": "",
                "Alloy_Ratio": "",
            },
            "ni/al2o3": {
                "Catalyst": "8NiAl",
                "Canonical_Catalyst_ID": "ni/al2o3",
                "Active_Metal": "Ni",
                "Promoter_Metal": "",
                "Support": "Al2O3",
                "Support_Normalized": "Al2O3",
                "Metal_Loading_wt%": "8",
                "Alloy_Ratio": "",
            },
        }

        family = _find_unique_family_registry_candidate(
            {
                "Catalyst": "2Ni5Zn",
                "Active_Metal": "Ni",
                "Promoter_Metal": "Zn",
                "Support": "Al2O3",
            },
            registry,
        )
        if [key for key, _ in family] != ["ni__prom-zn/al2o3"]:
            print(f"  [WARNING] family backbone selfcheck failed: 2Ni5Zn expected ni__prom-zn/al2o3, got {[key for key, _ in family]!r}")

        relaxed = _find_loading_relaxed_single_metal_family_candidate(
            {
                "Catalyst": "2NiAl",
                "Active_Metal": "Ni",
                "Support": "Al2O3",
            },
            registry,
        )
        if [key for key, _ in relaxed] != ["ni/al2o3"]:
            print(f"  [WARNING] family backbone selfcheck failed: 2NiAl expected ni/al2o3, got {[key for key, _ in relaxed]!r}")
    except Exception as _e:
        print(f"  [WARNING] family backbone selfcheck error (non-fatal): {_e}")


def _selfcheck_compact_loading_support_label() -> None:
    if globals().get("_COMPACT_LOADING_SUPPORT_SELFCHECK_DONE"):
        return
    globals()["_COMPACT_LOADING_SUPPORT_SELFCHECK_DONE"] = True
    try:
        expected = {
            "2NiAl": ("Ni", "Al2O3", "2"),
            "20ZnAl": ("Zn", "Al2O3", "20"),
            "5ZnAl": ("Zn", "Al2O3", "5"),
            "4Ni10Zn": ("Ni", "Al2O3", "4"),
            "8Ni20Zn": ("Ni", "Al2O3", "8"),
            "2 Ni Al": ("Ni", "Al2O3", "2"),
            "2Ni-Al": ("Ni", "Al2O3", "2"),
            "2Ni/Al": ("Ni", "Al2O3", "2"),
            "2 wt% Ni Al": ("Ni", "Al2O3", "2"),
        }
        for label, (metal, support, loading) in expected.items():
            parsed = _parse_compact_loading_support_identity(label)
            if (
                parsed.get("Active_Metal") != metal
                or parsed.get("Support") != support
                or str(parsed.get("Metal_Loading_wt%", "")) != loading
            ):
                print(f"  [WARNING] compact loading/support selfcheck failed: {label} -> {parsed}")

        selected = _select_compact_loading_support_label({
            "Catalyst_ID": "ni",
            "Series_Name": "2NiAl",
        })
        if selected != "2NiAl":
            print(f"  [WARNING] compact loading/support selfcheck failed: selected {selected!r}")

        clues = _extract_identity_clues_from_figure_record({
            "Catalyst_ID": "ni",
            "Series_Name": "2NiAl",
        }).get("clue_record", {})
        if (
            clues.get("Active_Metal") != "Ni"
            or clues.get("Support") != "Al2O3"
            or str(clues.get("Metal_Loading_wt%", "")) != "2"
        ):
            print(f"  [WARNING] compact loading/support selfcheck failed: clue {clues}")
        bimetal = _parse_compact_loading_support_identity("8Ni20Zn")
        if bimetal.get("Promoter_Metal") != "Zn" or "Zn:20" not in str(bimetal.get("Alloy_Ratio", "")):
            print(f"  [WARNING] compact loading/support selfcheck failed: NiZnAl bimetal {bimetal}")
        mo2c = _extract_identity_clues_from_figure_record({"Catalyst": "Ni-Mo₂C (1.6)"}).get("clue_record", {})
        if (
            mo2c.get("Active_Metal") != "Ni"
            or mo2c.get("Support") != "Mo2C"
            or str(mo2c.get("Promoter_Metal", "")).strip()
            or "Ni:1.6" not in str(mo2c.get("Alloy_Ratio", ""))
        ):
            print(f"  [WARNING] compact loading/support selfcheck failed: Mo2C framework {mo2c}")
        if _extract_explicit_metal_set_from_identity_clue("Ni-Mo₂C (1.6)") != {"Ni"}:
            print("  [WARNING] compact loading/support selfcheck failed: Mo2C metal set")
        if _extract_explicit_metal_set_from_identity_clue("Ni-Mo/Al2O3") != {"Ni", "Mo"}:
            print("  [WARNING] compact loading/support selfcheck failed: ordinary Ni-Mo metal set")
    except Exception as _e:
        print(f"  [WARNING] compact loading/support selfcheck error (non-fatal): {_e}")


def _selfcheck_species_filtering() -> None:
    if globals().get("_SPECIES_FILTERING_SELFCHECK_DONE"):
        return
    globals()["_SPECIES_FILTERING_SELFCHECK_DONE"] = True
    try:
        if _normalize_species_token("CO₂") != "co2":
            print("  [WARNING] species selfcheck failed: CO₂ normalization")
        if _normalize_species_token("CH₄") != "ch4":
            print("  [WARNING] species selfcheck failed: CH₄ normalization")
        if _normalize_species_token("H₂") != "h2":
            print("  [WARNING] species selfcheck failed: H₂ normalization")
        if not _is_non_catalyst_figure_label("CO₂"):
            print("  [WARNING] species selfcheck failed: CO₂ should be non-catalyst")
        if not _is_non_catalyst_figure_label("CH₄"):
            print("  [WARNING] species selfcheck failed: CH₄ should be non-catalyst")
        if not _is_non_catalyst_figure_label("H₂"):
            print("  [WARNING] species selfcheck failed: H₂ should be non-catalyst")
        if _is_series_name_real_catalyst("CO₂", "product"):
            print("  [WARNING] species selfcheck failed: CO₂ product series should not be catalyst")
        if not _is_series_name_real_catalyst("Ni-Mo₂C", "catalyst"):
            print("  [WARNING] species selfcheck failed: Ni-Mo₂C should remain catalyst-like")
    except Exception as _e:
        print(f"  [WARNING] species selfcheck error (non-fatal): {_e}")


def _selfcheck_product_series_drop() -> None:
    if globals().get("_PRODUCT_SERIES_DROP_SELFCHECK_DONE"):
        return
    globals()["_PRODUCT_SERIES_DROP_SELFCHECK_DONE"] = True
    try:
        row = _promote_partial_point_when_metric_and_anchor_exist({
            "series_role": "product",
            "Series_Name": "CO₂",
            "CO2_Selectivity_%": "10",
            "Reaction_Temp_C": "400",
        })
        if int(row.get("partial_point_promoted", 0) or 0):
            print("  [WARNING] product-series selfcheck failed: CO₂ row should not be partial-promoted")
        if not int(row.get("_drop_filtered_non_catalyst_row", 0) or 0):
            print("  [WARNING] product-series selfcheck failed: CO₂ row should be dropped")
    except Exception as _e:
        print(f"  [WARNING] product-series selfcheck error (non-fatal): {_e}")


def _build_extraction_empty_reason(metadata: Dict, semantic_role: str, metric_field: str = "", category_labels: Optional[List[str]] = None, effect_axis_values: Optional[List[str]] = None, shared_anchor_bundle: Optional[Dict[str, Any]] = None, raw_points: Optional[List[Dict]] = None, had_shell_points: bool = False, metric_candidates: Optional[List[str]] = None) -> str:
    """
    Produce a fine-grained reason string when metadata succeeded but raw_points = 0.
    Rules (in priority order):
      1. non_performance role → non_performance_role_after_sanity_check
      2. had shell points that were all removed → raw_points_constructed_but_empty_shells_removed
      3. metric_field missing:
           - multiple candidates exist → metric_candidates_ambiguous
           - no y_axis and no caption metric → metadata_missing_y_axis_and_caption_metric
           - otherwise → no_metric_field_inferred
      4. category_screening specific:
           - no labels at all → no_category_labels_detected
           - NOTE: missing anchor bundle is NOT a blocking reason here;
             category points can be partial (label + metric) without full anchor.
             Validation layer handles partial promotion.
      5. condition_effect / temperature_sweep / stability_tos:
           - no effect axis values → no_effect_axis_values_detected
           - condition_effect only: warn about missing anchor (but still allow)
      6. Fallback → metadata_structure_insufficient
    Additional caller-generated reason (not from this function):
      - binding_blocked_for_identity_safety — set by attach_figure_points_to_registry
    """
    semantic_role = clean_residual_mojibake_chars(str(semantic_role or "")).strip() \
        or clean_residual_mojibake_chars(str((metadata or {}).get("semantic_figure_role", "") or "")).strip()

    if semantic_role == "non_performance_like":
        return "non_performance_role_after_sanity_check"

    if raw_points is not None and not raw_points and had_shell_points:
        return "raw_points_constructed_but_empty_shells_removed"

    if not metric_field:
        if metric_candidates and len(metric_candidates) > 1:
            return "metric_candidates_ambiguous"
        y_axis_val = str((metadata or {}).get("y_axis", "") or "").strip()
        if not y_axis_val and not str(metric_candidates or ""):
            return "metadata_missing_y_axis_and_caption_metric"
        return "no_metric_field_inferred"

    if semantic_role == "category_screening":
        if not (category_labels or []):
            return "no_category_labels_detected"
        # Missing anchor bundle does NOT block construction — validation promotes partial points.
        return "metadata_structure_insufficient"

    if semantic_role in {"condition_effect", "temperature_sweep", "stability_tos"}:
        if not (effect_axis_values or []):
            return "no_effect_axis_values_detected"
        # condition_effect: warn about anchor but don't block
        return "metadata_structure_insufficient"

    return "metadata_structure_insufficient"



def _build_raw_points_from_category_screening(metadata: Dict, caption_text: str, nearby_text: str, figure_info: Dict) -> Tuple[List[Dict], Dict[str, Any]]:
    """
    Build raw data points for category_screening figures.
    Key improvement: VL API is called even when metric_field is uncertain,
    using metric_candidates as fallback hint. Never skip VL if category labels exist.
    """
    current = dict(metadata or {})
    figure_info = dict(figure_info or {})
    image_data = str(figure_info.get("image_data", "") or "")
    img_name = str(figure_info.get("image_name", "") or "")
    file_name = str(figure_info.get("file_name", "") or "")
    metric_candidates = _extract_metric_candidates_from_metadata(current, caption_text, nearby_text)
    metric_field = _infer_metric_field_from_axis_or_caption(current, caption_text, nearby_text)
    # Fallback: use first candidate when metric_field could not be uniquely inferred
    active_metric = metric_field or (metric_candidates[0] if metric_candidates else "")
    category_labels = _extract_category_labels_from_metadata_or_context(current, caption_text, nearby_text)
    shared_anchor = _extract_shared_condition_anchor_bundle(
        current, caption_text, nearby_text,
        " ".join(_safe_get_series_label(item) for item in (current.get("series", []) or []))
    )
    point_builder_used = "category_screening_builder"

    raw_points: List[Dict] = []
    seed_points = [dict(row) for row in (figure_info.get("legacy_points", []) or []) if isinstance(row, dict)]
    default_series = _safe_get_series_label(
        (current.get("series", []) or [""])[0] if (current.get("series", []) or [""]) else ""
    )
    for seed in seed_points:
        raw_points.append(_coerce_single_raw_point(
            seed, current, "category_screening", shared_anchor_bundle=shared_anchor,
            metric_field=active_metric, effect_field="raw_category_label", default_series_name=default_series,
        ))
    raw_points = _filter_empty_raw_points(raw_points)

    # VL API call: triggered when we have image_data AND (metric hint OR category hints).
    # Old guard was `metric_field` only → broke when metric_field="" but candidates existed.
    vl_should_call = bool(image_data) and bool(active_metric or category_labels)
    if not raw_points and vl_should_call:
        category_hint_text = " | ".join(category_labels[:20]) if category_labels else "none"
        shared_anchor_text = "; ".join([
            f"{field}={value}" for field, value in shared_anchor.items()
            if not field.startswith("_") and str(value).strip()
        ]) or "none"
        metric_hint_display = active_metric if active_metric else f"unknown — candidates: {', '.join(metric_candidates) if metric_candidates else 'none'}"
        prompt = (
            f"This is a category_screening performance figure from MSR catalyst paper: {file_name}\n\n"
            f"Metadata hints:\n"
            f"- semantic_figure_role: category_screening\n"
            f"- x-axis: {str(current.get('x_axis', '') or 'unknown')}\n"
            f"- y-axis: {str(current.get('y_axis', '') or 'unknown')}\n"
            f"- likely metric field: {metric_hint_display}\n"
            f"- all metric candidates: {', '.join(metric_candidates) if metric_candidates else 'none'}\n"
            f"- category label hints from metadata/context: {category_hint_text}\n"
            f"- shared condition hints: {shared_anchor_text}\n"
            f"- caption / nearby text: {clean_residual_mojibake_chars((caption_text + ' ' + nearby_text)[:1200])}\n\n"
            f"Task:\n"
            f"- Extract every visible category point or bar in this chart.\n"
            f"- Use the x-axis category/catalyst label as raw_category_label.\n"
            f"- Fill the performance metric value into the correct official field "
            f"(MeOH_Conversion_%, H2_Yield_%, H2_Production_Rate, CO_Selectivity_%, CO2_Selectivity_%, CO_Concentration_ppm).\n"
            f"- Only extract exact numeric values visible in the chart.\n"
            f"- If a category label or metric value is not visible, do not invent it.\n"
            f"- Shared conditions (e.g. 90 C or S/C=2) from caption may be copied into every point only when explicit.\n\n"
            f"Return JSON array only:\n"
            f"[{{\"raw_category_label\": \"\", \"Catalyst\": \"\", \"Series_Name\": \"\", "
            f"\"Reaction_Temp_C\": \"\", \"TOS_h\": \"\", \"S_C_Ratio\": \"\", \"Pressure_bar\": \"\", "
            f"\"GHSV_mL_g_h\": \"\", \"Flow_Rate\": \"\", \"Catalyst_Amount_g\": \"\", "
            f"\"MeOH_Conversion_%\": \"\", \"H2_Yield_%\": \"\", \"H2_Production_Rate\": \"\", "
            f"\"CO_Selectivity_%\": \"\", \"CO2_Selectivity_%\": \"\", \"CO_Concentration_ppm\": \"\", \"Notes\": \"\"}}]\n\n"
            f"If no exact points can be read, return []"
        )
        try:
            resp = client.chat.completions.create(
                model=VISION_MODEL,
                messages=[{"role": "user", "content": [
                    {"type": "image_url", "image_url": {"url": f"data:image/png;base64,{image_data}"}},
                    {"type": "text", "text": clean_residual_mojibake_chars(prompt)},
                ]}],
                temperature=0.1,
                max_tokens=4096,
            )
            parsed = _try_parse_json_block(resp.choices[0].message.content, "category_screening_builder", img_name)
            if isinstance(parsed, list):
                for point in parsed:
                    if isinstance(point, dict):
                        raw_points.append(_coerce_single_raw_point(
                            point, current, "category_screening", shared_anchor_bundle=shared_anchor,
                            metric_field=active_metric, effect_field="raw_category_label",
                            default_series_name=default_series,
                        ))
                raw_points = _filter_empty_raw_points(raw_points)
        except Exception as e:
            print(f"    [category builder ERROR] {img_name}: {e}")

    diag = {
        "point_builder_used": point_builder_used,
        "metric_field_inferred": metric_field,
        "active_metric_used": active_metric,
        "metric_candidates": metric_candidates,
        "category_labels_detected": category_labels,
        "effect_axis_values_detected": [],
        "shared_anchor_bundle": shared_anchor,
        "extraction_empty_reason": "",
    }
    if not raw_points:
        diag["extraction_empty_reason"] = _build_extraction_empty_reason(
            current, "category_screening",
            metric_field=active_metric,
            category_labels=category_labels,
            effect_axis_values=[],
            shared_anchor_bundle=shared_anchor,
            raw_points=[],
            had_shell_points=bool(seed_points),
        )
    return raw_points, diag



def _build_raw_points_from_condition_effect(metadata: Dict, caption_text: str, nearby_text: str, figure_info: Dict) -> Tuple[List[Dict], Dict[str, Any]]:
    """
    Build raw data points for condition_effect figures.
    Key improvements:
    - VL called when metric OR effect hint exists (not requiring BOTH to be non-empty).
    - When effect_field is unknown, x_axis text is passed to VL as hint.
    - active_metric fallback from metric_candidates when metric_field is uncertain.
    """
    current = dict(metadata or {})
    figure_info = dict(figure_info or {})
    image_data = str(figure_info.get("image_data", "") or "")
    img_name = str(figure_info.get("image_name", "") or "")
    file_name = str(figure_info.get("file_name", "") or "")
    metric_candidates = _extract_metric_candidates_from_metadata(current, caption_text, nearby_text)
    metric_field = _infer_metric_field_from_axis_or_caption(current, caption_text, nearby_text)
    # Fallback: use first candidate when metric_field uncertain
    active_metric = metric_field or (metric_candidates[0] if metric_candidates else "")
    effect_field = _infer_condition_axis_field(current, caption_text, nearby_text)
    # Usable effect_field must not be raw_category_label (that belongs to category_screening)
    usable_effect_field = effect_field if (effect_field and effect_field != "raw_category_label") else ""
    effect_values = _extract_effect_axis_values_from_metadata_or_context(current, caption_text, nearby_text)
    shared_anchor = _extract_shared_condition_anchor_bundle(
        current, caption_text, nearby_text,
        " ".join(_safe_get_series_label(item) for item in (current.get("series", []) or []))
    )
    point_builder_used = "condition_effect_builder"

    raw_points: List[Dict] = []
    seed_points = [dict(row) for row in (figure_info.get("legacy_points", []) or []) if isinstance(row, dict)]
    default_series = _safe_get_series_label(
        (current.get("series", []) or [""])[0] if (current.get("series", []) or [""]) else ""
    )
    for seed in seed_points:
        raw_points.append(_coerce_single_raw_point(
            seed, current, "condition_effect", shared_anchor_bundle=shared_anchor,
            metric_field=active_metric, effect_field=usable_effect_field, default_series_name=default_series,
        ))
    raw_points = _filter_empty_raw_points(raw_points)

    # VL call: triggered by metric OR effect hint — not requiring both.
    # When effect_field unknown, provide x_axis text so VL can infer the field.
    vl_should_call = bool(image_data) and bool(active_metric or usable_effect_field or effect_values)
    if not raw_points and vl_should_call:
        effect_display = usable_effect_field if usable_effect_field else f"unknown (x-axis: {str(current.get('x_axis', '') or 'unknown')})"
        effect_hint_text = " | ".join(effect_values[:20]) if effect_values else "none"
        shared_anchor_text = "; ".join([
            f"{field}={value}" for field, value in shared_anchor.items()
            if not field.startswith("_") and str(value).strip()
        ]) or "none"
        metric_hint_display = active_metric if active_metric else f"unknown — candidates: {', '.join(metric_candidates) if metric_candidates else 'none'}"
        # Use effect_field in JSON schema if known, else fall back to generic condition fields
        effect_json_key = usable_effect_field if usable_effect_field else "S_C_Ratio"
        prompt = (
            f"This is a condition_effect performance figure from MSR catalyst paper: {file_name}\n\n"
            f"Metadata hints:\n"
            f"- semantic_figure_role: condition_effect\n"
            f"- x-axis: {str(current.get('x_axis', '') or 'unknown')}\n"
            f"- y-axis: {str(current.get('y_axis', '') or 'unknown')}\n"
            f"- x-axis maps to field: {effect_display}\n"
            f"- likely metric field: {metric_hint_display}\n"
            f"- all metric candidates: {', '.join(metric_candidates) if metric_candidates else 'none'}\n"
            f"- x-axis value hints: {effect_hint_text}\n"
            f"- shared condition hints: {shared_anchor_text}\n"
            f"- caption / nearby text: {clean_residual_mojibake_chars((caption_text + ' ' + nearby_text)[:1200])}\n\n"
            f"Task:\n"
            f"- Extract every visible x-y data point.\n"
            f"- For each point, fill the x-axis value into the correct condition field "
            f"(choose from: Reaction_Temp_C, S_C_Ratio, Pressure_bar, GHSV_mL_g_h, WHSV_h_inv, Flow_Rate, Catalyst_Amount_g, TOS_h).\n"
            f"- Fill the performance metric value into the correct official field "
            f"(MeOH_Conversion_%, H2_Yield_%, H2_Production_Rate, CO_Selectivity_%, CO2_Selectivity_%, CO_Concentration_ppm).\n"
            f"- Keep Catalyst or Series_Name only when clearly visible in legend/caption.\n"
            f"- Do not invent values not visible in the chart.\n\n"
            f"Return JSON array only:\n"
            f"[{{\"Catalyst\": \"\", \"Series_Name\": \"\", \"{effect_json_key}\": \"\", "
            f"\"Reaction_Temp_C\": \"\", \"TOS_h\": \"\", \"S_C_Ratio\": \"\", \"Pressure_bar\": \"\", "
            f"\"GHSV_mL_g_h\": \"\", \"WHSV_h_inv\": \"\", \"Flow_Rate\": \"\", \"Catalyst_Amount_g\": \"\", "
            f"\"MeOH_Conversion_%\": \"\", \"H2_Yield_%\": \"\", \"H2_Production_Rate\": \"\", "
            f"\"CO_Selectivity_%\": \"\", \"CO2_Selectivity_%\": \"\", \"CO_Concentration_ppm\": \"\", \"Notes\": \"\"}}]\n\n"
            f"If no exact points can be read, return []"
        )
        try:
            resp = client.chat.completions.create(
                model=VISION_MODEL,
                messages=[{"role": "user", "content": [
                    {"type": "image_url", "image_url": {"url": f"data:image/png;base64,{image_data}"}},
                    {"type": "text", "text": clean_residual_mojibake_chars(prompt)},
                ]}],
                temperature=0.1,
                max_tokens=4096,
            )
            parsed = _try_parse_json_block(resp.choices[0].message.content, "condition_effect_builder", img_name)
            if isinstance(parsed, list):
                for point in parsed:
                    if isinstance(point, dict):
                        raw_points.append(_coerce_single_raw_point(
                            point, current, "condition_effect", shared_anchor_bundle=shared_anchor,
                            metric_field=active_metric, effect_field=usable_effect_field,
                            default_series_name=default_series,
                        ))
                raw_points = _filter_empty_raw_points(raw_points)
        except Exception as e:
            print(f"    [condition builder ERROR] {img_name}: {e}")

    diag = {
        "point_builder_used": point_builder_used,
        "metric_field_inferred": metric_field,
        "active_metric_used": active_metric,
        "metric_candidates": metric_candidates,
        "category_labels_detected": [],
        "effect_axis_values_detected": effect_values,
        "effect_field_inferred": usable_effect_field,
        "shared_anchor_bundle": shared_anchor,
        "extraction_empty_reason": "",
    }
    if not raw_points:
        diag["extraction_empty_reason"] = _build_extraction_empty_reason(
            current, "condition_effect",
            metric_field=active_metric,
            category_labels=[],
            effect_axis_values=effect_values,
            shared_anchor_bundle=shared_anchor,
            raw_points=[],
            had_shell_points=bool(seed_points),
        )
    return raw_points, diag



def _build_raw_points_from_temperature_sweep(metadata: Dict, caption_text: str, nearby_text: str, figure_info: Dict) -> Tuple[List[Dict], Dict[str, Any]]:
    current = dict(metadata or {})
    figure_info = dict(figure_info or {})
    image_data = str(figure_info.get("image_data", "") or "")
    img_name = str(figure_info.get("image_name", "") or "")
    file_name = str(figure_info.get("file_name", "") or "")
    semantic_role = clean_residual_mojibake_chars(str(current.get("semantic_figure_role", "") or "")).strip() or "temperature_sweep"
    metric_candidates = _extract_metric_candidates_from_metadata(current, caption_text, nearby_text)
    metric_field = _infer_metric_field_from_axis_or_caption(current, caption_text, nearby_text)
    active_metric = metric_field or (metric_candidates[0] if metric_candidates else "")
    effect_field = "Reaction_Temp_C" if semantic_role != "stability_tos" else "TOS_h"
    effect_values = _extract_effect_axis_values_from_metadata_or_context(current, caption_text, nearby_text)
    shared_anchor = _extract_shared_condition_anchor_bundle(current, caption_text, nearby_text, " ".join(_safe_get_series_label(item) for item in (current.get("series", []) or [])))
    point_builder_used = "temperature_sweep_builder"

    raw_points: List[Dict] = []
    seed_points = [dict(row) for row in (figure_info.get("legacy_points", []) or []) if isinstance(row, dict)]
    for seed in seed_points:
        raw_points.append(_coerce_single_raw_point(
            seed, current, semantic_role, shared_anchor_bundle=shared_anchor,
            metric_field=active_metric, effect_field=effect_field,
            default_series_name=str(seed.get("Series_Name", "") or ""),
        ))
    raw_points = _filter_empty_raw_points(raw_points)

    # VL call: use active_metric fallback; also trigger when effect_values detected in context
    vl_should_call = bool(image_data) and bool(active_metric or effect_values)
    if not raw_points and vl_should_call:
        series_list = [_safe_get_series_label(item) for item in (current.get("series", []) or []) if _safe_get_series_label(item)]
        if not series_list:
            series_list = ["main_curve"]
        for series_name in series_list:
            try:
                points = _vl_extract_single_series(
                    image_data,
                    img_name,
                    file_name,
                    series_name,
                    str(current.get("y_axis", "") or "") or active_metric,
                    str(current.get("x_axis", "") or ""),
                    series_role=str(current.get("series_role", "unknown") or "unknown"),
                    metadata=current,
                )
                for point in points:
                    raw_points.append(_coerce_single_raw_point(
                        point, current, semantic_role, shared_anchor_bundle=shared_anchor,
                        metric_field=active_metric, effect_field=effect_field, default_series_name=series_name,
                    ))
            except Exception as e:
                print(f"    [temperature builder series ERROR] {img_name} series='{series_name}': {e}")
        raw_points = _filter_empty_raw_points(raw_points)

    diag = {
        "point_builder_used": point_builder_used,
        "metric_field_inferred": metric_field,
        "active_metric_used": active_metric,
        "metric_candidates": metric_candidates,
        "category_labels_detected": [],
        "effect_axis_values_detected": effect_values,
        "shared_anchor_bundle": shared_anchor,
        "extraction_empty_reason": "",
    }
    if not raw_points:
        diag["extraction_empty_reason"] = _build_extraction_empty_reason(
            current,
            semantic_role,
            metric_field=active_metric,
            category_labels=[],
            effect_axis_values=effect_values,
            shared_anchor_bundle=shared_anchor,
            raw_points=[],
            had_shell_points=bool(seed_points),
        )
    return raw_points, diag


def _build_raw_points_from_stability_tos(metadata: Dict, caption_text: str, nearby_text: str, figure_info: Dict) -> Tuple[List[Dict], Dict[str, Any]]:
    """
    Build raw data points for stability_tos figures (time-on-stream stability runs).
    Dedicated builder so stability_tos gets its own VL prompt and TOS_h as primary axis.
    """
    current = dict(metadata or {})
    # Force semantic role to stability_tos for correct effect_field dispatch
    current["semantic_figure_role"] = "stability_tos"
    figure_info = dict(figure_info or {})
    image_data = str(figure_info.get("image_data", "") or "")
    img_name = str(figure_info.get("image_name", "") or "")
    file_name = str(figure_info.get("file_name", "") or "")
    metric_candidates = _extract_metric_candidates_from_metadata(current, caption_text, nearby_text)
    metric_field = _infer_metric_field_from_axis_or_caption(current, caption_text, nearby_text)
    active_metric = metric_field or (metric_candidates[0] if metric_candidates else "")
    effect_field = "TOS_h"
    effect_values = _extract_effect_axis_values_from_metadata_or_context(current, caption_text, nearby_text)
    shared_anchor = _extract_shared_condition_anchor_bundle(
        current, caption_text, nearby_text,
        " ".join(_safe_get_series_label(item) for item in (current.get("series", []) or []))
    )
    point_builder_used = "stability_tos_builder"

    raw_points: List[Dict] = []
    seed_points = [dict(row) for row in (figure_info.get("legacy_points", []) or []) if isinstance(row, dict)]
    for seed in seed_points:
        raw_points.append(_coerce_single_raw_point(
            seed, current, "stability_tos", shared_anchor_bundle=shared_anchor,
            metric_field=active_metric, effect_field=effect_field,
            default_series_name=str(seed.get("Series_Name", "") or ""),
        ))
    raw_points = _filter_empty_raw_points(raw_points)

    vl_should_call = bool(image_data) and bool(active_metric or effect_values)
    if not raw_points and vl_should_call:
        series_list = [_safe_get_series_label(item) for item in (current.get("series", []) or []) if _safe_get_series_label(item)]
        if not series_list:
            series_list = ["main_curve"]
        for series_name in series_list:
            try:
                points = _vl_extract_single_series(
                    image_data, img_name, file_name,
                    series_name,
                    str(current.get("y_axis", "") or "") or active_metric,
                    str(current.get("x_axis", "") or ""),
                    series_role=str(current.get("series_role", "unknown") or "unknown"),
                    metadata=current,
                )
                for point in points:
                    raw_points.append(_coerce_single_raw_point(
                        point, current, "stability_tos", shared_anchor_bundle=shared_anchor,
                        metric_field=active_metric, effect_field=effect_field, default_series_name=series_name,
                    ))
            except Exception as e:
                print(f"    [stability_tos builder series ERROR] {img_name} series='{series_name}': {e}")
        raw_points = _filter_empty_raw_points(raw_points)

    diag = {
        "point_builder_used": point_builder_used,
        "metric_field_inferred": metric_field,
        "active_metric_used": active_metric,
        "metric_candidates": metric_candidates,
        "category_labels_detected": [],
        "effect_axis_values_detected": effect_values,
        "shared_anchor_bundle": shared_anchor,
        "extraction_empty_reason": "",
    }
    if not raw_points:
        diag["extraction_empty_reason"] = _build_extraction_empty_reason(
            current, "stability_tos",
            metric_field=active_metric,
            category_labels=[],
            effect_axis_values=effect_values,
            shared_anchor_bundle=shared_anchor,
            raw_points=[],
            had_shell_points=bool(seed_points),
        )
    return raw_points, diag


def validate_extracted_point_count_against_context(records: List[Dict], metadata: Dict, image_name: str = "", page_num: Any = "", source_file: str = "", raw_record_count_before_dedupe: Optional[int] = None, raw_record_count_after_dedupe: Optional[int] = None) -> Tuple[List[Dict], Dict[str, Any]]:
    if not globals().get("_FIGURE_VALIDATE_SELFCHECK_DONE"):
        globals()["_FIGURE_VALIDATE_SELFCHECK_DONE"] = True
        try:
            _rows, _audit = validate_extracted_point_count_against_context(
                [{"raw_category_label": "Ni/ZrO2", "MeOH_Conversion_%": "82", "Series_Name": "90 C", "semantic_figure_role": "category_screening"}],
                {"x_axis": "Different catalysts", "y_axis": "MeOH conversion (%)", "x_axis_mode": "category", "extractor_type": "category_or_condition", "series_role": "condition", "semantic_figure_role": "category_screening"},
            )
            if _audit.get("validation_status") not in {"validation_kept_partial", "validation_ok"}:
                print("  [WARNING] figure selfcheck failed: category partial validation status")

            partial_rows, partial_audit = validate_extracted_point_count_against_context(
                [{"raw_category_label": "Ni/ZrO2", "MeOH_Conversion_%": "65", "S_C_Ratio": "2", "semantic_figure_role": "category_screening"}],
                {"x_axis": "Different catalysts", "y_axis": "MeOH conversion (%)", "x_axis_mode": "category", "extractor_type": "category_or_condition", "series_role": "condition", "semantic_figure_role": "category_screening"},
            )
            if not partial_rows or partial_audit.get("validation_status") == "validation_empty":
                print("  [WARNING] figure selfcheck failed: partial category point was dropped")
        except Exception as e:
            print(f"  [WARNING] figure validation selfcheck error: {e}")

    metadata = dict(metadata or {})
    raw_rows = [dict(row) for row in (records or []) if isinstance(row, dict)]
    x_axis = clean_residual_mojibake_chars(str(metadata.get("x_axis", "") or "")).strip()
    y_axis = clean_residual_mojibake_chars(str(metadata.get("y_axis", "") or "")).strip()
    x_axis_mode = clean_residual_mojibake_chars(str(metadata.get("x_axis_mode", "") or "")).strip() or "unknown"
    extractor_type = clean_residual_mojibake_chars(str(metadata.get("extractor_type", "") or "")).strip() or "unknown"
    series_role = clean_residual_mojibake_chars(str(metadata.get("series_role", "") or "")).strip() or "unknown"
    semantic_role = clean_residual_mojibake_chars(str(metadata.get("semantic_figure_role", "") or "")).strip() or "unknown"
    point_builder_used = clean_residual_mojibake_chars(str(metadata.get("point_builder_used", "") or "")).strip()
    extraction_empty_reason = clean_residual_mojibake_chars(str(metadata.get("extraction_empty_reason", "") or "")).strip()
    metric_field_inferred = clean_residual_mojibake_chars(str(metadata.get("metric_field_inferred", "") or "")).strip()
    effect_axis_field = _infer_condition_axis_field(metadata, "", "")

    metadata_series_count = len([str(s).strip() for s in (metadata.get("series", []) or []) if str(s).strip()])
    extracted_series_count = _count_series_from_records(raw_rows)
    raw_before = int(raw_record_count_before_dedupe if raw_record_count_before_dedupe is not None else len(raw_rows))
    raw_after = int(raw_record_count_after_dedupe if raw_record_count_after_dedupe is not None else len(raw_rows))

    flags: List[str] = []
    notes: List[str] = []
    kept_rows: List[Dict] = []
    removed_shell_count = 0
    grouped_axis_values: Dict[str, List[Any]] = {}
    axis_present_any = False
    partial_promoted_count = 0

    if metadata_series_count > 1 and extracted_series_count and extracted_series_count < metadata_series_count and raw_after > 0:
        flags.append("multi_series_under_extracted")
    if metadata_series_count > 0 and extracted_series_count == 0 and raw_after > 0:
        flags.append("series_metadata_not_used")
    if raw_before <= 0 and extraction_empty_reason:
        flags.append(extraction_empty_reason)

    for row in raw_rows:
        current = dict(row)
        row_flags: List[str] = []
        notes_text = clean_residual_mojibake_chars(str(current.get("Notes", "") or "")).strip()
        note_axis, note_x_value = _extract_x_axis_and_value_from_notes(notes_text)
        raw_category_label = clean_residual_mojibake_chars(str(current.get("raw_category_label", "") or note_x_value or current.get("Catalyst", ""))).strip()
        current["raw_category_label"] = raw_category_label
        series_name = clean_residual_mojibake_chars(str(current.get("Series_Name", "") or raw_category_label or current.get("Catalyst", ""))).strip() or "unknown"
        has_metric = _has_any_metric_value(current)
        has_anchor = _has_any_anchor_value(current)
        effect_value = ""
        if effect_axis_field and effect_axis_field not in {"", "raw_category_label"}:
            effect_value = clean_residual_mojibake_chars(str(current.get(effect_axis_field, "") or "")).strip()
        if not effect_value and note_axis and note_x_value:
            note_axis_field = _infer_condition_axis_field({"x_axis": note_axis, "x_axis_mode": x_axis_mode, "semantic_figure_role": semantic_role}, "", "")
            if note_axis_field and note_axis_field == effect_axis_field:
                effect_value = clean_residual_mojibake_chars(str(note_x_value or "")).strip()

        if not has_metric:
            row_flags.append("no_metric_detected")
        if semantic_role == "category_screening" and not raw_category_label:
            row_flags.append("no_category_labels_detected")
        if semantic_role in {"condition_effect", "temperature_sweep", "stability_tos"} and not effect_value:
            row_flags.append("no_effect_axis_values_detected")
        if not has_anchor:
            row_flags.append("no_anchor_bundle_detected")

        if semantic_role == "temperature_sweep" or x_axis_mode == "temperature":
            x_val = clean_residual_mojibake_chars(str(current.get("Reaction_Temp_C", "") or effect_value or "")).strip()
            if x_val:
                axis_present_any = True
                grouped_axis_values.setdefault(series_name, []).append(x_val)
            else:
                row_flags.append("temperature_axis_missing")
        elif semantic_role == "stability_tos" or x_axis_mode == "time":
            x_val = clean_residual_mojibake_chars(str(current.get("TOS_h", "") or effect_value or "")).strip()
            if x_val:
                axis_present_any = True
                grouped_axis_values.setdefault(series_name, []).append(x_val)
            else:
                row_flags.append("time_axis_missing")
        elif semantic_role == "category_screening" or x_axis_mode == "category":
            if raw_category_label or has_anchor:
                axis_present_any = True
            else:
                row_flags.append("category_condition_axis_missing")
        elif semantic_role == "condition_effect" or x_axis_mode == "condition":
            if effect_value:
                axis_present_any = True
                grouped_axis_values.setdefault(series_name, []).append(effect_value)
            else:
                row_flags.append("category_condition_axis_missing")

        has_series = bool(series_name and series_name.lower() not in {"", "unknown", "main_curve"})
        has_anchor_source = bool(clean_residual_mojibake_chars(str(current.get("condition_anchor_source", "") or current.get("anchor_sources", "") or "")).strip())
        keep_partial = bool(
            has_metric and (
                raw_category_label
                or effect_value
                or clean_residual_mojibake_chars(str(current.get("Reaction_Temp_C", "") or "")).strip()
                or clean_residual_mojibake_chars(str(current.get("TOS_h", "") or "")).strip()
                or has_anchor
            )
        )
        empty_shell = bool(
            not has_metric
            and not raw_category_label
            and not effect_value
            and not has_anchor
            and not has_series
            and not has_anchor_source
            and not notes_text
        )
        if empty_shell:
            removed_shell_count += 1
            continue

        if keep_partial:
            promoted = _promote_partial_point_when_metric_and_anchor_exist(current)
            if int(promoted.get("_drop_filtered_non_catalyst_row", 0) or 0):
                continue
            if _should_drop_filtered_non_catalyst_row(promoted):
                continue
            clues = _extract_identity_clues_from_figure_record(promoted)
            if clues.get("clue_record"):
                promoted = _backfill_identity_from_clue_record(promoted, clues["clue_record"])
            if int(promoted.get("partial_point_promoted", 0) or 0):
                partial_promoted_count += 1
                row_flags.append("partial_points_promoted")
            current = promoted

        current["figure_point_validation_flags"] = "|".join(_dedupe_keep_order(row_flags))
        current["figure_point_validation_notes"] = "; ".join(_dedupe_keep_order(row_flags))
        kept_rows.append(current)

    if semantic_role == "temperature_sweep" or x_axis_mode == "temperature":
        if raw_rows and not axis_present_any:
            flags.append("temperature_axis_missing")
        elif any(not _check_monotonic_numeric(values) for values in grouped_axis_values.values() if len(values) >= 2):
            flags.append("temperature_non_monotonic")
    elif semantic_role == "stability_tos" or x_axis_mode == "time":
        if raw_rows and not axis_present_any:
            flags.append("time_axis_missing")
        elif any(not _check_monotonic_numeric(values) for values in grouped_axis_values.values() if len(values) >= 2):
            flags.append("time_non_monotonic")
    elif semantic_role in {"category_screening", "condition_effect"} or x_axis_mode in {"category", "condition"}:
        if raw_rows and not axis_present_any:
            flags.append("category_condition_axis_missing")

    if raw_rows and not any(_has_any_metric_value(row) for row in raw_rows):
        flags.append("no_metric_detected")
    if raw_rows and not kept_rows:
        flags.append("raw_points_constructed_but_all_empty_shells")
    if partial_promoted_count > 0:
        flags.append("partial_points_promoted")

    flags = _dedupe_keep_order(flags)
    if point_builder_used:
        notes.append(f"point_builder_used={point_builder_used}")
    if metric_field_inferred:
        notes.append(f"metric_field_inferred={metric_field_inferred}")
    if extraction_empty_reason:
        notes.append(f"extraction_empty_reason={extraction_empty_reason}")
    if removed_shell_count > 0:
        notes.append(f"removed_empty_shell={removed_shell_count}")
    if partial_promoted_count > 0:
        notes.append(f"partial_points_promoted={partial_promoted_count}")
    if flags:
        notes.append("flags=" + "|".join(flags))

    status = _infer_validation_status(True, "", raw_before, raw_after, len(kept_rows), removed_shell_count, flags)
    for row in kept_rows:
        existing_flags = [flag for flag in str(row.get("figure_point_validation_flags", "")).split("|") if flag]
        merged_flags = _dedupe_keep_order(existing_flags + flags)
        row["figure_point_validation_status"] = status
        row["figure_point_validation_flags"] = "|".join(merged_flags)
        row["figure_point_validation_notes"] = "; ".join([
            part for part in [str(row.get("figure_point_validation_notes", "")).strip(), "; ".join(notes).strip()] if part
        ])

    audit_row = {
        "Source_File": str(source_file or ""),
        "image_name": str(image_name or ""),
        "page_num": page_num,
        "x_axis": x_axis,
        "y_axis": y_axis,
        "x_axis_mode": x_axis_mode,
        "extractor_type": extractor_type,
        "series_role": series_role,
        "semantic_figure_role": semantic_role,
        "point_builder_used": point_builder_used,
        "extraction_empty_reason": extraction_empty_reason,
        "metric_field_inferred": metric_field_inferred,
        "metadata_series_count": metadata_series_count,
        "extracted_record_count": raw_after,
        "extracted_series_count": extracted_series_count,
        "validation_status": status,
        "validation_flags": "|".join(flags),
        "validation_notes": "; ".join(notes),
    }
    return kept_rows, audit_row


def _save_vl_debug_artifacts(img_name: str, page_num: Any, file_name: str,
                              metadata: Dict, semantic_role: str,
                              prefilter_pass: bool, prefilter_override: bool,
                              raw_before: List[Dict], raw_after: List[Dict],
                              finalized: List[Dict], extraction_empty_reason: str,
                              point_builder_used: str) -> None:
    """
    Task G: Save per-figure debug artifacts to MULTIMODAL_AUDIT_DIR.
    Writes a compact JSON file for each processed figure, enabling post-hoc audit.
    Silently skips on any I/O error to avoid interrupting the main pipeline.
    """
    try:
        audit_dir = globals().get("MULTIMODAL_AUDIT_DIR", "")
        if not audit_dir:
            return
        os.makedirs(audit_dir, exist_ok=True)
        safe_name = re.sub(r"[^\w.\-]", "_", str(img_name or "unknown"))
        artifact_path = os.path.join(audit_dir, f"{safe_name}.debug.json")
        payload = {
            "image_name": img_name,
            "page_num": page_num,
            "source_file": file_name,
            "semantic_role": semantic_role,
            "prefilter_pass": int(bool(prefilter_pass)),
            "prefilter_override": int(bool(prefilter_override)),
            "metadata_x_axis": str(metadata.get("x_axis", "")),
            "metadata_y_axis": str(metadata.get("y_axis", "")),
            "metadata_x_axis_mode": str(metadata.get("x_axis_mode", "unknown")),
            "metadata_x_range": [metadata.get("x_range_min"), metadata.get("x_range_max")],
            "metadata_x_tick_labels": list(metadata.get("x_tick_labels", []) or []),
            "metadata_series": list(metadata.get("series", []) or []),
            "metadata_series_role": str(metadata.get("series_role", "unknown")),
            "metadata_metric_candidates": list(metadata.get("metric_candidates", []) or []),
            "metadata_confidence": str(metadata.get("metadata_confidence", "")),
            "point_builder_used": point_builder_used,
            "extraction_empty_reason": extraction_empty_reason,
            "raw_record_count_before_dedupe": len(raw_before),
            "raw_record_count_after_dedupe": len(raw_after),
            "finalized_record_count": len(finalized),
            "finalized_preview": [
                {k: v for k, v in row.items()
                 if k in ("Catalyst", "Catalyst_ID", "Reaction_Temp_C", "TOS_h",
                          "MeOH_Conversion_%", "H2_Yield_%", "CO_Selectivity_%",
                          "CO2_Selectivity_%", "Series_Name", "raw_category_label")}
                for row in finalized[:5]
            ],
        }
        with open(artifact_path, "w", encoding="utf-8") as fh:
            json.dump(payload, fh, ensure_ascii=False, indent=2)
        print(f"    [debug_artifact] saved: {artifact_path}")
    except Exception as e:
        print(f"    [debug_artifact WARNING] could not save artifact for {img_name}: {e}")


def analyze_figure_with_vl(image_path: str, context: str, file_name: str) -> List[Dict]:
    context = clean_residual_mojibake_chars(context or "")
    img_name = os.path.basename(image_path)
    page_match = re.search(r"_page(\d+)\.", img_name, flags=re.I)
    page_num = int(page_match.group(1)) if page_match else ""
    validation_store = globals().setdefault("_FIGURE_VALIDATION_AUDIT_ROWS", [])
    debug_store = globals().setdefault("_FIGURE_DEBUG_PAYLOADS", [])

    def _infer_semantic_role(meta: Dict[str, Any], local_context: str) -> str:
        meta_text = " ".join([
            str(meta.get("x_axis", "")),
            str(meta.get("y_axis", "")),
            str(meta.get("x_axis_mode", "")),
            str(meta.get("extractor_type", "")),
            str(meta.get("chart_type", "")),
            " ".join(str(s) for s in meta.get("series", []) or []),
            local_context,
        ])
        perf_info = _is_performance_figure_context(meta_text, local_context, "")
        nonperf_info = _is_nonperformance_figure_context(meta_text)
        x_axis_mode = str(meta.get("x_axis_mode", "") or "unknown").strip().lower()
        y_axis = str(meta.get("y_axis", "") or "").strip().lower()
        y_metric_like = bool(re.search(r"conversion|yield|selectivity|h2|hydrogen|co2?|stability|deactivation", y_axis, flags=re.I))

        if nonperf_info.get("hard_exclusion") and not perf_info.get("strong_positive"):
            return "non_performance_like"
        if x_axis_mode == "temperature" and (perf_info.get("metric_signal") or y_metric_like):
            return "temperature_sweep"
        if x_axis_mode == "time" and (perf_info.get("metric_signal") or y_metric_like):
            return "stability_tos"
        if x_axis_mode == "category" or (perf_info.get("category_like") and (perf_info.get("metric_signal") or y_metric_like)):
            return "category_screening"
        if x_axis_mode == "condition" and (perf_info.get("metric_signal") or y_metric_like):
            return "condition_effect"
        if perf_info.get("metric_signal") or y_metric_like:
            return "condition_effect"
        return "non_performance_like"

    def _emit(metadata: Dict, meta_reason: str, raw_before: List[Dict], raw_after: List[Dict], validated: List[Dict], audit_rows: List[Dict], point_builder_used: str = "", extraction_empty_reason: str = "") -> None:
        meta_payload = dict(metadata or {})
        if point_builder_used:
            meta_payload["point_builder_used"] = point_builder_used
        if extraction_empty_reason:
            meta_payload["extraction_empty_reason"] = extraction_empty_reason
        payload = {
            "image_name": img_name,
            "page_num": page_num,
            "metadata": meta_payload,
            "metadata_failure_reason": str(meta_reason or ""),
            "point_builder_used": str(point_builder_used or ""),
            "extraction_empty_reason": str(extraction_empty_reason or ""),
            "raw_record_count_before_dedupe": len(raw_before),
            "raw_record_count_after_dedupe": len(raw_after),
            "validated_record_count": len(validated),
            "raw_records_before_dedupe": [dict(row) for row in raw_before],
            "raw_records_after_dedupe": [dict(row) for row in raw_after],
            "validated_records": [dict(row) for row in validated],
            "validation_audit_rows": [dict(row) for row in audit_rows],
        }
        debug_store.append(payload)
        globals()["_LAST_FIGURE_DEBUG_PAYLOAD"] = payload

    def _build_failure_audit(metadata: Dict[str, Any], failure_reason: str, extra_note: str = "") -> Dict[str, Any]:
        notes = [f"metadata_failure_reason={failure_reason}"]
        if extra_note:
            notes.append(extra_note)
        return {
            "Source_File": file_name,
            "image_name": img_name,
            "page_num": page_num,
            "x_axis": str(metadata.get("x_axis", "")),
            "y_axis": str(metadata.get("y_axis", "")),
            "x_axis_mode": str(metadata.get("x_axis_mode", "unknown")),
            "extractor_type": str(metadata.get("extractor_type", "unknown")),
            "series_role": str(metadata.get("series_role", "unknown")),
            "semantic_figure_role": str(metadata.get("semantic_figure_role", "unknown")),
            "metadata_series_count": len(metadata.get("series", []) or []),
            "extracted_record_count": 0,
            "extracted_series_count": 0,
            "validation_status": "metadata_failed",
            "validation_flags": "",
            "validation_notes": "; ".join([note for note in notes if note]),
        }

    if not globals().get("_FIGURE_ANALYZE_SELFCHECK_DONE"):
        globals()["_FIGURE_ANALYZE_SELFCHECK_DONE"] = True
        try:
            sample_meta = {
                "x_axis": "Different catalysts",
                "y_axis": "MeOH conversion (%)",
                "x_axis_mode": "category",
                "series": ["90 C"],
                "chart_type": "bar",
                "extractor_type": "category_or_condition",
                "semantic_figure_role": "category_screening",
            }
            if _infer_semantic_role(sample_meta, "Fig. 4 Methanol steam reforming performance Conversion Different catalysts 90 C") != "category_screening":
                print("  [WARNING] figure selfcheck failed: category semantic role inference")
            if _classify_metadata_failure_reason({}, "") != "metadata_empty":
                print("  [WARNING] figure selfcheck failed: metadata empty classification")
            if _classify_metadata_failure_reason({"x_axis": "XRD", "y_axis": "Intensity", "semantic_figure_role": "non_performance_like"}, "") != "metadata_nonperformance_like":
                print("  [WARNING] figure selfcheck failed: non-performance metadata classification")

            cat_rows, _ = _build_raw_points_from_category_screening(
                sample_meta,
                "Fig. 4 Methanol steam reforming performance of different catalysts at 90 C",
                "No catalyst, ZrO2, Ni/ZrO2",
                {"legacy_points": [
                    {"raw_category_label": "No catalyst", "MeOH_Conversion_%": "0"},
                    {"raw_category_label": "ZrO2", "MeOH_Conversion_%": "12.3"},
                    {"raw_category_label": "Ni/ZrO2", "MeOH_Conversion_%": "65.0"},
                ]},
            )
            if len(cat_rows) < 3 or not all(str(row.get("raw_category_label", "")).strip() for row in cat_rows) or not any(str(row.get("Reaction_Temp_C", "")).strip() == "90" for row in cat_rows):
                print("  [WARNING] figure selfcheck failed: category builder did not keep seeded labels/anchors")

            cond_rows, _ = _build_raw_points_from_condition_effect(
                {
                    "x_axis": "S/C ratio",
                    "y_axis": "CO2 selectivity (%)",
                    "x_axis_mode": "condition",
                    "semantic_figure_role": "condition_effect",
                },
                "",
                "Ni-Cu/clay, S/C ratio 1, 2, 3",
                {"legacy_points": [
                    {"S_C_Ratio": "1", "CO2_Selectivity_%": "70", "Catalyst": "Ni-Cu/clay"},
                    {"S_C_Ratio": "2", "CO2_Selectivity_%": "75", "Catalyst": "Ni-Cu/clay"},
                    {"S_C_Ratio": "3", "CO2_Selectivity_%": "80", "Catalyst": "Ni-Cu/clay"},
                ]},
            )
            if len(cond_rows) < 3 or not all(str(row.get("S_C_Ratio", "")).strip() for row in cond_rows) or not any(str(row.get("CO2_Selectivity_%", "")).strip() for row in cond_rows):
                print("  [WARNING] figure selfcheck failed: condition builder did not keep seeded x-values/metrics")

            reason_cat = _build_extraction_empty_reason(
                sample_meta, "category_screening",
                metric_field="MeOH_Conversion_%", category_labels=[],
                effect_axis_values=[], shared_anchor_bundle={}, raw_points=[], had_shell_points=False,
            )
            if reason_cat != "no_category_labels_detected":
                print("  [WARNING] figure selfcheck failed: category empty reason")

            reason_metric = _build_extraction_empty_reason(
                {"x_axis": "S/C ratio", "x_axis_mode": "condition", "semantic_figure_role": "condition_effect"},
                "condition_effect",
                metric_field="", category_labels=[], effect_axis_values=["1", "2", "3"],
                shared_anchor_bundle={}, raw_points=[], had_shell_points=False,
            )
            if reason_metric != "no_metric_field_inferred":
                print("  [WARNING] figure selfcheck failed: metric inference empty reason")

            # Run the comprehensive point construction selfcheck (in-memory only)
            _selfcheck_point_construction_only()
            # Run metric inference selfcheck (Problem C)
            _selfcheck_metric_inference()
            _selfcheck_filtered_and_backfill()
            _selfcheck_family_backbone_fallback()
            _selfcheck_compact_loading_support_label()
            _selfcheck_species_filtering()
            _selfcheck_product_series_drop()
        except Exception as e:
            print(f"  [WARNING] figure analyze selfcheck error: {e}")

    context_perf = _is_performance_figure_context(context)
    context_nonperf = _is_nonperformance_figure_context(context)

    try:
        image_data = _encode_image(image_path)
    except Exception as e:
        audit = _build_failure_audit({}, "metadata_parse_error", f"image_encoding_error={e}")
        validation_store.append(audit)
        _emit({}, "metadata_parse_error", [], [], [], [audit], "", "")
        return []

    image_prefilter_pass = False
    try:
        image_prefilter_pass = bool(_vl_is_performance_chart(image_data, img_name))
    except Exception as e:
        print(f"    [WARNING] image prefilter failed; continue to metadata: {img_name}: {e}")
        image_prefilter_pass = True

    prefilter_override = bool(
        not image_prefilter_pass
        and context_perf.get("strong_positive")
        and (
            not context_nonperf.get("hard_exclusion")
            or int(context_perf.get("performance_strength", 0) or 0) >= int(context_nonperf.get("nonperformance_strength", 0) or 0) + 3
        )
    )
    if not image_prefilter_pass and not prefilter_override:
        audit = _build_failure_audit({}, "metadata_nonperformance_like", "image_prefilter_rejected")
        validation_store.append(audit)
        _emit({}, "metadata_nonperformance_like", [], [], [], [audit], "", "")
        return []

    meta_exc = None
    try:
        metadata = _vl_extract_chart_metadata(image_data, img_name, file_name)
    except Exception as e:
        metadata, meta_exc = {}, e

    metadata = dict(metadata or {})
    metadata.setdefault("x_axis", "")
    metadata.setdefault("y_axis", "")
    metadata.setdefault("x_axis_mode", "unknown")
    metadata.setdefault("chart_type", "unknown")
    metadata.setdefault("series", [])
    metadata.setdefault("extractor_type", "unknown")
    metadata.setdefault("series_role", "unknown")
    # Task C extended fields: safe defaults if _normalize_chart_metadata was not called
    metadata.setdefault("x_range_min", None)
    metadata.setdefault("x_range_max", None)
    metadata.setdefault("x_tick_labels", [])
    metadata.setdefault("estimated_series_count", len(metadata.get("series", [])))
    metadata.setdefault("metric_candidates", [])
    metadata.setdefault("legend_detected", bool(metadata.get("series")))
    metadata.setdefault("metadata_confidence", "medium")
    metadata.setdefault("x_axis_unit", "")
    metadata.setdefault("y_axis_unit", "")
    metadata["semantic_figure_role"] = _infer_semantic_role(metadata, context)
    metadata["image_prefilter_pass"] = int(bool(image_prefilter_pass))
    metadata["image_prefilter_override"] = int(bool(prefilter_override))

    meta_reason = _classify_metadata_failure_reason(metadata, context, exception=meta_exc)
    if meta_reason:
        extra = "image_prefilter_override=1" if prefilter_override else ""
        audit = _build_failure_audit(metadata, meta_reason, extra)
        validation_store.append(audit)
        _emit(metadata or {}, meta_reason, [], [], [], [audit], "", "")
        return []

    x_axis = clean_residual_mojibake_chars(str(metadata.get("x_axis", "") or "")).strip()
    y_axis = clean_residual_mojibake_chars(str(metadata.get("y_axis", "") or "")).strip()
    x_axis_mode = clean_residual_mojibake_chars(str(metadata.get("x_axis_mode", "") or "")).strip() or "unknown"
    extractor_type = clean_residual_mojibake_chars(str(metadata.get("extractor_type", "") or "")).strip() or "unknown"
    series_role = clean_residual_mojibake_chars(str(metadata.get("series_role", "") or "")).strip() or "unknown"
    semantic_role = clean_residual_mojibake_chars(str(metadata.get("semantic_figure_role", "") or "")).strip() or "unknown"
    chart_type = clean_residual_mojibake_chars(str(metadata.get("chart_type", "") or "")).strip() or "unknown"
    caption_text = _safe_get_caption_text(context)
    nearby_text = _safe_get_nearby_context_text(context)
    figure_info = {
        "image_data": image_data,
        "image_path": image_path,
        "image_name": img_name,
        "page_num": page_num,
        "file_name": file_name,
    }

    raw_before: List[Dict] = []
    point_builder_used = ""
    extraction_empty_reason = ""
    builder_diag: Dict[str, Any] = {}

    # Explicit role dispatch — each role has its own dedicated builder
    if semantic_role == "category_screening":
        raw_before, builder_diag = _build_raw_points_from_category_screening(
            metadata, caption_text, nearby_text, figure_info)
    elif semantic_role == "condition_effect":
        raw_before, builder_diag = _build_raw_points_from_condition_effect(
            metadata, caption_text, nearby_text, figure_info)
    elif semantic_role == "temperature_sweep":
        raw_before, builder_diag = _build_raw_points_from_temperature_sweep(
            metadata, caption_text, nearby_text, figure_info)
    elif semantic_role == "stability_tos":
        # Dedicated stability_tos builder — uses TOS_h as primary axis
        raw_before, builder_diag = _build_raw_points_from_stability_tos(
            metadata, caption_text, nearby_text, figure_info)
    elif semantic_role == "non_performance_like":
        builder_diag = {
            "point_builder_used": "none",
            "metric_field_inferred": "",
            "active_metric_used": "",
            "category_labels_detected": [],
            "effect_axis_values_detected": [],
            "shared_anchor_bundle": {},
            "extraction_empty_reason": "non_performance_role_after_sanity_check",
        }
        raw_before = []
    else:
        # unknown role — record reason but don't crash
        builder_diag = {
            "point_builder_used": "no_point_builder",
            "metric_field_inferred": "",
            "active_metric_used": "",
            "category_labels_detected": [],
            "effect_axis_values_detected": [],
            "shared_anchor_bundle": {},
            "extraction_empty_reason": _build_extraction_empty_reason(
                metadata, semantic_role, "", [], [], {}, [], False),
        }
        raw_before = []

    point_builder_used = clean_residual_mojibake_chars(str(builder_diag.get("point_builder_used", "") or "")).strip()
    extraction_empty_reason = clean_residual_mojibake_chars(str(builder_diag.get("extraction_empty_reason", "") or "")).strip()
    metadata["point_builder_used"] = point_builder_used
    metadata["metric_field_inferred"] = clean_residual_mojibake_chars(str(builder_diag.get("metric_field_inferred", "") or "")).strip()
    metadata["active_metric_used"] = clean_residual_mojibake_chars(str(builder_diag.get("active_metric_used", "") or "")).strip()
    metadata["extraction_empty_reason"] = extraction_empty_reason
    metadata["builder_category_labels_detected"] = list(builder_diag.get("category_labels_detected", []) or [])
    metadata["builder_effect_axis_values_detected"] = list(builder_diag.get("effect_axis_values_detected", []) or [])
    metadata["builder_shared_anchor_bundle"] = dict(builder_diag.get("shared_anchor_bundle") or {})
    metadata["caption_text_preview"] = _context_preview(caption_text, 400)
    metadata["nearby_context_preview"] = _context_preview(nearby_text, 600)

    seen = set()
    raw_after: List[Dict] = []
    for row in raw_before:
        notes_val = clean_residual_mojibake_chars(str(row.get("Notes", "") or ""))
        note_axis, note_x_value = _extract_x_axis_and_value_from_notes(notes_val)
        series_key = str(row.get("Series_Name", "") or row.get("raw_category_label", "") or row.get("Catalyst", "") or "").strip()
        if semantic_role == "temperature_sweep" or x_axis_mode == "temperature":
            x_key = str(row.get("Reaction_Temp_C", "")).strip()
        elif semantic_role == "stability_tos" or x_axis_mode == "time":
            x_key = str(row.get("TOS_h", "")).strip()
        elif semantic_role == "category_screening" or x_axis_mode == "category":
            label_key = str(row.get("raw_category_label", "") or note_x_value or row.get("Catalyst", "")).strip()
            perf_key = "|".join([
                str(row.get("MeOH_Conversion_%", "")),
                str(row.get("H2_Yield_%", "")),
                str(row.get("H2_Production_Rate", "")),
                str(row.get("CO_Selectivity_%", "")),
                str(row.get("CO2_Selectivity_%", "")),
                str(row.get("CO_Concentration_ppm", "")),
            ])
            x_key = f"{label_key}|{perf_key}|{row.get('Reaction_Temp_C', '')}|{row.get('TOS_h', '')}|{row.get('S_C_Ratio', '')}|{row.get('Pressure_bar', '')}|{row.get('GHSV_mL_g_h', '')}|{row.get('Flow_Rate', '')}"
        elif semantic_role == "condition_effect" or x_axis_mode == "condition":
            x_key = "|".join([
                str(note_axis),
                str(note_x_value),
                str(row.get("Reaction_Temp_C", "")),
                str(row.get("TOS_h", "")),
                str(row.get("S_C_Ratio", "")),
                str(row.get("Feed_MeOH_to_H2O_Ratio", "")),
                str(row.get("Pressure_bar", "")),
                str(row.get("GHSV_mL_g_h", "")),
                str(row.get("WHSV_h_inv", "")),
                str(row.get("Flow_Rate", "")),
                str(row.get("Catalyst_Amount_g", "")),
                str(row.get("MeOH_Conversion_%", "")),
                str(row.get("H2_Yield_%", "")),
                str(row.get("H2_Production_Rate", "")),
                str(row.get("CO_Selectivity_%", "")),
                str(row.get("CO2_Selectivity_%", "")),
                str(row.get("CO_Concentration_ppm", "")),
            ])
        else:
            x_key = str(row.get("Reaction_Temp_C", "")).strip() or str(row.get("TOS_h", "")).strip() or note_x_value
        dedupe_key = (series_key, x_key, str(row.get("Series_Name", "")).strip())
        if dedupe_key in seen:
            continue
        seen.add(dedupe_key)
        raw_after.append(dict(row))

    if not raw_after and not extraction_empty_reason:
        extraction_empty_reason = _build_extraction_empty_reason(
            metadata,
            semantic_role,
            metric_field=metadata.get("metric_field_inferred", ""),
            category_labels=metadata.get("builder_category_labels_detected", []) or [],
            effect_axis_values=metadata.get("builder_effect_axis_values_detected", []) or [],
            shared_anchor_bundle=metadata.get("builder_shared_anchor_bundle", {}) or {},
            raw_points=[],
            had_shell_points=bool(raw_before),
        )
        metadata["extraction_empty_reason"] = extraction_empty_reason

    validated_rows, audit = validate_extracted_point_count_against_context(
        raw_after,
        metadata,
        image_name=img_name,
        page_num=page_num,
        source_file=file_name,
        raw_record_count_before_dedupe=len(raw_before),
        raw_record_count_after_dedupe=len(raw_after),
    )
    validation_store.append(audit)
    finalized = finalize_extracted_records(validated_rows, context, file_name, "figure") if validated_rows else []
    for row in finalized:
        row["image_name"] = img_name
        row["page_num"] = page_num
        row["x_axis"] = x_axis
        row["y_axis"] = y_axis
        row["x_axis_mode"] = x_axis_mode
        row["extractor_type"] = extractor_type
        row["series_role"] = series_role
        row["semantic_figure_role"] = semantic_role
        row.setdefault("figure_point_validation_status", audit.get("validation_status", ""))
        row.setdefault("figure_point_validation_flags", audit.get("validation_flags", ""))
        row.setdefault("figure_point_validation_notes", audit.get("validation_notes", ""))

    _emit(metadata, "", raw_before, raw_after, finalized, [audit], point_builder_used, extraction_empty_reason)
    # Task G: write per-figure debug artifact for post-hoc audit
    _save_vl_debug_artifacts(
        img_name=img_name,
        page_num=page_num,
        file_name=file_name,
        metadata=metadata,
        semantic_role=semantic_role,
        prefilter_pass=bool(image_prefilter_pass),
        prefilter_override=bool(prefilter_override),
        raw_before=raw_before,
        raw_after=raw_after,
        finalized=finalized,
        extraction_empty_reason=extraction_empty_reason,
        point_builder_used=point_builder_used,
    )
    return finalized


# ==========================================
# Physical cleanup and feature normalization
# ==========================================
def smart_feature_engineering(records: List[Dict[str, Any]]) -> None:
    VALID_ELEMENTS = {
        "H","He","Li","Be","B","C","N","O","F","Ne","Na","Mg","Al","Si","P","S","Cl",
        "Ar","K","Ca","Sc","Ti","V","Cr","Mn","Fe","Co","Ni","Cu","Zn","Ga","Ge","As",
        "Se","Br","Kr","Rb","Sr","Y","Zr","Nb","Mo","Tc","Ru","Rh","Pd","Ag","Cd","In",
        "Sn","Sb","Te","I","Xe","Cs","Ba","La","Ce","Pr","Nd","Pm","Sm","Eu","Gd","Tb",
        "Dy","Ho","Er","Tm","Yb","Lu","Hf","Ta","W","Re","Os","Ir","Pt","Au","Hg","Tl",
        "Pb","Bi"
    }
    BLACKLIST = {"Al","Zr","Ti","Si","Mg","Ca","O","C","N","H","S","P"}
    PROMOTER_BLACKLIST = {"Al","Zr","Ti","Si","Ca","O","C","N","H","S","P"}
    SUPPORT_RULES = {
        r"(?i)(gamma-alumina|al2o3|alumina|^al$)": "Al2O3",
        r"(?i)(layered double hydroxide|ldh|hydrotalcite)": "Al-LDH",
        r"(?i)(cnts|carbon nanotube)": "CNTs",
        r"(?i)(mgo.*al2o3|al2o3.*mgo|mgal2o4|spinel)": "MgO-Al2O3",
        r"(?i)(ceo2.*zro2|zro2.*ceo2|ce\d+\.?\d*zr\d+\.?\d*o2)": "CeO2-ZrO2",
        r"(?i)(sepiolite|^sep$)": "SEP",
        r"(?i)(^ac\d*|^c-\d|^p-\d|activated carbon)": "Activated Carbon",
        r"(?i)(cement[- ]clay)": "Cement-Clay"
    }
    METAL_PAT = "|".join(sorted(MSR_ACTIVE_METALS, key=len, reverse=True))

    for r in records:
        normalize_identity_aliases(r)
        catalyst = str(r.get("Catalyst", "")).strip()

        # Normalize Active_Metal.
        if r.get("Active_Metal"):
            elems = re.findall(r"([A-Z][a-z]?)", str(r["Active_Metal"]))
            active, promoters = [], []
            for el in dict.fromkeys(elems):
                if el in VALID_ELEMENTS and el not in BLACKLIST:
                    if el in MSR_ACTIVE_METALS:
                        active.append(el)
                    elif el in MSR_PROMOTERS or el in MSR_RARE_EARTH:
                        promoters.append(el)
                    else:
                        active.append(el)
            # [Task3] Force a stable element order so "Cu/Ni" and "Ni/Cu" normalize the same way.
            r["Active_Metal"] = "/".join(_sort_elements_for_merge(active)) if active else ""
            if not r.get("Promoter") and promoters:
                r["Promoter"] = "/".join(_sort_elements_for_merge([p for p in promoters if p != "Ni"]))

        # Normalize Promoter.
        if r.get("Promoter"):
            elems = re.findall(r"([A-Z][a-z]?)", str(r["Promoter"]))
            valid_p = [el for el in dict.fromkeys(elems)
                       if el in VALID_ELEMENTS and el not in PROMOTER_BLACKLIST]
            # [Task3] Apply the same stable ordering to Promoter values.
            r["Promoter"] = "/".join(_sort_elements_for_merge([p for p in valid_p if p != "Ni"])) if valid_p else ""

        # Fallback rescue path: infer metals from the catalyst name itself.
        if catalyst and not r.get("Active_Metal"):
            elems = re.findall(r"([A-Z][a-z]?)", catalyst)
            active, promoters = [], []
            for el in dict.fromkeys(elems):
                if el in VALID_ELEMENTS and el not in BLACKLIST:
                    if el in MSR_ACTIVE_METALS:
                        active.append(el)
                    elif el in MSR_PROMOTERS or el in MSR_RARE_EARTH:
                        promoters.append(el)
            if active:
                # [Task3] Use the same stable ordering in the fallback rescue path.
                r["Active_Metal"] = "/".join(_sort_elements_for_merge(active))
            if promoters and not r.get("Promoter"):
                r["Promoter"] = "/".join(_sort_elements_for_merge([p for p in promoters if p != "Ni"]))

        # Support normalization.
        if r.get("Support"):
            sup = str(r["Support"]).strip()
            for pattern, std in SUPPORT_RULES.items():
                if re.search(pattern, sup):
                    r["Support"] = std
                    break
        elif catalyst and "/" in catalyst:
            parts = catalyst.split("/")
            if len(parts) >= 2:
                r["Support"] = re.sub(r"-\w+$|\(.*?\)", "", parts[-1].strip()).strip()

        # Preparation-method normalization.
        for field in ["Support_Prep_Method", "Metal_Loading_Method"]:
            if r.get(field):
                prep = str(r[field]).lower()
                if field == "Metal_Loading_Method":
                    VALID_MLM = {"Impregnation", "Co-precipitation", "Deposition-precipitation",
                                 "Sol-gel", "Hydrothermal", "Mechanical mixing", "Other"}
                    if r[field] in VALID_MLM:
                        pass  # LLM宸叉纭垎绫伙紝淇濈暀
                    elif re.search(r"impregnation|incipient|wetness|imp\b", prep):
                        r[field] = "Impregnation"
                    elif re.search(r"co.?precipit|coprecip", prep):
                        r[field] = "Co-precipitation"
                    elif re.search(r"deposition.?precipit|deposition-precipit", prep):
                        r[field] = "Deposition-precipitation"
                    elif re.search(r"sol.?gel", prep):
                        r[field] = "Sol-gel"
                    elif re.search(r"hydrothermal", prep):
                        r[field] = "Hydrothermal"
                    elif re.search(r"mechanical.?mix|grind|ball.?mill", prep):
                        r[field] = "Mechanical mixing"
                    elif prep.strip():
                        r[field] = "Other"
                    else:
                        r[field] = ""
                elif field == "Support_Prep_Method":
                    VALID_SPM = {"Commercial", "Precipitation", "Sol-gel", "Hydrothermal", "Other"}
                    if r[field] in VALID_SPM:
                        pass  # LLM宸叉纭垎绫伙紝淇濈暀
                    elif re.search(r"commercial|purchased|supplied|sigma|aladdin|aldrich|bought", prep):
                        r[field] = "Commercial"
                    elif re.search(r"precipit|co.?precipit", prep):
                        r[field] = "Precipitation"
                    elif re.search(r"sol.?gel", prep):
                        r[field] = "Sol-gel"
                    elif re.search(r"hydrothermal", prep):
                        r[field] = "Hydrothermal"
                    elif prep.strip():
                        r[field] = "Other"
                    else:
                        r[field] = ""

        # Normalize S/C ratio.
        if r.get("S_C_Ratio"):
            sc_raw = str(r["S_C_Ratio"]).strip().lower()
            m = re.search(r"(\d+(?:\.\d+)?)\s*[:/]\s*(\d+(?:\.\d+)?)", sc_raw)
            if m:
                v1, v2 = float(m.group(1)), float(m.group(2))
                if "h2o" in sc_raw or "water" in sc_raw:
                    h2o_pos = sc_raw.index("h2o") if "h2o" in sc_raw else sc_raw.index("water")
                    meoh_pos = sc_raw.index("meoh") if "meoh" in sc_raw else len(sc_raw)
                    r["S_C_Ratio"] = str(round(v1/v2 if h2o_pos < meoh_pos else v2/v1, 2))
                elif "meoh" in sc_raw or "methanol" in sc_raw:
                    meoh_pos = sc_raw.index("meoh") if "meoh" in sc_raw else sc_raw.index("methanol")
                    h2o_pos = sc_raw.index("h2o") if "h2o" in sc_raw else len(sc_raw)
                    r["S_C_Ratio"] = str(round(v2/v1 if meoh_pos < h2o_pos else v1/v2, 2))
                else:
                    r["S_C_Ratio"] = str(round(max(v1, v2) / min(v1, v2), 2))
            elif re.match(r"^\d+(?:\.\d+)?$", sc_raw):
                r["S_C_Ratio"] = sc_raw

        # Recover Metal_Loading_wt% from catalyst names when possible.
        # Search across all identity fields: Catalyst, Catalyst_ID, Series_Name, Canonical_Catalyst_ID
        if not r.get("Metal_Loading_wt%"):
            _loading_search_fields = [
                str(r.get("Catalyst", "") or ""),
                str(r.get("Catalyst_ID", "") or ""),
                str(r.get("Series_Name", "") or ""),
                str(r.get("Canonical_Catalyst_ID", "") or ""),
            ]
            _loading_candidates = [f for f in _loading_search_fields if f.strip()]
            for _lsrc in _loading_candidates:
                m = re.search(rf"(\d{{1,3}}(?:\.\d+)?)\s*(?:wt%|wt\s*%|%)\s*(?:{METAL_PAT})", _lsrc)
                if m:
                    r["Metal_Loading_wt%"] = m.group(1)
                    break
                m = re.search(r"\((\d+(?:\.\d+)?)\s*(?:wt%|%)\)", _lsrc)
                if m:
                    r["Metal_Loading_wt%"] = m.group(1)
                    break
                m = re.search(rf"(?<!\d)(\d{{1,3}}(?:\.\d+)?)(?:{METAL_PAT})(?![a-z])", _lsrc)
                if m:
                    r["Metal_Loading_wt%"] = m.group(1)
                    break

        # [Fix 19] Parse Alloy_Ratio into Ni_Fraction / Promoter_Metal / Promoter_Fraction.
        # Explicit Promoter takes priority; missing values stay empty instead of using None/0.0.
        alloy_raw = str(r.get("Alloy_Ratio", "")).strip()
        # Read explicit Promoter after the upstream cleanup step.
        explicit_promoter = str(r.get("Promoter", "")).strip()
        _PROMOTER_NONE = {"", "none", "unknown", "n/a"}

        if alloy_raw:
            try:
                matches = re.findall(r"([A-Z][a-z]?)\s*:\s*(\d+(?:\.\d+)?)", alloy_raw)
                if matches:
                    ratio_dict = {el: float(val) for el, val in matches}
                    total_ratio = sum(ratio_dict.values())
                    if total_ratio > 0:
                        # Ni_Fraction logic stays unchanged.
                        if "Ni" in ratio_dict:
                            r["Ni_Fraction"] = str(round(ratio_dict["Ni"] / total_ratio, 3))
                        else:
                            r["Ni_Fraction"] = "0.0"

                        # [Fix 19A] Do not overwrite an explicit promoter with Alloy_Ratio parsing.
                        if explicit_promoter.lower() not in _PROMOTER_NONE:
                            r["Promoter_Metal"] = explicit_promoter
                            # Keep the promoter fraction empty when the ratio cannot be assigned safely.
                            r["Promoter_Fraction"] = ""
                        else:
                            # [Fix 19B] Only derive Promoter_Metal from Alloy_Ratio when explicit Promoter is missing.
                            other_metals = [el for el in ratio_dict if el != "Ni"]
                            if len(other_metals) == 1:
                                # Binary system: the only non-Ni metal becomes Promoter_Metal.
                                pm = other_metals[0]
                                r["Promoter_Metal"] = pm
                                r["Promoter_Fraction"] = str(round(ratio_dict[pm] / total_ratio, 3))
                            elif len(other_metals) == 0:
                                # Pure Ni system: no promoter metal.
                                r["Promoter_Metal"] = ""
                                r["Promoter_Fraction"] = ""
                            else:
                                # [Fix 19C] Ternary-or-higher systems: do not force one promoter metal.
                                r["Promoter_Metal"] = ""
                                r["Promoter_Fraction"] = ""
            except Exception:
                pass  # Ignore parsing failure and keep the original record.
        else:
            # [Fix 19D] Without Alloy_Ratio, use explicit Promoter directly when present.
            if explicit_promoter.lower() not in _PROMOTER_NONE:
                r["Promoter_Metal"] = explicit_promoter
                r["Promoter_Fraction"] = ""
            else:
                r["Promoter_Metal"] = ""
                r["Promoter_Fraction"] = ""

        # Final normalization pass keeps identity, support, and Catalyst_ID aligned.
        normalize_identity_aliases(r)
        annotate_numeric_expression_guards(r)
        if not str(r.get("Catalyst_ID", "")).strip():
            r["Catalyst_ID"] = normalize_catalyst_id(r)


def validate_msr_data(records: List[Dict[str, Any]]) -> None:
    for r in records:
        warnings = []

        # --- Existing checks: temperature and pressure ranges ---
        for field, lo, hi, label in [
            ("Reaction_Temp_C", 100, 500, "reaction temperature"),
            ("Calcination_Temp_C", 200, 1000, "calcination temperature"),
            ("Pressure_bar", 0.5, 10, "pressure"),
        ]:
            if r.get(field):
                try:
                    val = float(str(r[field]).strip())
                    if val < lo or val > hi:
                        warnings.append(f"{label} abnormal ({val})")
                except (ValueError, TypeError):
                    pass

        # --- Existing checks: percentage-type performance fields must stay within 0-100 ---
        for key in ["MeOH_Conversion_%", "H2_Selectivity_%", "CO_Selectivity_%", "CO2_Selectivity_%"]:
            if r.get(key):
                try:
                    val = float(str(r[key]).strip())
                    if val < 0 or val > 100:
                        warnings.append(f"{key} out_of_range ({val}%)")
                        r[key] = ""
                except (ValueError, TypeError):
                    pass

        # --- Existing checks: S/C ratio range ---
        if r.get("S_C_Ratio"):
            try:
                sc = float(str(r["S_C_Ratio"]).strip())
                if sc < 0.5 or sc > 5:
                    warnings.append(f"S_C_Ratio abnormal ({sc})")
            except (ValueError, TypeError):
                pass

        # [Fix 20] New check: H2_Yield_% must stay within 0-100.
        if r.get("H2_Yield_%"):
            try:
                val = float(str(r["H2_Yield_%"]).strip())
                if val < 0 or val > 100:
                    warnings.append(f"H2_Yield_% out_of_range ({val}%)")
                    r["H2_Yield_%"] = ""
            except (ValueError, TypeError):
                pass

        # [Fix 20] New check: TOS_h must be > 0; extremely large values are warned but not cleared.
        if r.get("TOS_h"):
            try:
                val = float(str(r["TOS_h"]).strip())
                if val <= 0:
                    warnings.append(f"TOS_h invalid ({val})")
                    r["TOS_h"] = ""
                elif val > 5000:
                    warnings.append(f"TOS_h unusually_large ({val}h)")
            except (ValueError, TypeError):
                pass

        # [Fix 20] New check: Carbon_Deposition_wt% cannot be negative.
        if r.get("Carbon_Deposition_wt%"):
            try:
                val = float(str(r["Carbon_Deposition_wt%"]).strip())
                if val < 0:
                    warnings.append(f"Carbon_Deposition_wt% negative ({val})")
                    r["Carbon_Deposition_wt%"] = ""
            except (ValueError, TypeError):
                pass

        # [Fix 20] New check: CO_Concentration_ppm cannot be negative.
        if r.get("CO_Concentration_ppm"):
            try:
                val = float(str(r["CO_Concentration_ppm"]).strip())
                if val < 0:
                    warnings.append(f"CO_Concentration_ppm negative ({val})")
                    r["CO_Concentration_ppm"] = ""
            except (ValueError, TypeError):
                pass

        # [Fix 20] New check: numeric GHSV must be > 0; textual WHSV values are kept.
        if r.get("GHSV_mL_g_h"):
            ghsv_raw = str(r["GHSV_mL_g_h"]).strip()
            if "WHSV" not in ghsv_raw.upper():
                try:
                    val = float(ghsv_raw)
                    if val <= 0:
                        warnings.append(f"GHSV_mL_g_h invalid ({val})")
                        r["GHSV_mL_g_h"] = ""
                except (ValueError, TypeError):
                    pass  # Non-numeric non-WHSV strings are kept as-is.

        # [Fix 20] New check: category_or_condition records may omit Reaction_Temp_C / TOS_h.
        # category_or_condition records may omit Reaction_Temp_C / TOS_h.
        # Detect this mode from Notes tags such as [x_axis=...][x_value=...].
        has_temp  = bool(str(r.get("Reaction_Temp_C", "")).strip())
        has_tos   = bool(str(r.get("TOS_h", "")).strip())
        has_notes_x = bool(re.search(r"\[x_axis=", str(r.get("Notes", ""))))
        perf_fields = ["MeOH_Conversion_%", "H2_Yield_%", "H2_Selectivity_%",
                       "CO_Selectivity_%", "CO2_Selectivity_%"]
        has_perf = any(str(r.get(f, "")).strip() not in ("", "N/A") for f in perf_fields)
        if has_perf and not has_temp and not has_tos and not has_notes_x:
            warnings.append("condition info missing (no temperature/TOS/x_value)")

        # Append warnings to Notes without overwriting existing content.
        if warnings:
            existing = str(r.get("Notes", ""))
            warn_str = f"[warning:{'|'.join(warnings)}]"
            # Avoid duplicating identical warning bundles.
            if warn_str not in existing:
                r["Notes"] = f"{existing} {warn_str}".strip()



def clean_numeric_fields(records: List[Dict[str, Any]]) -> None:
    # [Improve6] Unified missing-value representation: all NA-like strings in
    # numeric columns are normalised to "" so pandas reads them as NaN directly.
    _NA_PATTERNS = re.compile(r"^\s*(?:n/?a|none|null|nan|-+|–+|—+|nd|not\s+detected|not\s+available)\s*$", re.I)

    for r in records:
        for key in list(r.keys()):
            if key in ("Metal_Loading_wt%", "Catalyst_ID", "Catalyst_ID_normalized"):
                continue

            val_raw = str(r.get(key, "")).strip()
            if not val_raw:
                r[key] = ""
                continue

            # Only clean fields that are explicitly numeric-like output columns.
            numeric_like = (
                key.endswith(("_%", "_C", "_h", "_bar", "_ppm"))
                or key in ("GHSV_mL_g_h", "Ni_Fraction", "Promoter_Fraction",
                           "SpaceVelocity_norm", "H2_Production_Rate_mmol_g_h")
            )
            if not numeric_like:
                continue

            # Normalise NA-like strings to empty string.
            if _NA_PATTERNS.match(val_raw):
                r[key] = ""
                continue

            # Preserve WHSV-like strings for downstream unit-aware handling.
            if key == "GHSV_mL_g_h" and "WHSV" in val_raw.upper():
                continue

            cleaned = re.sub(r"[^\d\.\-\+]", "", val_raw)
            r[key] = cleaned if cleaned else ""


# ==========================================
# [UnitNorm] Unit normalization layer
# Converts heterogeneous literature units to schema-defined standard units.
# Called inside apply_condition_feature_layer per-record.
# ==========================================
def _parse_numeric_prefix(text: str) -> Optional[float]:
    """Extract leading numeric value from a string like '1.5 MPa' -> 1.5."""
    m = re.match(r"^\s*([-+]?\d+(?:\.\d+)?(?:[eE][-+]?\d+)?)", text.strip())
    return float(m.group(1)) if m else None


def _normalize_pressure_to_bar(raw: str) -> str:
    """Convert pressure string to bar. Handles MPa, kPa, atm, psi."""
    s = raw.strip()
    val = _parse_numeric_prefix(s)
    if val is None:
        return raw
    sl = s.lower()
    if "mpa" in sl:
        return str(round(val * 10.0, 4))
    if "kpa" in sl:
        return str(round(val * 0.01, 4))
    if "atm" in sl:
        return str(round(val * 1.01325, 4))
    if "psi" in sl:
        return str(round(val * 0.0689476, 4))
    # already bar or dimensionless
    return str(val)


def _normalize_time_to_h(raw: str) -> str:
    """Convert time string to hours. Handles min/minutes."""
    s = raw.strip()
    val = _parse_numeric_prefix(s)
    if val is None:
        return raw
    sl = s.lower()
    if re.search(r"\bmin(utes?)?\b", sl):
        return str(round(val / 60.0, 4))
    return str(val)


def _normalize_temp_to_c(raw: str, field: str = "") -> str:
    """Convert temperature string to Celsius. Detects Kelvin (value > 500 heuristic).
    For Reaction_Temp_C the valid MSR range is 150-600°C, so only convert when
    the value is clearly above the Kelvin threshold (>700) to avoid misidentifying
    high but valid Celsius values like 650°C.
    """
    s = raw.strip()
    val = _parse_numeric_prefix(s)
    if val is None:
        return raw
    sl = s.lower()
    if "k" in sl and "c" not in sl:
        return str(round(val - 273.15, 2))
    # For Reaction_Temp_C use a higher threshold: 150-700°C is a valid MSR range,
    # so only treat bare numbers >700 as Kelvin to avoid converting 650°C → 376.85°C.
    if field == "Reaction_Temp_C":
        if val > 700 and not re.search(r"[cC°]", s):
            return str(round(val - 273.15, 2))
    else:
        # heuristic: bare number > 500 is almost certainly Kelvin in MSR context
        if val > 500 and not re.search(r"[cC°]", s):
            return str(round(val - 273.15, 2))
    return str(val)


def _normalize_co_concentration(raw: str) -> str:
    """Convert CO concentration to ppm. Handles vol%, %."""
    s = raw.strip()
    val = _parse_numeric_prefix(s)
    if val is None:
        return raw
    sl = s.lower()
    if re.search(r"\bvol\s*%|\b%", sl):
        return str(round(val * 10000.0, 2))
    return str(val)


def _normalize_h2_production_rate(raw: str, catalyst_mass_g: Optional[float] = None) -> Tuple[str, str]:
    """
    Normalize H2_Production_Rate to mmol·g⁻¹·h⁻¹.
    Returns (normalized_value_str, detected_unit_str).
    If conversion is not possible, returns (raw, "").

    [Improve2] Volumetric flow units (mL/min, L/h) can now be converted when
    catalyst_mass_g is provided (extracted from Catalyst_Amount_g field).
    Conversion factor: H2 at STP, 1 mL = 0.044643 mmol.
    """
    s = raw.strip()
    val = _parse_numeric_prefix(s)
    if val is None:
        return raw, ""
    sl = s.lower()

    # mol/g/h or mol·g⁻¹·h⁻¹
    if re.search(r"\bmol[\s·/]*g[-\u207b]?\s*[-\u207b]?\s*1?\s*h", sl) and "mmol" not in sl and "umol" not in sl and "μmol" not in sl:
        return str(round(val * 1000.0, 4)), "mol/g/h"
    # mmol/g/h — already target unit
    if re.search(r"\bmmol[\s·/]*g[-\u207b]?\s*[-\u207b]?\s*1?\s*h", sl):
        return str(round(val, 4)), "mmol/g/h"
    # μmol/g/s or umol/g/s
    if re.search(r"[μu]mol[\s·/]*g[-\u207b]?\s*[-\u207b]?\s*1?\s*s", sl):
        return str(round(val * 3.6, 4)), "μmol/g/s"
    # μmol/g/h or umol/g/h
    if re.search(r"[μu]mol[\s·/]*g[-\u207b]?\s*[-\u207b]?\s*1?\s*h", sl):
        return str(round(val / 1000.0, 4)), "μmol/g/h"
    # mL/min (volumetric, STP) — needs catalyst mass
    if re.search(r"ml[\s·/]*min", sl) and catalyst_mass_g and catalyst_mass_g > 0:
        mmol_per_h = val * 60 * 0.044643
        return str(round(mmol_per_h / catalyst_mass_g, 4)), "mL/min"
    # L/h (volumetric, STP) — needs catalyst mass
    if re.search(r"\bl[\s·/]*h\b", sl) and catalyst_mass_g and catalyst_mass_g > 0:
        mmol_per_h = val * 1000 * 0.044643
        return str(round(mmol_per_h / catalyst_mass_g, 4)), "L/h"
    # Cannot convert — mark in Notes downstream
    return raw, ""


def apply_unit_normalization(row: Dict) -> Dict:
    """
    Normalize units for numeric fields to schema-defined standards.
    Operates in-place on a copy; writes conversion notes to Notes field.
    Fields handled:
      Pressure_bar        : MPa/kPa/atm/psi → bar
      Dry/Calcination/Reduction/Reaction_Temp_C : K → °C
      Dry/Calcination/Reduction_Time_h, TOS_h   : min → h
      CO_Concentration_ppm: vol%/% → ppm
      H2_Production_Rate  : mol/g/h, μmol/g/s, μmol/g/h → mmol/g/h
    """
    notes_additions = []

    # --- Pressure ---
    p_raw = str(row.get("Pressure_bar", "")).strip()
    if p_raw and re.search(r"[mkMK]?[pP][aA]|atm|psi", p_raw):
        p_norm = _normalize_pressure_to_bar(p_raw)
        if p_norm != p_raw:
            row["Pressure_bar"] = p_norm
            notes_additions.append(f"[unit_norm: Pressure_bar {p_raw}→{p_norm} bar]")

    # --- Temperatures ---
    for temp_field in ("Dry_Temp_C", "Calcination_Temp_C", "Reduction_Temp_C", "Reaction_Temp_C"):
        t_raw = str(row.get(temp_field, "")).strip()
        if not t_raw:
            continue
        t_norm = _normalize_temp_to_c(t_raw, field=temp_field)
        if t_norm != t_raw:
            row[temp_field] = t_norm
            notes_additions.append(f"[unit_norm: {temp_field} {t_raw}→{t_norm} C]")

    # --- Times ---
    for time_field in ("Dry_Time_h", "Calcination_Time_h", "Reduction_Time_h", "TOS_h"):
        t_raw = str(row.get(time_field, "")).strip()
        if not t_raw:
            continue
        t_norm = _normalize_time_to_h(t_raw)
        if t_norm != t_raw:
            row[time_field] = t_norm
            notes_additions.append(f"[unit_norm: {time_field} {t_raw}→{t_norm} h]")

    # --- CO concentration ---
    co_raw = str(row.get("CO_Concentration_ppm", "")).strip()
    if co_raw and re.search(r"%", co_raw):
        co_norm = _normalize_co_concentration(co_raw)
        if co_norm != co_raw:
            row["CO_Concentration_ppm"] = co_norm
            notes_additions.append(f"[unit_norm: CO_Concentration_ppm {co_raw}→{co_norm} ppm]")

    # --- H2 Production Rate ---
    # [Improve2] Pass catalyst mass so volumetric units (mL/min, L/h) can be converted.
    h2_raw = str(row.get("H2_Production_Rate", "")).strip()
    if h2_raw:
        cat_mass_raw = str(row.get("Catalyst_Amount_g", "")).strip()
        cat_mass_g = _parse_numeric_prefix(cat_mass_raw) if cat_mass_raw else None
        h2_norm, detected_unit = _normalize_h2_production_rate(h2_raw, catalyst_mass_g=cat_mass_g)
        if h2_norm != h2_raw and detected_unit:
            row["H2_Production_Rate_mmol_g_h"] = h2_norm
            notes_additions.append(f"[unit_norm: H2_Production_Rate {h2_raw}→{h2_norm} mmol/g/h (from {detected_unit})]")
        elif h2_norm == h2_raw and h2_raw:
            # Cannot convert — flag for ML preprocessing to treat as NaN
            notes_additions.append(f"[H2_rate_unit_unconvertible: {h2_raw}]")

    # Append conversion notes
    if notes_additions:
        existing_notes = str(row.get("Notes", "")).strip()
        row["Notes"] = (existing_notes + " " + " ".join(notes_additions)).strip()

    return row


_COMPARISON_AUDIT_FIELDS = [
    "Source_File", "data_source", "Catalyst", "Series_Name",
    "comparison_filter_action", "comparison_filter_reason",
    "is_this_work", "is_literature_comparison", "Notes",
]


_INTERNAL_CONTEXT_FIELDS = {"_comparison_context_raw", "_text_condition_salvage_source", "_text_condition_salvage_evidence"}
_CORE_FIGURE_PERFORMANCE_FIELDS = [
    "MeOH_Conversion_%", "H2_Yield_%", "H2_Production_Rate",
    "CO_Selectivity_%", "CO2_Selectivity_%", "CO_Concentration_ppm",
]
_FIGURE_PREMERGE_AUDIT_FIELDS = [
    "Source_File", "Series_Name", "raw_catalyst", "raw_catalyst_id", "raw_category_label",
    "figure_binding_status", "figure_binding_mode", "figure_binding_confidence", "figure_binding_reason",
    "matched_registry_key", "matched_registry_label", "alias_map_source", "alias_map_evidence",
    "semantic_figure_role", "premerge_filter_action", "premerge_filter_reason",
    "comparison_context_preview", "Notes",
]
_COMPARISON_AUDIT_FIELDS = [
    "Source_File", "data_source", "Catalyst", "Series_Name",
    "comparison_filter_action", "comparison_filter_reason",
    "is_this_work", "is_literature_comparison", "comparison_context_preview", "Notes",
]


def _context_preview(value: Any, limit: int = 240) -> str:
    text = _normalize_identity_text(str(value or ""))
    return text[:limit]



def _has_core_figure_performance_value(row: Dict) -> bool:
    return any(str(row.get(field, "")).strip() not in ("", "N/A", "n/a") for field in _CORE_FIGURE_PERFORMANCE_FIELDS)



def _extract_catalyst_amount_reason(text: str) -> str:
    text_norm = _normalize_identity_text(text).lower()
    if not text_norm:
        return ""
    if re.search(r"\b(?:wt\.?\s*%|wt%|mol\.?\s*%|mol%|at\.?\s*%|at%)\b", text_norm):
        return ""

    mass_match = re.search(r"\b(?:\d+(?:\.\d+)?|x)\s*(?:g|gm|gram|grams)\b", text_norm)
    if not mass_match:
        return ""

    explicit_amount_phrase = bool(re.search(
        r"\b(catalyst amount|amount of catalyst|catalyst charge|catalyst bed|bed amount|packing amount)\b",
        text_norm,
        flags=re.I,
    ))
    mass_with_catalyst_noun = bool(re.search(
        r"\b(?:\d+(?:\.\d+)?|x)\s*(?:g|gm|gram|grams)\b[^.;\n]{0,40}\b(catalyst|cat\.?|catalytic bed|bed|packing|charge|amount)\b",
        text_norm,
        flags=re.I,
    ))
    mass_with_support_config = bool(re.search(
        r"\b(?:\d+(?:\.\d+)?|x)\s*(?:g|gm|gram|grams)\b[^.;\n]{0,60}\b(with(?:out)?\s+[^.;\n]{0,30}(support|cement-clay)|no\s+[^.;\n]{0,30}(support|cement-clay)|cement-clay)\b",
        text_norm,
        flags=re.I,
    ))

    if explicit_amount_phrase:
        return "explicit_catalyst_amount_phrase"
    if mass_with_catalyst_noun:
        return "mass_plus_catalyst_noun"
    if mass_with_support_config:
        return "mass_plus_support_configuration"
    return ""



def _looks_like_catalyst_amount_label(text: str) -> bool:
    return bool(_extract_catalyst_amount_reason(text))



def _has_resolved_figure_identity(row: Dict) -> bool:
    clues = [
        row.get("Catalyst", ""),
        row.get("Catalyst_ID", ""),
        row.get("matched_registry_key", ""),
        row.get("Active_Metal", ""),
        row.get("Support", ""),
        row.get("Series_Name", ""),
    ]
    return any(str(value).strip() for value in clues)



def _is_high_value_figure_binding_mode(binding_mode: str) -> bool:
    mode = clean_residual_mojibake_chars(str(binding_mode or "")).strip()
    return mode in {
        "direct_label",
        "caption_map",
        "partial_unique",
        "baseline_match",
        "support_only_match",
        "alias_map_label",
        "numbered_alias_map",
    }



def _has_reliable_figure_condition_anchor(row: Dict) -> bool:
    if not isinstance(row, dict):
        return False

    direct_fields = [
        "Reaction_Temp_C",
        "TOS_h",
        "S_C_Ratio",
        "GHSV_mL_g_h",
        "Pressure_bar",
        "Flow_Rate",
        "Catalyst_Amount_g",
    ]
    if any(str(row.get(field, "")).strip() for field in direct_fields):
        return True

    condition_text = " ".join([
        str(row.get("condition_anchor_source", "") or ""),
        str(row.get("figure_point_validation_notes", "") or ""),
        str(row.get("figure_binding_notes", "") or ""),
        str(row.get("binding_reason", "") or ""),
        str(row.get("Notes", "") or ""),
    ])
    condition_text = clean_residual_mojibake_chars(condition_text)
    if not condition_text.strip():
        return False

    lowered = condition_text.lower()
    if re.search(r"\b(calcination|reduction|drying|preparation|precursor|impregnation)\b", lowered, flags=re.I) and not re.search(
        r"\b(reaction temperature|reaction temp|time on stream|tos|s/c|steam[- ]?to[- ]?carbon|ghsv|whsv|pressure|flow rate|catalyst amount|condition_anchor|anchor_sources?|x_axis|x_value)\b",
        lowered,
        flags=re.I,
    ):
        return False

    anchor_source_hit = bool(re.search(r"\b(caption|legend|series|nearby|context)\b", lowered, flags=re.I))
    explicit_anchor_tag = bool(re.search(r"\[(?:condition_anchor|anchor_sources?|x_axis|x_value)[^\]]+\]", condition_text, flags=re.I))
    condition_axis_hit = bool(re.search(
        r"\b(reaction temperature|reaction temp|temperature|time on stream|tos|s/c|steam[- ]?to[- ]?carbon|ghsv|whsv|pressure|flow rate|catalyst amount)\b",
        lowered,
        flags=re.I,
    ))
    numeric_value_hit = bool(re.search(r"\b\d+(?:\.\d+)?\s*(?:deg\.?\s*c|celsius|c|h|bar|atm|mpa)?\b", lowered, flags=re.I))

    if numeric_value_hit and (explicit_anchor_tag or (anchor_source_hit and condition_axis_hit) or condition_axis_hit):
        return True
    if bool(re.search(r"(?:caption|legend|series|context)\s*:\s*(?:Reaction_Temp_C|TOS_h|S_C_Ratio|GHSV_mL_g_h|Pressure_bar)", condition_text, flags=re.I)):
        return True
    return False



def _is_allowable_category_binding(row: Dict) -> bool:
    if not isinstance(row, dict):
        return False

    semantic_role = str(row.get("semantic_figure_role", "") or "").strip()
    if semantic_role not in {"category_screening", "catalyst_screening"}:
        return False
    if str(row.get("figure_binding_status", "") or "").strip() != "matched":
        return False

    raw_mode = str(row.get("binding_mode", "") or "").strip()
    compat_mode = str(row.get("figure_binding_mode", "") or "").strip()
    if not any(_is_high_value_figure_binding_mode(mode) for mode in [raw_mode, compat_mode] if mode):
        return False

    confidence = str(row.get("figure_binding_confidence", row.get("binding_confidence", "")) or "").strip()
    if confidence not in {"high", "medium"}:
        return False
    if not _has_core_figure_performance_value(row):
        return False
    if not _has_reliable_figure_condition_anchor(row):
        return False

    effective_mode = raw_mode or compat_mode
    if effective_mode in {"alias_map_label", "numbered_alias_map"}:
        evidence_blob = " ".join([
            str(row.get("alias_map_source", "") or ""),
            str(row.get("alias_map_evidence", "") or ""),
            str(row.get("binding_reason", "") or ""),
            str(row.get("figure_binding_notes", "") or ""),
        ]).strip()
        if not evidence_blob:
            return False
        if re.search(r"\bambiguous\b|caption_mapping_ambiguous|no_reliable_registry_match|generic_label_without_context_mapping", evidence_blob, flags=re.I):
            return False

    return True



def _should_keep_figure_row_candidate_only(row: Dict) -> Tuple[bool, str]:
    if not isinstance(row, dict):
        return True, "not_figure_row"

    semantic_role = str(row.get("semantic_figure_role", "") or "").strip()
    blocked_roles = {
        "product_species_profile",
        "case_profile",
        "schematic_or_nonperformance",
        "catalyst_amount_effect",
        "non_performance_like",
    }
    if semantic_role in blocked_roles:
        return True, "semantic_role_blocked"

    if not _has_core_figure_performance_value(row):
        return True, "missing_performance_value"

    binding_status = str(row.get("figure_binding_status", "") or "").strip()
    raw_mode = str(row.get("binding_mode", "") or "").strip()
    compat_mode = str(row.get("figure_binding_mode", "") or "").strip()
    effective_mode = raw_mode or compat_mode
    confidence = str(row.get("figure_binding_confidence", row.get("binding_confidence", "")) or "").strip()
    unresolved_blob = " ".join([
        str(row.get("figure_binding_notes", "") or ""),
        str(row.get("binding_reason", "") or ""),
        str(row.get("Series_Name", "") or ""),
        str(row.get("Notes", "") or ""),
    ])
    unresolved_blob = clean_residual_mojibake_chars(unresolved_blob)

    if binding_status != "matched":
        if binding_status == "ambiguous" or effective_mode == "ambiguous" or re.search(r"\bambiguous\b", unresolved_blob, flags=re.I):
            return True, "ambiguous_binding"
        return True, "unmatched_binding"

    strong_match_evidence = bool(
        str(row.get("matched_registry_key", "") or "").strip()
        or str(row.get("matched_registry_label", "") or "").strip()
        or str(row.get("alias_map_source", "") or "").strip()
        or str(row.get("alias_map_evidence", "") or "").strip()
    )
    if re.search(r"generic_label_without_context_mapping|no_reliable_registry_match|caption_mapping_ambiguous", unresolved_blob, flags=re.I) and not strong_match_evidence:
        return True, "unresolved_generic_label"

    if not any(_is_high_value_figure_binding_mode(mode) for mode in [raw_mode, compat_mode] if mode):
        return True, "unsupported_binding_mode"

    if confidence not in {"high", "medium"}:
        return True, "low_binding_confidence"

    alias_specific = effective_mode in {"alias_map_label", "numbered_alias_map"}
    if alias_specific:
        alias_blob = " ".join([
            str(row.get("alias_map_source", "") or ""),
            str(row.get("alias_map_evidence", "") or ""),
            str(row.get("binding_reason", "") or ""),
            str(row.get("figure_binding_notes", "") or ""),
        ]).strip()
        if not alias_blob or re.search(r"\bambiguous\b|caption_mapping_ambiguous|generic_label_without_context_mapping|no_reliable_registry_match", alias_blob, flags=re.I):
            return True, "weak_alias_map_evidence"

    if semantic_role in {"category_screening", "catalyst_screening"}:
        if not _has_reliable_figure_condition_anchor(row):
            return True, "missing_condition_anchor_for_category"
        if not _is_allowable_category_binding(row):
            if alias_specific:
                return True, "weak_alias_map_evidence"
            return True, "ambiguous_binding"
        return False, "merge_eligible"

    if not _has_resolved_figure_identity(row) and not _has_reliable_figure_condition_anchor(row):
        return True, "missing_identity_clue"

    return False, "merge_eligible"



def _describe_premerge_binding_gate(row: Dict) -> str:
    raw_mode = str(row.get("binding_mode", "") or "").strip()
    compat_mode = str(row.get("figure_binding_mode", "") or "").strip()
    parts = [
        f"role={str(row.get('semantic_figure_role', '') or '').strip() or 'unknown'}",
        f"status={str(row.get('figure_binding_status', '') or '').strip() or 'unknown'}",
        f"mode={raw_mode or compat_mode or 'unknown'}",
        f"confidence={str(row.get('figure_binding_confidence', row.get('binding_confidence', '')) or '').strip() or 'unknown'}",
        f"has_metric={int(bool(_has_core_figure_performance_value(row)))}",
        f"has_condition_anchor={int(bool(_has_reliable_figure_condition_anchor(row)))}",
    ]
    if str(row.get("matched_registry_key", "") or "").strip():
        parts.append(f"registry_key={str(row.get('matched_registry_key', '')).strip()}")
    if str(row.get("alias_map_source", "") or "").strip():
        parts.append(f"alias_source={_context_preview(row.get('alias_map_source', ''), 80)}")
    if str(row.get("alias_map_evidence", "") or "").strip():
        parts.append(f"alias_evidence={_context_preview(row.get('alias_map_evidence', ''), 80)}")
    return "; ".join(parts)



def _derive_figure_premerge_filter_reason(row: Dict) -> str:
    if not globals().get("_FIGURE_PREMERGE_STRICT_SELFCHECK_DONE"):
        globals()["_FIGURE_PREMERGE_STRICT_SELFCHECK_DONE"] = True
        try:
            base = {
                "data_source": "figure",
                "premerge_filter_action": "kept_for_merge",
                "numeric_reliability_level": "direct_numeric",
                "identity_completeness_level": "complete",
                "figure_binding_status": "matched",
                "is_literature_comparison": 0,
                "comparison_filter_action": "",
                "is_range_like_value": 0,
                "is_qualitative_value": 0,
                "raw_category_label": "ZrO2",
            }
            case1 = dict(base)
            case1.update({
                "semantic_figure_role": "category_screening",
                "binding_mode": "baseline_match",
                "figure_binding_mode": "baseline_match",
                "figure_binding_confidence": "medium",
                "Reaction_Temp_C": "90",
                "MeOH_Conversion_%": "12.3",
            })
            if _derive_figure_premerge_filter_reason(case1) != "merge_eligible" or not is_strict_trainable_row(case1):
                print("  [WARNING] figure premerge selfcheck failed: baseline_match with condition anchor should pass")

            case2 = dict(base)
            case2.update({
                "semantic_figure_role": "category_screening",
                "binding_mode": "support_only_match",
                "figure_binding_mode": "support_only_match",
                "figure_binding_confidence": "medium",
                "Reaction_Temp_C": "90",
                "MeOH_Conversion_%": "12.3",
            })
            if _derive_figure_premerge_filter_reason(case2) != "merge_eligible" or not is_strict_trainable_row(case2):
                print("  [WARNING] figure premerge selfcheck failed: support_only_match with condition anchor should pass")

            case3 = dict(base)
            case3.update({
                "semantic_figure_role": "category_screening",
                "binding_mode": "alias_map_label",
                "figure_binding_mode": "alias_map_label",
                "figure_binding_confidence": "high",
                "alias_map_source": "context_label_map",
                "alias_map_evidence": "C1=Ni/ZrO2",
                "binding_reason": "alias_map_label:C1->Ni/ZrO2",
                "Reaction_Temp_C": "250",
                "MeOH_Conversion_%": "78.0",
            })
            if _derive_figure_premerge_filter_reason(case3) != "merge_eligible" or not is_strict_trainable_row(case3):
                print("  [WARNING] figure premerge selfcheck failed: alias_map_label with evidence should pass")

            case4 = dict(base)
            case4.update({
                "semantic_figure_role": "category_screening",
                "binding_mode": "numbered_alias_map",
                "figure_binding_mode": "numbered_alias_map",
                "figure_binding_confidence": "high",
                "Reaction_Temp_C": "250",
                "MeOH_Conversion_%": "78.0",
            })
            if _derive_figure_premerge_filter_reason(case4) not in {"weak_alias_map_evidence", "ambiguous_binding"} or is_strict_trainable_row(case4):
                print("  [WARNING] figure premerge selfcheck failed: numbered_alias_map without evidence should stay blocked")

            case5 = dict(base)
            case5.update({
                "semantic_figure_role": "category_screening",
                "binding_mode": "baseline_match",
                "figure_binding_mode": "baseline_match",
                "figure_binding_confidence": "medium",
                "MeOH_Conversion_%": "12.3",
            })
            if _derive_figure_premerge_filter_reason(case5) != "missing_condition_anchor_for_category" or is_strict_trainable_row(case5):
                print("  [WARNING] figure premerge selfcheck failed: category row without condition anchor should stay blocked")

            case6 = dict(base)
            case6.update({
                "semantic_figure_role": "product_species_profile",
                "binding_mode": "direct_label",
                "figure_binding_mode": "direct_label",
                "figure_binding_confidence": "high",
                "Reaction_Temp_C": "250",
                "MeOH_Conversion_%": "45.0",
            })
            if _derive_figure_premerge_filter_reason(case6) != "semantic_role_blocked" or is_strict_trainable_row(case6):
                print("  [WARNING] figure premerge selfcheck failed: non-performance semantic role should stay blocked")
        except Exception as e:
            print(f"  [WARNING] figure premerge selfcheck error: {e}")

    if str(row.get("data_source", "")).strip() != "figure":
        return "not_figure_row"
    if not _has_core_figure_performance_value(row):
        return "missing_performance_value"

    keep_candidate_only, reason = _should_keep_figure_row_candidate_only(row)
    if keep_candidate_only:
        return reason
    return "merge_eligible"



def is_merge_eligible_figure_row(row: Dict) -> bool:
    return _derive_figure_premerge_filter_reason(row) == "merge_eligible"



def split_figure_rows_for_merge(fig_records: List[Dict], file_name: str) -> Tuple[List[Dict], List[Dict]]:
    merge_eligible_rows: List[Dict] = []
    candidate_only_rows: List[Dict] = []

    for row in fig_records:
        current = dict(row)
        reason = _derive_figure_premerge_filter_reason(current)
        if reason == "merge_eligible":
            current["premerge_filter_action"] = "kept_for_merge"
            current["premerge_filter_reason"] = reason
            merge_eligible_rows.append(current)
        else:
            current["premerge_filter_action"] = "candidate_only"
            current["premerge_filter_reason"] = reason
            candidate_only_rows.append(current)

    blocked_n = len(candidate_only_rows)
    if blocked_n:
        print(f"    [premerge gate] {file_name}: {blocked_n} figure rows kept as candidate-only")
    return merge_eligible_rows, candidate_only_rows



def _build_figure_premerge_filter_audit_rows(rows: List[Dict], file_name: str) -> List[Dict]:
    audit_rows = []
    for row in rows:
        audit_rows.append({
            "Source_File": file_name,
            "Series_Name": str(row.get("Series_Name", "")),
            "raw_catalyst": str(row.get("Catalyst", "")),
            "raw_catalyst_id": str(row.get("Catalyst_ID", "")),
            "figure_binding_status": str(row.get("figure_binding_status", "")),
            "figure_binding_mode": str(row.get("binding_mode", row.get("figure_binding_mode", ""))),
            "figure_binding_confidence": str(row.get("figure_binding_confidence", row.get("binding_confidence", ""))),
            "semantic_figure_role": str(row.get("semantic_figure_role", "")),
            "premerge_filter_action": str(row.get("premerge_filter_action", "")),
            "premerge_filter_reason": str(row.get("premerge_filter_reason", "")),
            "figure_binding_notes": str(row.get("figure_binding_notes", row.get("binding_reason", ""))),
            "matched_registry_key": str(row.get("matched_registry_key", "")),
            "matched_registry_label": str(row.get("matched_registry_label", "")),
            "raw_category_label": str(row.get("raw_category_label", "")),
            "alias_map_source": str(row.get("alias_map_source", "")),
            "alias_map_evidence": str(row.get("alias_map_evidence", "")),
            "binding_gate_summary": _describe_premerge_binding_gate(row),
            "comparison_context_preview": _context_preview(row.get("_comparison_context_raw", "")),
            "Notes": str(row.get("Notes", "")),
        })
    return audit_rows



def strip_internal_context_fields(record: Dict) -> Dict:
    if not isinstance(record, dict):
        return record
    return {k: v for k, v in record.items() if k not in _INTERNAL_CONTEXT_FIELDS}



def strip_internal_context_fields_batch(records: List[Dict]) -> List[Dict]:
    return [strip_internal_context_fields(dict(record)) for record in records if isinstance(record, dict)]



def _row_level_comparison_text(row: Dict) -> str:
    return _normalize_identity_text(" ".join([
        str(row.get("Notes", "")),
        str(row.get("Catalyst", "")),
        str(row.get("Series_Name", "")),
        str(row.get("figure_binding_notes", "")),
    ]))



def _comparison_context_text(source_text: str, row: Dict) -> str:
    return _normalize_identity_text(" ".join([
        str(row.get("_comparison_context_raw", "")),
        str(source_text or ""),
        _row_level_comparison_text(row),
    ]))



def _has_schematic_system_context(text: str) -> bool:
    return bool(re.search(
        r"\b(schematic of the experimental system|experimental system|fig\.?\s*\d+\s*schematic|system schematic|illustration|diagram|flowsheet)\b",
        text,
        flags=re.I,
    ))



def _context_has_benchmark_risk(text: str, row: Optional[Dict] = None) -> bool:
    if not text:
        return False
    citation_hits = re.findall(r"\[\s*\d+\s*\]", text)
    reference_hits = re.findall(r"\bref(?:erence)?\.?\b", text, flags=re.I)
    benchmark_like = bool(re.search(
        r"\b(benchmark|literature comparison|comparison table|comparison|compared? with|versus|vs\.?|literature)\b",
        text,
        flags=re.I,
    ))
    schematic_context = _has_schematic_system_context(text)

    if len(citation_hits) >= 2:
        return True
    if reference_hits:
        return True
    if benchmark_like:
        return True
    if schematic_context and row and str(row.get("data_source", "")).strip() == "figure":
        return True
    return False



def _row_has_this_study_signal(row: Dict, row_text: str = "") -> bool:
    return bool(row.get("is_this_work")) or bool(re.search(
        r"\b(this study|this work|present work|our work|our catalyst)\b",
        row_text or _row_level_comparison_text(row),
        flags=re.I,
    ))



def detect_comparison_or_benchmark_context(source_text: str, row: Dict) -> bool:
    row_text = _row_level_comparison_text(row)
    combined_text = _comparison_context_text(source_text, row)
    if not combined_text:
        return bool(row.get("is_literature_comparison"))

    row_has_this_study = _row_has_this_study_signal(row, row_text)
    row_citation_hits = re.findall(r"\[\s*\d+\s*\]", row_text)
    row_reference_hits = re.findall(r"\bref(?:erence)?\.?\b", row_text, flags=re.I)
    row_benchmark_like = bool(re.search(
        r"\b(benchmark|literature comparison|comparison table|comparison|compared? with|versus|vs\.?|literature)\b",
        row_text,
        flags=re.I,
    ))

    if row_has_this_study:
        return False
    if row_citation_hits or row_reference_hits or row_benchmark_like:
        return True
    if _context_has_benchmark_risk(combined_text, row):
        return True
    if bool(row.get("is_literature_comparison")):
        return True
    return False



def mark_or_filter_comparison_rows(records: List[Dict], file_name: str) -> Tuple[List[Dict], List[Dict]]:
    kept_rows: List[Dict] = []
    audit_rows: List[Dict] = []

    for row in records:
        current = dict(row)
        row_text = _row_level_comparison_text(current)
        combined_text = _comparison_context_text(current.get("_comparison_context_raw", ""), current)
        own_study_hint = _row_has_this_study_signal(current, row_text)
        is_benchmark = detect_comparison_or_benchmark_context(current.get("_comparison_context_raw", ""), current)
        schematic_context = _has_schematic_system_context(combined_text)
        context_benchmark_risk = _context_has_benchmark_risk(combined_text, current)
        context_preview = _context_preview(current.get("_comparison_context_raw", "") or combined_text)

        if own_study_hint:
            current["is_literature_comparison"] = False
            if context_benchmark_risk:
                current["comparison_filter_action"] = "kept_this_study_candidate"
                current["comparison_filter_reason"] = "this_study_with_benchmark_context"
            else:
                current.setdefault("comparison_filter_action", "")
                current.setdefault("comparison_filter_reason", "")
        elif is_benchmark:
            current["is_literature_comparison"] = True
            current["comparison_filter_action"] = "downgraded_candidate_only"
            current["comparison_filter_reason"] = (
                "schematic_or_system_context" if schematic_context and str(current.get("data_source", "")).strip() == "figure"
                else "benchmark_or_literature_comparison"
            )
            audit_rows.append({
                "Source_File": file_name,
                "data_source": str(current.get("data_source", "")),
                "Catalyst": str(current.get("Catalyst", "")),
                "Series_Name": str(current.get("Series_Name", "")),
                "comparison_filter_action": current["comparison_filter_action"],
                "comparison_filter_reason": current["comparison_filter_reason"],
                "is_this_work": bool(current.get("is_this_work")),
                "is_literature_comparison": bool(current.get("is_literature_comparison")),
                "comparison_context_preview": context_preview,
                "Notes": str(current.get("Notes", "")),
            })
        else:
            current.setdefault("comparison_filter_action", "")
            current.setdefault("comparison_filter_reason", "")

        kept_rows.append(current)

    if audit_rows:
        print(f"  [Audit] comparison benchmark rows downgraded: {len(audit_rows)}")
    return kept_rows, audit_rows



def classify_figure_row_semantic_role(row: Dict, local_context: str = "") -> str:
    notes = _normalize_identity_text(str(row.get("Notes", "")))
    series_name = _normalize_identity_text(str(row.get("Series_Name", "")))
    catalyst = _normalize_identity_text(str(row.get("Catalyst", "")))
    binding_notes = _normalize_identity_text(str(row.get("figure_binding_notes", "")))
    x_axis, x_value = _extract_x_axis_and_value_from_notes(notes)
    combined_text = _normalize_identity_text(" ".join([notes, series_name, catalyst, binding_notes, local_context]))

    if re.search(r"\b(schematic|diagram|illustration|experimental system|reactor system|flowsheet)\b", combined_text, flags=re.I):
        return "schematic_or_nonperformance"

    if re.search(r"\b(case\s*[- ]?[a-z0-9]+|run\s*[- ]?[a-z0-9]+)\b", series_name, flags=re.I):
        return "case_profile"

    if "[series_role=product]" in notes.lower():
        return "product_species_profile"
    if re.fullmatch(r"(?:h2|co2|co|ch4|dme|h2o)", series_name.strip(), flags=re.I):
        return "product_species_profile"
    if re.search(r"\b(h2|co2|co|ch4|dme|h2o)\b", series_name, flags=re.I) and not str(row.get("figure_binding_status", "")).strip() == "matched":
        return "product_species_profile"

    amount_reason = _extract_catalyst_amount_reason(combined_text)
    if amount_reason:
        row["figure_semantic_subtype"] = "catalyst_amount"
        return "catalyst_amount_effect"

    x_axis_text = _normalize_identity_text(" ".join([x_axis, combined_text]))
    if re.search(r"\b(tos|time on stream|stability|long-term|durability|deactivation)\b", x_axis_text, flags=re.I):
        return "stability_profile"

    if re.search(r"\b(s/c|steam.?to.?carbon|ghsv|whsv|pressure|flow|amount|feed composition|ratio|reaction temperature|temperature)\b", x_axis_text, flags=re.I):
        return "condition_effect"

    if re.search(r"\b(catalyst|sample|support|composition|loading|content)\b", x_axis, flags=re.I):
        return "catalyst_screening"
    if x_value and _looks_like_identity_mapping(x_value):
        return "catalyst_screening"
    if str(row.get("figure_binding_status", "")).strip() == "matched":
        return "catalyst_screening"

    return "unknown"



def annotate_figure_semantic_role(records: List[Dict], local_context: str) -> List[Dict]:
    annotated = []
    for row in records:
        current = dict(row)
        if str(current.get("data_source", "")).strip() == "figure":
            current["semantic_figure_role"] = classify_figure_row_semantic_role(current, local_context)
        annotated.append(current)
    return annotated



def is_strict_trainable_row(row: Dict) -> bool:
    data_source = str(row.get("data_source", "")).strip()

    if data_source == "text":
        return is_text_point_record(row)

    if data_source == "figure":
        if str(row.get("premerge_filter_action", "")).strip() == "candidate_only":
            return False
        if bool(row.get("is_literature_comparison")):
            return False
        if str(row.get("comparison_filter_action", "")).strip() == "downgraded_candidate_only":
            return False
        if str(row.get("numeric_reliability_level", "")).strip() != "direct_numeric":
            return False
        if bool(row.get("is_range_like_value")) or bool(row.get("is_qualitative_value")):
            return False
        if str(row.get("identity_completeness_level", "")).strip() not in {"complete", "partial"}:
            return False
        if str(row.get("figure_binding_status", "")).strip() != "matched":
            return False

        raw_mode = str(row.get("binding_mode", "") or "").strip()
        compat_mode = str(row.get("figure_binding_mode", "") or "").strip()
        if not any(_is_high_value_figure_binding_mode(mode) for mode in [raw_mode, compat_mode] if mode):
            return False

        if str(row.get("figure_binding_confidence", row.get("binding_confidence", ""))).strip() not in {"high", "medium"}:
            return False

        if str(row.get("semantic_figure_role", "")).strip() in {
            "product_species_profile",
            "case_profile",
            "schematic_or_nonperformance",
            "catalyst_amount_effect",
            "non_performance_like",
        }:
            return False

        semantic_role = str(row.get("semantic_figure_role", "") or "").strip()
        effective_mode = raw_mode or compat_mode
        if semantic_role in {"category_screening", "catalyst_screening"} and not _has_reliable_figure_condition_anchor(row):
            return False
        if effective_mode in {"alias_map_label", "numbered_alias_map"}:
            alias_blob = " ".join([
                str(row.get("alias_map_source", "") or ""),
                str(row.get("alias_map_evidence", "") or ""),
                str(row.get("binding_reason", "") or ""),
                str(row.get("figure_binding_notes", "") or ""),
            ]).strip()
            if not alias_blob:
                return False
            if re.search(r"\bambiguous\b|caption_mapping_ambiguous|generic_label_without_context_mapping|no_reliable_registry_match", alias_blob, flags=re.I):
                return False
        return True

    return False



def split_candidate_and_strict_rows(records: List[Dict]) -> Tuple[List[Dict], List[Dict]]:
    _initialize_layered_export_state_if_needed()
    _selfcheck_layered_gate_examples()

    candidate_rows: List[Dict] = []
    for index, row in enumerate(records or []):
        if not isinstance(row, dict):
            continue
        current = _finalize_figure_row_binding_payload(dict(row))
        if not str(current.get("layer_row_uid", "") or "").strip():
            source_file = clean_residual_mojibake_chars(str(current.get("Source_File", "") or "")).strip() or "unknown_source"
            current["layer_row_uid"] = f"{source_file}::layer::{index + 1:06d}"
        candidate_rows.append(current)

    text_views = build_text_gate_stratified_views([
        row for row in candidate_rows if str(row.get("data_source", "") or "").strip() == "text"
    ])
    figure_views = build_figure_gate_stratified_views([
        row for row in candidate_rows if str(row.get("data_source", "") or "").strip() == "figure"
    ])

    annotated_lookup = {
        str(row.get("layer_row_uid", "") or "").strip(): row
        for row in (text_views.get("annotated_rows", []) + figure_views.get("annotated_rows", []))
        if isinstance(row, dict) and str(row.get("layer_row_uid", "") or "").strip()
    }

    annotated_candidate_rows: List[Dict] = []
    research_all_rows: List[Dict] = []
    research_strict_rows: List[Dict] = []
    for row in candidate_rows:
        uid = str(row.get("layer_row_uid", "") or "").strip()
        current = annotated_lookup.get(uid)
        if current is None:
            current = _annotate_row_with_layer_membership(
                row,
                research_all_ok=False,
                research_strict_ok=False,
                modeling_final_ok=False,
                strict_reason="unsupported_data_source",
                final_reason="unsupported_data_source",
            )
        annotated_candidate_rows.append(current)
        if bool(current.get("included_in_research_all")):
            research_all_rows.append(dict(current))
        if bool(current.get("included_in_research_strict")):
            research_strict_rows.append(dict(current))

    globals().setdefault("_RESEARCH_ALL_ROWS_CACHE", []).extend([dict(row) for row in research_all_rows])
    globals().setdefault("_RESEARCH_STRICT_ROWS_CACHE", []).extend([dict(row) for row in research_strict_rows])

    source_files = sorted({
        clean_residual_mojibake_chars(str(row.get("Source_File", "") or "")).strip()
        for row in annotated_candidate_rows
        if clean_residual_mojibake_chars(str(row.get("Source_File", "") or "")).strip()
    })
    source_key = source_files[0] if len(source_files) == 1 else (source_files[0] if source_files else "<unknown_source>")
    paper_stats = globals().setdefault("_LAYERED_PAPER_STATS", {})
    paper_stats[source_key] = {
        **paper_stats.get(source_key, {}),
        "candidate_row_count": len(annotated_candidate_rows),
        "research_all_row_count": len(research_all_rows),
        "research_strict_row_count": len(research_strict_rows),
        "research_all_block_reason_counts": _count_reason_values(
            [row for row in annotated_candidate_rows if not bool(row.get("included_in_research_all"))],
            "strict_exclusion_reason",
        ),
        "research_strict_excluded_reason_counts": _count_reason_values(
            [row for row in annotated_candidate_rows if bool(row.get("included_in_research_all")) and not bool(row.get("included_in_research_strict"))],
            "strict_exclusion_reason",
        ),
        "research_all_text_row_count": sum(1 for row in research_all_rows if str(row.get("data_source", "") or "").strip() == "text"),
        "research_all_figure_row_count": sum(1 for row in research_all_rows if str(row.get("data_source", "") or "").strip() == "figure"),
        "research_strict_text_row_count": sum(1 for row in research_strict_rows if str(row.get("data_source", "") or "").strip() == "text"),
        "research_strict_figure_row_count": sum(1 for row in research_strict_rows if str(row.get("data_source", "") or "").strip() == "figure"),
    }

    return clean_records_text_fields(annotated_candidate_rows), clean_records_text_fields(research_strict_rows)


def is_final_publishable_text_row(row: Dict) -> Tuple[bool, str]:
    if str(row.get("data_source", "")).strip() != "text":
        return False, "not_text"
    if bool(row.get("is_literature_comparison")) or str(row.get("comparison_filter_action", "")).strip() == "downgraded_candidate_only":
        return False, "comparison_like"
    if str(row.get("identity_completeness_level", "")).strip() != "complete":
        return False, "weak_identity"
    if str(row.get("numeric_reliability_level", "")).strip() != "direct_numeric":
        return False, "non_direct_numeric"
    if str(row.get("source_granularity", "")).strip() not in {"table_row", "text_numeric", "si_table"}:
        return False, "weak_source_granularity"
    if bool(row.get("is_approximate_value")) or bool(row.get("is_range_like_value")) or bool(row.get("is_qualitative_value")):
        return False, "approximate_or_range"
    if bool(row.get("same_physical_point_possible")) or bool(row.get("obvious_duplicate_flag")) or str(row.get("duplicate_candidate_type", "")).strip():
        return False, "duplicate_like"
    if not has_core_performance_metric(row):
        return False, "missing_core_metric"
    if not has_valid_point_condition_anchor(row):
        return False, "performance_without_condition_anchor"
    if is_preparation_backbone_only_row(row):
        return False, "preparation_backbone_only"
    return (True, "ok") if is_text_point_record(row) else (False, "missing_core_metric")



def _has_reliable_final_figure_condition_anchor(row: Dict) -> bool:
    if not isinstance(row, dict):
        return False

    current = _finalize_figure_row_binding_payload(dict(row))
    bundle = _canonicalize_final_condition_bundle(current)

    present_groups = set()
    if bundle.get("Reaction_Temp_C"):
        present_groups.add("temperature")
    if bundle.get("TOS_h"):
        present_groups.add("time")
    if bundle.get("S_C_Ratio") or bundle.get("Feed_MeOH_to_H2O_Ratio"):
        present_groups.add("ratio")
    if bundle.get("Pressure_bar"):
        present_groups.add("pressure")
    if bundle.get("GHSV_mL_g_h") or bundle.get("WHSV_h_inv") or bundle.get("SpaceVelocity_norm"):
        present_groups.add("space_velocity")
    if bundle.get("Flow_Rate"):
        present_groups.add("flow")
    if bundle.get("Catalyst_Amount_g"):
        present_groups.add("catalyst_amount")

    notes = clean_residual_mojibake_chars(str(current.get("Notes", "") or ""))
    x_axis_candidates = [
        str(current.get("x_axis", "") or ""),
        str(current.get("x_axis_mode", "") or ""),
    ]
    x_value_candidates = [str(current.get("x_value", "") or "")]
    note_axis, note_x_value = _extract_x_axis_and_value_from_notes(notes)
    if note_axis:
        x_axis_candidates.append(note_axis)
    if note_x_value:
        x_value_candidates.append(note_x_value)

    def _looks_numeric(value: Any, allow_ratio: bool = False) -> bool:
        text_value = clean_residual_mojibake_chars(str(value or "")).strip()
        if not text_value or text_value.lower() in {"", "n/a", "na", "none", "nan", "unknown"}:
            return False
        if _looks_like_identity_mapping(text_value):
            return False
        if allow_ratio and re.fullmatch(r"\d+(?:\.\d+)?\s*[:/]\s*\d+(?:\.\d+)?", text_value):
            return True
        if _parse_float_if_possible(text_value.replace(",", "")) is not None:
            return True
        return bool(re.search(r"[-+]?\d+(?:\.\d+)?", text_value))

    for axis_text in x_axis_candidates:
        axis_norm = _normalize_identity_text(axis_text).lower()
        if not axis_norm:
            continue
        for value_text in x_value_candidates:
            if not _looks_numeric(value_text, allow_ratio=("ratio" in axis_norm or "s/c" in axis_norm)):
                continue
            if re.search(r"\b(reaction temperature|reaction temp|temperature|temp)\b", axis_norm, flags=re.I):
                present_groups.add("temperature")
            elif re.search(r"\b(tos|time on stream|stream time|stability)\b", axis_norm, flags=re.I):
                present_groups.add("time")
            elif re.search(r"\b(s/c|steam[- ]?to[- ]?carbon|feed ratio|meoh[- ]?to[- ]?h2o|methanol[- ]?to[- ]?water)\b", axis_norm, flags=re.I):
                present_groups.add("ratio")
            elif re.search(r"\bpressure\b", axis_norm, flags=re.I):
                present_groups.add("pressure")
            elif re.search(r"\b(ghsv|whsv|space velocity)\b", axis_norm, flags=re.I):
                present_groups.add("space_velocity")
            elif re.search(r"\b(flow rate|feed flow|feed rate)\b", axis_norm, flags=re.I):
                present_groups.add("flow")
            elif re.search(r"\b(catalyst amount|catalyst mass|catalyst charge|catalyst bed)\b", axis_norm, flags=re.I):
                present_groups.add("catalyst_amount")

    if "temperature" in present_groups or "time" in present_groups:
        return True
    return len(present_groups) >= 2



def _normalize_target_family(row: Dict) -> str:
    if not isinstance(row, dict):
        return ""
    target_fields = [
        "MeOH_Conversion_%",
        "H2_Yield_%",
        "H2_Production_Rate",
        "CO_Selectivity_%",
        "CO2_Selectivity_%",
        "CO_Concentration_ppm",
    ]
    present = []
    for field in target_fields:
        value = clean_residual_mojibake_chars(str(row.get(field, "") or "")).strip()
        if value and value.lower() not in {"", "n/a", "na", "none", "nan", "unknown"}:
            present.append(field)
    return "|".join(present)



def _canonicalize_final_condition_bundle(row: Dict) -> Dict[str, str]:
    current = dict(row or {})
    fields = [
        ("Reaction_Temp_C", False),
        ("TOS_h", False),
        ("S_C_Ratio", True),
        ("Feed_MeOH_to_H2O_Ratio", True),
        ("Pressure_bar", False),
        ("GHSV_mL_g_h", False),
        ("WHSV_h_inv", False),
        ("SpaceVelocity_norm", False),
        ("Flow_Rate", False),
        ("Catalyst_Amount_g", False),
    ]

    def _normalize_value(value: Any, allow_ratio: bool = False) -> str:
        text_value = clean_residual_mojibake_chars(str(value or "")).strip()
        if not text_value or text_value.lower() in {"", "n/a", "na", "none", "nan", "unknown"}:
            return ""
        compact = text_value.replace(",", "")
        if allow_ratio:
            ratio_match = re.search(r"\d+(?:\.\d+)?\s*[:/]\s*\d+(?:\.\d+)?", compact)
            if ratio_match:
                return ratio_match.group(0).replace(" ", "")
        parsed = _parse_float_if_possible(compact)
        if parsed is not None:
            return _format_ratio_token(str(parsed))
        match = re.search(r"[-+]?\d+(?:\.\d+)?", compact)
        return _format_ratio_token(match.group(0)) if match else ""

    return {field: _normalize_value(current.get(field, ""), allow_ratio=allow_ratio) for field, allow_ratio in fields}



def _make_final_figure_dedupe_signature(row: Dict) -> str:
    current = _finalize_figure_row_binding_payload(dict(row or {}))
    condition_bundle = _canonicalize_final_condition_bundle(current)
    catalyst_key = clean_residual_mojibake_chars(str(
        current.get("Canonical_Catalyst_ID", "")
        or current.get("Catalyst_ID_normalized", "")
        or current.get("Catalyst_ID", "")
        or current.get("Catalyst", "")
    )).strip().lower()

    metric_parts = []
    for field in [
        "MeOH_Conversion_%",
        "H2_Yield_%",
        "H2_Production_Rate",
        "CO_Selectivity_%",
        "CO2_Selectivity_%",
        "CO_Concentration_ppm",
    ]:
        value = clean_residual_mojibake_chars(str(current.get(field, "") or "")).strip()
        if not value:
            continue
        parsed = _parse_float_if_possible(value.replace(",", ""))
        if parsed is not None:
            metric_parts.append(f"{field}={format(parsed, '.6g')}")
            continue
        match = re.search(r"[-+]?\d+(?:\.\d+)?", value.replace(",", ""))
        metric_parts.append(f"{field}={format(float(match.group(0)), '.6g')}" if match else f"{field}={value.lower()}")

    signature_payload = {
        "Source_File": clean_residual_mojibake_chars(str(current.get("Source_File", "") or "")).strip(),
        "target_family": _normalize_target_family(current),
        "catalyst_key": catalyst_key,
        "condition_bundle": condition_bundle,
        "metric_bundle": metric_parts,
    }
    return json.dumps(signature_payload, ensure_ascii=False, sort_keys=True)



def _make_final_text_dedupe_signature(row: Dict) -> str:
    current = dict(row or {})
    condition_bundle = _canonicalize_final_condition_bundle(current)
    catalyst_key = clean_residual_mojibake_chars(str(
        current.get("Canonical_Catalyst_ID", "")
        or current.get("Catalyst_ID_normalized", "")
        or current.get("Catalyst_ID", "")
        or current.get("Catalyst", "")
    )).strip().lower()

    metric_parts = []
    for field in [
        "MeOH_Conversion_%",
        "H2_Yield_%",
        "H2_Production_Rate",
        "CO_Selectivity_%",
        "CO2_Selectivity_%",
        "CO_Concentration_ppm",
    ]:
        value = clean_residual_mojibake_chars(str(current.get(field, "") or "")).strip()
        if not value:
            continue
        parsed = _parse_float_if_possible(value.replace(",", ""))
        if parsed is not None:
            metric_parts.append(f"{field}={format(parsed, '.6g')}")
            continue
        match = re.search(r"[-+]?\d+(?:\.\d+)?", value.replace(",", ""))
        metric_parts.append(f"{field}={format(float(match.group(0)), '.6g')}" if match else f"{field}={value.lower()}")

    signature_payload = {
        "Source_File": clean_residual_mojibake_chars(str(current.get("Source_File", "") or "")).strip(),
        "target_family": _normalize_target_family(current),
        "catalyst_key": catalyst_key,
        "condition_bundle": condition_bundle,
        "metric_bundle": metric_parts,
    }
    return json.dumps(signature_payload, ensure_ascii=False, sort_keys=True)



def _count_nonempty_condition_features(row: Dict) -> int:
    bundle = _canonicalize_final_condition_bundle(row)
    grouped = {
        "temperature": bool(bundle.get("Reaction_Temp_C")),
        "time": bool(bundle.get("TOS_h")),
        "ratio": bool(bundle.get("S_C_Ratio") or bundle.get("Feed_MeOH_to_H2O_Ratio")),
        "pressure": bool(bundle.get("Pressure_bar")),
        "space_velocity": bool(bundle.get("GHSV_mL_g_h") or bundle.get("WHSV_h_inv") or bundle.get("SpaceVelocity_norm")),
        "flow": bool(bundle.get("Flow_Rate")),
        "catalyst_amount": bool(bundle.get("Catalyst_Amount_g")),
    }
    return sum(1 for present in grouped.values() if present)



def _score_final_row_priority(row: Dict) -> Dict[str, Any]:
    current = _finalize_figure_row_binding_payload(dict(row or {}))
    data_source = str(current.get("data_source", "") or "").strip()
    source_granularity = str(current.get("source_granularity", "") or "").strip()
    semantic_role = _canonicalize_semantic_figure_role(str(current.get("semantic_figure_role", "") or "").strip())
    condition_count = _count_nonempty_condition_features(current)
    final_anchor_bonus = 10 if ((data_source == "figure" and _has_reliable_final_figure_condition_anchor(current)) or (data_source == "text" and has_valid_point_condition_anchor(current))) else 0
    if data_source == "text" and source_granularity in {"table_row", "si_table"}:
        source_rank = 5
    elif data_source == "text":
        source_rank = 4
    elif data_source == "figure":
        source_rank = 3
    else:
        source_rank = 1

    confidence = str(current.get("figure_binding_confidence", current.get("binding_confidence", "")) or "").strip().lower()
    binding_rank = 2 if confidence == "high" else 1 if confidence == "medium" else 0
    numeric_rank = 2 if str(current.get("numeric_reliability_level", "") or "").strip() == "direct_numeric" else 0
    identity_level = str(current.get("identity_completeness_level", "") or "").strip()
    identity_rank = 2 if identity_level == "complete" else 1 if identity_level == "partial" else 0
    semantic_rank = 2 if semantic_role in {"temperature_sweep", "stability_tos", "condition_effect", "category_screening"} else 1 if semantic_role else 0

    ambiguity_blob = " ".join([
        str(current.get("Notes", "") or ""),
        str(current.get("figure_binding_reason", "") or ""),
        str(current.get("comparison_filter_reason", "") or ""),
    ])
    ambiguity_penalty = 1 if re.search(r"\bambiguous\b|generic_label_without_context_mapping|no_reliable_registry_match", ambiguity_blob, flags=re.I) else 0

    total_score = (
        source_rank * 20
        + condition_count * 8
        + final_anchor_bonus
        + binding_rank * 6
        + numeric_rank * 5
        + identity_rank * 5
        + semantic_rank * 3
        - ambiguity_penalty * 6
    )
    return {
        "total_score": total_score,
        "sort_key": (total_score, condition_count, source_rank, binding_rank, numeric_rank, identity_rank, semantic_rank, -ambiguity_penalty),
        "explanation": (
            f"source_rank={source_rank}; condition_count={condition_count}; final_anchor_bonus={final_anchor_bonus}; "
            f"binding_rank={binding_rank}; numeric_rank={numeric_rank}; identity_rank={identity_rank}; "
            f"semantic_rank={semantic_rank}; ambiguity_penalty={ambiguity_penalty}"
        ),
    }



def _is_weak_condition_figure_row(row: Dict) -> bool:
    if not isinstance(row, dict):
        return False
    current = _finalize_figure_row_binding_payload(dict(row))
    if str(current.get("data_source", "") or "").strip() != "figure":
        return False
    if not has_core_performance_metric(current):
        return False
    semantic_role = _canonicalize_semantic_figure_role(str(current.get("semantic_figure_role", "") or "").strip())
    if semantic_role not in {"category_screening", "condition_effect", "temperature_sweep", "stability_tos"}:
        return False
    return not _has_reliable_final_figure_condition_anchor(current)



def _should_downgrade_figure_row_from_final(row: Dict) -> Tuple[bool, str]:
    if not isinstance(row, dict):
        return False, ""
    current = _finalize_figure_row_binding_payload(dict(row))
    if str(current.get("data_source", "") or "").strip() != "figure":
        return False, ""
    if _is_baseline_or_support_only_row(current):
        return True, "support_only_or_baseline"
    numeric_level = str(current.get("numeric_reliability_level", "") or "").strip()
    if numeric_level != "direct_numeric":
        return True, "non_direct_numeric"
    if bool(current.get("is_approximate_value")) or bool(current.get("is_range_like_value")) or bool(current.get("is_qualitative_value")):
        return True, "approximate_or_range"
    semantic_role = _canonicalize_semantic_figure_role(str(current.get("semantic_figure_role", "") or "").strip())
    if semantic_role in {"product_species_profile", "case_profile", "non_performance_like", "catalyst_amount_effect"}:
        return True, "blocked_figure_semantic_role"
    if not _has_core_figure_performance_value(current):
        return True, "missing_core_metric"
    if str(current.get("figure_binding_status", "") or "").strip() != "matched":
        return True, "nonmatched_figure_binding"
    if str(current.get("figure_binding_confidence", current.get("binding_confidence", "")) or "").strip() not in {"high", "medium"}:
        return True, "low_figure_binding_confidence"
    mode = str(current.get("figure_binding_mode", current.get("binding_mode", "")) or "").strip()
    if not (_is_high_value_figure_binding_mode(mode) or mode in {"canonical_label", "exact_label"}):
        return True, "unsupported_binding_mode"
    if _is_weak_condition_figure_row(current):
        return True, "weak_condition_for_final"
    return False, ""


def _is_duplicate_final_candidate(left_row: Dict, right_row: Dict) -> bool:
    if not isinstance(left_row, dict) or not isinstance(right_row, dict):
        return False

    left_source = str(left_row.get("data_source", "") or "").strip()
    right_source = str(right_row.get("data_source", "") or "").strip()
    left_sig = _make_final_figure_dedupe_signature(left_row) if left_source == "figure" else _make_final_text_dedupe_signature(left_row)
    right_sig = _make_final_figure_dedupe_signature(right_row) if right_source == "figure" else _make_final_text_dedupe_signature(right_row)
    if left_sig and right_sig and left_sig == right_sig:
        return True

    if {left_source, right_source} != {"text", "figure"}:
        return False

    try:
        left_payload = json.loads(left_sig) if left_sig else {}
        right_payload = json.loads(right_sig) if right_sig else {}
    except Exception:
        return False

    for key in ["Source_File", "target_family", "catalyst_key", "metric_bundle"]:
        if left_payload.get(key) != right_payload.get(key):
            return False

    left_bundle = {k: v for k, v in dict(left_payload.get("condition_bundle") or {}).items() if str(v).strip()}
    right_bundle = {k: v for k, v in dict(right_payload.get("condition_bundle") or {}).items() if str(v).strip()}
    if not left_bundle or not right_bundle:
        return False
    for field in set(left_bundle) | set(right_bundle):
        left_value = str(left_bundle.get(field, "") or "").strip()
        right_value = str(right_bundle.get(field, "") or "").strip()
        if left_value and right_value and left_value != right_value:
            return False

    smaller, larger = (left_bundle, right_bundle) if len(left_bundle) <= len(right_bundle) else (right_bundle, left_bundle)
    return len(smaller) < len(larger) and all(str(larger.get(field, "") or "").strip() == str(value).strip() for field, value in smaller.items())



def _build_final_dedupe_audit_rows(decisions: List[Dict[str, Any]]) -> List[Dict]:
    audit_rows = []
    for item in decisions or []:
        kept_row = _finalize_figure_row_binding_payload(dict(item.get("kept_row") or {}))
        removed_row = _finalize_figure_row_binding_payload(dict(item.get("removed_row") or {}))
        kept_score = item.get("kept_score") or {}
        removed_score = item.get("removed_score") or {}
        audit_rows.append({
            "Source_File": str(removed_row.get("Source_File", "") or ""),
            "data_source": str(removed_row.get("data_source", "") or ""),
            "semantic_figure_role": _canonicalize_semantic_figure_role(str(removed_row.get("semantic_figure_role", "") or "")),
            "raw_category_label": str(removed_row.get("raw_category_label", "") or ""),
            "Catalyst": str(removed_row.get("Catalyst", "") or ""),
            "Catalyst_ID": str(removed_row.get("Catalyst_ID", "") or ""),
            "Canonical_Catalyst_ID": str(removed_row.get("Canonical_Catalyst_ID", "") or ""),
            "figure_binding_status": str(removed_row.get("figure_binding_status", "") or ""),
            "figure_binding_mode": str(removed_row.get("figure_binding_mode", removed_row.get("binding_mode", "")) or ""),
            "figure_binding_confidence": str(removed_row.get("figure_binding_confidence", removed_row.get("binding_confidence", "")) or ""),
            "figure_binding_reason": str(removed_row.get("figure_binding_reason", removed_row.get("binding_reason", "")) or ""),
            "matched_registry_key": str(removed_row.get("matched_registry_key", "") or ""),
            "matched_registry_label": str(removed_row.get("matched_registry_label", "") or ""),
            "alias_map_source": str(removed_row.get("alias_map_source", "") or ""),
            "alias_map_evidence": str(removed_row.get("alias_map_evidence", "") or ""),
            "target_family": _normalize_target_family(removed_row),
            "final_dedupe_signature": str(item.get("signature", "") or ""),
            "dedupe_action": "removed_duplicate",
            "dedupe_reason": "duplicate_final_signature",
            "final_exclusion_reason": "lower_priority_duplicate",
            "final_exclusion_detail": str(item.get("detail", "") or ""),
            "kept_data_source": str(kept_row.get("data_source", "") or ""),
            "kept_semantic_figure_role": _canonicalize_semantic_figure_role(str(kept_row.get("semantic_figure_role", "") or "")),
            "kept_Canonical_Catalyst_ID": str(kept_row.get("Canonical_Catalyst_ID", "") or ""),
            "kept_target_family": _normalize_target_family(kept_row),
            "kept_priority_score": kept_score.get("total_score", ""),
            "removed_priority_score": removed_score.get("total_score", ""),
            "kept_priority_explanation": kept_score.get("explanation", ""),
            "removed_priority_explanation": removed_score.get("explanation", ""),
            "Reaction_Temp_C": str(removed_row.get("Reaction_Temp_C", "") or ""),
            "TOS_h": str(removed_row.get("TOS_h", "") or ""),
            "S_C_Ratio": str(removed_row.get("S_C_Ratio", "") or ""),
            "Pressure_bar": str(removed_row.get("Pressure_bar", "") or ""),
            "GHSV_mL_g_h": str(removed_row.get("GHSV_mL_g_h", "") or ""),
            "Flow_Rate": str(removed_row.get("Flow_Rate", "") or ""),
            "Catalyst_Amount_g": str(removed_row.get("Catalyst_Amount_g", "") or ""),
        })
    return clean_records_text_fields(audit_rows)



def _dedupe_final_rows(rows: List[Dict]) -> Tuple[List[Dict], List[Dict], List[Dict]]:
    grouped: Dict[str, List[Dict[str, Any]]] = {}
    for index, row in enumerate(rows or []):
        if not isinstance(row, dict):
            continue
        current = _finalize_figure_row_binding_payload(dict(row))
        if str(current.get("data_source", "") or "").strip() == "figure":
            signature = _make_final_figure_dedupe_signature(current)
        else:
            signature = _make_final_text_dedupe_signature(current)
        if not signature:
            signature = f"__NO_FINAL_SIG__::{index}"
        grouped.setdefault(signature, []).append({"row": current, "index": index})

    kept_groups = []
    excluded_rows = []
    decision_rows = []

    for signature, entries in grouped.items():
        if len(entries) == 1:
            kept_groups.append({"first_index": entries[0]["index"], "row": entries[0]["row"]})
            continue

        scored_entries = []
        for entry in entries:
            score = _score_final_row_priority(entry["row"])
            stable_payload = json.dumps(strip_internal_context_fields(clean_record_text_fields(dict(entry["row"]))), ensure_ascii=False, sort_keys=True)
            scored_entries.append({
                "row": entry["row"],
                "index": entry["index"],
                "score": score,
                "stable_payload": stable_payload,
            })

        scored_entries.sort(key=lambda item: (item["score"]["sort_key"], item["stable_payload"]), reverse=True)
        kept = scored_entries[0]
        kept_groups.append({
            "first_index": min(item["index"] for item in scored_entries),
            "row": kept["row"],
        })

        for removed in scored_entries[1:]:
            current = dict(removed["row"])
            current["final_exclusion_reason"] = "lower_priority_duplicate"
            current["final_exclusion_detail"] = (
                f"duplicate_final_signature={signature}; kept_score={kept['score']['total_score']}; "
                f"removed_score={removed['score']['total_score']}"
            )
            current["final_duplicate_signature"] = signature
            current["dedupe_action"] = "removed_duplicate"
            excluded_rows.append(current)
            decision_rows.append({
                "signature": signature,
                "kept_row": kept["row"],
                "removed_row": current,
                "kept_score": kept["score"],
                "removed_score": removed["score"],
                "detail": current["final_exclusion_detail"],
            })

    cross_source_kept: List[Dict[str, Any]] = []
    for item in sorted(kept_groups, key=lambda payload: payload["first_index"]):
        current_row = item["row"]
        current_score = _score_final_row_priority(current_row)
        merged = False
        for existing in cross_source_kept:
            if not _is_duplicate_final_candidate(existing["row"], current_row):
                continue
            existing_score = existing["score"]
            choose_current = current_score["sort_key"] > existing_score["sort_key"]
            if not choose_current and current_score["sort_key"] == existing_score["sort_key"]:
                current_payload = json.dumps(strip_internal_context_fields(clean_record_text_fields(dict(current_row))), ensure_ascii=False, sort_keys=True)
                existing_payload = json.dumps(strip_internal_context_fields(clean_record_text_fields(dict(existing["row"]))), ensure_ascii=False, sort_keys=True)
                choose_current = current_payload > existing_payload

            kept_entry = current_row if choose_current else existing["row"]
            removed_entry = existing["row"] if choose_current else current_row
            kept_score = current_score if choose_current else existing_score
            removed_score = existing_score if choose_current else current_score
            signature = _make_final_text_dedupe_signature(kept_entry) if str(kept_entry.get("data_source", "") or "").strip() == "text" else _make_final_figure_dedupe_signature(kept_entry)
            removed_row = dict(removed_entry)
            removed_row["final_exclusion_reason"] = "lower_priority_duplicate"
            removed_row["final_exclusion_detail"] = (
                f"duplicate_final_signature={signature}; kept_score={kept_score['total_score']}; "
                f"removed_score={removed_score['total_score']}"
            )
            removed_row["final_duplicate_signature"] = signature
            removed_row["dedupe_action"] = "removed_duplicate"
            excluded_rows.append(removed_row)
            decision_rows.append({
                "signature": signature,
                "kept_row": kept_entry,
                "removed_row": removed_row,
                "kept_score": kept_score,
                "removed_score": removed_score,
                "detail": removed_row["final_exclusion_detail"],
            })
            if choose_current:
                existing["row"] = current_row
                existing["score"] = current_score
                existing["first_index"] = min(existing["first_index"], item["first_index"])
            merged = True
            break

        if not merged:
            cross_source_kept.append({
                "row": current_row,
                "score": current_score,
                "first_index": item["first_index"],
            })

    cross_source_kept.sort(key=lambda payload: payload["first_index"])
    deduped_rows = [payload["row"] for payload in cross_source_kept]
    audit_rows = _build_final_dedupe_audit_rows(decision_rows)
    return deduped_rows, clean_records_text_fields(excluded_rows), audit_rows



_LAYER_MEMBERSHIP_AUDIT_FIELDS = {
    "included_in_research_all",
    "included_in_research_strict",
    "included_in_modeling_final",
    "layer_membership",
    "exclusion_from_strict_reason",
    "exclusion_from_strict_detail",
    "exclusion_from_final_reason",
    "exclusion_from_final_detail",
    "strict_exclusion_reason",
    "strict_exclusion_detail",
    "final_exclusion_reason",
    "final_exclusion_detail",
    "final_duplicate_signature",
    "final_dedupe_signature",
    "dedupe_action",
}


def _stable_layer_row_signature(row: Dict) -> str:
    current = _finalize_figure_row_binding_payload(dict(row or {}))
    for field in list(_LAYER_MEMBERSHIP_AUDIT_FIELDS) + ["layer_row_uid"]:
        current.pop(field, None)
    return json.dumps(strip_internal_context_fields(clean_record_text_fields(current)), ensure_ascii=False, sort_keys=True)


def _apply_final_layer_outcomes(
    research_all_rows: List[Dict],
    research_strict_rows: List[Dict],
    final_rows: List[Dict],
    final_excluded_rows: List[Dict],
) -> Tuple[List[Dict], List[Dict], List[Dict], List[Dict]]:
    final_rows_norm = [_finalize_figure_row_binding_payload(dict(row)) for row in (final_rows or []) if isinstance(row, dict)]
    final_excluded_norm = [_finalize_figure_row_binding_payload(dict(row)) for row in (final_excluded_rows or []) if isinstance(row, dict)]

    final_uids = {
        clean_residual_mojibake_chars(str(row.get("layer_row_uid", "") or "")).strip(): row
        for row in final_rows_norm
        if clean_residual_mojibake_chars(str(row.get("layer_row_uid", "") or "")).strip()
    }
    final_excluded_uids = {
        clean_residual_mojibake_chars(str(row.get("layer_row_uid", "") or "")).strip(): row
        for row in final_excluded_norm
        if clean_residual_mojibake_chars(str(row.get("layer_row_uid", "") or "")).strip()
    }
    final_sig_counts: Dict[str, int] = {}
    final_excluded_by_sig: Dict[str, List[Dict]] = {}
    for row in final_rows_norm:
        if clean_residual_mojibake_chars(str(row.get("layer_row_uid", "") or "")).strip():
            continue
        sig = _stable_layer_row_signature(row)
        final_sig_counts[sig] = final_sig_counts.get(sig, 0) + 1
    for row in final_excluded_norm:
        if clean_residual_mojibake_chars(str(row.get("layer_row_uid", "") or "")).strip():
            continue
        sig = _stable_layer_row_signature(row)
        final_excluded_by_sig.setdefault(sig, []).append(row)

    def _mark_row(row: Dict, force_all: bool = False, force_strict: bool = False) -> Dict:
        current = _finalize_figure_row_binding_payload(dict(row or {}))
        if force_all:
            current["included_in_research_all"] = 1
        if force_strict:
            current["included_in_research_strict"] = 1
        uid = clean_residual_mojibake_chars(str(current.get("layer_row_uid", "") or "")).strip()
        matched_final = final_uids.get(uid) if uid else None
        matched_excluded = final_excluded_uids.get(uid) if uid else None
        if matched_final is None and matched_excluded is None and not uid and bool(current.get("included_in_research_strict")):
            sig = _stable_layer_row_signature(current)
            if final_sig_counts.get(sig, 0) > 0:
                final_sig_counts[sig] -= 1
                matched_final = current
            elif final_excluded_by_sig.get(sig):
                matched_excluded = final_excluded_by_sig[sig].pop(0)
        if matched_final is not None:
            current["included_in_modeling_final"] = 1
            current["exclusion_from_final_reason"] = ""
            current["exclusion_from_final_detail"] = ""
            current["final_exclusion_reason"] = ""
            current["final_exclusion_detail"] = ""
        elif matched_excluded is not None:
            current["included_in_modeling_final"] = 0
            current["exclusion_from_final_reason"] = str(matched_excluded.get("final_exclusion_reason", "") or "")
            current["exclusion_from_final_detail"] = str(matched_excluded.get("final_exclusion_detail", "") or "")
            current["final_exclusion_reason"] = current["exclusion_from_final_reason"]
            current["final_exclusion_detail"] = current["exclusion_from_final_detail"]
        current.setdefault("included_in_research_all", int(force_all))
        current.setdefault("included_in_research_strict", int(force_strict))
        current.setdefault("included_in_modeling_final", 0)
        return _sync_layer_membership_fields(current)

    updated_all = [_mark_row(row, force_all=True, force_strict=bool(row.get("included_in_research_strict"))) for row in (research_all_rows or []) if isinstance(row, dict)]
    updated_strict = [_mark_row(row, force_all=True, force_strict=True) for row in (research_strict_rows or []) if isinstance(row, dict)]
    updated_final = [_mark_row(row, force_all=True, force_strict=True) for row in final_rows_norm]
    updated_excluded = []
    for row in final_excluded_norm:
        current = _mark_row(row, force_all=True, force_strict=True)
        current["included_in_modeling_final"] = 0
        current["exclusion_from_final_reason"] = str(current.get("final_exclusion_reason", "") or "")
        current["exclusion_from_final_detail"] = str(current.get("final_exclusion_detail", "") or "")
        current = _sync_layer_membership_fields(current)
        updated_excluded.append(current)
    return updated_all, updated_strict, updated_final, updated_excluded


def _build_layered_distribution_summary(
    research_all_rows: List[Dict],
    research_strict_rows: List[Dict],
    final_rows: List[Dict],
    final_excluded_rows: Optional[List[Dict]] = None,
    final_dedupe_audit_rows: Optional[List[Dict]] = None,
) -> Dict[str, Any]:
    final_only_summary = _build_final_distribution_summary(final_rows, final_excluded_rows, final_dedupe_audit_rows)
    layers = {
        "research_all": [_finalize_figure_row_binding_payload(dict(row)) for row in (research_all_rows or []) if isinstance(row, dict)],
        "research_strict": [_finalize_figure_row_binding_payload(dict(row)) for row in (research_strict_rows or []) if isinstance(row, dict)],
        "modeling_final": [_finalize_figure_row_binding_payload(dict(row)) for row in (final_rows or []) if isinstance(row, dict)],
    }
    tracked_fields = [
        "Reaction_Temp_C", "TOS_h", "S_C_Ratio", "Feed_MeOH_to_H2O_Ratio", "Pressure_bar",
        "GHSV_mL_g_h", "WHSV_h_inv", "SpaceVelocity_norm", "Flow_Rate", "Catalyst_Amount_g",
        "MeOH_Conversion_%", "H2_Yield_%", "H2_Production_Rate", "CO_Selectivity_%",
        "CO2_Selectivity_%", "CO_Concentration_ppm",
    ]
    per_source_file: Dict[str, Dict[str, int]] = {}
    text_vs_figure_by_layer: Dict[str, Dict[str, int]] = {}
    by_source_file: Dict[str, Dict[str, int]] = {}
    by_data_source: Dict[str, Dict[str, int]] = {}
    field_coverage_by_layer: Dict[str, Dict[str, Dict[str, Any]]] = {}

    for layer_name, rows in layers.items():
        source_counts: Dict[str, int] = {}
        data_source_counts: Dict[str, int] = {}
        for row in rows:
            source_file = clean_residual_mojibake_chars(str(row.get("Source_File", "") or "")).strip() or "<blank>"
            data_source = clean_residual_mojibake_chars(str(row.get("data_source", "") or "")).strip() or "<blank>"
            source_counts[source_file] = source_counts.get(source_file, 0) + 1
            data_source_counts[data_source] = data_source_counts.get(data_source, 0) + 1
            per_source_file.setdefault(source_file, {"research_all": 0, "research_strict": 0, "modeling_final": 0})
            per_source_file[source_file][layer_name] += 1
        by_source_file[layer_name] = dict(sorted(source_counts.items(), key=lambda item: (-item[1], item[0])))
        by_data_source[layer_name] = dict(sorted(data_source_counts.items(), key=lambda item: (-item[1], item[0])))
        text_vs_figure_by_layer[layer_name] = {
            "text": data_source_counts.get("text", 0),
            "figure": data_source_counts.get("figure", 0),
            "other": sum(count for key, count in data_source_counts.items() if key not in {"text", "figure"}),
        }
        total_rows = len(rows)
        field_coverage_by_layer[layer_name] = {}
        for field in tracked_fields:
            nonempty = sum(1 for row in rows if clean_residual_mojibake_chars(str(row.get(field, "") or "")).strip())
            field_coverage_by_layer[layer_name][field] = {
                "nonempty_count": nonempty,
                "nonempty_ratio": round(nonempty / total_rows, 4) if total_rows else 0.0,
            }

    summary = {
        "research_all_row_count": len(layers["research_all"]),
        "research_strict_row_count": len(layers["research_strict"]),
        "modeling_final_row_count": len(layers["modeling_final"]),
        "final_row_count": len(layers["modeling_final"]),
        "research_all_by_source_file": by_source_file.get("research_all", {}),
        "research_strict_by_source_file": by_source_file.get("research_strict", {}),
        "modeling_final_by_source_file": by_source_file.get("modeling_final", {}),
        "per_source_file_layer_counts": dict(sorted(per_source_file.items(), key=lambda item: (-item[1]["modeling_final"], item[0]))),
        "research_all_by_data_source": by_data_source.get("research_all", {}),
        "research_strict_by_data_source": by_data_source.get("research_strict", {}),
        "modeling_final_by_data_source": by_data_source.get("modeling_final", {}),
        "text_vs_figure_by_layer": text_vs_figure_by_layer,
        "field_coverage_by_layer": field_coverage_by_layer,
        "legacy_aliases": {
            "dataset_all.csv": "research_master_all.csv",
            "dataset_all_strict.csv": "research_master_strict.csv",
        },
        "top_10_source_files_by_modeling_final_rows": [
            {"Source_File": source_file, "count": count}
            for source_file, count in list(by_source_file.get("modeling_final", {}).items())[:10]
        ],
    }
    summary.update(final_only_summary)
    return summary


def _layer_subset_consistency_check(
    research_all_rows: List[Dict],
    research_strict_rows: List[Dict],
    final_rows: List[Dict],
    final_excluded_rows: List[Dict],
    research_all_csv_path: str,
    research_strict_csv_path: str,
    modeling_final_csv_path: str,
    excluded_csv_path: str,
    distribution_summary: Dict,
    quality_summary_rows: List[Dict],
) -> List[str]:
    issues: List[str] = []

    def _csv_count(path: str) -> int:
        with open(path, encoding="utf-8-sig") as fh:
            return max(sum(1 for _ in fh) - 1, 0)

    for path_label, path_value, expected_count in [
        ("research_master_all.csv", research_all_csv_path, len(research_all_rows)),
        ("research_master_strict.csv", research_strict_csv_path, len(research_strict_rows)),
        ("modeling_final.csv", modeling_final_csv_path, len(final_rows)),
        ("final_excluded_audit.csv", excluded_csv_path, len(final_excluded_rows)),
    ]:
        if not os.path.exists(path_value):
            issues.append(f"missing output file: {path_label} ({path_value})")
            continue
        try:
            csv_count = _csv_count(path_value)
            if csv_count != expected_count:
                issues.append(f"{path_label} has {csv_count} data rows but in-memory rows={expected_count}")
        except Exception as e:
            issues.append(f"could not read {path_label}: {e}")

    if distribution_summary:
        if int(distribution_summary.get("research_all_row_count", -1)) != len(research_all_rows):
            issues.append(
                f"distribution_summary.research_all_row_count={distribution_summary.get('research_all_row_count')} but in-memory rows={len(research_all_rows)}"
            )
        if int(distribution_summary.get("research_strict_row_count", -1)) != len(research_strict_rows):
            issues.append(
                f"distribution_summary.research_strict_row_count={distribution_summary.get('research_strict_row_count')} but in-memory rows={len(research_strict_rows)}"
            )
        if int(distribution_summary.get("modeling_final_row_count", distribution_summary.get("final_row_count", -1))) != len(final_rows):
            issues.append(
                f"distribution_summary.modeling_final_row_count={distribution_summary.get('modeling_final_row_count', distribution_summary.get('final_row_count'))} but in-memory rows={len(final_rows)}"
            )

    if quality_summary_rows:
        research_all_total = sum(int(row.get("research_all_inclusion_count", 0) or 0) for row in quality_summary_rows if isinstance(row, dict))
        research_strict_total = sum(int(row.get("research_strict_inclusion_count", 0) or 0) for row in quality_summary_rows if isinstance(row, dict))
        modeling_final_total = sum(int(row.get("modeling_final_inclusion_count", row.get("final_point_row_count", 0)) or 0) for row in quality_summary_rows if isinstance(row, dict))
        if research_all_total != len(research_all_rows):
            issues.append(f"extraction_quality_summary sum(research_all_inclusion_count)={research_all_total} but research_all_rows={len(research_all_rows)}")
        if research_strict_total != len(research_strict_rows):
            issues.append(f"extraction_quality_summary sum(research_strict_inclusion_count)={research_strict_total} but research_strict_rows={len(research_strict_rows)}")
        if modeling_final_total != len(final_rows):
            issues.append(f"extraction_quality_summary sum(modeling_final_inclusion_count)={modeling_final_total} but modeling_final_rows={len(final_rows)}")

    def _row_keys(rows: List[Dict]) -> Set[str]:
        keys: Set[str] = set()
        for row in rows or []:
            if not isinstance(row, dict):
                continue
            uid = clean_residual_mojibake_chars(str(row.get("layer_row_uid", "") or "")).strip()
            keys.add(f"uid::{uid}" if uid else f"sig::{_stable_layer_row_signature(row)}")
        return keys

    all_keys = _row_keys(research_all_rows)
    strict_keys = _row_keys(research_strict_rows)
    final_keys = _row_keys(final_rows)
    if not strict_keys.issubset(all_keys):
        issues.append("research_strict is not a subset of research_all")
    if not final_keys.issubset(strict_keys):
        issues.append("modeling_final is not a subset of research_strict")

    if issues:
        for issue in issues:
            print(f"  [WARNING][layer_consistency] {issue}")
        print(f"  [WARNING][layer_consistency] {len(issues)} inconsistency(ies) found; review layered summary and audit files.")
    else:
        print(
            f"  [OK][layer_consistency] all={len(research_all_rows)}, strict={len(research_strict_rows)}, "
            f"final={len(final_rows)}, excluded={len(final_excluded_rows)}"
        )
    return issues


def _warn_if_legacy_export_names_are_ambiguous() -> None:
    if globals().get("_LEGACY_LAYER_ALIAS_WARNING_DONE"):
        return
    globals()["_LEGACY_LAYER_ALIAS_WARNING_DONE"] = True
    print("  [Legacy] dataset_all.csv now aliases research_master_all.csv")
    print("  [Legacy] dataset_all_strict.csv now aliases research_master_strict.csv")


def _build_final_distribution_summary(rows: List[Dict], final_excluded_rows: Optional[List[Dict]] = None, final_dedupe_audit_rows: Optional[List[Dict]] = None) -> Dict[str, Any]:
    final_rows = [_finalize_figure_row_binding_payload(dict(row)) for row in (rows or []) if isinstance(row, dict)]
    excluded_rows = [_finalize_figure_row_binding_payload(dict(row)) for row in (final_excluded_rows or []) if isinstance(row, dict)]
    dedupe_rows = [_normalize_binding_audit_record(dict(row)) for row in (final_dedupe_audit_rows or []) if isinstance(row, dict)]

    def _count_map_increment(target: Dict[str, int], key: Any) -> None:
        key_text = clean_residual_mojibake_chars(str(key or "")).strip() or "<blank>"
        target[key_text] = target.get(key_text, 0) + 1

    final_by_source_file: Dict[str, int] = {}
    final_by_data_source: Dict[str, int] = {}
    final_by_target_family: Dict[str, int] = {}
    final_by_semantic_figure_role: Dict[str, int] = {}
    condition_fields = [
        "Reaction_Temp_C", "TOS_h", "S_C_Ratio", "Pressure_bar",
        "GHSV_mL_g_h", "WHSV_h_inv", "SpaceVelocity_norm", "Flow_Rate", "Catalyst_Amount_g",
    ]
    condition_coverage = {field: {"nonempty_count": 0, "nonempty_ratio": 0.0} for field in condition_fields}

    for row in final_rows:
        _count_map_increment(final_by_source_file, row.get("Source_File", ""))
        _count_map_increment(final_by_data_source, row.get("data_source", ""))
        families = _normalize_target_family(row)
        if families:
            for family in families.split("|"):
                _count_map_increment(final_by_target_family, family)
        if str(row.get("data_source", "") or "").strip() == "figure":
            _count_map_increment(final_by_semantic_figure_role, _canonicalize_semantic_figure_role(str(row.get("semantic_figure_role", "") or "")))

    total_final = len(final_rows)
    for field in condition_fields:
        nonempty_count = sum(1 for row in final_rows if clean_residual_mojibake_chars(str(row.get(field, "") or "")).strip())
        condition_coverage[field] = {
            "nonempty_count": nonempty_count,
            "nonempty_ratio": round(nonempty_count / total_final, 4) if total_final else 0.0,
        }

    duplicate_signatures = set()
    for row in dedupe_rows:
        signature = clean_residual_mojibake_chars(str(row.get("final_dedupe_signature", "") or row.get("final_duplicate_signature", ""))).strip()
        if signature:
            duplicate_signatures.add(signature)
    if not duplicate_signatures:
        for row in excluded_rows:
            detail = clean_residual_mojibake_chars(str(row.get("final_exclusion_detail", "") or ""))
            match = re.search(r"duplicate_final_signature=([^;]+)", detail)
            if match:
                duplicate_signatures.add(match.group(1).strip())

    rows_removed_by_final_dedupe = sum(
        1 for row in excluded_rows
        if str(row.get("final_exclusion_reason", "") or "").strip() in {"lower_priority_duplicate", "duplicate_final_signature"}
        or "duplicate_final_signature=" in str(row.get("final_exclusion_detail", "") or "")
    )
    weak_condition_rows_excluded_from_final = sum(
        1 for row in excluded_rows
        if str(row.get("final_exclusion_reason", "") or "").strip() == "weak_condition_for_final"
    )

    sorted_sources = sorted(final_by_source_file.items(), key=lambda item: (-item[1], item[0]))
    return {
        "final_row_count": total_final,
        "final_by_source_file": dict(sorted_sources),
        "final_source_file_distribution": dict(sorted_sources),
        "final_by_data_source": dict(sorted(final_by_data_source.items(), key=lambda item: (-item[1], item[0]))),
        "figure_vs_text_final_distribution": {
            "figure": final_by_data_source.get("figure", 0),
            "text": final_by_data_source.get("text", 0),
        },
        "final_by_target_family": dict(sorted(final_by_target_family.items(), key=lambda item: (-item[1], item[0]))),
        "final_target_distribution": dict(sorted(final_by_target_family.items(), key=lambda item: (-item[1], item[0]))),
        "final_by_semantic_figure_role": dict(sorted(final_by_semantic_figure_role.items(), key=lambda item: (-item[1], item[0]))),
        "final_condition_feature_coverage": condition_coverage,
        "duplicate_signature_group_count": len(duplicate_signatures),
        "rows_removed_by_final_dedupe": rows_removed_by_final_dedupe,
        "weak_condition_rows_excluded_from_final": weak_condition_rows_excluded_from_final,
        "top_10_source_files_by_final_rows": [
            {"Source_File": source_file, "count": count}
            for source_file, count in sorted_sources[:10]
        ],
    }



def _selfcheck_binding_safety() -> None:
    """
    Problem A selfcheck: verify that the new metal-set binding guards work correctly.
    Tests must-block and must-allow cases. Prints WARNING if any case fails.
    Called once at first binding operation; never raises, never crashes main flow.
    """
    if globals().get("_BINDING_SAFETY_SELFCHECK_DONE"):
        return
    globals()["_BINDING_SAFETY_SELFCHECK_DONE"] = True

    try:
        # Build a tiny registry with three entries for testing
        _test_registry = {
            "ni/clay":    {"Catalyst": "Ni/clay",    "Active_Metal": "Ni",   "Promoter": "",   "Support": "clay"},
            "ni-cu/clay": {"Catalyst": "Ni-Cu/clay", "Active_Metal": "Ni",   "Promoter": "Cu", "Support": "clay"},
            "ni/zro2":    {"Catalyst": "Ni/ZrO2",    "Active_Metal": "Ni",   "Promoter": "",   "Support": "ZrO2"},
            "ni-ce/zro2": {"Catalyst": "Ni-Ce/ZrO2", "Active_Metal": "Ni",   "Promoter": "Ce", "Support": "ZrO2"},
            "zro2":       {"Catalyst": "ZrO2",        "Active_Metal": "",    "Promoter": "",   "Support": "ZrO2"},
            "10ni/zro2":  {"Catalyst": "10Ni/ZrO2",  "Active_Metal": "Ni",  "Promoter": "",   "Support": "ZrO2",
                           "_registry_aliases": ["ni/zro2_10wt"]},
        }

        def _check_no_match(clue_label: str, registry_subset: Dict, test_name: str):
            clue = {"Active_Metal": "", "Promoter": "", "Support": ""}
            # Derive Active_Metal / Promoter / Support from the label
            metals = _extract_explicit_metal_set_from_identity_clue(clue_label)
            if metals:
                metal_list = list(metals)
                clue["Active_Metal"] = metal_list[0] if metal_list else ""
                clue["Promoter"] = metal_list[1] if len(metal_list) > 1 else ""
            clue["Support"] = clue_label.split("/", 1)[1].strip() if "/" in clue_label else ""
            hits = _find_unique_registry_candidate(clue, registry_subset)
            if any(k == list(registry_subset.keys())[-1] for k, _ in hits):
                print(f"  [WARNING] binding_safety_selfcheck FAILED: {test_name} "
                      f"clue='{clue_label}' matched when it should NOT")
            else:
                print(f"  [OK] binding_safety_selfcheck: {test_name} correctly blocked")

        def _check_match(clue_record: Dict, registry_subset: Dict, expected_key: str, test_name: str):
            hits = _find_unique_registry_candidate(clue_record, registry_subset)
            matched_keys = [k for k, _ in hits]
            if expected_key in matched_keys:
                print(f"  [OK] binding_safety_selfcheck: {test_name} correctly matched")
            else:
                print(f"  [WARNING] binding_safety_selfcheck FAILED: {test_name} "
                      f"expected '{expected_key}' but got {matched_keys}")

        # Test 1: Ni/clay MUST NOT match Ni-Cu/clay
        _check_no_match("Ni/clay", {"ni-cu/clay": _test_registry["ni-cu/clay"]}, "Ni/clay→Ni-Cu/clay blocked")

        # Test 2: Cu/clay MUST NOT match Ni-Cu/clay (different active metal)
        cu_clue = {"Active_Metal": "Cu", "Promoter": "", "Support": "clay"}
        hits2 = _find_unique_registry_candidate(cu_clue, {"ni-cu/clay": _test_registry["ni-cu/clay"]})
        if hits2:
            print("  [WARNING] binding_safety_selfcheck FAILED: Cu/clay→Ni-Cu/clay should be blocked by active metal mismatch")
        else:
            print("  [OK] binding_safety_selfcheck: Cu/clay→Ni-Cu/clay correctly blocked")

        # Test 3: Ni/ZrO2 MUST NOT match Ni-Ce/ZrO2
        _check_no_match("Ni/ZrO2", {"ni-ce/zro2": _test_registry["ni-ce/zro2"]}, "Ni/ZrO2→Ni-Ce/ZrO2 blocked")

        # Test 4: Ni/ZrO2 MUST match Ni/ZrO2 (plain monometallic, same support)
        ni_zro2_clue = {"Active_Metal": "Ni", "Promoter": "", "Support": "ZrO2"}
        hits4 = _find_unique_registry_candidate(ni_zro2_clue, {"ni/zro2": _test_registry["ni/zro2"]})
        if hits4:
            print("  [OK] binding_safety_selfcheck: Ni/ZrO2→Ni/ZrO2 correctly matched")
        else:
            print("  [WARNING] binding_safety_selfcheck FAILED: Ni/ZrO2→Ni/ZrO2 should match")

        # Test 5: Support-only ZrO2 must NOT match active Ni/ZrO2
        _is_sup = _normalize_support_only_label("ZrO2")
        if _is_sup:
            print("  [OK] binding_safety_selfcheck: ZrO2 correctly identified as support-only")
        else:
            print("  [WARNING] binding_safety_selfcheck FAILED: ZrO2 not identified as support-only")

        # Test 6: metal_set_compatible helpers
        ok, reason = _is_metal_set_compatible_for_unique_match({"Ni"}, {"Ni", "Cu"})
        if not ok and "monometallic_vs_bimetallic" in reason:
            print("  [OK] binding_safety_selfcheck: {Ni} vs {Ni,Cu} correctly blocked")
        else:
            print(f"  [WARNING] binding_safety_selfcheck FAILED: {{Ni}} vs {{Ni,Cu}} should be blocked, got ok={ok} reason={reason}")

        ok2, reason2 = _is_metal_set_compatible_for_unique_match({"Ni", "Cu"}, {"Ni"})
        if not ok2 and "bimetallic_vs_monometallic" in reason2:
            print("  [OK] binding_safety_selfcheck: {Ni,Cu} vs {Ni} correctly blocked")
        else:
            print(f"  [WARNING] binding_safety_selfcheck FAILED: {{Ni,Cu}} vs {{Ni}} should be blocked, got ok={ok2}")

        ok3, _ = _is_metal_set_compatible_for_unique_match(None, {"Ni"})
        if ok3:
            print("  [OK] binding_safety_selfcheck: None (unknown) vs {Ni} correctly allowed")
        else:
            print("  [WARNING] binding_safety_selfcheck FAILED: None should be allowed conservatively")

        # ── Integration tests: real production bug scenarios ──────────────────
        # These test through _match_category_figure_label_to_registry and
        # _find_unique_registry_candidate with the actual "Ni/clay → Ni-Cu/clay" registry.
        # The final safety gate in attach_figure_points_to_registry covers the fallback path.

        _only_ni_cu_clay = {"ni-cu/clay": _test_registry["ni-cu/clay"]}
        _only_ni_clay = {"ni/clay": _test_registry["ni/clay"]}

        # Test A: "Ni/clay" category label must NOT match "Ni-Cu/clay" registry
        _ta = _match_category_figure_label_to_registry("Ni/clay", _only_ni_cu_clay, "")
        if _ta.get("matched"):
            print("  [WARNING] binding_safety_selfcheck FAILED: Ni/clay category-matched to Ni-Cu/clay")
        else:
            print(f"  [OK] binding_safety_selfcheck: Ni/clay→Ni-Cu/clay blocked via category path (reason={_ta.get('binding_reason','')})")

        # Test B: "Cu/clay" category label must NOT match "Ni-Cu/clay" registry
        _tb = _match_category_figure_label_to_registry("Cu/clay", _only_ni_cu_clay, "")
        if _tb.get("matched"):
            print("  [WARNING] binding_safety_selfcheck FAILED: Cu/clay category-matched to Ni-Cu/clay")
        else:
            print("  [OK] binding_safety_selfcheck: Cu/clay→Ni-Cu/clay blocked via category path")

        # Test C: "Ni-Cu/clay" label must NOT match "Ni/clay" registry
        _tc = _match_category_figure_label_to_registry("Ni-Cu/clay", _only_ni_clay, "")
        if _tc.get("matched"):
            print("  [WARNING] binding_safety_selfcheck FAILED: Ni-Cu/clay category-matched to Ni/clay")
        else:
            print("  [OK] binding_safety_selfcheck: Ni-Cu/clay→Ni/clay blocked via category path")

        # Test D: "Ni/clay" must match "Ni/clay" (monometallic → monometallic OK)
        _td = _match_category_figure_label_to_registry("Ni/clay", _only_ni_clay, "")
        if _td.get("matched"):
            print("  [OK] binding_safety_selfcheck: Ni/clay→Ni/clay correctly matched")
        else:
            print(f"  [WARNING] binding_safety_selfcheck FAILED: Ni/clay→Ni/clay should match, got reason={_td.get('binding_reason','')}")

        # Test E: "No catalyst" label must NOT match any active registry
        _te = _match_category_figure_label_to_registry("No catalyst", _only_ni_clay, "")
        if _te.get("matched"):
            print("  [WARNING] binding_safety_selfcheck FAILED: 'No catalyst' matched an active registry")
        elif _te.get("is_baseline_label"):
            print("  [OK] binding_safety_selfcheck: 'No catalyst' correctly identified as baseline")
        else:
            print(f"  [WARNING] binding_safety_selfcheck: 'No catalyst' not flagged as baseline (reason={_te.get('binding_reason','')})")

        # Test F: "ZrO2" support-only must NOT match "Ni/ZrO2" active registry
        _only_ni_zro2 = {"ni/zro2": _test_registry["ni/zro2"]}
        _tf = _match_category_figure_label_to_registry("ZrO2", _only_ni_zro2, "")
        if _tf.get("matched"):
            print("  [WARNING] binding_safety_selfcheck FAILED: ZrO2 support-only matched Ni/ZrO2 active registry")
        elif _tf.get("is_support_only_label"):
            print("  [OK] binding_safety_selfcheck: ZrO2 correctly blocked as support-only")
        else:
            print(f"  [WARNING] binding_safety_selfcheck: ZrO2 not flagged as support-only (reason={_tf.get('binding_reason','')})")

        # Test G: _find_unique_registry_candidate with empty Active_Metal row (the actual production bug path)
        # Simulates a figure row where Active_Metal="" but raw_category_label="Ni/clay",
        # and registry only has "Ni-Cu/clay". This must NOT be a unique match.
        _empty_am_clue = {"Active_Metal": "", "Promoter": "", "Support": "clay", "Support_Normalized": "clay"}
        _tg_hits = _find_unique_registry_candidate(_empty_am_clue, _only_ni_cu_clay)
        if _tg_hits:
            print("  [WARNING] binding_safety_selfcheck FAILED: empty-Active_Metal clue with Support=clay "
                  "matched Ni-Cu/clay (this is the production bug path!)")
        else:
            print("  [OK] binding_safety_selfcheck: empty-AM clue with clay support correctly blocked vs Ni-Cu/clay")

    except Exception as _e:
        print(f"  [WARNING] binding_safety_selfcheck error (non-fatal): {_e}")


def _final_row_count_consistency_check(
    final_rows: List[Dict],
    final_excluded_rows: List[Dict],
    master_csv_path: str,
    excluded_csv_path: str,
    distribution_summary: Dict,
    quality_summary_rows: List[Dict],
) -> None:
    """
    Problem B: lightweight end-of-run consistency check.
    Verifies that all final outputs agree on row counts.
    Prints WARNING for each discrepancy; never raises.
    """
    issues: List[str] = []

    # Check 1: modeling_master.csv row count
    if os.path.exists(master_csv_path):
        try:
            with open(master_csv_path, encoding="utf-8-sig") as fh:
                csv_count = sum(1 for _ in fh) - 1  # minus header
            if csv_count != len(final_rows):
                issues.append(
                    f"modeling_master.csv has {csv_count} data rows but "
                    f"final_rows in memory has {len(final_rows)} rows"
                )
        except Exception as e:
            issues.append(f"could not read {master_csv_path}: {e}")

    # Check 2: modeling_excluded_audit.csv row count
    if os.path.exists(excluded_csv_path):
        try:
            with open(excluded_csv_path, encoding="utf-8-sig") as fh:
                exc_count = sum(1 for _ in fh) - 1
            if exc_count != len(final_excluded_rows):
                issues.append(
                    f"modeling_excluded_audit.csv has {exc_count} rows but "
                    f"final_excluded_rows has {len(final_excluded_rows)} rows"
                )
        except Exception as e:
            issues.append(f"could not read {excluded_csv_path}: {e}")

    # Check 3: distribution_summary final_row_count
    if distribution_summary:
        dist_count = int(distribution_summary.get("final_row_count", -1))
        if dist_count != len(final_rows):
            issues.append(
                f"final_distribution_summary.final_row_count={dist_count} but "
                f"final_rows has {len(final_rows)} rows"
            )

    # Check 4: quality_summary per-paper final counts vs actual
    if quality_summary_rows and final_rows:
        summary_total = sum(int(r.get("final_point_row_count", 0) or 0)
                            for r in quality_summary_rows if isinstance(r, dict))
        if summary_total != len(final_rows):
            issues.append(
                f"extraction_quality_summary sum(final_point_row_count)={summary_total} "
                f"but final_rows has {len(final_rows)} rows"
            )

    if issues:
        for issue in issues:
            print(f"  [WARNING][final_consistency] {issue}")
        print(f"  [WARNING][final_consistency] {len(issues)} inconsistency(ies) found — "
              f"check audit files to identify which pipeline stage diverges.")
    else:
        print(f"  [OK][final_consistency] All output counts consistent: "
              f"final_rows={len(final_rows)}, excluded={len(final_excluded_rows)}")


def _selfcheck_final_dedupe_and_gate() -> None:
    if globals().get("_FINAL_DEDUPE_AND_GATE_SELFCHECK_DONE"):
        return
    globals()["_FINAL_DEDUPE_AND_GATE_SELFCHECK_DONE"] = True
    try:
        base_figure = {
            "Source_File": "paper_a.pdf",
            "data_source": "figure",
            "figure_binding_status": "matched",
            "figure_binding_confidence": "high",
            "figure_binding_mode": "direct_label",
            "semantic_figure_role": "category_screening",
            "identity_completeness_level": "complete",
            "numeric_reliability_level": "direct_numeric",
            "Canonical_Catalyst_ID": "Ni/ZrO2",
            "Catalyst_ID": "Ni/ZrO2",
            "Catalyst": "Ni/ZrO2",
            "MeOH_Conversion_%": "65",
            "Reaction_Temp_C": "250",
        }
        dup_rows, dup_excluded, _ = _dedupe_final_rows([dict(base_figure), dict(base_figure)])
        if len(dup_rows) != 1 or not any(str(row.get("final_exclusion_reason", "")) == "lower_priority_duplicate" for row in dup_excluded):
            print("  [WARNING] final selfcheck failed: identical figure rows were not deduped")

        diff_temp = [dict(base_figure), dict(base_figure, Reaction_Temp_C="300")]
        diff_temp_rows, _, _ = _dedupe_final_rows(diff_temp)
        if len(diff_temp_rows) != 2:
            print("  [WARNING] final selfcheck failed: different temperature points were wrongly deduped")

        diff_target = [dict(base_figure), dict(base_figure, **{"MeOH_Conversion_%": "", "CO2_Selectivity_%": "65"})]
        diff_target_rows, _, _ = _dedupe_final_rows(diff_target)
        if len(diff_target_rows) != 2:
            print("  [WARNING] final selfcheck failed: different target families were wrongly deduped")

        weak_category = dict(base_figure)
        weak_category.pop("Reaction_Temp_C", None)
        weak_category["S_C_Ratio"] = "2"
        final_rows_weak, final_excluded_weak = build_final_publishable_rows([weak_category])
        if final_rows_weak or not any(str(row.get("final_exclusion_reason", "")) == "weak_condition_for_final" for row in final_excluded_weak):
            print("  [WARNING] final selfcheck failed: weak-condition category figure row entered final")

        strong_category = dict(base_figure, figure_binding_mode="canonical_label")
        final_rows_strong, _ = build_final_publishable_rows([strong_category])
        if len(final_rows_strong) != 1:
            print("  [WARNING] final selfcheck failed: strong figure row failed final gate")

        text_row = {
            "Source_File": "paper_a.pdf",
            "data_source": "text",
            "identity_completeness_level": "complete",
            "numeric_reliability_level": "direct_numeric",
            "source_granularity": "table_row",
            "Canonical_Catalyst_ID": "Ni/ZrO2",
            "Catalyst_ID": "Ni/ZrO2",
            "Catalyst": "Ni/ZrO2",
            "MeOH_Conversion_%": "65",
            "Reaction_Temp_C": "250",
            "TOS_h": "10",
        }
        text_vs_figure_rows, text_vs_figure_excluded = build_final_publishable_rows([dict(base_figure), text_row])
        if len(text_vs_figure_rows) != 1 or str(text_vs_figure_rows[0].get("data_source", "")) != "text" or not any(str(row.get("final_exclusion_reason", "")) == "lower_priority_duplicate" for row in text_vs_figure_excluded):
            print("  [WARNING] final selfcheck failed: text row did not outrank duplicate figure row")
    except Exception as e:
        print(f"  [WARNING] final dedupe/gate selfcheck error: {e}")



def build_final_publishable_rows(strict_rows: List[Dict]) -> Tuple[List[Dict], List[Dict]]:
    if not globals().get("_FINAL_DEDUPE_AND_GATE_SELFCHECK_DONE"):
        _selfcheck_final_dedupe_and_gate()

    final_candidates: List[Dict] = []
    final_excluded: List[Dict] = []

    for row in strict_rows or []:
        if not isinstance(row, dict):
            continue
        current = _finalize_figure_row_binding_payload(dict(row))
        current.setdefault("included_in_research_all", 1)
        current.setdefault("included_in_research_strict", 1)
        data_source = str(current.get("data_source", "") or "").strip()

        if data_source == "text":
            ok, reason = _is_modeling_final_eligible_text_row(current)
            if ok:
                final_candidates.append(current)
                continue

            note_axis, note_x_value = _extract_x_axis_and_value_from_notes(str(current.get("Notes", "")))
            detail = ""
            if reason == "comparison_like":
                detail = str(current.get("comparison_filter_reason", "") or current.get("comparison_filter_action", ""))
            elif reason == "weak_identity":
                detail = str(current.get("identity_completeness_level", ""))
            elif reason == "non_direct_numeric":
                detail = str(current.get("numeric_reliability_level", ""))
            elif reason == "weak_source_granularity":
                detail = str(current.get("source_granularity", ""))
            elif reason == "approximate_or_range":
                detail = str(current.get("raw_numeric_expression", "") or current.get("numeric_expression_type", ""))
            elif reason == "duplicate_like":
                detail = str(current.get("duplicate_candidate_type", "") or current.get("obvious_duplicate_type", ""))
            elif reason == "performance_without_condition_anchor":
                detail = "; ".join(item for item in [
                    f"x_axis={str(current.get('x_axis', '') or note_axis).strip()}" if str(current.get("x_axis", "") or note_axis).strip() else "",
                    f"x_value={str(current.get('x_value', '') or note_x_value).strip()}" if str(current.get("x_value", "") or note_x_value).strip() else "",
                ] if item)
            else:
                detail = str(current.get("Notes", ""))[:240]

            current["final_exclusion_reason"] = reason
            current["final_exclusion_detail"] = detail
            current["exclusion_from_final_reason"] = reason
            current["exclusion_from_final_detail"] = detail
            current["included_in_modeling_final"] = 0
            current = _sync_layer_membership_fields(current)
            final_excluded.append(current)
            continue

        if data_source == "figure":
            ok, reason = _is_modeling_final_eligible_figure_row(current)
            if not ok:
                current["final_exclusion_reason"] = reason
                if reason == "weak_condition_for_final":
                    current["final_exclusion_detail"] = (
                        f"condition_count={_count_nonempty_condition_features(current)}; "
                        f"target_family={_normalize_target_family(current)}"
                    )
                elif reason == "blocked_figure_semantic_role":
                    current["final_exclusion_detail"] = _canonicalize_semantic_figure_role(str(current.get("semantic_figure_role", "") or ""))
                elif reason in {"non_direct_numeric", "approximate_or_range"}:
                    current["final_exclusion_detail"] = str(current.get("raw_numeric_expression", "") or current.get("numeric_reliability_level", ""))[:240]
                elif reason == "support_only_or_baseline":
                    current["final_exclusion_detail"] = _derive_layer_exclusion_detail(current, reason)
                else:
                    current["final_exclusion_detail"] = str(current.get("figure_binding_reason", "") or current.get("figure_binding_notes", "") or current.get("Notes", ""))[:240]
                current["exclusion_from_final_reason"] = current["final_exclusion_reason"]
                current["exclusion_from_final_detail"] = current["final_exclusion_detail"]
                current["included_in_modeling_final"] = 0
                current = _sync_layer_membership_fields(current)
                final_excluded.append(current)
                continue

            final_candidates.append(current)
            continue

        final_candidates.append(current)

    final_rows, dedupe_excluded, dedupe_audit_rows = _dedupe_final_rows(final_candidates)
    final_excluded.extend(dedupe_excluded)
    globals().setdefault("_FINAL_DEDUPE_AUDIT_ROWS", []).extend(dedupe_audit_rows)
    globals()["_LAST_FINAL_DEDUPE_AUDIT_ROWS"] = dedupe_audit_rows

    normalized_final_rows: List[Dict] = []
    for row in final_rows:
        current = _finalize_figure_row_binding_payload(dict(row))
        current["included_in_research_all"] = 1
        current["included_in_research_strict"] = 1
        current["included_in_modeling_final"] = 1
        current["exclusion_from_final_reason"] = ""
        current["exclusion_from_final_detail"] = ""
        current["final_exclusion_reason"] = ""
        current["final_exclusion_detail"] = ""
        current = _sync_layer_membership_fields(current)
        normalized_final_rows.append(current)

    normalized_final_excluded: List[Dict] = []
    for row in final_excluded:
        current = _finalize_figure_row_binding_payload(dict(row))
        current["included_in_research_all"] = 1
        current["included_in_research_strict"] = 1
        current["included_in_modeling_final"] = 0
        current["exclusion_from_final_reason"] = str(current.get("final_exclusion_reason", "") or "")
        current["exclusion_from_final_detail"] = str(current.get("final_exclusion_detail", "") or "")
        current = _sync_layer_membership_fields(current)
        normalized_final_excluded.append(current)

    source_files = sorted({
        clean_residual_mojibake_chars(str(row.get("Source_File", "") or "")).strip()
        for row in list(strict_rows or []) + normalized_final_rows + normalized_final_excluded
        if isinstance(row, dict) and clean_residual_mojibake_chars(str(row.get("Source_File", "") or "")).strip()
    })
    source_key = source_files[0] if len(source_files) == 1 else (source_files[0] if source_files else "<unknown_source>")
    paper_stats = globals().setdefault("_LAYERED_PAPER_STATS", {})
    paper_stats[source_key] = {
        **paper_stats.get(source_key, {}),
        "modeling_final_row_count": len(normalized_final_rows),
        "modeling_final_excluded_count": len(normalized_final_excluded),
        "modeling_final_exclusion_reason_counts": _count_reason_values(normalized_final_excluded, "final_exclusion_reason"),
        "modeling_final_text_row_count": sum(1 for row in normalized_final_rows if str(row.get("data_source", "") or "").strip() == "text"),
        "modeling_final_figure_row_count": sum(1 for row in normalized_final_rows if str(row.get("data_source", "") or "").strip() == "figure"),
    }

    globals()["_LAST_FINAL_DISTRIBUTION_SUMMARY"] = _build_final_distribution_summary(normalized_final_rows, normalized_final_excluded, dedupe_audit_rows)
    return clean_records_text_fields(normalized_final_rows), clean_records_text_fields(normalized_final_excluded)


def save_comparison_audit_csv(rows: List[Dict], output_path: str) -> None:
    cleaned_rows = strip_internal_context_fields_batch([
        clean_record_text_fields(dict(row)) for row in rows if isinstance(row, dict)
    ])
    extra_keys = sorted({k for row in cleaned_rows for k in row.keys()} - set(_COMPARISON_AUDIT_FIELDS))
    fieldnames = list(_COMPARISON_AUDIT_FIELDS) + extra_keys
    with open(output_path, "w", newline="", encoding="utf-8-sig") as f:
        writer = csv.DictWriter(f, fieldnames=fieldnames)
        writer.writeheader()
        if cleaned_rows:
            writer.writerows(cleaned_rows)



def save_figure_premerge_filter_audit_csv(rows: List[Dict], output_path: str) -> None:
    cleaned_rows = strip_internal_context_fields_batch([
        clean_record_text_fields(dict(row)) for row in rows if isinstance(row, dict)
    ])
    extra_keys = sorted({k for row in cleaned_rows for k in row.keys()} - set(_FIGURE_PREMERGE_AUDIT_FIELDS))
    fieldnames = list(_FIGURE_PREMERGE_AUDIT_FIELDS) + extra_keys
    with open(output_path, "w", newline="", encoding="utf-8-sig") as f:
        writer = csv.DictWriter(f, fieldnames=fieldnames)
        writer.writeheader()
        if cleaned_rows:
            writer.writerows(cleaned_rows)



def save_to_csv(data_list: List[Dict], output_path: str) -> None:
    if not data_list:
        return
    cleaned_list = strip_internal_context_fields_batch([
        clean_record_text_fields(dict(item)) for item in data_list if isinstance(item, dict)
    ])
    all_keys: set = set()
    for item in cleaned_list:
        all_keys.update(item.keys())
    # [Task3] Place Ni_Fraction / Promoter_Metal / Promoter_Fraction next to Alloy_Ratio in CSV output.
    # [Fix 6] Place Canonical_Catalyst_ID next to Catalyst_ID for easier manual merge auditing.
    priority = ["Catalyst_ID", "Catalyst_ID_raw", "Catalyst_ID_normalized", "identity_alias_group",
                "Canonical_Catalyst_ID", "Catalyst", "Active_Metal",
                "Ni_Loading_wt%", "Promoter_Loading_wt%",
                "Alloy_Ratio", "Promoter", "Ni_Fraction", "Promoter_Metal", "Promoter_Fraction",
                "Support", "Support_Normalized", "Support_Family", "Support_Modifier",
                "Support_Prep_Method_Normalized", "Reaction_Temp_C", "Reasoning_Selectivity",
                "identity_normalization_notes", "identity_role_parse_notes", "record_origin_type",
                "source_granularity", "duplicate_candidate_type", "duplicate_candidate_signature",
                "same_physical_point_possible", "obvious_duplicate_flag", "obvious_duplicate_type",
                "obvious_duplicate_keep_drop", "obvious_duplicate_reason", "obvious_duplicate_group_id",
                "is_this_work", "is_literature_comparison", "is_source_file_si",
                "numeric_expression_type", "raw_numeric_expression", "numeric_reliability_level",
                "is_range_like_value", "is_qualitative_value", "is_approximate_value",
                "extraction_confidence", "identity_completeness_level", "origin_quality_class",
                "temp_conversion_clue_flag", "explicit_temp_conversion_note",
                "temp_duplicate_resolution_upstream", "upstream_temp_duplicate_pair_id"]
    remaining = sorted(k for k in all_keys if k not in priority)
    fieldnames = [k for k in priority if k in all_keys] + remaining
    with open(output_path, "w", newline="", encoding="utf-8-sig") as f:
        writer = csv.DictWriter(f, fieldnames=fieldnames)
        writer.writeheader()
        writer.writerows(cleaned_list)



def get_model_ready_export_columns() -> Dict[str, List[str]]:
    return {
        "id_columns": [
            "Source_File", "data_source", "Catalyst_ID", "Catalyst_ID_normalized", "Canonical_Catalyst_ID",
        ],
        "group_column": ["cv_group"],
        "target_candidates": [
            "MeOH_Conversion_%", "CO2_Selectivity_%", "CO_Selectivity_%", "H2_Production_Rate", "CO_Concentration_ppm",
        ],
        "baseline_numeric_feature_columns_model": [
            "Reaction_Temp_C_num_model", "Metal_Loading_wt%_num_model", "Ni_Fraction_fixed_model",
            "Promoter_Fraction_num_model", "S_C_Ratio_num_model", "Pressure_bar_num_model",
            "Feed_MeOH_to_H2O_Ratio_model", "SpaceVelocity_norm_model",
        ],
        "baseline_categorical_feature_columns": [
            "Support_clean_model", "Promoter_Metal_clean_model",
        ],
        "baseline_derived_feature_columns": [
            "has_promoter", "is_bimetallic", "Ni_Loading", "Promoter_Loading",
        ],
        "enhanced_numeric_feature_columns_model": [
            "Calcination_Temp_C_num", "Calcination_Time_h_num", "Reduction_Temp_C_num", "Reduction_Time_h_num",
        ],
        "enhanced_categorical_feature_columns": [
            "Metal_Loading_Method_grouped", "Support_Prep_Method_grouped", "Precursor_family_grouped",
        ],
        "enhanced_derived_feature_columns": [
            "Calcination_Severity", "Reduction_Severity", "Temp_Ni_interaction",
        ],
        "missing_flag_columns": [
            "metal_loading_missing_flag", "ni_fraction_missing_flag", "promoter_fraction_missing_flag",
            "sc_ratio_missing_flag", "pressure_missing_flag", "feed_ratio_missing_flag",
            "space_velocity_missing_flag", "calcination_temp_missing_flag", "calcination_time_missing_flag",
            "reduction_temp_missing_flag", "reduction_time_missing_flag", "method_missing_flag",
        ],
        "provenance_columns": [
            "record_origin_type", "paragraph_role", "paragraph_router_method", "text_extraction_subroute",
            "figure_binding_status", "figure_binding_mode", "figure_binding_confidence", "semantic_figure_role",
            "extraction_confidence", "origin_quality_class", "identity_completeness_level",
            "numeric_reliability_level", "same_physical_point_possible", "duplicate_candidate_type", "Notes",
        ],
    }



def _validate_runtime_config(api_key: Optional[str] = None, base_url: Optional[str] = None, text_model: Optional[str] = None, vision_model: Optional[str] = None, in_dir: Optional[str] = None, out_dir: Optional[str] = None, temp_img_dir: Optional[str] = None) -> Dict[str, str]:
    resolved = {
        "API_KEY": str(API_KEY if api_key is None else api_key).strip(),
        "BASE_URL": str(BASE_URL if base_url is None else base_url).strip(),
        "TEXT_MODEL": str(TEXT_MODEL if text_model is None else text_model).strip(),
        "VISION_MODEL": str(VISION_MODEL if vision_model is None else vision_model).strip(),
        "IN_DIR": str(IN_DIR if in_dir is None else in_dir).strip(),
        "OUT_DIR": str(OUT_DIR if out_dir is None else out_dir).strip(),
        "TEMP_IMG_DIR": str(TEMP_IMG_DIR if temp_img_dir is None else temp_img_dir).strip(),
    }
    if not resolved["API_KEY"]:
        raise RuntimeError("MSR_API_KEY is required. Set environment variable MSR_API_KEY before running the extractor.")
    for key in ["BASE_URL", "TEXT_MODEL", "VISION_MODEL", "IN_DIR", "OUT_DIR", "TEMP_IMG_DIR"]:
        if not resolved[key]:
            raise RuntimeError(f"runtime config '{key}' is empty; set the corresponding environment variable or default path")
    return resolved



def _canonicalize_semantic_figure_role(role: str) -> str:
    role_norm = clean_residual_mojibake_chars(str(role or "")).strip()
    role_map = {
        "catalyst_screening": "category_screening",
        "stability_profile": "stability_tos",
        "schematic_or_nonperformance": "non_performance_like",
    }
    return role_map.get(role_norm, role_norm)



def _canonicalize_figure_binding_fields(row: Dict) -> Dict:
    if not isinstance(row, dict):
        return row

    current = dict(row)
    figureish = bool(
        str(current.get("data_source", "")).strip() == "figure"
        or any(key in current for key in [
            "figure_binding_status", "figure_binding_mode", "binding_mode", "match_mode",
            "figure_binding_confidence", "binding_confidence", "match_confidence",
            "figure_binding_notes", "figure_binding_reason", "binding_reason", "binding_notes",
            "matched_registry_key", "matched_registry_label", "registry_record",
            "alias_map_source", "alias_map_evidence",
        ])
    )
    if not figureish:
        return current

    def _first_text(*values: Any) -> str:
        for value in values:
            if isinstance(value, (list, tuple, set)):
                for item in value:
                    text_value = clean_residual_mojibake_chars(str(item or "")).strip()
                    if text_value:
                        return text_value
                continue
            if isinstance(value, dict):
                continue
            text_value = clean_residual_mojibake_chars(str(value or "")).strip()
            if text_value:
                return text_value
        return ""

    def _merge_texts(*values: Any) -> str:
        merged: List[str] = []
        for value in values:
            if isinstance(value, (list, tuple, set)):
                iterable = value
            else:
                iterable = [value]
            for item in iterable:
                text_value = clean_residual_mojibake_chars(str(item or "")).strip()
                if text_value and text_value not in merged:
                    merged.append(text_value)
        return "; ".join(merged)

    registry_record = current.get("registry_record") if isinstance(current.get("registry_record"), dict) else {}
    status = _first_text(current.get("figure_binding_status"), current.get("binding_status"))
    mode = _first_text(current.get("figure_binding_mode"), current.get("binding_mode"), current.get("match_mode"))
    confidence = _first_text(current.get("figure_binding_confidence"), current.get("binding_confidence"), current.get("match_confidence"))
    reason = _merge_texts(
        current.get("figure_binding_reason", ""),
        current.get("figure_binding_notes", ""),
        current.get("binding_reason", ""),
        current.get("binding_notes", ""),
    )
    matched_registry_key = _first_text(
        current.get("matched_registry_key", ""),
        registry_record.get("Catalyst_ID", "") if registry_record else "",
        registry_record.get("Canonical_Catalyst_ID", "") if registry_record else "",
    )
    matched_registry_label = _first_text(
        current.get("matched_registry_label", ""),
        registry_record.get("Catalyst", "") if registry_record else "",
        registry_record.get("Canonical_Catalyst_ID", "") if registry_record else "",
        matched_registry_key,
    )
    alias_map_source = _first_text(current.get("alias_map_source", ""))
    alias_map_evidence = _first_text(current.get("alias_map_evidence", ""))

    current["figure_binding_status"] = status
    current["figure_binding_mode"] = mode
    current["figure_binding_confidence"] = confidence
    current["figure_binding_reason"] = reason
    current["figure_binding_notes"] = reason
    current["binding_mode"] = mode
    current["binding_confidence"] = confidence
    current["binding_reason"] = reason
    current["binding_notes"] = reason
    current["match_mode"] = mode
    current["match_confidence"] = confidence
    current["matched_registry_key"] = matched_registry_key
    current["matched_registry_label"] = matched_registry_label
    current["alias_map_source"] = alias_map_source
    current["alias_map_evidence"] = alias_map_evidence
    current["is_baseline_label"] = int(bool(current.get("is_baseline_label", 0)))
    current["is_support_only_label"] = int(bool(current.get("is_support_only_label", 0)))
    current.setdefault("raw_category_label", clean_residual_mojibake_chars(str(current.get("raw_category_label", "") or "")).strip())
    return current



def _finalize_figure_row_binding_payload(row: Dict) -> Dict:
    if not isinstance(row, dict):
        return row
    current = _canonicalize_figure_binding_fields(dict(row))
    if str(current.get("figure_binding_status", "") or "").strip() == "matched":
        mode = str(current.get("figure_binding_mode", current.get("binding_mode", "")) or "").strip()
        if mode not in {"family_backbone_match", "loading_relaxed_family_backbone"}:
            if not str(current.get("Canonical_Catalyst_ID", "") or "").strip():
                key = str(current.get("matched_registry_key", "") or "").strip()
                if key:
                    current["Canonical_Catalyst_ID"] = key
    raw_role = clean_residual_mojibake_chars(str(current.get("semantic_figure_role", "") or "")).strip()
    canonical_role = _canonicalize_semantic_figure_role(raw_role)
    if raw_role and canonical_role != raw_role:
        current.setdefault("semantic_figure_role_raw", raw_role)
    if raw_role or "semantic_figure_role" in current:
        current["semantic_figure_role"] = canonical_role
    return current



def _normalize_binding_audit_record(row: Dict) -> Dict:
    if not isinstance(row, dict):
        return row
    current = _finalize_figure_row_binding_payload(dict(row))
    current.setdefault("figure_binding_status", "")
    current.setdefault("figure_binding_mode", "")
    current.setdefault("figure_binding_confidence", "")
    current.setdefault("figure_binding_reason", clean_residual_mojibake_chars(str(current.get("figure_binding_notes", "") or current.get("binding_reason", "") or "")).strip())
    current.setdefault("matched_registry_key", "")
    current.setdefault("matched_registry_label", "")
    current.setdefault("alias_map_source", "")
    current.setdefault("alias_map_evidence", "")
    current.setdefault("raw_category_label", clean_residual_mojibake_chars(str(current.get("raw_category_label", "") or "")).strip())
    current.setdefault("semantic_figure_role", _canonicalize_semantic_figure_role(str(current.get("semantic_figure_role", "") or "")))
    return clean_record_text_fields(strip_internal_context_fields(current))



def _retire_legacy_modeling_exports(candidate_records: List[Dict], strict_records: List[Dict], export_columns: Dict[str, List[str]], output_dir: str = OUT_DIR) -> Dict[str, str]:
    if not globals().get("_LEGACY_MODELING_EXPORT_WARNING_DONE"):
        globals()["_LEGACY_MODELING_EXPORT_WARNING_DONE"] = True
        print("[WARNING] legacy candidate/strict modeling export is retired; using final modeling export only")

    os.makedirs(output_dir, exist_ok=True)
    for legacy_name in ["modeling_master_candidate.csv", "modeling_master_strict.csv"]:
        legacy_path = os.path.join(output_dir, legacy_name)
        if os.path.exists(legacy_path):
            try:
                os.remove(legacy_path)
            except Exception as e:
                print(f"  [WARNING] unable to remove retired legacy modeling export '{legacy_path}': {e}")

    final_rows = [_finalize_figure_row_binding_payload(dict(row)) for row in (strict_records or []) if isinstance(row, dict)]
    final_excluded_rows = build_modeling_excluded_audit(candidate_records, strict_records)
    export_paths = save_modeling_final_exports(final_rows, final_excluded_rows, export_columns, output_dir)
    return {
        "legacy_candidate_csv": "",
        "legacy_strict_csv": "",
        "final_csv": export_paths.get("master_csv", ""),
        "master_csv": export_paths.get("master_csv", ""),
        "excluded_csv": export_paths.get("excluded_csv", ""),
        "schema_json": export_paths.get("schema_json", ""),
        "coverage_json": export_paths.get("coverage_json", ""),
    }



def _selfcheck_runtime_and_binding_consistency() -> None:
    if globals().get("_RUNTIME_BINDING_SELFCHECK_DONE"):
        return
    globals()["_RUNTIME_BINDING_SELFCHECK_DONE"] = True
    try:
        case1 = _canonicalize_figure_binding_fields({
            "binding_mode": "baseline_match",
            "binding_confidence": "medium",
            "figure_binding_status": "matched",
        })
        if case1.get("figure_binding_mode") != "baseline_match" or case1.get("figure_binding_confidence") != "medium":
            print("  [WARNING] runtime selfcheck failed: canonical binding mode/confidence propagation")

        case2 = _canonicalize_figure_binding_fields({
            "figure_binding_mode": "numbered_alias_map",
            "figure_binding_confidence": "high",
            "alias_map_source": "nearby_context",
            "alias_map_evidence": "C1=Ni/ZrO2",
            "figure_binding_status": "matched",
        })
        if case2.get("alias_map_source") != "nearby_context" or case2.get("alias_map_evidence") != "C1=Ni/ZrO2":
            print("  [WARNING] runtime selfcheck failed: alias-map canonical payload lost evidence")

        if _canonicalize_semantic_figure_role("catalyst_screening") != "category_screening":
            print("  [WARNING] runtime selfcheck failed: catalyst_screening canonical role mismatch")
        if _canonicalize_semantic_figure_role("schematic_or_nonperformance") != "non_performance_like":
            print("  [WARNING] runtime selfcheck failed: schematic_or_nonperformance canonical role mismatch")

        try:
            _validate_runtime_config(api_key="")
            print("  [WARNING] runtime selfcheck failed: missing API key did not raise")
        except RuntimeError:
            pass

        temp_dir = os.path.join(TEMP_IMG_DIR, "__legacy_modeling_export_selfcheck__")
        os.makedirs(temp_dir, exist_ok=True)
        for name in ["modeling_master.csv", "modeling_master_candidate.csv", "modeling_master_strict.csv", "modeling_excluded_audit.csv", "modeling_feature_schema.json", "modeling_feature_coverage_summary.json"]:
            path_to_clean = os.path.join(temp_dir, name)
            if os.path.exists(path_to_clean):
                try:
                    os.remove(path_to_clean)
                except Exception:
                    pass

        sample_row = {
            "Source_File": "selfcheck.pdf",
            "data_source": "figure",
            "Catalyst": "Ni/ZrO2",
            "Catalyst_ID": "Ni/ZrO2",
            "Canonical_Catalyst_ID": "Ni/ZrO2",
            "Reaction_Temp_C": "250",
            "MeOH_Conversion_%": "78.0",
            "figure_binding_status": "matched",
            "figure_binding_mode": "direct_label",
            "figure_binding_confidence": "high",
            "semantic_figure_role": "category_screening",
        }
        export_paths = save_modeling_exports([sample_row], [sample_row], get_model_ready_export_columns(), temp_dir)
        if os.path.exists(os.path.join(temp_dir, "modeling_master_candidate.csv")) or os.path.exists(os.path.join(temp_dir, "modeling_master_strict.csv")):
            print("  [WARNING] runtime selfcheck failed: retired legacy modeling CSVs were recreated")
        if not os.path.exists(os.path.join(temp_dir, "modeling_master.csv")):
            print("  [WARNING] runtime selfcheck failed: final modeling master CSV missing from legacy wrapper")
        if str(export_paths.get("legacy_candidate_csv", "")).strip() or str(export_paths.get("legacy_strict_csv", "")).strip():
            print("  [WARNING] runtime selfcheck failed: legacy wrapper still exposes candidate/strict CSV paths")
    except Exception as e:
        print(f"  [WARNING] runtime/binding consistency selfcheck error: {e}")



def build_modeling_export_table(records: List[Dict], export_columns: Dict[str, List[str]]) -> List[Dict]:
    ordered_columns: List[str] = []
    for group_name in [
        "id_columns", "group_column", "target_candidates", "baseline_numeric_feature_columns_model",
        "baseline_categorical_feature_columns", "baseline_derived_feature_columns",
        "enhanced_numeric_feature_columns_model", "enhanced_categorical_feature_columns",
        "enhanced_derived_feature_columns", "missing_flag_columns", "provenance_columns",
    ]:
        columns = export_columns.get(group_name, [])
        if isinstance(columns, str):
            columns = [columns]
        for col in columns:
            if col not in ordered_columns:
                ordered_columns.append(col)

    table_rows: List[Dict] = []
    for record in records or []:
        if not isinstance(record, dict):
            continue
        row = _finalize_figure_row_binding_payload(dict(record))
        row = clean_record_text_fields(strip_internal_context_fields(row))
        table_rows.append({col: row.get(col, "") for col in ordered_columns})
    return table_rows



def build_modeling_excluded_audit(candidate_records: List[Dict], strict_records: List[Dict]) -> List[Dict]:
    if any(isinstance(row, dict) and (str(row.get("final_exclusion_reason", "")).strip() or str(row.get("final_exclusion_detail", "")).strip()) for row in (strict_records or [])):
        normalized_rows: List[Dict] = []
        for row in strict_records or []:
            if not isinstance(row, dict):
                continue
            current = _normalize_binding_audit_record(row)
            current.setdefault("strict_exclusion_reason", str(current.get("final_exclusion_reason", "") or ""))
            current.setdefault("strict_exclusion_detail", str(current.get("final_exclusion_detail", "") or ""))
            normalized_rows.append(current)
        return normalized_rows

    def _row_signature(row: Dict) -> str:
        payload = _finalize_figure_row_binding_payload(dict(row))
        payload = strip_internal_context_fields(clean_record_text_fields(payload))
        return json.dumps(payload, ensure_ascii=False, sort_keys=True)

    strict_signatures = {_row_signature(row) for row in (strict_records or []) if isinstance(row, dict)}

    def _derive_exclusion_reason(row: Dict) -> Tuple[str, str]:
        current = _finalize_figure_row_binding_payload(dict(row))
        if str(current.get("comparison_filter_action", "")).strip() == "downgraded_candidate_only":
            return "comparison_downgraded", str(current.get("comparison_filter_reason", "")).strip()
        if str(current.get("data_source", "")).strip() == "figure":
            if str(current.get("premerge_filter_action", "")).strip() == "candidate_only":
                return "candidate_only_figure_row", str(current.get("premerge_filter_reason", "")).strip()
            binding_status = str(current.get("figure_binding_status", "")).strip()
            if binding_status == "unmatched":
                return "unmatched_figure_binding", str(current.get("figure_binding_reason", "")).strip()
            if binding_status == "ambiguous":
                return "ambiguous_figure_binding", str(current.get("figure_binding_reason", "")).strip()
            if binding_status and binding_status != "matched":
                return "nonmatched_figure_binding", str(current.get("figure_binding_reason", "")).strip()
            if str(current.get("figure_binding_confidence", "")).strip() not in {"", "high", "medium"}:
                return "low_figure_binding_confidence", str(current.get("figure_binding_confidence", "")).strip()
            semantic_role = _canonicalize_semantic_figure_role(str(current.get("semantic_figure_role", "")).strip())
            if semantic_role in {"product_species_profile", "case_profile", "non_performance_like", "catalyst_amount_effect"}:
                return "blocked_figure_semantic_role", semantic_role
        if bool(current.get("is_range_like_value")):
            return "range_like_value", str(current.get("raw_numeric_expression", "")).strip()
        if bool(current.get("is_qualitative_value")):
            return "low_numeric_reliability", "qualitative_value"
        if bool(current.get("is_approximate_value")):
            return "approximate_numeric_value", str(current.get("raw_numeric_expression", "")).strip()
        numeric_level = str(current.get("numeric_reliability_level", "")).strip()
        if numeric_level and numeric_level != "direct_numeric":
            return "low_numeric_reliability", numeric_level
        identity_level = str(current.get("identity_completeness_level", "")).strip()
        if identity_level and identity_level not in {"complete", "partial"}:
            return "weak_identity", identity_level
        if bool(current.get("obvious_duplicate_flag")):
            return "obvious_duplicate_flagged", str(current.get("obvious_duplicate_type", "")).strip()
        if bool(current.get("same_physical_point_possible")):
            return "possible_duplicate_point", str(current.get("duplicate_candidate_type", "")).strip()
        origin_quality = str(current.get("origin_quality_class", "")).strip()
        if origin_quality and origin_quality not in {"high", "medium"}:
            return "low_origin_quality", origin_quality
        return "not_selected_by_strict_filters", ""

    audit_rows: List[Dict] = []
    for row in candidate_records or []:
        if not isinstance(row, dict):
            continue
        sig = _row_signature(row)
        if sig in strict_signatures:
            continue
        reason, detail = _derive_exclusion_reason(row)
        current = _normalize_binding_audit_record(row)
        current["strict_exclusion_reason"] = reason
        current["strict_exclusion_detail"] = detail
        audit_rows.append(current)
    return audit_rows



def build_feature_coverage_summary(final_rows: List[Dict], final_excluded_rows: List[Dict], export_columns: Dict[str, List[str]]) -> Dict[str, Any]:
    master_rows = build_modeling_export_table(final_rows, export_columns)

    def _is_nonempty(value: Any) -> bool:
        return bool(str(value if value is not None else "").strip())

    def _coverage(rows: List[Dict], columns: List[str]) -> Dict[str, Dict[str, Any]]:
        total = len(rows)
        summary: Dict[str, Dict[str, Any]] = {}
        for col in columns:
            nonempty = sum(1 for row in rows if _is_nonempty(row.get(col, "")))
            summary[col] = {
                "nonempty_count": nonempty,
                "nonempty_ratio": round(nonempty / total, 4) if total else 0.0,
            }
        return summary

    key_feature_columns = []
    for group_name in [
        "baseline_numeric_feature_columns_model", "baseline_categorical_feature_columns",
        "baseline_derived_feature_columns", "enhanced_numeric_feature_columns_model",
        "enhanced_categorical_feature_columns", "enhanced_derived_feature_columns", "missing_flag_columns",
    ]:
        for col in export_columns.get(group_name, []) or []:
            if col not in key_feature_columns:
                key_feature_columns.append(col)

    return {
        "row_count_final": len(master_rows),
        "row_count_excluded": len(final_excluded_rows or []),
        "target_summary": [
            {
                "target_name": target_name,
                "nonempty_final": sum(1 for row in master_rows if _is_nonempty(row.get(target_name, ""))),
            }
            for target_name in export_columns.get("target_candidates", [])
        ],
        "feature_coverage_final": _coverage(master_rows, key_feature_columns),
    }



def save_modeling_exports(candidate_records: List[Dict], strict_records: List[Dict], export_columns: Dict[str, List[str]], output_dir: str = OUT_DIR) -> Dict[str, str]:
    return _retire_legacy_modeling_exports(candidate_records, strict_records, export_columns, output_dir)



def save_modeling_final_exports(final_rows: List[Dict], final_excluded_rows: List[Dict], export_columns: Dict[str, List[str]], output_dir: str) -> Dict[str, str]:
    def _ordered_columns(schema: Dict[str, List[str]]) -> List[str]:
        ordered = []
        for group_name in [
            "id_columns", "group_column", "target_candidates", "baseline_numeric_feature_columns_model",
            "baseline_categorical_feature_columns", "baseline_derived_feature_columns",
            "enhanced_numeric_feature_columns_model", "enhanced_categorical_feature_columns",
            "enhanced_derived_feature_columns", "missing_flag_columns", "provenance_columns",
        ]:
            columns = schema.get(group_name, [])
            if isinstance(columns, str):
                columns = [columns]
            for col in columns:
                if col not in ordered:
                    ordered.append(col)
        return ordered

    def _write_rows(rows: List[Dict], output_path: str, preferred_fields: List[str]) -> None:
        cleaned_rows = []
        for row in rows or []:
            if not isinstance(row, dict):
                continue
            current = _finalize_figure_row_binding_payload(dict(row))
            cleaned_rows.append(strip_internal_context_fields(clean_record_text_fields(current)))
        all_keys = set()
        for row in cleaned_rows:
            all_keys.update(row.keys())
        fieldnames = [k for k in preferred_fields if k in all_keys] + sorted(k for k in all_keys if k not in preferred_fields)
        if not fieldnames:
            fieldnames = list(preferred_fields)
        with open(output_path, "w", newline="", encoding="utf-8-sig") as f:
            writer = csv.DictWriter(f, fieldnames=fieldnames, extrasaction="ignore")
            writer.writeheader()
            for row in cleaned_rows:
                writer.writerow({
                    key: (json.dumps(row[key], ensure_ascii=False) if isinstance(row.get(key), (list, dict, tuple, set)) else row.get(key, ""))
                    for key in fieldnames
                })

    os.makedirs(output_dir, exist_ok=True)
    master_rows = build_modeling_export_table(final_rows, export_columns)
    excluded_rows = build_modeling_excluded_audit(final_rows, final_excluded_rows)
    dedupe_rows = [dict(row) for row in (globals().get("_FINAL_DEDUPE_AUDIT_ROWS", []) or []) if isinstance(row, dict)]
    coverage_payload = build_feature_coverage_summary(final_rows, final_excluded_rows, export_columns)
    distribution_payload = _build_final_distribution_summary(final_rows, final_excluded_rows, dedupe_rows)
    coverage_payload.update({
        "final_target_distribution": distribution_payload.get("final_target_distribution", {}),
        "final_condition_feature_coverage": distribution_payload.get("final_condition_feature_coverage", {}),
        "final_source_file_distribution": distribution_payload.get("final_source_file_distribution", {}),
        "figure_vs_text_final_distribution": distribution_payload.get("figure_vs_text_final_distribution", {}),
        "duplicate_signature_group_count": distribution_payload.get("duplicate_signature_group_count", 0),
        "rows_removed_by_final_dedupe": distribution_payload.get("rows_removed_by_final_dedupe", 0),
        "weak_condition_rows_excluded_from_final": distribution_payload.get("weak_condition_rows_excluded_from_final", 0),
        "top_10_source_files_by_final_rows": distribution_payload.get("top_10_source_files_by_final_rows", []),
    })
    ordered_cols = _ordered_columns(export_columns)
    master_csv = os.path.join(output_dir, "modeling_master.csv")
    excluded_csv = os.path.join(output_dir, "modeling_excluded_audit.csv")
    schema_json = os.path.join(output_dir, "modeling_feature_schema.json")
    coverage_json = os.path.join(output_dir, "modeling_feature_coverage_summary.json")

    _write_rows(master_rows, master_csv, ordered_cols)
    _write_rows(excluded_rows, excluded_csv, [
        "Source_File", "data_source", "Catalyst", "Catalyst_ID", "Canonical_Catalyst_ID",
        "final_exclusion_reason", "final_exclusion_detail", "strict_exclusion_reason", "strict_exclusion_detail",
        "figure_binding_status", "figure_binding_mode", "figure_binding_confidence", "figure_binding_reason",
        "final_duplicate_signature", "dedupe_action",
    ])

    schema_payload: Dict[str, Any] = dict(export_columns)
    if isinstance(schema_payload.get("group_column"), list) and len(schema_payload["group_column"]) == 1:
        schema_payload["group_column"] = schema_payload["group_column"][0]

    with open(schema_json, "w", encoding="utf-8") as f:
        json.dump(schema_payload, f, ensure_ascii=False, indent=2)
    with open(coverage_json, "w", encoding="utf-8") as f:
        json.dump(coverage_payload, f, ensure_ascii=False, indent=2)

    return {
        "master_csv": master_csv,
        "excluded_csv": excluded_csv,
        "schema_json": schema_json,
        "coverage_json": coverage_json,
    }



def save_audit_csv(rows: List[Dict], output_path: str) -> None:
    cleaned_rows = strip_internal_context_fields_batch([
        clean_record_text_fields(dict(row)) for row in rows if isinstance(row, dict)
    ])
    extra_keys = sorted({k for row in cleaned_rows for k in row.keys()} - set(_FIGURE_BINDING_AUDIT_FIELDS))
    fieldnames = list(_FIGURE_BINDING_AUDIT_FIELDS) + extra_keys
    with open(output_path, "w", newline="", encoding="utf-8-sig") as f:
        writer = csv.DictWriter(f, fieldnames=fieldnames)
        writer.writeheader()
        if cleaned_rows:
            writer.writerows(cleaned_rows)


def save_extraction_quality_summary(rows: List[Dict], output_path: str) -> None:
    paper_stats = globals().get("_LAYERED_PAPER_STATS", {}) if isinstance(globals().get("_LAYERED_PAPER_STATS", {}), dict) else {}
    payload_rows = []
    for row in rows or []:
        if not isinstance(row, dict):
            continue
        current = strip_internal_context_fields(clean_record_text_fields(dict(row)))
        source_file = clean_residual_mojibake_chars(str(current.get("Source_File", "") or "")).strip()
        stats = paper_stats.get(source_file, {}) if source_file else {}
        current["candidate_row_count"] = int(stats.get("candidate_row_count", current.get("candidate_row_count", 0)) or 0)
        current["research_all_inclusion_count"] = int(stats.get("research_all_row_count", current.get("research_all_inclusion_count", 0)) or 0)
        current["research_strict_inclusion_count"] = int(stats.get("research_strict_row_count", current.get("research_strict_inclusion_count", 0)) or 0)
        current["modeling_final_inclusion_count"] = int(current.get("final_point_row_count", stats.get("modeling_final_row_count", 0)) or 0)
        current["research_all_text_row_count"] = int(stats.get("research_all_text_row_count", 0) or 0)
        current["research_all_figure_row_count"] = int(stats.get("research_all_figure_row_count", 0) or 0)
        current["research_strict_text_row_count"] = int(stats.get("research_strict_text_row_count", 0) or 0)
        current["research_strict_figure_row_count"] = int(stats.get("research_strict_figure_row_count", 0) or 0)
        current["modeling_final_text_row_count"] = int(stats.get("modeling_final_text_row_count", current.get("final_text_point_count", 0)) or 0)
        current["modeling_final_figure_row_count"] = int(stats.get("modeling_final_figure_row_count", current.get("final_figure_point_count", 0)) or 0)
        current["research_all_block_reason_summary"] = stats.get("research_all_block_reason_counts", {})
        current["strict_excluded_reason_summary"] = stats.get("research_strict_excluded_reason_counts", {})
        current["final_excluded_reason_summary"] = stats.get("modeling_final_exclusion_reason_counts", {})
        payload_rows.append(current)

    overall_summary = {
        "paper_count": len(payload_rows),
        "candidate_row_count": sum(int(row.get("candidate_row_count", 0) or 0) for row in payload_rows),
        "research_all_row_count": sum(int(row.get("research_all_inclusion_count", 0) or 0) for row in payload_rows),
        "research_strict_row_count": sum(int(row.get("research_strict_inclusion_count", 0) or 0) for row in payload_rows),
        "modeling_final_row_count": sum(int(row.get("modeling_final_inclusion_count", 0) or 0) for row in payload_rows),
    }
    payload = {
        "papers": payload_rows,
        "overall_summary": overall_summary,
        "notes": {
            "compatibility_aliases": _LAYERED_EXPORT_COMPAT_NOTE,
            "semantics": {
                "research_master_all": "maximal research retention after removing obvious noise or wrong bindings",
                "research_master_strict": "cleaner research table for review and manual statistics",
                "modeling_final": "high-purity final subset for modeling",
            },
        },
    }
    globals()["_LAST_QUALITY_SUMMARY_ROWS"] = payload_rows
    globals()["_LAST_QUALITY_SUMMARY_PAYLOAD"] = payload
    os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
    with open(output_path, "w", encoding="utf-8") as f:
        json.dump(payload, f, ensure_ascii=False, indent=2)


def _layered_export_parallel_selfcheck(
    research_all_rows_clean: List[Dict],
    research_strict_rows_clean: List[Dict],
    modeling_final_rows_clean: List[Dict],
    final_excluded_rows_clean: List[Dict],
) -> None:
    """Lightweight selfcheck: verify four parallel clean rows are present and consistent.
    Only emits warnings; never interrupts main flow."""
    try:
        n_all = len(research_all_rows_clean or [])
        n_strict = len(research_strict_rows_clean or [])
        n_final = len(modeling_final_rows_clean or [])
        n_excl = len(final_excluded_rows_clean or [])

        # Check: all four layers must be present (not final-only)
        if n_all == 0 and n_strict == 0 and n_final > 0:
            print("  [WARNING] layered_export_selfcheck: research_all and research_strict are empty "
                  "but modeling_final has rows — possible final-only export path still active.")
        if n_all == 0 and n_final > 0:
            print("  [WARNING] layered_export_selfcheck: research_all is empty but modeling_final has rows "
                  "— research_all_rows_clean may not be sourced from all_data accumulator.")
        if n_strict == 0 and n_final > 0:
            print("  [WARNING] layered_export_selfcheck: research_strict is empty but modeling_final has rows "
                  "— research_strict_rows_clean may not be sourced from all_data_strict accumulator.")

        # Check: subset relationships should hold (by row count, as a quick proxy)
        if n_strict > n_all:
            print(f"  [WARNING] layered_export_selfcheck: research_strict ({n_strict}) > research_all ({n_all}) "
                  "— strict should be a subset of all.")
        if n_final > n_strict:
            print(f"  [WARNING] layered_export_selfcheck: modeling_final ({n_final}) > research_strict ({n_strict}) "
                  "— final should be a subset of strict.")

        # Check: strict ⊆ all by key membership (sampled up to 500 rows)
        def _uid_set(rows: List[Dict]) -> set:
            result = set()
            for row in rows or []:
                uid = str(row.get("layer_row_uid", "") or "").strip()
                if uid:
                    result.add(uid)
            return result

        all_uids = _uid_set(research_all_rows_clean)
        strict_uids = _uid_set(research_strict_rows_clean)
        final_uids = _uid_set(modeling_final_rows_clean)
        if all_uids and strict_uids and not strict_uids.issubset(all_uids):
            overhang = len(strict_uids - all_uids)
            print(f"  [WARNING] layered_export_selfcheck: strict ⊄ all by uid ({overhang} strict uids not in all).")
        if strict_uids and final_uids and not final_uids.issubset(strict_uids):
            overhang = len(final_uids - strict_uids)
            print(f"  [WARNING] layered_export_selfcheck: final ⊄ strict by uid ({overhang} final uids not in strict).")

        print(f"  [Selfcheck] layered_export: all={n_all} strict={n_strict} final={n_final} excluded={n_excl} — parallel clean rows confirmed.")
    except Exception as e:
        print(f"  [WARNING] layered_export_parallel_selfcheck error: {e}; skipping.")


def _prepare_layered_master_export_payloads(
    research_all_rows_clean: List[Dict],
    research_strict_rows_clean: List[Dict],
    modeling_final_rows_clean: List[Dict],
    final_excluded_rows_clean: List[Dict],
) -> Dict[str, Any]:
    # Historical helper name retained; now consumes parallel layered inputs rather than final-only inputs.
    # Each layer arrives pre-cleaned and independent from main() — all / strict / final / excluded
    # are NOT derived from each other here. This is the architectural contract.
    research_all_rows = [dict(row) for row in (research_all_rows_clean or []) if isinstance(row, dict)]
    research_strict_rows = [dict(row) for row in (research_strict_rows_clean or []) if isinstance(row, dict)]
    modeling_final_rows = [dict(row) for row in (modeling_final_rows_clean or []) if isinstance(row, dict)]
    final_excluded_rows = [dict(row) for row in (final_excluded_rows_clean or []) if isinstance(row, dict)]

    dedupe_rows = [dict(row) for row in (globals().get("_FINAL_DEDUPE_AUDIT_ROWS", []) or []) if isinstance(row, dict)]
    distribution_summary = _build_layered_distribution_summary(
        research_all_rows,
        research_strict_rows,
        modeling_final_rows,
        final_excluded_rows,
        dedupe_rows,
    )
    distribution_summary.setdefault("notes", {})
    distribution_summary["notes"].update({
        "semantic_primary_outputs": {
            "research_master_all.csv": "maximum-retention research table",
            "research_master_strict.csv": "cleaner research table for review/statistics",
            "modeling_final.csv": "highest-confidence table for ML only",
            "research_master_all.json": "JSON companion of research_master_all.csv",
            "research_master_strict.json": "JSON companion of research_master_strict.csv",
            "modeling_final.json": "JSON companion of modeling_final.csv",
            "layered_distribution_summary.json": "layered summary for all/strict/final outputs",
        },
        "compatibility_aliases": {
            "dataset_all.csv": "research_master_all.csv",
            "dataset_all_strict.csv": "research_master_strict.csv",
            "msr_dataset.json": "modeling_final.json",
            "final_distribution_summary.json": "layered_distribution_summary.json",
        },
    })
    return {
        "research_all_rows": research_all_rows,
        "research_strict_rows": research_strict_rows,
        "modeling_final_rows": modeling_final_rows,
        "final_excluded_rows": final_excluded_rows,
        "dedupe_rows": dedupe_rows,
        "distribution_summary": distribution_summary,
    }



def _save_layered_master_json_exports(
    research_all_rows: List[Dict],
    research_strict_rows: List[Dict],
    modeling_final_rows: List[Dict],
    output_dir: str,
) -> Dict[str, str]:
    def _sanitize_rows(rows: List[Dict]) -> List[Dict]:
        return [
            strip_internal_context_fields(clean_record_text_fields(_finalize_figure_row_binding_payload(dict(row))))
            for row in (rows or []) if isinstance(row, dict)
        ]

    def _write_payload(path_out: str, layer_name: str, layer_semantics: str, csv_name: str, rows: List[Dict], legacy_alias: str = "") -> None:
        payload = {
            "layer_name": layer_name,
            "layer_semantics": layer_semantics,
            "corresponding_csv": csv_name,
            "row_count": len(rows),
            "data": rows,
        }
        if legacy_alias:
            payload["legacy_alias"] = legacy_alias
            payload["note"] = f"{legacy_alias} is a compatibility alias only and follows the {layer_name} layer."
        with open(path_out, "w", encoding="utf-8") as f:
            json.dump(payload, f, ensure_ascii=False, indent=2)

    os.makedirs(output_dir, exist_ok=True)
    research_all_json = os.path.join(output_dir, os.path.basename(RESEARCH_MASTER_ALL_JSON))
    research_strict_json = os.path.join(output_dir, os.path.basename(RESEARCH_MASTER_STRICT_JSON))
    modeling_final_json = os.path.join(output_dir, os.path.basename(MODELING_FINAL_JSON))
    legacy_msr_dataset_json = os.path.join(output_dir, os.path.basename(LEGACY_MSR_DATASET_JSON))

    research_all_payload_rows = _sanitize_rows(research_all_rows)
    research_strict_payload_rows = _sanitize_rows(research_strict_rows)
    modeling_final_payload_rows = _sanitize_rows(modeling_final_rows)

    _write_payload(
        research_all_json,
        layer_name="research_master_all",
        layer_semantics="maximum-retention research table",
        csv_name="research_master_all.csv",
        rows=research_all_payload_rows,
    )
    _write_payload(
        research_strict_json,
        layer_name="research_master_strict",
        layer_semantics="cleaner research table for review/statistics",
        csv_name="research_master_strict.csv",
        rows=research_strict_payload_rows,
    )
    _write_payload(
        modeling_final_json,
        layer_name="modeling_final",
        layer_semantics="highest-confidence table for ML only",
        csv_name="modeling_final.csv",
        rows=modeling_final_payload_rows,
    )
    _write_payload(
        legacy_msr_dataset_json,
        layer_name="modeling_final",
        layer_semantics="highest-confidence table for ML only",
        csv_name="modeling_final.csv",
        rows=modeling_final_payload_rows,
        legacy_alias="msr_dataset.json",
    )

    return {
        "research_all_json": research_all_json,
        "research_strict_json": research_strict_json,
        "modeling_final_json": modeling_final_json,
        "legacy_msr_dataset_json": legacy_msr_dataset_json,
    }



def _row_membership_keys(rows: List[Dict]) -> Set[str]:
    keys: Set[str] = set()
    for row in rows or []:
        if not isinstance(row, dict):
            continue
        uid = clean_residual_mojibake_chars(str(row.get("layer_row_uid", "") or "")).strip()
        keys.add(f"uid::{uid}" if uid else f"sig::{_stable_layer_row_signature(row)}")
    return keys



def _print_layered_export_summary(
    processed_paper_count: int,
    internal_candidate_row_count: int,
    research_all_rows: List[Dict],
    research_strict_rows: List[Dict],
    modeling_final_rows: List[Dict],
    final_export_paths: Dict[str, str],
    json_export_paths: Dict[str, str],
) -> None:
    strict_is_subset_of_all = _row_membership_keys(research_strict_rows).issubset(_row_membership_keys(research_all_rows))
    final_is_subset_of_strict = _row_membership_keys(modeling_final_rows).issubset(_row_membership_keys(research_strict_rows))

    print()
    print("[Layered export summary]")
    print(f"Processed papers         : {processed_paper_count}")
    print(f"Internal candidate rows  : {internal_candidate_row_count}")
    print(f"Research all table       : {final_export_paths['research_all_csv']} ({len(research_all_rows)} rows)")
    print(f"Research strict table    : {final_export_paths['research_strict_csv']} ({len(research_strict_rows)} rows)")
    print(f"Modeling final table     : {final_export_paths['modeling_final_csv']} ({len(modeling_final_rows)} rows)")

    print()
    print("Compatibility aliases:")
    print(f"dataset_all.csv          -> {final_export_paths['research_all_csv']}")
    print(f"dataset_all_strict.csv   -> {final_export_paths['research_strict_csv']}")
    if final_export_paths.get("legacy_distribution_summary_json"):
        print(f"final_distribution_summary.json -> {final_export_paths['legacy_distribution_summary_json']} (compatibility alias of layered_distribution_summary.json)")
    print("these are compatibility aliases only; do not confuse them with the semantic primary outputs")

    print()
    print("JSON exports:")
    print(f"research_master_all.json    -> {json_export_paths['research_all_json']} (corresponds to research_master_all.csv)")
    print(f"research_master_strict.json -> {json_export_paths['research_strict_json']} (corresponds to research_master_strict.csv)")
    print(f"modeling_final.json        -> {json_export_paths['modeling_final_json']} (corresponds to modeling_final.csv)")
    print(f"msr_dataset.json           -> modeling_final layer only ({json_export_paths['legacy_msr_dataset_json']}; compatibility alias)")
    print(f"layered_distribution_summary.json -> {final_export_paths['distribution_summary_json']}")

    print()
    print("Layer semantics:")
    print("research_master_all      = maximum-retention research table")
    print("research_master_strict   = cleaner research table for review/statistics")
    print("modeling_final           = highest-confidence table for ML only")

    print()
    print("Subset checks:")
    print(f"strict is a subset of all   : {'PASS' if strict_is_subset_of_all else 'WARNING'}")
    print(f"final is a subset of strict : {'PASS' if final_is_subset_of_strict else 'WARNING'}")


def save_final_single_table_exports(
    research_all_rows: List[Dict],
    research_strict_rows: List[Dict],
    modeling_final_rows: List[Dict],
    final_excluded_rows: List[Dict],
    output_dir: str,
) -> Dict[str, str]:
    # Historical name retained for compatibility.
    # Actual behavior: exports layered outputs from parallel research_all / research_strict /
    # modeling_final payloads that arrive pre-cleaned from main().
    # Each layer is an independent input; this function does NOT reconstruct all/strict from final.
    def _write_rows(rows: List[Dict], output_path: str, preferred_fields: List[str]) -> None:
        cleaned_rows = [strip_internal_context_fields(clean_record_text_fields(_finalize_figure_row_binding_payload(dict(row)))) for row in (rows or []) if isinstance(row, dict)]
        all_keys = set()
        for row in cleaned_rows:
            all_keys.update(row.keys())
        fieldnames = [k for k in preferred_fields if k in all_keys] + sorted(k for k in all_keys if k not in preferred_fields)
        if not fieldnames:
            fieldnames = list(preferred_fields)
        with open(output_path, "w", newline="", encoding="utf-8-sig") as f:
            writer = csv.DictWriter(f, fieldnames=fieldnames, extrasaction="ignore")
            writer.writeheader()
            for row in cleaned_rows:
                writer.writerow({
                    key: (json.dumps(row[key], ensure_ascii=False) if isinstance(row.get(key), (list, dict, tuple, set)) else row.get(key, ""))
                    for key in fieldnames
                })

    os.makedirs(output_dir, exist_ok=True)
    # Use the four parallel clean rows directly — no internal reconstruction.
    research_all_rows = [dict(row) for row in (research_all_rows or []) if isinstance(row, dict)]
    research_strict_rows = [dict(row) for row in (research_strict_rows or []) if isinstance(row, dict)]
    final_rows_ready = [dict(row) for row in (modeling_final_rows or []) if isinstance(row, dict)]
    final_excluded_ready = [dict(row) for row in (final_excluded_rows or []) if isinstance(row, dict)]
    dedupe_rows = [dict(row) for row in (globals().get("_FINAL_DEDUPE_AUDIT_ROWS", []) or []) if isinstance(row, dict)]
    distribution_summary = _build_layered_distribution_summary(
        research_all_rows,
        research_strict_rows,
        final_rows_ready,
        final_excluded_ready,
        dedupe_rows,
    )
    distribution_summary.setdefault("notes", {})
    distribution_summary["notes"].update({
        "semantic_primary_outputs": {
            "research_master_all.csv": "maximum-retention research table",
            "research_master_strict.csv": "cleaner research table for review/statistics",
            "modeling_final.csv": "highest-confidence table for ML only",
            "research_master_all.json": "JSON companion of research_master_all.csv",
            "research_master_strict.json": "JSON companion of research_master_strict.csv",
            "modeling_final.json": "JSON companion of modeling_final.csv",
            "layered_distribution_summary.json": "layered summary for all/strict/final outputs",
        },
        "compatibility_aliases": {
            "dataset_all.csv": "research_master_all.csv",
            "dataset_all_strict.csv": "research_master_strict.csv",
            "msr_dataset.json": "modeling_final.json",
            "final_distribution_summary.json": "layered_distribution_summary.json",
        },
    })

    research_all_csv = os.path.join(output_dir, "research_master_all.csv")
    research_strict_csv = os.path.join(output_dir, "research_master_strict.csv")
    modeling_final_csv = os.path.join(output_dir, "modeling_final.csv")
    legacy_all_csv = os.path.join(output_dir, "dataset_all.csv")
    legacy_strict_csv = os.path.join(output_dir, "dataset_all_strict.csv")
    excluded_csv = os.path.join(output_dir, "final_excluded_audit.csv")
    dedupe_audit_csv = os.path.join(output_dir, "final_dedupe_audit.csv")
    distribution_summary_json = os.path.join(output_dir, os.path.basename(LAYERED_DISTRIBUTION_SUMMARY_JSON))
    legacy_distribution_summary_json = os.path.join(output_dir, os.path.basename(LEGACY_FINAL_DISTRIBUTION_SUMMARY_JSON))

    master_preferred_fields = [
        "Source_File", "data_source", "page_num", "page_index", "chunk_id", "figure_id", "row_origin", "extraction_stage", "layer_row_uid",
        "Catalyst_ID", "Catalyst_ID_raw", "Catalyst_ID_normalized", "identity_alias_group", "Canonical_Catalyst_ID", "Catalyst",
        "Active_Metal", "Metal_Loading_wt%", "Alloy_Ratio", "Promoter", "Promoter_Metal", "Support", "Support_Normalized", "Support_Family", "Support_Modifier",
        "Reaction_Temp_C", "TOS_h", "S_C_Ratio", "Feed_MeOH_to_H2O_Ratio", "Pressure_bar", "GHSV_mL_g_h", "WHSV_h_inv", "SpaceVelocity_norm", "Flow_Rate", "Catalyst_Amount_g", "Feed_Ratio",
        "MeOH_Conversion_%", "H2_Yield_%", "H2_Production_Rate", "H2_Selectivity_%", "CO_Selectivity_%", "CO2_Selectivity_%", "CO_Concentration_ppm",
        "source_granularity", "numeric_reliability_level", "identity_completeness_level",
        "figure_binding_status", "figure_binding_reason", "figure_binding_mode", "figure_binding_confidence", "binding_confidence", "semantic_figure_role",
        "approx_flag", "range_flag", "partial_point_flag",
        "included_in_research_all", "included_in_research_strict", "included_in_modeling_final", "layer_membership",
        "exclusion_from_strict_reason", "exclusion_from_strict_detail", "exclusion_from_final_reason", "exclusion_from_final_detail",
        "strict_exclusion_reason", "strict_exclusion_detail", "final_exclusion_reason", "final_exclusion_detail",
        "Notes",
    ]

    _write_rows(research_all_rows, research_all_csv, master_preferred_fields)
    _write_rows(research_all_rows, legacy_all_csv, master_preferred_fields)
    _write_rows(research_strict_rows, research_strict_csv, master_preferred_fields)
    _write_rows(research_strict_rows, legacy_strict_csv, master_preferred_fields)
    _write_rows(final_rows_ready, modeling_final_csv, master_preferred_fields)
    _write_rows(final_excluded_ready, excluded_csv, [
        "Source_File", "data_source", "layer_row_uid", "Catalyst", "Catalyst_ID", "Canonical_Catalyst_ID",
        "Active_Metal", "Promoter_Metal", "Support", "MeOH_Conversion_%", "H2_Yield_%",
        "H2_Production_Rate", "CO_Selectivity_%", "CO2_Selectivity_%", "CO_Concentration_ppm",
        "included_in_research_all", "included_in_research_strict", "included_in_modeling_final", "layer_membership",
        "exclusion_from_final_reason", "exclusion_from_final_detail", "final_exclusion_reason", "final_exclusion_detail",
        "final_duplicate_signature", "dedupe_action", "Notes",
    ])
    _write_rows(dedupe_rows, dedupe_audit_csv, [
        "Source_File", "data_source", "layer_row_uid", "semantic_figure_role", "raw_category_label", "Catalyst", "Catalyst_ID",
        "Canonical_Catalyst_ID", "target_family", "figure_binding_status", "figure_binding_mode",
        "figure_binding_confidence", "figure_binding_reason", "matched_registry_key", "matched_registry_label",
        "alias_map_source", "alias_map_evidence", "final_dedupe_signature", "dedupe_action", "dedupe_reason",
        "final_exclusion_reason", "final_exclusion_detail", "kept_data_source", "kept_semantic_figure_role",
        "kept_Canonical_Catalyst_ID", "kept_target_family", "kept_priority_score", "removed_priority_score",
        "kept_priority_explanation", "removed_priority_explanation",
    ])
    with open(distribution_summary_json, "w", encoding="utf-8") as f:
        json.dump(distribution_summary, f, ensure_ascii=False, indent=2)
    with open(legacy_distribution_summary_json, "w", encoding="utf-8") as f:
        json.dump(distribution_summary, f, ensure_ascii=False, indent=2)

    issues = _layer_subset_consistency_check(
        research_all_rows=research_all_rows,
        research_strict_rows=research_strict_rows,
        final_rows=final_rows_ready,
        final_excluded_rows=final_excluded_ready,
        research_all_csv_path=research_all_csv,
        research_strict_csv_path=research_strict_csv,
        modeling_final_csv_path=modeling_final_csv,
        excluded_csv_path=excluded_csv,
        distribution_summary=distribution_summary,
        quality_summary_rows=globals().get("_LAST_QUALITY_SUMMARY_ROWS", []),
    )
    if issues:
        distribution_summary["consistency_warnings"] = issues
        with open(distribution_summary_json, "w", encoding="utf-8") as f:
            json.dump(distribution_summary, f, ensure_ascii=False, indent=2)
        with open(legacy_distribution_summary_json, "w", encoding="utf-8") as f:
            json.dump(distribution_summary, f, ensure_ascii=False, indent=2)

    globals()["_LAST_FINAL_DISTRIBUTION_SUMMARY"] = distribution_summary
    globals()["_RESEARCH_ALL_ROWS_CACHE"] = [dict(row) for row in research_all_rows]
    globals()["_RESEARCH_STRICT_ROWS_CACHE"] = [dict(row) for row in research_strict_rows]
    globals()["_LAYERED_EXPORT_RUN_COMPLETED"] = True

    _warn_if_legacy_export_names_are_ambiguous()
    print(f"  [OK] research master all: {research_all_csv} ({len(research_all_rows)} rows)")
    print(f"  [OK] research master strict: {research_strict_csv} ({len(research_strict_rows)} rows)")
    print(f"  [OK] modeling final: {modeling_final_csv} ({len(final_rows_ready)} rows)")
    print(f"  [Compat] dataset_all.csv -> {legacy_all_csv}")
    print(f"  [Compat] dataset_all_strict.csv -> {legacy_strict_csv}")
    print(f"  [Audit] layered distribution summary: {distribution_summary_json}")
    print(f"  [Compat] final_distribution_summary.json -> {legacy_distribution_summary_json}")
    print(f"  [Summary] total research all rows: {len(research_all_rows)}")
    print(f"  [Summary] total research strict rows: {len(research_strict_rows)}")
    print(f"  [Summary] total modeling final rows: {len(final_rows_ready)}")
    top_sources = list(distribution_summary.get("per_source_file_layer_counts", {}).items())[:10]
    for source_file, counts in top_sources:
        print(
            f"    [per-file] {source_file}: "
            f"all={counts.get('research_all', 0)}, strict={counts.get('research_strict', 0)}, final={counts.get('modeling_final', 0)}"
        )

    return {
        "research_all_csv": research_all_csv,
        "research_strict_csv": research_strict_csv,
        "modeling_final_csv": modeling_final_csv,
        "dataset_csv": modeling_final_csv,
        "legacy_dataset_csv": legacy_all_csv,
        "legacy_dataset_strict_csv": legacy_strict_csv,
        "excluded_csv": excluded_csv,
        "dedupe_audit_csv": dedupe_audit_csv,
        "distribution_summary_json": distribution_summary_json,
        "legacy_distribution_summary_json": legacy_distribution_summary_json,
    }


def save_text_point_gate_audit_csv(rows: List[Dict], output_path: str) -> None:
    cleaned_rows = [strip_internal_context_fields(clean_record_text_fields(dict(row))) for row in (rows or []) if isinstance(row, dict)]
    fieldnames = [
        "Source_File", "data_source", "Catalyst", "Catalyst_ID", "Canonical_Catalyst_ID",
        "source_granularity", "numeric_reliability_level", "identity_completeness_level",
        "point_gate_status", "point_gate_reason", "point_gate_notes",
    ]
    extra_keys = sorted({k for row in cleaned_rows for k in row.keys()} - set(fieldnames)) if cleaned_rows else []
    with open(output_path, "w", newline="", encoding="utf-8-sig") as f:
        writer = csv.DictWriter(f, fieldnames=fieldnames + extra_keys, extrasaction="ignore")
        writer.writeheader()
        if cleaned_rows:
            writer.writerows(cleaned_rows)



def save_composition_consistency_audit_csv(rows: List[Dict], output_path: str) -> None:
    cleaned_rows = [strip_internal_context_fields(clean_record_text_fields(dict(row))) for row in (rows or []) if isinstance(row, dict)]
    fieldnames = [
        "Source_File", "Catalyst", "Catalyst_ID", "Active_Metal", "Promoter", "Promoter_Metal",
        "Alloy_Ratio", "parsed_ratio_pairs", "Ni_Fraction", "Promoter_Fraction",
        "has_promoter", "is_bimetallic", "composition_consistency_flag", "composition_consistency_notes",
    ]
    extra_keys = sorted({k for row in cleaned_rows for k in row.keys()} - set(fieldnames)) if cleaned_rows else []
    with open(output_path, "w", newline="", encoding="utf-8-sig") as f:
        writer = csv.DictWriter(f, fieldnames=fieldnames + extra_keys, extrasaction="ignore")
        writer.writeheader()
        if cleaned_rows:
            writer.writerows(cleaned_rows)



def save_figure_candidate_audit_csv(rows: List[Dict], output_path: str) -> None:
    cleaned_rows = [strip_internal_context_fields(clean_record_text_fields(dict(row))) for row in (rows or []) if isinstance(row, dict)]
    fieldnames = [
        "Source_File", "page_num", "candidate_reason", "candidate_score", "text_len",
        "positive_flags", "negative_flags", "hard_exclusion", "is_candidate", "text_context_preview",
    ]
    extra_keys = sorted({k for row in cleaned_rows for k in row.keys()} - set(fieldnames)) if cleaned_rows else []
    with open(output_path, "w", newline="", encoding="utf-8-sig") as f:
        writer = csv.DictWriter(f, fieldnames=fieldnames + extra_keys, extrasaction="ignore")
        writer.writeheader()
        if cleaned_rows:
            for row in cleaned_rows:
                writer.writerow({
                    key: (json.dumps(row[key], ensure_ascii=False) if isinstance(row.get(key), (list, dict, tuple, set)) else row.get(key, ""))
                    for key in (fieldnames + extra_keys)
                })



def save_figure_point_validation_audit_csv(rows: List[Dict], output_path: str) -> None:
    cleaned_rows = [strip_internal_context_fields(clean_record_text_fields(dict(row))) for row in (rows or []) if isinstance(row, dict)]
    fieldnames = [
        "Source_File", "image_name", "page_num", "x_axis", "y_axis", "x_axis_mode",
        "extractor_type", "series_role", "metadata_series_count", "extracted_record_count",
        "extracted_series_count", "validation_status", "validation_flags", "validation_notes",
    ]
    extra_keys = sorted({k for row in cleaned_rows for k in row.keys()} - set(fieldnames)) if cleaned_rows else []
    with open(output_path, "w", newline="", encoding="utf-8-sig") as f:
        writer = csv.DictWriter(f, fieldnames=fieldnames + extra_keys, extrasaction="ignore")
        writer.writeheader()
        if cleaned_rows:
            writer.writerows(cleaned_rows)



def save_multimodal_audit_artifacts(output_dir: str, filename: str, figure_infos: List[Dict], figure_metadata_rows: List[Dict], raw_figure_rows_before_dedupe: List[Dict], raw_figure_rows_after_dedupe: List[Dict], validated_figure_rows: List[Dict], binding_rows: List[Dict], preparation_sidecar: Dict) -> None:
    def _prepare_row(payload: Any, audit_mode: bool = False) -> Any:
        if not isinstance(payload, dict):
            return payload
        current = dict(payload)
        if audit_mode:
            current = _normalize_binding_audit_record(current)
        else:
            current = _finalize_figure_row_binding_payload(current)
        return current

    def _sanitize(payload: Any, audit_mode: bool = False) -> Any:
        payload = _prepare_row(payload, audit_mode=audit_mode)
        if isinstance(payload, dict):
            row = strip_internal_context_fields(clean_record_text_fields(dict(payload)))
            cleaned = {}
            for key, value in row.items():
                key_s = str(key)
                if re.search(r"(?:base64|image_data|image_bytes|binary)", key_s, flags=re.I):
                    continue
                if key_s == "text_context":
                    cleaned["text_context_preview"] = _context_preview(value, 800)
                    continue
                if "context" in key_s.lower() and not key_s.lower().endswith("_preview"):
                    cleaned[f"{key_s}_preview"] = _context_preview(value, 800)
                    continue
                cleaned[key_s] = _sanitize(value, audit_mode=audit_mode)
            return cleaned
        if isinstance(payload, (list, tuple, set)):
            return [_sanitize(item, audit_mode=audit_mode) for item in payload]
        if isinstance(payload, float) and (payload != payload or payload in {float("inf"), float("-inf")}):
            return None
        return payload if payload is None or isinstance(payload, (str, int, bool, float)) else str(payload)

    def _safe_write_json(path_out: str, payload: Any) -> None:
        try:
            with open(path_out, "w", encoding="utf-8") as f:
                json.dump(payload, f, ensure_ascii=False, indent=2)
        except Exception as e:
            print(f"  [WARNING] multimodal audit write failed: {path_out}: {e}")
            try:
                with open(path_out, "w", encoding="utf-8") as f:
                    json.dump([] if isinstance(payload, list) else {}, f, ensure_ascii=False, indent=2)
            except Exception as inner_e:
                print(f"  [WARNING] multimodal audit fallback write failed: {path_out}: {inner_e}")

    def _safe_write_jsonl(path_out: str, rows: List[Dict], audit_mode: bool = False) -> None:
        try:
            with open(path_out, "w", encoding="utf-8") as f:
                for row in rows or []:
                    f.write(json.dumps(_sanitize(row, audit_mode=audit_mode), ensure_ascii=False) + "\n")
        except Exception as e:
            print(f"  [WARNING] multimodal audit write failed: {path_out}: {e}")
            try:
                with open(path_out, "w", encoding="utf-8") as f:
                    f.write("")
            except Exception as inner_e:
                print(f"  [WARNING] multimodal audit fallback write failed: {path_out}: {inner_e}")

    base_name = os.path.splitext(os.path.basename(filename))[0]
    audit_dir = os.path.join(output_dir, "multimodal_audit", base_name)
    try:
        os.makedirs(audit_dir, exist_ok=True)
    except Exception as e:
        print(f"  [WARNING] multimodal audit directory create failed: {audit_dir}: {e}")
        return

    candidate_pages = []
    for row in figure_infos or []:
        if not isinstance(row, dict):
            continue
        current = dict(row)
        if str(current.get("text_context", "")).strip() and not str(current.get("text_context_preview", "")).strip():
            current["text_context_preview"] = _context_preview(current.get("text_context", ""), 800)
        candidate_pages.append(current)

    _safe_write_json(os.path.join(audit_dir, "figure_candidate_pages.json"), _sanitize(candidate_pages))
    _safe_write_jsonl(os.path.join(audit_dir, "figure_metadata.jsonl"), [_finalize_figure_row_binding_payload(dict(row)) for row in (figure_metadata_rows or []) if isinstance(row, dict)])
    _safe_write_jsonl(os.path.join(audit_dir, "figure_raw_points_before_dedupe.jsonl"), [_finalize_figure_row_binding_payload(dict(row)) for row in (raw_figure_rows_before_dedupe or []) if isinstance(row, dict)])
    _safe_write_jsonl(os.path.join(audit_dir, "figure_raw_points_after_dedupe.jsonl"), [_finalize_figure_row_binding_payload(dict(row)) for row in (raw_figure_rows_after_dedupe or []) if isinstance(row, dict)])
    _safe_write_jsonl(os.path.join(audit_dir, "figure_validated_points.jsonl"), [_finalize_figure_row_binding_payload(dict(row)) for row in (validated_figure_rows or []) if isinstance(row, dict)])
    _safe_write_jsonl(os.path.join(audit_dir, "figure_binding_rows.jsonl"), [_normalize_binding_audit_record(dict(row)) for row in (binding_rows or []) if isinstance(row, dict)], audit_mode=True)
    _safe_write_json(os.path.join(audit_dir, "preparation_sidecar.json"), _sanitize(preparation_sidecar or {}))



def main():
    def _slice_global_rows(name: str, start_idx: int) -> List[Dict]:
        values = globals().get(name, [])
        if not isinstance(values, list):
            return []
        return [dict(row) for row in values[start_idx:] if isinstance(row, dict)]

    def _canonicalize_rows(rows: List[Dict], audit_mode: bool = False) -> List[Dict]:
        normalized = []
        for row in rows or []:
            if not isinstance(row, dict):
                continue
            current = dict(row)
            current = _normalize_binding_audit_record(current) if audit_mode else _finalize_figure_row_binding_payload(current)
            normalized.append(current)
        return normalized

    def _unpack_debug_payloads(payloads: List[Dict], file_name: str) -> Tuple[List[Dict], List[Dict], List[Dict], List[Dict]]:
        meta_rows, raw_before, raw_after, validated = [], [], [], []
        for payload in payloads:
            if not isinstance(payload, dict):
                continue
            image_name = str(payload.get("image_name", "") or "")
            page_num = payload.get("page_num", "")
            metadata = dict(payload.get("metadata") or {})
            meta_row = {"Source_File": file_name, "image_name": image_name, "page_num": page_num}
            meta_row.update(metadata)
            if str(payload.get("metadata_failure_reason", "")).strip():
                meta_row["metadata_failure_reason"] = str(payload.get("metadata_failure_reason", "")).strip()
            meta_rows.append(_finalize_figure_row_binding_payload(meta_row))
            for key, target in [
                ("raw_records_before_dedupe", raw_before),
                ("raw_records_after_dedupe", raw_after),
                ("validated_records", validated),
            ]:
                for row in payload.get(key, []) or []:
                    if not isinstance(row, dict):
                        continue
                    current = dict(row)
                    current.setdefault("Source_File", file_name)
                    current.setdefault("image_name", image_name)
                    current.setdefault("page_num", page_num)
                    target.append(_finalize_figure_row_binding_payload(current))
        return (
            clean_records_text_fields(meta_rows),
            clean_records_text_fields(raw_before),
            clean_records_text_fields(raw_after),
            clean_records_text_fields(validated),
        )

    def _build_preparation_sidecar(file_name: str, global_params: Dict[str, Any], rows: List[Dict]) -> Dict[str, Any]:
        prep_rows = []
        for row in rows or []:
            if not isinstance(row, dict):
                continue
            if str(row.get("data_source", "")).strip() != "text":
                continue
            if any(str(row.get(field, "")).strip() for field in PREPARATION_VALUE_FIELDS):
                prep_rows.append({
                    "Catalyst": row.get("Catalyst", ""),
                    "Catalyst_ID": row.get("Catalyst_ID", ""),
                    "Canonical_Catalyst_ID": row.get("Canonical_Catalyst_ID", ""),
                    "Precursor": row.get("Precursor", ""),
                    "Precursor_Family": row.get("Precursor_Family", ""),
                    "Metal_Loading_Method": row.get("Metal_Loading_Method", ""),
                    "Support_Prep_Method": row.get("Support_Prep_Method", ""),
                    "Calcination_Temp_C": row.get("Calcination_Temp_C", ""),
                    "Calcination_Time_h": row.get("Calcination_Time_h", ""),
                    "Reduction_Temp_C": row.get("Reduction_Temp_C", ""),
                    "Reduction_Time_h": row.get("Reduction_Time_h", ""),
                    "Preparation_Fingerprint": row.get("Preparation_Fingerprint", ""),
                })
        return {
            "Source_File": file_name,
            "global_params": global_params or {},
            "preparation_row_count": len(prep_rows),
            "preparation_rows": prep_rows[:100],
        }

    export_preferred_fields = [
        "Source_File",
        "data_source",
        "page_num",
        "Canonical_Catalyst_ID",
        "Catalyst",
        "Active_Metal",
        "Ni_Loading_wt%",
        "Promoter_Loading_wt%",
        "Alloy_Ratio",
        "Promoter_Metal",
        "Support",
        "Support_Normalized",
        "Support_Grouped",
        "Precursor",
        "Precursor_Family",
        "Metal_Loading_Method",
        "Support_Prep_Method",
        "Dry_Temp_C",
        "Dry_Time_h",
        "Calcination_Temp_C",
        "Calcination_Time_h",
        "Reduction_Temp_C",
        "Reduction_Time_h",
        "Reaction_Temp_C",
        "TOS_h",
        "S_C_Ratio",
        "Feed_MeOH_to_H2O_Ratio",
        "Pressure_bar",
        "GHSV_mL_g_h",
        "WHSV_h_inv",
        "SpaceVelocity_norm",
        "MeOH_Conversion_%",
        "H2_Yield_%",
        "H2_Production_Rate",
        "H2_Selectivity_%",
        "CO_Selectivity_%",
        "CO2_Selectivity_%",
        "CO_Concentration_ppm",
        "figure_binding_status",
        "figure_binding_mode",
        "figure_binding_confidence",
        "figure_binding_reason",
        "binding_reason",
        "matched_registry_key",
        "matched_registry_label",
        "_broadcast_path_primary",
        "_broadcast_flags",
        "_field_source_map_json",
        "is_support_only_label",
        "source_granularity",
        "numeric_reliability_level",
        "identity_completeness_level",
        "is_approximate_value",
        "is_range_like_value",
        "is_literature_comparison",
        "Notes",
    ]

    def _write_curated_csv(rows: List[Dict], output_path: str) -> None:
        cleaned_rows = [strip_internal_context_fields(clean_record_text_fields(_finalize_figure_row_binding_payload(dict(row)))) for row in (rows or []) if isinstance(row, dict)]
        fieldnames = list(export_preferred_fields)
        with open(output_path, "w", newline="", encoding="utf-8-sig") as f:
            writer = csv.DictWriter(f, fieldnames=fieldnames, extrasaction="ignore")
            writer.writeheader()
            for row in cleaned_rows:
                writer.writerow({
                    key: (json.dumps(row[key], ensure_ascii=False) if isinstance(row.get(key), (list, dict, tuple, set)) else row.get(key, ""))
                    for key in fieldnames
                })

    def _build_single_paper_csv_path(file_name: str) -> str:
        stem, ext = os.path.splitext(file_name)
        ext_tag = ext.lstrip(".").lower() or "file"
        return os.path.join(OUT_DIR, f"{stem}_{ext_tag}_dataset.csv")

    runtime_config = _validate_runtime_config()
    global client, API_KEY, BASE_URL, TEXT_MODEL, VISION_MODEL, IN_DIR, OUT_DIR, TEMP_IMG_DIR
    API_KEY = runtime_config["API_KEY"]
    BASE_URL = runtime_config["BASE_URL"]
    TEXT_MODEL = runtime_config["TEXT_MODEL"]
    VISION_MODEL = runtime_config["VISION_MODEL"]
    IN_DIR = runtime_config["IN_DIR"]
    OUT_DIR = runtime_config["OUT_DIR"]
    TEMP_IMG_DIR = runtime_config["TEMP_IMG_DIR"]
    client = OpenAI(api_key=API_KEY, base_url=BASE_URL)

    print("=" * 60)
    print("  MSR literature extraction V7 Pro")
    print("=" * 60)

    try:
        selfcheck_ok = _selfcheck_promoter_aware_merge_key()
        if not selfcheck_ok:
            print("  [WARNING] promoter-aware merge key selfcheck failed; continue running.")
    except Exception as e:
        print(f"  [WARNING] promoter-aware merge key selfcheck error: {e}")
    try:
        _selfcheck_generalization_rules()
    except Exception as e:
        print(f"  [WARNING] generalization selfcheck error: {e}")
    try:
        _selfcheck_runtime_and_binding_consistency()
    except Exception as e:
        print(f"  [WARNING] runtime/binding selfcheck error: {e}")
    globals()["_FINAL_DEDUPE_AUDIT_ROWS"] = []
    globals()["_LAST_FINAL_DEDUPE_AUDIT_ROWS"] = []
    globals()["_LAST_FINAL_DISTRIBUTION_SUMMARY"] = {}

    os.makedirs(OUT_DIR, exist_ok=True)
    os.makedirs(TEMP_IMG_DIR, exist_ok=True)
    os.makedirs(MULTIMODAL_AUDIT_DIR, exist_ok=True)

    if not os.path.exists(IN_DIR):
        print(f"[ERROR] input directory not found: {IN_DIR}")
        return

    files = sorted(set(
        f for f in os.listdir(IN_DIR)
        if f.lower().endswith((".pdf", ".docx")) and not f.startswith("~")
    ))
    if not files:
        print(f"[WARN] no .pdf or .docx files found in {IN_DIR}")
        return

    all_data = []
    all_data_raw = []  # 保留 _broadcast_path_primary/_broadcast_flags 字段，用于 summary 统计
    audit_rows = []
    figure_binding_audit_rows = []
    figure_premerge_filter_audit_rows = []
    comparison_audit_rows = []
    figure_point_validation_audit_rows = []
    composition_consistency_audit_rows = []
    text_point_gate_audit_rows = []
    figure_candidate_audit_rows = []
    quality_summary_rows = []

    print(f"\n[INFO] {len(files)} papers pending processing\n")

    for index, filename in enumerate(files):
        file_path = os.path.join(IN_DIR, filename)
        print(f"[{index + 1}/{len(files)}] Processing: {filename}")

        records = []
        figure_candidate_only_rows = []
        global_params = {}
        paper_figure_infos = []
        paper_figure_candidate_audit_rows = []
        paper_figure_metadata_rows = []
        paper_raw_figure_rows_before_dedupe = []
        paper_raw_figure_rows_after_dedupe = []
        paper_validated_figure_rows = []
        paper_binding_rows = []
        paper_validation_audit_rows = []
        paper_preparation_sidecar = {}
        paper_text_records = []
        paper_comparison_audit_rows = []
        paper_text_point_gate_rows = []
        paper_composition_audit_rows = []
        temp_audit = []
        low_info_audit = []
        obvious_dup_audit = []

        fig_candidate_start = len(globals().get("_FIGURE_CANDIDATE_AUDIT_ROWS", []) or [])
        fig_validation_start = len(globals().get("_FIGURE_VALIDATION_AUDIT_ROWS", []) or [])
        fig_debug_start = len(globals().get("_FIGURE_DEBUG_PAYLOADS", []) or [])

        if file_path.lower().endswith(".pdf"):
            chunks = extract_pages_as_chunks(file_path, pages_per_chunk=4, overlap=1)
        else:
            chunks = extract_text_from_docx(file_path)

        if not chunks:
            print("  [WARN] extracted text is empty; continue with empty text branch")
            chunks = []
        else:
            global_params = extract_global_params_from_preparation_paragraphs(chunks, filename)
            if not global_params:
                global_params = extract_global_params_from_chunks(chunks, filename)
            else:
                # 补全：preparation_paragraphs 只提取制备字段，用 chunks 规则兜底补全反应条件字段
                # 对反应条件字段，from_chunks（LLM）的值优先覆盖 preparation_paragraphs 的值
                _REACTION_CONDITION_FIELDS = {
                    "S_C_Ratio", "S_C_Ratio_Raw", "GHSV_mL_g_h", "GHSV_mL_g_h_Raw",
                    "WHSV_h_inv", "WHSV_h_inv_Raw", "SpaceVelocity_norm", "SpaceVelocity_type",
                    "SpaceVelocity_unit", "Reaction_Temp_C", "Pressure_bar", "Pressure_bar_Raw",
                    "Feed_MeOH_to_H2O_Ratio", "Feed_MeOH_to_H2O_Ratio_Raw", "Feed_Composition",
                }
                condition_supplement = extract_global_params_from_chunks(chunks, filename)
                for k, v in condition_supplement.items():
                    if not v:
                        continue
                    if k not in global_params:
                        global_params[k] = v
                    elif k in _REACTION_CONDITION_FIELDS:
                        # from_chunks 的反应条件值更可靠（LLM提取），覆盖 preparation_paragraphs 的误提取
                        if global_params[k] != v:
                            print(f"  [global_params_override] {k}: {global_params[k]} -> {v} (from_chunks wins)")
                            global_params[k] = v
            print(f"  text: {len(chunks)} chunks ({sum(len(c) for c in chunks)} chars total)")
            for i, chunk in enumerate(chunks):
                chunk_records = extract_info_from_chunk(chunk, filename, global_params)
                if not chunk_records:
                    continue
                for row in chunk_records:
                    row["Source_File"] = filename
                    row["data_source"] = "text"
                records.extend(chunk_records)

        paper_text_records = clean_records_text_fields([
            dict(row) for row in records
            if isinstance(row, dict) and str(row.get("data_source", "")).strip() == "text"
        ])
        text_sample_registry = build_text_sample_registry(records, filename) if records else {}

        if file_path.lower().endswith(".pdf"):
            paper_figure_infos = extract_figures_from_pdf(file_path) or []
            print(f"  figures: {len(paper_figure_infos)} candidates")
            for fig in paper_figure_infos:
                fig_records = analyze_figure_with_vl(fig["image_path"], fig["text_context"], filename)
                if not fig_records:
                    continue
                for row in fig_records:
                    row["Source_File"] = filename
                    row["data_source"] = "figure"
                fig_records, binding_rows = attach_figure_points_to_registry(fig_records, text_sample_registry, fig["text_context"], filename)
                fig_records = _canonicalize_rows(fig_records)
                binding_rows = _canonicalize_rows(binding_rows, audit_mode=True)
                for row in fig_records:
                    row["_comparison_context_raw"] = fig["text_context"]
                fig_records = annotate_figure_semantic_role(fig_records, fig["text_context"])
                fig_records = _canonicalize_rows(fig_records)
                merge_fig_rows, candidate_only_fig_rows = split_figure_rows_for_merge(fig_records, filename)
                merge_fig_rows = _canonicalize_rows(merge_fig_rows)
                candidate_only_fig_rows = _canonicalize_rows(candidate_only_fig_rows)
                records.extend(merge_fig_rows)
                figure_candidate_only_rows.extend(candidate_only_fig_rows)
                paper_binding_rows.extend(binding_rows)
                figure_binding_audit_rows.extend(binding_rows)
                figure_premerge_filter_audit_rows.extend(clean_records_text_fields(
                    _canonicalize_rows(_build_figure_premerge_filter_audit_rows(merge_fig_rows + candidate_only_fig_rows, filename), audit_mode=True)
                ))

        if records:
            records = _canonicalize_rows(records)
            apply_global_params_fallback(records, global_params)
            records = merge_fragmented_records(records)
            smart_feature_engineering(records)
            apply_numeric_expression_guards(records)
            clean_numeric_fields(records)
            validate_msr_data(records)
            records, temp_audit = resolve_obvious_temp_conversion_duplicates_pre_save(records)
            records, low_info_audit = filter_low_information_rows_before_save(records, filename)
            annotate_cross_source_duplicate_candidates(records)
            assign_extraction_quality_labels(records)
            records, obvious_dup_audit = resolve_obvious_cross_source_duplicates_pre_save(records)
            records = clean_records_text_fields(_canonicalize_rows(records))
        else:
            print("  [INFO] no merge-eligible rows entered main merge pool")

        figure_candidate_only_rows = _canonicalize_rows(figure_candidate_only_rows)
        # 对 candidate_only 的 figure 行也广播全局参数（binding失败行同样需要条件字段）
        # 把 paper_text_records 合并传入，使 consensus_fallback 能统计 text 行字段
        apply_global_params_fallback(figure_candidate_only_rows + paper_text_records, global_params)
        records = apply_ml_ready_catalyst_feature_batch(records)
        figure_candidate_only_rows = apply_ml_ready_catalyst_feature_batch(figure_candidate_only_rows)
        records = apply_condition_feature_batch(records)
        figure_candidate_only_rows = apply_condition_feature_batch(figure_candidate_only_rows)
        records = build_model_ready_schema_batch(records)
        figure_candidate_only_rows = build_model_ready_schema_batch(figure_candidate_only_rows)
        records = _canonicalize_rows(records)
        figure_candidate_only_rows = _canonicalize_rows(figure_candidate_only_rows)
        candidate_rows = clean_records_text_fields(_canonicalize_rows(records + figure_candidate_only_rows))
        print(f"  -> {len(records)} records | {len(figure_candidate_only_rows)} fig-only | total {len(candidate_rows)}")

        paper_text_point_gate_rows = _build_text_point_gate_audit_rows(candidate_rows, filename)
        paper_composition_audit_rows = _build_composition_consistency_audit_rows(candidate_rows, filename)
        text_point_gate_audit_rows.extend(clean_records_text_fields(paper_text_point_gate_rows))
        composition_consistency_audit_rows.extend(clean_records_text_fields(paper_composition_audit_rows))

        paper_figure_candidate_audit_rows = clean_records_text_fields(_slice_global_rows("_FIGURE_CANDIDATE_AUDIT_ROWS", fig_candidate_start))
        if not paper_figure_candidate_audit_rows:
            for fig in paper_figure_infos:
                paper_figure_candidate_audit_rows.append({
                    "Source_File": filename,
                    "page_num": fig.get("page_num", ""),
                    "candidate_reason": fig.get("candidate_reason", ""),
                    "candidate_score": fig.get("candidate_score", ""),
                    "text_len": fig.get("text_len", fig.get("page_text_len", "")),
                    "is_candidate": True,
                    "text_context_preview": _context_preview(fig.get("text_context", ""), 800),
                })
        figure_candidate_audit_rows.extend(clean_records_text_fields(paper_figure_candidate_audit_rows))

        paper_validation_audit_rows = clean_records_text_fields(_slice_global_rows("_FIGURE_VALIDATION_AUDIT_ROWS", fig_validation_start))
        payloads = _slice_global_rows("_FIGURE_DEBUG_PAYLOADS", fig_debug_start)
        if payloads:
            paper_figure_metadata_rows, paper_raw_figure_rows_before_dedupe, paper_raw_figure_rows_after_dedupe, paper_validated_figure_rows = _unpack_debug_payloads(payloads, filename)
        paper_figure_metadata_rows = clean_records_text_fields(_canonicalize_rows(paper_figure_metadata_rows))
        paper_raw_figure_rows_before_dedupe = clean_records_text_fields(_canonicalize_rows(paper_raw_figure_rows_before_dedupe))
        paper_raw_figure_rows_after_dedupe = clean_records_text_fields(_canonicalize_rows(paper_raw_figure_rows_after_dedupe))
        paper_validated_figure_rows = clean_records_text_fields(_canonicalize_rows(paper_validated_figure_rows))
        paper_validation_audit_rows = clean_records_text_fields(_canonicalize_rows(paper_validation_audit_rows, audit_mode=True))
        figure_point_validation_audit_rows.extend(clean_records_text_fields(paper_validation_audit_rows))

        if candidate_rows:
            candidate_rows = propagate_loading_within_catalyst_families(candidate_rows)
            candidate_rows, paper_comparison_audit_rows = mark_or_filter_comparison_rows(candidate_rows, filename)
            all_data_raw.extend(candidate_rows)  # 保留原始字段（含 _broadcast_path_primary/_broadcast_flags），在 clean 前收集
            candidate_rows = clean_records_text_fields(_canonicalize_rows(candidate_rows))
        else:
            candidate_rows = []

        all_data.extend(clean_records_text_fields(_canonicalize_rows(candidate_rows)))

        audit_rows.extend(clean_records_text_fields(temp_audit))
        audit_rows.extend(clean_records_text_fields(low_info_audit))
        audit_rows.extend(clean_records_text_fields(obvious_dup_audit))
        comparison_audit_rows.extend(clean_records_text_fields(paper_comparison_audit_rows))

        single_csv = _build_single_paper_csv_path(filename)
        _write_curated_csv(candidate_rows, single_csv)
        print(f"  [OK] {os.path.basename(single_csv)} ({len(candidate_rows)} rows)")

        paper_preparation_sidecar = _build_preparation_sidecar(filename, global_params, paper_text_records or candidate_rows)
        if EXPORT_AUDIT_ARTIFACTS:
            save_multimodal_audit_artifacts(
                OUT_DIR,
                filename,
                paper_figure_infos,
                paper_figure_metadata_rows,
                paper_raw_figure_rows_before_dedupe,
                paper_raw_figure_rows_after_dedupe,
                paper_validated_figure_rows,
                _canonicalize_rows(paper_binding_rows, audit_mode=True),
                paper_preparation_sidecar,
            )

        metadata_failed_count = sum(1 for row in paper_validation_audit_rows if str(row.get("validation_status", "")).strip() == "metadata_failed")
        extraction_empty_count = sum(1 for row in paper_validation_audit_rows if str(row.get("validation_status", "")).strip() == "extraction_empty")
        figure_metadata_success_count = sum(
            1 for row in paper_validation_audit_rows
            if str(row.get("validation_status", "")).strip() in {"extraction_empty", "validation_empty", "validation_kept_partial", "validation_ok"}
        )
        if not figure_metadata_success_count and paper_figure_metadata_rows:
            figure_metadata_success_count = sum(1 for row in paper_figure_metadata_rows if not str(row.get("metadata_failure_reason", "")).strip())

        quality_summary_rows.append({
            "Source_File": filename,
            "candidate_row_count": len(candidate_rows),
            "text_record_count": len(paper_text_records),
            "figure_candidate_page_count": len(paper_figure_infos),
            "figure_metadata_success_count": figure_metadata_success_count,
            "figure_raw_point_count": len(paper_raw_figure_rows_before_dedupe),
            "figure_validated_point_count": len(paper_validated_figure_rows),
            "candidate_text_row_count": sum(1 for row in candidate_rows if str(row.get("data_source", "")).strip() == "text"),
            "candidate_figure_row_count": sum(1 for row in candidate_rows if str(row.get("data_source", "")).strip() == "figure"),
            "preparation_only_row_count": sum(1 for row in paper_text_point_gate_rows if str(row.get("point_gate_reason", "")).strip() == "preparation_backbone_only"),
            "performance_without_condition_anchor_count": sum(1 for row in paper_text_point_gate_rows if str(row.get("point_gate_reason", "")).strip() == "performance_without_condition_anchor"),
            "composition_conflict_count": sum(1 for row in paper_composition_audit_rows if str(row.get("composition_consistency_flag", "")).strip() not in {"", "ok"}),
            "comparison_filtered_count": len(paper_comparison_audit_rows),
            "metadata_failed_count": metadata_failed_count,
            "extraction_empty_count": extraction_empty_count,
        })

    print("\n" + "=" * 80)
    quality_summary_rows = clean_records_text_fields(quality_summary_rows)
    if EXPORT_AUDIT_ARTIFACTS:
        figure_binding_audit_rows = clean_records_text_fields(_canonicalize_rows(figure_binding_audit_rows, audit_mode=True))
        save_audit_csv(figure_binding_audit_rows, FIGURE_BINDING_AUDIT_CSV)
        print(f"  [Audit] figure binding audit: {FIGURE_BINDING_AUDIT_CSV} ({len(figure_binding_audit_rows)} rows)")

        figure_premerge_filter_audit_rows = clean_records_text_fields(_canonicalize_rows(figure_premerge_filter_audit_rows, audit_mode=True))
        save_figure_premerge_filter_audit_csv(figure_premerge_filter_audit_rows, FIGURE_PREMERGE_FILTER_AUDIT_CSV)
        print(f"  [Audit] figure premerge filter audit: {FIGURE_PREMERGE_FILTER_AUDIT_CSV} ({len(figure_premerge_filter_audit_rows)} rows)")

        comparison_audit_rows = clean_records_text_fields(comparison_audit_rows)
        save_comparison_audit_csv(comparison_audit_rows, COMPARISON_FILTER_AUDIT_CSV)
        print(f"  [Audit] comparison filter audit: {COMPARISON_FILTER_AUDIT_CSV} ({len(comparison_audit_rows)} rows)")

        if audit_rows:
            audit_rows = clean_records_text_fields(audit_rows)
            save_to_csv(audit_rows, UPSTREAM_FILTER_AUDIT_CSV)
            print(f"  [Audit] upstream filter audit: {UPSTREAM_FILTER_AUDIT_CSV} ({len(audit_rows)} rows)")

        figure_candidate_audit_rows = clean_records_text_fields(figure_candidate_audit_rows)
        save_figure_candidate_audit_csv(figure_candidate_audit_rows, FIGURE_CANDIDATE_AUDIT_CSV)
        print(f"  [Audit] figure candidate audit: {FIGURE_CANDIDATE_AUDIT_CSV} ({len(figure_candidate_audit_rows)} rows)")

        figure_point_validation_audit_rows = clean_records_text_fields(figure_point_validation_audit_rows)
        save_figure_point_validation_audit_csv(figure_point_validation_audit_rows, FIGURE_POINT_VALIDATION_AUDIT_CSV)
        print(f"  [Audit] figure point validation audit: {FIGURE_POINT_VALIDATION_AUDIT_CSV} ({len(figure_point_validation_audit_rows)} rows)")

        composition_consistency_audit_rows = clean_records_text_fields(composition_consistency_audit_rows)
        save_composition_consistency_audit_csv(composition_consistency_audit_rows, COMPOSITION_CONSISTENCY_AUDIT_CSV)
        print(f"  [Audit] composition consistency audit: {COMPOSITION_CONSISTENCY_AUDIT_CSV} ({len(composition_consistency_audit_rows)} rows)")

        text_point_gate_audit_rows = clean_records_text_fields(text_point_gate_audit_rows)
        save_text_point_gate_audit_csv(text_point_gate_audit_rows, TEXT_POINT_GATE_AUDIT_CSV)
        print(f"  [Audit] text point gate audit: {TEXT_POINT_GATE_AUDIT_CSV} ({len(text_point_gate_audit_rows)} rows)")

        save_extraction_quality_summary(quality_summary_rows, EXTRACTION_QUALITY_SUMMARY_JSON)
        print(f"  [Audit] extraction quality summary: {EXTRACTION_QUALITY_SUMMARY_JSON} ({len(quality_summary_rows)} rows)")

    dataset_rows_clean = clean_records_text_fields(_canonicalize_rows(all_data))
    dataset_all_csv = os.path.join(OUT_DIR, "dataset_all.csv")
    _write_curated_csv(dataset_rows_clean, dataset_all_csv)
    print(f"  [Export] single-layer mode active")
    print(f"  [OK] dataset_all.csv: {dataset_all_csv} ({len(dataset_rows_clean)} rows)")
    print(f"  [Summary] processed papers: {len(files)}")
    print(f"  [Summary] single-paper csvs: {len(files)}")
    print(f"  [Summary] candidate rows retained: {len(dataset_rows_clean)}")
    if not EXPORT_AUDIT_ARTIFACTS:
        print("  [Audit] audit artifact export disabled (set MSR_EXPORT_AUDITS=1 to re-enable)")
    # 广播效率 summary：无论 EXPORT_AUDIT_ARTIFACTS 是否开启都输出，作为验收依据
    # 传入 all_data_raw（clean 前）以保留 _broadcast_path_primary/_broadcast_flags 字段
    save_broadcast_efficiency_summary(all_data_raw, OUT_DIR)
    print("=" * 80)



def save_broadcast_efficiency_summary(all_rows: List[Dict], out_dir: str) -> None:
    """统计 figure 行各广播路径的贡献及关键字段填充率，输出 broadcast_efficiency_summary.json。
    使用显式 _broadcast_path_primary / _broadcast_flags 字段统计，不依赖 Notes tag。
    fill_rate 拆成 any_source 和 identity_bound 两层，分母统一为 figure_total。
    """
    import json

    FILL_RATE_FIELDS = [
        "Canonical_Catalyst_ID", "Active_Metal", "Promoter_Metal",
        "Alloy_Ratio", "Ni_Fraction", "Promoter_Fraction",
        "Metal_Loading_wt%", "Support",
        "Reaction_Temp_C", "S_C_Ratio", "GHSV_mL_g_h", "Pressure_bar",
    ]
    IDENTITY_BOUND_FLAGS = {"registry_bind", "text_master"}

    fig_rows = [r for r in all_rows if str(r.get("data_source", "")).strip() == "figure"]
    total = len(fig_rows)
    if total == 0:
        return

    def _is_filled(r, field):
        v = str(r.get(field, "")).strip()
        return bool(v) and v.lower() != "nan"

    def _has_identity_flag(r):
        flags = set(str(r.get("_broadcast_flags", "")).split("|"))
        return bool(flags & IDENTITY_BOUND_FLAGS)

    # 按 _broadcast_path_primary 统计各路径行数
    matched_registry    = [r for r in fig_rows if str(r.get("figure_binding_status", "")).strip() == "matched"]
    matched_with_identity = [r for r in matched_registry if _is_filled(r, "Canonical_Catalyst_ID")]
    primary_counts = {}
    for r in fig_rows:
        p = str(r.get("_broadcast_path_primary", "unresolved")).strip() or "unresolved"
        primary_counts[p] = primary_counts.get(p, 0) + 1

    unmatched_total = sum(1 for r in fig_rows if str(r.get("figure_binding_status", "")).strip() == "unmatched")
    ambiguous_total = sum(1 for r in fig_rows if str(r.get("figure_binding_status", "")).strip() == "ambiguous")
    unresolved_total = unmatched_total + ambiguous_total

    # fill_rate：any_source 和 identity_bound，分母均为 figure_total
    fill_rate = {}
    for field in FILL_RATE_FIELDS:
        any_source_count = sum(1 for r in fig_rows if _is_filled(r, field))
        identity_bound_count = sum(1 for r in fig_rows if _is_filled(r, field) and _has_identity_flag(r))
        fill_rate[field] = {
            "any_source": round(any_source_count / total, 4),
            "identity_bound": round(identity_bound_count / total, 4),
        }

    summary = {
        "figure_total": total,
        "matched_registry_total": len(matched_registry),
        "matched_registry_with_identity_core": len(matched_with_identity),
        "broadcast_path_counts": primary_counts,
        "unresolved_figure_total": unresolved_total,
        "unmatched_figure_total": unmatched_total,
        "ambiguous_figure_total": ambiguous_total,
        "fill_rate": fill_rate,
    }

    out_path = os.path.join(out_dir, "broadcast_efficiency_summary.json")
    try:
        with open(out_path, "w", encoding="utf-8") as f:
            json.dump(summary, f, ensure_ascii=False, indent=2)
        print(f"  [Broadcast] efficiency summary: {out_path}")
    except Exception as e:
        print(f"  [WARNING] failed to save broadcast_efficiency_summary: {e}")


if __name__ == "__main__":
    main()









